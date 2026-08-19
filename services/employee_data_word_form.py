"""Controlled Word questionnaire generation and extraction.

The DOCX uses Word content controls (SDTs) with stable tags.  This lets the
system read answers without relying on a human-readable layout or OCR.
"""
from __future__ import annotations

import io
import zipfile
from datetime import datetime
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from services.employee_data_import import EmployeeDataImportError


WORD_SCHEMA = "EMP-DATA-WORD/V1.0"
PLACEHOLDER = "انقر هنا للكتابة"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}

SECTIONS = (
    ("تعريف الموظف", (("employee_no", "الرقم الوظيفي"), ("full_name_quad", "الاسم الرباعي"), ("timeclock_code", "كود ساعة الدوام (رقم الهوية أو الرقم الوظيفي)"))),
    ("البيانات الشخصية", (("identity_type_lookup_id", "نوع وثيقة الهوية"), ("national_id", "رقم الهوية/الجواز"), ("birth_date", "تاريخ الميلاد YYYY-MM-DD"), ("gender_lookup_id", "الجنس"), ("marital_status_lookup_id", "الحالة الاجتماعية"), ("religion_lookup_id", "الديانة"), ("disability_lookup_id", "حالة الإعاقة"), ("home_governorate_lookup_id", "محافظة السكن"), ("locality_lookup_id", "التجمع السكاني"), ("address", "العنوان"), ("phone", "الهاتف"), ("mobile", "رقم الجوال"), ("email", "البريد الإلكتروني"))),
    ("بيانات العمل والتعيين", (("work_governorate_lookup_id", "محافظة العمل"), ("work_location_lookup_id", "موقع العمل"), ("employee_status_lookup_id", "حالة الموظف"), ("status_date", "تاريخ سريان الحالة YYYY-MM-DD"), ("status_note", "ملاحظة حالة الموظف"), ("shift_lookup_id", "الوردية"), ("hourly_number", "الرقم في الساعة"), ("organization_id", "الإدارة العامة"), ("directorate_id", "الدائرة"), ("department_id", "القسم"), ("division_id", "الشعبة"), ("direct_manager_user_id", "المسؤول المباشر"), ("project_lookup_id", "المشروع"), ("appointment_type_lookup_id", "نوع التعيين"), ("hire_date", "تاريخ التعيين YYYY-MM-DD"), ("last_promotion_date", "تاريخ آخر ترقية YYYY-MM-DD"), ("job_category_lookup_id", "الفئة الوظيفية"), ("job_grade_lookup_id", "الدرجة الوظيفية"), ("job_title_lookup_id", "المسمى الوظيفي"), ("admin_title_lookup_id", "المسمى الإداري"))),
    ("البيانات المالية", (("bank_lookup_id", "البنك"), ("bank_account", "رقم الحساب / IBAN"), ("notes", "ملاحظات عامة"))),
)

REPEATED = {
    "dependents": ("dependent", ("full_name", "صلة القرابة", "national_id", "gender", "birth_date", "allowance"), ("full_name", "relation_lookup_id", "national_id", "gender_lookup_id", "birth_date", "allowance")),
    "qualifications": ("qualification", ("الدرجة العلمية", "التخصص", "التقدير", "تاريخ المؤهل", "الجامعة/المعهد", "الدولة", "ملاحظات"), ("degree_lookup_id", "specialization_lookup_id", "grade_lookup_id", "qualification_date", "university_lookup_id", "country_lookup_id", "notes")),
}


def _sdt(tag: str):
    sdt = OxmlElement("w:sdt")
    props = OxmlElement("w:sdtPr")
    alias = OxmlElement("w:alias"); alias.set(qn("w:val"), tag); props.append(alias)
    tag_el = OxmlElement("w:tag"); tag_el.set(qn("w:val"), tag); props.append(tag_el)
    text = OxmlElement("w:text"); props.append(text)
    content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r"); text_el = OxmlElement("w:t"); text_el.text = PLACEHOLDER; run.append(text_el); content.append(run)
    sdt.append(props); sdt.append(content)
    return sdt


def _field(cell, tag: str):
    paragraph = cell.paragraphs[0]
    paragraph._p.clear_content()
    paragraph._p.append(_sdt(tag))


def build_employee_word_form() -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Cm(1.5)
    section.left_margin = section.right_margin = Cm(1.5)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial"); normal.font.size = Pt(10)
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("استبيان استكمال ملف الموظف — نموذج Word تفاعلي").bold = True
    doc.add_paragraph(f"{WORD_SCHEMA} | اكتب داخل الحقول فقط، ولا تغيّر عناوين الحقول. بعد الحفظ بصيغة DOCX سلّم الملف للموارد البشرية.")
    for title_text, fields in SECTIONS:
        doc.add_heading(title_text, level=1)
        table = doc.add_table(rows=0, cols=2); table.style = "Table Grid"
        for key, label in fields:
            cells = table.add_row().cells
            cells[0].text = label
            _field(cells[1], f"f:{key}")
    for name, (prefix, headers, keys) in REPEATED.items():
        doc.add_heading("التابعون" if name == "dependents" else "المؤهلات العلمية", level=1)
        table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"
        for cell, header in zip(table.rows[0].cells, headers): cell.text = header
        for number in range(1, 4):
            cells = table.add_row().cells
            for cell, key in zip(cells, keys): _field(cell, f"t:{prefix}:{number}:{key}")
    doc.add_paragraph("إقرار الموظف: أقر بأن المعلومات المدخلة صحيحة حسب علمي.")
    doc.add_paragraph("الاسم والتوقيع: __________________________   التاريخ: __________________")
    buffer = io.BytesIO(); doc.save(buffer); return buffer.getvalue()


def parse_employee_word_form(raw: bytes) -> dict:
    try:
        with zipfile.ZipFile(io.BytesIO(raw)) as archive:
            root = ET.fromstring(archive.read("word/document.xml"))
    except Exception as exc:
        raise EmployeeDataImportError("ملف Word غير صالح أو تالف.") from exc
    text = " ".join(root.itertext())
    if WORD_SCHEMA not in text:
        raise EmployeeDataImportError("هذا ليس نموذج Word المعتمد لاستكمال بيانات الموظف.")
    fields, table_rows = {}, {"dependents": {}, "qualifications": {}}
    for sdt in root.findall(".//w:sdt", NS):
        tag = sdt.find("./w:sdtPr/w:tag", NS)
        if tag is None: continue
        value = " ".join(t.text or "" for t in sdt.findall(".//w:sdtContent//w:t", NS)).strip()
        if not value or value == PLACEHOLDER: continue
        marker = tag.get(f"{{{W_NS}}}val", "")
        if marker.startswith("f:"):
            fields[marker[2:]] = [{"value": value, "occurrence": 1}]
        elif marker.startswith("t:"):
            _, prefix, row, key = marker.split(":", 3)
            collection = "dependents" if prefix == "dependent" else "qualifications"
            table_rows[collection].setdefault(int(row), {})[f"{prefix}.{key}"] = value
    return {"schema": "EMP-DATA-FORM/V1.1", "exported_at": datetime.utcnow().isoformat(), "employee": {}, "fields": fields, "tables": {key: list(rows.values()) for key, rows in table_rows.items() if rows}, "selections": []}
