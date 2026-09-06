"""Word export helpers for employee follow-up reports."""

from __future__ import annotations

import io
import zipfile
from pathlib import Path


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def is_valid_docx(path: str | Path) -> bool:
    try:
        with zipfile.ZipFile(path) as archive:
            return "word/document.xml" in archive.namelist()
    except (OSError, zipfile.BadZipFile):
        return False


def _set_rtl(
    paragraph,
    *,
    bold: bool = False,
    size: int = 12,
    color: str | None = None,
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(6)
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        r_pr = run._element.get_or_add_rPr()
        fonts = r_pr.get_or_add_rFonts()
        fonts.set(qn("w:ascii"), "Arial")
        fonts.set(qn("w:hAnsi"), "Arial")
        fonts.set(qn("w:cs"), "Arial")
        rtl = r_pr.find(qn("w:rtl"))
        if rtl is None:
            rtl = OxmlElement("w:rtl")
            r_pr.append(rtl)
        rtl.set(qn("w:val"), "1")


def _paragraph(document, value: str, *, bold: bool = False, size: int = 12):
    paragraph = document.add_paragraph(value or "-")
    _set_rtl(paragraph, bold=bold, size=size, color="000000" if bold else None)
    return paragraph


def _date_label(value) -> str:
    return value.strftime("%Y-%m-%d") if value else "-"


def _status_label(value: str | None) -> str:
    return {
        "COMPLETED": "منجز",
        "INCOMPLETE": "غير مكتمل",
        "IN_PROGRESS": "قيد التنفيذ",
    }.get((value or "").upper(), value or "-")


def _set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge in ("top", "start", "bottom", "end"):
        margin = margins.find(qn(f"w:{edge}"))
        if margin is None:
            margin = OxmlElement(f"w:{edge}")
            margins.append(margin)
        margin.set(qn("w:w"), "110")
        margin.set(qn("w:type"), "dxa")


def _set_cell(cell, value: str, *, header: bool = False) -> None:
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

    cell.text = str(value or "-")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)
    if header:
        _set_cell_shading(cell, "1F4E78")
    for paragraph in cell.paragraphs:
        _set_rtl(paragraph, bold=header, size=11, color="FFFFFF" if header else None)


def _set_table_rtl(table, widths: tuple[float, ...]) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches

    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = False
    table_pr = table._tbl.tblPr
    bidi_visual = table_pr.find(qn("w:bidiVisual"))
    if bidi_visual is None:
        bidi_visual = OxmlElement("w:bidiVisual")
        table_pr.append(bidi_visual)
    bidi_visual.set(qn("w:val"), "1")

    borders = table_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "D9D9D9")

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if index < len(widths):
                cell.width = Inches(widths[index])


def _add_rtl_table(document, headers: tuple[str, ...], rows, widths: tuple[float, ...]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_table_rtl(table, widths)
    for cell, label in zip(table.rows[0].cells, headers):
        _set_cell(cell, label, header=True)
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            _set_cell(cell, value)
        _set_table_rtl(table, widths)
    return table


def build_followup_docx(report, template_path: str | Path | None = None) -> bytes:
    """Build a reviewable report document and preserve an uploaded letterhead."""
    from docx import Document
    from docx.shared import Pt

    path = Path(template_path) if template_path else None
    document = Document(str(path)) if path and path.is_file() else Document()
    try:
        normal = document.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(12)
    except KeyError:
        pass

    title = document.add_paragraph("تقرير إنجاز الموظف")
    try:
        title.style = document.styles["Title"]
    except KeyError:
        pass
    _set_rtl(title, bold=True, size=18, color="000000")

    details_rows = []
    for label, value in (
        ("الموظف", getattr(getattr(report, "employee", None), "full_name", "-")),
        ("المدير المباشر", getattr(getattr(report, "manager", None), "full_name", "-")),
        ("فترة التقرير", f"{_date_label(report.period_start)} إلى {_date_label(report.period_end)}"),
        ("الحالة", _status_label(report.status)),
    ):
        details_rows.append((label, value))
    _add_rtl_table(document, ("البيان", "التفاصيل"), details_rows, (1.7, 4.8))

    _paragraph(document, "ملخص الإنجازات", bold=True, size=15)
    completed_items = [
        item
        for item in (report.items or [])
        if getattr(item, "is_included", True)
        and (getattr(item, "status", "") or "").upper() == "COMPLETED"
    ]
    accomplishment_rows = [
        (getattr(item, "title", None) or "مهمة منجزة", _date_label(getattr(item, "completed_on", None)))
        for item in completed_items
    ] or [("لا توجد مهام منجزة خلال فترة التقرير.", "-")]
    _add_rtl_table(document, ("المهمة", "التاريخ"), accomplishment_rows, (5.0, 1.5))

    _paragraph(document, "ملخص الموظف", bold=True, size=15)
    _paragraph(document, report.employee_summary or report.ai_summary or "-")
    _paragraph(document, "التحديات أو الاحتياجات", bold=True, size=15)
    _paragraph(document, report.challenges or "-")
    _paragraph(document, "المطلوب من المدير", bold=True, size=15)
    _paragraph(document, report.manager_request or "-")

    if report.manager_comment or report.manager_rating:
        _paragraph(document, "مراجعة المدير", bold=True, size=15)
        _paragraph(document, report.manager_comment or "-")
        _paragraph(document, f"التقييم المختصر: {report.manager_rating or '-'}")

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()
