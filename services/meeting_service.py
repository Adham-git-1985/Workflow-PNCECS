"""Small, framework-independent helpers for portal meetings."""

from __future__ import annotations

from collections.abc import Iterable
from io import BytesIO
from pathlib import Path
from typing import Any
import zipfile

from services.ole_package import build_ole_package


RECORDED_ATTENDANCE_LABELS = {
    "ATTENDED": "حضر",
    "ABSENT": "تغيب",
}


EMBEDDED_ATTACHMENT_RELATIONSHIP = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject"
)
EMBEDDED_ATTACHMENT_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.oleObject"
)


def add_embedded_attachments_to_docx(
    document: Any,
    attachments: Iterable[tuple[str, str | Path]],
    *,
    icon_path: str | Path,
) -> int:
    """Insert files into a DOCX as Word ``Package`` OLE objects.

    Each source file becomes part of the DOCX package itself. The visible
    filename remains ordinary document text, while double-clicking the icon in
    Microsoft Word opens the embedded copy.
    """

    from docx.opc.packuri import PackURI
    from docx.opc.part import Part
    from docx.oxml import parse_xml
    from docx.shared import Pt
    from markupsafe import escape

    embedded_count = 0
    for original_name, source_path in attachments:
        path = Path(source_path)
        if not path.is_file():
            raise FileNotFoundError(f"Meeting attachment is missing: {path.name}")

        payload = path.read_bytes()
        ole_blob = build_ole_package(original_name, payload)
        partname = document.part.package.next_partname(
            "/word/embeddings/oleObject%d.bin"
        )
        ole_part = Part(
            PackURI(str(partname)),
            EMBEDDED_ATTACHMENT_CONTENT_TYPE,
            ole_blob,
            document.part.package,
        )
        ole_rid = document.part.relate_to(
            ole_part,
            EMBEDDED_ATTACHMENT_RELATIONSHIP,
        )
        image_rid, _image = document.part.get_or_add_image(str(icon_path))

        embedded_count += 1
        shape_id = f"_x0000_i{1100 + embedded_count}"
        shape_type_id = f"_x0000_t75_{embedded_count}"
        object_id = f"_{1848043000 + embedded_count}"
        object_namespaces = (
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
            'xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:o="urn:schemas-microsoft-com:office:office"'
        )
        object_xml = (
            f'<w:object {object_namespaces} '
            'w:dxaOrig="720" w:dyaOrig="720">'
            f'<v:shapetype id="{shape_type_id}" coordsize="21600,21600" '
            'o:spt="75" o:preferrelative="t" path="m@4@5l@4@11@9@11@9@5xe" '
            'filled="f" stroked="f">'
            '<v:stroke joinstyle="miter"/>'
            '<v:formulas><v:f eqn="if lineDrawn pixelLineWidth 0"/>'
            '<v:f eqn="sum @0 1 0"/><v:f eqn="sum 0 0 @1"/>'
            '<v:f eqn="prod @2 1 2"/><v:f eqn="prod @3 21600 pixelWidth"/>'
            '<v:f eqn="prod @3 21600 pixelHeight"/><v:f eqn="sum @0 0 1"/>'
            '<v:f eqn="prod @6 1 2"/><v:f eqn="prod @7 21600 pixelWidth"/>'
            '<v:f eqn="sum @8 21600 0"/><v:f eqn="prod @7 21600 pixelHeight"/>'
            '<v:f eqn="sum @10 21600 0"/></v:formulas>'
            '<v:path o:extrusionok="f" gradientshapeok="t" o:connecttype="rect"/>'
            '<o:lock v:ext="edit" aspectratio="t"/>'
            '</v:shapetype>'
            f'<v:shape id="{shape_id}" type="#{shape_type_id}" '
            'style="width:28pt;height:28pt" o:ole="">'
            f'<v:imagedata r:id="{escape(image_rid)}" o:title=""/>'
            '</v:shape>'
            f'<o:OLEObject Type="Embed" ProgID="Package" ShapeID="{shape_id}" '
            f'DrawAspect="Content" ObjectID="{object_id}" r:id="{escape(ole_rid)}"/>'
            '</w:object>'
        )

        paragraph = document.add_paragraph()
        paragraph.add_run()._r.append(parse_xml(object_xml))
        label = paragraph.add_run(f"  {original_name}")
        label.font.size = Pt(14)

        # Keep the label accessible even if a non-Word viewer does not activate OLE.
        try:
            paragraph.style = document.styles["Normal"]
        except Exception:
            pass

    return embedded_count


def recorded_attendance_label(status: object) -> str:
    """Return the binary attendance label used in official minutes."""

    normalized = str(status or "").strip().upper()
    return RECORDED_ATTENDANCE_LABELS.get(normalized, "تغيب")


def validate_docx_package(data: bytes) -> None:
    """Reject incomplete or corrupt DOCX packages before download."""

    if not data:
        raise ValueError("The DOCX package is empty.")

    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            names = set(archive.namelist())
            required = {"[Content_Types].xml", "_rels/.rels", "word/document.xml"}
            if not required.issubset(names):
                raise ValueError("The DOCX package is missing required parts.")
            if archive.testzip() is not None:
                raise ValueError("The DOCX package contains a corrupt entry.")
    except (OSError, zipfile.BadZipFile) as exc:
        raise ValueError("The DOCX package is invalid.") from exc


def validate_embedded_attachments(data: bytes, expected_count: int) -> None:
    """Confirm a generated DOCX contains every expected embedded OLE object."""

    validate_docx_package(data)
    with zipfile.ZipFile(BytesIO(data)) as archive:
        names = archive.namelist()
        embedded = [
            name for name in names
            if name.startswith("word/embeddings/oleObject") and name.endswith(".bin")
        ]
        if len(embedded) != int(expected_count):
            raise ValueError("The DOCX package is missing an embedded attachment.")
        for name in embedded:
            blob = archive.read(name)
            if not blob.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):
                raise ValueError("An embedded attachment is not a valid OLE package.")



def normalize_agenda_order(
    submitted_ids: Iterable[object],
    current_ids: Iterable[object],
) -> list[int]:
    """Validate and normalize a complete agenda ordering.

    The submitted order must contain every current agenda item exactly once.
    This prevents a stale or manipulated form from moving an item that belongs
    to another meeting or silently dropping newly added items.
    """

    try:
        submitted = [int(value) for value in submitted_ids]
        current = [int(value) for value in current_ids]
    except (TypeError, ValueError) as exc:
        raise ValueError("Agenda order contains an invalid item id.") from exc

    if len(current) != len(set(current)):
        raise ValueError("Current agenda contains duplicate item ids.")
    if len(submitted) != len(set(submitted)):
        raise ValueError("Agenda order contains duplicate item ids.")
    if len(submitted) != len(current) or set(submitted) != set(current):
        raise ValueError("Agenda order does not match the current meeting agenda.")

    return submitted
