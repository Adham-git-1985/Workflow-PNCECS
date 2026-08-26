from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from services.transport_forms import _date_text, _plain, _time_text


FONT_NAME = "Sakkal Majalla"
FONT_SIZE = Pt(16)
PAGE_WIDTH_MM = 210
PAGE_HEIGHT_MM = 297
MARGIN_LEFT_MM = 18
MARGIN_RIGHT_MM = 18
CONTENT_WIDTH_MM = PAGE_WIDTH_MM - MARGIN_LEFT_MM - MARGIN_RIGHT_MM


def _set_run_font(run, *, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), FONT_NAME)
    for size_tag in ("sz", "szCs"):
        size_node = r_pr.find(qn(f"w:{size_tag}"))
        if size_node is None:
            size_node = OxmlElement(f"w:{size_tag}")
            r_pr.append(size_node)
        size_node.set(qn("w:val"), "32")


def _set_paragraph_rtl(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:bidi")) is None:
        p_pr.append(OxmlElement("w:bidi"))


def _format_paragraph(
    paragraph,
    *,
    alignment=WD_ALIGN_PARAGRAPH.RIGHT,
    before: float = 0,
    after: float = 2,
    keep_with_next: bool = False,
) -> None:
    # In Word, w:bidi mirrors physical left/right alignment. Using LEFT for an
    # RTL paragraph therefore places Arabic text against the visual right edge.
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if alignment == WD_ALIGN_PARAGRAPH.RIGHT else alignment
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.keep_with_next = keep_with_next
    _set_paragraph_rtl(paragraph)


def _add_paragraph(
    container,
    text,
    *,
    bold: bool = False,
    alignment=WD_ALIGN_PARAGRAPH.RIGHT,
    before: float = 0,
    after: float = 2,
    keep_with_next: bool = False,
):
    paragraph = container.add_paragraph()
    _format_paragraph(
        paragraph,
        alignment=alignment,
        before=before,
        after=after,
        keep_with_next=keep_with_next,
    )
    run = paragraph.add_run(_plain(text))
    _set_run_font(run, bold=bold)
    return paragraph


def _image_dimensions(payload: bytes, max_width_mm: float, max_height_mm: float) -> tuple[float, float]:
    from PIL import Image

    with Image.open(BytesIO(payload)) as image:
        width, height = image.size
    scale = min(max_width_mm / max(width, 1), max_height_mm / max(height, 1))
    return width * scale, height * scale


def _letterhead_fragments(letterhead_path: str | None) -> tuple[bytes | None, bytes | None]:
    path = Path(letterhead_path) if letterhead_path else None
    if not path or not path.is_file():
        return None, None
    if path.suffix.lower() in {".png", ".jpg", ".jpeg"}:
        return path.read_bytes(), None
    if path.suffix.lower() != ".pdf":
        return None, None

    import fitz
    from PIL import Image

    source = fitz.open(str(path))
    try:
        page = source[0]
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        page_image = Image.open(BytesIO(pixmap.tobytes("png"))).convert("RGB")
        width, height = page_image.size
        header_image = page_image.crop((0, 0, width, max(1, round(height * 0.20))))
        footer_image = page_image.crop((0, round(height * 0.86), width, height))
        header_output = BytesIO()
        footer_output = BytesIO()
        header_image.save(header_output, format="PNG", optimize=True)
        footer_image.save(footer_output, format="PNG", optimize=True)
        return header_output.getvalue(), footer_output.getvalue()
    finally:
        source.close()


def _add_centered_picture(container, payload: bytes, *, max_width_mm: float, max_height_mm: float) -> None:
    paragraph = container.paragraphs[0]
    _format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    width_mm, height_mm = _image_dimensions(payload, max_width_mm, max_height_mm)
    run = paragraph.add_run()
    run.add_picture(BytesIO(payload), width=Mm(width_mm), height=Mm(height_mm))


def _set_document_defaults(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = FONT_SIZE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1
    r_pr = normal.element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), FONT_NAME)
    for size_tag in ("sz", "szCs"):
        size_node = r_pr.find(qn(f"w:{size_tag}"))
        if size_node is None:
            size_node = OxmlElement(f"w:{size_tag}")
            r_pr.append(size_node)
        size_node.set(qn("w:val"), "32")


def _new_document(letterhead_path: str | None) -> Document:
    document = Document()
    _set_document_defaults(document)
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.left_margin = Mm(MARGIN_LEFT_MM)
    section.right_margin = Mm(MARGIN_RIGHT_MM)
    section.top_margin = Mm(47)
    section.bottom_margin = Mm(21)
    section.header_distance = Mm(4)
    section.footer_distance = Mm(5)

    header_payload, footer_payload = _letterhead_fragments(letterhead_path)
    if header_payload:
        _add_centered_picture(section.header, header_payload, max_width_mm=CONTENT_WIDTH_MM, max_height_mm=40)
    else:
        header = section.header.paragraphs[0]
        _format_paragraph(header, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=0)
        run = header.add_run("دولة فلسطين - اللجنة الوطنية الفلسطينية للتربية والثقافة والعلوم")
        _set_run_font(run, bold=True)
    if footer_payload:
        _add_centered_picture(section.footer, footer_payload, max_width_mm=CONTENT_WIDTH_MM, max_height_mm=15)

    document.core_properties.title = "نماذج الحركة والنقل"
    document.core_properties.author = "نظام مسار"
    return document


def _set_cell_margins(cell, *, top: int = 55, start: int = 90, bottom: int = 55, end: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = tc_pr.find(qn("w:shd"))
    if shade is None:
        shade = OxmlElement("w:shd")
        tc_pr.append(shade)
    shade.set(qn("w:fill"), fill)


def _remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def _configure_table_geometry(table, widths_mm: list[float], *, with_grid: bool = True) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    table_width_dxa = round(sum(widths_mm) / 25.4 * 1440)
    widths_dxa = [round(width_mm / 25.4 * 1440) for width_mm in widths_mm]
    widths_dxa[-1] += table_width_dxa - sum(widths_dxa)
    tbl_width = tbl_pr.first_child_found_in("w:tblW")
    tbl_width.set(qn("w:w"), str(table_width_dxa))
    tbl_width.set(qn("w:type"), "dxa")
    tbl_indent = OxmlElement("w:tblInd")
    tbl_indent.set(qn("w:w"), "90")
    tbl_indent.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_indent)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_dxa in widths_dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width_dxa))
        grid.append(grid_column)

    for row in table.rows:
        row_pr = row._tr.get_or_add_trPr()
        row_pr.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            width_dxa = widths_dxa[index]
            cell.width = Mm(widths_mm[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.first_child_found_in("w:tcW")
            tc_width.set(qn("w:w"), str(width_dxa))
            tc_width.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)

    if not with_grid:
        _remove_table_borders(table)


def _fill_cell(cell, text, *, bold: bool = False, shaded: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    paragraph = cell.paragraphs[0]
    _format_paragraph(paragraph, alignment=alignment, after=0)
    run = paragraph.add_run(_plain(text, " "))
    _set_run_font(run, bold=bold)
    if shaded:
        _shade_cell(cell, "F1F3F5")


def _add_grid_table(document: Document, rows: list[list], widths_mm: list[float], *, header_rows: int = 1):
    table = document.add_table(rows=len(rows), cols=len(widths_mm))
    table.style = "Table Grid"
    _configure_table_geometry(table, widths_mm)
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            _fill_cell(
                table.cell(row_index, column_index),
                value,
                bold=row_index < header_rows,
                shaded=row_index < header_rows,
            )
    for row_index in range(min(header_rows, len(table.rows))):
        tr_pr = table.rows[row_index]._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:tblHeader"))
    return table


def _add_spacer(document: Document, points: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(points)
    paragraph.paragraph_format.line_spacing = Pt(1)


def _save_document(document: Document) -> bytes:
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_vehicle_license_docx(vehicle, letterhead_path: str | None = None) -> bytes:
    document = _new_document(letterhead_path)
    vehicle_type = " ".join(filter(None, [
        _plain(getattr(vehicle, "vehicle_type", None), ""),
        _plain(getattr(vehicle, "model", None), ""),
    ])).strip() or "-"
    assigned_to = _plain(getattr(vehicle, "assigned_to", None), _plain(getattr(vehicle, "label", None)))

    _add_paragraph(document, "كتاب ترخيص مركبة حكومية", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=4, keep_with_next=True)
    _add_paragraph(document, f"\u200fالرقم: ح/ترخيص/{getattr(vehicle, 'id', '')}", keep_with_next=True)
    _add_paragraph(document, f"\u200fالتاريخ \u202a{_date_text(None)}\u202c", after=5, keep_with_next=True)
    _add_paragraph(document, "الأخ المحترم / مدير عام النقل الحكومي", bold=True, keep_with_next=True)
    _add_paragraph(document, "وزارة النقل والمواصلات", bold=True, after=4, keep_with_next=True)
    _add_paragraph(
        document,
        "الموضوع: ترخيص مركبة حكومية خاصة باللجنة الوطنية الفلسطينية للتربية والثقافة والعلوم",
        bold=True,
        after=3,
        keep_with_next=True,
    )
    _add_paragraph(
        document,
        "تهديكم اللجنة الوطنية الفلسطينية للتربية والثقافة والعلوم أطيب تحياتها، وبالإشارة إلى الموضوع أعلاه، نرجو التكرم بإصدار رخصة المركبة الحكومية المذكورة بياناتها أدناه.",
        after=4,
        keep_with_next=True,
    )
    _add_grid_table(document, [
        ["المستخدم", "سنة الإنتاج", "النوع", "الرقم الحكومي"],
        [assigned_to, getattr(vehicle, "year", None), vehicle_type, getattr(vehicle, "plate_no", None)],
    ], [46, 34, 52, 42])
    _add_spacer(document, 4)
    _add_paragraph(
        document,
        "علمًا بأنه لا مانع لدينا من خصم رسوم الترخيص السنوية من حساب اللجنة لدى وزارة المالية والتخطيط.",
        after=5,
    )
    _add_paragraph(document, "وتفضلوا بقبول فائق الاحترام والتقدير،", bold=True, after=18)
    _add_paragraph(document, "الأمين العام", bold=True)
    return _save_document(document)


def build_maintenance_request_docx(maintenance, items, letterhead_path: str | None = None) -> bytes:
    document = _new_document(letterhead_path)
    vehicle = maintenance.vehicle
    date_value = _date_text(getattr(maintenance, "invoice_day", None) or getattr(maintenance, "created_at", None))
    vehicle_type = " ".join(filter(None, [
        _plain(getattr(vehicle, "vehicle_type", None), ""),
        _plain(getattr(vehicle, "model", None), ""),
    ])).strip() or "-"

    _add_paragraph(document, "طلب صيانة مركبة", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=4, keep_with_next=True)
    _add_paragraph(document, f"\u200fالتاريخ \u202a{date_value}\u202c", after=3, keep_with_next=True)
    _add_paragraph(
        document,
        "الأخ / أمين عام اللجنة الوطنية الفلسطينية للتربية والثقافة والعلوم",
        bold=True,
        keep_with_next=True,
    )
    _add_paragraph(document, "الموضوع: طلب صيانة مركبة", bold=True, keep_with_next=True)
    _add_paragraph(document, "يرجى التكرم بالموافقة على إجراء أعمال الصيانة التالية:", after=3, keep_with_next=True)
    _add_grid_table(document, [
        ["سنة الإنتاج", "رقم العداد", "رقم المحرك", "رقم الشاصي", "رقم المركبة", "نوع المركبة"],
        [
            getattr(vehicle, "year", None),
            getattr(vehicle, "odometer_no", None) or getattr(vehicle, "current_odometer", None),
            getattr(vehicle, "engine_no", None),
            getattr(vehicle, "chassis_no", None),
            getattr(vehicle, "plate_no", None),
            vehicle_type,
        ],
    ], [26, 27, 30, 32, 27, 32])
    _add_spacer(document, 3)

    requested = []
    for item in list(items or [])[:8]:
        label = getattr(getattr(item, "maintenance_type_lookup", None), "label", None)
        note = getattr(item, "note", None)
        requested.append(" - ".join(part for part in (label, note) if part) or "بند صيانة")
    if not requested and getattr(maintenance, "notes", None):
        requested = [line.strip() for line in str(maintenance.notes).splitlines() if line.strip()][:8]
    requested.extend([""] * (8 - len(requested)))
    rows = [["البيانات المطلوبة لأعمال الصيانة", "الرقم"]]
    rows.extend([[description or " ", index] for index, description in enumerate(requested, start=1)])
    _add_grid_table(document, rows, [156, 18])
    _add_spacer(document, 8)
    signatures = document.add_table(rows=1, cols=2)
    _configure_table_geometry(signatures, [87, 87], with_grid=False)
    _fill_cell(signatures.cell(0, 0), "توقيع أمين عام اللجنة الوطنية", bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _fill_cell(signatures.cell(0, 1), "توقيع مسؤول الحركة", bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    return _save_document(document)


def build_movement_permit_docx(permit, trip=None, letterhead_path: str | None = None) -> bytes:
    document = _new_document(letterhead_path)
    vehicle = getattr(permit, "vehicle", None)
    driver = getattr(permit, "driver", None)
    requester = getattr(permit, "requester", None)
    origin = getattr(getattr(permit, "origin_zone", None), "name", None) or getattr(permit, "origin_text", None)
    destination = getattr(getattr(permit, "dest_zone", None), "name", None) or getattr(permit, "dest_text", None)
    route_text = f"من: {_plain(origin)}\nإلى: {_plain(destination)}"
    if getattr(permit, "purpose", None):
        route_text += f"\nالغرض: {_plain(permit.purpose)}"
    depart_at = getattr(trip, "started_at", None) or getattr(permit, "depart_at", None)
    return_at = getattr(trip, "ended_at", None) or getattr(permit, "return_at", None)
    order_no = getattr(trip, "order_no", None) or getattr(permit, "ref_no", None) or getattr(permit, "id", None)
    requester_name = (
        getattr(requester, "full_name", None)
        or getattr(requester, "name", None)
        or getattr(requester, "email", None)
    )

    _add_paragraph(document, "نموذج رقم (1)", bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, after=2, keep_with_next=True)
    _add_paragraph(
        document,
        f"تصريح أمر حركة رقم ({_plain(order_no)})",
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        after=3,
        keep_with_next=True,
    )
    _add_paragraph(document, f"\u200fاليوم والتاريخ \u202a{_date_text(depart_at)}\u202c", after=3, keep_with_next=True)
    fields = [
        ("اسم السائق", getattr(driver, "name", None)),
        ("رقم السيارة", getattr(vehicle, "plate_no", None)),
        ("اسم الموظف المكلف بالمهمة", requester_name),
        ("خط سير الرحلة مع العنوان المستهدف", route_text),
        ("ساعة بدء المهمة", _time_text(depart_at)),
        ("ساعة نهاية المهمة", _time_text(return_at)),
        ("رقم العداد في بداية المهمة", getattr(trip, "start_odometer", None) if trip else getattr(vehicle, "current_odometer", None)),
        ("رقم العداد في نهاية المهمة", getattr(trip, "end_odometer", None) if trip else None),
    ]
    rows = [[value, label] for label, value in fields]
    table = _add_grid_table(document, rows, [118, 56], header_rows=0)
    for row in table.rows:
        _shade_cell(row.cells[1], "F5F5F5")
        for run in row.cells[1].paragraphs[0].runs:
            _set_run_font(run, bold=True)
    _add_spacer(document, 9)
    _add_paragraph(document, "توقيع المكلف بالمهمة: ______________________________", bold=True, after=8)
    _add_paragraph(document, "توقيع السائق: ______________________________________", bold=True, after=10)
    _add_paragraph(document, "اعتماد مسؤول الحركة: _______________________________", bold=True)
    return _save_document(document)
