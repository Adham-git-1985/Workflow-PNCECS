"""Local extraction and suggestions for manually uploaded correspondence.

The intake pipeline deliberately stays deterministic and local: uploaded bytes
are inspected in memory, converted to text, and matched against the live
correspondence lookups supplied by the caller.  Nothing in this module calls an
external AI service or persists the uploaded file.
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date, datetime
from email import policy
from email.parser import BytesParser
from email.utils import getaddresses, parseaddr
from io import BytesIO
from pathlib import Path
import html
import mimetypes
import os
import re
import shutil
import struct
import subprocess
import tempfile
import time
import unicodedata
import zipfile


SMART_INTAKE_EXTENSIONS = frozenset({
    ".csv",
    ".doc",
    ".docx",
    ".eml",
    ".htm",
    ".html",
    ".json",
    ".pdf",
    ".pptx",
    ".txt",
    ".xlsx",
    ".xml",
})

IMAGE_EXTENSIONS = frozenset({
    ".bmp",
    ".gif",
    ".jpeg",
    ".jpg",
    ".png",
    ".tif",
    ".tiff",
    ".webp",
})

_ARABIC_DIACRITICS = re.compile(
    r"[\u0610-\u061A\u064B-\u065F\u0670\u06D6-\u06ED]"
)
_MEANINGLESS_TOKENS = {
    "اداره",
    "الاداره",
    "دائره",
    "الدائره",
    "قسم",
    "القسم",
    "شعبه",
    "الشعبه",
    "فريق",
    "الفريق",
    "مسار",
    "المسار",
    "وحده",
    "الوحده",
}


class CorrespondenceIntakeError(ValueError):
    """A safe, user-facing smart-intake error."""

    def __init__(self, message: str, *, code: str, status_code: int = 400):
        super().__init__(message)
        self.message = message
        self.code = code
        self.status_code = status_code


@dataclass(frozen=True)
class OcrConfig:
    """Bounded local OCR settings for Tesseract."""

    enabled: bool = False
    command: str = "tesseract"
    languages: str = "ara+eng"
    max_pages: int = 10
    dpi: int = 200
    timeout_seconds: float = 45.0
    max_image_pixels: int = 40_000_000


@dataclass(frozen=True)
class ExtractedEmailAttachment:
    """One attachment extracted from an uploaded EML message."""

    filename: str
    payload: bytes
    mimetype: str | None


@dataclass(frozen=True)
class EmailAttachmentExtraction:
    """Bounded EML attachment extraction result, including safe warnings."""

    attachments: tuple[ExtractedEmailAttachment, ...]
    warnings: tuple[str, ...]


class _OcrRuntimeError(RuntimeError):
    pass


def _bounded_ocr_config(config: OcrConfig | None) -> OcrConfig:
    config = config or OcrConfig()
    languages = re.sub(r"[^A-Za-z0-9_+\-]", "", str(config.languages or ""))
    return OcrConfig(
        enabled=bool(config.enabled),
        command=str(config.command or "tesseract").strip() or "tesseract",
        languages=languages or "ara+eng",
        max_pages=max(1, min(int(config.max_pages or 1), 50)),
        dpi=max(120, min(int(config.dpi or 200), 400)),
        timeout_seconds=max(5.0, min(float(config.timeout_seconds or 45), 300.0)),
        max_image_pixels=max(1_000_000, min(int(config.max_image_pixels or 1), 100_000_000)),
    )


def _resolve_tesseract_command(config: OcrConfig) -> str | None:
    if not config.enabled:
        return None
    command = config.command
    if os.path.isfile(command):
        return os.path.abspath(command)
    return shutil.which(command)


def _image_to_png_frames(payload: bytes, config: OcrConfig) -> list[bytes]:
    """Normalize bounded image frames before handing them to OCR."""
    try:
        from PIL import Image, ImageOps, ImageSequence

        image = Image.open(BytesIO(payload))
    except Exception as exc:
        raise CorrespondenceIntakeError(
            "تعذر فتح الصورة للتحليل.", code="INVALID_IMAGE", status_code=422
        ) from exc

    frames: list[bytes] = []
    try:
        for frame_number, source_frame in enumerate(ImageSequence.Iterator(image), start=1):
            if frame_number > config.max_pages:
                break
            frame = ImageOps.exif_transpose(source_frame.copy())
            width, height = frame.size
            if width <= 0 or height <= 0 or width * height > config.max_image_pixels:
                raise CorrespondenceIntakeError(
                    "أبعاد الصورة أكبر من الحد الآمن للتحليل.",
                    code="IMAGE_TOO_LARGE",
                    status_code=413,
                )
            frame = ImageOps.grayscale(frame)
            if max(frame.size) < 1800:
                scale = min(3.0, 1800.0 / max(frame.size))
                frame = frame.resize(
                    (max(1, int(frame.width * scale)), max(1, int(frame.height * scale))),
                    Image.Resampling.LANCZOS,
                )
            frame = ImageOps.autocontrast(frame)
            output = BytesIO()
            frame.save(output, format="PNG", optimize=True)
            frames.append(output.getvalue())
    except CorrespondenceIntakeError:
        raise
    except Exception as exc:
        raise CorrespondenceIntakeError(
            "تعذر تجهيز الصورة للتحليل.", code="INVALID_IMAGE", status_code=422
        ) from exc
    finally:
        image.close()
    return frames


def _run_tesseract_png(
    png_payload: bytes,
    command: str,
    config: OcrConfig,
    timeout_seconds: float,
) -> str:
    """Run Tesseract without a shell and return its UTF-8 output."""
    with tempfile.TemporaryDirectory(prefix="corr_ocr_") as temporary_directory:
        image_path = os.path.join(temporary_directory, "page.png")
        with open(image_path, "wb") as image_file:
            image_file.write(png_payload)
        try:
            completed = subprocess.run(
                [
                    command,
                    image_path,
                    "stdout",
                    "-l",
                    config.languages,
                    "--psm",
                    "6",
                ],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=max(1.0, timeout_seconds),
                check=False,
                shell=False,
            )
        except subprocess.TimeoutExpired as exc:
            raise _OcrRuntimeError("انتهت مهلة OCR قبل اكتمال التحليل.") from exc
        except OSError as exc:
            raise _OcrRuntimeError("تعذر تشغيل محرك OCR المحلي.") from exc

    if completed.returncode != 0:
        detail = completed.stderr.decode("utf-8", errors="replace").strip()
        if "failed loading language" in detail.casefold() or "error opening data file" in detail.casefold():
            raise _OcrRuntimeError(
                "حزم لغات OCR المطلوبة غير مثبتة (العربية والإنجليزية)."
            )
        raise _OcrRuntimeError("فشل محرك OCR في قراءة المرفق.")
    return completed.stdout.decode("utf-8", errors="replace")


def _ocr_png_frames(
    frames: list[bytes],
    config: OcrConfig,
    *,
    command: str,
) -> tuple[str, list[str]]:
    parts: list[str] = []
    warnings: list[str] = []
    deadline = time.monotonic() + config.timeout_seconds
    for frame_number, png_payload in enumerate(frames[:config.max_pages], start=1):
        remaining = deadline - time.monotonic()
        if remaining <= 0:
            warnings.append("توقّف OCR بعد بلوغ المهلة الإجمالية للتحليل.")
            break
        try:
            page_text = _run_tesseract_png(png_payload, command, config, remaining)
        except _OcrRuntimeError as exc:
            warnings.append(str(exc))
            break
        if page_text.strip():
            parts.append(page_text)
    return "\n".join(parts), warnings


def _human_size_limit(max_bytes: int) -> str:
    """Return a concise Arabic label for a configured byte limit."""
    if max_bytes >= 1024 * 1024:
        return f"{max_bytes / (1024 * 1024):g} ميجابايت"
    if max_bytes >= 1024:
        return f"{max_bytes / 1024:g} كيلوبايت"
    return f"{max_bytes} بايت"


def read_limited_upload(file_storage, max_bytes: int) -> bytes:
    """Read an upload without allowing its compressed size to exceed the limit."""
    if not file_storage:
        raise CorrespondenceIntakeError(
            "اختر مرفقاً لتحليله.", code="MISSING_FILE", status_code=400
        )

    try:
        declared_size = int(getattr(file_storage, "content_length", 0) or 0)
    except (TypeError, ValueError):
        declared_size = 0
    if declared_size > max_bytes:
        raise CorrespondenceIntakeError(
            f"حجم المرفق أكبر من الحد المسموح للتحليل الذكي "
            f"({_human_size_limit(max_bytes)}). يمكنك حفظه كمرفق دون تحليله.",
            code="FILE_TOO_LARGE",
            status_code=413,
        )

    stream = getattr(file_storage, "stream", file_storage)
    try:
        stream.seek(0)
    except (AttributeError, OSError):
        pass

    chunks: list[bytes] = []
    total = 0
    while True:
        chunk = stream.read(min(64 * 1024, max_bytes + 1 - total))
        if not chunk:
            break
        chunks.append(chunk)
        total += len(chunk)
        if total > max_bytes:
            raise CorrespondenceIntakeError(
                f"حجم المرفق أكبر من الحد المسموح للتحليل الذكي "
                f"({_human_size_limit(max_bytes)}). يمكنك حفظه كمرفق دون تحليله.",
                code="FILE_TOO_LARGE",
                status_code=413,
            )

    payload = b"".join(chunks)
    if not payload:
        raise CorrespondenceIntakeError(
            "المرفق فارغ ولا يمكن تحليله.", code="EMPTY_FILE", status_code=400
        )
    return payload


def _safe_filename(filename: str | None) -> str:
    name = str(filename or "").replace("\\", "/").rsplit("/", 1)[-1].strip()
    return name or "attachment"


def _normalize(value: str | None) -> str:
    value = unicodedata.normalize("NFKC", str(value or ""))
    value = _ARABIC_DIACRITICS.sub("", value)
    value = (
        value.replace("أ", "ا")
        .replace("إ", "ا")
        .replace("آ", "ا")
        .replace("ى", "ي")
        .replace("ؤ", "و")
        .replace("ئ", "ي")
        .replace("ة", "ه")
    )
    value = re.sub(r"[^\w\u0600-\u06FF]+", " ", value, flags=re.UNICODE)
    return re.sub(r"\s+", " ", value).strip().casefold()


def _clean_text(value: str | None, max_chars: int) -> str:
    value = unicodedata.normalize("NFKC", str(value or "")).replace("\x00", "")
    lines: list[str] = []
    previous_blank = False
    for raw_line in value.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = re.sub(r"[\t \u00A0]+", " ", raw_line).strip()
        if not line:
            if lines and not previous_blank:
                lines.append("")
            previous_blank = True
            continue
        lines.append(line)
        previous_blank = False
        if sum(len(item) + 1 for item in lines) >= max_chars:
            break
    return "\n".join(lines).strip()[:max_chars]


def _decode_text(payload: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "cp1256"):
        try:
            return payload.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return payload.decode("utf-8", errors="replace")


def _preflight_zip(payload: bytes, *, max_uncompressed_bytes: int = 50 * 1024 * 1024) -> None:
    """Reject malformed, encrypted, or excessively expanded Office packages."""
    try:
        with zipfile.ZipFile(BytesIO(payload)) as package:
            total = 0
            for info in package.infolist():
                if info.flag_bits & 0x1:
                    raise CorrespondenceIntakeError(
                        "لا يمكن تحليل ملف Office مشفّر.",
                        code="ENCRYPTED_FILE",
                        status_code=422,
                    )
                total += max(0, int(info.file_size or 0))
                if total > max_uncompressed_bytes:
                    raise CorrespondenceIntakeError(
                        "حجم محتويات ملف Office بعد فك الضغط أكبر من الحد الآمن.",
                        code="EXPANDED_FILE_TOO_LARGE",
                        status_code=413,
                    )
    except CorrespondenceIntakeError:
        raise
    except (OSError, zipfile.BadZipFile) as exc:
        raise CorrespondenceIntakeError(
            "ملف Office تالف أو ليس بالصيغة المتوقعة.",
            code="INVALID_OFFICE_FILE",
            status_code=422,
        ) from exc


def _extract_pdf(
    payload: bytes,
    max_chars: int,
    max_pages: int,
    ocr_config: OcrConfig | None = None,
) -> tuple[str, list[str], dict]:
    ocr_config = _bounded_ocr_config(ocr_config)
    tesseract_command = _resolve_tesseract_command(ocr_config)
    try:
        import fitz

        document = fitz.open(stream=payload, filetype="pdf")
    except Exception as exc:
        raise CorrespondenceIntakeError(
            "تعذر فتح ملف PDF. قد يكون تالفاً أو محمياً بكلمة مرور.",
            code="INVALID_PDF",
            status_code=422,
        ) from exc

    warnings: list[str] = []
    if getattr(document, "needs_pass", False):
        document.close()
        raise CorrespondenceIntakeError(
            "ملف PDF محمي بكلمة مرور ولا يمكن تحليله.",
            code="ENCRYPTED_FILE",
            status_code=422,
        )

    page_count = int(getattr(document, "page_count", 0) or 0)
    if page_count > max_pages:
        warnings.append(f"تم تحليل أول {max_pages} صفحة فقط من أصل {page_count} صفحة.")

    parts: list[str] = []
    current_length = 0
    ocr_attempted_pages = 0
    ocr_pages = 0
    scanned_pages = 0
    ocr_deadline = time.monotonic() + ocr_config.timeout_seconds
    ocr_failure_reported = False
    ocr_oversize_reported = False
    try:
        for page_number in range(min(page_count, max_pages)):
            page = document.load_page(page_number)
            page_text = page.get_text("text") or ""
            if not page_text.strip():
                scanned_pages += 1
                if (
                    tesseract_command
                    and ocr_attempted_pages < ocr_config.max_pages
                    and not ocr_failure_reported
                ):
                    remaining = ocr_deadline - time.monotonic()
                    if remaining <= 0:
                        warnings.append("توقّف OCR بعد بلوغ المهلة الإجمالية للتحليل.")
                        ocr_failure_reported = True
                    else:
                        try:
                            estimated_width = max(1, int(page.rect.width * ocr_config.dpi / 72))
                            estimated_height = max(1, int(page.rect.height * ocr_config.dpi / 72))
                            if (
                                estimated_width * estimated_height
                                > ocr_config.max_image_pixels
                            ):
                                if not ocr_oversize_reported:
                                    warnings.append(
                                        "تجاوزت أبعاد إحدى صفحات PDF الحد الآمن لـOCR؛ تم تخطيها."
                                    )
                                    ocr_oversize_reported = True
                                continue
                            ocr_attempted_pages += 1
                            pixmap = page.get_pixmap(
                                dpi=ocr_config.dpi,
                                colorspace=fitz.csGRAY,
                                alpha=False,
                            )
                            page_text = _run_tesseract_png(
                                pixmap.tobytes("png"),
                                tesseract_command,
                                ocr_config,
                                remaining,
                            )
                            if page_text.strip():
                                ocr_pages += 1
                        except _OcrRuntimeError as exc:
                            warnings.append(str(exc))
                            ocr_failure_reported = True
                        except Exception:
                            warnings.append("تعذر تجهيز إحدى صفحات PDF لإجراء OCR.")
                            ocr_failure_reported = True
            if page_text.strip():
                parts.append(page_text)
                current_length += len(page_text)
            if current_length >= max_chars:
                warnings.append("تم اختصار النص المستخرج إلى الحد الآمن للتحليل.")
                break
    finally:
        document.close()

    text = _clean_text("\n".join(parts), max_chars)
    if not text:
        if ocr_config.enabled and not tesseract_command:
            warnings.append(
                "ملف PDF ممسوح ضوئياً، لكن محرك Tesseract المحلي غير متوفر؛ أدخل البيانات يدوياً أو ثبّت OCR."
            )
        else:
            warnings.append(
                "لم يُعثر على نص داخل PDF بعد التحليل؛ راجع جودة المسح أو أدخل البيانات يدوياً."
            )
    elif ocr_pages:
        warnings.append(f"تم استخدام OCR المحلي لاستخراج النص من {ocr_pages} صفحة ممسوحة.")
    if scanned_pages > ocr_attempted_pages and ocr_attempted_pages >= ocr_config.max_pages:
        warnings.append(f"تم تطبيق OCR على أول {ocr_config.max_pages} صفحة ممسوحة فقط.")
    return text, warnings, {
        "page_count": page_count,
        "ocr_enabled": ocr_config.enabled,
        "ocr_available": bool(tesseract_command),
        "ocr_used": bool(ocr_pages),
        "ocr_pages": ocr_pages,
        "ocr_attempted_pages": ocr_attempted_pages,
    }


def _extract_image_with_ocr(
    payload: bytes,
    max_chars: int,
    ocr_config: OcrConfig | None,
) -> tuple[str, list[str], dict]:
    ocr_config = _bounded_ocr_config(ocr_config)
    command = _resolve_tesseract_command(ocr_config)
    if not ocr_config.enabled:
        return "", [
            "OCR المحلي غير مفعّل؛ استخدم اسم الملف كمقترح وأدخل بقية البيانات يدوياً."
        ], {
            "ocr_enabled": False,
            "ocr_available": False,
            "ocr_used": False,
            "ocr_pages": 0,
        }
    if not command:
        return "", [
            "محرك Tesseract المحلي غير متوفر؛ أدخل البيانات يدوياً أو ثبّت OCR العربي."
        ], {
            "ocr_enabled": True,
            "ocr_available": False,
            "ocr_used": False,
            "ocr_pages": 0,
        }

    frames = _image_to_png_frames(payload, ocr_config)
    text, warnings = _ocr_png_frames(frames, ocr_config, command=command)
    text = _clean_text(text, max_chars)
    if text:
        warnings.append(
            f"تم استخدام OCR المحلي لاستخراج النص من {len(frames)} صورة/صفحة."
        )
    else:
        warnings.append(
            "لم يُعثر على نص واضح في الصورة بعد OCR؛ راجع جودة الصورة أو أدخل البيانات يدوياً."
        )
    return text, warnings, {
        "ocr_enabled": True,
        "ocr_available": True,
        "ocr_used": bool(text),
        "ocr_pages": len(frames) if text else 0,
    }


def _extract_docx(payload: bytes, max_chars: int) -> tuple[str, list[str], dict]:
    _preflight_zip(payload)
    try:
        from docx import Document

        document = Document(BytesIO(payload))
    except Exception as exc:
        raise CorrespondenceIntakeError(
            "تعذر قراءة ملف Word.", code="INVALID_DOCX", status_code=422
        ) from exc

    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables[:30]:
        for row in table.rows[:200]:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
            if sum(len(item) + 1 for item in parts) >= max_chars:
                break

    properties = document.core_properties
    metadata = {
        "subject": (properties.subject or properties.title or "").strip(),
        "sender": (properties.author or "").strip(),
    }
    return _clean_text("\n".join(parts), max_chars), [], metadata


def _resolve_legacy_word_converter() -> str | None:
    """Find a local LibreOffice executable without invoking a shell."""
    configured = str(os.getenv("CORR_INTAKE_LIBREOFFICE_CMD") or "").strip()
    candidates = [
        (shutil.which(configured) or configured) if configured else "",
        shutil.which("soffice") or "",
        shutil.which("libreoffice") or "",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    return next(
        (os.path.abspath(path) for path in candidates if path and os.path.isfile(path)),
        None,
    )


def _extract_rtf_text(payload: bytes, max_chars: int) -> str:
    """Extract reviewable text from an RTF document saved with a .doc suffix."""
    source = payload.decode("latin-1", errors="replace")
    source = re.sub(r"\\par[d]?\b", "\n", source, flags=re.IGNORECASE)
    source = re.sub(r"\\tab\b", "\t", source, flags=re.IGNORECASE)

    def unicode_character(match) -> str:
        value = int(match.group(1))
        if value < 0:
            value += 65536
        try:
            return chr(value)
        except ValueError:
            return ""

    source = re.sub(r"\\u(-?\d+)\??", unicode_character, source)

    def encoded_character(match) -> str:
        try:
            return bytes([int(match.group(1), 16)]).decode("cp1256")
        except (ValueError, UnicodeDecodeError):
            return ""

    source = re.sub(r"\\'([0-9a-fA-F]{2})", encoded_character, source)
    source = re.sub(r"\\[A-Za-z]+-?\d* ?", "", source)
    source = source.replace(r"\{", "{").replace(r"\}", "}").replace(r"\\", "\\")
    source = source.replace("{", "").replace("}", "")
    return _clean_text(source, max_chars)


def _compound_file_stream(payload: bytes, stream_name: str) -> bytes | None:
    """Read one stream from an OLE Compound File without external packages."""
    signature = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
    if len(payload) < 512 or not payload.startswith(signature):
        return None

    free_sector = 0xFFFFFFFF
    end_of_chain = 0xFFFFFFFE
    special_sector_min = 0xFFFFFFFA

    try:
        read_u16 = lambda offset: struct.unpack_from("<H", payload, offset)[0]
        read_u32 = lambda offset: struct.unpack_from("<I", payload, offset)[0]
        if read_u16(28) != 0xFFFE:
            return None
        major_version = read_u16(26)
        sector_size = 1 << read_u16(30)
        mini_sector_size = 1 << read_u16(32)
        if sector_size not in (512, 4096) or mini_sector_size != 64:
            return None

        sector_count = max(0, len(payload) // sector_size - 1)

        def sector(sector_id: int) -> bytes:
            if not 0 <= sector_id < sector_count:
                raise ValueError("OLE sector is outside the file")
            start = (sector_id + 1) * sector_size
            return payload[start : start + sector_size]

        fat_sector_count = read_u32(44)
        fat_sector_ids = [
            read_u32(76 + index * 4)
            for index in range(109)
            if read_u32(76 + index * 4) < special_sector_min
        ]

        difat_sector_id = read_u32(68)
        difat_sector_count = read_u32(72)
        seen_difat: set[int] = set()
        entries_per_difat_sector = sector_size // 4
        for _ in range(min(difat_sector_count, sector_count)):
            if (
                difat_sector_id >= special_sector_min
                or difat_sector_id in seen_difat
            ):
                break
            seen_difat.add(difat_sector_id)
            values = struct.unpack(
                f"<{entries_per_difat_sector}I", sector(difat_sector_id)
            )
            fat_sector_ids.extend(
                value for value in values[:-1] if value < special_sector_min
            )
            difat_sector_id = values[-1]

        fat_sector_ids = fat_sector_ids[:fat_sector_count]
        if not fat_sector_ids:
            return None
        entries_per_sector = sector_size // 4
        fat: list[int] = []
        for sector_id in fat_sector_ids:
            fat.extend(struct.unpack(f"<{entries_per_sector}I", sector(sector_id)))

        def chain(first_sector_id: int, allocation_table: list[int]) -> list[int]:
            result: list[int] = []
            seen: set[int] = set()
            sector_id = first_sector_id
            limit = len(allocation_table)
            while (
                sector_id < special_sector_min
                and sector_id not in seen
                and len(result) < limit
            ):
                if sector_id >= len(allocation_table):
                    raise ValueError("OLE allocation chain is invalid")
                result.append(sector_id)
                seen.add(sector_id)
                sector_id = allocation_table[sector_id]
            if sector_id not in (end_of_chain, free_sector) and sector_id < special_sector_min:
                raise ValueError("OLE allocation chain is cyclic")
            return result

        def regular_stream(first_sector_id: int, size: int | None = None) -> bytes:
            value = b"".join(sector(item) for item in chain(first_sector_id, fat))
            return value if size is None else value[:size]

        directory_bytes = regular_stream(read_u32(48))
        entries: dict[str, tuple[int, int, int]] = {}
        root_entry: tuple[int, int, int] | None = None
        for offset in range(0, len(directory_bytes), 128):
            entry = directory_bytes[offset : offset + 128]
            if len(entry) < 128:
                break
            name_length = struct.unpack_from("<H", entry, 64)[0]
            entry_type = entry[66]
            if not 2 <= name_length <= 64 or name_length % 2:
                continue
            name = entry[: name_length - 2].decode("utf-16le", errors="strict")
            first_sector_id = struct.unpack_from("<I", entry, 116)[0]
            stream_size = struct.unpack_from("<Q", entry, 120)[0]
            if major_version == 3:
                stream_size &= 0xFFFFFFFF
            details = (entry_type, first_sector_id, stream_size)
            entries[name.casefold()] = details
            if entry_type == 5:
                root_entry = details

        details = entries.get(stream_name.casefold())
        if details is None or details[0] != 2:
            return None
        _, first_sector_id, stream_size = details
        if stream_size > len(payload):
            return None

        mini_stream_cutoff = read_u32(56)
        if stream_size >= mini_stream_cutoff:
            return regular_stream(first_sector_id, stream_size)

        if root_entry is None:
            return None
        _, root_first_sector_id, root_size = root_entry
        root_stream = regular_stream(root_first_sector_id, root_size)
        mini_fat_bytes = regular_stream(read_u32(60), read_u32(64) * sector_size)
        mini_fat = list(
            struct.unpack(
                f"<{len(mini_fat_bytes) // 4}I",
                mini_fat_bytes[: len(mini_fat_bytes) // 4 * 4],
            )
        )
        mini_parts: list[bytes] = []
        for mini_sector_id in chain(first_sector_id, mini_fat):
            start = mini_sector_id * mini_sector_size
            if start >= len(root_stream):
                raise ValueError("OLE mini stream is invalid")
            mini_parts.append(root_stream[start : start + mini_sector_size])
        return b"".join(mini_parts)[:stream_size]
    except (OverflowError, UnicodeDecodeError, ValueError, struct.error):
        return None


def _legacy_doc_candidate_strings(payload: bytes) -> list[tuple[int, str]]:
    """Return likely human-readable strings embedded in a binary .doc stream."""
    unicode_candidates: list[tuple[int, str]] = []

    def allowed_character(character: str) -> bool:
        codepoint = ord(character)
        return (
            codepoint == 7  # Word table cell/end-of-row marker.
            or character in "\r\n\t "
            or character in "،؛؟,.!?():/\\-_@'\"%&#+=|[]{}"
            or 0x2010 <= codepoint <= 0x2015
            or 0x2018 <= codepoint <= 0x201F
            or 0x30 <= codepoint <= 0x39
            or 0x41 <= codepoint <= 0x5A
            or 0x61 <= codepoint <= 0x7A
            or 0x00C0 <= codepoint <= 0x024F
            or 0x0600 <= codepoint <= 0x06FF
            or 0x0750 <= codepoint <= 0x077F
            or 0x08A0 <= codepoint <= 0x08FF
            or 0xFB50 <= codepoint <= 0xFDFF
            or 0xFE70 <= codepoint <= 0xFEFF
        )

    def useful_text(value: str) -> str:
        value = value.replace("\u200e", "").replace("\u200f", "")
        value = re.sub(r"\x07\s*\x07", "\n", value)
        value = value.replace("\x07", " | ")
        value = re.sub(r"[\x00-\x06\x08\x0b\x0c\x0e-\x1f]+", " ", value)
        value = re.sub(r"[\t \u00a0]+", " ", value).strip(" \t|_-\x00")
        if len(value) < 4 or len(value) > 10_000:
            return ""
        letters = sum(character.isalpha() for character in value)
        if letters < 3:
            return ""

        alphanumeric = [
            character.casefold() for character in value if character.isalnum()
        ]
        distinct_alphanumeric = len(set(alphanumeric))
        if len(alphanumeric) >= 3 and distinct_alphanumeric == 1:
            return ""
        if len(alphanumeric) >= 6 and distinct_alphanumeric <= 2:
            return ""
        if len(alphanumeric) >= 30 and distinct_alphanumeric < 8:
            return ""
        if len(value) > 32 and not any(character.isspace() for character in value):
            return ""

        maximum_run = 0
        current_run = 0
        previous = ""
        for character in alphanumeric:
            if character == previous:
                current_run += 1
            else:
                previous = character
                current_run = 1
            maximum_run = max(maximum_run, current_run)
        if len(alphanumeric) >= 12 and maximum_run > max(12, len(alphanumeric) * 0.45):
            return ""

        normalized = re.sub(r"\s+", " ", value).casefold()
        metadata_markers = (
            "default paragraph font",
            "document summary information",
            "schemas.microsoft.com",
            "summaryinformation",
            "table grid",
            "table normal",
            "worddocument",
            "xmlns",
        )
        if any(marker in normalized for marker in metadata_markers):
            return ""
        return value

    # Uncompressed Word pieces commonly store Unicode text as UTF-16LE.
    for alignment in (0, 1):
        start = None
        buffer: list[str] = []
        for offset in range(alignment, len(payload) - 1, 2):
            codepoint = payload[offset] | (payload[offset + 1] << 8)
            character = chr(codepoint)
            if allowed_character(character):
                if start is None:
                    start = offset
                buffer.append(character)
                continue
            if buffer:
                value = useful_text("".join(buffer))
                if value:
                    unicode_candidates.append((start or 0, value))
            start = None
            buffer = []
        if buffer:
            value = useful_text("".join(buffer))
            if value:
                unicode_candidates.append((start or 0, value))

    # Avoid interpreting binary bytes a second time when clear Unicode content exists.
    unicode_letter_count = sum(
        sum(character.isalpha() for character in value)
        for _offset, value in unicode_candidates
    )
    if unicode_letter_count >= 100:
        return sorted(unicode_candidates, key=lambda item: item[0])

    candidates = list(unicode_candidates)
    # Compressed Word pieces store characters in a single-byte Windows codepage.
    for match in re.finditer(rb"[\x09\x0A\x0D\x20-\x7E\x80-\xFF]{6,}", payload):
        raw = match.group(0)
        for encoding in ("cp1256", "cp1252"):
            try:
                decoded = raw.decode(encoding)
                value = useful_text(
                    "".join(character for character in decoded if allowed_character(character))
                )
            except UnicodeDecodeError:
                value = ""
            if value:
                candidates.append((match.start(), value))
                break

    return sorted(candidates, key=lambda item: item[0])


def _extract_legacy_doc_fallback(payload: bytes, max_chars: int) -> str:
    document_stream = _compound_file_stream(payload, "WordDocument")
    searchable_payload = document_stream if document_stream is not None else payload
    parts: list[str] = []
    seen: set[str] = set()
    ignored = {
        "worddocument",
        "summaryinformation",
        "documentsummaryinformation",
        "microsoft office word",
    }
    for _offset, value in _legacy_doc_candidate_strings(searchable_payload):
        normalized = re.sub(r"\s+", " ", value).strip().casefold()
        if not normalized or normalized in ignored or normalized in seen:
            continue
        seen.add(normalized)
        if parts:
            previous_last_line = parts[-1].rstrip().rsplit("\n", 1)[-1]
            current_first_line = value.lstrip().split("\n", 1)[0]
            if previous_last_line.count("|") == current_first_line.count("|") == 1:
                parts[-1] = f"{parts[-1].rstrip()} {value.lstrip()}"
                continue
        parts.append(value)
        if sum(len(part) + 1 for part in parts) >= max_chars:
            break
    return _clean_text("\n".join(parts), max_chars)


def _convert_legacy_doc_to_docx(payload: bytes, converter: str) -> bytes:
    with tempfile.TemporaryDirectory(prefix="corr_legacy_doc_") as temporary_directory:
        input_path = os.path.join(temporary_directory, "legacy.doc")
        output_path = os.path.join(temporary_directory, "legacy.docx")
        profile_path = Path(temporary_directory, "libreoffice-profile")
        profile_path.mkdir()
        with open(input_path, "wb") as handle:
            handle.write(payload)
        try:
            completed = subprocess.run(
                [
                    converter,
                    f"-env:UserInstallation={profile_path.as_uri()}",
                    "--headless",
                    "--nologo",
                    "--nodefault",
                    "--nofirststartwizard",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    temporary_directory,
                    input_path,
                ],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=60,
                check=False,
                shell=False,
            )
        except (OSError, subprocess.TimeoutExpired) as exc:
            raise CorrespondenceIntakeError(
                "تعذر تشغيل محول Word المحلي.",
                code="LEGACY_DOC_CONVERSION_FAILED",
                status_code=422,
            ) from exc
        if completed.returncode != 0 or not os.path.isfile(output_path):
            raise CorrespondenceIntakeError(
                "تعذر تحويل ملف Word القديم إلى صيغة قابلة للتحليل.",
                code="LEGACY_DOC_CONVERSION_FAILED",
                status_code=422,
            )
        if os.path.getsize(output_path) > 50 * 1024 * 1024:
            raise CorrespondenceIntakeError(
                "حجم ملف Word بعد التحويل أكبر من الحد الآمن.",
                code="EXPANDED_FILE_TOO_LARGE",
                status_code=413,
            )
        with open(output_path, "rb") as handle:
            return handle.read()


def _extract_legacy_doc(payload: bytes, max_chars: int) -> tuple[str, list[str], dict]:
    # Some systems give an OOXML or RTF file the old .doc extension.
    if payload.startswith(b"PK"):
        text, warnings, metadata = _extract_docx(payload, max_chars)
        return text, ["تم التعرف على الملف كـWord حديث رغم امتداده .doc.", *warnings], metadata
    if payload.lstrip().startswith(b"{\\rtf"):
        return _extract_rtf_text(payload, max_chars), [], {}

    ole_signature = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
    if not payload.startswith(ole_signature):
        raise CorrespondenceIntakeError(
            "ملف Word القديم تالف أو ليس بصيغة .doc المتوقعة.",
            code="INVALID_DOC",
            status_code=422,
        )

    converter = _resolve_legacy_word_converter()
    conversion_warning = ""
    if converter:
        try:
            converted_payload = _convert_legacy_doc_to_docx(payload, converter)
            text, warnings, metadata = _extract_docx(converted_payload, max_chars)
            return text, ["تم تحويل ملف .doc محليًا قبل تحليله.", *warnings], metadata
        except CorrespondenceIntakeError as exc:
            conversion_warning = exc.message

    text = _extract_legacy_doc_fallback(payload, max_chars)
    warnings = []
    if text:
        warnings.append(
            "تم استخراج نص ملف .doc مباشرةً بالطريقة الاحتياطية؛ راجع المقترحات قبل اعتمادها."
        )
    else:
        warnings.append(
            conversion_warning
            or "لم يُعثر على نص واضح داخل ملف .doc؛ ثبّت LibreOffice على الخادم لتحويل الملفات القديمة بدقة أعلى."
        )
    return text, warnings, {}


def _extract_xlsx(payload: bytes, max_chars: int) -> tuple[str, list[str], dict]:
    _preflight_zip(payload)
    try:
        from openpyxl import load_workbook

        workbook = load_workbook(BytesIO(payload), read_only=True, data_only=True)
    except Exception as exc:
        raise CorrespondenceIntakeError(
            "تعذر قراءة ملف Excel.", code="INVALID_XLSX", status_code=422
        ) from exc

    parts: list[str] = []
    warnings: list[str] = []
    try:
        worksheets = workbook.worksheets[:5]
        if len(workbook.worksheets) > len(worksheets):
            warnings.append("تم تحليل أول خمس أوراق عمل فقط.")
        for sheet in worksheets:
            parts.append(f"[{sheet.title}]")
            for row_number, row in enumerate(
                sheet.iter_rows(min_row=1, max_row=200, max_col=40, values_only=True),
                start=1,
            ):
                values = [str(value).strip() for value in row if value not in (None, "")]
                if values:
                    parts.append(" | ".join(values))
                if row_number >= 200 or sum(len(item) + 1 for item in parts) >= max_chars:
                    break
            if sum(len(item) + 1 for item in parts) >= max_chars:
                warnings.append("تم اختصار بيانات Excel إلى الحد الآمن للتحليل.")
                break
    finally:
        workbook.close()
    return _clean_text("\n".join(parts), max_chars), warnings, {}


def _extract_pptx(payload: bytes, max_chars: int) -> tuple[str, list[str], dict]:
    _preflight_zip(payload)
    try:
        with zipfile.ZipFile(BytesIO(payload)) as package:
            slide_names = sorted(
                (
                    name
                    for name in package.namelist()
                    if re.fullmatch(r"ppt/slides/slide\d+\.xml", name)
                ),
                key=lambda name: int(re.search(r"(\d+)\.xml$", name).group(1)),
            )
            parts: list[str] = []
            for slide_name in slide_names[:100]:
                xml_text = package.read(slide_name).decode("utf-8", errors="replace")
                slide_parts = [
                    html.unescape(re.sub(r"<[^>]+>", "", match)).strip()
                    for match in re.findall(r"<a:t(?:\s[^>]*)?>.*?</a:t>", xml_text, re.DOTALL)
                ]
                slide_parts = [part for part in slide_parts if part]
                if slide_parts:
                    parts.append(" | ".join(slide_parts))
                if sum(len(item) + 1 for item in parts) >= max_chars:
                    break
    except (OSError, zipfile.BadZipFile) as exc:
        raise CorrespondenceIntakeError(
            "تعذر قراءة ملف PowerPoint.", code="INVALID_PPTX", status_code=422
        ) from exc
    return _clean_text("\n".join(parts), max_chars), [], {}


def _parse_eml_message(payload: bytes):
    try:
        return BytesParser(policy=policy.default).parsebytes(payload)
    except Exception as exc:
        raise CorrespondenceIntakeError(
            "تعذر قراءة ملف البريد الإلكتروني.", code="INVALID_EML", status_code=422
        ) from exc


def _email_attachment_filename(part, ordinal: int) -> str:
    filename = str(part.get_filename() or "").strip()
    filename = filename.replace("\x00", "").replace("\\", "/").rsplit("/", 1)[-1].strip()
    if filename:
        return filename

    mimetype = str(part.get_content_type() or "").strip().lower()
    suffix = mimetypes.guess_extension(mimetype) or ""
    return f"مرفق البريد {ordinal}{suffix}"


def extract_eml_attachments(
    payload: bytes,
    *,
    max_attachments: int = 50,
    max_total_bytes: int = 25 * 1024 * 1024,
    max_attachment_bytes: int = 25 * 1024 * 1024,
) -> EmailAttachmentExtraction:
    """Extract EML file attachments with bounded count and payload size.

    Named inline parts are included because mail clients frequently use those
    for documents and images shown in the message. The parent EML itself is
    deliberately not included; callers keep it as the original upload.
    """
    message = _parse_eml_message(payload)
    max_attachments = max(1, min(int(max_attachments or 1), 200))
    max_total_bytes = max(1, int(max_total_bytes or 1))
    max_attachment_bytes = max(1, int(max_attachment_bytes or 1))

    attachments: list[ExtractedEmailAttachment] = []
    warnings: list[str] = []
    total_bytes = 0

    for part in message.walk() if message.is_multipart() else [message]:
        disposition = (part.get_content_disposition() or "").lower()
        has_filename = bool(part.get_filename())
        if disposition != "attachment" and not has_filename:
            continue

        ordinal = len(attachments) + 1
        filename = _email_attachment_filename(part, ordinal)
        try:
            if part.is_multipart():
                attachment_payload = part.as_bytes(policy=policy.default)
            else:
                attachment_payload = part.get_payload(decode=True)
        except Exception:
            attachment_payload = None

        if attachment_payload is None:
            warnings.append(f"تعذر استخراج مرفق البريد: {filename}.")
            continue
        if len(attachment_payload) > max_attachment_bytes:
            warnings.append(
                f"لم يُستخرج المرفق {filename} لأنه يتجاوز الحد المسموح لكل مرفق."
            )
            continue
        if total_bytes + len(attachment_payload) > max_total_bytes:
            warnings.append(
                "لم تُستخرج بقية مرفقات البريد لأنها تجاوزت الحجم الإجمالي المسموح."
            )
            break
        if len(attachments) >= max_attachments:
            warnings.append(
                "لم تُستخرج بقية مرفقات البريد لأنها تجاوزت العدد المسموح."
            )
            break

        attachments.append(
            ExtractedEmailAttachment(
                filename=filename,
                payload=attachment_payload,
                mimetype=str(part.get_content_type() or "").strip() or None,
            )
        )
        total_bytes += len(attachment_payload)

    return EmailAttachmentExtraction(tuple(attachments), tuple(warnings))


def _extract_eml(payload: bytes, max_chars: int) -> tuple[str, list[str], dict]:
    message = _parse_eml_message(payload)
    parts: list[str] = []
    for part in message.walk() if message.is_multipart() else [message]:
        if part.get_content_disposition() == "attachment":
            continue
        content_type = part.get_content_type()
        if content_type not in {"text/plain", "text/html"}:
            continue
        try:
            content = part.get_content()
        except Exception:
            raw = part.get_payload(decode=True) or b""
            content = _decode_text(raw)
        if content_type == "text/html":
            try:
                from bs4 import BeautifulSoup

                content = BeautifulSoup(str(content), "html.parser").get_text("\n")
            except Exception:
                content = re.sub(r"<[^>]+>", " ", str(content))
        parts.append(str(content))
        if sum(len(item) + 1 for item in parts) >= max_chars:
            break

    sender_name, sender_address = parseaddr(str(message.get("From") or ""))
    recipients = getaddresses(message.get_all("To", []))
    recipient_name = ""
    if recipients:
        recipient_name, recipient_address = recipients[0]
        recipient_name = (recipient_name or recipient_address or "").strip()
    metadata = {
        "subject": str(message.get("Subject") or "").strip(),
        "sender": (sender_name or sender_address or "").strip(),
        "recipient": recipient_name,
        "message_date": str(message.get("Date") or "").strip(),
    }
    return _clean_text("\n".join(parts), max_chars), [], metadata


def extract_attachment_text(
    payload: bytes,
    filename: str,
    *,
    max_chars: int = 20_000,
    max_pdf_pages: int = 40,
    ocr_config: OcrConfig | None = None,
) -> dict:
    """Extract bounded text and metadata from a supported attachment."""
    safe_name = _safe_filename(filename)
    extension = Path(safe_name).suffix.lower()
    warnings: list[str] = []
    metadata: dict = {}

    if extension == ".pdf":
        text, warnings, metadata = _extract_pdf(
            payload,
            max_chars,
            max_pdf_pages,
            ocr_config,
        )
        format_label = "PDF"
    elif extension == ".docx":
        text, warnings, metadata = _extract_docx(payload, max_chars)
        format_label = "Word"
    elif extension == ".doc":
        text, warnings, metadata = _extract_legacy_doc(payload, max_chars)
        format_label = "Word (DOC)"
    elif extension == ".xlsx":
        text, warnings, metadata = _extract_xlsx(payload, max_chars)
        format_label = "Excel"
    elif extension == ".pptx":
        text, warnings, metadata = _extract_pptx(payload, max_chars)
        format_label = "PowerPoint"
    elif extension == ".eml":
        text, warnings, metadata = _extract_eml(payload, max_chars)
        format_label = "Email"
    elif extension in {".txt", ".csv", ".json", ".xml", ".html", ".htm"}:
        decoded = _decode_text(payload)
        if extension in {".html", ".htm"}:
            try:
                from bs4 import BeautifulSoup

                decoded = BeautifulSoup(decoded, "html.parser").get_text("\n")
            except Exception:
                decoded = re.sub(r"<[^>]+>", " ", decoded)
        text = _clean_text(decoded, max_chars)
        format_label = "Text"
    elif extension in IMAGE_EXTENSIONS:
        text, warnings, metadata = _extract_image_with_ocr(
            payload,
            max_chars,
            ocr_config,
        )
        format_label = "Image"
    else:
        raise CorrespondenceIntakeError(
            "صيغة المرفق غير مدعومة للتحليل الذكي حالياً. يمكن رفعها وحفظها بالطريقة العادية.",
            code="UNSUPPORTED_FILE_TYPE",
            status_code=415,
        )

    return {
        "filename": safe_name,
        "extension": extension,
        "format": format_label,
        "text": text,
        "metadata": metadata,
        "warnings": warnings,
        "ocr": {
            "enabled": bool(metadata.get("ocr_enabled", False)),
            "available": bool(metadata.get("ocr_available", False)),
            "used": bool(metadata.get("ocr_used", False)),
            "pages": int(metadata.get("ocr_pages", 0) or 0),
        },
    }


def _suggestion(value, confidence: float, reason: str, **extra) -> dict | None:
    if value in (None, ""):
        return None
    result = {
        "value": value,
        "confidence": round(max(0.0, min(1.0, float(confidence))), 2),
        "reason": reason,
    }
    result.update(extra)
    return result


def _choice_label(choice: dict) -> str:
    return str(choice.get("match_label") or choice.get("label") or choice.get("value") or "").strip()


def _choice_match_text(choice: dict) -> str:
    return str(choice.get("match_text") or _choice_label(choice)).strip()


def _match_choice(text: str, choices: list[dict]) -> tuple[dict | None, float]:
    normalized_text = _normalize(text)
    normalized_head = _normalize(text[:2500])
    best_choice = None
    best_score = 0.0
    best_length = 0
    for choice in choices:
        label = _choice_match_text(choice)
        normalized_label = _normalize(re.sub(r"\s*\([^)]*\)\s*$", "", label))
        if len(normalized_label) < 3:
            continue
        score = 0.0
        if normalized_label in normalized_head:
            score = 0.96
        elif normalized_label in normalized_text:
            score = 0.9
        else:
            tokens = {
                token
                for token in normalized_label.split()
                if len(token) >= 3 and token not in _MEANINGLESS_TOKENS
            }
            if len(tokens) >= 2:
                overlap = sum(1 for token in tokens if token in normalized_text)
                if overlap == len(tokens):
                    score = 0.78
                elif overlap / len(tokens) >= 0.75:
                    score = 0.68
        if score > best_score or (
            score == best_score and score > 0 and len(normalized_label) > best_length
        ):
            best_choice = choice
            best_score = score
            best_length = len(normalized_label)
    return best_choice, best_score


def _match_known_value(value: str, choices: list[dict]) -> dict | None:
    normalized_value = _normalize(value)
    if not normalized_value:
        return None
    best = None
    best_length = 0
    for choice in choices:
        label = _choice_label(choice)
        normalized_label = _normalize(label)
        if not normalized_label:
            continue
        if normalized_label == normalized_value or (
            len(normalized_label) >= 4
            and (normalized_label in normalized_value or normalized_value in normalized_label)
        ):
            if len(normalized_label) > best_length:
                best = choice
                best_length = len(normalized_label)
    return best


def _explicit_line_value(text: str, labels: tuple[str, ...], max_length: int = 250) -> str:
    label_pattern = "|".join(re.escape(label) for label in labels)
    pattern = re.compile(
        rf"^(?:{label_pattern})[ \t]*[:：\-][ \t]*(.+)$",
        flags=re.IGNORECASE | re.MULTILINE,
    )
    match = pattern.search(text[:8000])
    if not match:
        return ""
    value = re.sub(r"\s+", " ", match.group(1)).strip(" -:：|\t")
    return value[:max_length]


def _suggest_subject(
    text: str,
    filename: str,
    metadata: dict,
    *,
    default_value: str = "وارد جديد",
) -> dict:
    metadata_subject = str(metadata.get("subject") or "").strip()
    if metadata_subject:
        return _suggestion(metadata_subject[:500], 0.99, "عنوان البريد أو بيانات المستند")

    explicit = _explicit_line_value(text, ("الموضوع", "الموضوع /", "بخصوص", "Subject"), 500)
    if explicit:
        return _suggestion(explicit, 0.97, "حقل الموضوع داخل المستند")

    stem = Path(filename).stem
    stem = re.sub(r"[_\-]+", " ", stem).strip()
    if stem:
        return _suggestion(stem[:500], 0.45, "اسم المرفق؛ يحتاج إلى مراجعة")
    return _suggestion(default_value, 0.2, "قيمة افتراضية؛ تحتاج إلى مراجعة")


def _suggest_sender(text: str, metadata: dict, sender_choices: list[dict]) -> dict | None:
    explicit = str(metadata.get("sender") or "").strip() or _explicit_line_value(
        text,
        ("الجهة المرسلة", "المرسل", "صادر عن", "من", "From"),
        200,
    )
    if explicit:
        known = _match_known_value(explicit, sender_choices)
        if known:
            return _suggestion(
                _choice_label(known),
                0.98,
                "جهة مرسلة مطابقة لدليل الجهات",
                select_value=str(known.get("value") or _choice_label(known)),
                is_known=True,
            )
        return _suggestion(
            explicit,
            0.82,
            "جهة مرسلة مذكورة في المستند وتحتاج إلى مراجعة",
            select_value="__OTHER__",
            is_known=False,
        )

    known, score = _match_choice(text[:6000], sender_choices)
    if known:
        return _suggestion(
            _choice_label(known),
            score,
            "مطابقة اسم جهة من دليل المرسلين",
            select_value=str(known.get("value") or _choice_label(known)),
            is_known=True,
        )
    return None


def _suggest_recipient(text: str, metadata: dict, recipient_choices: list[dict]) -> dict | None:
    explicit = str(metadata.get("recipient") or "").strip() or _explicit_line_value(
        text,
        ("الجهة المستلمة", "المرسل إليه", "موجه إلى", "إلى", "To"),
        200,
    )
    if explicit:
        known = _match_known_value(explicit, recipient_choices)
        if known:
            return _suggestion(
                _choice_label(known),
                0.98,
                "جهة مستلمة مطابقة لدليل الجهات",
                select_value=str(known.get("value") or _choice_label(known)),
                is_known=True,
            )
        return _suggestion(
            explicit,
            0.82,
            "جهة مستلمة مذكورة في المستند وتحتاج إلى مراجعة",
            select_value="__OTHER__",
            is_known=False,
        )

    known, score = _match_choice(text[:6000], recipient_choices)
    if known:
        return _suggestion(
            _choice_label(known),
            score,
            "مطابقة اسم جهة من دليل المستلمين",
            select_value=str(known.get("value") or _choice_label(known)),
            is_known=True,
        )
    return None


def _suggest_category(text: str, category_choices: list[dict]) -> dict | None:
    choice, score = _match_choice(text, category_choices)
    if choice:
        return _suggestion(
            _choice_label(choice),
            score,
            "مطابقة تصنيف مسجل مع محتوى المستند",
            select_value=str(choice.get("value") or ""),
        )

    normalized_text = _normalize(text)
    aliases = {
        "FIN": ("مالي", "ماليه", "موازنه", "فاتوره", "دفع"),
        "HR": ("موارد بشريه", "موظف", "توظيف", "اجازه", "راتب"),
        "LEGAL": ("قانوني", "قانونيه", "عقد", "اتفاقيه", "مذكره تفاهم"),
        "PROCUREMENT": ("عطاء", "مشتريات", "توريد", "مورد"),
    }
    for code, keywords in aliases.items():
        if not any(keyword in normalized_text for keyword in keywords):
            continue
        for candidate in category_choices:
            candidate_code = str(candidate.get("value") or "").upper()
            if candidate_code == code:
                return _suggestion(
                    _choice_label(candidate),
                    0.76,
                    "تصنيف مقترح من كلمات المستند",
                    select_value=candidate_code,
                )

    for candidate in category_choices:
        if str(candidate.get("value") or "").upper() == "GENERAL":
            return _suggestion(
                _choice_label(candidate),
                0.35,
                "التصنيف العام الافتراضي؛ يحتاج إلى مراجعة",
                select_value="GENERAL",
            )
    return None


def _suggest_due_date(text: str) -> dict | None:
    patterns = (
        r"(?:الموعد النهائي|الرد قبل|قبل تاريخ|بحد اقصى|بحد أقصى)[ \t]*[:：\-]?[ \t]*(\d{4}[-/]\d{1,2}[-/]\d{1,2})",
        r"(?:الموعد النهائي|الرد قبل|قبل تاريخ|بحد اقصى|بحد أقصى)[ \t]*[:：\-]?[ \t]*(\d{1,2}[-/]\d{1,2}[-/]\d{4})",
    )
    for pattern in patterns:
        match = re.search(pattern, text[:12000], flags=re.IGNORECASE)
        if not match:
            continue
        raw_value = match.group(1).replace("/", "-")
        for date_format in ("%Y-%m-%d", "%d-%m-%Y"):
            try:
                parsed = datetime.strptime(raw_value, date_format).date()
                return _suggestion(parsed.isoformat(), 0.92, "موعد نهائي مذكور في المستند")
            except ValueError:
                continue
    return None


def _suggest_reference(text: str) -> dict | None:
    pattern = re.compile(
        r"(?:رقم الكتاب|رقم المرجع|الرقم المرجعي|رقم الصادر|Reference(?: No\.?)?|Ref(?: No\.?)?)[ \t]*[:：#\-][ \t]*([A-Za-z0-9\u0600-\u06FF/_\-.]{2,80})",
        flags=re.IGNORECASE,
    )
    match = pattern.search(text[:8000])
    if not match:
        return None
    return _suggestion(match.group(1).strip(" .-"), 0.94, "رقم مرجع المستند")


def _suggest_priority(text: str) -> dict:
    normalized_text = _normalize(text[:12000])
    if any(token in normalized_text for token in ("عاجل جدا", "عاجل", "فوري", "علي وجه السرعه")):
        return _suggestion("URGENT", 0.93, "وجود عبارة تدل على الاستعجال")
    if any(token in normalized_text for token in ("هام", "مهم", "اولويه عاليه")):
        return _suggestion("HIGH", 0.82, "وجود عبارة تدل على الأهمية")
    return _suggestion("NORMAL", 0.65, "لم تظهر علامة واضحة على الاستعجال")


def _suggest_confidentiality(text: str) -> dict:
    normalized_text = _normalize(text[:12000])
    normalized_text = normalized_text.replace("غير سري", "")
    if any(token in normalized_text for token in ("سري للغايه", "سري", "خاص وسري")):
        return _suggestion("SECRET", 0.95, "وجود علامة سرية في المستند")
    return _suggestion("NORMAL", 0.65, "لم تظهر علامة سرية واضحة")


def _suggest_workflow(text: str, subject: str, workflow_choices: list[dict]) -> dict | None:
    combined = f"{subject}\n{text}"
    choice, score = _match_choice(combined, workflow_choices)
    if choice:
        return _suggestion(
            _choice_label(choice),
            score,
            "مطابقة اسم قالب المسار مع موضوع المستند",
            select_value=str(choice.get("value") or ""),
        )
    if len(workflow_choices) == 1:
        choice = workflow_choices[0]
        return _suggestion(
            _choice_label(choice),
            0.55,
            "قالب المسار النشط الوحيد؛ يحتاج إلى مراجعة",
            select_value=str(choice.get("value") or ""),
        )
    return None


def analyze_correspondence_attachment(
    payload: bytes,
    filename: str,
    *,
    sender_choices: list[dict] | None = None,
    recipient_choices: list[dict] | None = None,
    category_choices: list[dict] | None = None,
    competence_choices: list[dict] | None = None,
    workflow_choices: list[dict] | None = None,
    received_date: str | None = None,
    sent_date: str | None = None,
    direction: str = "IN",
    max_text_chars: int = 20_000,
    max_pdf_pages: int = 40,
    ocr_config: OcrConfig | None = None,
) -> dict:
    """Extract an attachment into reviewable inbound or outbound suggestions."""
    extracted = extract_attachment_text(
        payload,
        filename,
        max_chars=max_text_chars,
        max_pdf_pages=max_pdf_pages,
        ocr_config=ocr_config,
    )
    text = extracted["text"]
    metadata = extracted["metadata"]
    sender_choices = list(sender_choices or [])
    recipient_choices = list(recipient_choices or [])
    category_choices = list(category_choices or [])
    competence_choices = list(competence_choices or [])
    workflow_choices = list(workflow_choices or [])

    direction = str(direction or "IN").strip().upper()
    is_outbound = direction == "OUT"
    subject = _suggest_subject(
        text,
        extracted["filename"],
        metadata,
        default_value="صادر جديد" if is_outbound else "وارد جديد",
    )
    party = (
        _suggest_recipient(text, metadata, recipient_choices)
        if is_outbound
        else _suggest_sender(text, metadata, sender_choices)
    )
    category = _suggest_category(text, category_choices)

    competence_choice, competence_score = _match_choice(text, competence_choices)
    competence = None
    if competence_choice:
        competence = _suggestion(
            _choice_label(competence_choice),
            competence_score,
            "مطابقة جهة اختصاص من الهيكل التنظيمي",
            select_value=str(competence_choice.get("value") or ""),
        )

    workflow = _suggest_workflow(text, str(subject.get("value") or ""), workflow_choices)
    date_key = "sent_date" if is_outbound else "received_date"
    party_key = "recipient" if is_outbound else "sender"
    suggestions = {
        date_key: _suggestion(
            (sent_date if is_outbound else received_date) or date.today().isoformat(),
            1.0,
            (
                "تاريخ الإرسال الحالي؛ عدّله عند الحاجة"
                if is_outbound
                else "تاريخ الاستلام الحالي؛ عدّله عند الحاجة"
            ),
        ),
        "subject": subject,
        party_key: party,
        "category": category,
        "competence": competence,
        "workflow_template": workflow,
        "priority": _suggest_priority(text),
        "confidentiality": _suggest_confidentiality(text),
        "due_date": _suggest_due_date(text),
        "document_reference": _suggest_reference(text),
        "mail_scope": _suggestion(
            "EXTERNAL",
            0.75,
            "المرفق مسجل كبريد صادر خارجي" if is_outbound else "المرفق مسجل كبريد وارد خارجي",
        ),
        "body": _suggestion(text[:8000], 0.88, "النص المستخرج محلياً من المرفق") if text else None,
    }

    return {
        "filename": extracted["filename"],
        "format": extracted["format"],
        "extracted_characters": len(text),
        "preview": text[:1200],
        "warnings": extracted["warnings"],
        "ocr": extracted["ocr"],
        "suggestions": {key: value for key, value in suggestions.items() if value},
        "privacy": "LOCAL_ONLY",
    }


def analyze_workflow_attachment(
    payload: bytes,
    filename: str,
    *,
    request_type_choices: list[dict] | None = None,
    workflow_choices: list[dict] | None = None,
    max_text_chars: int = 20_000,
    max_pdf_pages: int = 40,
    ocr_config: OcrConfig | None = None,
) -> dict:
    """Extract one attachment into reviewable fields for a workflow request."""
    extracted = extract_attachment_text(
        payload,
        filename,
        max_chars=max_text_chars,
        max_pdf_pages=max_pdf_pages,
        ocr_config=ocr_config,
    )
    text = extracted["text"]
    metadata = extracted["metadata"]
    request_type_choices = list(request_type_choices or [])
    workflow_choices = list(workflow_choices or [])

    title = _suggest_subject(text, extracted["filename"], metadata)
    if title:
        title["value"] = str(title.get("value") or "")[:200]

    combined = f"{str((title or {}).get('value') or '')}\n{text}"
    request_type_choice, request_type_score = _match_choice(
        combined,
        request_type_choices,
    )
    request_type = None
    if request_type_choice:
        request_type = _suggestion(
            _choice_label(request_type_choice),
            request_type_score,
            "مطابقة نوع طلب مسجل مع محتوى المرفق",
            select_value=str(request_type_choice.get("value") or ""),
        )

    workflow = _suggest_workflow(
        text,
        str((title or {}).get("value") or ""),
        workflow_choices,
    )
    suggestions = {
        "title": title,
        "description": (
            _suggestion(text[:8000], 0.88, "النص المستخرج محلياً من المرفق")
            if text
            else None
        ),
        "request_type": request_type,
        "workflow_template": workflow,
    }

    return {
        "filename": extracted["filename"],
        "format": extracted["format"],
        "extracted_characters": len(text),
        "preview": text[:1200],
        "warnings": extracted["warnings"],
        "ocr": extracted["ocr"],
        "suggestions": {key: value for key, value in suggestions.items() if value},
        "privacy": "LOCAL_ONLY",
    }
