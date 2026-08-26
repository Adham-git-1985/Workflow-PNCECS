import tempfile
import unittest
import unicodedata
from datetime import datetime
from io import BytesIO
from pathlib import Path

import fitz
from PIL import Image
from docx import Document
from docx.oxml.ns import qn
from flask import Flask, g
from flask_login import LoginManager
from jinja2 import ChoiceLoader, DictLoader

from extensions import db
from models import (
    HRLookupItem,
    SystemSetting,
    TransportDriver,
    TransportMaintenance,
    TransportMaintenanceItem,
    TransportPermit,
    TransportTrip,
    TransportVehicle,
    User,
)
from portal import portal_bp
from services.transport_forms import _shape_transport_text


class TransportReadyFormsTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.temp_dir = tempfile.TemporaryDirectory()
        project_root = Path(__file__).resolve().parents[1]
        cls.app = Flask(
            __name__,
            instance_path=cls.temp_dir.name,
            template_folder=str(project_root / "templates"),
            static_folder=str(project_root / "static"),
        )
        cls.app.config.update(
            TESTING=True,
            SECRET_KEY="transport-forms-test",
            SQLALCHEMY_DATABASE_URI="sqlite:///:memory:",
            SQLALCHEMY_TRACK_MODIFICATIONS=False,
        )
        db.init_app(cls.app)
        login_manager = LoginManager()
        login_manager.init_app(cls.app)

        @login_manager.user_loader
        def load_user(user_id):
            try:
                return db.session.get(User, int(user_id))
            except (TypeError, ValueError):
                return None

        cls.app.register_blueprint(portal_bp)
        cls.app.jinja_loader = ChoiceLoader([
            DictLoader({"portal/layout.html": "{% block content %}{% endblock %}"}),
            cls.app.jinja_loader,
        ])
        cls.context = cls.app.app_context()
        cls.context.push()
        db.create_all()

    @classmethod
    def tearDownClass(cls):
        db.session.remove()
        cls.context.pop()
        cls.temp_dir.cleanup()

    def setUp(self):
        db.session.remove()
        db.drop_all()
        db.create_all()
        self.admin = User(
            email="transport-admin@example.test",
            name="مسؤول الحركة",
            password_hash="not-used-in-test",
            role="SUPER_ADMIN",
        )
        self.employee = User(
            email="transport-employee@example.test",
            name="موظف الحركة",
            password_hash="not-used-in-test",
            role="EMPLOYEE",
        )
        self.driver = TransportDriver(name="سائق الاختبار", status="ACTIVE")
        self.vehicle = TransportVehicle(
            plate_no="31-456-78",
            label="مركبة الاختبار",
            vehicle_type="مركبة حكومية",
            model="Corolla",
            year=2024,
            chassis_no="CHASSIS-123",
            engine_no="ENGINE-456",
            odometer_no="METER-789",
            assigned_to="دائرة الشؤون الإدارية",
            current_odometer=1500,
            status="ACTIVE",
        )
        db.session.add_all([self.admin, self.employee, self.driver, self.vehicle])
        db.session.flush()

        maintenance_type = HRLookupItem(
            category="TRANSPORT_MAINT_TYPE",
            code="OIL",
            name_ar="تغيير الزيت والفلتر",
            is_active=True,
        )
        db.session.add(maintenance_type)
        db.session.flush()
        self.maintenance = TransportMaintenance(
            vehicle_id=self.vehicle.id,
            invoice_day="2026-08-26",
            notes="فحص عام",
            created_by_id=self.admin.id,
        )
        db.session.add(self.maintenance)
        db.session.flush()
        db.session.add(TransportMaintenanceItem(
            maintenance_id=self.maintenance.id,
            maintenance_type_lookup_id=maintenance_type.id,
            note="حسب المواصفات",
        ))

        self.permit = TransportPermit(
            requester_user_id=self.admin.id,
            vehicle_id=self.vehicle.id,
            driver_id=self.driver.id,
            origin_text="مقر اللجنة",
            dest_text="وزارة التربية والتعليم",
            purpose="تسليم وثائق رسمية",
            depart_at=datetime(2026, 8, 26, 9, 0),
            return_at=datetime(2026, 8, 26, 12, 0),
            status="APPROVED",
            approval_stage="DONE",
        )
        db.session.add(self.permit)
        db.session.flush()
        db.session.add(TransportTrip(
            permit_id=self.permit.id,
            vehicle_id=self.vehicle.id,
            driver_id=self.driver.id,
            started_at=datetime(2026, 8, 26, 9, 5),
            ended_at=datetime(2026, 8, 26, 11, 55),
            start_odometer=1500,
            end_odometer=1530,
            order_no="1/2026",
            is_deleted=False,
        ))
        self.employee_permit = TransportPermit(
            requester_user_id=self.employee.id,
            vehicle_id=self.vehicle.id,
            driver_id=self.driver.id,
            origin_text="مقر اللجنة",
            dest_text="مقر المهمة",
            purpose="مهمة الموظف الخاصة",
            status="APPROVED",
            approval_stage="DONE",
        )
        db.session.add(self.employee_permit)
        db.session.commit()

    def _login(self, client, user_id: int):
        g.pop("_login_user", None)
        with client.session_transaction() as session:
            session.clear()
            session["_user_id"] = str(user_id)
            session["_fresh"] = True

    def _assert_single_page_pdf(self, response):
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.mimetype, "application/pdf")
        self.assertTrue(response.data.startswith(b"%PDF"))
        document = fitz.open(stream=response.data, filetype="pdf")
        try:
            self.assertEqual(document.page_count, 1)
            font_names = {
                str(font[3]).lower()
                for page in document
                for font in page.get_fonts(full=True)
            }
            self.assertTrue(any("sakkalmajalla" in name for name in font_names), font_names)
            self.assertTrue(any("sakkalmajalla-bold" in name for name in font_names), font_names)
        finally:
            document.close()

    def _assert_sakkal_docx(self, response):
        self.assertEqual(response.status_code, 200)
        self.assertEqual(
            response.mimetype,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertTrue(response.data.startswith(b"PK"))
        document = Document(BytesIO(response.data))
        runs = []
        containers = [document, document.sections[0].header, document.sections[0].footer]
        for container in containers:
            for paragraph in container.paragraphs:
                runs.extend(run for run in paragraph.runs if run.text.strip())
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            runs.extend(run for run in paragraph.runs if run.text.strip())
        self.assertTrue(runs)
        for run in runs:
            self.assertEqual(run.font.name, "Sakkal Majalla")
            self.assertIsNotNone(run.font.size)
            self.assertAlmostEqual(run.font.size.pt, 16, places=2)
        arabic_runs = [
            run
            for run in runs
            if any(unicodedata.bidirectional(char) in {"R", "AL"} for char in run.text)
        ]
        self.assertTrue(arabic_runs)
        for run in arabic_runs:
            self.assertTrue(run.font.rtl, run.text)

    def test_pdf_arabic_names_keep_standard_ligatures(self):
        shaped = _shape_transport_text("ادهم محمدوصفي عبدالله حنون")
        self.assertIn("\ufdf2", shaped)

    def test_word_movement_form_isolates_parentheses_and_draws_signature_lines(self):
        with self.app.test_client() as client:
            self._login(client, self.admin.id)
            response = client.get(
                f"/portal/transport/forms/permit/{self.permit.id}.docx"
            )

        self._assert_sakkal_docx(response)
        document = Document(BytesIO(response.data))
        body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("\u202a(1)\u202c", body_text)
        self.assertIn("\u202a(1/2026)\u202c", body_text)
        self.assertNotIn("_", body_text)

        signature_tables = document.tables[-3:]
        self.assertEqual(len(signature_tables), 3)
        for table in signature_tables:
            borders = table.cell(0, 0)._tc.tcPr.find(qn("w:tcBorders"))
            self.assertIsNotNone(borders)
            bottom = borders.find(qn("w:bottom"))
            self.assertIsNotNone(bottom)
            self.assertEqual(bottom.get(qn("w:val")), "single")

    def test_forms_page_and_three_generated_pdfs(self):
        with self.app.test_client() as client:
            self._login(client, self.admin.id)
            page = client.get("/portal/transport/forms")
            self.assertEqual(page.status_code, 200)
            body = page.get_data(as_text=True)
            self.assertIn("كتاب ترخيص مركبة حكومية", body)
            self.assertIn("طلب صيانة مركبة", body)
            self.assertIn("تصريح أمر حركة", body)
            self.assertIn("Word", body)

            self._assert_single_page_pdf(client.get(
                f"/portal/transport/forms/vehicle-license/{self.vehicle.id}.pdf"
            ))
            self._assert_single_page_pdf(client.get(
                f"/portal/transport/forms/maintenance/{self.maintenance.id}.pdf"
            ))
            self._assert_single_page_pdf(client.get(
                f"/portal/transport/forms/permit/{self.permit.id}.pdf"
            ))
            self._assert_sakkal_docx(client.get(
                f"/portal/transport/forms/vehicle-license/{self.vehicle.id}.docx"
            ))
            self._assert_sakkal_docx(client.get(
                f"/portal/transport/forms/maintenance/{self.maintenance.id}.docx"
            ))
            self._assert_sakkal_docx(client.get(
                f"/portal/transport/forms/permit/{self.permit.id}.docx"
            ))

    def test_uploaded_letterhead_is_saved_and_used(self):
        image_bytes = BytesIO()
        Image.new("RGB", (1200, 180), "white").save(image_bytes, format="PNG")
        image_bytes.seek(0)

        with self.app.test_client() as client:
            self._login(client, self.admin.id)
            response = client.post(
                "/portal/transport/forms/letterhead",
                data={"letterhead_file": (image_bytes, "official-header.png")},
                content_type="multipart/form-data",
            )
            self.assertEqual(response.status_code, 302)
            stored = SystemSetting.query.filter_by(key="TRANSPORT_FORMS_LETTERHEAD_PATH").first()
            self.assertIsNotNone(stored)
            stored_path = Path(self.temp_dir.name) / stored.value
            self.assertTrue(stored_path.is_file())

            preview = client.get("/portal/transport/forms/letterhead/view")
            self.assertEqual(preview.status_code, 200)
            self.assertEqual(preview.mimetype, "image/png")
            preview.close()

            generated = client.get(
                f"/portal/transport/forms/vehicle-license/{self.vehicle.id}.pdf"
            )
            self._assert_single_page_pdf(generated)
            generated_docx = client.get(
                f"/portal/transport/forms/vehicle-license/{self.vehicle.id}.docx"
            )
            self._assert_sakkal_docx(generated_docx)

    def test_pdf_letterhead_is_merged_as_full_page_background(self):
        letterhead = fitz.open()
        page = letterhead.new_page(width=595.276, height=841.89)
        page.insert_text((40, 40), "LETTERHEAD-BACKGROUND")
        payload = BytesIO(letterhead.tobytes())
        letterhead.close()

        with self.app.test_client() as client:
            self._login(client, self.admin.id)
            upload = client.post(
                "/portal/transport/forms/letterhead",
                data={"letterhead_file": (payload, "official-letterhead.pdf")},
                content_type="multipart/form-data",
            )
            self.assertEqual(upload.status_code, 302)

            generated = client.get(
                f"/portal/transport/forms/vehicle-license/{self.vehicle.id}.pdf"
            )
            self._assert_single_page_pdf(generated)
            document = fitz.open(stream=generated.data, filetype="pdf")
            try:
                self.assertIn("LETTERHEAD-BACKGROUND", document[0].get_text())
            finally:
                document.close()
            generated_docx = client.get(
                f"/portal/transport/forms/vehicle-license/{self.vehicle.id}.docx"
            )
            self._assert_sakkal_docx(generated_docx)

    def test_employee_sees_and_generates_only_own_permit_form(self):
        with self.app.test_client() as client:
            self._login(client, self.employee.id)
            page = client.get("/portal/transport/forms")
            self.assertEqual(page.status_code, 200)
            body = page.get_data(as_text=True)
            self.assertIn("مهمة الموظف الخاصة", body)
            self.assertNotIn("تسليم وثائق رسمية", body)

            own_pdf = client.get(
                f"/portal/transport/forms/permit/{self.employee_permit.id}.pdf"
            )
            self._assert_single_page_pdf(own_pdf)
            own_docx = client.get(
                f"/portal/transport/forms/permit/{self.employee_permit.id}.docx"
            )
            self._assert_sakkal_docx(own_docx)
            forbidden = client.get(
                f"/portal/transport/forms/vehicle-license/{self.vehicle.id}.pdf"
            )
            self.assertEqual(forbidden.status_code, 403)
            forbidden_docx = client.get(
                f"/portal/transport/forms/vehicle-license/{self.vehicle.id}.docx"
            )
            self.assertEqual(forbidden_docx.status_code, 403)


if __name__ == "__main__":
    unittest.main()
