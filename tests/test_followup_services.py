from datetime import date
from io import BytesIO
from types import SimpleNamespace
import unittest

from docx import Document
from docx.oxml.ns import qn

from services.followup_assistant import build_followup_analysis
from services.followup_docx import build_followup_docx


class FollowupServicesTests(unittest.TestCase):
    def _item(self, item_id, title, status="COMPLETED", included=True):
        return SimpleNamespace(
            id=item_id,
            title=title,
            description="تفاصيل البند",
            status=status,
            is_included=included,
            completed_on=date(2026, 9, 1),
            ai_suggestion=None,
            duplicate_hint=None,
        )

    def test_analysis_marks_duplicates_and_incomplete_items(self):
        analysis = build_followup_analysis([
            self._item(1, "إنجاز التقرير"),
            self._item(2, "إنجاز التقرير"),
            self._item(3, "متابعة الطلب", status="IN_PROGRESS"),
        ])

        self.assertEqual(analysis["duplicate_ids"], {1, 2})
        self.assertIn("يحتاج متابعة", analysis["summary"])
        self.assertIn("غير مكتملة", analysis["notes"])
        self.assertIn("تم إنجاز", analysis["suggestions"][1])

    def test_docx_export_contains_report_details(self):
        report = SimpleNamespace(
            employee=SimpleNamespace(full_name="موظف تجريبي"),
            manager=SimpleNamespace(full_name="مدير تجريبي"),
            period_start=date(2026, 9, 1),
            period_end=date(2026, 9, 5),
            status="REVIEWED",
            items=[self._item(1, "إنجاز التقرير")],
            employee_summary="ملخص الإنجاز",
            ai_summary=None,
            challenges="لا توجد",
            manager_request="دعم بسيط",
            manager_comment="تمت المراجعة",
            manager_rating="GOOD",
        )

        document = Document(BytesIO(build_followup_docx(report)))
        text = " ".join(paragraph.text for paragraph in document.paragraphs)
        table_text = " ".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        accomplishments_table = document.tables[-1]

        self.assertIn("تقرير إنجاز الموظف", text)
        self.assertIn("موظف تجريبي", table_text)
        self.assertIn("إنجاز التقرير", table_text)
        self.assertIn("تمت المراجعة", text)
        self.assertEqual(
            [cell.text for cell in accomplishments_table.rows[0].cells],
            ["المهمة", "التاريخ"],
        )
        self.assertEqual(
            [cell.text for cell in accomplishments_table.rows[1].cells],
            ["إنجاز التقرير", "2026-09-01"],
        )
        self.assertIsNotNone(accomplishments_table._tbl.tblPr.find(qn("w:bidiVisual")))

    def test_docx_export_keeps_an_empty_accomplishments_table(self):
        report = SimpleNamespace(
            employee=SimpleNamespace(full_name="موظف تجريبي"),
            manager=SimpleNamespace(full_name="مدير تجريبي"),
            period_start=date(2026, 9, 1),
            period_end=date(2026, 9, 5),
            status="DRAFT",
            items=[],
            employee_summary=None,
            ai_summary=None,
            challenges=None,
            manager_request=None,
            manager_comment=None,
            manager_rating=None,
        )

        document = Document(BytesIO(build_followup_docx(report)))
        accomplishments_table = document.tables[-1]

        self.assertEqual(
            [cell.text for cell in accomplishments_table.rows[0].cells],
            ["المهمة", "التاريخ"],
        )
        self.assertIn("لا توجد مهام منجزة", accomplishments_table.rows[1].cells[0].text)


if __name__ == "__main__":
    unittest.main()
