import tempfile
import unittest
from datetime import datetime
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from portal.routes import (
    _build_meeting_minutes_docx,
    _docx_add_heading,
    _docx_add_text,
    _docx_set_cell,
    _meeting_text_should_use_ltr,
    _meeting_minutes_filename,
)


class MeetingMinutesFormattingTests(unittest.TestCase):
    def test_direction_uses_first_strong_language_character(self):
        self.assertTrue(_meeting_text_should_use_ltr("English text ثم نص عربي"))
        self.assertFalse(_meeting_text_should_use_ltr("نص عربي then English text"))
        self.assertTrue(_meeting_text_should_use_ltr("2026 - English text"))
        self.assertFalse(_meeting_text_should_use_ltr("2026 - نص عربي"))

    def test_filename_is_readable_and_uses_meeting_datetime(self):
        meeting = SimpleNamespace(
            id=12,
            title="اجتماع مجلس الإدارة",
            start_at=datetime(2026, 8, 16, 14, 5),
        )

        self.assertEqual(
            _meeting_minutes_filename(meeting, "docx"),
            "محضر اجتماع - اجتماع مجلس الإدارة - بتاريخ 16-08-2026 - الساعة 14-05.docx",
        )

    def test_filename_removes_characters_forbidden_by_windows(self):
        meeting = SimpleNamespace(
            id=7,
            title='  مراجعة: الخطة / المرحلة الأولى؟  ',
            start_at=datetime(2026, 1, 2, 9, 3),
        )

        filename = _meeting_minutes_filename(meeting, ".pdf")

        self.assertEqual(
            filename,
            "محضر اجتماع - مراجعة الخطة المرحلة الأولى - بتاريخ 02-01-2026 - الساعة 09-03.pdf",
        )
        self.assertFalse(any(character in filename for character in '<>:"/\\|?*'))

    def test_docx_uses_meeting_title_as_first_visible_title(self):
        meeting = SimpleNamespace(
            id=12,
            title="اجتماع مجلس الإدارة",
            attachments=[],
        )
        context = {
            "id": 12,
            "title": meeting.title,
            "description": "",
            "location": "قاعة الاجتماعات",
            "status": "مجدول",
            "when": "2026-08-16 14:05",
            "organizer": "مدير الاجتماع",
            "created_at": "2026-08-15 10:00",
            "updated_at": "2026-08-15 10:00",
            "minutes_text": "نص المحضر",
            "decisions_text": "لا توجد قرارات.",
            "participants": [],
            "agenda_items": [],
            "tasks": [],
            "generated_at": "2026-08-16 15:00",
            "meeting_url": "http://example.test/portal/meetings/12",
        }

        with tempfile.TemporaryDirectory() as temp_dir:
            with (
                patch("portal.routes._meeting_minutes_context", return_value=context),
                patch("portal.routes._meeting_letterhead_path", return_value=None),
                patch("portal.routes._meeting_minutes_attachment_dir", return_value=Path(temp_dir)),
            ):
                docx_bytes = _build_meeting_minutes_docx(meeting)

        document = Document(BytesIO(docx_bytes))
        visible_paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]

        self.assertEqual(document.core_properties.title, meeting.title)
        self.assertEqual(visible_paragraphs[0].text, meeting.title)
        self.assertEqual(visible_paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertIsNotNone(visible_paragraphs[0]._p.get_or_add_pPr().find(qn("w:bidi")))
        self.assertEqual(visible_paragraphs[1].text, "بيانات الاجتماع")

    def test_docx_heading_hierarchy_and_content_follow_language_direction(self):
        document = Document()

        main_heading = _docx_add_heading(document, "العنوان الرئيسي", level=1)
        subheading = _docx_add_heading(document, "العنوان الفرعي", level=2)
        _docx_add_text(document, "محتوى عربي\nEnglish content")
        arabic_content = document.paragraphs[2]
        english_content = document.paragraphs[3]

        self.assertEqual(main_heading.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(subheading.alignment, WD_ALIGN_PARAGRAPH.RIGHT)
        self.assertEqual(arabic_content.alignment, WD_ALIGN_PARAGRAPH.RIGHT)
        self.assertEqual(english_content.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertIsNotNone(subheading._p.get_or_add_pPr().find(qn("w:bidi")))
        self.assertIsNotNone(arabic_content._p.get_or_add_pPr().find(qn("w:bidi")))
        self.assertIsNone(english_content._p.get_or_add_pPr().find(qn("w:bidi")))

    def test_docx_table_content_follows_language_direction(self):
        document = Document()
        table = document.add_table(rows=1, cols=2)

        _docx_set_cell(table.cell(0, 0), "محتوى عربي")
        _docx_set_cell(table.cell(0, 1), "English content")
        arabic_paragraph = table.cell(0, 0).paragraphs[0]
        english_paragraph = table.cell(0, 1).paragraphs[0]

        self.assertEqual(arabic_paragraph.alignment, WD_ALIGN_PARAGRAPH.RIGHT)
        self.assertEqual(english_paragraph.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertIsNotNone(arabic_paragraph._p.get_or_add_pPr().find(qn("w:bidi")))
        self.assertIsNone(english_paragraph._p.get_or_add_pPr().find(qn("w:bidi")))


if __name__ == "__main__":
    unittest.main()
