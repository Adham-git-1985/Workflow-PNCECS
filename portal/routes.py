from __future__ import annotations

from functools import wraps

from datetime import datetime, date, timedelta
import os
import re
import io
import uuid
import csv
import json
import mimetypes
import unicodedata
from pathlib import Path
from io import BytesIO

from flask import (
    render_template, request, redirect, url_for, flash, abort, current_app,
    send_file, send_from_directory, has_request_context
)

from utils.portal_search import apply_search_all_columns

from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.graphics.barcode import qr

from flask_login import login_required, current_user

# Used for safe filesystem storage of uploaded filenames. We keep the original
# name for display (Arabic-friendly) but use secure_filename as a fallback when
# needed.
from werkzeug.utils import secure_filename

from . import portal_bp
from extensions import db
from sqlalchemy import or_, and_, text, func
from sqlalchemy.sql import exists
from sqlalchemy.exc import OperationalError
from utils.perms import perm_required

# Backward-compatible alias: some routes historically used @require_permissions(...)
# while the canonical decorator in this project is utils.perms.perm_required.
require_permissions = perm_required

from utils.events import emit_event
from utils.org_dynamic import build_org_node_picker_tree
from models import (
    User,
    EmployeeFile,
    EmployeeAttachment,
    EmployeeEvaluationRun,
    HRLookupItem,
    EmployeeDependent,
    EmployeeQualification,
    EmployeeSecondment,
    Team,
    OrgUnitManager,
    OrgUnitAssignment,
    OrgNode,
    OrgNodeAssignment,
    Organization,
    Directorate,
    Department,
    Section,
    Division,
    Unit,
    WorkSchedule,
    WorkScheduleDay,
    EmployeeScheduleAssignment,
    WorkPolicy,
    WorkAssignment,
    HRPermissionType,
    HRLeaveType,
    HRPermissionRequest,
    HRLeaveRequest,
    HRLeaveAttachment,
    HRLeaveBalance,
    HRMonthlyPermissionAllowance,
    HRLeaveGradeEntitlement,
    AttendanceDailySummary,
    SystemSetting,
    AttendanceImportBatch,
    AttendanceEvent,
    InboundMail,
    OutboundMail,
    CorrAttachment,
    CorrCategory,
    CorrParty,
    CorrCounter,
    SavedFilter,
    AuditLog,
    StoreCategory,
    StoreFile,
    StoreFilePermission,
    PortalAccessRequest,
    PortalCircular,
    Delegation,
    UserPermission,
    Notification,
    Message,
    MessageRecipient,
    HRSSWorkflowDefinition,
    HRSSWorkflowStepDefinition,
    HRSSRequest,
    HRSSRequestApproval,
    HRSSRequestAttachment,
    HRDisciplinaryCase,
    HRDisciplinaryAction,
    HRDisciplinaryAttachment,
    HRDoc,
    HRDocVersion,
    HRPerformanceForm,
    HRPerformanceSection,
    HRPerformanceQuestion,
    HRPerformanceCycle,
    HRPerformanceAssignment,
    HRAttendanceSpecialCase,
    HRAttendanceClosing,
    HRAttendanceDeductionConfig,
    HRAttendanceDeductionRun,
    HRAttendanceDeductionItem,
    HROfficialMission,
    HROfficialOccasion,
    HRRoom,
    HRRoomBooking,

    # HR Training
    HRTrainingCourse,
    HRTrainingProgram,
    HRTrainingCondition,
    HRTrainingAttachment,
    HRTrainingEnrollment,


    # Inventory Store (Warehouse module)
    InvWarehouse,
    InvItemCategory,
    InvItem,
    InvIssueVoucher,
    InvIssueVoucherLine,
    InvIssueVoucherAttachment,
    InvInboundVoucher,
    InvInboundVoucherLine,
    InvInboundVoucherAttachment,
    InvScrapVoucher,
    InvScrapVoucherLine,
    InvScrapVoucherAttachment,
    InvReturnVoucher,
    InvReturnVoucherLine,
    InvReturnVoucherAttachment,
    InvRequest,

    InvSupplier,
    InvUnit,
    InvRoom,
    InvStocktakeVoucher,
    InvStocktakeVoucherLine,
    InvStocktakeVoucherAttachment,
    InvCustodyVoucher,
    InvCustodyVoucherLine,
    InvCustodyVoucherAttachment,
    InvRoomRequester,
    InvWarehousePermission,
)

# -------------------------
# Permissions (Portal)
# -------------------------
# Canonical (CRUD-like) keys (preferred)
PORTAL_READ = "PORTAL_READ"
PORTAL_ADMIN_READ = "PORTAL_ADMIN_READ"
PORTAL_ADMIN_PERMISSIONS_MANAGE = "PORTAL_ADMIN_PERMISSIONS_MANAGE"

PORTAL_CIRCULARS_MANAGE = "PORTAL_CIRCULARS_MANAGE"

HR_READ = "HR_READ"
HR_ATT_READ = "HR_ATTENDANCE_READ"
HR_ATT_CREATE = "HR_ATTENDANCE_CREATE"
HR_ATT_EXPORT = "HR_ATTENDANCE_EXPORT"


# Backward-compat variable names (some decorators use the full names)
HR_ATTENDANCE_READ = HR_ATT_READ
HR_ATTENDANCE_CREATE = HR_ATT_CREATE
HR_ATTENDANCE_EXPORT = HR_ATT_EXPORT

# HR Reports
HR_REPORTS_VIEW = "HR_REPORTS_VIEW"
HR_REPORTS_EXPORT = "HR_REPORTS_EXPORT"
HR_EMP_READ = "HR_EMPLOYEE_READ"
HR_EMP_MANAGE = "HR_EMPLOYEE_MANAGE"
HR_EMP_ATTACH = "HR_EMPLOYEE_ATTACHMENTS_MANAGE"
HR_ORG_READ = "HR_ORGSTRUCTURE_READ"
HR_ORG_MANAGE = "HR_ORGSTRUCTURE_MANAGE"

HR_MASTERDATA_MANAGE = "HR_MASTERDATA_MANAGE"

# Employee HR requests (self-service + approvals)
HR_REQUESTS_READ = "HR_REQUESTS_READ"
HR_REQUESTS_CREATE = "HR_REQUESTS_CREATE"
HR_REQUESTS_APPROVE = "HR_REQUESTS_APPROVE"
HR_REQUESTS_VIEW_ALL = "HR_REQUESTS_VIEW_ALL"

# HR Self-Service (Light Workflow)
HR_SS_READ = "HR_SS_READ"
HR_SS_CREATE = "HR_SS_CREATE"
HR_SS_APPROVE = "HR_SS_APPROVE"
HR_SS_WORKFLOWS_MANAGE = "HR_SS_WORKFLOWS_MANAGE"

# HR Discipline & Docs
HR_DISCIPLINE_READ = "HR_DISCIPLINE_READ"
HR_DISCIPLINE_MANAGE = "HR_DISCIPLINE_MANAGE"
HR_DOCS_READ = "HR_DOCS_READ"
HR_DOCS_MANAGE = "HR_DOCS_MANAGE"

# HR Payslip
HR_PAYSLIP_VIEW = "HR_PAYSLIP_VIEW"
# Backward-compat alias (some presets may refer to HR_DOCS_VIEW)
HR_DOCS_VIEW = HR_DOCS_READ

# Store (المستودع)
STORE_READ = "STORE_READ"
STORE_MANAGE = "STORE_MANAGE"
STORE_EXPORT = "STORE_EXPORT"

# HR Performance (360)
HR_PERF_READ = "HR_PERFORMANCE_READ"
HR_PERF_SUBMIT = "HR_PERFORMANCE_SUBMIT"
HR_PERF_MANAGE = "HR_PERFORMANCE_MANAGE"

# HR System Evaluation (KPI-based)
HR_SYSTEM_EVALUATION_VIEW = "HR_SYSTEM_EVALUATION_VIEW"
HR_PERF_EXPORT = "HR_PERFORMANCE_EXPORT"

# Reports / Compliance
PORTAL_REPORTS_READ = "PORTAL_REPORTS_READ"
PORTAL_REPORTS_EXPORT = "PORTAL_REPORTS_EXPORT"
PORTAL_AUDIT_READ = "PORTAL_AUDIT_READ"

PORTAL_INTEGRATIONS_MANAGE = "PORTAL_INTEGRATIONS_MANAGE"

CORR_READ = "CORR_READ"
CORR_CREATE = "CORR_CREATE"
CORR_UPDATE = "CORR_UPDATE"
CORR_DELETE = "CORR_DELETE"
CORR_EXPORT = "CORR_EXPORT"
CORR_LOOKUPS_MANAGE = "CORR_LOOKUPS_MANAGE"
CORR_MANAGE = "CORR_MANAGE"  # legacy-ish (kept)

# Legacy keys (still accepted via User.has_perm aliases)
PORTAL_VIEW = "PORTAL_VIEW"
HR_ATT_IMPORT = "HR_ATTENDANCE_IMPORT"
CORR_VIEW = "CORR_VIEW"
CORR_IN_CREATE = "CORR_IN_CREATE"
CORR_OUT_CREATE = "CORR_OUT_CREATE"



def _perm(p: str):
    return perm_required(p)



def _perm_any(*perms: str):
    """Require ANY of the given permission keys (OR)."""
    def deco(f):
        @wraps(f)
        def wrapper(*args, **kwargs):
            if not current_user.is_authenticated:
                abort(401)

            # IMPORTANT: Delegation must NOT reduce the privileges of a SUPER/ADMIN account.
            # Always evaluate the real logged-in user first.
            base_user = current_user
            try:
                role_raw = (getattr(base_user, 'role', '') or '').strip().upper().replace('-', '_').replace(' ', '_')
                role_raw = unicodedata.normalize('NFKC', role_raw)
                role_raw = ''.join(ch for ch in role_raw if (ch.isalnum() or ch == '_'))
                if role_raw.startswith('SUPER'):
                    return f(*args, **kwargs)
                if role_raw == 'ADMIN':
                    # keep going: ADMIN bypass is handled by User.has_perm for portal keys
                    pass
                if hasattr(base_user, 'has_role') and (base_user.has_role('SUPERADMIN') or base_user.has_role('SUPER_ADMIN')):
                    return f(*args, **kwargs)
            except Exception:
                pass

            # Delegation-aware effective user: if the current user is a delegatee,
            # permission checks should apply to the delegator.
            user = base_user
            try:
                from utils.permissions import get_effective_user  # local import (avoid cycles)
                if callable(get_effective_user):
                    user = get_effective_user() or base_user
            except Exception:
                user = base_user

            try:
                role_raw = (getattr(user, 'role', '') or '').strip().upper().replace('-', '_').replace(' ', '_')
                role_raw = unicodedata.normalize('NFKC', role_raw)
                role_raw = ''.join(ch for ch in role_raw if (ch.isalnum() or ch == '_'))
                if role_raw.startswith('SUPER'):
                    return f(*args, **kwargs)
                if hasattr(user, 'has_role') and (user.has_role('SUPERADMIN') or user.has_role('SUPER_ADMIN')):
                    return f(*args, **kwargs)
            except Exception:
                pass

            # Check BOTH the real logged-in user and (if enabled) the effective (delegator) user.
            # This prevents delegation from unexpectedly *reducing* access when the base user already has the permission.
            candidates = [base_user]
            try:
                if user is not None and getattr(user, 'id', None) is not None and getattr(base_user, 'id', None) != getattr(user, 'id', None):
                    candidates.append(user)
            except Exception:
                candidates = [base_user]

            def _has_any(pkey: str) -> bool:
                for cand in candidates:
                    hp = getattr(cand, 'has_perm', None)
                    if not callable(hp):
                        continue
                    try:
                        if hp(pkey):
                            return True
                    except Exception:
                        continue
                return False

            if not any(callable(getattr(c, 'has_perm', None)) for c in candidates):
                abort(403)

            for p in perms:
                if _has_any(p):
                    return f(*args, **kwargs)

            abort(403)
        return wrapper
    return deco

def _can_manage_corr() -> bool:
    try:
        return bool(current_user.has_perm(CORR_MANAGE))
    except Exception:
        return False


def _safe_count(qry):
    """Safely count query results (handles missing tables during early setups)."""
    try:
        return int(qry.count())
    except Exception:
        return 0


# ======================
# Portal Excel utilities (standard export/import)
# ======================

PORTAL_EXCEL_EXPORT_ENDPOINTS: set[str] = set()
PORTAL_EXCEL_IMPORT_ENDPOINTS: set[str] = set()
PORTAL_EXCEL_IMPORT_META: dict[str, dict] = {}
PORTAL_EXCEL_IMPORT_HANDLERS: dict[str, callable] = {}
PORTAL_EXCEL_IMPORT_HANDLERS: dict[str, callable] = {}


def _portal_register_excel_export(endpoint: str) -> None:
    PORTAL_EXCEL_EXPORT_ENDPOINTS.add(endpoint)


def _portal_register_excel_import(endpoint: str, *, meta: dict | None = None, handler=None) -> None:
    PORTAL_EXCEL_IMPORT_ENDPOINTS.add(endpoint)
    if meta:
        PORTAL_EXCEL_IMPORT_META[endpoint] = meta
    if handler:
        PORTAL_EXCEL_IMPORT_HANDLERS[endpoint] = handler


def _xlsx_from_dicts(data: list[dict], columns: list[tuple[str, str]], sheet_name: str) -> bytes:
    """Build XLSX bytes from list of dicts + (key,label) columns.

    This keeps older export blocks short and consistent.
    """
    from utils.excel import make_xlsx_bytes
    headers = [label for _, label in columns]
    rows = []
    for d in data:
        rows.append([d.get(key, "") for key, _ in columns])
    return make_xlsx_bytes(sheet_name, headers, rows)


def _normalize_header(s: str) -> str:
    s = (s or "").strip().lower()
    for ch in ['-', '/', '\\']:
        s = s.replace(ch, ' ')
    return ' '.join(s.split())


def _read_xlsx_dicts(file_storage, *, header_map: dict[str, str] | None = None, sheet_index: int = 0) -> list[dict]:
    """Read an XLSX file into a list of dicts using the first row as headers.

    - header_map maps normalized headers to desired keys (aliases).
    - Empty rows are skipped.
    """
    from openpyxl import load_workbook

    wb = load_workbook(file_storage, data_only=True)
    ws = wb.worksheets[sheet_index]
    it = ws.iter_rows(values_only=True)
    try:
        header_row = next(it)
    except StopIteration:
        return []

    headers: list[str] = []
    for h in header_row:
        hh = _normalize_header(str(h or ""))
        if header_map and hh in header_map:
            hh = header_map[hh]
        headers.append(hh)

    out: list[dict] = []
    for row in it:
        if not row:
            continue
        if all(v is None or str(v).strip() == "" for v in row):
            continue
        item: dict = {}
        for i, val in enumerate(row):
            if i >= len(headers):
                break
            key = headers[i]
            if not key:
                continue
            item[key] = val
        out.append(item)
    return out


def _safe_next_url_portal(next_url: str | None) -> str:
    """Allow only same-site relative redirect targets."""
    if not next_url:
        return url_for('portal.index')
    target = (next_url or '').strip()
    if not target.startswith('/'):
        return url_for('portal.index')
    if target.startswith('//') or '://' in target or '\\' in target:
        return url_for('portal.index')
    return target


def portal_excel_can_export(endpoint: str | None = None) -> bool:
    ep = endpoint or request.endpoint
    return bool(ep and ep in PORTAL_EXCEL_EXPORT_ENDPOINTS)


def portal_excel_can_import(endpoint: str | None = None) -> bool:
    ep = endpoint or request.endpoint
    return bool(ep and ep in PORTAL_EXCEL_IMPORT_ENDPOINTS)


def portal_excel_import_meta(endpoint: str | None = None) -> dict:
    ep = endpoint or request.endpoint
    return dict(PORTAL_EXCEL_IMPORT_META.get(ep or '', {}))


def _portal_flags():
    """Compute portal module access flags for the current user."""
    def has(key: str) -> bool:
        try:
            return bool(current_user.has_perm(key))
        except Exception:
            return False

    can_corr = has(CORR_READ)
    can_att = has(HR_ATT_READ)
    can_hr = any([
        has(HR_READ),
        has(HR_ATT_READ),
        has(HR_EMP_READ),
        has(HR_ORG_READ),
        has(HR_MASTERDATA_MANAGE),
        has(HR_REQUESTS_READ),
        has(HR_REQUESTS_CREATE),
        has(HR_REQUESTS_APPROVE),
        has(HR_SS_READ),
        has(HR_SS_CREATE),
        has(HR_SS_APPROVE),
        has(HR_SS_WORKFLOWS_MANAGE),
        has(HR_DOCS_READ),
        has(HR_DOCS_MANAGE),
        has(HR_REQUESTS_VIEW_ALL),
    ])
    can_store_manage = has('STORE_MANAGE')
    can_store = has('STORE_READ') or can_store_manage

    # Transport / Fleet
    can_transport = any([
        has('TRANSPORT_READ'),
        has('TRANSPORT_CREATE'),
        has('TRANSPORT_UPDATE'),
        has('TRANSPORT_DELETE'),
        has('TRANSPORT_APPROVE'),
        has('TRANSPORT_TRACKING_READ'),
        has('TRANSPORT_TRACKING_MANAGE'),
    ])
    can_transport_manage = any([has('TRANSPORT_CREATE'), has('TRANSPORT_UPDATE'), has('TRANSPORT_DELETE')])
    can_transport_track = any([has('TRANSPORT_TRACKING_READ'), has('TRANSPORT_TRACKING_MANAGE')])

    can_portal_admin = has(PORTAL_ADMIN_READ) or has(PORTAL_ADMIN_PERMISSIONS_MANAGE)
    can_approve = (has(HR_REQUESTS_APPROVE) or has(HR_REQUESTS_VIEW_ALL) or has(HR_SS_APPROVE) or has(HR_SS_WORKFLOWS_MANAGE))

    return {
        'can_corr': can_corr,
        'can_att': can_att,
        'can_hr': can_hr,
        'can_store': can_store,
        'can_store_manage': can_store_manage,
        'can_transport': can_transport,
        'can_transport_manage': can_transport_manage,
        'can_transport_track': can_transport_track,
        'can_portal_admin': can_portal_admin,
        'can_approve': can_approve,
        'can_corr_create': has(CORR_CREATE),
        'can_hr_req_create': has(HR_REQUESTS_CREATE),
    }



# -------------------------
# Helpers (Portal)
# -------------------------
ALLOWED_CORR_EXTS = {
    ".pdf", ".png", ".jpg", ".jpeg",
    ".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx",
    ".txt", ".csv",
    ".zip", ".rar",
}

# Store repository allowed extensions (keep same safe list as correspondence for now)
ALLOWED_STORE_EXTS = set(ALLOWED_CORR_EXTS)

def _clean_suffix(filename: str) -> str:
    # Normalize suffix like '.pdf' and strip bidi/zero-width chars.
    s = filename or ""
    try:
        s = str(s).split("/")[-1].split("\\")[-1]
    except Exception:
        s = str(filename or "")

    suf = (Path(s).suffix or "").strip().lower()

    # remove unicode format characters (e.g., RTL marks) and whitespace
    try:
        suf = "".join(ch for ch in suf if unicodedata.category(ch) != "Cf")
    except Exception:
        pass

    suf = suf.strip()

    # keep only dot + alnum
    suf = "".join(ch for ch in suf if (ch == "." or ch.isalnum()))

    if suf and not suf.startswith("."):
        suf = "." + suf
    return suf


def _clean_ext(filename: str) -> str:
    suf = _clean_suffix(filename)
    return suf[1:] if suf.startswith(".") else suf


# Avoid doing PRAGMA checks on every request.
_STORE_SCHEMA_CHECKED = False


def _store_storage_dir() -> str:
    base = os.path.join(current_app.instance_path, "uploads", "store")
    os.makedirs(base, exist_ok=True)
    return base


def _allowed_store_file(filename: str) -> bool:
    ext = _clean_suffix(filename)
    return ext in ALLOWED_STORE_EXTS


def _ensure_store_seed():
    """Seed a few default categories if none exist."""
    try:
        if StoreCategory.query.count() == 0:
            defaults = [
                "عام",
                "نماذج",
                "تعليمات",
                "تقارير",
                "موارد بشرية",
            ]
            for n in defaults:
                if not StoreCategory.query.filter_by(name=n).first():
                    db.session.add(StoreCategory(name=n, is_active=True))
            db.session.commit()
    except Exception:
        db.session.rollback()

    _STORE_SCHEMA_CHECKED = True


def _sqlite_table_columns(table_name: str) -> set[str]:
    """Return current SQLite table columns (best-effort)."""
    try:
        rows = db.session.execute(text(f"PRAGMA table_info({table_name})")).fetchall()
        return {r[1] for r in rows}
    except Exception:
        return set()


def _ensure_store_schema():
    """Best-effort schema upgrades for SQLite without Alembic.

    - Create missing tables.
    - Add missing columns required by newer portal store features.
    """
    global _STORE_SCHEMA_CHECKED
    if _STORE_SCHEMA_CHECKED:
        return

    # 1) Ensure tables exist
    try:
        db.create_all()
    except Exception:
        pass

    # 2) Add new columns (SQLite supports ADD COLUMN)
    try:
        cols = _sqlite_table_columns("store_category")
        if cols and "parent_id" not in cols:
            db.session.execute(text("ALTER TABLE store_category ADD COLUMN parent_id INTEGER"))
            db.session.commit()
    except Exception:
        db.session.rollback()

    _STORE_SCHEMA_CHECKED = True

    # 3) Helpful indexes (no-op if exists)
    try:
        db.session.execute(text("CREATE INDEX IF NOT EXISTS ix_store_category_parent_id ON store_category(parent_id)"))
        db.session.commit()
    except Exception:
        db.session.rollback()


def _ensure_store_ready():
    """Best-effort: ensure store tables exist (create_all already runs at startup)."""
    try:
        # Touch tables; if missing, create_all then seed.
        _ = StoreCategory.query.limit(1).all()
    except OperationalError:
        try:
            db.create_all()
        except Exception:
            pass
    # Apply small schema upgrades when running without migrations
    _ensure_store_schema()
    _ensure_store_seed()


def _store_user_role_upper() -> str | None:
    try:
        r = (getattr(current_user, "role", None) or "").strip().upper()
        return r or None
    except Exception:
        return None


def _store_shared_query_for_user():
    """Query StoreFilePermission rows relevant to current user."""
    role = _store_user_role_upper()
    now = datetime.utcnow()
    return (
        StoreFilePermission.query
        .filter(or_(StoreFilePermission.user_id == current_user.id, StoreFilePermission.role == role))
        .filter(or_(StoreFilePermission.expires_at.is_(None), StoreFilePermission.expires_at > now))
    )


def _store_can_access_file(row: StoreFile) -> bool:
    """Access rule:
    - STORE_MANAGE: all
    - STORE_READ: all (non-deleted)
    - Otherwise: only files shared to the user/role
    """
    try:
        if current_user.has_perm("STORE_MANAGE"):
            return True
        if current_user.has_perm("STORE_READ"):
            return True
    except Exception:
        pass

    try:
        perm = (
            _store_shared_query_for_user()
            .filter(StoreFilePermission.file_id == row.id)
            .first()
        )
        return perm is not None
    except Exception:
        return False


def _store_can_download_file(row: StoreFile) -> bool:
    try:
        if current_user.has_perm("STORE_MANAGE") or current_user.has_perm("STORE_READ"):
            return True
    except Exception:
        pass

    try:
        perm = (
            _store_shared_query_for_user()
            .filter(StoreFilePermission.file_id == row.id)
            .first()
        )
        return bool(perm and perm.can_download)
    except Exception:
        return False


def _store_category_options() -> list[tuple[int, str]]:
    """Return flat category list with indentation for selects."""
    try:
        cats = (
            StoreCategory.query
            .filter(StoreCategory.is_active.is_(True))
            .order_by(StoreCategory.name.asc())
            .all()
        )
    except Exception:
        return []

    by_parent: dict[int | None, list[StoreCategory]] = {}
    for c in cats:
        by_parent.setdefault(getattr(c, "parent_id", None), []).append(c)

    # sort each bucket
    for k in by_parent.keys():
        by_parent[k] = sorted(by_parent[k], key=lambda x: (0 if x.is_active else 1, (x.name or "")))

    out: list[tuple[int, str]] = []

    def walk(pid: int | None, prefix: str = ""):
        for c in by_parent.get(pid, []):
            label = f"{prefix}{c.name}"
            out.append((c.id, label))
            walk(c.id, prefix + "— ")

    walk(None, "")
    return out



@portal_bp.app_context_processor
def _inject_portal_context():
    """Make portal flags and lightweight badges available to all portal templates."""
    if not getattr(current_user, 'is_authenticated', False):
        return {}

    flags = _portal_flags()

    approvals_pending = 0
    if flags.get('can_approve'):
        try:
            if current_user.has_perm(HR_REQUESTS_VIEW_ALL):
                approvals_pending = (
                    _safe_count(HRLeaveRequest.query.filter(HRLeaveRequest.status == 'SUBMITTED'))
                    + _safe_count(HRPermissionRequest.query.filter(HRPermissionRequest.status == 'SUBMITTED'))
                )
            else:
                approvals_pending = (
                    _safe_count(HRLeaveRequest.query.filter(HRLeaveRequest.status == 'SUBMITTED', HRLeaveRequest.approver_user_id == current_user.id))
                    + _safe_count(HRPermissionRequest.query.filter(HRPermissionRequest.status == 'SUBMITTED', HRPermissionRequest.approver_user_id == current_user.id))
                )
        except Exception:
            approvals_pending = 0

    # Include HR Self-Service approvals pending
    try:
        if current_user.has_perm(HR_SS_APPROVE) or current_user.has_perm(HR_SS_WORKFLOWS_MANAGE) or current_user.has_perm(HR_REQUESTS_VIEW_ALL):
            role = (getattr(current_user, 'role', None) or '').strip()
            q_ss = HRSSRequestApproval.query.filter(HRSSRequestApproval.status == 'PENDING')
            if not (current_user.has_perm(HR_SS_WORKFLOWS_MANAGE) or current_user.has_perm(HR_REQUESTS_VIEW_ALL)):
                q_ss = q_ss.filter(or_(
                    HRSSRequestApproval.approver_user_id == current_user.id,
                    func.upper(HRSSRequestApproval.approver_role) == func.upper(role)
                ))
            approvals_pending += _safe_count(q_ss)
    except Exception:
        pass

    # Store: count files shared with me (user or role). This enables a "Shared files" shortcut
    # even when STORE_READ is not granted.
    store_shared_count = 0
    try:
        _ensure_store_schema()
        role = (getattr(current_user, 'role', None) or '').strip().upper() or None
        now = datetime.utcnow()
        q_shared = (
            StoreFilePermission.query
            .join(StoreFile, StoreFilePermission.file_id == StoreFile.id)
            .filter(StoreFile.is_deleted == False)  # noqa: E712
            .filter(or_(
                StoreFilePermission.user_id == current_user.id,
                StoreFilePermission.role == role,
            ))
            .filter(or_(StoreFilePermission.expires_at.is_(None), StoreFilePermission.expires_at > now))
        )
        store_shared_count = _safe_count(q_shared)
    except Exception:
        store_shared_count = 0

    
    # Portal access requests badge (pending)
    access_pending = 0
    my_pending = 0
    try:
        my_pending = _safe_count(PortalAccessRequest.query.filter(PortalAccessRequest.user_id == current_user.id, PortalAccessRequest.status == 'PENDING'))
    except Exception:
        my_pending = 0

    try:
        if current_user.has_perm(PORTAL_ADMIN_PERMISSIONS_MANAGE):
            access_pending = _safe_count(PortalAccessRequest.query.filter(PortalAccessRequest.status == 'PENDING'))
    except Exception:
        access_pending = 0

    # Portal notifications badge (reuse global Notification table)
    notif_unread = 0
    try:
        notif_unread = _safe_count(
            Notification.query
            .filter(Notification.user_id == current_user.id)
            .filter(Notification.is_mirror.is_(False))
            .filter(Notification.source == 'portal')
            .filter(Notification.is_read == False)  # noqa: E712
        )
    except Exception:
        notif_unread = 0

    return {
        'portal_flags': flags,
        'portal_excel_can_export': portal_excel_can_export,
        'portal_excel_can_import': portal_excel_can_import,
        'portal_excel_import_meta': portal_excel_import_meta,
        'portal_approvals_pending': approvals_pending,
        'portal_store_shared_count': store_shared_count,
        'portal_access_requests_pending': access_pending,
        'portal_my_access_requests_pending': my_pending,
        'portal_notif_unread': notif_unread,
    }


def _get_multi_arg(name: str, upper: bool = False) -> list[str]:
    """Read multi-valued query params.

    Supports:
      - repeated params: ?category=A&category=B
      - comma-separated: ?category=A,B
    """
    vals = request.args.getlist(name) or []
    if len(vals) == 1 and isinstance(vals[0], str) and "," in vals[0]:
        vals = [v.strip() for v in vals[0].split(",")]
    out = []
    for v in vals:
        if not v:
            continue
        s = (v or "").strip()
        if not s:
            continue
        out.append(s.upper() if upper else s)
    return out

def _corr_storage_dir() -> str:
    base = os.path.join(current_app.instance_path, "uploads", "correspondence")
    os.makedirs(base, exist_ok=True)
    return base

def _allowed_file(filename: str) -> bool:
    ext = _clean_suffix(filename)
    return ext in ALLOWED_CORR_EXTS

def _save_corr_files(files, inbound_id: int | None = None, outbound_id: int | None = None) -> int:
    """Save one or more uploaded files as CorrAttachment(s). Returns how many were saved."""
    saved = 0
    if not files:
        return 0
    storage = _corr_storage_dir()
    for f in files:
        if not f or not getattr(f, "filename", ""):
            continue
        if not _allowed_file(f.filename):
            continue
        original_name = f.filename
        ext = _clean_suffix(original_name)
        prefix = "IN" if inbound_id else "OUT"
        rid = inbound_id or outbound_id
        # uuid is imported as the standard library module; use uuid.uuid4() directly
        stored_name = f"{prefix}_{rid}_{uuid.uuid4().hex}{ext}"
        file_path = os.path.join(storage, stored_name)
        f.save(file_path)

        att = CorrAttachment(
            inbound_id=inbound_id,
            outbound_id=outbound_id,
            original_name=original_name,
            stored_name=stored_name,
            uploaded_by_id=current_user.id,
            uploaded_at=datetime.utcnow(),
        )
        db.session.add(att)
        saved += 1
    return saved

# --- Leave Attachments ---

def _leave_storage_dir() -> str:
    base = os.path.join(current_app.instance_path, "uploads", "leaves")
    os.makedirs(base, exist_ok=True)
    return base

def _leave_upload_dir(req_id: int) -> str:
    base = os.path.join(_leave_storage_dir(), str(req_id))
    os.makedirs(base, exist_ok=True)
    return base

def _save_leave_files(files, req_id: int, doc_type: str | None = None) -> int:
    """Save one or more uploaded files as HRLeaveAttachment(s). Returns how many were saved."""
    saved = 0
    if not files:
        return 0
    folder = _leave_upload_dir(req_id)
    for f in files:
        if not f or not getattr(f, "filename", ""):
            continue
        if not _allowed_file(f.filename):
            continue
        original_name = f.filename
        ext = _clean_suffix(original_name)
        stored_name = f"LEAVE_{req_id}_{uuid.uuid4().hex}{ext}"
        file_path = os.path.join(folder, stored_name)
        f.save(file_path)

        att = HRLeaveAttachment(
            request_id=req_id,
            doc_type=(doc_type or None),
            original_name=original_name,
            stored_name=stored_name,
            uploaded_by_id=current_user.id,
            uploaded_at=datetime.utcnow(),
        )
        db.session.add(att)
        saved += 1
    return saved

def _corr_year_from_date(date_s: str) -> int:
    """Extract year from YYYY-MM-DD; fallback to current year."""
    try:
        return int((date_s or "")[:4])
    except Exception:
        return datetime.utcnow().year


def _corr_next_ref(kind: str, date_s: str, category: str) -> str:
    """Generate next reference number using CorrCounter.

    Pattern: IN-YYYY-0001 / OUT-YYYY-0001
    (Category is used for partitioning the counter only.)
    """
    k = (kind or "IN").strip().upper()
    if k not in ("IN", "OUT"):
        k = "IN"
    year = _corr_year_from_date(date_s)
    cat = (category or "GENERAL").strip().upper() or "GENERAL"

    try:
        row = CorrCounter.query.filter_by(kind=k, year=year, category=cat).first()
    except Exception:
        # In case the counter table doesn't exist yet (old DB), attempt to create tables.
        try:
            db.create_all()
            row = CorrCounter.query.filter_by(kind=k, year=year, category=cat).first()
        except Exception:
            return f"{k}-{year}-{uuid.uuid4().hex[:6].upper()}"

    if not row:
        row = CorrCounter(kind=k, year=year, category=cat, last_no=0)
        db.session.add(row)
        db.session.flush()

    try:
        row.last_no = int(row.last_no or 0) + 1
    except Exception:
        row.last_no = 1

    n = int(row.last_no or 1)
    return f"{k}-{year}-{n:04d}"




# -------------------------
# Portal Home
# -------------------------
@portal_bp.route("/")
@login_required
@_perm(PORTAL_READ)
def index():
    flags = _portal_flags()

    # Pending access requests per service (shown inside disabled cards)
    access_reqs: dict[str, PortalAccessRequest] = {}
    try:
        pending_rows = (
            PortalAccessRequest.query
            .filter(PortalAccessRequest.user_id == current_user.id)
            .filter(PortalAccessRequest.status == "PENDING")
            .order_by(PortalAccessRequest.created_at.desc(), PortalAccessRequest.id.desc())
            .all()
        )
        for r in pending_rows or []:
            svc = (getattr(r, "service", "") or "").strip().lower()
            if svc and svc not in access_reqs:
                access_reqs[svc] = r
    except Exception:
        access_reqs = {}

    stats = {
        'corr_in_30': 0,
        'corr_out_30': 0,
        'att_days_month': 0,
        'my_requests_open': 0,
        'approvals_pending': 0,
        'transport_trips_month': 0,
        'transport_permits_pending': 0,
        'transport_vehicles': 0,
    }

    # --- Correspondence quick stats (last 30 days) ---
    if flags.get('can_corr'):
        try:
            since = datetime.utcnow() - timedelta(days=30)
            stats['corr_in_30'] = _safe_count(InboundMail.query.filter(InboundMail.created_at >= since))
            stats['corr_out_30'] = _safe_count(OutboundMail.query.filter(OutboundMail.created_at >= since))
        except Exception:
            pass

    # --- Attendance quick stat (days in current month for current user) ---
    if flags.get('can_att'):
        try:
            now = datetime.now()
            ym = f"{now.year:04d}-{now.month:02d}"
            stats['att_days_month'] = _safe_count(
                AttendanceDailySummary.query
                .filter(AttendanceDailySummary.user_id == current_user.id)
                .filter(AttendanceDailySummary.day.like(ym + '%'))
            )
        except Exception:
            pass

    # --- HR quick stats (my open requests + approvals pending) ---
    if flags.get('can_hr'):
        try:
            stats['my_requests_open'] = (
                _safe_count(HRLeaveRequest.query.filter(HRLeaveRequest.user_id == current_user.id, HRLeaveRequest.status == 'SUBMITTED'))
                + _safe_count(HRPermissionRequest.query.filter(HRPermissionRequest.user_id == current_user.id, HRPermissionRequest.status == 'SUBMITTED'))
            )
        except Exception:
            pass

    if flags.get('can_approve'):
        try:
            # use injected badge logic (but also pass explicitly for index template)
            if current_user.has_perm(HR_REQUESTS_VIEW_ALL):
                stats['approvals_pending'] = (
                    _safe_count(HRLeaveRequest.query.filter(HRLeaveRequest.status == 'SUBMITTED'))
                    + _safe_count(HRPermissionRequest.query.filter(HRPermissionRequest.status == 'SUBMITTED'))
                )
            else:
                stats['approvals_pending'] = (
                    _safe_count(HRLeaveRequest.query.filter(HRLeaveRequest.status == 'SUBMITTED', HRLeaveRequest.approver_user_id == current_user.id))
                    + _safe_count(HRPermissionRequest.query.filter(HRPermissionRequest.status == 'SUBMITTED', HRPermissionRequest.approver_user_id == current_user.id))
                )
        except Exception:
            pass


    # --- Transport quick stats (this month + pending permits) ---
    if flags.get('can_transport'):
        try:
            from models import TransportTrip, TransportPermit, TransportVehicle
            now = datetime.now()
            month_start = datetime(now.year, now.month, 1)
            if now.month == 12:
                next_month = datetime(now.year + 1, 1, 1)
            else:
                next_month = datetime(now.year, now.month + 1, 1)

            stats['transport_trips_month'] = _safe_count(
                TransportTrip.query.filter(TransportTrip.started_at >= month_start, TransportTrip.started_at < next_month)
            )
            stats['transport_permits_pending'] = _safe_count(
                TransportPermit.query.filter(TransportPermit.status == 'SUBMITTED')
            )
            stats['transport_vehicles'] = _safe_count(TransportVehicle.query)
        except Exception:
            pass

    

    # --- Circulars (last 5) ---
    last_circulars = []
    try:
        last_circulars = (PortalCircular.query
                          .order_by(PortalCircular.created_at.desc(), PortalCircular.id.desc())
                          .limit(5)
                          .all())
    except Exception:
        last_circulars = []

    return render_template("portal/index.html", flags=flags, stats=stats, access_reqs=access_reqs, last_circulars=last_circulars)




# -------------------------
# Portal entry (used from Masar sidebar for all users)
# -------------------------
@portal_bp.route("/entry")
@login_required
def portal_entry():
    """Public entry point from Masar.

    - Always visible in Masar sidebar.
    - If user has PORTAL access → redirect to Portal home.
    - Otherwise → show friendly "no access" page (instead of raw 403).
    """
    try:
        if current_user.has_perm(PORTAL_READ) or current_user.has_perm(PORTAL_VIEW):
            return redirect(url_for("portal.index"))
    except Exception:
        pass

    # Show last portal-access request (if any) so the user understands the status.
    last_req = None
    try:
        last_req = (
            PortalAccessRequest.query
            .filter(PortalAccessRequest.user_id == current_user.id)
            .filter(PortalAccessRequest.service == "portal")
            .order_by(PortalAccessRequest.created_at.desc(), PortalAccessRequest.id.desc())
            .first()
        )
    except Exception:
        last_req = None

    return render_template("portal/no_access.html", last_req=last_req)


# -------------------------
# Portal: Access Requests (طلب صلاحية)
# -------------------------

def _access_service_defs():
    """Service definitions: labels + recommended permission bundles."""
    return {
        "portal": {
            "label": "دخول البوابة",
            "options": [
                ("READ", "تفعيل الدخول للبوابة الإدارية", [PORTAL_READ]),
            ],
        },
        "corr": {
            "label": "الصادر/الوارد",
            "options": [
                ("READ", "عرض المراسلات", [CORR_READ]),
                ("EDIT", "تسجيل/تعديل مراسلات", [CORR_READ, CORR_CREATE, CORR_UPDATE]),
            ],
        },
        "attendance": {
            "label": "الدوام",
            "options": [
                ("READ", "عرض الدوام", [HR_ATT_READ]),
                ("MANAGE", "استيراد/إدارة الدوام", [HR_ATT_READ, HR_ATT_CREATE]),
            ],
        },
        "hr": {
            "label": "الموارد البشرية",
            "options": [
                ("EMPLOYEE", "خدمات الموظف (طلبات + وثائق HR)", [HR_READ, HR_REQUESTS_READ, HR_REQUESTS_CREATE, HR_SS_READ, HR_SS_CREATE, HR_DOCS_READ]),
                ("SYS_EVAL", "عرض التقييم النظامي (شهري/سنوي)", [HR_SYSTEM_EVALUATION_VIEW]),
                ("SELF_SERVICE", "الطلبات الداخلية (Self-Service)", [HR_SS_READ, HR_SS_CREATE]),
                ("DOCS", "وثائق HR", [HR_DOCS_READ]),
                ("APPROVE", "اعتماد الطلبات (مدير/HR)", [HR_READ, HR_REQUESTS_APPROVE, HR_SS_APPROVE]),
            ],
        },
        "store": {
            "label": "المستودع",
            "options": [
                ("READ", "عرض المستودع", ["STORE_READ"]),
                ("MANAGE", "رفع/إدارة المستودع", ["STORE_READ", "STORE_MANAGE"]),
            ],
        },
        "transport": {
            "label": "الحركة والنقل",
            "options": [
                ("READ", "عرض وحدة الحركة والنقل", ["TRANSPORT_READ"]),
                ("MANAGE", "إدارة الحركة (سيارات/سائقين/أذون/رحلات)", ["TRANSPORT_READ", "TRANSPORT_CREATE", "TRANSPORT_UPDATE", "TRANSPORT_DELETE"]),
                ("APPROVE", "اعتماد أذون الحركة", ["TRANSPORT_READ", "TRANSPORT_APPROVE"]),
                ("TRACKING", "التتبع (الخيار C)", ["TRANSPORT_READ", "TRANSPORT_TRACKING_READ", "TRANSPORT_TRACKING_MANAGE"]),
            ],
        },
    }


def _normalize_keys(keys: list[str]) -> list[str]:
    out = []
    seen = set()
    for k in keys or []:
        kk = (k or "").strip().upper()
        if not kk:
            continue
        if kk in seen:
            continue
        seen.add(kk)
        out.append(kk)
    return out


def _portal_admin_user_ids() -> list[int]:
    """Find users who should receive access requests notifications.

    We include:
    - Users with explicit UserPermission PORTAL_ADMIN_PERMISSIONS_MANAGE
    - Users whose role has RolePermission PORTAL_ADMIN_PERMISSIONS_MANAGE
    - Built-in ADMIN/SUPER_ADMIN roles (if used)
    """
    from sqlalchemy import func
    from models import RolePermission

    ids: set[int] = set()

    # Explicit user permissions
    try:
        rows = (
            UserPermission.query
            .filter(UserPermission.key == PORTAL_ADMIN_PERMISSIONS_MANAGE)
            .filter(UserPermission.is_allowed == True)  # noqa: E712
            .all()
        )
        for r in rows:
            if r.user_id:
                ids.add(int(r.user_id))
    except Exception:
        pass

    # Role-based permissions
    try:
        roles = (
            RolePermission.query
            .filter(RolePermission.permission == PORTAL_ADMIN_PERMISSIONS_MANAGE)
            .all()
        )
        role_codes = { (r.role or "").strip().upper() for r in roles if (r.role or "").strip() }
        if role_codes:
            urows = (
                db.session.query(User.id)
                .filter(func.upper(User.role).in_(list(role_codes)))
                .all()
            )
            for (uid,) in urows:
                ids.add(int(uid))
    except Exception:
        pass

    # Built-in roles (fallback)
    try:
        urows = (
            db.session.query(User.id)
            .filter(func.upper(User.role).in_(["SUPER_ADMIN", "SUPERADMIN"]))
            .all()
        )
        for (uid,) in urows:
            ids.add(int(uid))
    except Exception:
        pass

    return sorted(ids)


def _user_hierarchy(u: User) -> tuple[int | None, int | None, int | None]:
    """Return (org_id, directorate_id, department_id) for a user (best-effort)."""
    dept_id = getattr(u, "department_id", None)
    dir_id = getattr(u, "directorate_id", None)
    org_id = None

    try:
        if dept_id and not dir_id:
            dep = Department.query.get(int(dept_id))
            if dep and getattr(dep, "directorate_id", None):
                dir_id = dep.directorate_id
    except Exception:
        pass

    try:
        if dir_id:
            d = Directorate.query.get(int(dir_id))
            if d and getattr(d, "organization_id", None):
                org_id = d.organization_id
        elif dept_id:
            dep = Department.query.get(int(dept_id))
            if dep and getattr(dep, "directorate_id", None):
                d = Directorate.query.get(int(dep.directorate_id))
                if d and getattr(d, "organization_id", None):
                    org_id = d.organization_id
    except Exception:
        pass

    return org_id, dir_id, dept_id


def _choose_access_request_assignee(requester: User) -> int | None:
    """Pick a portal admin to handle this request (best-effort).

    Routing logic (simple, practical):
      - Prefer portal admins in the same department
      - Otherwise prefer same directorate
      - Otherwise prefer same organization
      - Otherwise fallback to the first portal admin
    """
    admin_ids = _portal_admin_user_ids()
    if not admin_ids:
        return None

    req_org, req_dir, req_dept = _user_hierarchy(requester)

    admins: list[User] = []
    try:
        admins = User.query.filter(User.id.in_(admin_ids)).all()
    except Exception:
        admins = []

    # Fallback to first id if we can't load admin users
    if not admins:
        return int(admin_ids[0])

    best_id = int(admin_ids[0])
    best_score = -1
    for a in admins:
        if not a or getattr(a, "id", None) is None:
            continue
        if getattr(a, "id", None) == getattr(requester, "id", None):
            continue
        a_org, a_dir, a_dept = _user_hierarchy(a)
        score = 0
        if req_dept and a_dept and int(a_dept) == int(req_dept):
            score += 3
        if req_dir and a_dir and int(a_dir) == int(req_dir):
            score += 2
        if req_org and a_org and int(a_org) == int(req_org):
            score += 1

        if score > best_score:
            best_score = score
            best_id = int(a.id)

    return best_id


def _access_option(service: str, opt_code: str):
    defs = _access_service_defs()
    svc = defs.get(service)
    if not svc:
        return None
    for code, label, keys in svc.get("options", []):
        if (code or "").strip().upper() == (opt_code or "").strip().upper():
            return (code, label, keys)
    return None


@portal_bp.route("/access/request", methods=["GET", "POST"])
@login_required
def access_request_new():
    defs = _access_service_defs()

    # Super Admin should never need to request portal permissions.
    try:
        if getattr(current_user, "email", "").strip().lower() == "superadmin@pncecs.org" or current_user.has_role("SUPERADMIN") or current_user.has_role("SUPER_ADMIN"):
            flash("حساب السوبر أدمن يمتلك الصلاحيات افتراضياً ولا يحتاج إلى طلب صلاحية.", "info")
            return redirect(url_for("portal.index"))
    except Exception:
        pass

    # Users who do NOT have portal access (PORTAL_READ/PORTAL_VIEW) should still
    # be able to request "Portal entry" permission from the no-access page.
    try:
        can_enter_portal = bool(current_user.has_perm(PORTAL_READ) or current_user.has_perm(PORTAL_VIEW))
    except Exception:
        can_enter_portal = False

    service = (request.args.get("service") or request.form.get("service") or "").strip().lower()
    if not can_enter_portal:
        # Force portal entry request for users without access
        service = "portal"
    else:
        if service not in defs:
            service = "corr"

    requested = (request.args.get("opt") or request.form.get("opt") or "").strip().upper()
    if not requested:
        requested = defs[service]["options"][0][0]

    opt = _access_option(service, requested)
    if not opt:
        requested = defs[service]["options"][0][0]
        opt = _access_option(service, requested)

    if request.method == "POST":
        note = (request.form.get("note") or "").strip() or None
        opt_code = (request.form.get("opt") or "").strip().upper()
        opt = _access_option(service, opt_code)
        if not opt:
            flash("الخيار غير صالح.", "danger")
            return redirect(url_for("portal.access_request_new", service=service))

        code, opt_label, keys = opt
        keys_norm = _normalize_keys(list(keys or []))

        if not keys_norm:
            flash("لا توجد صلاحيات مطلوبة لهذا الخيار.", "warning")
            return redirect(url_for("portal.access_request_new", service=service))

        keys_str = ",".join(keys_norm)

        # prevent duplicates
        try:
            existing = (
                PortalAccessRequest.query
                .filter_by(user_id=current_user.id, service=service, requested_keys=keys_str, status="PENDING")
                .first()
            )
            if existing:
                flash("يوجد طلب صلاحية مشابه قيد المراجعة بالفعل.", "info")
                return redirect(url_for("portal.my_access_requests"))
        except Exception:
            pass

        # Route/assign this request to a specific portal admin (best-effort)
        assigned_to_user_id = None
        try:
            assigned_to_user_id = _choose_access_request_assignee(current_user)
        except Exception:
            assigned_to_user_id = None

        req_row = PortalAccessRequest(
            user_id=current_user.id,
            service=service,
            requested_keys=keys_str,
            note=note,
            status="PENDING",
            created_at=datetime.utcnow(),
            assigned_to_user_id=assigned_to_user_id,
        )
        db.session.add(req_row)
        db.session.flush()

        # Audit
        try:
            db.session.add(AuditLog(
                user_id=current_user.id,
                action="PORTAL_ACCESS_REQUEST_CREATE",
                note=f"طلب صلاحية: {defs[service]['label']} ({opt_label})",
                target_type="PORTAL_ACCESS_REQUEST",
                target_id=req_row.id,
            ))
        except Exception:
            pass

        # Notify assignee (preferred) + fallback to all portal admins
        msg = f"طلب صلاحية بوابة: {(current_user.name or current_user.email)} — {defs[service]['label']} ({opt_label})"
        try:
            notified: set[int] = set()

            # Prefer a routed assignee (if found)
            if assigned_to_user_id and int(assigned_to_user_id) != int(current_user.id):
                db.session.add(Notification(
                    user_id=int(assigned_to_user_id),
                    message=msg,
                    type="PORTAL",
                    source='portal',
                    is_read=False,
                    created_at=datetime.utcnow(),
                ))
                notified.add(int(assigned_to_user_id))

            # Also notify all portal admins (except the requester and duplicates)
            for uid in _portal_admin_user_ids():
                if int(uid) == int(current_user.id):
                    continue
                if int(uid) in notified:
                    continue
                db.session.add(Notification(
                    user_id=int(uid),
                    message=msg,
                    type="PORTAL",
                    source='portal',
                    is_read=False,
                    created_at=datetime.utcnow(),
                ))
        except Exception:
            pass

        db.session.commit()
        flash("تم إرسال طلب الصلاحية للإدارة.", "success")
        return redirect(url_for("portal.my_access_requests"))

    return render_template(
        "portal/access_request_new.html",
        service=service,
        defs=defs,
        selected_opt=requested,
        can_enter_portal=can_enter_portal,
    )


@portal_bp.route("/access/requests")
@login_required
def my_access_requests():
    defs = _access_service_defs()
    try:
        can_enter_portal = bool(current_user.has_perm(PORTAL_READ) or current_user.has_perm(PORTAL_VIEW))
    except Exception:
        can_enter_portal = False

    qry = PortalAccessRequest.query.filter(PortalAccessRequest.user_id == current_user.id)

    # If user doesn't have portal access yet, only show the portal-entry requests.
    if not can_enter_portal:
        qry = qry.filter(PortalAccessRequest.service == "portal")

    rows = []
    try:
        rows = qry.order_by(PortalAccessRequest.created_at.desc(), PortalAccessRequest.id.desc()).all()
    except Exception:
        rows = []

    return render_template("portal/access_requests_mine.html", rows=rows, defs=defs, can_enter_portal=can_enter_portal)


# -------------------------


# -------------------------
# Portal: Circulars (التعميمات)
# -------------------------

@portal_bp.route("/circulars")
@login_required
@_perm(PORTAL_READ)
def circulars_list():
    """List circulars for portal users."""
    rows = []
    try:
        rows = (PortalCircular.query
                .order_by(PortalCircular.created_at.desc(), PortalCircular.id.desc())
                .limit(200)
                .all())
    except Exception:
        rows = []

    can_manage = False
    try:
        can_manage = bool(current_user.has_perm(PORTAL_CIRCULARS_MANAGE))
    except Exception:
        can_manage = False

    return render_template("portal/circulars_list.html", rows=rows, can_manage=can_manage)


@portal_bp.route("/circulars/<int:circular_id>")
@login_required
@_perm(PORTAL_READ)
def circular_view(circular_id: int):
    row = PortalCircular.query.get_or_404(circular_id)
    return render_template("portal/circular_view.html", row=row)


@portal_bp.route("/circulars/new", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_CIRCULARS_MANAGE)
def circular_new():
    """Create a new circular and notify all users (source='portal')."""

    if request.method == "POST":
        title = (request.form.get("title") or "").strip()
        body = (request.form.get("body") or "").strip()
        is_urgent = (request.form.get("is_urgent") or "").strip() in ("1", "on", "true", "True")

        if not title:
            flash("عنوان التعميم مطلوب.", "danger")
            return render_template("portal/circular_new.html", title=title, body=body, is_urgent=is_urgent)
        if not body:
            flash("نص التعميم مطلوب.", "danger")
            return render_template("portal/circular_new.html", title=title, body=body, is_urgent=is_urgent)

        try:
            circ = PortalCircular(
                title=title[:200],
                body=body,
                is_urgent=bool(is_urgent),
                created_by_user_id=getattr(current_user, "id", None),
                created_at=datetime.utcnow(),
            )
            db.session.add(circ)
            db.session.flush()  # get circ.id

            notif_type = "URGENT" if is_urgent else "INFO"
            msg = f"تعميم جديد: {title}".strip()
            if len(msg) > 250:
                msg = msg[:250]

            user_ids = []
            try:
                user_ids = [r[0] for r in db.session.query(User.id).all()]
            except Exception:
                user_ids = []

            now = datetime.utcnow()
            notifs = [
                Notification(
                    user_id=uid,
                    message=msg,
                    type=notif_type,
                    source="portal",
                    created_at=now,
                    is_read=False,
                    is_mirror=False,
                )
                for uid in user_ids
            ]

            if notifs:
                try:
                    db.session.bulk_save_objects(notifs)
                except Exception:
                    # Fallback to add_all for compatibility
                    db.session.add_all(notifs)

            db.session.commit()
            flash("تم إصدار التعميم وإرسال إشعار للجميع.", "success")
            return redirect(url_for("portal.circular_view", circular_id=circ.id))

        except Exception:
            try:
                db.session.rollback()
            except Exception:
                pass
            flash("حدث خطأ أثناء إصدار التعميم.", "danger")

    # GET
    return render_template("portal/circular_new.html", title="", body="", is_urgent=True)

# Portal: Notifications (simple inbox)
# -------------------------

@portal_bp.route("/notifications", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_READ)
def portal_notifications():
    """Show user notifications inside the portal UI.

    We reuse the global Notification table (already used across Masar).
    """
    if request.method == "POST":
        action = (request.form.get("action") or "").strip().upper()
        try:
            if action == "READ_ALL":
                Notification.query.filter(
                    Notification.user_id == current_user.id,
                    Notification.is_mirror.is_(False),
                    Notification.source == 'portal',
                    Notification.is_read == False,  # noqa: E712
                ).update({"is_read": True})
                db.session.commit()
                flash("تم تعليم جميع الإشعارات كمقروءة.", "success")
            elif action == "READ_ONE":
                nid = int(request.form.get("id") or 0)
                n = Notification.query.get(nid)
                if n and n.user_id == current_user.id and (getattr(n, 'source', None) == 'portal') and (not getattr(n, 'is_mirror', False)):
                    n.is_read = True
                    db.session.commit()
                return redirect(url_for("portal.portal_notifications"))
        except Exception:
            try:
                db.session.rollback()
            except Exception:
                pass

    unread_only = (request.args.get("unread") or "").strip() == "1"
    q = (Notification.query
         .filter(Notification.user_id == current_user.id)
         .filter(Notification.is_mirror.is_(False))
         .filter(Notification.source == 'portal'))
    if unread_only:
        q = q.filter(Notification.is_read == False)  # noqa: E712
    rows = q.order_by(Notification.created_at.desc(), Notification.id.desc()).limit(120).all()

    unread_count = 0
    try:
        unread_count = _safe_count(Notification.query
            .filter(Notification.user_id == current_user.id)
            .filter(Notification.is_mirror.is_(False))
            .filter(Notification.source == 'portal')
            .filter(Notification.is_read == False))
    except Exception:
        unread_count = 0

    return render_template(
        "portal/notifications.html",
        rows=rows,
        unread_only=unread_only,
        unread_count=unread_count,
    )


@portal_bp.route("/access/requests/<int:req_id>/cancel", methods=["POST"])
@login_required
def my_access_request_cancel(req_id: int):
    req_row = PortalAccessRequest.query.get_or_404(req_id)
    if req_row.user_id != current_user.id:
        abort(403)

    # Users without portal entry can only manage their portal-entry request.
    try:
        can_enter_portal = bool(current_user.has_perm(PORTAL_READ) or current_user.has_perm(PORTAL_VIEW))
    except Exception:
        can_enter_portal = False
    if not can_enter_portal and (req_row.service or "") != "portal":
        abort(403)

    if req_row.status != "PENDING":
        flash("لا يمكن إلغاء طلب غير قيد المراجعة.", "warning")
        return redirect(url_for("portal.my_access_requests"))

    req_row.status = "CANCELLED"
    req_row.decided_at = datetime.utcnow()
    req_row.decided_by_id = current_user.id
    req_row.decision_note = (request.form.get("decision_note") or "").strip() or "تم الإلغاء من قبل الموظف"

    try:
        db.session.add(AuditLog(
            user_id=current_user.id,
            action="PORTAL_ACCESS_REQUEST_CANCEL",
            note=f"إلغاء طلب صلاحية #{req_row.id}",
            target_type="PORTAL_ACCESS_REQUEST",
            target_id=req_row.id,
        ))
    except Exception:
        pass

    db.session.commit()
    flash("تم إلغاء الطلب.", "success")
    return redirect(url_for("portal.my_access_requests"))


@portal_bp.route("/admin/access-requests")
@login_required
@_perm(PORTAL_ADMIN_PERMISSIONS_MANAGE)
def admin_access_requests():
    defs = _access_service_defs()

    status = (request.args.get("status") or "PENDING").strip().upper()
    if status not in ("PENDING", "APPROVED", "REJECTED", "CANCELLED", "ALL"):
        status = "PENDING"

    service = (request.args.get("service") or "").strip().lower()
    q = (request.args.get("q") or "").strip()
    mine = (request.args.get("mine") or "").strip() == "1"

    qry = PortalAccessRequest.query.join(User, PortalAccessRequest.user_id == User.id)
    if status != "ALL":
        qry = qry.filter(PortalAccessRequest.status == status)
    if service and service in defs:
        qry = qry.filter(PortalAccessRequest.service == service)
    if q:
        # Search across all request columns + requester name/email.
        qry = apply_search_all_columns(
            qry,
            PortalAccessRequest,
            q,
            extra_columns=[User.email, User.name],
        )

    if mine:
        try:
            role = (getattr(current_user, "role", None) or "").strip().upper() or None
        except Exception:
            role = None
        conds = [PortalAccessRequest.assigned_to_user_id == current_user.id]
        if role:
            conds.append(PortalAccessRequest.assigned_role == role)
        qry = qry.filter(or_(*conds))

    rows = qry.order_by(PortalAccessRequest.created_at.desc(), PortalAccessRequest.id.desc()).limit(400).all()

    return render_template(
        "portal/admin/access_requests.html",
        rows=rows,
        defs=defs,
        status=status,
        service=service,
        q=q,
        mine=mine,
    )


def _grant_user_permissions(user_id: int, keys: list[str]) -> int:
    """Grant requested permission keys to a user (UserPermission). Returns count."""
    from portal.perm_defs import ALL_KEYS as ALLOWED

    allowed = { (k or "").strip().upper() for k in (ALLOWED or []) }
    keys_norm = [k for k in _normalize_keys(keys or []) if k in allowed]

    applied = 0
    for key in keys_norm:
        row = UserPermission.query.filter_by(user_id=user_id, key=key).first()
        if not row:
            db.session.add(UserPermission(user_id=user_id, key=key, is_allowed=True))
            applied += 1
        else:
            if not bool(row.is_allowed):
                row.is_allowed = True
                applied += 1
    return applied


@portal_bp.route("/admin/access-requests/<int:req_id>", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_ADMIN_PERMISSIONS_MANAGE)
def admin_access_request_view(req_id: int):
    defs = _access_service_defs()
    req_row = PortalAccessRequest.query.get_or_404(req_id)
    user = User.query.get(req_row.user_id)

    if request.method == "POST":
        action = (request.form.get("action") or "").strip().lower()
        decision_note = (request.form.get("decision_note") or "").strip() or None

        if req_row.status != "PENDING":
            flash("تمت معالجة هذا الطلب مسبقاً.", "info")
            return redirect(url_for("portal.admin_access_request_view", req_id=req_row.id))

        if action == "approve":
            # Apply permissions
            applied = 0
            try:
                applied = _grant_user_permissions(req_row.user_id, req_row.keys_list)
            except Exception:
                db.session.rollback()
                flash("فشل منح الصلاحيات. تأكد من الجداول.", "danger")
                return redirect(url_for("portal.admin_access_request_view", req_id=req_row.id))

            req_row.status = "APPROVED"
            req_row.decided_at = datetime.utcnow()
            req_row.decided_by_id = current_user.id
            req_row.decision_note = decision_note

            try:
                db.session.add(AuditLog(
                    user_id=current_user.id,
                    action="PORTAL_ACCESS_REQUEST_APPROVE",
                    note=f"موافقة على طلب صلاحية #{req_row.id} (منح {applied} صلاحية)",
                    target_type="PORTAL_ACCESS_REQUEST",
                    target_id=req_row.id,
                ))
            except Exception:
                pass

            # Notify requester
            try:
                svc_label = defs.get(req_row.service, {}).get("label", req_row.service)
                db.session.add(Notification(
                    user_id=req_row.user_id,
                    message=f"تمت الموافقة على طلب صلاحية: {svc_label}.",
                    type="SUCCESS",
                    source='portal',
                    is_read=False,
                    created_at=datetime.utcnow(),
                ))
            except Exception:
                pass

            db.session.commit()
            flash("تمت الموافقة ومنح الصلاحيات.", "success")
            return redirect(url_for("portal.admin_access_requests"))

        if action == "reject":
            req_row.status = "REJECTED"
            req_row.decided_at = datetime.utcnow()
            req_row.decided_by_id = current_user.id
            req_row.decision_note = decision_note

            try:
                db.session.add(AuditLog(
                    user_id=current_user.id,
                    action="PORTAL_ACCESS_REQUEST_REJECT",
                    note=f"رفض طلب صلاحية #{req_row.id}",
                    target_type="PORTAL_ACCESS_REQUEST",
                    target_id=req_row.id,
                ))
            except Exception:
                pass

            try:
                svc_label = defs.get(req_row.service, {}).get("label", req_row.service)
                msg = f"تم رفض طلب صلاحية: {svc_label}."
                if decision_note:
                    msg += f" السبب: {decision_note}"
                db.session.add(Notification(
                    user_id=req_row.user_id,
                    message=msg,
                    type="WARNING",
                    source='portal',
                    is_read=False,
                    created_at=datetime.utcnow(),
                ))
            except Exception:
                pass

            db.session.commit()
            flash("تم رفض الطلب.", "success")
            return redirect(url_for("portal.admin_access_requests"))

        flash("إجراء غير معروف.", "danger")
        return redirect(url_for("portal.admin_access_request_view", req_id=req_row.id))

    return render_template(
        "portal/admin/access_request_view.html",
        r=req_row,
        user=user,
        defs=defs,
    )


# ======================
# Portal: Excel Import (standard UI)
# ======================
@portal_bp.route("/excel/import", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_READ)
def portal_excel_import():
    key = (request.args.get("key") or "").strip()
    next_url = _safe_next_url_portal(request.args.get("next") or request.referrer)

    if not key:
        flash("لم يتم تحديد نوع الاستيراد.", "warning")
        return redirect(next_url)

    if not portal_excel_can_import(key) or key not in PORTAL_EXCEL_IMPORT_HANDLERS:
        flash("لا يوجد استيراد Excel متاح لهذه الصفحة.", "warning")
        return redirect(next_url)

    meta = portal_excel_import_meta(key)
    required_perm = meta.get("required_perm")
    if required_perm and not ((current_user.has_role("SUPERADMIN") or current_user.has_role("SUPER_ADMIN")) or current_user.has_perm(required_perm)):
        abort(403)

    modes = meta.get("modes") or {"default": {"label": "استيراد"}}
    default_mode = meta.get("default_mode") or (list(modes.keys())[0] if modes else "default")

    if request.method == "POST":
        mode = (request.form.get("mode") or default_mode).strip()
        file = request.files.get("file")
        if not file or not getattr(file, "filename", ""):
            flash("اختر ملف Excel أولاً.", "warning")
            return redirect(url_for("portal.portal_excel_import", key=key, next=next_url))

        handler = PORTAL_EXCEL_IMPORT_HANDLERS.get(key)
        try:
            result = handler(file, mode=mode, user=current_user)
            db.session.commit()
            inserted = result.get("inserted", 0)
            updated = result.get("updated", 0)
            skipped = result.get("skipped", 0)
            flash(f"تم الاستيراد بنجاح. (إضافة: {inserted}، تحديث: {updated}، تجاهل: {skipped})", "success")
        except Exception as e:
            db.session.rollback()
            flash(f"فشل الاستيراد: {e}", "danger")

        return redirect(next_url)

    return render_template("portal/excel_import.html", key=key, meta=meta, next=next_url)


@portal_bp.route("/excel/template")
@login_required
@_perm(PORTAL_READ)
def portal_excel_template():
    key = (request.args.get("key") or "").strip()
    mode = (request.args.get("mode") or "default").strip()

    if not key:
        abort(404)

    meta = portal_excel_import_meta(key)
    required_perm = meta.get("required_perm")
    if required_perm and not ((current_user.has_role("SUPERADMIN") or current_user.has_role("SUPER_ADMIN")) or current_user.has_perm(required_perm)):
        abort(403)

    modes = meta.get("modes") or {}
    m = modes.get(mode) or modes.get(meta.get("default_mode") or "") or {}
    cols = m.get("columns") or meta.get("columns") or []

    from utils.excel import make_xlsx_bytes, make_xlsx_bytes_multi
    filename = (m.get("template_filename") or meta.get("template_filename") or "import_template.xlsx")

    # Special case: permissions template supports multi-sheet file (role + user)
    if key == 'portal.portal_admin_permissions' and mode in ('both', 'all', 'multi'):
        tables = [
            ("RolePermissions", ["role", "permission"], []),
            ("UserPermissions", ["user_email", "permission", "is_allowed"], []),
        ]
        xbytes = make_xlsx_bytes_multi(tables)
    else:
        xbytes = make_xlsx_bytes("TEMPLATE", cols, [])
    return send_file(
        io.BytesIO(xbytes),
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=filename,
    )


# -------- Excel Importers (per page) --------

def _import_corr_categories(stream, *, mode: str, user):
    from models import CorrCategory
    header_map = {
        _normalize_header('code'): 'code',
        _normalize_header('كود'): 'code',
        _normalize_header('الرمز'): 'code',
        _normalize_header('name_ar'): 'name_ar',
        _normalize_header('الاسم ar'): 'name_ar',
        _normalize_header('الاسم (ar)'): 'name_ar',
        _normalize_header('الاسم عربي'): 'name_ar',
        _normalize_header('name_en'): 'name_en',
        _normalize_header('الاسم en'): 'name_en',
        _normalize_header('الاسم (en)'): 'name_en',
        _normalize_header('active'): 'is_active',
        _normalize_header('is_active'): 'is_active',
        _normalize_header('نشط'): 'is_active',
    }
    rows = _read_xlsx_dicts(stream, header_map=header_map)
    inserted = updated = 0
    for r in rows:
        code = str(r.get('code') or '').strip()
        if not code:
            continue
        obj = CorrCategory.query.filter_by(code=code).first()
        if obj is None:
            obj = CorrCategory(code=code)
            db.session.add(obj)
            inserted += 1
        else:
            updated += 1
        if r.get('name_ar') is not None:
            obj.name_ar = str(r.get('name_ar') or '').strip()
        if r.get('name_en') is not None:
            obj.name_en = str(r.get('name_en') or '').strip()
        if r.get('is_active') is not None:
            v = str(r.get('is_active') or '').strip().lower()
            obj.is_active = v in ('1','true','yes','y','نعم','فعال','نشط')
    return {'inserted': inserted, 'updated': updated}


def _import_corr_parties(stream, *, mode: str, user):
    from models import CorrParty
    header_map = {
        _normalize_header('kind'): 'kind',
        _normalize_header('type'): 'kind',
        _normalize_header('النوع'): 'kind',
        _normalize_header('name_ar'): 'name_ar',
        _normalize_header('الاسم (ar)'): 'name_ar',
        _normalize_header('الاسم عربي'): 'name_ar',
        _normalize_header('name_en'): 'name_en',
        _normalize_header('الاسم (en)'): 'name_en',
        _normalize_header('active'): 'is_active',
        _normalize_header('is_active'): 'is_active',
        _normalize_header('نشط'): 'is_active',
    }
    rows = _read_xlsx_dicts(stream, header_map=header_map)
    inserted = updated = 0
    for r in rows:
        kind = str(r.get('kind') or '').strip().upper()
        if kind not in ('SENDER','RECIPIENT','BOTH'):
            # fallback for Arabic values
            if kind in ('مرسل','صادر'):
                kind = 'SENDER'
            elif kind in ('مستلم','وارد'):
                kind = 'RECIPIENT'
            else:
                kind = 'BOTH'
        name_ar = str(r.get('name_ar') or '').strip()
        if not name_ar:
            continue
        obj = CorrParty.query.filter_by(kind=kind, name_ar=name_ar).first()
        if obj is None:
            obj = CorrParty(kind=kind, name_ar=name_ar)
            db.session.add(obj)
            inserted += 1
        else:
            updated += 1
        if r.get('name_en') is not None:
            obj.name_en = str(r.get('name_en') or '').strip()
        if r.get('is_active') is not None:
            v = str(r.get('is_active') or '').strip().lower()
            obj.is_active = v in ('1','true','yes','y','نعم','فعال','نشط')
    return {'inserted': inserted, 'updated': updated}


def _import_store_categories(stream, *, mode: str, user):
    from models import StoreCategory
    header_map = {
        _normalize_header('name'): 'name',
        _normalize_header('اسم'): 'name',
        _normalize_header('name_ar'): 'name',
        _normalize_header('parent'): 'parent',
        _normalize_header('parent_name'): 'parent',
        _normalize_header('الأب'): 'parent',
        _normalize_header('active'): 'is_active',
        _normalize_header('is_active'): 'is_active',
        _normalize_header('نشط'): 'is_active',
    }
    rows = _read_xlsx_dicts(stream, header_map=header_map)
    inserted = updated = 0
    for r in rows:
        name = str(r.get('name') or '').strip()
        if not name:
            continue
        obj = StoreCategory.query.filter_by(name=name).first()
        if obj is None:
            obj = StoreCategory(name=name)
            db.session.add(obj)
            inserted += 1
        else:
            updated += 1
        parent_name = str(r.get('parent') or '').strip()
        if parent_name:
            parent = StoreCategory.query.filter_by(name=parent_name).first()
            obj.parent = parent
        if r.get('is_active') is not None:
            v = str(r.get('is_active') or '').strip().lower()
            obj.is_active = v in ('1','true','yes','y','نعم','فعال','نشط')
    return {'inserted': inserted, 'updated': updated}


def _import_employee_timeclock(stream, *, mode: str, user):
    from models import User, EmployeeFile
    header_map = {
        _normalize_header('email'): 'email',
        _normalize_header('user_email'): 'email',
        _normalize_header('البريد'): 'email',
        _normalize_header('timeclock_code'): 'timeclock_code',
        _normalize_header('attendance_code'): 'timeclock_code',
        _normalize_header('كود الدوام'): 'timeclock_code',
        _normalize_header('code'): 'timeclock_code',
    }
    rows = _read_xlsx_dicts(stream, header_map=header_map)
    updated = 0
    for r in rows:
        email = str(r.get('email') or '').strip().lower()
        if not email:
            continue
        u = User.query.filter(db.func.lower(User.email) == email).first()
        if not u:
            continue

        code = str(r.get('timeclock_code') or '').strip()
        if not code:
            continue
        # keep it strict: 9 digits
        if (not code.isdigit()) or len(code) != 9:
            continue

        emp = EmployeeFile.query.filter_by(user_id=u.id).first()
        if not emp:
            emp = EmployeeFile(user_id=u.id, created_at=datetime.utcnow(), updated_at=datetime.utcnow())
            db.session.add(emp)

        emp.timeclock_code = code
        emp.updated_at = datetime.utcnow()
        try:
            if user and getattr(user, 'id', None):
                emp.updated_by_id = int(user.id)
        except Exception:
            pass

        updated += 1

    return {'updated': updated}



def _import_permissions(stream, *, mode: str, user):
    from models import RolePermission, UserPermission, User
    header_map = {
        _normalize_header('role'): 'role',
        _normalize_header('permission'): 'permission',
        _normalize_header('perm'): 'permission',
        _normalize_header('key'): 'permission',
        _normalize_header('user_email'): 'user_email',
        _normalize_header('email'): 'user_email',
        _normalize_header('allowed'): 'is_allowed',
        _normalize_header('is_allowed'): 'is_allowed',
        _normalize_header('yes'): 'is_allowed',
        _normalize_header('نعم'): 'is_allowed',
    }
    inserted = updated = 0

    def _apply_role_rows(rows: list[dict]) -> tuple[int, int]:
        ins = upd = 0
        for r in rows:
            role = str(r.get('role') or '').strip()
            perm = str(r.get('permission') or '').strip().upper()
            if not role or not perm:
                continue
            obj = RolePermission.query.filter_by(role=role, permission=perm).first()
            if obj is None:
                db.session.add(RolePermission(role=role, permission=perm))
                ins += 1
            else:
                upd += 1
        return ins, upd

    def _apply_user_rows(rows: list[dict]) -> tuple[int, int]:
        ins = upd = 0
        for r in rows:
            email = str(r.get('user_email') or '').strip().lower()
            perm = str(r.get('permission') or '').strip().upper()
            if not email or not perm:
                continue
            u = User.query.filter(db.func.lower(User.email) == email).first()
            if not u:
                continue

            obj = UserPermission.query.filter_by(user_id=u.id, key=perm).first()
            if obj is None:
                obj = UserPermission(user_id=u.id, key=perm)
                db.session.add(obj)
                ins += 1
            else:
                upd += 1

            if r.get('is_allowed') is not None:
                v = str(r.get('is_allowed') or '').strip().lower()
                obj.is_allowed = v in ('1', 'true', 'yes', 'y', 'نعم', 'allowed', 'ok')
        return ins, upd

    # ---- single-sheet modes ----
    if mode in ('role', 'user'):
        rows = _read_xlsx_dicts(stream, header_map=header_map)
        if mode == 'role':
            ins, upd = _apply_role_rows(rows)
        else:
            ins, upd = _apply_user_rows(rows)
        inserted += ins
        updated += upd
        return {'inserted': inserted, 'updated': updated}

    # ---- multi-sheet mode: both ----
    # Expected workbook contains 2 sheets:
    #  - Sheet 1: role permissions   (role, permission)
    #  - Sheet 2: user permissions   (user_email, permission, is_allowed)
    if mode in ('both', 'all', 'multi'):
        raw = stream.read()
        role_rows = _read_xlsx_dicts(io.BytesIO(raw), header_map=header_map, sheet_index=0)
        user_rows = _read_xlsx_dicts(io.BytesIO(raw), header_map=header_map, sheet_index=1)
        ins1, upd1 = _apply_role_rows(role_rows)
        ins2, upd2 = _apply_user_rows(user_rows)
        inserted += (ins1 + ins2)
        updated += (upd1 + upd2)
        return {'inserted': inserted, 'updated': updated}

    # fallback
    rows = _read_xlsx_dicts(stream, header_map=header_map)
    ins, upd = _apply_role_rows(rows)
    inserted += ins
    updated += upd
    return {'inserted': inserted, 'updated': updated}


# Bind importers + metadata
PORTAL_EXCEL_IMPORT_HANDLERS.update({
    'portal.portal_admin_corr_categories': _import_corr_categories,
    'portal.portal_admin_corr_parties': _import_corr_parties,
    'portal.store_categories': _import_store_categories,
    'portal.hr_employees': _import_employee_timeclock,
    'portal.portal_admin_permissions': _import_permissions,
})

_portal_register_excel_import('portal.portal_admin_corr_categories', meta={
    'title': 'تصنيفات المراسلات',
    'required_perm': CORR_LOOKUPS_MANAGE,
    'template_filename': 'corr_categories_template.xlsx',
    'modes': {
        'default': {
            'label': 'تصنيفات',
            'columns': ['code','name_ar','name_en','is_active'],
        }
    }
})

_portal_register_excel_import('portal.portal_admin_corr_parties', meta={
    'title': 'الجهات (مرسل/مستلم)',
    'required_perm': CORR_LOOKUPS_MANAGE,
    'template_filename': 'corr_parties_template.xlsx',
    'modes': {
        'default': {
            'label': 'جهات',
            'columns': ['kind','name_ar','name_en','is_active'],
        }
    }
})

_portal_register_excel_import('portal.store_categories', meta={
    'title': 'تصنيفات المستودع',
    'required_perm': "STORE_MANAGE",
    'template_filename': 'store_categories_template.xlsx',
    'modes': {
        'default': {
            'label': 'تصنيفات',
            'columns': ['name','parent','is_active'],
        }
    }
})

_portal_register_excel_import('portal.hr_employees', meta={
    'title': 'تحديث كود الدوام للموظفين',
    'required_perm': HR_EMP_MANAGE,
    'template_filename': 'employees_timeclock_template.xlsx',
    'modes': {
        'default': {
            'label': 'تحديث كود الدوام',
            'columns': ['email','timeclock_code'],
        }
    }
})

_portal_register_excel_import('portal.portal_admin_permissions', meta={
    'title': 'صلاحيات البوابة',
    'required_perm': PORTAL_ADMIN_PERMISSIONS_MANAGE,
    'template_filename': 'portal_permissions_template.xlsx',
    # default: one file with 2 sheets (role + user)
    'default_mode': 'both',
    'modes': {
        'both': {
            'label': 'ملف واحد (صلاحيات الدور + المستخدم)',
            'template_filename': 'portal_permissions_template.xlsx',
            'columns': [
                'Sheet1: role, permission',
                'Sheet2: user_email, permission, is_allowed',
            ],
        },
        'role': {
            'label': 'صلاحيات حسب الدور',
            'template_filename': 'role_permissions_template.xlsx',
            'columns': ['role','permission'],
        },
        'user': {
            'label': 'صلاحيات حسب المستخدم',
            'template_filename': 'user_permissions_template.xlsx',
            'columns': ['user_email','permission','is_allowed'],
        }
    }
})


# -------------------------
# HR Home (Portal)
# -------------------------
@portal_bp.route("/hr")
@login_required
@_perm(PORTAL_READ)
def hr_home():
    # Allow access if user has any HR-related permission.
    hr_keys = [
        HR_READ,
        HR_SS_READ, HR_SS_CREATE, HR_SS_APPROVE,
        HR_DOCS_READ, HR_DOCS_MANAGE,
        HR_PERF_READ, HR_PERF_SUBMIT, HR_PERF_MANAGE, HR_PERF_EXPORT,
        HR_DISCIPLINE_READ, HR_DISCIPLINE_MANAGE,
        HR_ATT_READ, HR_ATT_CREATE, HR_ATT_EXPORT,
        HR_EMP_READ, HR_EMP_MANAGE, HR_EMP_ATTACH,
        HR_ORG_READ, HR_ORG_MANAGE,
        HR_MASTERDATA_MANAGE,
    ]
    allowed = False
    try:
        allowed = any(current_user.has_perm(k) for k in hr_keys)
    except Exception:
        allowed = False
    if not allowed:
        abort(403)

    # تبسيط التجربة:
    # إذا كان المستخدم لا يملك أي صلاحيات إدارية في HR، نحوله إلى صفحة "ملفي" داخل HR.
    manage_keys = [
        HR_ATT_CREATE,
        HR_REQUESTS_APPROVE,
        HR_SS_APPROVE, HR_SS_WORKFLOWS_MANAGE,
        HR_DISCIPLINE_READ, HR_DISCIPLINE_MANAGE,
        HR_DOCS_MANAGE,
        HR_PERF_MANAGE, HR_PERF_EXPORT,
        HR_EMP_READ, HR_EMP_MANAGE, HR_EMP_ATTACH,
        HR_ORG_READ, HR_ORG_MANAGE,
        HR_MASTERDATA_MANAGE,
    ]

    try:
        has_manage = any(current_user.has_perm(k) for k in manage_keys)
    except Exception:
        has_manage = False

    # موظف عادي: HR self-service فقط
    if not has_manage:
        return redirect(url_for("portal.hr_me_home"))

    # Build sections in the same order as the reference system, and avoid a crowded UI.
    # Order (as requested): الموظفين → الدوام → الإجازات والمهام → التقارير → البرامج الفرعية → لوحة التحكم
    sections_order = [
        ("الموظفين", "bi-people"),
        ("الدوام", "bi-clock-history"),
        ("الإجازات والمهام", "bi-calendar2-check"),
        ("التقارير", "bi-file-earmark-text"),
        ("البرامج الفرعية", "bi-grid-3x3-gap"),
        ("لوحة التحكم", "bi-gear"),
    ]
    _sec_map = {t: [] for t, _ in sections_order}

    def add_item(perm_key, title, desc, icon, endpoint, section_title, *, when=True):
        """Append a link tile into a section if permission is granted."""
        try:
            if when and current_user.has_perm(perm_key):
                _sec_map.setdefault(section_title, []).append({
                    "title": title,
                    "desc": desc,
                    "icon": icon,
                    "url": url_for(endpoint),
                })
        except Exception:
            pass

    # --- Self-service shortcuts (kept inside the HR hub for convenience) ---
    add_item(HR_READ, "إجازاتي", "عرض أنواع الإجازات والمتابعة (تطوير لاحقاً).", "bi-calendar2-check", "portal.hr_my_leaves", "الإجازات والمهام")
    add_item(HR_READ, "مغادراتي", "عرض أنواع المغادرات والمتابعة (تطوير لاحقاً).", "bi-door-open", "portal.hr_my_permissions", "الإجازات والمهام")
    add_item(HR_READ, "قسيمة الراتب", "عرض قسائم الراتب بصيغة PDF (عرض فقط).", "bi-file-earmark-pdf", "portal.hr_my_payslips", "البرامج الفرعية")
    add_item(HR_SS_READ, "الطلبات الداخلية", "طلبات شهادة/تحديث بيانات/رفع مستندات.", "bi-ui-checks", "portal.hr_ss_home", "البرامج الفرعية")
    add_item(HR_DOCS_READ, "وثائق HR", "سياسات ونماذج بآخر نسخة معتمدة.", "bi-journal-text", "portal.hr_docs_home", "البرامج الفرعية")
    add_item(HR_PERF_READ, "الأداء والتقييم", "تقييم 360 (مدير/زملاء/ذاتي) حسب التكليف.", "bi-graph-up", "portal.hr_perf_home", "البرامج الفرعية")

    # --- Attendance (viewer vs importer) ---
    try:
        if current_user.has_perm(HR_ATT_READ):
            endpoint = "portal.hr_attendance_import" if current_user.has_perm(HR_ATT_CREATE) else "portal.hr_my_attendance"
            _sec_map["الدوام"].append({
                "title": "الدوام",
                "desc": "ملف ساعة الدوام، الأحداث، التقارير.",
                "icon": "bi-clock-history",
                "url": url_for(endpoint),
            })
    except Exception:
        pass

    # --- HR Admin / management items ---
    add_item(HR_EMP_READ, "ملفات الموظفين", "ربط كود الساعة ومتابعة ملفات الموظفين.", "bi-person-badge", "portal.hr_employees", "الموظفين")
    # قسائم الرواتب (للـ HR Admin)
    add_item(
        HR_EMP_ATTACH,
        "رفع قسائم الرواتب (شهرياً)",
        "رفع جميع قسائم الموظفين دفعة واحدة لشهر محدد كمسودات (لا يتم الإرسال تلقائياً).",
        "bi-upload",
        "portal.hr_payslips_bulk_upload",
        "الموظفين",
    )
    add_item(
        HR_EMP_ATTACH,
        "إرسال قسائم الرواتب",
        "إرسال/نشر قسائم الرواتب المسودة لشهر محدد عبر المراسلات.",
        "bi-send",
        "portal.hr_payslips_send",
        "الموظفين",
    )
    add_item(HR_ORG_READ, "الهيكل التنظيمي", "عرض المنظمات/الإدارات/الدوائر/الأقسام والفرق.", "bi-diagram-3", "portal.hr_org_structure", "لوحة التحكم")
    add_item(HR_ORG_MANAGE, "تعيين تبعية الموظفين (هيكلية موحدة)", "ربط الموظفين بعناصر الهيكلية الموحدة لاستخدامها في المسارات والموافقات.", "bi-person-badge", "portal.hr_org_node_assignments", "لوحة التحكم")
    add_item(HR_MASTERDATA_MANAGE, "إعدادات الدوام", "إعدادات الدوام/الإجازات/المغادرات والجداول.", "bi-gear", "portal.hr_masterdata_index", "لوحة التحكم")
    add_item(HR_REQUESTS_APPROVE, "الموافقات", "اعتماد/رفض طلبات الموظفين.", "bi-check2-square", "portal.hr_approvals", "الإجازات والمهام")
    # Reports
    try:
        if current_user.has_perm(HR_REQUESTS_VIEW_ALL) or current_user.has_perm(HR_REPORTS_VIEW):
            _sec_map["التقارير"].append({
                "title": "تقرير الإجازات",
                "desc": "تقرير طلبات الإجازات السنوي.",
                "icon": "bi-file-earmark-text",
                "url": url_for("portal.hr_leaves_report"),
            })
            _sec_map["التقارير"].append({
                "title": "تقرير شهري (إجازات + مغادرات)",
                "desc": "تقرير شهري يخصم المغادرات الزائدة ويعرض أرصدة الإجازات.",
                "icon": "bi-calendar2-week",
                "url": url_for("portal.hr_monthly_leave_report"),
            })
        if current_user.has_perm(HR_MASTERDATA_MANAGE) or current_user.has_perm(HR_REQUESTS_VIEW_ALL) or current_user.has_perm(HR_REPORTS_VIEW):
            _sec_map["التقارير"].append({
                "title": "أرصدة الإجازات",
                "desc": "عرض الأرصدة السنوية (المعتمدة/المستهلك/المتبقي).",
                "icon": "bi-wallet2",
                "url": url_for("portal.hr_leave_balances"),
            })
    except Exception:
        pass

    # Smart alerts (attendance/leave/pending requests)
    try:
        if current_user.has_perm(HR_ATT_READ) or current_user.has_perm(HR_REQUESTS_VIEW_ALL):
            _sec_map["لوحة التحكم"].append({
                "title": "تنبيهات ذكية",
                "desc": "تنبيهات التأخير ورصيد الإجازات والطلبات المعلقة.",
                "icon": "bi-bell",
                "url": url_for("portal.hr_alerts"),
            })
    except Exception:
        pass

    add_item(HR_SS_APPROVE, "موافقات الطلبات الداخلية", "اعتماد/رفض طلبات Self-Service.", "bi-check2-circle", "portal.hr_ss_approvals", "البرامج الفرعية")
    add_item(HR_SS_WORKFLOWS_MANAGE, "إعداد سير الطلبات الداخلية", "تعريف الخطوات ومن يعتمد كل نوع طلب.", "bi-diagram-2", "portal.portal_admin_hr_ss_workflows", "لوحة التحكم")
    add_item(HR_DISCIPLINE_READ, "الانضباط والشؤون القانونية", "مخالفات/إنذارات/تحقيقات مع مرفقات.", "bi-shield-exclamation", "portal.hr_discipline_home", "البرامج الفرعية")
    add_item(HR_DOCS_MANAGE, "وثائق HR (إدارة)", "رفع نسخ جديدة وإصدار معتمد.", "bi-folder2-open", "portal.hr_docs_admin", "البرامج الفرعية")
    add_item(HR_PERF_MANAGE, "الأداء والتقييم (إدارة)", "إدارة النماذج والدورات وتوليد التكليفات.", "bi-clipboard2-check", "portal.portal_admin_hr_perf_dashboard", "البرامج الفرعية")

    sections = []
    for title, icon in sections_order:
        items = _sec_map.get(title) or []
        if items:
            sections.append({"title": title, "icon": icon, "items": items})

    return render_template("portal/hr/index.html", sections=sections)


@portal_bp.route("/hr/reports")
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_reports_home():
    return render_template("portal/hr/reports_home.html")


# -------------------------
# HR Reports - Employees / Attendance
# -------------------------

def _parse_yyyy_mm_dd(v: str | None):
    try:
        s = (v or '').strip()
        if not s:
            return None
        return datetime.strptime(s, '%Y-%m-%d').date()
    except Exception:
        return None


def _add_years_safe(d: date, years: int):
    """Add years to a date while handling Feb 29."""
    try:
        return d.replace(year=d.year + years)
    except Exception:
        # fallback: move to last valid day of month
        try:
            # try Feb 28 for Feb 29
            if d.month == 2 and d.day == 29:
                return d.replace(year=d.year + years, day=28)
        except Exception:
            pass
        # generic fallback
        return date(d.year + years, d.month, min(d.day, 28))


def _hr_lookup_options(category: str):
    # HRLookupItem fields are name_ar/name_en (no generic "name").
    # Keep ordering stable and Arabic-friendly.
    return (
        HRLookupItem.query
        .filter_by(category=category, is_active=True)
        .order_by(
            HRLookupItem.sort_order.asc(),
            HRLookupItem.name_ar.asc(),
            HRLookupItem.id.asc(),
        )
        .all()
    )


def _filtered_user_ids(employee_id=None, work_location_id=None, appointment_type_id=None, organization_id=None, directorate_id=None, department_id=None, division_id=None):
    q = EmployeeFile.query
    if employee_id:
        q = q.filter(EmployeeFile.user_id == employee_id)
    if work_location_id:
        q = q.filter(EmployeeFile.work_location_lookup_id == work_location_id)
    if appointment_type_id:
        q = q.filter(EmployeeFile.appointment_type_lookup_id == appointment_type_id)

    if organization_id:
        q = q.filter(EmployeeFile.organization_id == organization_id)
    if directorate_id:
        q = q.filter(EmployeeFile.directorate_id == directorate_id)
    if department_id:
        q = q.filter(EmployeeFile.department_id == department_id)
    if division_id:
        q = q.filter(EmployeeFile.division_id == division_id)
    return [r.user_id for r in q.with_entities(EmployeeFile.user_id).all()]


def _cmp_ok(value: float, op: str, threshold: float):
    if op == 'eq':
        return value == threshold
    if op == 'gt':
        return value > threshold
    if op == 'gte':
        return value >= threshold
    if op == 'lt':
        return value < threshold
    if op == 'lte':
        return value <= threshold
    return True


def _export_xlsx(filename: str, headers: list[str], rows: list[list]):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = 'Report'
    ws.append(headers)
    for r in rows:
        ws.append(r)
    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)
    return send_file(bio, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', as_attachment=True, download_name=filename)


@portal_bp.route('/hr/reports/employees/promotions', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_promotions():
    from_date = _parse_yyyy_mm_dd(request.args.get('from') or request.args.get('from_date'))
    to_date = _parse_yyyy_mm_dd(request.args.get('to') or request.args.get('to_date'))
    work_location_id = (request.args.get('work_location_id') or '').strip()
    work_location_id = int(work_location_id) if work_location_id.isdigit() else None

    years = 4  # default policy
    q = EmployeeFile.query.join(User, User.id == EmployeeFile.user_id)
    if work_location_id:
        q = q.filter(EmployeeFile.work_location_lookup_id == work_location_id)
    q = q.order_by(func.coalesce(EmployeeFile.full_name_quad, User.name, User.email).asc())

    rows_view = []
    for ef in q.all():
        base = _parse_yyyy_mm_dd(ef.last_promotion_date) or _parse_yyyy_mm_dd(ef.hire_date)
        if not base:
            continue
        due = _add_years_safe(base, years)
        if from_date and due < from_date:
            continue
        if to_date and due > to_date:
            continue
        rows_view.append({
            'user': ef.user,
            'ef': ef,
            'base_date': base,
            'due_date': due,
        })

    work_locations = _hr_lookup_options('WORK_LOCATION')
    work_location_map = {x.id: x.name for x in (work_locations or [])}

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)
        headers = ['الرقم الوظيفي', 'الموظف', 'موقع العمل', 'تاريخ آخر ترقية/تعيين', 'تاريخ الاستحقاق']
        xrows = []
        for r in rows_view:
            loc = ''
            try:
                loc = next((x.name for x in work_locations if x.id == r['ef'].work_location_lookup_id), '')
            except Exception:
                loc = ''
            xrows.append([
                r['ef'].employee_no or '',
                (r['ef'].full_name_quad or r['user'].full_name or r['user'].name or r['user'].email),
                loc,
                r['base_date'].strftime('%Y-%m-%d'),
                r['due_date'].strftime('%Y-%m-%d'),
            ])
        return _export_xlsx('hr_promotions_due.xlsx', headers, xrows)

    return render_template(
        'portal/hr/reports_promotions.html',
        rows=rows_view,
        work_locations=work_locations,
        work_location_map=work_location_map,
        selected_work_location_id=work_location_id,
        from_date=from_date.strftime('%Y-%m-%d') if from_date else '',
        to_date=to_date.strftime('%Y-%m-%d') if to_date else '',
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
    )


@portal_bp.route('/hr/reports/employees/retirement', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_retirement():
    from_date = _parse_yyyy_mm_dd(request.args.get('from') or request.args.get('from_date'))
    to_date = _parse_yyyy_mm_dd(request.args.get('to') or request.args.get('to_date'))
    work_location_id = (request.args.get('work_location_id') or '').strip()
    work_location_id = int(work_location_id) if work_location_id.isdigit() else None

    retirement_age = 60
    q = EmployeeFile.query.join(User, User.id == EmployeeFile.user_id)
    if work_location_id:
        q = q.filter(EmployeeFile.work_location_lookup_id == work_location_id)
    q = q.order_by(func.coalesce(EmployeeFile.full_name_quad, User.name, User.email).asc())

    rows_view = []
    for ef in q.all():
        bd = _parse_yyyy_mm_dd(ef.birth_date)
        if not bd:
            continue
        ret_date = _add_years_safe(bd, retirement_age)
        if from_date and ret_date < from_date:
            continue
        if to_date and ret_date > to_date:
            continue
        rows_view.append({'user': ef.user, 'ef': ef, 'birth_date': bd, 'retirement_date': ret_date})

    work_locations = _hr_lookup_options('WORK_LOCATION')
    work_location_map = {x.id: x.name for x in (work_locations or [])}

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)
        headers = ['الرقم الوظيفي', 'الموظف', 'موقع العمل', 'تاريخ الميلاد', 'تاريخ التقاعد المتوقع']
        xrows = []
        for r in rows_view:
            loc = ''
            try:
                loc = next((x.name for x in work_locations if x.id == r['ef'].work_location_lookup_id), '')
            except Exception:
                loc = ''
            xrows.append([
                r['ef'].employee_no or '',
                (r['ef'].full_name_quad or r['user'].full_name or r['user'].name or r['user'].email),
                loc,
                r['birth_date'].strftime('%Y-%m-%d'),
                r['retirement_date'].strftime('%Y-%m-%d'),
            ])
        return _export_xlsx('hr_retirement.xlsx', headers, xrows)

    return render_template(
        'portal/hr/reports_retirement.html',
        rows=rows_view,
        work_locations=work_locations,
        work_location_map=work_location_map,
        selected_work_location_id=work_location_id,
        from_date=from_date.strftime('%Y-%m-%d') if from_date else '',
        to_date=to_date.strftime('%Y-%m-%d') if to_date else '',
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
    )


@portal_bp.route('/hr/reports/attendance/last-login', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_last_login():
    employee_id = (request.args.get('user_id') or '').strip()
    employee_id = int(employee_id) if employee_id.isdigit() else None
    work_location_id = (request.args.get('work_location_id') or '').strip()
    work_location_id = int(work_location_id) if work_location_id.isdigit() else None
    from_date = _parse_yyyy_mm_dd(request.args.get('from') or request.args.get('from_date'))
    to_date = _parse_yyyy_mm_dd(request.args.get('to') or request.args.get('to_date'))

    user_ids = _filtered_user_ids(employee_id=employee_id, work_location_id=work_location_id)
    q = User.query
    if user_ids:
        q = q.filter(User.id.in_(user_ids))
    else:
        # if filters produce no ids, show none
        q = q.filter(text('1=0'))

    q = q.filter(User.last_login_success_at.isnot(None)).order_by(User.last_login_success_at.desc())
    if from_date:
        q = q.filter(User.last_login_success_at >= datetime.combine(from_date, datetime.min.time()))
    if to_date:
        q = q.filter(User.last_login_success_at < datetime.combine(to_date + timedelta(days=1), datetime.min.time()))

    rows = q.limit(500).all()

    users = _list_hr_users()
    work_locations = _hr_lookup_options('WORK_LOCATION')

    loc_map = {x.id: x.name for x in (work_locations or [])}

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)
        headers = ['الموظف', 'الإيميل', 'موقع العمل', 'آخر دخول ناجح', 'IP', 'User-Agent']
        xrows = []
        for u in rows:
            loc = ''
            try:
                ef = u.employee_file
                if ef and ef.work_location_lookup_id:
                    loc = next((x.name for x in work_locations if x.id == ef.work_location_lookup_id), '')
            except Exception:
                loc = ''
            xrows.append([
                u.full_name or u.name or u.email,
                u.email,
                loc,
                u.last_login_success_at.strftime('%Y-%m-%d %H:%M') if u.last_login_success_at else '',
                u.last_login_success_ip or '',
                (u.last_login_success_ua or ''),
            ])
        return _export_xlsx('hr_last_login.xlsx', headers, xrows)

    return render_template(
        'portal/hr/reports_last_login.html',
        rows=rows,
        users=users,
        selected_user_id=employee_id,
        work_locations=work_locations,
        selected_work_location_id=work_location_id,
        from_date=from_date.strftime('%Y-%m-%d') if from_date else '',
        to_date=to_date.strftime('%Y-%m-%d') if to_date else '',
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
    )


@portal_bp.route('/hr/reports/attendance/attendance-permissions', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_attendance_permissions():
    employee_id = (request.args.get('user_id') or '').strip()
    employee_id = int(employee_id) if employee_id.isdigit() else None
    work_location_id = (request.args.get('work_location_id') or '').strip()
    work_location_id = int(work_location_id) if work_location_id.isdigit() else None
    from_date = _parse_yyyy_mm_dd(request.args.get('from') or request.args.get('from_date'))
    to_date = _parse_yyyy_mm_dd(request.args.get('to') or request.args.get('to_date'))

    appointment_type_id = (request.args.get('appointment_type_id') or '').strip()
    appointment_type_id = int(appointment_type_id) if appointment_type_id.isdigit() else None

    move_kind = (request.args.get('move_kind') or '').strip()  # attendance|private|official|''
    dur_op = (request.args.get('dur_op') or '').strip()  # eq|gt|gte|lt|lte|''
    hours_raw = (request.args.get('hours') or '').strip()
    try:
        hours_val = float(hours_raw) if hours_raw else None
    except Exception:
        hours_val = None

    users = _list_hr_users()
    work_locations = _hr_lookup_options('WORK_LOCATION')
    appointment_types = _hr_lookup_options('APPOINTMENT_TYPE')

    user_ids = _filtered_user_ids(employee_id=employee_id, work_location_id=work_location_id, appointment_type_id=appointment_type_id)
    if not user_ids:
        rows_view = []
    else:
        rows_view = []

        # Attendance rows
        if move_kind in ('', 'attendance'):
            aq = AttendanceDailySummary.query.filter(AttendanceDailySummary.user_id.in_(user_ids))
            if from_date:
                aq = aq.filter(AttendanceDailySummary.day >= from_date.strftime('%Y-%m-%d'))
            if to_date:
                aq = aq.filter(AttendanceDailySummary.day <= to_date.strftime('%Y-%m-%d'))
            for a in aq.order_by(AttendanceDailySummary.day.desc()).limit(2000).all():
                dur_h = None
                try:
                    if a.work_minutes is not None:
                        dur_h = float(a.work_minutes) / 60.0
                except Exception:
                    dur_h = None
                if hours_val is not None and dur_h is not None and dur_op:
                    if not _cmp_ok(dur_h, dur_op, hours_val):
                        continue
                u = None
                try:
                    u = User.query.get(a.user_id)
                except Exception:
                    u = None
                rows_view.append({
                    'date': a.day,
                    'user': u,
                    'kind': 'ATTENDANCE',
                    'label': 'الدوام',
                    'from': a.first_in,
                    'to': a.last_out,
                    'hours': dur_h,
                    'note': a.status or '',
                })

        # Permission rows
        if move_kind in ('', 'private', 'official'):
            pq = HRPermissionRequest.query.join(HRPermissionType, HRPermissionType.id == HRPermissionRequest.permission_type_id).filter(HRPermissionRequest.user_id.in_(user_ids))
            if from_date:
                pq = pq.filter(HRPermissionRequest.day >= from_date.strftime('%Y-%m-%d'))
            if to_date:
                pq = pq.filter(HRPermissionRequest.day <= to_date.strftime('%Y-%m-%d'))

            if move_kind == 'private':
                pq = pq.filter(or_(HRPermissionType.counts_as_work == False, HRPermissionType.counts_as_work.is_(None)))  # noqa: E712
            elif move_kind == 'official':
                pq = pq.filter(HRPermissionType.counts_as_work == True)  # noqa: E712

            for r in pq.order_by(HRPermissionRequest.day.desc()).limit(2000).all():
                dur_h = None
                try:
                    if r.hours is not None:
                        dur_h = float(r.hours)
                except Exception:
                    dur_h = None
                if dur_h is None:
                    try:
                        ft = _parse_hhmm(r.from_time)
                        tt = _parse_hhmm(r.to_time)
                        if ft and tt:
                            mins = (datetime.combine(date.today(), tt) - datetime.combine(date.today(), ft)).total_seconds() / 60.0
                            dur_h = max(0.0, mins / 60.0)
                    except Exception:
                        dur_h = None

                if hours_val is not None and dur_h is not None and dur_op:
                    if not _cmp_ok(dur_h, dur_op, hours_val):
                        continue

                u = None
                try:
                    u = User.query.get(r.user_id)
                except Exception:
                    u = None

                label = 'مغادرة رسمية' if (getattr(r.permission_type, 'counts_as_work', False)) else 'مغادرة خاصة'
                rows_view.append({
                    'date': r.day,
                    'user': u,
                    'kind': 'PERMISSION',
                    'label': label,
                    'from': r.from_time,
                    'to': r.to_time,
                    'hours': dur_h,
                    'note': (r.reason or ''),
                })

        # Sort combined
        def _sort_key(x):
            return (x.get('date') or '', (x.get('user').full_name if x.get('user') else '') or '')
        rows_view = sorted(rows_view, key=_sort_key, reverse=True)

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)
        headers = ['التاريخ', 'الموظف', 'نوع الحركة', 'من', 'إلى', 'المدة (ساعات)', 'ملاحظات']
        xrows = []
        for r in rows_view:
            u = r.get('user')
            xrows.append([
                r.get('date') or '',
                (u.full_name or u.name or u.email) if u else '',
                r.get('label') or '',
                r.get('from') or '',
                r.get('to') or '',
                round(float(r.get('hours') or 0), 2) if r.get('hours') is not None else '',
                r.get('note') or '',
            ])
        return _export_xlsx('hr_attendance_permissions.xlsx', headers, xrows)

    return render_template(
        'portal/hr/reports_attendance_permissions.html',
        rows=rows_view,
        users=users,
        selected_user_id=employee_id,
        work_locations=work_locations,
        selected_work_location_id=work_location_id,
        appointment_types=appointment_types,
        selected_appointment_type_id=appointment_type_id,
        from_date=from_date.strftime('%Y-%m-%d') if from_date else '',
        to_date=to_date.strftime('%Y-%m-%d') if to_date else '',
        move_kind=move_kind,
        dur_op=dur_op,
        hours_value=hours_raw,
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
    )


@portal_bp.route('/hr/reports/attendance/delay', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_delay():
    employee_id = (request.args.get('user_id') or '').strip()
    employee_id = int(employee_id) if employee_id.isdigit() else None
    work_location_id = (request.args.get('work_location_id') or '').strip()
    work_location_id = int(work_location_id) if work_location_id.isdigit() else None
    from_date = _parse_yyyy_mm_dd(request.args.get('from') or request.args.get('from_date'))
    to_date = _parse_yyyy_mm_dd(request.args.get('to') or request.args.get('to_date'))

    delay_type = (request.args.get('delay_type') or '').strip()  # late|early|''
    dur_op = (request.args.get('dur_op') or '').strip()
    hours_raw = (request.args.get('hours') or '').strip()
    try:
        hours_val = float(hours_raw) if hours_raw else None
    except Exception:
        hours_val = None

    users = _list_hr_users()
    work_locations = _hr_lookup_options('WORK_LOCATION')

    user_ids = _filtered_user_ids(employee_id=employee_id, work_location_id=work_location_id)
    if not user_ids:
        rows_view = []
    else:
        q = AttendanceDailySummary.query.filter(AttendanceDailySummary.user_id.in_(user_ids))
        if from_date:
            q = q.filter(AttendanceDailySummary.day >= from_date.strftime('%Y-%m-%d'))
        if to_date:
            q = q.filter(AttendanceDailySummary.day <= to_date.strftime('%Y-%m-%d'))
        q = q.order_by(AttendanceDailySummary.day.desc()).limit(2000)

        rows_view = []
        for a in q.all():
            late = int(a.late_minutes or 0)
            early = int(a.early_leave_minutes or 0)
            if delay_type == 'late':
                minutes = late
                if minutes <= 0:
                    continue
            elif delay_type == 'early':
                minutes = early
                if minutes <= 0:
                    continue
            else:
                minutes = late + early
                if minutes <= 0:
                    continue
            hours = float(minutes) / 60.0
            if hours_val is not None and dur_op:
                if not _cmp_ok(hours, dur_op, hours_val):
                    continue
            u = None
            try:
                u = User.query.get(a.user_id)
            except Exception:
                u = None
            rows_view.append({
                'date': a.day,
                'user': u,
                'late_minutes': late,
                'early_minutes': early,
                'hours': hours,
            })

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)
        headers = ['التاريخ', 'الموظف', 'تأخير صباحي (دقيقة)', 'خروج مبكر (دقيقة)', 'المدة (ساعات)']
        xrows = []
        for r in rows_view:
            u = r.get('user')
            xrows.append([
                r.get('date') or '',
                (u.full_name or u.name or u.email) if u else '',
                r.get('late_minutes') or 0,
                r.get('early_minutes') or 0,
                round(float(r.get('hours') or 0), 2),
            ])
        return _export_xlsx('hr_delay.xlsx', headers, xrows)

    return render_template(
        'portal/hr/reports_delay.html',
        rows=rows_view,
        users=users,
        selected_user_id=employee_id,
        work_locations=work_locations,
        selected_work_location_id=work_location_id,
        from_date=from_date.strftime('%Y-%m-%d') if from_date else '',
        to_date=to_date.strftime('%Y-%m-%d') if to_date else '',
        delay_type=delay_type,
        dur_op=dur_op,
        hours_value=hours_raw,
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
    )


@portal_bp.route('/hr/reports/attendance/employee', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_employee_attendance():
    employee_id = (request.args.get('user_id') or '').strip()
    employee_id = int(employee_id) if employee_id.isdigit() else None
    work_location_id = (request.args.get('work_location_id') or '').strip()
    work_location_id = int(work_location_id) if work_location_id.isdigit() else None
    from_date = _parse_yyyy_mm_dd(request.args.get('from') or request.args.get('from_date'))
    to_date = _parse_yyyy_mm_dd(request.args.get('to') or request.args.get('to_date'))

    users = _list_hr_users()
    work_locations = _hr_lookup_options('WORK_LOCATION')

    rows_view = []
    if employee_id:
        user_ids = _filtered_user_ids(employee_id=employee_id, work_location_id=work_location_id)
        if user_ids:
            q = AttendanceDailySummary.query.filter(AttendanceDailySummary.user_id == employee_id)
            if from_date:
                q = q.filter(AttendanceDailySummary.day >= from_date.strftime('%Y-%m-%d'))
            if to_date:
                q = q.filter(AttendanceDailySummary.day <= to_date.strftime('%Y-%m-%d'))
            q = q.order_by(AttendanceDailySummary.day.desc()).limit(2000)
            for a in q.all():
                rows_view.append(a)

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)
        headers = ['التاريخ', 'أول دخول', 'آخر خروج', 'ساعات العمل', 'تأخير (دقيقة)', 'خروج مبكر (دقيقة)', 'عمل إضافي (دقيقة)', 'الحالة']
        xrows = []
        for a in rows_view:
            hours = round(float(a.work_minutes or 0) / 60.0, 2) if a.work_minutes is not None else ''
            xrows.append([a.day, a.first_in or '', a.last_out or '', hours, int(a.late_minutes or 0), int(a.early_leave_minutes or 0), int(a.overtime_minutes or 0), a.status or ''])
        return _export_xlsx('hr_employee_attendance.xlsx', headers, xrows)

    selected_user = None
    if employee_id:
        try:
            selected_user = User.query.get(employee_id)
        except Exception:
            selected_user = None

    return render_template(
        'portal/hr/reports_employee_attendance.html',
        rows=rows_view,
        users=users,
        selected_user=selected_user,
        selected_user_id=employee_id,
        work_locations=work_locations,
        selected_work_location_id=work_location_id,
        from_date=from_date.strftime('%Y-%m-%d') if from_date else '',
        to_date=to_date.strftime('%Y-%m-%d') if to_date else '',
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
    )



@portal_bp.route('/hr/reports/attendance/employees', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_employees_attendance():
    """تقرير دوام موظفين: موقع العمل + الهيكل التنظيمي + نوع التعيين + نطاق التاريخ."""

    work_location_id = (request.args.get('work_location_id') or '').strip()
    work_location_id = int(work_location_id) if work_location_id.isdigit() else None

    organization_id = (request.args.get('organization_id') or '').strip()
    organization_id = int(organization_id) if organization_id.isdigit() else None

    directorate_id = (request.args.get('directorate_id') or '').strip()
    directorate_id = int(directorate_id) if directorate_id.isdigit() else None

    department_id = (request.args.get('department_id') or '').strip()
    department_id = int(department_id) if department_id.isdigit() else None

    division_id = (request.args.get('division_id') or '').strip()
    division_id = int(division_id) if division_id.isdigit() else None

    appointment_type_id = (request.args.get('appointment_type_id') or '').strip()
    appointment_type_id = int(appointment_type_id) if appointment_type_id.isdigit() else None

    from_date = _parse_yyyy_mm_dd(request.args.get('from') or request.args.get('from_date'))
    to_date = _parse_yyyy_mm_dd(request.args.get('to') or request.args.get('to_date'))

    # Default range: current month → today (avoids pulling a huge dataset accidentally)
    if (not from_date) and (not to_date):
        try:
            today = date.today()
            from_date = date(today.year, today.month, 1)
            to_date = today
        except Exception:
            pass

    work_locations = _hr_lookup_options('WORK_LOCATION')
    loc_map = {x.id: x.name for x in (work_locations or [])}
    appointment_types = _hr_lookup_options('APPOINTMENT_TYPE')

    orgs = Organization.query.filter(Organization.is_active.is_(True)).order_by(Organization.name_ar.asc()).all()
    dirs = Directorate.query.filter(Directorate.is_active.is_(True)).order_by(Directorate.name_ar.asc()).all()
    depts = Department.query.filter(Department.is_active.is_(True)).order_by(Department.name_ar.asc()).all()
    divs = Division.query.filter(Division.is_active.is_(True)).order_by(Division.name_ar.asc()).all()

    user_ids = _filtered_user_ids(
        work_location_id=work_location_id,
        appointment_type_id=appointment_type_id,
        organization_id=organization_id,
        directorate_id=directorate_id,
        department_id=department_id,
        division_id=division_id,
    )

    rows = []
    if user_ids:
        q = AttendanceDailySummary.query.filter(AttendanceDailySummary.user_id.in_(user_ids))
        if from_date:
            q = q.filter(AttendanceDailySummary.day >= from_date.strftime('%Y-%m-%d'))
        if to_date:
            q = q.filter(AttendanceDailySummary.day <= to_date.strftime('%Y-%m-%d'))
        # Most recent first
        recs = q.order_by(AttendanceDailySummary.day.desc(), AttendanceDailySummary.user_id.asc()).limit(5000).all()

        for a in recs:
            u = getattr(a, 'user', None)
            ef = getattr(u, 'employee_file', None) if u else None
            rows.append({
                'day': a.day,
                'user_id': a.user_id,
                'employee_no': getattr(ef, 'employee_no', '') if ef else '',
                'name': (getattr(ef, 'full_name_quad', None) or getattr(u, 'full_name', None) or getattr(u, 'name', None) or getattr(u, 'email', None) or '') if u else '',
                'work_location': loc_map.get(getattr(ef, 'work_location_lookup_id', None), '') if ef else '',
                'organization': getattr(getattr(ef, 'organization', None), 'name_ar', None) if ef else None,
                'directorate': getattr(getattr(ef, 'directorate', None), 'name_ar', None) if ef else None,
                'department': getattr(getattr(ef, 'department', None), 'name_ar', None) if ef else None,
                'division': getattr(getattr(ef, 'division', None), 'name_ar', None) if ef else None,
                'first_in': a.first_in,
                'last_out': a.last_out,
                'work_minutes': int(a.work_minutes or 0),
                'late_minutes': int(a.late_minutes or 0),
                'early_minutes': int(a.early_leave_minutes or 0),
                'overtime_minutes': int(a.overtime_minutes or 0),
                'status': a.status or '',
            })

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)

        def _fmt_dt(dt):
            try:
                return dt.strftime('%Y-%m-%d %H:%M:%S') if dt else ''
            except Exception:
                return ''

        headers = ['التاريخ', 'الرقم الوظيفي', 'الموظف', 'موقع العمل', 'الإدارة العامة', 'الدائرة', 'القسم', 'الشعبة', 'أول دخول', 'آخر خروج', 'ساعات الدوام', 'دقائق تأخير', 'دقائق خروج مبكر', 'دقائق إضافي', 'الحالة']
        xrows = []
        for r in rows:
            xrows.append([
                r.get('day') or '',
                r.get('employee_no') or '',
                r.get('name') or '',
                r.get('work_location') or '',
                r.get('organization') or '',
                r.get('directorate') or '',
                r.get('department') or '',
                r.get('division') or '',
                _fmt_dt(r.get('first_in')),
                _fmt_dt(r.get('last_out')),
                round(float(r.get('work_minutes') or 0) / 60.0, 2),
                r.get('late_minutes') or 0,
                r.get('early_minutes') or 0,
                r.get('overtime_minutes') or 0,
                r.get('status') or '',
            ])
        return _export_xlsx('hr_employees_attendance.xlsx', headers, xrows)

    return render_template(
        'portal/hr/reports_employees_attendance.html',
        rows=rows,
        work_locations=work_locations,
        selected_work_location_id=work_location_id,
        orgs=orgs,
        dirs=dirs,
        depts=depts,
        divs=divs,
        organization_id=organization_id,
        directorate_id=directorate_id,
        department_id=department_id,
        division_id=division_id,
        appointment_types=appointment_types,
        selected_appointment_type_id=appointment_type_id,
        from_date=from_date.strftime('%Y-%m-%d') if from_date else '',
        to_date=to_date.strftime('%Y-%m-%d') if to_date else '',
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
    )


@portal_bp.route('/hr/reports/attendance/summary', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_attendance_summary():
    """تقرير دوام تجميعي: (اختياري) الموظف + موقع العمل + الهيكل التنظيمي + نطاق التاريخ."""

    employee_id = (request.args.get('user_id') or '').strip()
    employee_id = int(employee_id) if employee_id.isdigit() else None

    work_location_id = (request.args.get('work_location_id') or '').strip()
    work_location_id = int(work_location_id) if work_location_id.isdigit() else None

    organization_id = (request.args.get('organization_id') or '').strip()
    organization_id = int(organization_id) if organization_id.isdigit() else None

    directorate_id = (request.args.get('directorate_id') or '').strip()
    directorate_id = int(directorate_id) if directorate_id.isdigit() else None

    department_id = (request.args.get('department_id') or '').strip()
    department_id = int(department_id) if department_id.isdigit() else None

    division_id = (request.args.get('division_id') or '').strip()
    division_id = int(division_id) if division_id.isdigit() else None

    from_date = _parse_yyyy_mm_dd(request.args.get('from') or request.args.get('from_date'))
    to_date = _parse_yyyy_mm_dd(request.args.get('to') or request.args.get('to_date'))

    if (not from_date) and (not to_date):
        try:
            today = date.today()
            from_date = date(today.year, today.month, 1)
            to_date = today
        except Exception:
            pass

    users = _list_hr_users()
    work_locations = _hr_lookup_options('WORK_LOCATION')
    loc_map = {x.id: x.name for x in (work_locations or [])}

    orgs = Organization.query.filter(Organization.is_active.is_(True)).order_by(Organization.name_ar.asc()).all()
    dirs = Directorate.query.filter(Directorate.is_active.is_(True)).order_by(Directorate.name_ar.asc()).all()
    depts = Department.query.filter(Department.is_active.is_(True)).order_by(Department.name_ar.asc()).all()
    divs = Division.query.filter(Division.is_active.is_(True)).order_by(Division.name_ar.asc()).all()

    user_ids = _filtered_user_ids(
        employee_id=employee_id,
        work_location_id=work_location_id,
        organization_id=organization_id,
        directorate_id=directorate_id,
        department_id=department_id,
        division_id=division_id,
    )

    rows = []
    if user_ids:
        q = AttendanceDailySummary.query.filter(AttendanceDailySummary.user_id.in_(user_ids))
        if from_date:
            q = q.filter(AttendanceDailySummary.day >= from_date.strftime('%Y-%m-%d'))
        if to_date:
            q = q.filter(AttendanceDailySummary.day <= to_date.strftime('%Y-%m-%d'))
        recs = q.all()

        # preload users
        ulist = User.query.filter(User.id.in_(user_ids)).all()
        umap = {u.id: u for u in ulist}

        agg = {}
        for a in recs:
            uid = a.user_id
            g = agg.setdefault(uid, {
                'days': 0,
                'absent': 0,
                'incomplete': 0,
                'work_minutes': 0,
                'late_minutes': 0,
                'early_minutes': 0,
                'overtime_minutes': 0,
            })
            g['days'] += 1
            st = (a.status or '').upper()
            if st == 'ABSENT' or (not a.first_in and not a.last_out):
                g['absent'] += 1
            elif st == 'INCOMPLETE':
                g['incomplete'] += 1
            g['work_minutes'] += int(a.work_minutes or 0)
            g['late_minutes'] += int(a.late_minutes or 0)
            g['early_minutes'] += int(a.early_leave_minutes or 0)
            g['overtime_minutes'] += int(a.overtime_minutes or 0)

        for uid, g in agg.items():
            u = umap.get(uid)
            ef = getattr(u, 'employee_file', None) if u else None
            rows.append({
                'user_id': uid,
                'employee_no': getattr(ef, 'employee_no', '') if ef else '',
                'name': (getattr(ef, 'full_name_quad', None) or getattr(u, 'full_name', None) or getattr(u, 'name', None) or getattr(u, 'email', None) or '') if u else '',
                'work_location': loc_map.get(getattr(ef, 'work_location_lookup_id', None), '') if ef else '',
                'organization': getattr(getattr(ef, 'organization', None), 'name_ar', None) if ef else None,
                'directorate': getattr(getattr(ef, 'directorate', None), 'name_ar', None) if ef else None,
                'department': getattr(getattr(ef, 'department', None), 'name_ar', None) if ef else None,
                'division': getattr(getattr(ef, 'division', None), 'name_ar', None) if ef else None,
                **g,
            })

        rows.sort(key=lambda r: (r.get('name') or ''), reverse=False)

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)
        headers = ['من تاريخ', 'إلى تاريخ', 'الرقم الوظيفي', 'الموظف', 'موقع العمل', 'الإدارة العامة', 'الدائرة', 'القسم', 'الشعبة', 'عدد الأيام', 'غياب', 'غير مكتمل', 'ساعات الدوام', 'دقائق تأخير', 'دقائق خروج مبكر', 'دقائق إضافي']
        xrows = []
        for r in rows:
            xrows.append([
                from_date.strftime('%Y-%m-%d') if from_date else '',
                to_date.strftime('%Y-%m-%d') if to_date else '',
                r.get('employee_no') or '',
                r.get('name') or '',
                r.get('work_location') or '',
                r.get('organization') or '',
                r.get('directorate') or '',
                r.get('department') or '',
                r.get('division') or '',
                r.get('days') or 0,
                r.get('absent') or 0,
                r.get('incomplete') or 0,
                round(float(r.get('work_minutes') or 0) / 60.0, 2),
                r.get('late_minutes') or 0,
                r.get('early_minutes') or 0,
                r.get('overtime_minutes') or 0,
            ])
        return _export_xlsx('hr_attendance_summary.xlsx', headers, xrows)

    return render_template(
        'portal/hr/reports_attendance_summary.html',
        rows=rows,
        users=users,
        selected_user_id=employee_id,
        work_locations=work_locations,
        selected_work_location_id=work_location_id,
        orgs=orgs,
        dirs=dirs,
        depts=depts,
        divs=divs,
        organization_id=organization_id,
        directorate_id=directorate_id,
        department_id=department_id,
        division_id=division_id,
        from_date=from_date.strftime('%Y-%m-%d') if from_date else '',
        to_date=to_date.strftime('%Y-%m-%d') if to_date else '',
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
    )



@portal_bp.route('/hr/reports/attendance/absence', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_absence():
    """تقرير الغياب: يعتمد على AttendanceDailySummary (أيام بدون بصمات أو حالة ABSENT)."""
    employee_id = (request.args.get('user_id') or '').strip()
    employee_id = int(employee_id) if employee_id.isdigit() else None
    work_location_id = (request.args.get('work_location_id') or '').strip()
    work_location_id = int(work_location_id) if work_location_id.isdigit() else None
    from_date = _parse_yyyy_mm_dd(request.args.get('from') or request.args.get('from_date'))
    to_date = _parse_yyyy_mm_dd(request.args.get('to') or request.args.get('to_date'))

    users = _list_hr_users()
    work_locations = _hr_lookup_options('WORK_LOCATION')

    loc_map = {x.id: x.name for x in (work_locations or [])}

    user_ids = _filtered_user_ids(employee_id=employee_id, work_location_id=work_location_id)
    rows_view = []

    if user_ids:
        q = AttendanceDailySummary.query.filter(AttendanceDailySummary.user_id.in_(user_ids))
        if from_date:
            q = q.filter(AttendanceDailySummary.day >= from_date.strftime('%Y-%m-%d'))
        if to_date:
            q = q.filter(AttendanceDailySummary.day <= to_date.strftime('%Y-%m-%d'))

        # غياب = بدون بصمة دخول وخروج، أو حالة ABSENT (إن تم ضبطها عبر الحالات الخاصة)
        q = q.filter(or_(func.upper(AttendanceDailySummary.status) == 'ABSENT',
                         and_(AttendanceDailySummary.first_in.is_(None), AttendanceDailySummary.last_out.is_(None))))

        rows_view = q.order_by(AttendanceDailySummary.day.desc()).limit(2000).all()

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)

        headers = ['التاريخ', 'الرقم الوظيفي', 'الموظف', 'موقع العمل', 'أول دخول', 'آخر خروج', 'الحالة']
        xrows = []
        for a in rows_view:
            u = getattr(a, 'user', None)
            ef = getattr(u, 'employee_file', None) if u else None
            loc = loc_map.get(getattr(ef, 'work_location_lookup_id', None), '') if ef else ''
            name = (getattr(ef, 'full_name_quad', None) or getattr(u, 'full_name', None) or getattr(u, 'name', None) or getattr(u, 'email', None) or '') if u else ''
            emp_no = getattr(ef, 'employee_no', '') if ef else ''
            fi = a.first_in.isoformat(sep=' ', timespec='minutes') if a.first_in else ''
            lo = a.last_out.isoformat(sep=' ', timespec='minutes') if a.last_out else ''
            st = (a.status or '')
            if (not fi) and (not lo) and (st or '').upper() == 'INCOMPLETE':
                st = 'ABSENT'
            xrows.append([a.day or '', emp_no, name, loc, fi, lo, st])

        return _export_xlsx('hr_absence.xlsx', headers, xrows)

    return render_template(
        'portal/hr/reports_absence.html',
        rows=rows_view,
        loc_map=loc_map,
        users=users,
        selected_user_id=employee_id,
        work_locations=work_locations,
        selected_work_location_id=work_location_id,
        from_date=from_date.strftime('%Y-%m-%d') if from_date else '',
        to_date=to_date.strftime('%Y-%m-%d') if to_date else '',
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
    )


@portal_bp.route('/hr/reports/attendance/fingerprints', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_fingerprints():
    """تقرير البصمات (الحركات): الموظف/موقع العمل/نطاق التاريخ + (مقروءة/غير مقروءة).

    تعريف (مقروءة): تم احتساب اليوم ضمن AttendanceDailySummary.
    تعريف (غير مقروءة): توجد حركات ولكن لم يتم احتساب اليوم بعد.
    """
    employee_id = (request.args.get('user_id') or '').strip()
    employee_id = int(employee_id) if employee_id.isdigit() else None
    work_location_id = (request.args.get('work_location_id') or '').strip()
    work_location_id = int(work_location_id) if work_location_id.isdigit() else None
    from_date = _parse_yyyy_mm_dd(request.args.get('from') or request.args.get('from_date'))
    to_date = _parse_yyyy_mm_dd(request.args.get('to') or request.args.get('to_date'))

    move_read = (request.args.get('move_read') or '').strip().lower()  # read|unread|''

    users = _list_hr_users()
    work_locations = _hr_lookup_options('WORK_LOCATION')

    # Candidate users from employee file filters
    user_ids = _filtered_user_ids(employee_id=employee_id, work_location_id=work_location_id)

    events = []
    if user_ids:
        q = AttendanceEvent.query.filter(AttendanceEvent.user_id.in_(user_ids))
        if from_date:
            dt_from = datetime.fromisoformat(from_date.strftime('%Y-%m-%d') + 'T00:00:00')
            q = q.filter(AttendanceEvent.event_dt >= dt_from)
        if to_date:
            dt_to = datetime.fromisoformat(to_date.strftime('%Y-%m-%d') + 'T23:59:59')
            q = q.filter(AttendanceEvent.event_dt <= dt_to)

        # Fetch enough rows then apply read/unread in python (DB-portable)
        limit = 10000 if move_read in ('read', 'unread') else 5000
        events = q.order_by(AttendanceEvent.event_dt.desc()).limit(limit).all()

    # Build summary keys for read/unread detection
    keys = set()
    if events:
        uids = sorted({e.user_id for e in events})
        days = sorted({e.event_dt.date().isoformat() for e in events if e.event_dt})
        if uids and days:
            for uid, day in (AttendanceDailySummary.query
                             .filter(AttendanceDailySummary.user_id.in_(uids))
                             .filter(AttendanceDailySummary.day.in_(days))
                             .with_entities(AttendanceDailySummary.user_id, AttendanceDailySummary.day)
                             .all()):
                keys.add((uid, day))

    # lookup maps
    loc_map = {x.id: x.name for x in (work_locations or [])}

    rows_view = []
    for e in (events or []):
        day = e.event_dt.date().isoformat() if e.event_dt else ''
        is_read = (e.user_id, day) in keys if day else False

        if move_read == 'read' and not is_read:
            continue
        if move_read == 'unread' and is_read:
            continue

        u = getattr(e, 'user', None)
        ef = getattr(u, 'employee_file', None) if u else None
        loc = loc_map.get(getattr(ef, 'work_location_lookup_id', None), '') if ef else ''
        name = (getattr(ef, 'full_name_quad', None) or getattr(u, 'full_name', None) or getattr(u, 'name', None) or getattr(u, 'email', None) or '') if u else ''

        rows_view.append({
            'day': day,
            'event_dt': e.event_dt,
            'event_type': e.event_type,
            'device_id': e.device_id,
            'batch_id': e.batch_id,
            'raw': e.raw_line,
            'user': u,
            'name': name,
            'employee_no': getattr(ef, 'employee_no', '') if ef else '',
            'work_location': loc,
            'is_read': is_read,
        })

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)
        headers = ['التاريخ', 'الوقت', 'الرقم الوظيفي', 'الموظف', 'موقع العمل', 'نوع الحركة', 'الجهاز', 'دفعة', 'الحالة (مقروءة؟)', 'Raw']
        xrows = []
        for r in rows_view:
            dt = r.get('event_dt')
            day = r.get('day') or ''
            tm = dt.strftime('%H:%M:%S') if dt else ''
            et = (r.get('event_type') or '').upper()
            et_label = 'دخول' if et == 'I' else ('خروج' if et == 'O' else et)
            xrows.append([
                day,
                tm,
                r.get('employee_no') or '',
                r.get('name') or '',
                r.get('work_location') or '',
                et_label,
                r.get('device_id') or '',
                r.get('batch_id') or '',
                'مقروءة' if r.get('is_read') else 'غير مقروءة',
                r.get('raw') or '',
            ])
        return _export_xlsx('hr_fingerprints.xlsx', headers, xrows)

    return render_template(
        'portal/hr/reports_fingerprints.html',
        rows=rows_view,
        users=users,
        selected_user_id=employee_id,
        work_locations=work_locations,
        selected_work_location_id=work_location_id,
        from_date=from_date.strftime('%Y-%m-%d') if from_date else '',
        to_date=to_date.strftime('%Y-%m-%d') if to_date else '',
        move_read=move_read,
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
    )


# ------------------------------------------
# HR Reports - Attendance (Extra)
# ------------------------------------------

@portal_bp.route('/hr/reports/attendance/permissions-fingerprint', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_permissions_fingerprint():
    """تقرير المغادرات مع البصمة.

    - يعتمد على HRPermissionRequest + AttendanceEvent.
    - الهدف هو اكتشاف المشاكل:
        * بلا مشاكل
        * مغادرة مدخلة بدون بصمة
        * مغادرة بنوع والبصمة بنوع آخر
        * بصمة مدخلة بدون إدخال إذن مغادرة
    """
    employee_id = (request.args.get('user_id') or '').strip()
    employee_id = int(employee_id) if employee_id.isdigit() else None
    work_location_id = (request.args.get('work_location_id') or '').strip()
    work_location_id = int(work_location_id) if work_location_id.isdigit() else None
    from_date = _parse_yyyy_mm_dd(request.args.get('from') or request.args.get('from_date'))
    to_date = _parse_yyyy_mm_dd(request.args.get('to') or request.args.get('to_date'))

    perm_type_id = (request.args.get('permission_type_id') or '').strip()
    perm_type_id = int(perm_type_id) if perm_type_id.isdigit() else None

    issue = (request.args.get('issue') or '').strip().lower()  # ok|perm_no_fp|mismatch|fp_no_perm|''

    users = _list_hr_users()
    work_locations = _hr_lookup_options('WORK_LOCATION')
    perm_types = HRPermissionType.query.filter(HRPermissionType.is_active.is_(True)).order_by(HRPermissionType.name_ar.asc()).all()

    loc_map = {x.id: x.name for x in (work_locations or [])}
    ptype_map = {x.id: (x.name_ar or x.code) for x in (perm_types or [])}

    rows_view = []

    # Need a date range to make this report meaningful
    if not from_date or not to_date:
        # render empty with filters
        if (request.args.get('export') or '').lower() == 'xlsx':
            abort(400)
        return render_template(
            'portal/hr/reports_permissions_fingerprint.html',
            rows=rows_view,
            users=users,
            selected_user_id=employee_id,
            work_locations=work_locations,
            selected_work_location_id=work_location_id,
            perm_types=perm_types,
            selected_permission_type_id=perm_type_id,
            from_date=from_date.strftime('%Y-%m-%d') if from_date else '',
            to_date=to_date.strftime('%Y-%m-%d') if to_date else '',
            issue=issue,
            can_export=current_user.has_perm(HR_REPORTS_EXPORT),
        )

    user_ids = _filtered_user_ids(employee_id=employee_id, work_location_id=work_location_id)

    if not user_ids:
        # nothing
        pass
    else:
        # Fetch permissions in range
        pq = HRPermissionRequest.query.filter(HRPermissionRequest.user_id.in_(user_ids))
        pq = pq.filter(HRPermissionRequest.day >= from_date.strftime('%Y-%m-%d'))
        pq = pq.filter(HRPermissionRequest.day <= to_date.strftime('%Y-%m-%d'))
        if perm_type_id:
            pq = pq.filter(HRPermissionRequest.permission_type_id == perm_type_id)
        pq = pq.filter(or_(HRPermissionRequest.status.is_(None), func.upper(HRPermissionRequest.status) != 'CANCELLED'))
        perms = pq.order_by(HRPermissionRequest.day.desc()).limit(5000).all()

        # Prefetch events in range (+tolerance)
        tol_min = 30
        start_dt = datetime.fromisoformat(from_date.strftime('%Y-%m-%d') + 'T00:00:00') - timedelta(minutes=tol_min)
        end_dt = datetime.fromisoformat(to_date.strftime('%Y-%m-%d') + 'T23:59:59') + timedelta(minutes=tol_min)

        eq = AttendanceEvent.query.filter(AttendanceEvent.user_id.in_(user_ids))
        eq = eq.filter(AttendanceEvent.event_dt >= start_dt)
        eq = eq.filter(AttendanceEvent.event_dt <= end_dt)
        events = eq.order_by(AttendanceEvent.event_dt.asc()).limit(200000).all()

        # Group events by (user_id, day)
        ev_map = {}
        for e in events:
            if not e.event_dt:
                continue
            d = e.event_dt.date().isoformat()
            ev_map.setdefault((e.user_id, d), []).append(e)

        # Helper: nearest event in window
        def _nearest(user_id:int, day_str:str, target_dt:datetime, window_minutes:int=30):
            lst = ev_map.get((user_id, day_str)) or []
            best = None
            best_diff = None
            for e in lst:
                if not e.event_dt:
                    continue
                diff = abs((e.event_dt - target_dt).total_seconds())
                if diff <= window_minutes * 60:
                    if best is None or diff < best_diff:
                        best = e
                        best_diff = diff
            return best

        # Permissions -> rows
        perm_windows = {}  # (uid,day) -> list[(from_dt,to_dt)]
        for r in perms:
            day_str = (r.day or '')
            uid = r.user_id
            if not day_str:
                continue
            f = (r.from_time or '').strip()
            t = (r.to_time or '').strip()
            dt_f = None
            dt_t = None
            try:
                if f and len(f) == 5:
                    dt_f = datetime.fromisoformat(day_str + 'T' + f + ':00')
            except Exception:
                dt_f = None
            try:
                if t and len(t) == 5:
                    dt_t = datetime.fromisoformat(day_str + 'T' + t + ':00')
            except Exception:
                dt_t = None

            if dt_f and dt_t and dt_t < dt_f:
                # ignore bad ranges
                dt_t = None

            if dt_f or dt_t:
                perm_windows.setdefault((uid, day_str), []).append((dt_f, dt_t, r))

            ev_from = _nearest(uid, day_str, dt_f, tol_min) if dt_f else None
            ev_to = _nearest(uid, day_str, dt_t, tol_min) if dt_t else None

            # expected types
            exp_from = 'O' if dt_f else None
            exp_to = 'I' if dt_t else None

            # compute issue
            issue_code = 'ok'
            if (dt_f and not ev_from) and (dt_t and not ev_to):
                issue_code = 'perm_no_fp'
            else:
                mism = False
                if exp_from and ev_from and (ev_from.event_type or '').upper() not in (exp_from, ''):
                    mism = True
                if exp_to and ev_to and (ev_to.event_type or '').upper() not in (exp_to, ''):
                    mism = True
                if mism:
                    issue_code = 'mismatch'
                # if one side missing
                if issue_code == 'ok':
                    if (dt_f and not ev_from) or (dt_t and not ev_to):
                        issue_code = 'perm_no_fp'

            if issue and issue_code != issue:
                continue

            # Map user info
            u = None
            try:
                u = User.query.get(uid)
            except Exception:
                u = None
            ef = getattr(u, 'employee_file', None) if u else None
            loc = loc_map.get(getattr(ef, 'work_location_lookup_id', None), '') if ef else ''
            name = (getattr(ef, 'full_name_quad', None) or getattr(u, 'full_name', None) or getattr(u, 'name', None) or getattr(u, 'email', None) or '') if u else ''

            rows_view.append({
                'row_kind': 'PERMISSION',
                'day': day_str,
                'user_id': uid,
                'employee_no': getattr(ef, 'employee_no', '') if ef else '',
                'name': name,
                'work_location': loc,
                'permission_type': ptype_map.get(getattr(r, 'permission_type_id', None), ''),
                'permission_type_id': getattr(r, 'permission_type_id', None),
                'from_time': (r.from_time or ''),
                'to_time': (r.to_time or ''),
                'fp_from_dt': ev_from.event_dt if ev_from else None,
                'fp_from_type': (ev_from.event_type or '').upper() if ev_from else '',
                'fp_to_dt': ev_to.event_dt if ev_to else None,
                'fp_to_type': (ev_to.event_type or '').upper() if ev_to else '',
                'issue': issue_code,
                'note': (r.note or ''),
            })

        # Fingerprint-only: detect mid-day OUT->IN pair without permission
        if issue in ('', 'fp_no_perm'):
            for (uid, day_str), lst in ev_map.items():
                # keep within main date range
                if day_str < from_date.strftime('%Y-%m-%d') or day_str > to_date.strftime('%Y-%m-%d'):
                    continue
                # user filter already applied
                # generate simple pairs
                lst_sorted = sorted([e for e in lst if e and e.event_dt], key=lambda x: x.event_dt)
                for i in range(len(lst_sorted) - 1):
                    e1 = lst_sorted[i]
                    e2 = lst_sorted[i+1]
                    t1 = (e1.event_type or '').upper()
                    t2 = (e2.event_type or '').upper()
                    if t1 != 'O' or t2 != 'I':
                        continue
                    # within 6 hours
                    delta_h = (e2.event_dt - e1.event_dt).total_seconds() / 3600.0
                    if delta_h <= 0 or delta_h > 6:
                        continue
                    # ignore if at start/end of day (likely attendance in/out)
                    if e1.event_dt.hour < 9 or e2.event_dt.hour > 17:
                        continue

                    # check overlap with any permission in same day
                    has_perm = False
                    for (dt_f, dt_t, r) in (perm_windows.get((uid, day_str)) or []):
                        # if permission has a window, compare proximity
                        try:
                            if dt_f and abs((dt_f - e1.event_dt).total_seconds()) <= tol_min*60:
                                has_perm = True
                                break
                            if dt_t and abs((dt_t - e2.event_dt).total_seconds()) <= tol_min*60:
                                has_perm = True
                                break
                        except Exception:
                            pass
                    if has_perm:
                        continue

                    # build row
                    u = None
                    try:
                        u = User.query.get(uid)
                    except Exception:
                        u = None
                    ef = getattr(u, 'employee_file', None) if u else None
                    loc = loc_map.get(getattr(ef, 'work_location_lookup_id', None), '') if ef else ''
                    name = (getattr(ef, 'full_name_quad', None) or getattr(u, 'full_name', None) or getattr(u, 'name', None) or getattr(u, 'email', None) or '') if u else ''

                    rows_view.append({
                        'row_kind': 'FINGERPRINT_ONLY',
                        'day': day_str,
                        'user_id': uid,
                        'employee_no': getattr(ef, 'employee_no', '') if ef else '',
                        'name': name,
                        'work_location': loc,
                        'permission_type': '',
                        'permission_type_id': None,
                        'from_time': '',
                        'to_time': '',
                        'fp_from_dt': e1.event_dt,
                        'fp_from_type': 'O',
                        'fp_to_dt': e2.event_dt,
                        'fp_to_type': 'I',
                        'issue': 'fp_no_perm',
                        'note': '',
                    })

        # sort (date desc, name)
        rows_view.sort(key=lambda r: (r.get('day') or '', r.get('name') or ''), reverse=True)

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)
        headers = ['التاريخ', 'الرقم الوظيفي', 'الموظف', 'موقع العمل', 'نوع المغادرة', 'من (المغادرة)', 'إلى (العودة)', 'بصمة من', 'نوع بصمة من', 'بصمة إلى', 'نوع بصمة إلى', 'المشكلة']
        xrows = []
        for r in rows_view:
            def _fmt_dt(dt):
                try:
                    return dt.strftime('%Y-%m-%d %H:%M:%S') if dt else ''
                except Exception:
                    return ''
            issue_map = {
                'ok': 'بلا مشاكل',
                'perm_no_fp': 'مغادرة مدخلة بدون بصمة',
                'mismatch': 'مغادرة بنوع والبصمة بنوع آخر',
                'fp_no_perm': 'بصمة مدخلة بدون إدخال إذن مغادرة',
            }
            xrows.append([
                r.get('day') or '',
                r.get('employee_no') or '',
                r.get('name') or '',
                r.get('work_location') or '',
                r.get('permission_type') or '',
                r.get('from_time') or '',
                r.get('to_time') or '',
                _fmt_dt(r.get('fp_from_dt')),
                r.get('fp_from_type') or '',
                _fmt_dt(r.get('fp_to_dt')),
                r.get('fp_to_type') or '',
                issue_map.get(r.get('issue') or '', r.get('issue') or ''),
            ])
        return _export_xlsx('hr_permissions_fingerprint.xlsx', headers, xrows)

    return render_template(
        'portal/hr/reports_permissions_fingerprint.html',
        rows=rows_view,
        users=users,
        selected_user_id=employee_id,
        work_locations=work_locations,
        selected_work_location_id=work_location_id,
        perm_types=perm_types,
        selected_permission_type_id=perm_type_id,
        from_date=from_date.strftime('%Y-%m-%d') if from_date else '',
        to_date=to_date.strftime('%Y-%m-%d') if to_date else '',
        issue=issue,
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
    )


@portal_bp.route('/hr/reports/attendance/daily', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_attendance_daily():
    """التقرير اليومي (شهري): موقع العمل + الإدارة/الدائرة/القسم/الشعبة + الشهر والسنة."""
    work_location_id = (request.args.get('work_location_id') or '').strip()
    work_location_id = int(work_location_id) if work_location_id.isdigit() else None

    organization_id = (request.args.get('organization_id') or '').strip()
    organization_id = int(organization_id) if organization_id.isdigit() else None

    directorate_id = (request.args.get('directorate_id') or '').strip()
    directorate_id = int(directorate_id) if directorate_id.isdigit() else None

    department_id = (request.args.get('department_id') or '').strip()
    department_id = int(department_id) if department_id.isdigit() else None

    division_id = (request.args.get('division_id') or '').strip()
    division_id = int(division_id) if division_id.isdigit() else None

    year_raw = (request.args.get('year') or '').strip()
    month_raw = (request.args.get('month') or '').strip()

    try:
        year = int(year_raw) if year_raw else datetime.utcnow().year
    except Exception:
        year = datetime.utcnow().year
    try:
        month = int(month_raw) if month_raw else datetime.utcnow().month
    except Exception:
        month = datetime.utcnow().month

    work_locations = _hr_lookup_options('WORK_LOCATION')
    loc_map = {x.id: x.name for x in (work_locations or [])}

    orgs = Organization.query.filter(Organization.is_active.is_(True)).order_by(Organization.name_ar.asc()).all()
    dirs = Directorate.query.filter(Directorate.is_active.is_(True)).order_by(Directorate.name_ar.asc()).all()
    depts = Department.query.filter(Department.is_active.is_(True)).order_by(Department.name_ar.asc()).all()
    divs = Division.query.filter(Division.is_active.is_(True)).order_by(Division.name_ar.asc()).all()

    # Month range
    try:
        start_day = date(year, month, 1)
    except Exception:
        start_day = date(datetime.utcnow().year, datetime.utcnow().month, 1)
    # end_day: next month - 1
    if start_day.month == 12:
        end_day = date(start_day.year + 1, 1, 1) - timedelta(days=1)
    else:
        end_day = date(start_day.year, start_day.month + 1, 1) - timedelta(days=1)

    user_ids = _filtered_user_ids(
        work_location_id=work_location_id,
        organization_id=organization_id,
        directorate_id=directorate_id,
        department_id=department_id,
        division_id=division_id,
    )

    # build summary per employee
    rows = []
    if user_ids:
        q = AttendanceDailySummary.query.filter(AttendanceDailySummary.user_id.in_(user_ids))
        q = q.filter(AttendanceDailySummary.day >= start_day.strftime('%Y-%m-%d'))
        q = q.filter(AttendanceDailySummary.day <= end_day.strftime('%Y-%m-%d'))
        recs = q.all()

        # preload users
        ulist = User.query.filter(User.id.in_(user_ids)).all()
        umap = {u.id: u for u in ulist}

        # aggregate
        agg = {}
        for a in recs:
            uid = a.user_id
            g = agg.setdefault(uid, {
                'days': 0,
                'absent': 0,
                'incomplete': 0,
                'work_minutes': 0,
                'late_minutes': 0,
                'early_minutes': 0,
                'overtime_minutes': 0,
            })
            g['days'] += 1
            st = (a.status or '').upper()
            if st == 'ABSENT' or (not a.first_in and not a.last_out):
                g['absent'] += 1
            elif st == 'INCOMPLETE':
                g['incomplete'] += 1
            g['work_minutes'] += int(a.work_minutes or 0)
            g['late_minutes'] += int(a.late_minutes or 0)
            g['early_minutes'] += int(a.early_leave_minutes or 0)
            g['overtime_minutes'] += int(a.overtime_minutes or 0)

        for uid, g in agg.items():
            u = umap.get(uid)
            ef = getattr(u, 'employee_file', None) if u else None
            rows.append({
                'user_id': uid,
                'employee_no': getattr(ef, 'employee_no', '') if ef else '',
                'name': (getattr(ef, 'full_name_quad', None) or getattr(u, 'full_name', None) or getattr(u, 'name', None) or getattr(u, 'email', None) or '') if u else '',
                'work_location': loc_map.get(getattr(ef, 'work_location_lookup_id', None), '') if ef else '',
                'organization': getattr(getattr(ef, 'organization', None), 'name_ar', None) if ef else None,
                'directorate': getattr(getattr(ef, 'directorate', None), 'name_ar', None) if ef else None,
                'department': getattr(getattr(ef, 'department', None), 'name_ar', None) if ef else None,
                'division': getattr(getattr(ef, 'division', None), 'name_ar', None) if ef else None,
                **g,
            })

        rows.sort(key=lambda r: (r.get('name') or ''), reverse=False)

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)
        headers = ['الشهر', 'الرقم الوظيفي', 'الموظف', 'موقع العمل', 'الإدارة', 'الدائرة', 'القسم', 'الشعبة', 'عدد الأيام', 'غياب', 'غير مكتمل', 'ساعات الدوام', 'دقائق تأخير', 'دقائق خروج مبكر', 'دقائق إضافي']
        xrows = []
        for r in rows:
            xrows.append([
                f"{year:04d}-{month:02d}",
                r.get('employee_no') or '',
                r.get('name') or '',
                r.get('work_location') or '',
                r.get('organization') or '',
                r.get('directorate') or '',
                r.get('department') or '',
                r.get('division') or '',
                r.get('days') or 0,
                r.get('absent') or 0,
                r.get('incomplete') or 0,
                round(float(r.get('work_minutes') or 0)/60.0, 2),
                r.get('late_minutes') or 0,
                r.get('early_minutes') or 0,
                r.get('overtime_minutes') or 0,
            ])
        return _export_xlsx('hr_daily_report.xlsx', headers, xrows)

    return render_template(
        'portal/hr/reports_daily.html',
        rows=rows,
        work_locations=work_locations,
        selected_work_location_id=work_location_id,
        orgs=orgs,
        dirs=dirs,
        depts=depts,
        divs=divs,
        organization_id=organization_id,
        directorate_id=directorate_id,
        department_id=department_id,
        division_id=division_id,
        year=year,
        month=month,
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
    )


@portal_bp.route('/hr/reports/attendance/edits', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_attendance_edits():
    """تقرير مراقبة تعديلات الدوام."""
    editor_id = (request.args.get('editor_id') or '').strip()
    editor_id = int(editor_id) if editor_id.isdigit() else None

    target_user_id = (request.args.get('user_id') or '').strip()
    target_user_id = int(target_user_id) if target_user_id.isdigit() else None

    date_from = _parse_yyyy_mm_dd(request.args.get('from') or request.args.get('from_date'))
    date_to = _parse_yyyy_mm_dd(request.args.get('to') or request.args.get('to_date'))

    edit_type = (request.args.get('edit_type') or '').strip().lower()  # manual|delete|''
    only_affect_time = (request.args.get('affect_time') or '').strip().lower() in ('1', 'true', 'on', 'yes')

    users = _list_hr_users()

    rows = []

    # Manual edits: HRAttendanceSpecialCase
    if edit_type in ('', 'manual'):
        q = HRAttendanceSpecialCase.query
        if editor_id:
            q = q.filter(HRAttendanceSpecialCase.created_by_id == editor_id)
        if target_user_id:
            q = q.filter(HRAttendanceSpecialCase.user_id == target_user_id)
        if date_from:
            q = q.filter(HRAttendanceSpecialCase.created_at >= datetime.fromisoformat(date_from.strftime('%Y-%m-%d') + 'T00:00:00'))
        if date_to:
            q = q.filter(HRAttendanceSpecialCase.created_at <= datetime.fromisoformat(date_to.strftime('%Y-%m-%d') + 'T23:59:59'))

        sp_rows = q.order_by(HRAttendanceSpecialCase.created_at.desc()).limit(3000).all()

        for r in sp_rows:
            affects = False
            try:
                if r.start_time or r.end_time:
                    affects = True
            except Exception:
                pass
            try:
                if r.allow_morning_minutes is not None or r.allow_evening_minutes is not None:
                    affects = True
            except Exception:
                pass
            try:
                if (r.field or '').upper() in ('WORK_MINUTES', 'LATE_MINUTES', 'EARLY_LEAVE_MINUTES', 'OVERTIME_MINUTES'):
                    affects = True
            except Exception:
                pass
            if only_affect_time and not affects:
                continue

            rows.append({
                'ts': r.created_at,
                'edit_type': 'manual',
                'edited_by': getattr(r, 'created_by', None),
                'target_user': getattr(r, 'user', None),
                'affects_time': affects,
                'details': f"{(r.kind or '').upper()} / {(r.status or r.field or '')} ({r.day} إلى {r.day_to or r.day})",
                'note': (r.note or ''),
            })

    # Delete movements: use AuditLog if exists
    if edit_type in ('', 'delete'):
        aq = AuditLog.query
        if editor_id:
            aq = aq.filter(AuditLog.user_id == editor_id)
        if target_user_id:
            aq = aq.filter(or_(AuditLog.target_id == target_user_id, AuditLog.note.ilike(f"%{target_user_id}%")))
        if date_from:
            aq = aq.filter(AuditLog.created_at >= datetime.fromisoformat(date_from.strftime('%Y-%m-%d') + 'T00:00:00'))
        if date_to:
            aq = aq.filter(AuditLog.created_at <= datetime.fromisoformat(date_to.strftime('%Y-%m-%d') + 'T23:59:59'))

        # heuristic: actions that mention attendance + delete
        logs = aq.order_by(AuditLog.created_at.desc()).limit(3000).all()
        for l in logs:
            act = (l.action or '').lower()
            note = (l.note or '').lower()
            if ('attendance' not in act and 'attendance' not in note and 'بصم' not in note and 'حركة' not in note):
                continue
            if ('delete' not in act and 'حذف' not in note and 'remove' not in act):
                continue
            # affects time by definition
            if only_affect_time is False or True:
                pass
            rows.append({
                'ts': l.created_at,
                'edit_type': 'delete',
                'edited_by': getattr(l, 'user', None),
                'target_user': None,
                'affects_time': True,
                'details': l.action or 'DELETE',
                'note': l.note or '',
            })

    rows.sort(key=lambda r: (r.get('ts') or datetime.min), reverse=True)

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)
        headers = ['تاريخ التعديل', 'نوع التعديل', 'من قام بالتعديل', 'الموظف', 'تأثير على التوقيت', 'تفاصيل', 'ملاحظات']
        xrows = []
        for r in rows:
            ts = r.get('ts')
            xrows.append([
                ts.strftime('%Y-%m-%d %H:%M:%S') if ts else '',
                'تعديل يدوي' if r.get('edit_type') == 'manual' else 'حذف حركات',
                (getattr(r.get('edited_by'), 'full_name', None) or getattr(r.get('edited_by'), 'name', None) or getattr(r.get('edited_by'), 'email', None) or '') if r.get('edited_by') else '',
                (getattr(r.get('target_user'), 'full_name', None) or getattr(r.get('target_user'), 'name', None) or getattr(r.get('target_user'), 'email', None) or '') if r.get('target_user') else '',
                'نعم' if r.get('affects_time') else 'لا',
                r.get('details') or '',
                r.get('note') or '',
            ])
        return _export_xlsx('hr_attendance_edits.xlsx', headers, xrows)

    return render_template(
        'portal/hr/reports_attendance_edits.html',
        rows=rows,
        users=users,
        editor_id=editor_id,
        target_user_id=target_user_id,
        date_from=date_from.strftime('%Y-%m-%d') if date_from else '',
        date_to=date_to.strftime('%Y-%m-%d') if date_to else '',
        edit_type=edit_type,
        only_affect_time=only_affect_time,
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
    )


@portal_bp.route('/hr/reports/attendance/negative-movements', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_negative_movements():
    """تقرير الحركات السالبة: موقع العمل + من/إلى."""
    work_location_id = (request.args.get('work_location_id') or '').strip()
    work_location_id = int(work_location_id) if work_location_id.isdigit() else None
    from_date = _parse_yyyy_mm_dd(request.args.get('from') or request.args.get('from_date'))
    to_date = _parse_yyyy_mm_dd(request.args.get('to') or request.args.get('to_date'))

    work_locations = _hr_lookup_options('WORK_LOCATION')
    loc_map = {x.id: x.name for x in (work_locations or [])}

    user_ids = _filtered_user_ids(work_location_id=work_location_id)

    rows = []
    if user_ids:
        q = AttendanceDailySummary.query.filter(AttendanceDailySummary.user_id.in_(user_ids))
        if from_date:
            q = q.filter(AttendanceDailySummary.day >= from_date.strftime('%Y-%m-%d'))
        if to_date:
            q = q.filter(AttendanceDailySummary.day <= to_date.strftime('%Y-%m-%d'))

        q = q.filter(or_(AttendanceDailySummary.work_minutes < 0,
                         AttendanceDailySummary.late_minutes < 0,
                         AttendanceDailySummary.early_leave_minutes < 0,
                         AttendanceDailySummary.overtime_minutes < 0))

        rows = q.order_by(AttendanceDailySummary.day.desc()).limit(5000).all()

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)
        headers = ['التاريخ', 'الرقم الوظيفي', 'الموظف', 'موقع العمل', 'ساعات الدوام', 'تأخير', 'خروج مبكر', 'إضافي', 'الحالة']
        xrows = []
        for a in rows:
            u = getattr(a, 'user', None)
            ef = getattr(u, 'employee_file', None) if u else None
            loc = loc_map.get(getattr(ef, 'work_location_lookup_id', None), '') if ef else ''
            name = (getattr(ef, 'full_name_quad', None) or getattr(u, 'full_name', None) or getattr(u, 'name', None) or getattr(u, 'email', None) or '') if u else ''
            emp_no = getattr(ef, 'employee_no', '') if ef else ''
            xrows.append([
                a.day or '',
                emp_no,
                name,
                loc,
                round(float(a.work_minutes or 0)/60.0, 2),
                a.late_minutes,
                a.early_leave_minutes,
                a.overtime_minutes,
                a.status or '',
            ])
        return _export_xlsx('hr_negative_movements.xlsx', headers, xrows)

    return render_template(
        'portal/hr/reports_negative_movements.html',
        rows=rows,
        work_locations=work_locations,
        selected_work_location_id=work_location_id,
        from_date=from_date.strftime('%Y-%m-%d') if from_date else '',
        to_date=to_date.strftime('%Y-%m-%d') if to_date else '',
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
    )
@portal_bp.route('/hr/reports/attendance/diwan', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_diwan():
    work_location_id = (request.args.get('work_location_id') or '').strip()
    work_location_id = int(work_location_id) if work_location_id.isdigit() else None
    year_raw = (request.args.get('year') or '').strip()
    month_raw = (request.args.get('month') or '').strip()
    appointment_type_id = (request.args.get('appointment_type_id') or '').strip()
    appointment_type_id = int(appointment_type_id) if appointment_type_id.isdigit() else None

    # Optional: single-employee report (choose employee from the UI)
    user_id_raw = (request.args.get('user_id') or '').strip()
    selected_user_id = int(user_id_raw) if user_id_raw.isdigit() else None

    try:
        year = int(year_raw) if year_raw else datetime.utcnow().year
    except Exception:
        year = datetime.utcnow().year
    try:
        month = int(month_raw) if month_raw else datetime.utcnow().month
    except Exception:
        month = datetime.utcnow().month

    view = (request.args.get('view') or '').strip().lower()
    users = _list_hr_users()
    work_locations = _hr_lookup_options('WORK_LOCATION')
    appointment_types = _hr_lookup_options('APPOINTMENT_TYPE')

    # Boundaries
    try:
        start = date(year, month, 1)
        if month == 12:
            end = date(year + 1, 1, 1) - timedelta(days=1)
        else:
            end = date(year, month + 1, 1) - timedelta(days=1)
    except Exception:
        start = date(datetime.utcnow().year, datetime.utcnow().month, 1)
        end = start

    # If a specific employee is selected, generate the report for that employee only
    # (even if they have no attendance/leave rows).
    if selected_user_id:
        user_ids = [selected_user_id]
    else:
        user_ids = _filtered_user_ids(work_location_id=work_location_id, appointment_type_id=appointment_type_id)
    rows_view = []
    if user_ids:
        q = AttendanceDailySummary.query.filter(AttendanceDailySummary.user_id.in_(user_ids))            .filter(AttendanceDailySummary.day >= start.strftime('%Y-%m-%d'))            .filter(AttendanceDailySummary.day <= end.strftime('%Y-%m-%d'))

        # Aggregate in python (sqlite safe)
        agg = {}
        for a in q.all():
            d = agg.setdefault(a.user_id, {'work': 0, 'late': 0, 'early': 0, 'overtime': 0, 'days': 0, 'absent': 0})
            d['days'] += 1
            d['work'] += int(a.work_minutes or 0)
            d['late'] += int(a.late_minutes or 0)
            d['early'] += int(a.early_leave_minutes or 0)
            d['overtime'] += int(a.overtime_minutes or 0)
            if (a.status or '').upper() == 'ABSENT':
                d['absent'] += 1

        for uid, d in agg.items():
            u = None
            try:
                u = User.query.get(uid)
            except Exception:
                u = None
            ef = getattr(u, 'employee_file', None) if u else None
            rows_view.append({
                'user': u,
                'employee_no': (ef.employee_no if ef else ''),
                'name': (ef.full_name_quad if ef and ef.full_name_quad else (u.full_name if u else '')),
                'work_hours': round(d['work'] / 60.0, 2),
                'late_hours': round(d['late'] / 60.0, 2),
                'early_hours': round(d['early'] / 60.0, 2),
                'overtime_hours': round(d['overtime'] / 60.0, 2),
                'days': d['days'],
                'absent_days': d['absent'],
            })
        rows_view.sort(key=lambda x: (x.get('name') or ''), reverse=False)

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)
        headers = ['الرقم الوظيفي', 'الموظف', 'أيام', 'غياب', 'ساعات عمل', 'ساعات تأخير', 'ساعات خروج مبكر', 'ساعات عمل إضافي']
        xrows = []
        for r in rows_view:
            xrows.append([
                r.get('employee_no') or '',
                r.get('name') or '',
                r.get('days') or 0,
                r.get('absent_days') or 0,
                r.get('work_hours') or 0,
                r.get('late_hours') or 0,
                r.get('early_hours') or 0,
                r.get('overtime_hours') or 0,
            ])
        return _export_xlsx('hr_diwan.xlsx', headers, xrows)


    work_location_name = 'الكل'
    if work_location_id:
        for x in work_locations:
            if x.id == work_location_id:
                work_location_name = (getattr(x, 'name', None) or getattr(x, 'label', None) or '').strip() or work_location_name
                break

    appointment_type_name = 'بدون اختيار'
    if appointment_type_id:
        for x in appointment_types:
            if x.id == appointment_type_id:
                appointment_type_name = (getattr(x, 'name', None) or getattr(x, 'label', None) or '').strip() or appointment_type_name
                break

    # -------------------------
    # Diwan (Council) view: generate per-employee "Diwan" report data
    # (based on the attached Diwan Word template)
    # -------------------------
    diwan_reports = []
    if view in ('diwan', 'council', 'diwan_report'):
        # Prefer employees that appear in the month summary; fallback to filtered employee ids.
        try:
            ids_from_rows = [getattr(r.get('user'), 'id', None) for r in (rows_view or []) if r.get('user')]
            employee_ids = [int(x) for x in ids_from_rows if x]
        except Exception:
            employee_ids = []
        if not employee_ids:
            employee_ids = list(user_ids or [])

        # Lookup maps for grade/title
        try:
            _grade_items = _hr_lookup_items_for_category('JOB_GRADE')
        except Exception:
            _grade_items = []
        try:
            _title_items = _hr_lookup_items_for_category('JOB_TITLE')
        except Exception:
            _title_items = []
        grade_map = {int(x.id): (getattr(x, 'label', None) or getattr(x, 'name', None) or '') for x in (_grade_items or []) if getattr(x, 'id', None)}
        title_map = {int(x.id): (getattr(x, 'label', None) or getattr(x, 'name', None) or '') for x in (_title_items or []) if getattr(x, 'id', None)}

        # Leave types by keywords (Arabic + code fallbacks)
        leave_types = []
        try:
            leave_types = HRLeaveType.query.filter_by(is_active=True).order_by(HRLeaveType.id.asc()).all()
        except Exception:
            leave_types = []

        def _pick_leave_type(ar_kw: list[str], code_kw: list[str]):
            for lt in (leave_types or []):
                try:
                    n = (getattr(lt, 'name_ar', '') or '').strip()
                    c = (getattr(lt, 'code', '') or '').upper().strip()
                    if any(k in n for k in ar_kw) or any(k in c for k in code_kw):
                        return lt
                except Exception:
                    continue
            return None

        lt_annual = _pick_leave_type(['سنوي'], ['ANNUAL', 'ANNUAL_LEAVE', 'PERSONAL'])
        lt_sick = _pick_leave_type(['مرض'], ['SICK', 'MEDICAL'])
        lt_casual = _pick_leave_type(['عارض'], ['CASUAL', 'EMERGENCY'])
        lt_no_pay = _pick_leave_type(['بدون راتب'], ['NO_PAY', 'UNPAID'])
        lt_hajj = _pick_leave_type(['حج'], ['HAJJ'])
        lt_maternity = _pick_leave_type(['أمومة', 'امومة'], ['MATERNITY'])

        y_start = date(year, 1, 1)
        y_end = date(year, 12, 31)
        as_of = min(date.today(), y_end) if year == date.today().year else y_end

        # Prefetch users
        users_map = {}
        try:
            users_map = {u.id: u for u in User.query.filter(User.id.in_(employee_ids)).all()}
        except Exception:
            users_map = {}

        def _fmt_date_dmy(d: date | None) -> str:
            try:
                return d.strftime('%d/%m/%Y') if d else ''
            except Exception:
                return ''

        def _leave_rows(uid: int, lt: HRLeaveType | None, with_days: bool = True):
            if not lt:
                return []
            out = []
            try:
                q = (HRLeaveRequest.query
                     .filter(HRLeaveRequest.user_id == uid)
                     .filter(HRLeaveRequest.leave_type_id == lt.id)
                     .filter(HRLeaveRequest.status.in_(['APPROVED', 'CANCELLED']))
                     .order_by(HRLeaveRequest.start_date.asc(), HRLeaveRequest.id.asc()))
                for r in q.all():
                    # Count CANCELLED only if it was cancelled from APPROVED
                    if (r.status or '').upper() == 'CANCELLED':
                        if (r.cancelled_from_status or '').upper() != 'APPROVED':
                            continue
                    s0 = _parse_yyyy_mm_dd(getattr(r, 'start_date', None))
                    e0 = _parse_yyyy_mm_dd(getattr(r, 'end_date', None))
                    if not s0 or not e0:
                        continue

                    # Effective end (respect cancellation effective date)
                    e_eff = e0
                    if (r.status or '').upper() == 'CANCELLED' and getattr(r, 'cancel_effective_date', None):
                        ce = _parse_yyyy_mm_dd(r.cancel_effective_date)
                        if ce:
                            e_eff = min(e_eff, ce)

                    # Intersect with year
                    s = max(s0, y_start)
                    e = min(e_eff, y_end)
                    if e < s:
                        continue

                    row = {'from': _fmt_date_dmy(s), 'to': _fmt_date_dmy(e)}
                    if with_days:
                        row['days'] = int((e - s).days + 1)
                    out.append(row)
            except Exception:
                return []
            return out

        for uid in employee_ids:
            u = users_map.get(uid)
            if not u:
                continue
            ef = getattr(u, 'employee_file', None)

            # Header values
            ministry = None
            try:
                ministry = (getattr(getattr(ef, 'organization', None), 'name_ar', None) or '').strip()
            except Exception:
                ministry = None
            if not ministry:
                ministry = 'اللجنة الوطنية الفلسطينية للتربية والثقافة والعلوم'

            unit = ''
            try:
                unit = (getattr(getattr(ef, 'department', None), 'name_ar', None) or '').strip()
            except Exception:
                unit = ''
            if not unit:
                try:
                    unit = (getattr(getattr(ef, 'directorate', None), 'name_ar', None) or '').strip()
                except Exception:
                    unit = ''
            if not unit:
                unit = work_location_name if work_location_name and work_location_name != 'الكل' else ''

            name = (getattr(ef, 'full_name_quad', None) or getattr(u, 'full_name', None) or getattr(u, 'name', None) or getattr(u, 'email', None) or '').strip()
            employee_no = (getattr(ef, 'employee_no', None) or '').strip()
            personal_no = (getattr(ef, 'national_id', None) or '').strip()
            hire_date = (getattr(ef, 'hire_date', None) or '').strip()
            try:
                hd = _parse_yyyy_mm_dd(hire_date)
                hire_date = _fmt_date_dmy(hd) if hd else hire_date
            except Exception:
                pass

            grade = ''
            try:
                gid = getattr(ef, 'job_grade_lookup_id', None)
                grade = grade_map.get(int(gid), '') if gid else ''
            except Exception:
                grade = ''

            job_title = ''
            try:
                tid = getattr(ef, 'job_title_lookup_id', None)
                job_title = title_map.get(int(tid), '') if tid else ''
            except Exception:
                job_title = ''
            if not job_title:
                try:
                    tid = getattr(ef, 'admin_title_lookup_id', None)
                    job_title = title_map.get(int(tid), '') if tid else ''
                except Exception:
                    job_title = ''

            # Entitlement/remaining for annual leave
            entitled_days = 0
            used_days = 0.0
            remaining_days = 0.0
            if lt_annual:
                try:
                    entitled_days = int(_leave_entitlement_days(uid, lt_annual, year) or 0)
                except Exception:
                    entitled_days = 0
                try:
                    used_days = float(_leave_used_days_as_of(uid, lt_annual.id, year, as_of) or 0.0)
                except Exception:
                    used_days = 0.0
                remaining_days = float(entitled_days) - float(used_days)

            diwan_reports.append({
                'ministry': ministry,
                'unit': unit,
                'report_date': _fmt_date_dmy(date.today()),
                'employee_name': name,
                'employee_no': employee_no,
                'personal_no': personal_no,
                'hire_date': hire_date,
                'grade': grade,
                'job_title': job_title,
                'entitled_days': entitled_days,
                'annual_remaining': round(float(remaining_days), 2),
                'annual_rows': _leave_rows(uid, lt_annual, with_days=True),
                'sick_rows': _leave_rows(uid, lt_sick, with_days=True),
                'casual_rows': _leave_rows(uid, lt_casual, with_days=True),
                'no_pay_rows': _leave_rows(uid, lt_no_pay, with_days=True),
                'hajj_rows': _leave_rows(uid, lt_hajj, with_days=False),
                'maternity_rows': _leave_rows(uid, lt_maternity, with_days=False),
            })

    
    # Ensure the Diwan view always shows a report "page" even if there are no employees/results.
    # This keeps the layout visible (per user request) and allows exporting an empty template too.
    if view in ('diwan', 'council', 'diwan_report') and not diwan_reports:
        diwan_reports = [{
            'ministry': '',
            'unit': '',
            'report_date': _fmt_date_dmy(date.today()),
            'employee_name': '',
            'employee_no': '',
            'personal_no': '',
            'hire_date': '',
            'grade': '',
            'job_title': '',
            'entitled_days': '',
            'annual_remaining': '',
            'annual_rows': [],
            'sick_rows': [],
            'casual_rows': [],
            'no_pay_rows': [],
            'hajj_rows': [],
            'maternity_rows': [],
        }]

    # Word export (DOCX) — uses the official Diwan template (provided by user).
    export_kind = (request.args.get('export') or '').strip().lower()
    if export_kind in ('docx', 'word'):
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)

        def _safe_part(s: str) -> str:
            try:
                s = (s or '').strip()
            except Exception:
                s = ''
            s = re.sub(r'[^0-9A-Za-z_\-]+', '_', s)
            s = re.sub(r'_+', '_', s).strip('_')
            return s or 'report'

        def _build_docx_bytes(rep: dict) -> bytes:
            from docx import Document  # python-docx
            tpl_path = os.path.join(current_app.root_path, 'assets', 'templates', 'hr', 'diwan_leave_template.docx')
            if not os.path.exists(tpl_path):
                # fallback (in case the folder structure changes)
                tpl_path = os.path.join(current_app.root_path, 'assets', 'diwan_leave_template.docx')
            doc = Document(tpl_path)

            # --- Helpers that preserve the original Word template layout ---
            def _replace_first_dots(text: str, value: str) -> str:
                try:
                    if not value:
                        return text
                    return re.sub(r'\.{3,}', str(value), text, count=1)
                except Exception:
                    return text

            def _fill_cell_dots(cell, value: str):
                """Replace the first dots-sequence inside a cell without rewriting the whole cell."""
                try:
                    if value in (None, ''):
                        return
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if run.text and re.search(r'\.{3,}', run.text):
                                run.text = _replace_first_dots(run.text, str(value))
                                return
                except Exception:
                    pass

            def _write_cell(table, r, c, value):
                """Write text into a (usually empty) cell while keeping paragraph formatting."""
                try:
                    cell = table.cell(r, c)
                    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph('')
                    if p.runs:
                        p.runs[0].text = str(value or '')
                        for rr in p.runs[1:]:
                            rr.text = ''
                    else:
                        p.add_run(str(value or ''))
                except Exception:
                    pass

            # Top line (same as template) — fill placeholders without rewriting the whole paragraph
            dt = (rep.get('report_date') or '').strip()  # expected dd/mm/yyyy
            dd = mm = yy = ''
            try:
                m = re.match(r'^(\d{1,2})/(\d{1,2})/(\d{4})$', dt)
                if m:
                    dd = str(m.group(1)).zfill(2)
                    mm = str(m.group(2)).zfill(2)
                    yy = str(m.group(3))[-2:]
            except Exception:
                pass

            ministry = (rep.get('ministry') or '').strip()
            unit = (rep.get('unit') or '').strip()
            try:
                if doc.paragraphs:
                    p0 = doc.paragraphs[0]
                    for run in p0.runs:
                        # Ministry + unit
                        if '......................' in (run.text or '') and ministry:
                            run.text = run.text.replace('......................', ministry)
                        if '..................' in (run.text or '') and unit:
                            run.text = run.text.replace('..................', unit)
                        # Date pieces (keep original slashes/spaces)
                        if 'التاريخ:' in (run.text or '') and '....' in (run.text or '') and dd:
                            run.text = run.text.replace('....', dd)
                        if '......' in (run.text or '') and mm:
                            run.text = run.text.replace('......', mm)
                        if (run.text or '').strip() == '20' and yy:
                            run.text = f"20{yy}"
            except Exception:
                pass

            # Info table (employee details) — replace dots only
            try:
                t0 = doc.tables[0]
                _fill_cell_dots(t0.cell(0, 0), rep.get('employee_name') or '')
                _fill_cell_dots(t0.cell(0, 1), rep.get('personal_no') or '')
                _fill_cell_dots(t0.cell(1, 0), rep.get('hire_date') or '')
                _fill_cell_dots(t0.cell(1, 1), rep.get('grade') or '')
                _fill_cell_dots(t0.cell(2, 0), rep.get('job_title') or '')
                _fill_cell_dots(t0.cell(2, 1), str(rep.get('entitled_days') if rep.get('entitled_days') not in (None, '') else ''))
            except Exception:
                pass

            try:
                t1 = doc.tables[1]

                # Annual/Sick/Casual blocks — fill into the pre-made blank rows (2..10)
                max_rows = 9
                annual = rep.get('annual_rows') or []
                sick = rep.get('sick_rows') or []
                casual = rep.get('casual_rows') or []

                for i in range(max_rows):
                    rr = 2 + i
                    # annual (cols: from=14, to=13, days=12, remaining=11 on first row)
                    a = annual[i] if i < len(annual) else {}
                    _write_cell(t1, rr, 14, a.get('from'))
                    _write_cell(t1, rr, 13, a.get('to'))
                    _write_cell(t1, rr, 12, a.get('days'))
                    if i == 0:
                        _write_cell(t1, rr, 11, rep.get('annual_remaining'))
                    # sick (cols: from=9, to=8, days=7)
                    s = sick[i] if i < len(sick) else {}
                    _write_cell(t1, rr, 9, s.get('from'))
                    _write_cell(t1, rr, 8, s.get('to'))
                    _write_cell(t1, rr, 7, s.get('days'))
                    # casual (cols: from=5, to=3/4, days=2)
                    c = casual[i] if i < len(casual) else {}
                    _write_cell(t1, rr, 5, c.get('from'))
                    _write_cell(t1, rr, 3, c.get('to'))
                    _write_cell(t1, rr, 4, c.get('to'))
                    _write_cell(t1, rr, 2, c.get('days'))

                # No-pay — rows 13..15
                no_pay = rep.get('no_pay_rows') or []
                for i in range(3):
                    rr = 13 + i
                    r0 = no_pay[i] if i < len(no_pay) else {}
                    _write_cell(t1, rr, 5, r0.get('from'))
                    _write_cell(t1, rr, 3, r0.get('to'))
                    _write_cell(t1, rr, 4, r0.get('to'))
                    _write_cell(t1, rr, 2, r0.get('days'))

                # Hajj — row 18 (single range)
                hajj = (rep.get('hajj_rows') or [])
                h0 = hajj[0] if hajj else {}
                _write_cell(t1, 18, 5, h0.get('from'))
                _write_cell(t1, 18, 4, h0.get('from'))
                _write_cell(t1, 18, 3, h0.get('to'))
                _write_cell(t1, 18, 2, h0.get('to'))

                # Maternity — row 21 (single range)
                mat = (rep.get('maternity_rows') or [])
                m0 = mat[0] if mat else {}
                _write_cell(t1, 21, 5, m0.get('from'))
                _write_cell(t1, 21, 4, m0.get('from'))
                _write_cell(t1, 21, 3, m0.get('to'))
                _write_cell(t1, 21, 2, m0.get('to'))

            except Exception:
                pass

            bio = BytesIO()
            doc.save(bio)
            bio.seek(0)
            return bio.getvalue()

        # One employee -> direct DOCX, many -> ZIP of DOCX
        try:
            reps = diwan_reports or []
            if len(reps) == 1:
                rep = reps[0]
                base = (rep.get('employee_no') or rep.get('personal_no') or 'diwan')
                fn = f"diwan_{_safe_part(base)}_{year}.docx"
                return send_file(BytesIO(_build_docx_bytes(rep)),
                                 mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                 as_attachment=True,
                                 download_name=fn)
            else:
                zbio = BytesIO()
                import zipfile as _zipfile
                with _zipfile.ZipFile(zbio, 'w', _zipfile.ZIP_DEFLATED) as zf:
                    for i, rep in enumerate(reps, start=1):
                        base = (rep.get('employee_no') or rep.get('personal_no') or str(i))
                        fn = f"diwan_{_safe_part(base)}_{year}.docx"
                        zf.writestr(fn, _build_docx_bytes(rep))
                zbio.seek(0)
                zip_name = f"diwan_reports_{year}_{month:02d}.zip"
                return send_file(zbio, mimetype='application/zip', as_attachment=True, download_name=zip_name)
        except Exception:
            # If anything goes wrong, fall back to showing the HTML view (no hard crash).
            flash('تعذر إنشاء ملف Word في الوقت الحالي. تأكد من تثبيت python-docx وأن قالب تقرير الديوان موجود.', 'danger')

    template_name = 'portal/hr/reports_diwan_council.html' if view in ('diwan', 'council', 'diwan_report') else 'portal/hr/reports_diwan.html'

    return render_template(
        template_name,
        rows=rows_view,
        diwan_reports=diwan_reports,
        users=users,
        selected_user_id=selected_user_id,
        work_locations=work_locations,
        selected_work_location_id=work_location_id,
        appointment_types=appointment_types,
        selected_appointment_type_id=appointment_type_id,
        year=year,
        month=month,
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
        view=view,
        work_location_name=work_location_name,
        appointment_type_name=appointment_type_name,
        generated_on=datetime.now(),
    )
@portal_bp.route("/hr/coming-soon")
@login_required
def hr_coming_soon():
    """Placeholder pages used from the HR sub-navigation."""
    title = (request.args.get("title") or "قيد التطوير").strip()
    hint = (request.args.get("hint") or "هذه الصفحة قيد التطوير حالياً.").strip()
    return render_template("portal/hr/coming_soon.html", title=title, hint=hint)


@portal_bp.route("/hr/employees/new")
@login_required
@require_permissions(HR_EMP_MANAGE)
def hr_employee_new():
    """HR shortcut for creating a new employee (reuses the main Users module)."""
    return redirect(url_for("users.create_user"))


# -------------------------
# HR - Employee self service
# -------------------------


def _find_direct_manager(u: User):
    """Return the direct manager (or deputy) for a user based on OrgUnitManager.

    Priority: Department → Directorate → Organization.
    """
    try:
        dept = Department.query.get(u.department_id) if getattr(u, "department_id", None) else None
        dir_ = None
        if getattr(u, "directorate_id", None):
            dir_ = Directorate.query.get(u.directorate_id)
        if not dir_ and dept:
            dir_ = dept.directorate
        org = dir_.organization if dir_ else None

        candidates = [
            ("DEPARTMENT", dept.id if dept else None),
            ("DIRECTORATE", dir_.id if dir_ else None),
            ("ORGANIZATION", org.id if org else None),
        ]
        for unit_type, unit_id in candidates:
            if not unit_id:
                continue
            row = OrgUnitManager.query.filter_by(unit_type=unit_type, unit_id=unit_id).first()
            if not row:
                continue
            # Prefer manager then deputy, and avoid returning the same user
            if row.manager_user_id and row.manager_user_id != u.id:
                return row.manager_user
            if row.deputy_user_id and row.deputy_user_id != u.id:
                return row.deputy_user
    except Exception:
        pass
    return None


def _parse_yyyy_mm_dd(s):
    """Parse YYYY-MM-DD into date.

    Accepts: str, date, datetime, None.
    """
    try:
        if s is None:
            return None
        # Already a date/datetime
        from datetime import datetime as _dt
        from datetime import date as _date
        if isinstance(s, _dt):
            return s.date()
        if isinstance(s, _date):
            return s
        s2 = (str(s) or '').strip()
        if not s2:
            return None
        # allow 'YYYY-MM-DD HH:MM:SS'
        if ' ' in s2:
            s2 = s2.split(' ')[0]
        return datetime.strptime(s2, "%Y-%m-%d").date()
    except Exception:
        return None



def _parse_hhmm(s: str):
    s = (s or "").strip()
    if not s:
        return None
    try:
        # accept HH:MM
        datetime.strptime(s, "%H:%M")
        return s
    except Exception:
        return None



# -------------------------
# HR - Pages previously marked as "قيد التطوير" (Phase 1)
# -------------------------

def _hr_can_manage() -> bool:
    try:
        return bool(current_user.has_perm(HR_MASTERDATA_MANAGE) or current_user.has_perm(HR_EMP_MANAGE))
    except Exception:
        return False


def _hr_can_manage_attendance() -> bool:
    try:
        # Managers (HR_REQUESTS_APPROVE) can also manage attendance-related actions
        return bool(
            current_user.has_perm(HR_MASTERDATA_MANAGE)
            or current_user.has_perm(HR_ATT_CREATE)
            or current_user.has_perm(HR_REQUESTS_VIEW_ALL)
            or current_user.has_perm(HR_REQUESTS_APPROVE)
        )
    except Exception:
        return False


def _hr_lookup_items_for_category(category: str):
    """Return active HRLookupItem rows for a given category.

    This helper is used by attendance/duam screens for dropdown filters
    (e.g., WORK_GOVERNORATE, WORK_LOCATION).
    """
    cat = (category or '').strip().upper()
    if not cat:
        return []
    try:
        return (
            HRLookupItem.query
            .filter(HRLookupItem.category == cat)
            .filter(HRLookupItem.is_active.is_(True))
            .order_by(HRLookupItem.sort_order.asc(), HRLookupItem.name_ar.asc())
            .all()
        )
    except Exception:
        return []


def _ym_to_range(year: int, month: int):
    try:
        y = int(year)
        m = int(month)
        if m < 1 or m > 12:
            return None, None
        start = date(y, m, 1)
        if m == 12:
            end = date(y + 1, 1, 1) - timedelta(days=1)
        else:
            end = date(y, m + 1, 1) - timedelta(days=1)
        return start, end
    except Exception:
        return None, None


def _as_yyyy_mm_dd(d: date | None) -> str:
    try:
        return d.strftime('%Y-%m-%d') if d else ''
    except Exception:
        return ''


def _list_hr_users():
    try:
        return User.query.order_by(func.coalesce(User.name, User.email).asc()).all()
    except Exception:
        return []


@portal_bp.route('/hr/employees/followup')
@login_required
@_perm_any(HR_EMP_READ, HR_EMP_MANAGE, HR_REQUESTS_VIEW_ALL)
def hr_employee_followup():
    """متابعة الإدخال: إظهار الموظفين الذين ينقصهم حقول أساسية في ملف الموظف."""
    q = (request.args.get('q') or '').strip()
    only_missing = (request.args.get('only_missing') or '1').strip()  # default yes

    page = int(request.args.get('page') or 1)
    per_page = 50

    # Outer join employee_file for missing checks
    base = db.session.query(User, EmployeeFile).outerjoin(EmployeeFile, EmployeeFile.user_id == User.id)

    # Best-effort: exclude system/service accounts
    try:
        base = base.filter((func.lower(User.email).notlike('%@example.com')))
    except Exception:
        pass

    if q:
        base = base.filter(or_(User.name.ilike(f'%{q}%'), User.email.ilike(f'%{q}%')))

    rows = []
    for u, ef in base.order_by(func.coalesce(User.name, User.email).asc()).all():
        missing = []
        # employee file
        if not ef or not (ef.employee_no or '').strip():
            missing.append('الرقم الوظيفي')
        if not ef or not (ef.timeclock_code or '').strip():
            missing.append('كود الساعة')
        if not ef or not (ef.national_id or '').strip():
            missing.append('رقم الهوية')
        if not ef or not (ef.phone or '').strip():
            missing.append('الهاتف')
        if not (u.job_title or '').strip() and (not ef or not (ef.job_title or '').strip()):
            missing.append('المسمى الوظيفي')

        if only_missing == '1' and not missing:
            continue

        rows.append({'user': u, 'ef': ef, 'missing': missing})

    # manual pagination (rows already materialized to compute missing)
    total = len(rows)
    start_i = (page - 1) * per_page
    end_i = start_i + per_page
    page_rows = rows[start_i:end_i]

    pages = max(1, (total + per_page - 1) // per_page)

    return render_template(
        'portal/hr/employees_followup.html',
        q=q,
        only_missing=only_missing,
        page=page,
        pages=pages,
        total=total,
        rows=page_rows,
    )


@portal_bp.route('/hr/attendance/special')
@login_required
@_perm_any(HR_ATT_READ, HR_REQUESTS_VIEW_ALL, HR_READ)
def hr_att_special_log():
    """سجل الحالات الخاصة (حالة/استثناء) - مطابق للشاشات المرفقة."""
    kind = (request.args.get('kind') or '').strip().upper()
    date_from = (request.args.get('date_from') or '').strip()
    date_to = (request.args.get('date_to') or '').strip()
    target_kind = (request.args.get('target_kind') or 'USER').strip().upper()
    user_id = (request.args.get('user_id') or '').strip()
    status_filter = (request.args.get('status') or '').strip().upper()
    exception_filter = (request.args.get('exception') or '').strip().upper()

    q = HRAttendanceSpecialCase.query

    if kind in ('STATUS', 'EXCEPTION'):
        q = q.filter(HRAttendanceSpecialCase.kind == kind)

    # date overlap filtering (string compare works for YYYY-MM-DD)
    if date_from and len(date_from) == 10:
        q = q.filter(or_(HRAttendanceSpecialCase.day_to == None, HRAttendanceSpecialCase.day_to >= date_from))
    if date_to and len(date_to) == 10:
        q = q.filter(HRAttendanceSpecialCase.day <= date_to)

    if target_kind:
        q = q.filter(or_(HRAttendanceSpecialCase.target_kind == None, HRAttendanceSpecialCase.target_kind == target_kind))

    if user_id:
        try:
            q = q.filter(HRAttendanceSpecialCase.user_id == int(user_id))
        except Exception:
            pass

    if kind == 'STATUS' and status_filter:
        q = q.filter(HRAttendanceSpecialCase.status == status_filter)
    if kind == 'EXCEPTION' and exception_filter:
        q = q.filter(HRAttendanceSpecialCase.field == exception_filter)

    rows = q.order_by(HRAttendanceSpecialCase.created_at.desc()).limit(500).all()

    return render_template(
        'portal/hr/att_special_log.html',
        rows=rows,
        kind=kind,
        date_from=date_from,
        date_to=date_to,
        target_kind=target_kind,
        user_id=user_id,
        status_filter=status_filter,
        exception_filter=exception_filter,
        users=_list_hr_users(),
        can_manage=_hr_can_manage_attendance(),
    )


@portal_bp.route('/hr/attendance/special/status/new', methods=['GET', 'POST'])
@login_required
@_perm_any(HR_ATT_READ, HR_REQUESTS_VIEW_ALL, HR_READ)
def hr_att_special_status_new():
    """إدخال/تعديل حالة خاصة."""
    if request.method == 'POST':
        if not _hr_can_manage_attendance():
            abort(403)

        uid = (request.form.get('user_id') or '').strip()
        target_kind = (request.form.get('target_kind') or 'USER').strip().upper()
        status_type = (request.form.get('status_type') or '').strip().upper()
        day_from = (request.form.get('day_from') or '').strip()
        day_to = (request.form.get('day_to') or '').strip()

        start_time = _parse_hhmm(request.form.get('start_time') or '')
        end_time = _parse_hhmm(request.form.get('end_time') or '')

        def _int_or_none(name:str):
            v = (request.form.get(name) or '').strip()
            if v == '':
                return None
            try:
                return int(v)
            except Exception:
                return None

        allow_morning = _int_or_none('allow_morning_minutes')
        allow_evening = _int_or_none('allow_evening_minutes')
        note = (request.form.get('note') or '').strip()
        action = (request.form.get('action') or 'save').strip().lower()

        try:
            uid_int = int(uid)
        except Exception:
            flash('اختر موظفاً بشكل صحيح.', 'danger')
            return redirect(url_for('portal.hr_att_special_status_new'))

        d1 = _parse_yyyy_mm_dd(day_from)
        d2 = _parse_yyyy_mm_dd(day_to) if day_to else d1
        if not d1 or not d2 or d2 < d1:
            flash('نطاق التاريخ غير صحيح.', 'danger')
            return redirect(url_for('portal.hr_att_special_status_new'))

        if not status_type:
            flash('اختر نوع الحالة.', 'danger')
            return redirect(url_for('portal.hr_att_special_status_new'))

        row = HRAttendanceSpecialCase(
            user_id=uid_int,
            target_kind=target_kind or 'USER',
            day=day_from,
            day_to=(day_to or day_from),
            kind='STATUS',
            status=status_type,
            start_time=start_time or None,
            end_time=end_time or None,
            allow_morning_minutes=allow_morning,
            allow_evening_minutes=allow_evening,
            note=note or None,
            applied=True,
            created_by_id=getattr(current_user, 'id', None),
        )
        db.session.add(row)
        db.session.commit()

        flash('تم حفظ الحالة الخاصة.', 'success')
        if action == 'calc':
            flash('ملاحظة: احتساب الدوام سيتم تطبيقه عند بناء منطق الاحتساب. تم حفظ السجل الآن.', 'info')

        return redirect(url_for('portal.hr_att_special_log', kind='STATUS', date_from=day_from[:10], date_to=(day_to or day_from)[:10], user_id=uid_int))

    return render_template(
        'portal/hr/att_special_status_new.html',
        users=_list_hr_users(),
        today=_as_yyyy_mm_dd(date.today()),
    )


@portal_bp.route('/hr/attendance/special/exception/new', methods=['GET', 'POST'])
@login_required
@_perm_any(HR_ATT_READ, HR_REQUESTS_VIEW_ALL, HR_READ)
def hr_att_special_exception_new():
    """إدخال/تعديل استثناء."""
    if request.method == 'POST':
        if not _hr_can_manage_attendance():
            abort(403)

        uid = (request.form.get('user_id') or '').strip()
        target_kind = (request.form.get('target_kind') or 'USER').strip().upper()
        exc_type = (request.form.get('exception_type') or '').strip().upper()
        day_from = (request.form.get('day_from') or '').strip()
        day_to = (request.form.get('day_to') or '').strip()
        note = (request.form.get('note') or '').strip()

        try:
            uid_int = int(uid)
        except Exception:
            flash('اختر موظفاً بشكل صحيح.', 'danger')
            return redirect(url_for('portal.hr_att_special_exception_new'))

        d1 = _parse_yyyy_mm_dd(day_from)
        d2 = _parse_yyyy_mm_dd(day_to) if day_to else d1
        if not d1 or not d2 or d2 < d1:
            flash('نطاق التاريخ غير صحيح.', 'danger')
            return redirect(url_for('portal.hr_att_special_exception_new'))

        if not exc_type:
            flash('اختر نوع الاستثناء.', 'danger')
            return redirect(url_for('portal.hr_att_special_exception_new'))

        row = HRAttendanceSpecialCase(
            user_id=uid_int,
            target_kind=target_kind or 'USER',
            day=day_from,
            day_to=(day_to or day_from),
            kind='EXCEPTION',
            field=exc_type,
            note=note or None,
            applied=True,
            created_by_id=getattr(current_user, 'id', None),
        )
        db.session.add(row)
        db.session.commit()

        flash('تم حفظ الاستثناء.', 'success')
        return redirect(url_for('portal.hr_att_special_log', kind='EXCEPTION', date_from=day_from[:10], date_to=(day_to or day_from)[:10], user_id=uid_int))

    return render_template(
        'portal/hr/att_special_exception_new.html',
        users=_list_hr_users(),
        today=_as_yyyy_mm_dd(date.today()),
    )


@portal_bp.route('/hr/attendance/special/<int:row_id>/delete', methods=['POST'])
@login_required
@_perm_any(HR_ATT_READ, HR_REQUESTS_VIEW_ALL, HR_READ)
def hr_att_special_delete(row_id: int):
    """حذف سجل حالة/استثناء من سجل الحالات الخاصة."""
    if not _hr_can_manage_attendance():
        abort(403)

    row = HRAttendanceSpecialCase.query.get_or_404(int(row_id))
    kind = (row.kind or '').strip().upper()

    db.session.delete(row)
    db.session.commit()

    flash('تم حذف السجل.', 'success')

    # prefer returning to the same filtered listing page
    ref = request.referrer or ''
    if '/portal/hr/attendance/special' in ref:
        return redirect(ref)

    return redirect(url_for('portal.hr_att_special_log', kind=kind or None))


@portal_bp.route('/hr/attendance/closing', methods=['GET', 'POST'])
@login_required
@_perm_any(HR_ATT_READ, HR_REQUESTS_VIEW_ALL, HR_READ)
def hr_att_mass_closing():
    """الإغلاق الجماعي لحركات مفتوحة (مغادرات + الدوام).

    ملاحظة: حالياً نسجّل الإغلاق كسجل إداري (يُستخدم لاحقاً في استثناءات الخصم/الحسابات).
    """
    govs = _hr_lookup_items_for_category('WORK_GOVERNORATE')
    locs = _hr_lookup_items_for_category('WORK_LOCATION')

    if request.method == 'POST':
        if not _hr_can_manage_attendance():
            abort(403)

        day = (request.form.get('day') or '').strip()
        d = _parse_yyyy_mm_dd(day)
        if not d:
            flash('التاريخ غير صحيح.', 'danger')
            return redirect(url_for('portal.hr_att_mass_closing'))

        def _int(name:str):
            v = (request.form.get(name) or '').strip()
            if v == '':
                return None
            try:
                return int(v)
            except Exception:
                return None

        gov_id = _int('work_governorate_lookup_id')
        loc_id = _int('work_location_lookup_id')

        close_attendance = (request.form.get('close_attendance') == '1')
        close_permissions = (request.form.get('close_permissions') == '1')
        if not close_attendance and not close_permissions:
            # default to both if user didn't check
            close_attendance = True
            close_permissions = True

        note = (request.form.get('note') or '').strip()

        row = HRAttendanceClosing(
            day_from=day,
            day_to=day,
            work_governorate_lookup_id=gov_id,
            work_location_lookup_id=loc_id,
            close_attendance=close_attendance,
            close_permissions=close_permissions,
            note=note or None,
            created_by_id=getattr(current_user, 'id', None),
        )
        db.session.add(row)
        db.session.commit()
        flash('تم تنفيذ الإغلاق الجماعي وحفظ السجل.', 'success')
        return redirect(url_for('portal.hr_att_mass_closing'))

    # GET: list the latest closings
    rows = HRAttendanceClosing.query.order_by(HRAttendanceClosing.created_at.desc()).limit(200).all()
    return render_template(
        'portal/hr/att_closing.html',
        rows=rows,
        govs=govs,
        locs=locs,
        today=_as_yyyy_mm_dd(date.today()),
        can_manage=_hr_can_manage_attendance(),
    )


@portal_bp.route('/hr/attendance/closing/<int:row_id>/delete', methods=['POST'])
@login_required
@_perm_any(HR_ATT_READ, HR_REQUESTS_VIEW_ALL, HR_READ)
def hr_att_closing_delete(row_id: int):
    if not _hr_can_manage_attendance():
        abort(403)
    row = HRAttendanceClosing.query.get_or_404(row_id)
    db.session.delete(row)
    db.session.commit()
    flash('تم حذف فترة الإغلاق.', 'success')
    return redirect(url_for('portal.hr_att_mass_closing'))


@portal_bp.route('/hr/deductions/settings', methods=['GET', 'POST'])
@login_required
@_perm_any(HR_ATT_READ, HR_REPORTS_VIEW, HR_READ, HR_REQUESTS_VIEW_ALL, HR_REQUESTS_APPROVE)
def hr_deductions_settings():
    """إعدادات الخصم (مطابقة للشاشات)."""
    cfg = HRAttendanceDeductionConfig.query.order_by(HRAttendanceDeductionConfig.updated_at.desc()).first()

    style_opts = [
        ('AGGREGATE', 'احتساب يوم خصم عن كل 7 ساعات من كافة الأنواع مجتمعة'),
        ('PER_CATEGORY', 'احتساب يوم خصم عن كل 7 ساعات لكل تصنيف على حدة'),
    ]
    source_opts = [('LEAVE', 'من الإجازات'), ('SALARY', 'من الراتب')]
    carry_opts = [('CARRY_TO_NEXT', 'تحويل من الشهر الحالي إلى القادم'), ('WITHIN_MONTH', 'ضمن نفس الشهر فقط')]

    if request.method == 'POST':
        if not _hr_can_manage():
            abort(403)

        def _f(name: str):
            v = (request.form.get(name) or '').strip()
            if v == '':
                return None
            try:
                return float(v)
            except Exception:
                return None

        # legacy values (we keep them for backward compatibility)
        late_v = _f('late_minute_value')
        early_v = _f('early_minute_value')
        if early_v is None:
            early_v = _f('early_leave_minute_value')
        absent_v = _f('absent_day_value')

        def _pick(name: str, default: str, allowed: set[str]):
            v = (request.form.get(name) or '').strip().upper()
            return v if v in allowed else default

        deduction_style = _pick('deduction_style', 'AGGREGATE', {'AGGREGATE', 'PER_CATEGORY'})
        late_source = _pick('late_source', 'SALARY', {'LEAVE', 'SALARY'})
        early_source = _pick('early_source', 'SALARY', {'LEAVE', 'SALARY'})
        special_permission_source = _pick('special_permission_source', 'SALARY', {'LEAVE', 'SALARY'})
        unauthorized_permission_source = _pick('unauthorized_permission_source', 'SALARY', {'LEAVE', 'SALARY'})
        carry_method = _pick('carry_method', 'CARRY_TO_NEXT', {'CARRY_TO_NEXT', 'WITHIN_MONTH'})

        hours_per_day = _f('hours_per_day') or 7.0
        if hours_per_day <= 0:
            hours_per_day = 7.0

        if not cfg:
            cfg = HRAttendanceDeductionConfig()
            db.session.add(cfg)

        cfg.late_minute_value = late_v
        cfg.early_minute_value = early_v
        cfg.absent_day_value = absent_v

        cfg.deduction_style = deduction_style
        cfg.late_source = late_source
        cfg.early_source = early_source
        cfg.special_permission_source = special_permission_source
        cfg.unauthorized_permission_source = unauthorized_permission_source
        cfg.carry_method = carry_method
        cfg.hours_per_day = hours_per_day

        cfg.updated_at = datetime.utcnow()
        db.session.commit()

        flash('تم حفظ إعدادات الخصم.', 'success')
        return redirect(url_for('portal.hr_deductions_settings'))

    return render_template(
        'portal/hr/deductions_settings.html',
        cfg=cfg,
        style_opts=style_opts,
        source_opts=source_opts,
        carry_opts=carry_opts,
        can_manage=_hr_can_manage(),
    )


@portal_bp.route('/hr/deductions/run', methods=['GET', 'POST'])
@login_required
@_perm_any(HR_ATT_READ, HR_REPORTS_VIEW, HR_READ, HR_REQUESTS_VIEW_ALL, HR_REQUESTS_APPROVE)
def hr_deductions_run():
    """تنفيذ الخصم (مع تعيين الموظفين من نفس الشاشة)."""
    cfg = HRAttendanceDeductionConfig.query.order_by(HRAttendanceDeductionConfig.updated_at.desc()).first()
    if not cfg:
        # Create default config if missing
        cfg = HRAttendanceDeductionConfig(
            late_minute_value=0.0,
            early_minute_value=0.0,
            absent_day_value=0.0,
            deduction_style='AGGREGATE',
            late_source='SALARY',
            early_source='SALARY',
            special_permission_source='SALARY',
            unauthorized_permission_source='SALARY',
            carry_method='CARRY_TO_NEXT',
            hours_per_day=7.0,
        )
        db.session.add(cfg)
        db.session.commit()

    # Filters for employee selection
    govs = _hr_lookup_items_for_category('WORK_GOVERNORATE')
    locs = _hr_lookup_items_for_category('WORK_LOCATION')
    gov_id = (request.values.get('work_governorate_lookup_id') or '').strip()
    loc_id = (request.values.get('work_location_lookup_id') or '').strip()
    qtxt = (request.values.get('q') or '').strip()

    year = int((request.values.get('year') or date.today().year) or date.today().year)
    month = int((request.values.get('month') or date.today().month) or date.today().month)

    # NOTE: User.is_active is a Flask-Login property (not a DB column) in this project,
    # so it cannot be used in SQLAlchemy filters.
    users_q = User.query.join(EmployeeFile, EmployeeFile.user_id == User.id)
    if gov_id.isdigit():
        users_q = users_q.filter(EmployeeFile.work_governorate_lookup_id == int(gov_id))
    if loc_id.isdigit():
        users_q = users_q.filter(EmployeeFile.work_location_lookup_id == int(loc_id))
    if qtxt:
        like = f"%{qtxt}%"
        # EmployeeFile uses employee_no (not emp_no)
        users_q = users_q.filter(or_(User.name.ilike(like), User.email.ilike(like), EmployeeFile.employee_no.ilike(like)))

    users = users_q.order_by(User.name.asc()).limit(500).all()

    if request.method == 'POST':
        if not _hr_can_manage_attendance():
            abort(403)

        selected_ids = request.form.getlist('user_ids')
        selected_ids = [int(x) for x in selected_ids if str(x).isdigit()]
        if not selected_ids:
            flash('اختر موظفاً واحداً على الأقل لتنفيذ الخصم.', 'danger')
            return redirect(url_for('portal.hr_deductions_run', year=year, month=month, work_governorate_lookup_id=gov_id, work_location_lookup_id=loc_id, q=qtxt))

        # Aggregate attendance minutes for the selected employees
        agg = db.session.query(
            AttendanceDailySummary.user_id,
            func.coalesce(func.sum(AttendanceDailySummary.late_minutes), 0),
            func.coalesce(func.sum(AttendanceDailySummary.early_leave_minutes), 0),
            func.coalesce(func.sum(AttendanceDailySummary.absent), 0),
        ).filter(
            AttendanceDailySummary.year == year,
            AttendanceDailySummary.month == month,
            AttendanceDailySummary.user_id.in_(selected_ids),
        ).group_by(AttendanceDailySummary.user_id).all()

        minutes_per_day = int(round((cfg.hours_per_day or 7.0) * 60))
        if minutes_per_day <= 0:
            minutes_per_day = 420

        run = HRAttendanceDeductionRun(
            year=year,
            month=month,
            config_snapshot_json=json.dumps({
                'deduction_style': cfg.deduction_style,
                'hours_per_day': cfg.hours_per_day,
                'late_source': cfg.late_source,
                'early_source': cfg.early_source,
                'special_permission_source': cfg.special_permission_source,
                'unauthorized_permission_source': cfg.unauthorized_permission_source,
                'carry_method': cfg.carry_method,
            }, ensure_ascii=False),
            created_by_id=getattr(current_user, 'id', None),
        )
        db.session.add(run)
        db.session.flush()

        items_payload = []
        total_days = 0

        for uid, late_min, early_min, absent_days in agg:
            late_min = int(late_min or 0)
            early_min = int(early_min or 0)
            absent_days = int(absent_days or 0)

            if (cfg.deduction_style or 'AGGREGATE').upper() == 'PER_CATEGORY':
                days_from_minutes = (late_min // minutes_per_day) + (early_min // minutes_per_day)
                remainder_minutes = (late_min % minutes_per_day) + (early_min % minutes_per_day)
            else:
                total_minutes = late_min + early_min
                days_from_minutes = total_minutes // minutes_per_day
                remainder_minutes = total_minutes % minutes_per_day

            if (cfg.carry_method or 'CARRY_TO_NEXT').upper() == 'WITHIN_MONTH':
                # within month only: we ignore remainder in totals
                remainder_minutes = 0

            days_total = float(days_from_minutes + absent_days)
            total_days += int(days_from_minutes) + absent_days

            item = HRAttendanceDeductionItem(
                run_id=run.id,
                user_id=uid,
                late_minutes=late_min,
                early_leave_minutes=early_min,
                absent_days=absent_days,
                amount=days_total,  # amount = days (for now)
                note=f"دقائق متبقية: {remainder_minutes}",
            )
            db.session.add(item)
            items_payload.append({
                'user_id': uid,
                'late_minutes': late_min,
                'early_leave_minutes': early_min,
                'absent_days': absent_days,
                'deduction_days': days_total,
                'remainder_minutes': remainder_minutes,
            })

        run.totals_json = json.dumps({
            'employees_count': len(items_payload),
            'total_deduction_days': total_days,
            'minutes_per_day': minutes_per_day,
            'items': items_payload,
        }, ensure_ascii=False)
        db.session.commit()

        flash('تم تنفيذ الخصم للموظفين المحددين.', 'success')
        return redirect(url_for('portal.hr_deductions_view', run_id=run.id))

    return render_template(
        'portal/hr/deductions_run.html',
        cfg=cfg,
        year=year,
        month=month,
        govs=govs,
        locs=locs,
        work_governorate_lookup_id=gov_id,
        work_location_lookup_id=loc_id,
        q=qtxt,
        users=users,
    )


@portal_bp.route('/hr/deductions/<int:run_id>')
@login_required
@_perm_any(HR_ATT_READ, HR_REPORTS_VIEW, HR_READ, HR_REQUESTS_VIEW_ALL, HR_REQUESTS_APPROVE)
def hr_deductions_view(run_id: int):
    run = HRAttendanceDeductionRun.query.get_or_404(run_id)

    export = (request.args.get('export') or '').strip().lower()
    if export in ('1', 'csv'):
        # CSV export
        import csv, io
        buf = io.StringIO()
        w = csv.writer(buf)
        w.writerow(['Employee', 'Email', 'Late minutes', 'Early minutes', 'Absent days', 'Amount'])
        for it in run.items:
            u = it.user
            w.writerow([
                (getattr(u, 'name', '') or '').strip(),
                (getattr(u, 'email', '') or '').strip(),
                it.late_minutes,
                it.early_leave_minutes,
                it.absent_days,
                it.amount,
            ])
        data = buf.getvalue().encode('utf-8-sig')
        return send_file(
            io.BytesIO(data),
            mimetype='text/csv',
            as_attachment=True,
            download_name=f"deductions_{run.year:04d}_{run.month:02d}.csv",
        )

    return render_template('portal/hr/deductions_view.html', run=run)


@portal_bp.route('/hr/deductions/yearly')
@login_required
@_perm_any(HR_ATT_READ, HR_REPORTS_VIEW, HR_READ, HR_REQUESTS_VIEW_ALL, HR_REQUESTS_APPROVE)
def hr_deductions_yearly():
    year = int(request.args.get('year') or date.today().year)
    # Map month -> latest run
    runs = HRAttendanceDeductionRun.query.filter_by(year=year).order_by(HRAttendanceDeductionRun.month.asc(), HRAttendanceDeductionRun.created_at.desc()).all()
    latest_by_month = {}
    for r in runs:
        if r.month not in latest_by_month:
            latest_by_month[r.month] = r
    return render_template('portal/hr/deductions_yearly.html', year=year, latest_by_month=latest_by_month)


@portal_bp.route('/hr/deductions/log')
@login_required
@_perm_any(HR_ATT_READ, HR_REPORTS_VIEW, HR_READ, HR_REQUESTS_VIEW_ALL, HR_REQUESTS_APPROVE)
def hr_deductions_log():
    """سجل الخصومات.

    تم إنشاء هذا المسار لأن القالب موجود (`deductions_log.html`) وكان
    ضمن قائمة الموارد البشرية، لكن الـ endpoint لم يكن مُعرّفاً مما تسبب
    في BuildError عند فتح /portal/hr.
    """
    year = (request.args.get('year') or '').strip()
    month = (request.args.get('month') or '').strip()
    status = (request.args.get('status') or '').strip().upper()

    q = HRAttendanceDeductionRun.query
    if year.isdigit():
        q = q.filter(HRAttendanceDeductionRun.year == int(year))
    if month.isdigit():
        q = q.filter(HRAttendanceDeductionRun.month == int(month))
    if status in ('DRAFT', 'FINAL'):
        q = q.filter(HRAttendanceDeductionRun.status == status)

    rows = q.order_by(HRAttendanceDeductionRun.created_at.desc()).limit(300).all()
    return render_template('portal/hr/deductions_log.html', rows=rows)


@portal_bp.route('/hr/attendance/monthly-schedule/new', methods=['GET', 'POST'])
@login_required
@_perm_any(HR_MASTERDATA_MANAGE, HR_REQUESTS_VIEW_ALL, HR_READ)
def hr_monthly_schedule_new():
    """إضافة جدول شهري: واجهة مبسطة لإنشاء WorkAssignment بفترة شهر."""
    schedules = WorkSchedule.query.order_by(WorkSchedule.name.asc()).all()
    departments = Department.query.order_by(Department.name_ar.asc()).all()
    users = _list_hr_users()

    if request.method == 'POST':
        if not _hr_can_manage():
            abort(403)

        name = (request.form.get('name') or '').strip() or 'جدول شهري'
        schedule_id = request.form.get('schedule_id')
        day_from = (request.form.get('day_from') or '').strip()
        day_to = (request.form.get('day_to') or '').strip()
        target_user_id = (request.form.get('target_user_id') or '').strip()
        target_role = (request.form.get('target_role') or '').strip()
        target_dept_id = (request.form.get('target_dept_id') or '').strip()

        d1 = _parse_yyyy_mm_dd(day_from)
        d2 = _parse_yyyy_mm_dd(day_to)
        if not d1 or not d2 or d2 < d1:
            flash('نطاق التاريخ غير صحيح.', 'danger')
            return redirect(url_for('portal.hr_monthly_schedule_new'))

        try:
            schedule_id_int = int(schedule_id)
        except Exception:
            flash('اختر قالب دوام.', 'danger')
            return redirect(url_for('portal.hr_monthly_schedule_new'))

        # choose one target
        target_type = None
        target_value = None
        if target_user_id:
            target_type = 'USER'
            try:
                target_value = int(target_user_id)
            except Exception:
                target_value = None
        elif target_dept_id:
            target_type = 'DEPARTMENT'
            try:
                target_value = int(target_dept_id)
            except Exception:
                target_value = None
        elif target_role:
            target_type = 'ROLE'
            target_value = target_role

        if not target_type or target_value in (None, ''):
            flash('اختر مستهدف الجدول (موظف/قسم/دور).', 'danger')
            return redirect(url_for('portal.hr_monthly_schedule_new'))

        wa = WorkAssignment(
            name=name,
            schedule_id=schedule_id_int,
            start_date=day_from,
            end_date=day_to,
            target_type=target_type,
            target_value=str(target_value),
            created_by_id=getattr(current_user, 'id', None),
        )
        db.session.add(wa)
        db.session.commit()
        flash('تم إنشاء الجدول الشهري.', 'success')
        return redirect(url_for('portal.hr_monthly_schedule_log'))

    return render_template(
        'portal/hr/monthly_schedule_new.html',
        schedules=schedules,
        departments=departments,
        users=users,
        today=_as_yyyy_mm_dd(date.today()),
    )


@portal_bp.route('/hr/attendance/monthly-schedule')
@login_required
@_perm_any(HR_MASTERDATA_MANAGE, HR_REQUESTS_VIEW_ALL, HR_READ)
def hr_monthly_schedule_log():
    rows = WorkAssignment.query.order_by(WorkAssignment.id.desc()).limit(200).all()
    return render_template('portal/hr/monthly_schedule_log.html', rows=rows)


@portal_bp.route('/hr/missions')
@login_required
@_perm_any(HR_READ, HR_REQUESTS_VIEW_ALL, HR_MASTERDATA_MANAGE)
def hr_official_missions():
    missions = HROfficialMission.query.order_by(HROfficialMission.start_day.desc()).limit(200).all()
    return render_template('portal/hr/missions_log.html', missions=missions, can_manage=_hr_can_manage())


@portal_bp.route('/hr/holidays/weekly', methods=['GET', 'POST'])
@login_required
@_perm_any(HR_READ, HR_REQUESTS_VIEW_ALL, HR_MASTERDATA_MANAGE)
def hr_weekly_holidays():
    """العطل الأسبوعية: حفظ إعداد على شكل bitmask في SystemSetting."""
    key = 'HR_WEEKLY_HOLIDAYS_MASK'
    row = SystemSetting.query.filter_by(key=key).first()
    mask = 0
    try:
        mask = int((row.value or '0').strip()) if row else 0
    except Exception:
        mask = 0

    if request.method == 'POST':
        if not _hr_can_manage():
            abort(403)
        new_mask = 0
        for i in range(7):
            if (request.form.get(f'd{i}') or '') == '1':
                new_mask |= (1 << i)
        if not row:
            row = SystemSetting(key=key, value=str(new_mask))
            db.session.add(row)
        else:
            row.value = str(new_mask)
        db.session.commit()
        flash('تم حفظ العطل الأسبوعية.', 'success')
        return redirect(url_for('portal.hr_weekly_holidays'))

    days = [
        ('الاثنين', 0),
        ('الثلاثاء', 1),
        ('الأربعاء', 2),
        ('الخميس', 3),
        ('الجمعة', 4),
        ('السبت', 5),
        ('الأحد', 6),
    ]

    checked = {idx: bool(mask & (1 << idx)) for _, idx in days}

    return render_template('portal/hr/weekly_holidays.html', days=days, checked=checked, can_manage=_hr_can_manage())


@portal_bp.route('/hr/rooms')
@login_required
@_perm_any(HR_READ, HR_REQUESTS_VIEW_ALL, HR_MASTERDATA_MANAGE)
def hr_room_bookings():
    day = (request.args.get('day') or _as_yyyy_mm_dd(date.today())).strip()
    rooms = HRRoom.query.filter_by(is_active=True).order_by(HRRoom.name.asc()).all()
    bookings = HRRoomBooking.query.filter_by(day=day).order_by(HRRoomBooking.start_time.asc()).all()
    return render_template('portal/hr/rooms_bookings.html', rooms=rooms, bookings=bookings, day=day, can_manage=_hr_can_manage())


@portal_bp.route('/hr/rooms/new', methods=['GET', 'POST'])
@login_required
@_perm_any(HR_READ, HR_REQUESTS_VIEW_ALL, HR_MASTERDATA_MANAGE)
def hr_room_booking_new():
    rooms = HRRoom.query.filter_by(is_active=True).order_by(HRRoom.name.asc()).all()
    if request.method == 'POST':
        # booking is allowed for HR users; management required only for rooms management
        room_id = request.form.get('room_id')
        title = (request.form.get('title') or '').strip()
        day = (request.form.get('day') or '').strip()
        start_time = _parse_hhmm(request.form.get('start_time') or '')
        end_time = _parse_hhmm(request.form.get('end_time') or '')
        note = (request.form.get('note') or '').strip()

        if not title or not _parse_yyyy_mm_dd(day) or not room_id:
            flash('القاعة/العنوان/التاريخ مطلوبة.', 'danger')
            return redirect(url_for('portal.hr_room_booking_new'))
        try:
            room_id = int(room_id)
        except Exception:
            flash('اختر قاعة صحيحة.', 'danger')
            return redirect(url_for('portal.hr_room_booking_new'))

        row = HRRoomBooking(
            room_id=room_id,
            booked_by_id=getattr(current_user, 'id', None),
            title=title,
            day=day,
            start_time=start_time,
            end_time=end_time,
            note=note or None,
        )
        db.session.add(row)
        db.session.commit()
        flash('تم الحجز.', 'success')
        return redirect(url_for('portal.hr_room_bookings', day=day))

    return render_template('portal/hr/rooms_booking_new.html', rooms=rooms, today=_as_yyyy_mm_dd(date.today()))


@portal_bp.route('/hr/rooms/manage', methods=['GET', 'POST'])
@login_required
@_perm_any(HR_MASTERDATA_MANAGE, HR_REQUESTS_VIEW_ALL, HR_READ)
def hr_rooms_manage():
    if request.method == 'POST':
        if not _hr_can_manage():
            abort(403)
        name = (request.form.get('name') or '').strip()
        location = (request.form.get('location') or '').strip()
        capacity = (request.form.get('capacity') or '').strip()
        if not name:
            flash('اسم القاعة مطلوب.', 'danger')
            return redirect(url_for('portal.hr_rooms_manage'))
        try:
            cap = int(capacity) if capacity else None
        except Exception:
            cap = None
        room = HRRoom(name=name, location=location or None, capacity=cap)
        db.session.add(room)
        db.session.commit()
        flash('تم إضافة القاعة.', 'success')
        return redirect(url_for('portal.hr_rooms_manage'))

    rooms = HRRoom.query.order_by(HRRoom.is_active.desc(), HRRoom.name.asc()).all()
    return render_template('portal/hr/rooms_manage.html', rooms=rooms, can_manage=_hr_can_manage())


@portal_bp.route('/hr/rooms/manage/<int:room_id>/toggle', methods=['POST'])
@login_required
@_perm_any(HR_MASTERDATA_MANAGE, HR_REQUESTS_VIEW_ALL, HR_READ)
def hr_rooms_toggle(room_id: int):
    if not _hr_can_manage():
        abort(403)
    room = HRRoom.query.get_or_404(room_id)
    room.is_active = not bool(room.is_active)
    db.session.commit()
    flash('تم تحديث حالة القاعة.', 'success')
    return redirect(url_for('portal.hr_rooms_manage'))


@portal_bp.route('/hr/system-screens')
@login_required
@_perm_any(HR_READ, HR_REQUESTS_VIEW_ALL, HR_MASTERDATA_MANAGE)
def hr_system_screens():
    """شاشات النظام: صفحة روابط تشخيصية/مختصرة."""
    return render_template('portal/hr/system_screens.html')

@portal_bp.route("/hr/me")
@login_required
@_perm(PORTAL_READ)
def hr_me_home():
    """Employee HR dashboard (simple, one page)."""
    # This page itself is gated by PORTAL_READ, but each section is gated by its own permission.
    can_hr = False
    can_att = False
    can_ss = False
    can_docs = False
    can_training = False
    can_sys_eval = False
    try:
        can_hr = current_user.has_perm(HR_READ)
        can_att = current_user.has_perm(HR_ATT_READ)
        can_ss = current_user.has_perm(HR_SS_READ)
        can_docs = current_user.has_perm(HR_DOCS_READ)
        can_training = bool(current_user.has_perm(HR_READ) or current_user.has_perm(HR_SS_READ) or current_user.has_perm(HR_SS_CREATE))
        can_sys_eval = current_user.has_perm(HR_SYSTEM_EVALUATION_VIEW)
    except Exception:
        pass

    # Quick stats (best-effort; never break the dashboard)
    uid = current_user.id

    # Payslips
    payslips_count = 0
    payslips_latest_label = ""
    if can_hr:
        try:
            _ensure_employee_attachment_payslip_schema()
            payslips_count = (
                EmployeeAttachment.query
                .filter_by(user_id=uid, attachment_type="PAYSLIP")
                .filter(_payslip_visible_expr())
                .count()
            )
            latest = (
                EmployeeAttachment.query
                .filter_by(user_id=uid, attachment_type="PAYSLIP")
                .filter(_payslip_visible_expr())
                .order_by(
                    func.coalesce(EmployeeAttachment.payslip_year, 0).desc(),
                    func.coalesce(EmployeeAttachment.payslip_month, 0).desc(),
                    EmployeeAttachment.uploaded_at.desc(),
                )
                .first()
            )
            if latest and getattr(latest, 'payslip_year', None) and getattr(latest, 'payslip_month', None):
                payslips_latest_label = f"{int(latest.payslip_month):02d}/{int(latest.payslip_year)}"
        except Exception:
            payslips_count = 0

    # Leaves & Permissions (requests)
    leaves_total = leaves_pending = 0
    permissions_total = permissions_pending = 0
    if can_hr:
        try:
            leaves_total = HrLeaveRequest.query.filter_by(user_id=uid).count()
            leaves_pending = HrLeaveRequest.query.filter_by(user_id=uid, status="PENDING").count()
        except Exception:
            pass
        try:
            permissions_total = HrPermissionRequest.query.filter_by(user_id=uid).count()
            permissions_pending = HrPermissionRequest.query.filter_by(user_id=uid, status="PENDING").count()
        except Exception:
            pass

    # Trainings
    trainings_total = 0
    trainings_upcoming = 0
    if can_training:
        try:
            trainings_total = HRTrainingEnrollment.query.filter_by(user_id=uid).count()
            today = date.today()
            trainings_upcoming = (
                HRTrainingEnrollment.query
                .join(HRTrainingProgram, HRTrainingProgram.id == HRTrainingEnrollment.program_id)
                .filter(HRTrainingEnrollment.user_id == uid)
                .filter(or_(HRTrainingProgram.start_date == None, HRTrainingProgram.start_date >= today.strftime('%Y-%m-%d')))
                .count()
            )
        except Exception:
            trainings_total = 0
            trainings_upcoming = 0

    # Attendance events
    attendance_events_today = 0
    attendance_events_month = 0
    if can_att:
        try:
            today = date.today()
            attendance_events_today = AttendanceEvent.query.filter(
                AttendanceEvent.user_id == uid,
                func.date(AttendanceEvent.event_dt) == today,
            ).count()
            attendance_events_month = AttendanceEvent.query.filter(
                AttendanceEvent.user_id == uid,
                func.strftime('%Y-%m', AttendanceEvent.event_dt) == today.strftime('%Y-%m')
            ).count()
        except Exception:
            pass

    # System evaluations (monthly + annual) - show only if user has explicit permission
    latest_monthly_eval = None
    latest_annual_eval = None
    if can_sys_eval:
        try:
            latest_monthly_eval = (
                EmployeeEvaluationRun.query
                .filter(EmployeeEvaluationRun.user_id == uid)
                .filter(EmployeeEvaluationRun.period_type == "MONTHLY")
                .order_by(EmployeeEvaluationRun.year.desc(), EmployeeEvaluationRun.month.desc(), EmployeeEvaluationRun.created_at.desc())
                .first()
            )
        except Exception:
            latest_monthly_eval = None
        try:
            latest_annual_eval = (
                EmployeeEvaluationRun.query
                .filter(EmployeeEvaluationRun.user_id == uid)
                .filter(EmployeeEvaluationRun.period_type == "ANNUAL")
                .order_by(EmployeeEvaluationRun.year.desc(), EmployeeEvaluationRun.created_at.desc())
                .first()
            )
        except Exception:
            latest_annual_eval = None

    return render_template(
        "portal/hr/me_home.html",
        can_att=can_att,
        can_hr=can_hr,
        can_training=can_training,
        can_ss=can_ss,
        can_docs=can_docs,
        can_sys_eval=can_sys_eval,
        attendance_events_today=attendance_events_today,
        attendance_events_month=attendance_events_month,
        leaves_total=leaves_total,
        leaves_pending=leaves_pending,
        permissions_total=permissions_total,
        permissions_pending=permissions_pending,
        payslips_count=payslips_count,
        payslips_latest_label=payslips_latest_label,
        trainings_total=trainings_total,
        trainings_upcoming=trainings_upcoming,
        latest_monthly_eval=latest_monthly_eval,
        latest_annual_eval=latest_annual_eval,
    )


@portal_bp.route("/hr/me/attendance")
@login_required
@_perm(PORTAL_READ)
def hr_my_attendance():
    """Employee attendance view (read-only)."""
    # Gate by attendance read permission
    try:
        if not current_user.has_perm(HR_ATT_READ):
            abort(403)
    except Exception:
        abort(403)

    emp_file = EmployeeFile.query.filter_by(user_id=current_user.id).first()
    code = (emp_file.timeclock_code if emp_file else None)

    rows = (
        AttendanceDailySummary.query
        .filter(AttendanceDailySummary.user_id == current_user.id)
        .order_by(AttendanceDailySummary.day.desc())
        .limit(45)
        .all()
    )

    return render_template(
        "portal/hr/my_attendance.html",
        timeclock_code=code,
        rows=rows,
    )


@portal_bp.route("/hr/me/leaves")
@login_required
@_perm(PORTAL_READ)
def hr_my_leaves():
    """Employee leaves (types + requests)."""
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    types = HRLeaveType.query.filter_by(is_active=True).order_by(HRLeaveType.name_ar.asc()).all()

    # Requests list (only if employee has requests read)
    can_requests = False
    can_create = False
    try:
        can_requests = current_user.has_perm(HR_REQUESTS_READ)
        can_create = current_user.has_perm(HR_REQUESTS_CREATE)
    except Exception:
        pass

    reqs = []
    if can_requests:
        reqs = (
            HRLeaveRequest.query
            .filter(HRLeaveRequest.user_id == current_user.id)
            .order_by(HRLeaveRequest.created_at.desc())
            .limit(50)
            .all()
        )


    # Attachments per request
    atts_map = {}
    try:
        ids = [r.id for r in (reqs or [])]
        if ids:
            atts = HRLeaveAttachment.query.filter(HRLeaveAttachment.request_id.in_(ids)).order_by(HRLeaveAttachment.id.desc()).all()
            for a in atts:
                atts_map.setdefault(a.request_id, []).append(a)
    except Exception:
        atts_map = {}

    today_str = date.today().strftime("%Y-%m-%d")

    return render_template(
        "portal/hr/my_leaves.html",
        types=types,
        reqs=reqs,
        can_create=can_create,
        can_requests=can_requests,
        today_str=today_str,
        atts_map=atts_map,
    )


@portal_bp.route("/hr/me/permissions")
@login_required
@_perm(PORTAL_READ)
def hr_my_permissions():
    """Employee departures/permissions (types + requests)."""
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    # Admin/Manager log view (same menu item "السجل")
    can_view_all = False
    can_approve = False
    try:
        can_view_all = bool(
            current_user.has_perm(HR_REQUESTS_VIEW_ALL)
            or current_user.has_role("SUPER_ADMIN")
            or current_user.has_role("SUPERADMIN")
        )
        can_approve = bool(current_user.has_perm(HR_REQUESTS_APPROVE) or can_view_all)
    except Exception:
        can_view_all = False
        can_approve = False

    if can_approve:
        """Admin/Manager departures log.

        المطلوب حسب الشاشات المرفقة:
        - الأدمن يرى الجميع
        - المدير يرى فقط موظفيه (direct reports) + نفسه
        """

        # Filters (match the attached DUAM screens)
        govs = _hr_lookup_items_for_category('WORK_GOVERNORATE')
        locs = _hr_lookup_items_for_category('WORK_LOCATION')
        types = HRPermissionType.query.filter_by(is_active=True).order_by(HRPermissionType.name_ar.asc()).all()

        user_id = (request.args.get('user_id') or '').strip()
        type_id = (request.args.get('type_id') or '').strip()
        date_from = (request.args.get('date_from') or '').strip()
        date_to = (request.args.get('date_to') or '').strip()
        status = (request.args.get('status') or '').strip().upper()
        work_governorate_lookup_id = (request.args.get('work_governorate_lookup_id') or '').strip()
        work_location_lookup_id = (request.args.get('work_location_lookup_id') or '').strip()
        created_by_id = (request.args.get('created_by_id') or '').strip()
        q_text = (request.args.get('q') or '').strip()

        # Base query (we join User/EmployeeFile once to support manager scope + search + work location filters)
        q = (
            HRPermissionRequest.query
            .join(User, User.id == HRPermissionRequest.user_id)
            .outerjoin(EmployeeFile, EmployeeFile.user_id == User.id)
        )

        # Manager scope: only direct reports + self (unless view_all)
        if not can_view_all:
            q = q.filter(
                or_(
                    EmployeeFile.direct_manager_user_id == getattr(current_user, 'id', None),
                    HRPermissionRequest.user_id == getattr(current_user, 'id', None),
                )
            )

        if user_id.isdigit():
            q = q.filter(HRPermissionRequest.user_id == int(user_id))
        if type_id.isdigit():
            q = q.filter(HRPermissionRequest.permission_type_id == int(type_id))
        if status:
            q = q.filter(HRPermissionRequest.status == status)

        if date_from:
            q = q.filter(HRPermissionRequest.day >= date_from)
        if date_to:
            q = q.filter(HRPermissionRequest.day <= date_to)

        if work_governorate_lookup_id.isdigit():
            q = q.filter(EmployeeFile.work_governorate_lookup_id == int(work_governorate_lookup_id))
        if work_location_lookup_id.isdigit():
            q = q.filter(EmployeeFile.work_location_lookup_id == int(work_location_lookup_id))

        if created_by_id.isdigit():
            q = q.filter(HRPermissionRequest.created_by_id == int(created_by_id))

        if q_text:
            like = f"%{q_text}%"
            q = q.filter(or_(User.name.ilike(like), User.email.ilike(like), HRPermissionRequest.note.ilike(like)))

        reqs = q.order_by(HRPermissionRequest.created_at.desc()).limit(500).all()

        # Users map for dropdowns/table (respect scope)
        if can_view_all:
            ulist = (
                User.query
                .outerjoin(EmployeeFile, EmployeeFile.user_id == User.id)
                .order_by(User.name.asc(), User.email.asc())
                .limit(4000)
                .all()
            )
        else:
            ulist = (
                User.query
                .join(EmployeeFile, EmployeeFile.user_id == User.id)
                .filter(EmployeeFile.direct_manager_user_id == getattr(current_user, 'id', None))
                .order_by(User.name.asc(), User.email.asc())
                .all()
            )
            # Ensure manager appears in list
            try:
                me = User.query.get(getattr(current_user, 'id', None))
                if me and all(u.id != me.id for u in ulist):
                    ulist = [me] + ulist
            except Exception:
                pass

        users_map = {u.id: u for u in (ulist or [])}

        return render_template(
            'portal/hr/permissions_log.html',
            reqs=reqs,
            users=users_map,
            types=types,
            govs=govs,
            locs=locs,
            work_governorate_lookup_id=work_governorate_lookup_id,
            work_location_lookup_id=work_location_lookup_id,
            user_id=user_id,
            type_id=type_id,
            date_from=date_from,
            date_to=date_to,
            status=status,
            created_by_id=created_by_id,
            q=q_text,
            can_view_all=can_view_all,
        )

    types = HRPermissionType.query.filter_by(is_active=True).order_by(HRPermissionType.name_ar.asc()).all()

    can_requests = False
    can_create = False
    try:
        can_requests = current_user.has_perm(HR_REQUESTS_READ)
        can_create = current_user.has_perm(HR_REQUESTS_CREATE)
    except Exception:
        pass

    reqs = []
    if can_requests:
        reqs = (
            HRPermissionRequest.query
            .filter(HRPermissionRequest.user_id == current_user.id)
            .order_by(HRPermissionRequest.created_at.desc())
            .limit(50)
            .all()
        )

    return render_template(
        "portal/hr/my_permissions.html",
        types=types,
        reqs=reqs,
        can_create=can_create,
        can_requests=can_requests,
    )


# -------------------------
# HR - Employee requests (Leaves / Permissions)
# -------------------------

@portal_bp.route("/hr/me/leaves/new", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_READ)
def hr_leave_request_new():
    """Submit a leave request.

    Simplicity rules:
    - One form, one submit.
    - Approval goes to direct manager if configured, otherwise stays pending.
    """
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    # Must have requests create to submit
    try:
        if not current_user.has_perm(HR_REQUESTS_CREATE):
            abort(403)
    except Exception:
        abort(403)

    types = HRLeaveType.query.filter_by(is_active=True).order_by(HRLeaveType.name_ar.asc()).all()
    types_meta = {str(t.id): {"is_external": bool(getattr(t, "is_external", False)),
                                   "requires_documents": bool(getattr(t, "requires_documents", False)),
                                   "documents_hint": (getattr(t, "documents_hint", None) or "")} for t in types}

    if request.method == "POST":
        leave_type_id = (request.form.get("leave_type_id") or "").strip()
        start_s = (request.form.get("start_date") or "").strip()
        end_s = (request.form.get("end_date") or "").strip()
        note = (request.form.get("note") or "").strip()

        days_s = (request.form.get("days") or "").strip()
        leave_place = (request.form.get("leave_place") or "").strip()

        # External leave fields (all optional)
        travel_country = (request.form.get("travel_country") or "").strip() or None
        travel_city = (request.form.get("travel_city") or "").strip() or None
        travel_address = (request.form.get("travel_address") or "").strip() or None
        travel_contact_phone = (request.form.get("travel_contact_phone") or "").strip() or None
        travel_purpose = (request.form.get("travel_purpose") or "").strip() or None
        border_crossing = (request.form.get("border_crossing") or "").strip() or None

        # Attachments (optional; may be required by leave type)
        raw_files = []
        try:
            raw_files = request.files.getlist('attachments') or []
        except Exception:
            raw_files = []
        valid_files = [f for f in raw_files if f and getattr(f, 'filename', '')]

        lt = None
        try:
            lt = HRLeaveType.query.get(int(leave_type_id))
        except Exception:
            lt = None

        start_d = _parse_yyyy_mm_dd(start_s)
        end_d = _parse_yyyy_mm_dd(end_s)

        if not lt or not lt.is_active:
            flash("اختر نوع الإجازة.", "danger")
            return render_template("portal/hr/leave_request_new.html", types=types, types_meta=types_meta)

        if not start_d or not end_d:
            flash("تأكد من إدخال تاريخ البداية والنهاية.", "danger")
            return render_template("portal/hr/leave_request_new.html", types=types, types_meta=types_meta)

        if end_d < start_d:
            flash("تاريخ النهاية يجب أن يكون بعد تاريخ البداية.", "danger")
            return render_template("portal/hr/leave_request_new.html", types=types, types_meta=types_meta)

        auto_days = _calc_leave_days_excluding_off(start_s, end_s)
        if not auto_days:
            auto_days = (end_d - start_d).days + 1
        try:
            days = int(days_s) if days_s else int(auto_days)
        except Exception:
            days = int(auto_days)


        # Enforce documents if the leave type requires them
        if getattr(lt, 'requires_documents', False) and not valid_files:
            flash('هذا النوع من الإجازات يتطلب إرفاق تقرير/مستند.', 'danger')
            return render_template('portal/hr/leave_request_new.html', types=types, types_meta=types_meta)

        # Enforce max days if set (with optional exceptional max)
        exceptional = False
        if lt.max_days and days > int(lt.max_days):
            ex = getattr(lt, "exception_max_days", None)
            if ex and days <= int(ex):
                exceptional = True
                flash(f"تنبيه: مدة الإجازة ({days} يوم) تتجاوز الحد الطبيعي ({lt.max_days})، وسيتم التعامل معها كحالة استثنائية (قد تتطلب اعتماد HR).", "warning")
            else:
                flash(f"عدد الأيام يتجاوز الحد الأقصى لهذا النوع ({lt.max_days}).", "danger")
                return render_template("portal/hr/leave_request_new.html", types=types, types_meta=types_meta)

        mgr = _find_direct_manager(current_user)

        status = "SUBMITTED"
        submitted_at = datetime.utcnow()
        decided_at = None
        decided_by_id = None

        if lt.requires_approval is False and not exceptional:
            status = "APPROVED"
            decided_at = submitted_at
            decided_by_id = current_user.id

        is_external = bool(getattr(lt, "is_external", False))

        lp = (leave_place or '').strip().upper()
        leave_place_val = lp if lp in ('INTERNAL','EXTERNAL') else None
        if is_external and not leave_place_val:
            leave_place_val = 'EXTERNAL'

        req = HRLeaveRequest(
            user_id=current_user.id,
            leave_type_id=lt.id,
            start_date=start_d.strftime("%Y-%m-%d"),
            end_date=end_d.strftime("%Y-%m-%d"),
            days=days,
            leave_place=leave_place_val,
            entered_by='SELF',
            created_by_id=getattr(current_user,'id',None),
            note=note or None,

            # store only if external leave type
            travel_country=travel_country if is_external else None,
            travel_city=travel_city if is_external else None,
            travel_address=travel_address if is_external else None,
            travel_contact_phone=travel_contact_phone if is_external else None,
            travel_purpose=travel_purpose if is_external else None,
            border_crossing=border_crossing if is_external else None,

            status=status,
            submitted_at=submitted_at,
            approver_user_id=(mgr.id if mgr else None),
            decided_at=decided_at,
            decided_by_id=decided_by_id,
        )
        db.session.add(req)
        db.session.flush()

        # Save attachments after request id is available
        saved_atts = 0
        try:
            saved_atts = _save_leave_files(valid_files, req.id, doc_type="REPORT")
        except Exception:
            saved_atts = 0

        db.session.commit()

        if status == "APPROVED":
            flash("تم اعتماد طلب الإجازة تلقائياً.", "success")
        else:
            flash("تم إرسال طلب الإجازة بنجاح.", "success")
            if not mgr:
                flash("ملاحظة: لم يتم تحديد مدير مباشر في الهيكل التنظيمي. يمكنك متابعة حالة الطلب، أو طلب من HR تحديد المدير.", "warning")
        return redirect(url_for("portal.hr_my_leaves"))

    return render_template("portal/hr/leave_request_new.html", types=types, types_meta=types_meta)

@portal_bp.route("/hr/me/leaves/<int:req_id>/cancel", methods=["POST"])
@login_required
@_perm(PORTAL_READ)
def hr_leave_request_cancel(req_id: int):
    """Employee cancels a leave request.

    Rules:
    - DRAFT/SUBMITTED: can cancel anytime.
    - APPROVED: can cancel only before the leave start date.
    - After the leave starts: only HR manager can cancel.
    """
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
        if not current_user.has_perm(HR_REQUESTS_CREATE):
            abort(403)
    except Exception:
        abort(403)

    r = HRLeaveRequest.query.get_or_404(req_id)
    if r.user_id != current_user.id:
        abort(403)

    st = (r.status or "").upper()
    today_str = date.today().strftime("%Y-%m-%d")

    allowed = False
    if st in {"SUBMITTED", "DRAFT"}:
        allowed = True
    elif st == "APPROVED":
        # cancel allowed only before start date
        if r.start_date and r.start_date > today_str:
            allowed = True

    if not allowed:
        flash("لا يمكن إلغاء هذا الطلب في حالته الحالية. بعد بدء الإجازة لا يمكن الإلغاء إلا بواسطة مدير الموارد البشرية.", "warning")
        return redirect(url_for("portal.hr_my_leaves"))

    prev = st or None
    r.status = "CANCELLED"
    r.cancelled_from_status = prev
    r.cancelled_at = datetime.utcnow()
    r.cancelled_by_id = current_user.id
    r.cancel_note = "إلغاء من الموظف"
    r.cancel_effective_date = today_str
    r.updated_at = datetime.utcnow()
    db.session.commit()

    flash("تم إلغاء الطلب.", "success")
    return redirect(url_for("portal.hr_my_leaves"))



@portal_bp.route("/hr/me/leaves/<int:req_id>/attachments/upload", methods=["POST"])
@login_required
@_perm(PORTAL_READ)
def hr_leave_attachments_upload(req_id: int):
    """Upload one or more attachments for a leave request (employee side)."""
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
        if not current_user.has_perm(HR_REQUESTS_CREATE):
            abort(403)
    except Exception:
        abort(403)

    r = HRLeaveRequest.query.get_or_404(req_id)
    if r.user_id != current_user.id:
        abort(403)

    st = (r.status or "").upper()
    if st in {"REJECTED", "CANCELLED"}:
        flash("لا يمكن رفع مرفقات على طلب مُغلق.", "warning")
        return redirect(url_for("portal.hr_my_leaves"))

    files = request.files.getlist("attachments") or []
    valid_files = [f for f in files if f and getattr(f, "filename", "") and _allowed_file(f.filename)]
    if not valid_files:
        flash("لم يتم اختيار ملفات صالحة. الامتدادات المسموحة: " + ", ".join(sorted([e.lstrip('.') for e in ALLOWED_CORR_EXTS])), "danger")
        return redirect(url_for("portal.hr_my_leaves"))

    try:
        saved = _save_leave_files(valid_files, req_id, doc_type="REPORT")
        db.session.commit()
        flash(f"تم رفع {saved} مرفق(ات).", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر رفع المرفقات.", "danger")

    return redirect(url_for("portal.hr_my_leaves"))


@portal_bp.route("/hr/leaves/<int:req_id>/attachments/<int:att_id>/download")
@login_required
@_perm(PORTAL_READ)
def hr_leave_attachment_download(req_id: int, att_id: int):
    """Download a leave request attachment.

    Allowed for:
    - owner employee
    - approver
    - HR/Admin with view-all permissions
    """
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    r = HRLeaveRequest.query.get_or_404(req_id)
    att = HRLeaveAttachment.query.filter_by(id=att_id, request_id=req_id).first_or_404()

    allowed = (r.user_id == current_user.id) or (r.approver_user_id == current_user.id)
    try:
        if current_user.has_perm(HR_REQUESTS_VIEW_ALL) or current_user.has_perm(HR_REQUESTS_APPROVE) or current_user.has_perm(HR_EMP_MANAGE) or current_user.has_perm(HR_EMP_READ):
            allowed = True
    except Exception:
        pass

    if not allowed:
        abort(403)

    folder = _leave_upload_dir(req_id)
    return send_from_directory(folder, att.stored_name, as_attachment=True, download_name=(att.original_name or att.stored_name))


@portal_bp.route("/hr/me/permissions/new", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_READ)
def hr_permission_request_new():
    """Create a permission (departure) request.

    - If a normal employee opens this page: employee is auto-assigned (cannot choose another).
    - If HR/Admin/Manager opens this page: they may choose an employee and enter it on their behalf.
    """
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
        if not current_user.has_perm(HR_REQUESTS_CREATE):
            abort(403)
    except Exception:
        abort(403)

    # HRPermissionType uses name_ar/name_en (no generic "name" column)
    types = HRPermissionType.query.order_by(HRPermissionType.name_ar.asc()).all()

    # Admin/Manager can choose the employee
    can_pick_any = False
    can_pick_team = False
    try:
        can_pick_any = bool(
            current_user.has_perm(HR_REQUESTS_VIEW_ALL)
            or current_user.has_role("SUPER_ADMIN")
            or current_user.has_role("SUPERADMIN")
        )
        can_pick_team = bool(current_user.has_perm(HR_REQUESTS_APPROVE) and not can_pick_any)
    except Exception:
        can_pick_any = False
        can_pick_team = False

    is_admin_entry = bool(can_pick_any or can_pick_team)

    users = []
    if can_pick_any:
        users = (
            User.query.join(EmployeeFile, EmployeeFile.user_id == User.id)
            .order_by(User.name.asc(), User.email.asc())
            .all()
        )
    elif can_pick_team:
        # Limit to direct reports for managers
        users = (
            User.query.join(EmployeeFile, EmployeeFile.user_id == User.id)
            .filter(EmployeeFile.direct_manager_user_id == getattr(current_user, 'id', None))
            .order_by(User.name.asc(), User.email.asc())
            .all()
        )
        # Ensure manager can still pick themselves
        if getattr(current_user, 'id', None) and all(u.id != current_user.id for u in users):
            try:
                me = User.query.get(current_user.id)
                if me:
                    users = [me] + users
            except Exception:
                pass

    # Default selected user
    selected_user_id = current_user.id
    if is_admin_entry:
        try:
            selected_user_id = int(request.args.get("user_id") or request.form.get("user_id") or current_user.id)
        except Exception:
            selected_user_id = current_user.id

    if request.method == "POST":
        try:
            permission_type_id = int(request.form.get("permission_type_id") or 0)
        except Exception:
            permission_type_id = 0

        # Who is the target employee?
        target_user_id = current_user.id
        if is_admin_entry:
            try:
                target_user_id = int(request.form.get("user_id") or selected_user_id)
            except Exception:
                target_user_id = current_user.id

            # Manager scope check (cannot pick outside team)
            if can_pick_team and (target_user_id != getattr(current_user, 'id', None)):
                try:
                    ok = EmployeeFile.query.filter_by(user_id=target_user_id, direct_manager_user_id=current_user.id).first() is not None
                except Exception:
                    ok = False
                if not ok:
                    flash('لا يمكنك إدخال مغادرة لموظف خارج صلاحيتك.', 'danger')
                    return render_template(
                        "portal/hr/permission_request_new.html",
                        types=types,
                        users=users,
                        is_admin_entry=is_admin_entry,
                        selected_user_id=selected_user_id,
                    )

        day = (request.form.get("day") or "").strip()
        from_time = (request.form.get("from_time") or "").strip()
        to_time = (request.form.get("to_time") or "").strip()
        note = (request.form.get("note") or "").strip()

        # Basic validation
        if not permission_type_id:
            flash("نوع المغادرة مطلوب.", "danger")
            return render_template(
                "portal/hr/permission_request_new.html",
                types=types,
                users=users,
                is_admin_entry=is_admin_entry,
                selected_user_id=selected_user_id,
            )
        if not day:
            flash("التاريخ مطلوب.", "danger")
            return render_template(
                "portal/hr/permission_request_new.html",
                types=types,
                users=users,
                is_admin_entry=is_admin_entry,
                selected_user_id=selected_user_id,
            )

        # Parse times (allow empty -> treated as 00:00)
        try:
            ft = _parse_hhmm(from_time)
            tt = _parse_hhmm(to_time)
        except Exception:
            flash("صيغة الوقت غير صحيحة. استخدم HH:MM.", "danger")
            return render_template(
                "portal/hr/permission_request_new.html",
                types=types,
                users=users,
                is_admin_entry=is_admin_entry,
                selected_user_id=selected_user_id,
            )

        pt = HRPermissionType.query.get(permission_type_id)
        if not pt:
            flash("نوع المغادرة غير موجود.", "danger")
            return render_template(
                "portal/hr/permission_request_new.html",
                types=types,
                users=users,
                is_admin_entry=is_admin_entry,
                selected_user_id=selected_user_id,
            )

        # Determine status
        status = "SUBMITTED"
        decided_by_id = None
        decided_at = None
        approver_user_id = None

        # For admin-entry, consider it approved directly
        if is_admin_entry:
            status = "APPROVED"
            decided_by_id = current_user.id
            decided_at = datetime.utcnow()
        else:
            if not pt.requires_approval:
                status = "APPROVED"
                decided_by_id = current_user.id
                decided_at = datetime.utcnow()
            else:
                approver_user_id = _hr_find_approver_for_user(target_user_id)

        req = HRPermissionRequest(
            user_id=target_user_id,
            permission_type_id=permission_type_id,
            day=day,
            from_time=f"{ft.hour:02d}:{ft.minute:02d}" if ft else None,
            to_time=f"{tt.hour:02d}:{tt.minute:02d}" if tt else None,
            note=note or None,
            status=status,
            submitted_at=datetime.utcnow() if status in {"SUBMITTED", "APPROVED"} else None,
            approver_user_id=approver_user_id,
            decided_by_id=decided_by_id,
            decided_at=decided_at,
            created_by_id=current_user.id,
        )
        db.session.add(req)
        db.session.commit()

        # Optional attachment
        file = request.files.get("attachment")
        if file and getattr(file, "filename", ""):
            try:
                import os
                from uuid import uuid4
                from werkzeug.utils import secure_filename

                base = os.path.join(current_app.instance_path, "uploads", "permissions", str(req.id))
                os.makedirs(base, exist_ok=True)
                original = file.filename
                stored = f"{uuid4().hex}_{secure_filename(original)}"
                full_path = os.path.join(base, stored)
                file.save(full_path)

                req.attachment_name = original
                req.attachment_path = os.path.relpath(full_path, current_app.instance_path)
                db.session.commit()
            except Exception:
                # Attachment failure shouldn't block saving
                db.session.rollback()
                flash("تم حفظ المغادرة، لكن تعذّر حفظ المرفق.", "warning")

        flash("تم حفظ المغادرة.", "success")

        # Redirect: admin/manager goes to log; employee goes to my list
        return redirect(url_for("portal.hr_my_permissions"))

    return render_template(
        "portal/hr/permission_request_new.html",
        types=types,
        users=users,
        is_admin_entry=is_admin_entry,
        selected_user_id=selected_user_id,
    )


@portal_bp.route("/hr/permissions/<int:req_id>/edit", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_READ)
def hr_permission_request_edit(req_id: int):
    """Edit a departure (permission) record.

    Rules:
    - Employee: can edit own request only while status is DRAFT/SUBMITTED.
    - Manager: can edit requests for their direct reports (EmployeeFile.direct_manager_user_id == current_user.id).
    - HR/Admin: can edit any request.
    """
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    r = HRPermissionRequest.query.get_or_404(req_id)

    # Resolve role-based scope
    can_view_all = False
    can_approve = False
    try:
        can_view_all = bool(
            current_user.has_perm(HR_REQUESTS_VIEW_ALL)
            or current_user.has_role("SUPER_ADMIN")
            or current_user.has_role("SUPERADMIN")
        )
        can_approve = bool(current_user.has_perm(HR_REQUESTS_APPROVE) or can_view_all)
    except Exception:
        can_view_all = False
        can_approve = False

    is_owner = (r.user_id == getattr(current_user, 'id', None))
    status = (r.status or '').upper()

    # Permission check
    if can_view_all:
        pass
    elif can_approve and (not can_view_all):
        # Manager scope: direct reports only (+ self)
        ok = False
        if r.user_id == getattr(current_user, 'id', None):
            ok = True
        else:
            try:
                ok = EmployeeFile.query.filter_by(user_id=r.user_id, direct_manager_user_id=current_user.id).first() is not None
            except Exception:
                ok = False
        if not ok:
            abort(403)
    elif is_owner and status in {'DRAFT', 'SUBMITTED'}:
        try:
            if not current_user.has_perm(HR_REQUESTS_CREATE):
                abort(403)
        except Exception:
            abort(403)
    else:
        abort(403)

    # HRPermissionType uses name_ar/name_en (no generic "name" column)
    types = HRPermissionType.query.order_by(HRPermissionType.name_ar.asc()).all()

    # Employee picker for admin/manager
    can_pick_any = bool(can_view_all)
    can_pick_team = bool((not can_pick_any) and can_approve)

    users = []
    if can_pick_any:
        users = (
            User.query.join(EmployeeFile, EmployeeFile.user_id == User.id)
            .order_by(User.name.asc(), User.email.asc())
            .all()
        )
    elif can_pick_team:
        users = (
            User.query.join(EmployeeFile, EmployeeFile.user_id == User.id)
            .filter(EmployeeFile.direct_manager_user_id == getattr(current_user, 'id', None))
            .order_by(User.name.asc(), User.email.asc())
            .all()
        )
        if getattr(current_user, 'id', None) and all(u.id != current_user.id for u in users):
            try:
                me = User.query.get(current_user.id)
                if me:
                    users = [me] + users
            except Exception:
                pass

    selected_user_id = r.user_id

    if request.method == 'POST':
        try:
            permission_type_id = int(request.form.get('permission_type_id') or 0)
        except Exception:
            permission_type_id = 0

        # Target user
        target_user_id = r.user_id
        if can_pick_any or can_pick_team:
            try:
                target_user_id = int(request.form.get('user_id') or r.user_id)
            except Exception:
                target_user_id = r.user_id

            # Manager scope check
            if can_pick_team and (target_user_id != getattr(current_user, 'id', None)):
                try:
                    ok = EmployeeFile.query.filter_by(user_id=target_user_id, direct_manager_user_id=current_user.id).first() is not None
                except Exception:
                    ok = False
                if not ok:
                    flash('لا يمكنك تعديل مغادرة لموظف خارج صلاحيتك.', 'danger')
                    return render_template(
                        'portal/hr/permission_request_new.html',
                        types=types,
                        users=users,
                        is_admin_entry=True,
                        selected_user_id=r.user_id,
                        edit_mode=True,
                        req=r,
                    )

        day = (request.form.get('day') or '').strip()
        from_time = (request.form.get('from_time') or '').strip()
        to_time = (request.form.get('to_time') or '').strip()
        note = (request.form.get('note') or '').strip()

        if not permission_type_id:
            flash('نوع المغادرة مطلوب.', 'danger')
            return redirect(url_for('portal.hr_permission_request_edit', req_id=req_id))
        if not day:
            flash('التاريخ مطلوب.', 'danger')
            return redirect(url_for('portal.hr_permission_request_edit', req_id=req_id))

        try:
            ft = _parse_hhmm(from_time)
            tt = _parse_hhmm(to_time)
        except Exception:
            flash('صيغة الوقت غير صحيحة. استخدم HH:MM.', 'danger')
            return redirect(url_for('portal.hr_permission_request_edit', req_id=req_id))

        # Apply updates
        r.user_id = target_user_id
        r.permission_type_id = permission_type_id
        r.day = day
        r.from_time = f"{ft.hour:02d}:{ft.minute:02d}" if ft else None
        r.to_time = f"{tt.hour:02d}:{tt.minute:02d}" if tt else None
        r.note = note or None
        r.updated_at = datetime.utcnow()

        # Optional attachment replace
        file = request.files.get('attachment')
        if file and getattr(file, 'filename', ''):
            try:
                import os
                from uuid import uuid4
                from werkzeug.utils import secure_filename

                base = os.path.join(current_app.instance_path, 'uploads', 'permissions', str(r.id))
                os.makedirs(base, exist_ok=True)
                original = file.filename
                stored = f"{uuid4().hex}_{secure_filename(original)}"
                full_path = os.path.join(base, stored)
                file.save(full_path)

                r.attachment_name = original
                r.attachment_path = os.path.relpath(full_path, current_app.instance_path)
            except Exception:
                db.session.rollback()
                flash('تم تعديل المغادرة، لكن تعذّر حفظ المرفق.', 'warning')

        db.session.commit()
        flash('تم تعديل المغادرة.', 'success')
        return redirect(url_for('portal.hr_my_permissions'))

    # GET
    return render_template(
        'portal/hr/permission_request_new.html',
        types=types,
        users=users,
        is_admin_entry=bool(can_pick_any or can_pick_team),
        selected_user_id=selected_user_id,
        edit_mode=True,
        req=r,
    )


@portal_bp.route("/hr/me/permissions/<int:req_id>/cancel", methods=["POST"])
@login_required
@_perm(PORTAL_READ)
def hr_permission_request_cancel(req_id: int):
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
        if not current_user.has_perm(HR_REQUESTS_CREATE):
            abort(403)
    except Exception:
        abort(403)

    r = HRPermissionRequest.query.get_or_404(req_id)
    if r.user_id != current_user.id:
        abort(403)
    if (r.status or "").upper() not in {"SUBMITTED", "DRAFT"}:
        flash("لا يمكن إلغاء هذا الطلب.", "warning")
        return redirect(url_for("portal.hr_my_permissions"))

    prev = (r.status or "").upper() or None

    r.status = "CANCELLED"
    r.cancelled_from_status = prev
    r.cancelled_at = datetime.utcnow()
    r.cancelled_by_id = current_user.id
    r.updated_at = datetime.utcnow()
    db.session.commit()
    flash("تم إلغاء الطلب.", "success")
    return redirect(url_for("portal.hr_my_permissions"))


# -------------------------
# HR - Approvals (Manager / HR)
# -------------------------

@portal_bp.route("/hr/approvals")
@login_required
@_perm(PORTAL_READ)
def hr_approvals():
    """Approvals inbox for managers/HR."""
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    can_view_all = False
    can_approve = False
    try:
        can_view_all = current_user.has_perm(HR_REQUESTS_VIEW_ALL)
        can_approve = current_user.has_perm(HR_REQUESTS_APPROVE) or can_view_all
    except Exception:
        pass

    if not can_approve:
        abort(403)

    status = (request.args.get("status") or "SUBMITTED").upper()
    allowed_status = {"SUBMITTED", "APPROVED", "REJECTED", "CANCELLED", "ALL"}
    if status not in allowed_status:
        status = "SUBMITTED"

    def _base(q):
        if not can_view_all:
            q = q.filter_by(approver_user_id=current_user.id)
        if status != "ALL":
            q = q.filter_by(status=status)
        return q

    leave_reqs = _base(HRLeaveRequest.query).order_by(HRLeaveRequest.created_at.desc()).limit(200).all()
    perm_reqs = _base(HRPermissionRequest.query).order_by(HRPermissionRequest.created_at.desc()).limit(200).all()

    return render_template(
        "portal/hr/approvals.html",
        status=status,
        can_view_all=can_view_all,
        leave_reqs=leave_reqs,
        perm_reqs=perm_reqs,
    )


@portal_bp.route("/hr/approvals/leaves/<int:req_id>", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_READ)
def hr_approval_leave(req_id: int):
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    can_view_all = False
    can_approve = False
    try:
        can_view_all = current_user.has_perm(HR_REQUESTS_VIEW_ALL)
        can_approve = current_user.has_perm(HR_REQUESTS_APPROVE) or can_view_all
    except Exception:
        pass
    if not can_approve:
        abort(403)

    r = HRLeaveRequest.query.get_or_404(req_id)
    if not can_view_all and r.approver_user_id != current_user.id:
        abort(403)

    if request.method == "POST":
        action = (request.form.get("action") or "").strip().upper()
        note = (request.form.get("decision_note") or "").strip()
        if (r.status or "").upper() != "SUBMITTED":
            flash("هذا الطلب ليس بانتظار الاعتماد.", "warning")
            return redirect(url_for("portal.hr_approval_leave", req_id=req_id))

        if action not in {"APPROVE", "REJECT"}:
            flash("إجراء غير صحيح.", "danger")
            return redirect(url_for("portal.hr_approval_leave", req_id=req_id))

        if action == "APPROVE":
            lt = HRLeaveType.query.get(r.leave_type_id) if r.leave_type_id else None
            if lt and lt.max_days and r.days and int(r.days) > int(lt.max_days):
                ex = getattr(lt, "exception_max_days", None)
                if not ex or int(r.days) > int(ex):
                    flash(f"عدد الأيام يتجاوز الحد الأقصى لهذا النوع ({lt.max_days}) ولا يوجد حد استثنائي يسمح بذلك.", "danger")
                    return redirect(url_for("portal.hr_approval_leave", req_id=req_id))

                # Optional restriction: only HR manager can approve exceptional durations
                if getattr(lt, "exception_requires_hr", False):
                    hr_ok = False
                    try:
                        hr_ok = current_user.has_perm(HR_EMP_MANAGE) or current_user.has_perm(HR_REQUESTS_VIEW_ALL) or current_user.has_role("ADMIN")
                    except Exception:
                        hr_ok = False
                    if not hr_ok:
                        flash("هذه حالة استثنائية وتحتاج اعتماد مدير الموارد البشرية.", "warning")
                        return redirect(url_for("portal.hr_approval_leave", req_id=req_id))

                if getattr(lt, "exception_requires_note", False) and not note:
                    flash("ملاحظة القرار مطلوبة لاعتماد الحالة الاستثنائية.", "danger")
                    return redirect(url_for("portal.hr_approval_leave", req_id=req_id))

        r.status = "APPROVED" if action == "APPROVE" else "REJECTED"
        r.decided_at = datetime.utcnow()
        r.decided_by_id = current_user.id
        r.decision_note = note or None
        r.updated_at = datetime.utcnow()
        db.session.commit()

        flash("تم تحديث حالة الطلب.", "success")
        return redirect(url_for("portal.hr_approvals"))

    today_str = date.today().strftime("%Y-%m-%d")
    started = bool(r.start_date and r.start_date <= today_str)
    can_hr_cancel = False
    try:
        can_hr_cancel = current_user.has_perm(HR_EMP_MANAGE) or current_user.has_perm(HR_REQUESTS_VIEW_ALL) or current_user.has_role("ADMIN")
    except Exception:
        can_hr_cancel = False

    attachments = HRLeaveAttachment.query.filter_by(request_id=r.id).order_by(HRLeaveAttachment.id.desc()).all()

    return render_template("portal/hr/approval_leave.html", r=r, today_str=today_str, started=started, can_hr_cancel=can_hr_cancel, attachments=attachments)


@portal_bp.route("/hr/approvals/leaves/<int:req_id>/cancel", methods=["POST"])
@login_required
@_perm(PORTAL_READ)
def hr_leave_cancel_by_hr(req_id: int):
    """HR manager cancels a leave request (even after start date).

    Cancellation stops any future deduction; past days remain counted.
    """
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
        if not (current_user.has_perm(HR_EMP_MANAGE) or current_user.has_perm(HR_REQUESTS_VIEW_ALL) or current_user.has_role("ADMIN")):
            abort(403)
    except Exception:
        abort(403)

    r = HRLeaveRequest.query.get_or_404(req_id)
    st = (r.status or "").upper()
    if st not in {"APPROVED", "SUBMITTED", "DRAFT"}:
        flash("لا يمكن إلغاء هذا الطلب في حالته الحالية.", "warning")
        return redirect(url_for("portal.hr_approval_leave", req_id=req_id))

    today_str = date.today().strftime("%Y-%m-%d")
    prev = st or None

    r.status = "CANCELLED"
    r.cancelled_from_status = prev
    r.cancelled_at = datetime.utcnow()
    r.cancelled_by_id = current_user.id
    r.cancel_note = (request.form.get("cancel_note") or "").strip() or "إلغاء من مدير الموارد البشرية"
    r.cancel_effective_date = today_str
    r.updated_at = datetime.utcnow()
    db.session.commit()

    flash("تم إلغاء الطلب.", "success")
    return redirect(url_for("portal.hr_approval_leave", req_id=req_id))


@portal_bp.route("/hr/approvals/permissions/<int:req_id>", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_READ)
def hr_approval_permission(req_id: int):
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    can_view_all = False
    can_approve = False
    try:
        can_view_all = current_user.has_perm(HR_REQUESTS_VIEW_ALL)
        can_approve = current_user.has_perm(HR_REQUESTS_APPROVE) or can_view_all
    except Exception:
        pass
    if not can_approve:
        abort(403)

    r = HRPermissionRequest.query.get_or_404(req_id)
    if not can_view_all and r.approver_user_id != current_user.id:
        abort(403)

    if request.method == "POST":
        action = (request.form.get("action") or "").strip().upper()
        note = (request.form.get("decision_note") or "").strip()
        if (r.status or "").upper() != "SUBMITTED":
            flash("هذا الطلب ليس بانتظار الاعتماد.", "warning")
            return redirect(url_for("portal.hr_approval_permission", req_id=req_id))

        if action not in {"APPROVE", "REJECT"}:
            flash("إجراء غير صحيح.", "danger")
            return redirect(url_for("portal.hr_approval_permission", req_id=req_id))

        r.status = "APPROVED" if action == "APPROVE" else "REJECTED"
        r.decided_at = datetime.utcnow()
        r.decided_by_id = current_user.id
        r.decision_note = note or None
        r.updated_at = datetime.utcnow()
        db.session.commit()

        flash("تم تحديث حالة الطلب.", "success")
        return redirect(url_for("portal.hr_approvals"))

    return render_template("portal/hr/approval_permission.html", r=r)


def _ensure_employee_attachment_payslip_schema() -> None:
    """Runtime DB upgrade (SQLite): make sure payslip fields exist.

    Users may run an old sqlite DB file created before payslip support.
    This function attempts to add missing columns at runtime (SQLite only).

    Columns ensured:
      - payslip_year, payslip_month
      - is_published (1=published, 0=draft)
      - published_at, published_by_id

    Notes:
      - For backward compatibility, existing rows are treated as published.
    """
    try:
        if getattr(db.engine.dialect, "name", "") != "sqlite":
            return
    except Exception:
        return

    def _cols() -> set[str]:
        try:
            rows = db.session.execute(text("PRAGMA table_info(employee_attachment)")).all()
            return {str(r[1]).strip().lower() for r in (rows or []) if r and r[1]}
        except Exception:
            return set()

    wanted = {
        "payslip_year": "INTEGER",
        "payslip_month": "INTEGER",
        # draft/published flag (treat NULL as published for older data)
        "is_published": "INTEGER DEFAULT 1",
        "published_at": "DATETIME",
        "published_by_id": "INTEGER",
    }

    cols = _cols()
    missing = [c for c in wanted.keys() if c not in cols]
    if not missing:
        # still backfill NULLs
        try:
            if "is_published" in cols:
                db.session.execute(text("UPDATE employee_attachment SET is_published=1 WHERE is_published IS NULL"))
                db.session.commit()
        except Exception:
            db.session.rollback()
        return

    import time

    for attempt in range(5):
        try:
            cols = _cols()
            for col, typ in wanted.items():
                if col.lower() in cols:
                    continue
                db.session.execute(text(f"ALTER TABLE employee_attachment ADD COLUMN {col} {typ}"))

            # Backfill: make older rows published
            cols = _cols()
            if "is_published" in cols:
                db.session.execute(text("UPDATE employee_attachment SET is_published=1 WHERE is_published IS NULL"))

            db.session.commit()
            return
        except Exception as e:
            db.session.rollback()
            msg = str(e).lower()
            if "locked" in msg or "busy" in msg:
                time.sleep(0.3 * (attempt + 1))
                continue
            return


def _payslip_visible_expr():
    """SQL expression: payslip visible to employee.

    Treat NULL as published for backward compatibility.
    """
    return or_(EmployeeAttachment.is_published == True, EmployeeAttachment.is_published.is_(None))
@portal_bp.route("/hr/me/payslips")
@login_required
@_perm(PORTAL_READ)
def hr_my_payslips():
    """Employee payslips list (PDF inline view)."""
    _ensure_employee_attachment_payslip_schema()
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    # Optional filters
    year = (request.args.get("year") or "").strip()
    month = (request.args.get("month") or "").strip()

    q = (
        EmployeeAttachment.query
        .filter(EmployeeAttachment.user_id == current_user.id)
        .filter(EmployeeAttachment.attachment_type == "PAYSLIP")
        .filter(_payslip_visible_expr())
    )

    f_year = None
    f_month = None
    try:
        if year:
            f_year = int(year)
            q = q.filter(EmployeeAttachment.payslip_year == f_year)
        if month:
            f_month = int(month)
            q = q.filter(EmployeeAttachment.payslip_month == f_month)
    except Exception:
        f_year = None
        f_month = None

    # SQLite compatibility: avoid NULLS LAST syntax by using CASE ordering.
    has_year = db.case((EmployeeAttachment.payslip_year.isnot(None), 1), else_=0)
    has_month = db.case((EmployeeAttachment.payslip_month.isnot(None), 1), else_=0)
    slips = q.order_by(
        has_year.desc(),
        EmployeeAttachment.payslip_year.desc(),
        has_month.desc(),
        EmployeeAttachment.payslip_month.desc(),
        EmployeeAttachment.uploaded_at.desc(),
    ).all()

    # Distinct years for dropdown
    years = []
    try:
        years = [r[0] for r in (
            db.session.query(EmployeeAttachment.payslip_year)
            .filter(EmployeeAttachment.user_id == current_user.id)
            .filter(EmployeeAttachment.attachment_type == "PAYSLIP")
            .filter(_payslip_visible_expr())
            .filter(EmployeeAttachment.payslip_year.isnot(None))
            .distinct()
            .order_by(EmployeeAttachment.payslip_year.desc())
            .all()
        ) if r[0] is not None]
    except Exception:
        years = []

    # Quick helpers (current month + latest)
    now_year = None
    now_month = None
    latest_slip = None
    try:
        now = datetime.now()
        now_year = int(now.year)
        now_month = int(now.month)
        has_year = db.case((EmployeeAttachment.payslip_year.isnot(None), 1), else_=0)
        has_month = db.case((EmployeeAttachment.payslip_month.isnot(None), 1), else_=0)
        latest_slip = (
            EmployeeAttachment.query
            .filter(EmployeeAttachment.user_id == current_user.id)
            .filter(EmployeeAttachment.attachment_type == "PAYSLIP")
            .filter(_payslip_visible_expr())
            .order_by(
                has_year.desc(),
                EmployeeAttachment.payslip_year.desc(),
                has_month.desc(),
                EmployeeAttachment.payslip_month.desc(),
                EmployeeAttachment.uploaded_at.desc(),
            )
            .first()
        )
    except Exception:
        now_year = None
        now_month = None
        latest_slip = None

    return render_template(
        "portal/hr/my_payslips.html",
        slips=slips,
        years=years,
        sel_year=f_year,
        sel_month=f_month,
        now_year=now_year,
        now_month=now_month,
        latest_slip=latest_slip,
    )


@portal_bp.route("/hr/me/payslips/<int:att_id>/view")
@login_required
@_perm(PORTAL_READ)
def hr_my_payslip_view(att_id: int):
    """Inline view for a payslip PDF that belongs to current user."""
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    a = EmployeeAttachment.query.get_or_404(att_id)
    if a.user_id != current_user.id or (a.attachment_type or "").upper() != "PAYSLIP":
        abort(403)

    # Draft payslips are not visible until HR publishes/sends them
    try:
        if getattr(a, "is_published", True) is False:
            abort(404)
    except Exception:
        pass

    if not a.stored_name:
        abort(404)

    # NOTE: new uploads are stored under instance/uploads/employees/<user_id>/
    # Keep a fallback to legacy static/uploads/employee for older deployments.
    try:
        dirp = _employee_upload_dir(a.user_id)
        fp = dirp / a.stored_name
        if fp.exists():
            return send_from_directory(str(dirp), a.stored_name, as_attachment=False)
    except Exception:
        pass

    legacy_dir = os.path.join(current_app.root_path, "static", "uploads", "employee")
    return send_from_directory(legacy_dir, a.stored_name, as_attachment=False)


@portal_bp.route("/hr/me/payslips/current")
@login_required
@_perm(PORTAL_READ)
def hr_my_payslip_current():
    """Quick open: current month payslip (if exists)."""
    _ensure_employee_attachment_payslip_schema()
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    now = datetime.now()
    y = int(now.year)
    m = int(now.month)

    slip = (
        EmployeeAttachment.query
        .filter(EmployeeAttachment.user_id == current_user.id)
        .filter(EmployeeAttachment.attachment_type == "PAYSLIP")
        .filter(_payslip_visible_expr())
        .filter(EmployeeAttachment.payslip_year == y)
        .filter(EmployeeAttachment.payslip_month == m)
        .order_by(EmployeeAttachment.uploaded_at.desc())
        .first()
    )

    if slip:
        return redirect(url_for("portal.hr_my_payslip_view", att_id=slip.id))

    flash("لا توجد قسيمة راتب لهذا الشهر.", "warning")
    return redirect(url_for("portal.hr_my_payslips", year=y, month=m))


@portal_bp.route("/hr/me/payslips/latest")
@login_required
@_perm(PORTAL_READ)
def hr_my_payslip_latest():
    """Quick open: latest available payslip."""
    _ensure_employee_attachment_payslip_schema()
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    has_year = db.case((EmployeeAttachment.payslip_year.isnot(None), 1), else_=0)
    has_month = db.case((EmployeeAttachment.payslip_month.isnot(None), 1), else_=0)
    slip = (
        EmployeeAttachment.query
        .filter(EmployeeAttachment.user_id == current_user.id)
        .filter(EmployeeAttachment.attachment_type == "PAYSLIP")
        .filter(_payslip_visible_expr())
        .order_by(
            has_year.desc(),
            EmployeeAttachment.payslip_year.desc(),
            has_month.desc(),
            EmployeeAttachment.payslip_month.desc(),
            EmployeeAttachment.uploaded_at.desc(),
        )
        .first()
    )

    if slip:
        return redirect(url_for("portal.hr_my_payslip_view", att_id=slip.id))

    flash("لا توجد قسائم راتب مرفوعة لك بعد.", "info")
    return redirect(url_for("portal.hr_my_payslips"))


# -------------------------
# Employee: System Evaluations (KPI-based)
# -------------------------
@portal_bp.route("/hr/me/system-evaluations")
@login_required
@_perm(PORTAL_READ)
def hr_my_system_evaluations():
    """Employee: list my monthly + annual system evaluations (if permitted)."""
    try:
        if not current_user.has_perm(HR_SYSTEM_EVALUATION_VIEW):
            abort(403)
    except Exception:
        abort(403)

    uid = current_user.id

    monthly = (
        EmployeeEvaluationRun.query
        .filter(EmployeeEvaluationRun.user_id == uid)
        .filter(EmployeeEvaluationRun.period_type == "MONTHLY")
        .order_by(EmployeeEvaluationRun.year.desc(), EmployeeEvaluationRun.month.desc(), EmployeeEvaluationRun.created_at.desc())
        .limit(24)
        .all()
    )

    annual = (
        EmployeeEvaluationRun.query
        .filter(EmployeeEvaluationRun.user_id == uid)
        .filter(EmployeeEvaluationRun.period_type == "ANNUAL")
        .order_by(EmployeeEvaluationRun.year.desc(), EmployeeEvaluationRun.created_at.desc())
        .limit(6)
        .all()
    )

    return render_template(
        "portal/hr/my_system_evaluations.html",
        monthly=monthly,
        annual=annual,
    )


@portal_bp.route("/hr/me/system-evaluations/<int:run_id>")
@login_required
@_perm(PORTAL_READ)
def hr_my_system_evaluation_view(run_id: int):
    """Employee: view one system evaluation details (must belong to me)."""
    try:
        if not current_user.has_perm(HR_SYSTEM_EVALUATION_VIEW):
            abort(403)
    except Exception:
        abort(403)

    run = EmployeeEvaluationRun.query.get_or_404(run_id)
    if run.user_id != current_user.id:
        abort(403)

    breakdown = {}
    try:
        breakdown = json.loads(run.breakdown_json) if run.breakdown_json else {}
    except Exception:
        breakdown = {}

    breakdown_pretty = ""
    try:
        breakdown_pretty = json.dumps(breakdown, ensure_ascii=False, indent=2)
    except Exception:
        breakdown_pretty = ""

    return render_template(
        "portal/hr/my_system_evaluation_view.html",
        run=run,
        breakdown=breakdown,
        breakdown_pretty=breakdown_pretty,
    )


@portal_bp.route("/hr/payslips/bulk-upload", methods=["GET", "POST"])
@login_required
@_perm(HR_EMP_ATTACH)
def hr_payslips_bulk_upload():
    """HR Admin: bulk upload payslips for a given month.

    Rules:
      - Upload multiple PDF files.
      - Each filename must be the employee number (employee_no) e.g. 12345.pdf
      - Payslips are saved as DRAFTS (do NOT send automatically).
      - HR can publish/send later from: HR → "إرسال قسائم الرواتب".
    """
    _ensure_employee_attachment_payslip_schema()

    now = datetime.now()
    # Allow pre-select via querystring ?year=YYYY&month=MM
    try:
        default_year = int((request.args.get("year") or now.year))
    except Exception:
        default_year = int(now.year)
    try:
        default_month = int((request.args.get("month") or now.month))
    except Exception:
        default_month = int(now.month)
    if not (2000 <= default_year <= 2100):
        default_year = int(now.year)
    if not (1 <= default_month <= 12):
        default_month = int(now.month)

    results = {
        "saved": 0,
        "skipped": 0,
        "errors": [],
    }

    if request.method == "POST":
        y = (request.form.get("year") or "").strip()
        m = (request.form.get("month") or "").strip()

        try:
            year = int(y)
            month = int(m)
        except Exception:
            flash("الرجاء اختيار سنة/شهر صحيحين.", "danger")
            return render_template(
                "portal/hr/payslips_bulk_upload.html",
                default_year=default_year,
                default_month=default_month,
                results=results,
            )

        if not (2000 <= year <= 2100) or not (1 <= month <= 12):
            flash("السنة/الشهر غير صالحين.", "danger")
            return render_template(
                "portal/hr/payslips_bulk_upload.html",
                default_year=default_year,
                default_month=default_month,
                results=results,
            )

        files = request.files.getlist("files") or []
        if not files:
            flash("اختر ملفات PDF لقسائم الرواتب.", "warning")
            return render_template(
                "portal/hr/payslips_bulk_upload.html",
                default_year=year,
                default_month=month,
                results=results,
            )

        def _norm_emp_no(filename: str) -> str:
            name = Path(filename or "").name
            stem = (Path(name).stem or "").strip()
            try:
                stem = "".join(ch for ch in stem if unicodedata.category(ch) != "Cf")
            except Exception:
                pass
            stem = stem.strip()
            stem = re.split(r"[\s_\-]+", stem)[0].strip()
            mm = re.search(r"(\d+)", stem)
            if mm:
                return mm.group(1)
            return stem

        for f in files:
            if not f or not getattr(f, "filename", ""):
                continue

            original = Path(f.filename).name
            ext = _clean_suffix(original)
            if ext != ".pdf":
                results["skipped"] += 1
                results["errors"].append(f"{original}: يجب أن يكون الملف PDF.")
                continue

            emp_no = _norm_emp_no(original)
            if not emp_no:
                results["skipped"] += 1
                results["errors"].append(f"{original}: لم أستطع استخراج الرقم الوظيفي من اسم الملف.")
                continue

            emp = (
                EmployeeFile.query
                .filter(func.trim(EmployeeFile.employee_no) == emp_no)
                .first()
            )
            if not emp:
                results["skipped"] += 1
                results["errors"].append(f"{original}: لا يوجد موظف برقم وظيفي ({emp_no}).")
                continue

            user_id = int(emp.user_id)

            stored = f"{uuid.uuid4().hex}{ext}"
            dirp = _employee_upload_dir(user_id)
            try:
                f.save(dirp / stored)
            except Exception as e:
                results["skipped"] += 1
                results["errors"].append(f"{original}: تعذر حفظ الملف ({str(e)}).")
                continue

            att = EmployeeAttachment.query.filter_by(
                user_id=user_id,
                attachment_type="PAYSLIP",
                payslip_year=year,
                payslip_month=month,
            ).first()

            if att:
                try:
                    old_fp = dirp / (att.stored_name or "")
                    if old_fp.exists():
                        old_fp.unlink()
                except Exception:
                    pass
                att.original_name = original
                att.stored_name = stored
                att.note = None
                att.uploaded_by_id = current_user.id
                att.uploaded_at = datetime.utcnow()
            else:
                att = EmployeeAttachment(
                    user_id=user_id,
                    attachment_type="PAYSLIP",
                    original_name=original,
                    stored_name=stored,
                    note=None,
                    payslip_year=year,
                    payslip_month=month,
                    uploaded_by_id=current_user.id,
                    uploaded_at=datetime.utcnow(),
                )
                db.session.add(att)

            # Draft by default (not visible to employee until sent)
            try:
                att.is_published = False
                att.published_at = None
                att.published_by_id = None
            except Exception:
                pass

            results["saved"] += 1

        try:
            _portal_audit(
                action="HR_PAYSLIPS_BULK_UPLOAD",
                note=f"رفع مسودات قسائم رواتب شهر {year:04d}-{month:02d} (تم حفظ {results['saved']})",
                target_type="PAYSLIP",
                target_id=0,
            )
        except Exception:
            pass

        try:
            db.session.commit()
            flash(
                f"تمت العملية: حفظ {results['saved']} مسودة، تخطي {results['skipped']}.",
                "success",
            )
        except Exception:
            db.session.rollback()
            flash("حدث خطأ أثناء الحفظ في قاعدة البيانات.", "danger")

        return render_template(
            "portal/hr/payslips_bulk_upload.html",
            default_year=year,
            default_month=month,
            results=results,
        )

    return render_template(
        "portal/hr/payslips_bulk_upload.html",
        default_year=default_year,
        default_month=default_month,
        results=results,
    )


@portal_bp.route("/hr/payslips/send", methods=["GET", "POST"])
@login_required
@_perm(HR_EMP_ATTACH)
def hr_payslips_send():
    """HR Admin: publish/send payslips (per month) via internal messages."""
    _ensure_employee_attachment_payslip_schema()

    now = datetime.now()
    year_s = (request.values.get("year") or str(now.year)).strip()
    month_s = (request.values.get("month") or str(now.month)).strip()

    try:
        year = int(year_s)
        month = int(month_s)
    except Exception:
        year = int(now.year)
        month = int(now.month)

    if not (2000 <= year <= 2100) or not (1 <= month <= 12):
        year = int(now.year)
        month = int(now.month)

    results = {
        "sent": 0,
        "skipped": 0,
        "errors": [],
    }

    def _is_published(att: EmployeeAttachment) -> bool:
        try:
            v = getattr(att, "is_published", True)
            if v is None:
                return True
            return bool(v)
        except Exception:
            return True

    def _send_one(att: EmployeeAttachment) -> None:
        if not att:
            return
        if (att.attachment_type or "").upper() != "PAYSLIP":
            results["skipped"] += 1
            return
        if _is_published(att):
            results["skipped"] += 1
            return

        try:
            db.session.flush()
            slip_url = url_for("portal.hr_my_payslip_view", att_id=att.id, _external=True)
            subject = f"قسيمة راتب {year:04d}-{month:02d}"
            body = (
                f"تم إرسال قسيمة راتب شهر {year:04d}-{month:02d} عبر النظام.\n\n"
                f"عرض القسيمة: {slip_url}\n\n"
                "ملاحظة: الرابط يعمل فقط لصاحب القسيمة."
            )

            msg = Message(
                sender_id=current_user.id,
                subject=subject,
                body=body,
                target_kind="USER",
                target_id=int(att.user_id),
                created_at=datetime.utcnow(),
            )
            db.session.add(msg)
            db.session.flush()
            db.session.add(MessageRecipient(
                message_id=msg.id,
                recipient_user_id=int(att.user_id),
                is_read=False,
                read_at=None,
                is_deleted=False,
                deleted_at=None,
            ))

            emit_event(
                actor_id=current_user.id,
                action="PAYSLIP_SENT",
                message=f"وصلتك قسيمة راتب {year:04d}-{month:02d} عبر المراسلات.",
                target_type="PAYSLIP",
                target_id=int(att.id),
                notify_user_id=int(att.user_id),
                level="INFO",
                auto_commit=False,
            )

            # Mark published
            try:
                att.is_published = True
                att.published_at = datetime.utcnow()
                att.published_by_id = current_user.id
            except Exception:
                pass

            results["sent"] += 1
        except Exception as e:
            results["errors"].append(f"{att.original_name or 'Payslip'}: تعذر الإرسال ({str(e)}).")

    if request.method == "POST":
        action = (request.form.get("action") or "").strip()

        base_q = (
            EmployeeAttachment.query
            .filter(EmployeeAttachment.attachment_type == "PAYSLIP")
            .filter(EmployeeAttachment.payslip_year == year)
            .filter(EmployeeAttachment.payslip_month == month)
        )

        targets = []
        try:
            if action == "send_one":
                att_id = int(request.form.get("att_id") or 0)
                if att_id:
                    targets = base_q.filter(EmployeeAttachment.id == att_id).all()
            elif action == "send_selected":
                ids = []
                for v in (request.form.getlist("att_ids") or []):
                    try:
                        ids.append(int(v))
                    except Exception:
                        pass
                if ids:
                    targets = base_q.filter(EmployeeAttachment.id.in_(ids)).all()
            elif action == "send_all":
                # only drafts
                targets = base_q.filter(EmployeeAttachment.is_published == False).all()  # noqa: E712
            else:
                targets = []
        except Exception:
            targets = []

        for att in targets:
            _send_one(att)

        try:
            _portal_audit(
                action="HR_PAYSLIPS_SEND",
                note=f"إرسال قسائم رواتب شهر {year:04d}-{month:02d} (تم إرسال {results['sent']})",
                target_type="PAYSLIP",
                target_id=0,
            )
        except Exception:
            pass

        try:
            db.session.commit()
            if results["sent"] > 0:
                flash(f"تم إرسال {results['sent']} قسيمة عبر المراسلات.", "success")
            else:
                flash("لا يوجد قسائم مسودة للإرسال.", "info")
        except Exception:
            db.session.rollback()
            flash("حدث خطأ أثناء الحفظ/الإرسال.", "danger")

    # List for this month
    q = (
        EmployeeAttachment.query
        .filter(EmployeeAttachment.attachment_type == "PAYSLIP")
        .filter(EmployeeAttachment.payslip_year == year)
        .filter(EmployeeAttachment.payslip_month == month)
        .order_by(EmployeeAttachment.uploaded_at.desc())
    )
    slips = q.all()

    drafts_count = 0
    sent_count = 0
    try:
        drafts_count = q.filter(EmployeeAttachment.is_published == False).count()  # noqa: E712
        sent_count = q.filter(EmployeeAttachment.is_published == True).count()  # noqa: E712
    except Exception:
        # If column missing for any reason, show 0/0
        drafts_count = 0
        sent_count = 0

    return render_template(
        "portal/hr/payslips_send.html",
        year=year,
        month=month,
        slips=slips,
        drafts_count=drafts_count,
        sent_count=sent_count,
        results=results,
    )


@portal_bp.route("/hr/payslips/<int:att_id>/admin-view")
@login_required
@_perm(HR_EMP_ATTACH)
def hr_payslips_admin_view(att_id: int):
    """HR Admin: preview/download any payslip (draft or published)."""
    _ensure_employee_attachment_payslip_schema()
    a = EmployeeAttachment.query.get_or_404(att_id)
    if (a.attachment_type or "").upper() != "PAYSLIP":
        abort(404)
    if not a.stored_name:
        abort(404)

    try:
        dirp = _employee_upload_dir(a.user_id)
        fp = dirp / a.stored_name
        if fp.exists():
            return send_from_directory(str(dirp), a.stored_name, as_attachment=False)
    except Exception:
        pass

    legacy_dir = os.path.join(current_app.root_path, "static", "uploads", "employee")
    return send_from_directory(legacy_dir, a.stored_name, as_attachment=False)


# -------------------------
# Store (Portal Repository)
# -------------------------
@portal_bp.route("/store")
@login_required
@_perm(PORTAL_READ)
def store_home():
    """Repository list with search + folders + sharing.

    Access:
    - Users with STORE_READ/STORE_MANAGE can browse all files.
    - Users without those perms can still access files that are shared with them.
    """
    _ensure_store_ready()

    # Capabilities
    can_manage = False
    can_read_all = False
    try:
        can_manage = bool(current_user.has_perm("STORE_MANAGE"))
        can_read_all = can_manage or bool(current_user.has_perm("STORE_READ"))
    except Exception:
        can_manage = False
        can_read_all = False

    # View mode: all vs shared
    view = (request.args.get("view") or "").strip().lower()  # all|shared
    if view not in {"all", "shared"}:
        view = "all" if can_read_all else "shared"
    if not can_read_all:
        view = "shared"

    # Filters
    q = (request.args.get("q") or "").strip()
    folder = (request.args.get("folder") or "").strip() or (request.args.get("cat") or "").strip()
    page = int(request.args.get("page") or 1)
    per_page = 12

    current_folder = None
    folder_id = None
    if folder:
        try:
            folder_id = int(folder)
            current_folder = StoreCategory.query.get(folder_id)
        except Exception:
            folder_id = None
            current_folder = None

    # Build query
    qry = StoreFile.query.filter(StoreFile.is_deleted == False)  # noqa: E712

    if view == "shared":
        sq = _store_shared_query_for_user().subquery()
        qry = qry.join(sq, sq.c.file_id == StoreFile.id)
        qry = qry.distinct()

    if q:
        # Search across all StoreFile columns (not only title/description).
        qry = apply_search_all_columns(qry, StoreFile, q)

    if folder_id:
        qry = qry.filter(StoreFile.category_id == folder_id)

    qry = qry.order_by(StoreFile.uploaded_at.desc())

    # Excel export for current filters
    if (request.args.get("export") or "").strip() in {"1", "excel"}:
        try:
            rows = qry.limit(5000).all()
            data = []
            for r in rows:
                data.append({
                    "id": r.id,
                    "title": r.title,
                    "category": getattr(getattr(r, "category", None), "name", None),
                    "original_name": r.original_name,
                    "uploaded_by": getattr(getattr(r, "uploaded_by", None), "name", None) or getattr(getattr(r, "uploaded_by", None), "email", None),
                    "uploaded_at": r.uploaded_at.strftime("%Y-%m-%d %H:%M") if r.uploaded_at else "",
                })

            xlsx = _xlsx_from_dicts(
                data,
                columns=[
                    ("id", "ID"),
                    ("title", "العنوان"),
                    ("category", "المجلد"),
                    ("original_name", "اسم الملف"),
                    ("uploaded_by", "الرافع"),
                    ("uploaded_at", "تاريخ الرفع"),
                ],
                sheet_name="Store",
            )

            return send_file(
                io.BytesIO(xlsx),
                as_attachment=True,
                download_name="store_files.xlsx",
                mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        except Exception:
            current_app.logger.exception("Store export failed")
            flash("تعذر تصدير Excel.", "danger")

    pagination = qry.paginate(page=page, per_page=per_page, error_out=False)

    items = list(pagination.items)
    for r in items:
        try:
            setattr(r, "_can_download", _store_can_download_file(r))
        except Exception:
            setattr(r, "_can_download", False)

    # Folder navigation
    try:
        folders = (
            StoreCategory.query
            .filter(StoreCategory.parent_id == folder_id)
            .filter(StoreCategory.is_active == True)  # noqa: E712
            .order_by(StoreCategory.name.asc())
            .all()
        )
    except Exception:
        folders = []

    # Breadcrumb
    breadcrumb = []
    try:
        cur = current_folder
        while cur:
            breadcrumb.append(cur)
            cur = getattr(cur, "parent", None)
        breadcrumb = list(reversed(breadcrumb))
    except Exception:
        breadcrumb = []

    category_options = _store_category_options()

    return render_template(
        "portal/store/index.html",
        q=q,
        view=view,
        folder_id=folder_id,
        breadcrumb=breadcrumb,
        folders=folders,
        category_options=category_options,
        pagination=pagination,
        items=items,
        can_manage=can_manage,
        can_read_all=can_read_all,
    )


@portal_bp.route("/store/shared")
@login_required
@_perm(PORTAL_READ)
def store_shared():
    """Shortcut page: files shared with me."""
    return redirect(url_for("portal.store_home", view="shared"))


@portal_bp.route("/store/files/<int:file_id>/share", methods=["GET", "POST"])
@login_required
@_perm("STORE_MANAGE")
def store_file_share(file_id: int):
    """Share a repository file to a user or a role."""
    _ensure_store_ready()

    row = StoreFile.query.get_or_404(file_id)
    if row.is_deleted:
        abort(404)

    if request.method == "POST":
        mode = (request.form.get("mode") or "").strip().lower()  # user|role
        can_download = bool(request.form.get("can_download") == "1")
        expires = (request.form.get("expires_at") or "").strip()

        expires_at = None
        if expires:
            try:
                # HTML datetime-local: 2026-01-30T12:30
                expires_at = datetime.strptime(expires, "%Y-%m-%dT%H:%M")
            except Exception:
                expires_at = None

        user_obj = None
        role_val = None

        if mode == "user":
            email = (request.form.get("user_email") or "").strip().lower()
            if not email:
                flash("اكتب البريد الإلكتروني للمستخدم.", "danger")
                return redirect(url_for("portal.store_file_share", file_id=file_id))
            user_obj = User.query.filter(func.lower(User.email) == email).first()
            if not user_obj:
                flash("لا يوجد مستخدم بهذا البريد.", "warning")
                return redirect(url_for("portal.store_file_share", file_id=file_id))

        elif mode == "role":
            role_val = (request.form.get("role") or "").strip().upper()
            if not role_val:
                flash("اختر دوراً.", "danger")
                return redirect(url_for("portal.store_file_share", file_id=file_id))
        else:
            flash("اختر نوع المشاركة (مستخدم/دور).", "danger")
            return redirect(url_for("portal.store_file_share", file_id=file_id))

        # Upsert
        q = StoreFilePermission.query.filter(StoreFilePermission.file_id == row.id)
        if user_obj:
            q = q.filter(StoreFilePermission.user_id == user_obj.id)
        else:
            q = q.filter(StoreFilePermission.role == role_val)

        perm = q.first()
        if not perm:
            perm = StoreFilePermission(
                file_id=row.id,
                user_id=user_obj.id if user_obj else None,
                role=role_val,
                can_download=can_download,
                shared_by=getattr(current_user, "id", None),
                expires_at=expires_at,
            )
            db.session.add(perm)
        else:
            perm.can_download = can_download
            perm.expires_at = expires_at
            perm.shared_by = getattr(current_user, "id", None)

        # Audit
        try:
            dest = (user_obj.email if user_obj else f'ROLE:{role_val}')
            db.session.add(AuditLog(user_id=current_user.id, action='STORE_SHARE', note=f'مشاركة ملف مستودع: {row.original_name} -> {dest}', target_type='STORE_FILE', target_id=row.id, created_at=datetime.utcnow()))
        except Exception:
            pass

        db.session.commit()
        flash("تم حفظ المشاركة.", "success")
        return redirect(url_for("portal.store_file_share", file_id=file_id))

    # GET
    role_list = []
    try:
        role_list = [r[0] for r in db.session.query(func.upper(User.role)).filter(User.role.isnot(None)).distinct().all()]
        role_list = sorted({(x or "").strip().upper() for x in role_list if (x or "").strip()})
    except Exception:
        role_list = []

    perms = []
    try:
        perms = (
            StoreFilePermission.query
            .filter(StoreFilePermission.file_id == row.id)
            .order_by(StoreFilePermission.shared_at.desc())
            .all()
        )
    except Exception:
        perms = []

    return render_template(
        "portal/store/share.html",
        f=row,
        perms=perms,
        role_list=role_list,
    )


@portal_bp.route("/store/files/<int:file_id>/share/<int:perm_id>/delete", methods=["POST"])
@login_required
@_perm("STORE_MANAGE")
def store_file_share_delete(file_id: int, perm_id: int):
    _ensure_store_ready()
    row = StoreFile.query.get_or_404(file_id)
    if row.is_deleted:
        abort(404)

    perm = StoreFilePermission.query.get_or_404(perm_id)
    if perm.file_id != row.id:
        abort(400)

    db.session.delete(perm)
    # Audit
    try:
        who = (perm.user.email if getattr(perm, 'user', None) else f'ROLE:{perm.role}')
        db.session.add(AuditLog(user_id=current_user.id, action='STORE_UNSHARE', note=f'إلغاء مشاركة ملف مستودع: {row.original_name} -> {who}', target_type='STORE_FILE', target_id=row.id, created_at=datetime.utcnow()))
    except Exception:
        pass
    db.session.commit()
    flash("تم إلغاء المشاركة.", "success")
    return redirect(url_for("portal.store_file_share", file_id=file_id))


@portal_bp.route("/store/upload", methods=["POST"])
@login_required
@_perm("STORE_MANAGE")
def store_upload():
    _ensure_store_ready()

    f = request.files.get("file")
    title = (request.form.get("title") or "").strip() or None
    desc = (request.form.get("description") or "").strip() or None
    cat_id = (request.form.get("category_id") or "").strip()
    cat_new = (request.form.get("category_new") or "").strip()
    parent_id = (request.form.get("parent_id") or "").strip()  # optional (create new folder under this)

    if not f or not getattr(f, "filename", ""):
        flash("اختر ملفاً للرفع.", "danger")
        return redirect(url_for("portal.store_home"))

    if not _allowed_store_file(f.filename):
        flash("نوع الملف غير مسموح.", "danger")
        return redirect(url_for("portal.store_home"))

    # Category resolution
    category = None
    if cat_new:
        cat_new = cat_new[:120]
        category = StoreCategory.query.filter_by(name=cat_new).first()
        if not category:
            pid = None
            if parent_id:
                try:
                    pid = int(parent_id)
                except Exception:
                    pid = None
            category = StoreCategory(name=cat_new, is_active=True, parent_id=pid)
            db.session.add(category)
            db.session.flush()
    elif cat_id:
        try:
            category = StoreCategory.query.get(int(cat_id))
        except Exception:
            category = None

    storage = _store_storage_dir()
    original_name = f.filename
    ext = _clean_suffix(original_name)
    # uuid is imported as the standard library module; use uuid.uuid4() directly
    stored_name = f"STORE_{uuid.uuid4().hex}{ext}"
    file_path = os.path.join(storage, stored_name)
    f.save(file_path)

    mt = mimetypes.guess_type(original_name)[0] or "application/octet-stream"
    try:
        size = os.path.getsize(file_path)
    except Exception:
        size = None

    row = StoreFile(
        title=title,
        description=desc,
        original_name=original_name,
        stored_name=stored_name,
        file_path=file_path,
        mime_type=mt,
        file_size=size,
        category_id=(category.id if category else None),
        uploader_id=current_user.id,
        uploaded_at=datetime.utcnow(),
    )
    db.session.add(row)
    db.session.flush()

    # Audit
    try:
        db.session.add(AuditLog(user_id=current_user.id, action='STORE_UPLOAD', note=f'رفع ملف مستودع: {original_name}', target_type='STORE_FILE', target_id=row.id, created_at=datetime.utcnow()))
    except Exception:
        pass

    try:
        db.session.commit()
        flash("تم رفع الملف إلى المستودع.", "success")
    except Exception:
        db.session.rollback()
        # Cleanup file on failure
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception:
            pass
        flash("تعذر رفع الملف.", "danger")

    return redirect(url_for("portal.store_home"))


@portal_bp.route("/store/files/<int:file_id>/view")
@login_required
@_perm(PORTAL_READ)
def store_file_view(file_id: int):
    _ensure_store_ready()
    row = StoreFile.query.get_or_404(file_id)
    if row.is_deleted:
        abort(404)
    if not _store_can_access_file(row):
        abort(403)

    # Inline only for safe types
    ext = (Path(row.original_name).suffix or "").lower()
    inline_ok = ext in {".pdf", ".png", ".jpg", ".jpeg"}
    if not inline_ok:
        return redirect(url_for("portal.store_file_download", file_id=file_id))

    return send_file(
        row.file_path,
        mimetype=row.mime_type or None,
        as_attachment=False,
        download_name=row.original_name,
    )


@portal_bp.route("/store/files/<int:file_id>/download")
@login_required
@_perm(PORTAL_READ)
def store_file_download(file_id: int):
    _ensure_store_ready()
    row = StoreFile.query.get_or_404(file_id)
    if row.is_deleted:
        abort(404)
    if not _store_can_access_file(row):
        abort(403)
    if not _store_can_download_file(row):
        abort(403)

    return send_file(
        row.file_path,
        mimetype=row.mime_type or None,
        as_attachment=True,
        download_name=row.original_name,
    )


@portal_bp.route("/store/files/<int:file_id>/delete", methods=["POST"])
@login_required
@_perm("STORE_MANAGE")
def store_file_delete(file_id: int):
    _ensure_store_ready()
    row = StoreFile.query.get_or_404(file_id)
    if row.is_deleted:
        flash("الملف محذوف مسبقاً.", "info")
        return redirect(url_for("portal.store_home"))

    row.is_deleted = True
    row.deleted_at = datetime.utcnow()
    row.deleted_by = current_user.id
    try:
        db.session.add(AuditLog(user_id=current_user.id, action='STORE_FILE_DELETE', note=f'حذف ملف مستودع: {row.original_name}', target_type='STORE_FILE', target_id=row.id, created_at=datetime.utcnow()))
    except Exception:
        pass
    try:
        db.session.commit()
        flash("تم حذف الملف.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر حذف الملف.", "danger")
    return redirect(url_for("portal.store_home"))


@portal_bp.route("/store/categories", methods=["GET", "POST"])
@login_required
@_perm("STORE_MANAGE")
def store_categories():
    _ensure_store_ready()
    if request.method == "POST":
        name = (request.form.get("name") or "").strip()
        parent_id = (request.form.get("parent_id") or "").strip()
        if not name:
            flash("اكتب اسم التصنيف.", "danger")
            return redirect(url_for("portal.store_categories"))
        name = name[:120]
        if StoreCategory.query.filter_by(name=name).first():
            flash("هذا التصنيف موجود.", "warning")
            return redirect(url_for("portal.store_categories"))
        pid = None
        if parent_id:
            try:
                pid = int(parent_id)
            except Exception:
                pid = None
        db.session.add(StoreCategory(name=name, is_active=True, parent_id=pid))
        db.session.commit()
        flash("تم إضافة التصنيف.", "success")
        return redirect(url_for("portal.store_categories"))

    cats = StoreCategory.query.order_by(StoreCategory.is_active.desc(), StoreCategory.name.asc()).all()
    category_options = _store_category_options()
    # Excel export
    if (request.args.get("export") or "").strip() in {"1", "excel"}:
        data = []
        for c in cats:
            data.append({
                "id": c.id,
                "name": c.name,
                "parent": (c.parent.name if c.parent else ""),
                "is_active": "نعم" if c.is_active else "لا",
            })

        xlsx_bytes = _xlsx_from_dicts(
            data,
            columns=[
                ("id", "ID"),
                ("name", "الاسم"),
                ("parent", "التصنيف الأب"),
                ("is_active", "نشط"),
            ],
            sheet_name="Store Categories",
        )

        return send_file(
            io.BytesIO(xlsx_bytes),
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name="store_categories.xlsx",
        )

    return render_template("portal/store/categories.html", cats=cats, category_options=category_options)


@portal_bp.route("/store/categories/<int:cat_id>/toggle", methods=["POST"])
@login_required
@_perm("STORE_MANAGE")
def store_category_toggle(cat_id: int):
    _ensure_store_ready()
    c = StoreCategory.query.get_or_404(cat_id)
    c.is_active = not bool(c.is_active)
    db.session.commit()
    flash("تم تحديث حالة التصنيف.", "success")
    return redirect(url_for("portal.store_categories"))

# -------------------------
# Saved Filters (Portal)
# -------------------------
@portal_bp.route("/filters/save", methods=["POST"])
@login_required
@_perm(PORTAL_READ)
def portal_save_filter():
    name = (request.form.get("name") or "").strip()
    scope = (request.form.get("scope") or "").strip().upper()
    path = (request.form.get("path") or "").strip()
    qs = (request.form.get("query_string") or "").strip().lstrip("?")

    if not name or not scope or not path:
        flash("اسم الفلتر والمجال مطلوبان.", "danger")
        return redirect(request.referrer or url_for("portal.index"))

    # prevent open redirects: keep only local paths
    if not path.startswith("/"):
        path = "/" + path
    if path.startswith("//") or "://" in path:
        flash("مسار غير صالح.", "danger")
        return redirect(request.referrer or url_for("portal.index"))

    # Normalize
    name = name[:120]
    if len(qs) > 4000:
        qs = qs[:4000]

    existing = SavedFilter.query.filter_by(owner_id=current_user.id, scope=scope, name=name).first()
    if existing:
        existing.path = path
        existing.query_string = qs
    else:
        db.session.add(SavedFilter(
            owner_id=current_user.id,
            scope=scope,
            name=name,
            path=path,
            query_string=qs,
            created_at=datetime.utcnow(),
        ))
    db.session.commit()
    flash("تم حفظ الفلتر.", "success")
    return redirect(path + (("?" + qs) if qs else ""))


@portal_bp.route("/filters/<int:filter_id>/apply")
@login_required
@_perm(PORTAL_READ)
def portal_apply_filter(filter_id: int):
    f = SavedFilter.query.get_or_404(filter_id)
    if int(f.owner_id) != int(current_user.id):
        abort(403)
    return redirect(f.path + (("?" + (f.query_string or "")) if (f.query_string or "") else ""))


@portal_bp.route("/filters/<int:filter_id>/delete", methods=["POST"])
@login_required
@_perm(PORTAL_READ)
def portal_delete_filter(filter_id: int):
    f = SavedFilter.query.get_or_404(filter_id)
    if int(f.owner_id) != int(current_user.id):
        abort(403)
    back = request.form.get("back") or f.path
    try:
        db.session.delete(f)
        db.session.commit()
        flash("تم حذف الفلتر.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر حذف الفلتر.", "danger")
    return redirect(back)


# -------------------------
# HR: Employees (Files)
# -------------------------


def _portal_audit(
    action: str,
    note: str = "",
    target_type: str | None = None,
    target_id: int | None = None,
    user_id: int | None = None,
):
    """Audit helper that is safe to call without an active request context.

    - If called within a request and the user is authenticated, we record delegation context.
    - If called from a background job / script, pass user_id explicitly.
    """
    base_user_id = user_id
    on_behalf_of_id = None
    delegation_id = None

    if base_user_id is None and has_request_context():
        try:
            if current_user.is_authenticated:
                base_user_id = int(current_user.id)
        except Exception:
            base_user_id = None

    # Background fallback: use configured system user id if present.
    if base_user_id is None:
        try:
            v = _setting_get("TIMECLK_IMPORTED_BY_USER_ID")
            base_user_id = int(v) if v and str(v).strip().isdigit() else None
        except Exception:
            base_user_id = None

    # Try to capture delegation context only within a request.
    if has_request_context():
        try:
            from utils.permissions import get_effective_user, get_active_delegation

            eff = get_effective_user()
            d = get_active_delegation()
            if current_user.is_authenticated and eff and getattr(eff, "id", None) and int(eff.id) != int(current_user.id):
                on_behalf_of_id = int(eff.id)
                if d and getattr(d, "id", None):
                    delegation_id = int(d.id)
        except Exception:
            pass

    # If we still can't determine a user, skip silently (avoid crashing background jobs).
    if base_user_id is None:
        return

    db.session.add(AuditLog(
        action=action,
        note=note,
        user_id=base_user_id,
        on_behalf_of_id=on_behalf_of_id,
        delegation_id=delegation_id,
        target_type=target_type,
        target_id=target_id,
        created_at=datetime.utcnow(),
    ))


@portal_bp.route("/hr/employees")
@login_required
@_perm(HR_EMP_READ)
def hr_employees():
    q = (request.args.get("q") or "").strip()

    qry = (
        User.query
        .outerjoin(EmployeeFile, EmployeeFile.user_id == User.id)
    )

    if q:
        qry = apply_search_all_columns(
            qry,
            User,
            q,
            exclude_columns={"password_hash"},
            extra_columns=[EmployeeFile.employee_no, EmployeeFile.timeclock_code],
        )

    users = qry.order_by(User.name.asc().nullslast(), User.email.asc()).limit(500).all()

    user_ids = [u.id for u in users] if users else []
    files = {
        f.user_id: f
        for f in EmployeeFile.query.filter(EmployeeFile.user_id.in_(user_ids)).all()
    } if user_ids else {}


    export = (request.args.get("export") or "").strip().lower()
    if export in {"1", "true", "excel", "xlsx"}:
        rows = []
        for u in users:
            f = files.get(u.id) if files else None
            rows.append({
                "id": u.id,
                "name": getattr(u, 'full_name', None) or getattr(u, 'name', None) or "",
                "email": u.email or "",
                "role": u.role or "",
                "employee_no": (getattr(f, 'employee_no', '') or ''),
                "timeclock_code": (getattr(f, 'timeclock_code', '') or ''),
            })
        xlsx_bytes = _xlsx_from_dicts(
            rows,
            columns=[
                ("id", "ID"),
                ("name", "الموظف"),
                ("email", "البريد الإلكتروني"),
                ("role", "الدور"),
                ("employee_no", "رقم الموظف"),
                ("timeclock_code", "كود الساعة"),
            ],
            sheet_name="Employees",
        )
        return send_file(
            io.BytesIO(xlsx_bytes),
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name="hr_employees.xlsx",
        )

    return render_template("portal/hr/employees.html", users=users, q=q, files=files)




def _employee_upload_dir(user_id: int) -> Path:
    base = Path(current_app.instance_path) / "uploads" / "employees" / str(user_id)
    base.mkdir(parents=True, exist_ok=True)
    return base


# -------------------------
# HR: Employee File - Lookups
# -------------------------

# Fixed set of lookup categories used by the employee file.
EMP_LOOKUP_CATEGORIES = [
    ("IDENTITY_TYPE", "نوع الهوية"),
    ("GENDER", "الجنس"),
    ("MARITAL_STATUS", "الحالة الاجتماعية"),
    ("RELIGION", "الديانة"),
    ("DISABILITY", "الإعاقة"),
    ("HOME_GOV", "محافظة السكن"),
    ("LOCALITY", "التجمع السكاني"),
    ("WORK_GOV", "محافظة العمل"),
    ("WORK_LOCATION", "موقع العمل"),
    ("EMP_STATUS", "حالة الموظف"),
    ("SHIFT", "الوردية"),
    ("PROJECT", "المشروع"),
    ("APPOINTMENT_TYPE", "نوع التعيين"),
    ("JOB_CATEGORY", "الفئة"),
    ("JOB_GRADE", "الدرجة"),
    ("JOB_TITLE", "المسمى الوظيفي"),
    ("ADMIN_TITLE", "المسمى الإداري"),
    ("BANK", "البنك"),
    ("DEP_RELATION", "العلاقة"),
    ("QUAL_DEGREE", "الدرجة العلمية"),
    ("QUAL_SPECIALIZATION", "التخصص"),
    ("QUAL_GRADE", "التقدير"),
    ("UNIVERSITY", "الجامعة"),
    ("COUNTRY", "الدولة"),
    ("ATTACH_TYPE", "نوع المرفق"),
]

EMP_LOOKUP_LABEL = {k: v for k, v in EMP_LOOKUP_CATEGORIES}


def _lookup_items(category: str, *, active_only: bool = True):
    q = HRLookupItem.query.filter_by(category=(category or '').upper().strip())
    if active_only:
        q = q.filter(HRLookupItem.is_active.is_(True))
    return q.order_by(HRLookupItem.sort_order.asc(), HRLookupItem.name_ar.asc(), HRLookupItem.id.asc()).all()


def _lookup_label_by_id() -> dict[int, str]:
    rows = HRLookupItem.query.all()
    return {r.id: r.label for r in rows if r and r.id}


def _to_int(v):
    try:
        s = (str(v) if v is not None else '').strip()
        return int(s) if s else None
    except Exception:
        return None


def _to_float(v):
    try:
        s = (str(v) if v is not None else '').strip()
        return float(s) if s else None
    except Exception:
        return None


def _to_str(v):
    s = (v or '').strip() if isinstance(v, str) else (str(v).strip() if v is not None else '')
    return s or None


def _employee_nav_counts(user_id: int) -> dict:
    try:
        deps = EmployeeDependent.query.filter_by(user_id=user_id).count()
    except Exception:
        deps = 0
    try:
        quals = EmployeeQualification.query.filter_by(user_id=user_id).count()
    except Exception:
        quals = 0
    try:
        secs = EmployeeSecondment.query.filter_by(user_id=user_id).count()
    except Exception:
        secs = 0
    try:
        atts = EmployeeAttachment.query.filter_by(user_id=user_id).filter(EmployeeAttachment.attachment_type != 'PAYSLIP').count()
    except Exception:
        atts = 0
    try:
        slips = EmployeeAttachment.query.filter_by(user_id=user_id, attachment_type='PAYSLIP').count()
    except Exception:
        slips = 0
    return {"deps": deps, "quals": quals, "secondments": secs, "atts": atts, "payslips": slips}


@portal_bp.route("/hr/employees/<int:user_id>", methods=["GET", "POST"])
@login_required
@_perm(HR_EMP_READ)
def hr_employee_file(user_id: int):
    """Employee file (page 1: basic information).

    Note: other categories are separate pages to keep the UX clean.
    """
    u = User.query.get_or_404(user_id)

    emp_file = EmployeeFile.query.filter_by(user_id=user_id).first()

    if request.method == "POST":
        if not current_user.has_perm(HR_EMP_MANAGE):
            abort(403)

        if not emp_file:
            emp_file = EmployeeFile(user_id=user_id, created_at=datetime.utcnow(), updated_at=datetime.utcnow())
            db.session.add(emp_file)

        # --- attendance mapping (timeclock) ---
        tc = _to_str(request.form.get("timeclock_code"))
        if tc:
            if (not tc.isdigit()) or len(tc) != 9:
                flash("كود الساعة يجب أن يكون 9 أرقام.", "danger")
                return redirect(request.url)
            emp_file.timeclock_code = tc
        else:
            emp_file.timeclock_code = None

        # header
        emp_file.employee_no = _to_str(request.form.get("employee_no"))
        emp_file.full_name_quad = _to_str(request.form.get("full_name_quad"))
        if emp_file.full_name_quad:
            u.name = emp_file.full_name_quad

        # (1) basic - personal
        emp_file.identity_type_lookup_id = _to_int(request.form.get("identity_type_lookup_id"))
        emp_file.national_id = _to_str(request.form.get("national_id"))
        emp_file.gender_lookup_id = _to_int(request.form.get("gender_lookup_id"))
        emp_file.marital_status_lookup_id = _to_int(request.form.get("marital_status_lookup_id"))
        emp_file.birth_date = _to_str(request.form.get("birth_date"))
        emp_file.religion_lookup_id = _to_int(request.form.get("religion_lookup_id"))
        emp_file.disability_lookup_id = _to_int(request.form.get("disability_lookup_id"))
        emp_file.home_governorate_lookup_id = _to_int(request.form.get("home_governorate_lookup_id"))
        emp_file.locality_lookup_id = _to_int(request.form.get("locality_lookup_id"))
        emp_file.address = _to_str(request.form.get("address"))
        emp_file.phone = _to_str(request.form.get("phone"))
        emp_file.mobile = _to_str(request.form.get("mobile"))
        emp_file.email = _to_str(request.form.get("email"))

        # (1) basic - work
        emp_file.work_governorate_lookup_id = _to_int(request.form.get("work_governorate_lookup_id"))
        emp_file.work_location_lookup_id = _to_int(request.form.get("work_location_lookup_id"))
        emp_file.employee_status_lookup_id = _to_int(request.form.get("employee_status_lookup_id"))
        emp_file.status_date = _to_str(request.form.get("status_date"))
        emp_file.status_note = _to_str(request.form.get("status_note"))
        emp_file.shift_lookup_id = _to_int(request.form.get("shift_lookup_id"))
        emp_file.hourly_number = _to_float(request.form.get("hourly_number"))

        # (1) placement
        emp_file.organization_id = _to_int(request.form.get("organization_id"))
        emp_file.directorate_id = _to_int(request.form.get("directorate_id"))
        emp_file.department_id = _to_int(request.form.get("department_id"))
        emp_file.division_id = _to_int(request.form.get("division_id"))
        emp_file.direct_manager_user_id = _to_int(request.form.get("direct_manager_user_id"))

        emp_file.project_lookup_id = _to_int(request.form.get("project_lookup_id"))
        emp_file.appointment_type_lookup_id = _to_int(request.form.get("appointment_type_lookup_id"))
        emp_file.hire_date = _to_str(request.form.get("hire_date"))
        emp_file.last_promotion_date = _to_str(request.form.get("last_promotion_date"))

        emp_file.job_category_lookup_id = _to_int(request.form.get("job_category_lookup_id"))
        emp_file.job_grade_lookup_id = _to_int(request.form.get("job_grade_lookup_id"))
        emp_file.job_title_lookup_id = _to_int(request.form.get("job_title_lookup_id"))
        emp_file.admin_title_lookup_id = _to_int(request.form.get("admin_title_lookup_id"))

        # (1) bank
        emp_file.bank_lookup_id = _to_int(request.form.get("bank_lookup_id"))
        emp_file.bank_account = _to_str(request.form.get("bank_account"))

        emp_file.notes = _to_str(request.form.get("notes"))

        emp_file.updated_at = datetime.utcnow()
        emp_file.updated_by_id = current_user.id

        _portal_audit(
            action="HR_EMPLOYEE_UPDATE",
            note=f"تحديث ملف الموظف: {u.email}",
            target_type="EMPLOYEE_FILE",
            target_id=user_id,
        )

        try:
            db.session.commit()
            flash("تم حفظ ملف الموظف.", "success")
        except Exception:
            db.session.rollback()
            flash("تعذر حفظ ملف الموظف.", "danger")

        return redirect(request.url)

    # nav counts
    counts = _employee_nav_counts(user_id)

    # lookups
    lookups = {
        "IDENTITY_TYPE": _lookup_items("IDENTITY_TYPE"),
        "GENDER": _lookup_items("GENDER"),
        "MARITAL_STATUS": _lookup_items("MARITAL_STATUS"),
        "RELIGION": _lookup_items("RELIGION"),
        "DISABILITY": _lookup_items("DISABILITY"),
        "HOME_GOV": _lookup_items("HOME_GOV"),
        "LOCALITY": _lookup_items("LOCALITY"),
        "WORK_GOV": _lookup_items("WORK_GOV"),
        "WORK_LOCATION": _lookup_items("WORK_LOCATION"),
        "EMP_STATUS": _lookup_items("EMP_STATUS"),
        "SHIFT": _lookup_items("SHIFT"),
        "PROJECT": _lookup_items("PROJECT"),
        "APPOINTMENT_TYPE": _lookup_items("APPOINTMENT_TYPE"),
        "JOB_CATEGORY": _lookup_items("JOB_CATEGORY"),
        "JOB_GRADE": _lookup_items("JOB_GRADE"),
        "JOB_TITLE": _lookup_items("JOB_TITLE"),
        "ADMIN_TITLE": _lookup_items("ADMIN_TITLE"),
        "BANK": _lookup_items("BANK"),
    }

    lookup_by_id = _lookup_label_by_id()

    # org structure dropdowns
    orgs = Organization.query.filter(Organization.is_active.is_(True)).order_by(Organization.name_ar.asc()).all()
    dirs = Directorate.query.filter(Directorate.is_active.is_(True)).order_by(Directorate.name_ar.asc()).all()
    depts = Department.query.filter(Department.is_active.is_(True)).order_by(Department.name_ar.asc()).all()
    divs = Division.query.filter(Division.is_active.is_(True)).order_by(Division.name_ar.asc()).all()

    managers = User.query.order_by(User.name.asc().nullslast(), User.email.asc()).limit(3000).all()

    logs = (
        AuditLog.query
        .filter(AuditLog.target_type.in_(['EMPLOYEE_FILE', 'EMPLOYEE_ATTACHMENT']), AuditLog.target_id == user_id)
        .order_by(AuditLog.created_at.desc(), AuditLog.id.desc())
        .limit(20)
        .all()
    )

    return render_template(
        "portal/hr/employee/basic.html",
        u=u,
        emp_file=emp_file,
        counts=counts,
        lookups=lookups,
        lookup_by_id=lookup_by_id,
        orgs=orgs,
        dirs=dirs,
        depts=depts,
        divs=divs,
        managers=managers,
        logs=logs,
    )


@portal_bp.route("/hr/employees/<int:user_id>/dependents")
@login_required
@_perm(HR_EMP_READ)
def hr_employee_dependents(user_id: int):
    u = User.query.get_or_404(user_id)
    emp_file = EmployeeFile.query.filter_by(user_id=user_id).first()
    counts = _employee_nav_counts(user_id)

    items = EmployeeDependent.query.filter_by(user_id=user_id).order_by(EmployeeDependent.id.desc()).all()

    # lookups
    lookup_by_id = _lookup_label_by_id()
    lookups = {
        "DEP_RELATION": _lookup_items("DEP_RELATION"),
        "GENDER": _lookup_items("GENDER"),
    }

    return render_template(
        "portal/hr/employee/dependents.html",
        u=u,
        emp_file=emp_file,
        counts=counts,
        items=items,
        lookup_label=lookup_by_id,
        lookups=lookups,
        lookup_by_id=lookup_by_id,
    )

@portal_bp.route("/hr/employees/<int:user_id>/dependents/new", methods=["GET", "POST"])
@login_required
@_perm(HR_EMP_READ)
def hr_employee_dependent_new(user_id: int):
    u = User.query.get_or_404(user_id)
    emp_file = EmployeeFile.query.filter_by(user_id=user_id).first()
    counts = _employee_nav_counts(user_id)

    lookups = {
        "DEP_RELATION": _lookup_items("DEP_RELATION"),
        "GENDER": _lookup_items("GENDER"),
    }

    if request.method == "POST":
        if not current_user.has_perm(HR_EMP_MANAGE):
            abort(403)

        d = EmployeeDependent(
            user_id=user_id,
            full_name=_to_str(request.form.get('full_name')),
            relation_lookup_id=_to_int(request.form.get('relation_lookup_id')),
            national_id=_to_str(request.form.get('national_id')),
            gender_lookup_id=_to_int(request.form.get('gender_lookup_id')),
            birth_date=_to_str(request.form.get('birth_date')),
            allowance=_to_float(request.form.get('allowance')),
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow(),
            updated_by_id=current_user.id,
        )
        db.session.add(d)
        _portal_audit(
            action="HR_EMPLOYEE_DEPENDENT_CREATE",
            note=f"إضافة تابع: {d.full_name or ''}",
            target_type="EMPLOYEE_FILE",
            target_id=user_id,
        )
        try:
            db.session.commit()
            flash("تمت إضافة البيانات.", "success")
            return redirect(url_for('portal.hr_employee_dependents', user_id=user_id))
        except Exception:
            db.session.rollback()
            flash("تعذر الحفظ.", "danger")

    return render_template(
        "portal/hr/employee/dependent_form.html",
        u=u,
        emp_file=emp_file,
        counts=counts,
        lookups=lookups,
        item=None,
        is_edit=False,
    )


@portal_bp.route("/hr/employees/<int:user_id>/dependents/<int:dep_id>/edit", methods=["GET", "POST"])
@login_required
@_perm(HR_EMP_READ)
def hr_employee_dependent_edit(user_id: int, dep_id: int):
    u = User.query.get_or_404(user_id)
    emp_file = EmployeeFile.query.filter_by(user_id=user_id).first()
    counts = _employee_nav_counts(user_id)

    dep = EmployeeDependent.query.filter_by(id=dep_id, user_id=user_id).first_or_404()

    lookups = {
        "DEP_RELATION": _lookup_items("DEP_RELATION"),
        "GENDER": _lookup_items("GENDER"),
    }

    if request.method == "POST":
        if not current_user.has_perm(HR_EMP_MANAGE):
            abort(403)

        dep.full_name = _to_str(request.form.get('full_name'))
        dep.relation_lookup_id = _to_int(request.form.get('relation_lookup_id'))
        dep.national_id = _to_str(request.form.get('national_id'))
        dep.gender_lookup_id = _to_int(request.form.get('gender_lookup_id'))
        dep.birth_date = _to_str(request.form.get('birth_date'))
        dep.allowance = _to_float(request.form.get('allowance'))
        dep.updated_at = datetime.utcnow()
        dep.updated_by_id = current_user.id

        _portal_audit(
            action="HR_EMPLOYEE_DEPENDENT_UPDATE",
            note=f"تعديل تابع: {dep.full_name or ''}",
            target_type="EMPLOYEE_FILE",
            target_id=user_id,
        )
        try:
            db.session.commit()
            flash("تم تحديث البيانات.", "success")
            return redirect(url_for('portal.hr_employee_dependents', user_id=user_id))
        except Exception:
            db.session.rollback()
            flash("تعذر الحفظ.", "danger")

    return render_template(
        "portal/hr/employee/dependent_form.html",
        u=u,
        emp_file=emp_file,
        counts=counts,
        lookups=lookups,
        item=dep,
        is_edit=True,
    )


@portal_bp.route("/hr/employees/<int:user_id>/dependents/<int:dep_id>/delete", methods=["POST"])
@login_required
@_perm(HR_EMP_READ)
def hr_employee_dependent_delete(user_id: int, dep_id: int):
    if not current_user.has_perm(HR_EMP_MANAGE):
        abort(403)

    dep = EmployeeDependent.query.filter_by(id=dep_id, user_id=user_id).first_or_404()
    db.session.delete(dep)
    _portal_audit(
        action="HR_EMPLOYEE_DEPENDENT_DELETE",
        note=f"حذف تابع: {dep.full_name or ''}",
        target_type="EMPLOYEE_FILE",
        target_id=user_id,
    )
    try:
        db.session.commit()
        flash("تم الحذف.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر الحذف.", "danger")
    return redirect(url_for('portal.hr_employee_dependents', user_id=user_id))


@portal_bp.route("/hr/employees/<int:user_id>/qualifications")
@login_required
@_perm(HR_EMP_READ)
def hr_employee_qualifications(user_id: int):
    u = User.query.get_or_404(user_id)
    emp_file = EmployeeFile.query.filter_by(user_id=user_id).first()
    counts = _employee_nav_counts(user_id)

    items = EmployeeQualification.query.filter_by(user_id=user_id).order_by(EmployeeQualification.id.desc()).all()
    lookup_by_id = _lookup_label_by_id()

    return render_template(
        "portal/hr/employee/qualifications.html",
        u=u,
        emp_file=emp_file,
        counts=counts,
        items=items,
        lookup_label=lookup_by_id,
        lookup_by_id=lookup_by_id,
    )

@portal_bp.route("/hr/employees/<int:user_id>/qualifications/new", methods=["GET", "POST"])
@login_required
@_perm(HR_EMP_READ)
def hr_employee_qualification_new(user_id: int):
    u = User.query.get_or_404(user_id)
    emp_file = EmployeeFile.query.filter_by(user_id=user_id).first()
    counts = _employee_nav_counts(user_id)

    lookups = {
        "QUAL_DEGREE": _lookup_items("QUAL_DEGREE"),
        "QUAL_SPECIALIZATION": _lookup_items("QUAL_SPECIALIZATION"),
        "QUAL_GRADE": _lookup_items("QUAL_GRADE"),
        "UNIVERSITY": _lookup_items("UNIVERSITY"),
        "COUNTRY": _lookup_items("COUNTRY"),
    }

    if request.method == "POST":
        if not current_user.has_perm(HR_EMP_MANAGE):
            abort(403)

        q = EmployeeQualification(
            user_id=user_id,
            degree_lookup_id=_to_int(request.form.get('degree_lookup_id')),
            specialization_lookup_id=_to_int(request.form.get('specialization_lookup_id')),
            grade_lookup_id=_to_int(request.form.get('grade_lookup_id')),
            qualification_date=_to_str(request.form.get('qualification_date')),
            university_lookup_id=_to_int(request.form.get('university_lookup_id')),
            country_lookup_id=_to_int(request.form.get('country_lookup_id')),
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow(),
            updated_by_id=current_user.id,
        )
        db.session.add(q)
        _portal_audit(
            action="HR_EMPLOYEE_QUAL_CREATE",
            note="إضافة مؤهل علمي",
            target_type="EMPLOYEE_FILE",
            target_id=user_id,
        )
        try:
            db.session.commit()
            flash("تمت إضافة المؤهل.", "success")
            return redirect(url_for('portal.hr_employee_qualifications', user_id=user_id))
        except Exception:
            db.session.rollback()
            flash("تعذر الحفظ.", "danger")

    return render_template(
        "portal/hr/employee/qualification_form.html",
        u=u,
        emp_file=emp_file,
        counts=counts,
        lookups=lookups,
        item=None,
        is_edit=False,
    )


@portal_bp.route("/hr/employees/<int:user_id>/qualifications/<int:qual_id>/edit", methods=["GET", "POST"])
@login_required
@_perm(HR_EMP_READ)
def hr_employee_qualification_edit(user_id: int, qual_id: int):
    u = User.query.get_or_404(user_id)
    emp_file = EmployeeFile.query.filter_by(user_id=user_id).first()
    counts = _employee_nav_counts(user_id)

    qual = EmployeeQualification.query.filter_by(id=qual_id, user_id=user_id).first_or_404()

    lookups = {
        "QUAL_DEGREE": _lookup_items("QUAL_DEGREE"),
        "QUAL_SPECIALIZATION": _lookup_items("QUAL_SPECIALIZATION"),
        "QUAL_GRADE": _lookup_items("QUAL_GRADE"),
        "UNIVERSITY": _lookup_items("UNIVERSITY"),
        "COUNTRY": _lookup_items("COUNTRY"),
    }

    if request.method == "POST":
        if not current_user.has_perm(HR_EMP_MANAGE):
            abort(403)

        qual.degree_lookup_id = _to_int(request.form.get('degree_lookup_id'))
        qual.specialization_lookup_id = _to_int(request.form.get('specialization_lookup_id'))
        qual.grade_lookup_id = _to_int(request.form.get('grade_lookup_id'))
        qual.qualification_date = _to_str(request.form.get('qualification_date'))
        qual.university_lookup_id = _to_int(request.form.get('university_lookup_id'))
        qual.country_lookup_id = _to_int(request.form.get('country_lookup_id'))
        qual.updated_at = datetime.utcnow()
        qual.updated_by_id = current_user.id

        _portal_audit(
            action="HR_EMPLOYEE_QUAL_UPDATE",
            note="تعديل مؤهل علمي",
            target_type="EMPLOYEE_FILE",
            target_id=user_id,
        )
        try:
            db.session.commit()
            flash("تم تحديث المؤهل.", "success")
            return redirect(url_for('portal.hr_employee_qualifications', user_id=user_id))
        except Exception:
            db.session.rollback()
            flash("تعذر الحفظ.", "danger")

    return render_template(
        "portal/hr/employee/qualification_form.html",
        u=u,
        emp_file=emp_file,
        counts=counts,
        lookups=lookups,
        item=qual,
        is_edit=True,
    )


@portal_bp.route("/hr/employees/<int:user_id>/qualifications/<int:qual_id>/delete", methods=["POST"])
@login_required
@_perm(HR_EMP_READ)
def hr_employee_qualification_delete(user_id: int, qual_id: int):
    if not current_user.has_perm(HR_EMP_MANAGE):
        abort(403)

    qual = EmployeeQualification.query.filter_by(id=qual_id, user_id=user_id).first_or_404()
    db.session.delete(qual)

    _portal_audit(
        action="HR_EMPLOYEE_QUAL_DELETE",
        note="حذف مؤهل علمي",
        target_type="EMPLOYEE_FILE",
        target_id=user_id,
    )

    try:
        db.session.commit()
        flash("تم الحذف.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر الحذف.", "danger")

    return redirect(url_for('portal.hr_employee_qualifications', user_id=user_id))


@portal_bp.route("/hr/employees/<int:user_id>/secondments")
@login_required
@_perm(HR_EMP_READ)
def hr_employee_secondments(user_id: int):
    u = User.query.get_or_404(user_id)
    emp_file = EmployeeFile.query.filter_by(user_id=user_id).first()
    counts = _employee_nav_counts(user_id)

    items = EmployeeSecondment.query.filter_by(user_id=user_id).order_by(EmployeeSecondment.id.desc()).all()

    lookup_by_id = _lookup_label_by_id()

    # org structure dropdowns + maps
    orgs = Organization.query.filter(Organization.is_active.is_(True)).order_by(Organization.name_ar.asc()).all()
    dirs = Directorate.query.filter(Directorate.is_active.is_(True)).order_by(Directorate.name_ar.asc()).all()
    depts = Department.query.filter(Department.is_active.is_(True)).order_by(Department.name_ar.asc()).all()
    divs = Division.query.filter(Division.is_active.is_(True)).order_by(Division.name_ar.asc()).all()

    managers = User.query.order_by(User.name.asc().nullslast(), User.email.asc()).limit(3000).all()

    org_map = {o.id: o.name_ar for o in orgs}
    dir_map = {d.id: d.name_ar for d in dirs}
    dept_map = {d.id: d.name_ar for d in depts}
    div_map = {v.id: v.name_ar for v in divs}
    manager_map = {m.id: (m.name or m.email) for m in managers}

    return render_template(
        "portal/hr/employee/secondments.html",
        u=u,
        emp_file=emp_file,
        counts=counts,
        items=items,
        lookup_label=lookup_by_id,
        org_map=org_map,
        dir_map=dir_map,
        dept_map=dept_map,
        div_map=div_map,
        manager_map=manager_map,
    )

@portal_bp.route("/hr/employees/<int:user_id>/secondments/new", methods=["GET", "POST"])
@login_required
@_perm(HR_EMP_READ)
def hr_employee_secondment_new(user_id: int):
    u = User.query.get_or_404(user_id)
    emp_file = EmployeeFile.query.filter_by(user_id=user_id).first()
    counts = _employee_nav_counts(user_id)

    orgs = Organization.query.filter(Organization.is_active.is_(True)).order_by(Organization.name_ar.asc()).all()
    dirs = Directorate.query.filter(Directorate.is_active.is_(True)).order_by(Directorate.name_ar.asc()).all()
    depts = Department.query.filter(Department.is_active.is_(True)).order_by(Department.name_ar.asc()).all()
    divs = Division.query.filter(Division.is_active.is_(True)).order_by(Division.name_ar.asc()).all()
    managers = User.query.order_by(User.name.asc().nullslast(), User.email.asc()).limit(3000).all()

    lookups = {
        "WORK_GOV": _lookup_items("WORK_GOV"),
        "WORK_LOCATION": _lookup_items("WORK_LOCATION"),
        "ADMIN_TITLE": _lookup_items("ADMIN_TITLE"),
    }

    if request.method == "POST":
        if not current_user.has_perm(HR_EMP_MANAGE):
            abort(403)

        s = EmployeeSecondment(
            user_id=user_id,
            date_from=_to_str(request.form.get('date_from')),
            date_to=_to_str(request.form.get('date_to')),
            organization_id=_to_int(request.form.get('organization_id')),
            directorate_id=_to_int(request.form.get('directorate_id')),
            department_id=_to_int(request.form.get('department_id')),
            division_id=_to_int(request.form.get('division_id')),
            direct_manager_user_id=_to_int(request.form.get('direct_manager_user_id')),
            work_governorate_lookup_id=_to_int(request.form.get('work_governorate_lookup_id')),
            work_location_lookup_id=_to_int(request.form.get('work_location_lookup_id')),
            admin_title_lookup_id=_to_int(request.form.get('admin_title_lookup_id')),
            details=_to_str(request.form.get('details')),
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow(),
            updated_by_id=current_user.id,
        )
        db.session.add(s)

        _portal_audit(
            action="HR_EMPLOYEE_SECONDMENT_CREATE",
            note="إضافة تكليف",
            target_type="EMPLOYEE_FILE",
            target_id=user_id,
        )

        try:
            db.session.commit()
            flash("تمت إضافة التكليف.", "success")
            return redirect(url_for('portal.hr_employee_secondments', user_id=user_id))
        except Exception:
            db.session.rollback()
            flash("تعذر الحفظ.", "danger")

    return render_template(
        "portal/hr/employee/secondment_form.html",
        u=u,
        emp_file=emp_file,
        counts=counts,
        orgs=orgs,
        dirs=dirs,
        depts=depts,
        divs=divs,
        managers=managers,
        lookups=lookups,
        item=None,
        is_edit=False,
    )


@portal_bp.route("/hr/employees/<int:user_id>/secondments/<int:sec_id>/edit", methods=["GET", "POST"])
@login_required
@_perm(HR_EMP_READ)
def hr_employee_secondment_edit(user_id: int, sec_id: int):
    u = User.query.get_or_404(user_id)
    emp_file = EmployeeFile.query.filter_by(user_id=user_id).first()
    counts = _employee_nav_counts(user_id)

    secondment = EmployeeSecondment.query.filter_by(id=sec_id, user_id=user_id).first_or_404()

    orgs = Organization.query.filter(Organization.is_active.is_(True)).order_by(Organization.name_ar.asc()).all()
    dirs = Directorate.query.filter(Directorate.is_active.is_(True)).order_by(Directorate.name_ar.asc()).all()
    depts = Department.query.filter(Department.is_active.is_(True)).order_by(Department.name_ar.asc()).all()
    divs = Division.query.filter(Division.is_active.is_(True)).order_by(Division.name_ar.asc()).all()
    managers = User.query.order_by(User.name.asc().nullslast(), User.email.asc()).limit(3000).all()

    lookups = {
        "WORK_GOV": _lookup_items("WORK_GOV"),
        "WORK_LOCATION": _lookup_items("WORK_LOCATION"),
        "ADMIN_TITLE": _lookup_items("ADMIN_TITLE"),
    }

    if request.method == "POST":
        if not current_user.has_perm(HR_EMP_MANAGE):
            abort(403)

        secondment.date_from = _to_str(request.form.get('date_from'))
        secondment.date_to = _to_str(request.form.get('date_to'))
        secondment.organization_id = _to_int(request.form.get('organization_id'))
        secondment.directorate_id = _to_int(request.form.get('directorate_id'))
        secondment.department_id = _to_int(request.form.get('department_id'))
        secondment.division_id = _to_int(request.form.get('division_id'))
        secondment.direct_manager_user_id = _to_int(request.form.get('direct_manager_user_id'))
        secondment.work_governorate_lookup_id = _to_int(request.form.get('work_governorate_lookup_id'))
        secondment.work_location_lookup_id = _to_int(request.form.get('work_location_lookup_id'))
        secondment.admin_title_lookup_id = _to_int(request.form.get('admin_title_lookup_id'))
        secondment.details = _to_str(request.form.get('details'))

        secondment.updated_at = datetime.utcnow()
        secondment.updated_by_id = current_user.id

        _portal_audit(
            action="HR_EMPLOYEE_SECONDMENT_UPDATE",
            note="تعديل تكليف",
            target_type="EMPLOYEE_FILE",
            target_id=user_id,
        )

        try:
            db.session.commit()
            flash("تم التحديث.", "success")
            return redirect(url_for('portal.hr_employee_secondments', user_id=user_id))
        except Exception:
            db.session.rollback()
            flash("تعذر الحفظ.", "danger")

    return render_template(
        "portal/hr/employee/secondment_form.html",
        u=u,
        emp_file=emp_file,
        counts=counts,
        orgs=orgs,
        dirs=dirs,
        depts=depts,
        divs=divs,
        managers=managers,
        lookups=lookups,
        item=secondment,
        is_edit=True,
    )


@portal_bp.route("/hr/employees/<int:user_id>/secondments/<int:sec_id>/delete", methods=["POST"])
@login_required
@_perm(HR_EMP_READ)
def hr_employee_secondment_delete(user_id: int, sec_id: int):
    if not current_user.has_perm(HR_EMP_MANAGE):
        abort(403)

    sec = EmployeeSecondment.query.filter_by(id=sec_id, user_id=user_id).first_or_404()
    db.session.delete(sec)

    _portal_audit(
        action="HR_EMPLOYEE_SECONDMENT_DELETE",
        note="حذف تكليف",
        target_type="EMPLOYEE_FILE",
        target_id=user_id,
    )

    try:
        db.session.commit()
        flash("تم الحذف.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر الحذف.", "danger")

    return redirect(url_for('portal.hr_employee_secondments', user_id=user_id))


@portal_bp.route("/hr/employees/<int:user_id>/attachments")
@login_required
@_perm(HR_EMP_READ)
def hr_employee_attachments(user_id: int):
    u = User.query.get_or_404(user_id)
    emp_file = EmployeeFile.query.filter_by(user_id=user_id).first()
    counts = _employee_nav_counts(user_id)

    lookup_by_id = _lookup_label_by_id()
    lookups = {"ATTACH_TYPE": _lookup_items("ATTACH_TYPE")}

    items = (
        EmployeeAttachment.query
        .filter_by(user_id=user_id)
        .filter(EmployeeAttachment.attachment_type != 'PAYSLIP')
        .order_by(EmployeeAttachment.uploaded_at.desc())
        .all()
    )

    return render_template(
        "portal/hr/employee/attachments.html",
        u=u,
        emp_file=emp_file,
        counts=counts,
        items=items,
        lookups=lookups,
        lookup_label=lookup_by_id,
        lookup_by_id=lookup_by_id,
    )

@portal_bp.route("/hr/employees/<int:user_id>/payslips")
@login_required
@_perm(HR_EMP_READ)
def hr_employee_payslips(user_id: int):
    u = User.query.get_or_404(user_id)
    emp_file = EmployeeFile.query.filter_by(user_id=user_id).first()
    counts = _employee_nav_counts(user_id)

    payslips = (
        EmployeeAttachment.query
        .filter_by(user_id=user_id, attachment_type='PAYSLIP')
        .order_by(EmployeeAttachment.payslip_year.desc().nullslast(), EmployeeAttachment.payslip_month.desc().nullslast(), EmployeeAttachment.uploaded_at.desc())
        .all()
    )

    return render_template(
        "portal/hr/employee/payslips.html",
        u=u,
        emp_file=emp_file,
        counts=counts,
        items=payslips,
    )


@portal_bp.route("/hr/employees/<int:user_id>/payslips/<int:att_id>/publish", methods=["POST"])
@login_required
@_perm(HR_EMP_ATTACH)
def hr_employee_payslip_publish(user_id: int, att_id: int):
    att = EmployeeAttachment.query.filter_by(id=att_id, user_id=user_id, attachment_type='PAYSLIP').first_or_404()
    att.is_published = True
    att.published_at = datetime.utcnow()
    att.published_by_id = current_user.id
    _portal_audit(
        action="HR_PAYSLIP_PUBLISH",
        note=f"نشر قسيمة راتب: {att.payslip_period_label}",
        target_type="EMPLOYEE_ATTACHMENT",
        target_id=user_id,
    )
    try:
        db.session.commit()
        flash("تم نشر القسيمة.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر النشر.", "danger")
    return redirect(url_for('portal.hr_employee_payslips', user_id=user_id))


@portal_bp.route("/hr/employees/<int:user_id>/payslips/<int:att_id>/unpublish", methods=["POST"])
@login_required
@_perm(HR_EMP_ATTACH)
def hr_employee_payslip_unpublish(user_id: int, att_id: int):
    att = EmployeeAttachment.query.filter_by(id=att_id, user_id=user_id, attachment_type='PAYSLIP').first_or_404()
    att.is_published = False
    att.published_at = None
    att.published_by_id = None
    _portal_audit(
        action="HR_PAYSLIP_UNPUBLISH",
        note=f"إلغاء نشر قسيمة راتب: {att.payslip_period_label}",
        target_type="EMPLOYEE_ATTACHMENT",
        target_id=user_id,
    )
    try:
        db.session.commit()
        flash("تم إلغاء النشر.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر العملية.", "danger")
    return redirect(url_for('portal.hr_employee_payslips', user_id=user_id))


@portal_bp.route("/hr/employees/<int:user_id>/attachments/upload", methods=["POST"])
@login_required
@_perm(HR_EMP_ATTACH)
def hr_employee_attachments_upload(user_id: int):
    User.query.get_or_404(user_id)

    files = request.files.getlist("files") or []
    if not files:
        one = request.files.get("file")
        if one and getattr(one, "filename", ""):
            files = [one]
    if not files:
        flash("اختر ملفًا واحدًا على الأقل.", "danger")
        return redirect(url_for("portal.hr_employee_attachments", user_id=user_id))

    next_url = (request.form.get('next') or '').strip() or None

    # mode: PAYSLIP or OTHER
    attachment_type = (request.form.get("attachment_type") or "OTHER").strip().upper() or "OTHER"
    note = (request.form.get("note") or "").strip() or None

    attachment_type_lookup_id = _to_int(request.form.get('attachment_type_lookup_id'))

    # Payslip period (required for PAYSLIP)
    payslip_year = None
    payslip_month = None
    if attachment_type == "PAYSLIP":
        y = (request.form.get("payslip_year") or "").strip()
        m = (request.form.get("payslip_month") or "").strip()
        if not y or not m:
            flash("لقسيمة الراتب: الرجاء اختيار السنة والشهر.", "warning")
            return redirect(next_url or url_for("portal.hr_employee_payslips", user_id=user_id))
        try:
            payslip_year = int(y)
            payslip_month = int(m)
        except Exception:
            flash("الشهر/السنة غير صالحين لقسيمة الراتب.", "danger")
            return redirect(next_url or url_for("portal.hr_employee_payslips", user_id=user_id))
        if not (1 <= payslip_month <= 12):
            flash("الشهر يجب أن يكون بين 1 و 12.", "danger")
            return redirect(next_url or url_for("portal.hr_employee_payslips", user_id=user_id))
        if not (2000 <= payslip_year <= 2100):
            flash("السنة غير منطقية.", "danger")
            return redirect(next_url or url_for("portal.hr_employee_payslips", user_id=user_id))

    saved = 0
    for f in files:
        if not f or not getattr(f, "filename", ""):
            continue
        original = Path(f.filename).name
        ext = _clean_suffix(original)
        if ext and ext not in ALLOWED_CORR_EXTS:
            flash(f"امتداد غير مسموح: {ext}", "danger")
            continue

        stored = f"{uuid.uuid4().hex}{ext}" if ext else uuid.uuid4().hex
        dirp = _employee_upload_dir(user_id)
        f.save(dirp / stored)

        if attachment_type == "PAYSLIP" and payslip_year and payslip_month:
            existing = EmployeeAttachment.query.filter_by(
                user_id=user_id,
                attachment_type="PAYSLIP",
                payslip_year=payslip_year,
                payslip_month=payslip_month,
            ).first()
            if existing:
                # delete old file
                try:
                    old_fp = dirp / (existing.stored_name or "")
                    if old_fp.exists():
                        old_fp.unlink()
                except Exception:
                    pass
                existing.original_name = original
                existing.stored_name = stored
                existing.note = note
                existing.uploaded_at = datetime.utcnow()
                existing.uploaded_by_id = current_user.id
                existing.is_published = False
                existing.published_at = None
                existing.published_by_id = None
                saved += 1
                continue

        att = EmployeeAttachment(
            user_id=user_id,
            attachment_type=("PAYSLIP" if attachment_type == "PAYSLIP" else "OTHER"),
            attachment_type_lookup_id=(None if attachment_type == "PAYSLIP" else attachment_type_lookup_id),
            original_name=original,
            stored_name=stored,
            note=note,
            payslip_year=payslip_year if attachment_type == "PAYSLIP" else None,
            payslip_month=payslip_month if attachment_type == "PAYSLIP" else None,
            is_published=(False if attachment_type == "PAYSLIP" else True),
            uploaded_at=datetime.utcnow(),
            uploaded_by_id=current_user.id,
        )
        db.session.add(att)
        saved += 1

    _portal_audit(
        action="HR_EMPLOYEE_ATTACHMENT_UPLOAD",
        note=f"رفع {saved} مرفق(ات)",
        target_type="EMPLOYEE_ATTACHMENT",
        target_id=user_id,
    )

    try:
        db.session.commit()
        flash("تم رفع المرفقات.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر رفع المرفقات.", "danger")

    # Redirect back to caller page
    if next_url:
        return redirect(next_url)
    if attachment_type == 'PAYSLIP':
        return redirect(url_for('portal.hr_employee_payslips', user_id=user_id))
    return redirect(url_for('portal.hr_employee_attachments', user_id=user_id))


@portal_bp.route("/hr/employees/<int:user_id>/attachments/<int:att_id>/download")
@login_required
@_perm(HR_EMP_READ)
def hr_employee_attachment_download(user_id: int, att_id: int):
    att = EmployeeAttachment.query.filter_by(id=att_id, user_id=user_id).first_or_404()
    dirp = _employee_upload_dir(user_id)
    return send_from_directory(dirp, att.stored_name, as_attachment=True, download_name=att.original_name)



@portal_bp.route("/hr/employees/<int:user_id>/attachments/<int:att_id>/edit", methods=["GET", "POST"])
@login_required
@_perm(HR_EMP_ATTACH)
def hr_employee_attachment_edit(user_id: int, att_id: int):
    u = User.query.get_or_404(user_id)
    a = EmployeeAttachment.query.filter_by(id=att_id, user_id=user_id).first_or_404()
    if a.attachment_type == "PAYSLIP":
        flash("لا يمكن تعديل نوع قسيمة الراتب من هنا.", "warning")
        return redirect(url_for("portal.hr_employee_payslips", user_id=user_id))

    lookups = {
        "ATTACH_TYPE": _hr_lookup_items("ATTACH_TYPE"),
    }

    if request.method == "POST":
        a.attachment_type_lookup_id = _i(request.form.get("attachment_type_lookup_id"))
        a.note = _s(request.form.get("note"))
        a.updated_by_id = current_user.id
        db.session.commit()
        flash("تم تحديث المرفق.", "success")
        return redirect(url_for("portal.hr_employee_attachments", user_id=user_id))

    return render_template(
        "portal/hr/employee/attachment_edit.html",
        u=u,
        emp_file=u.employee_file,
        a=a,
        lookups=lookups,
        active_tab="attachments",
    )


@portal_bp.route("/hr/employees/<int:user_id>/attachments/<int:att_id>/delete", methods=["POST"])
@login_required
@_perm(HR_EMP_ATTACH)
def hr_employee_attachment_delete(user_id: int, att_id: int):
    att = EmployeeAttachment.query.filter_by(id=att_id, user_id=user_id).first_or_404()

    dirp = _employee_upload_dir(user_id)
    try:
        fp = dirp / att.stored_name
        if fp.exists():
            fp.unlink()
    except Exception:
        pass

    db.session.delete(att)
    _portal_audit(
        action="HR_EMPLOYEE_ATTACHMENT_DELETE",
        note=f"حذف مرفق: {att.original_name}",
        target_type="EMPLOYEE_ATTACHMENT",
        target_id=user_id,
    )

    try:
        db.session.commit()
        flash("تم حذف المرفق.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر حذف المرفق.", "danger")

    # redirect based on type
    if (att.attachment_type or '').upper() == 'PAYSLIP':
        return redirect(url_for("portal.hr_employee_payslips", user_id=user_id))
    return redirect(url_for("portal.hr_employee_attachments", user_id=user_id))

# -------------------------
# HR: Attendance import
# -------------------------
def _parse_timeclock_line(line: str):
    """
    Supported formats:

    (A) Fixed-length (30 chars):
        YYYY-MM-DD + 15 digits (emp_code 9 + HHMMSS 6) + I/O + device_id(4)
        Example: 2023-02-05000080439073700I1002

    (B) CSV-like exports (common with daily *.CSV files):
        Tries to parse date + time + emp_code + I/O (+ optional device_id)
    """
    if not line:
        return None
    s = line.strip()

    # (A) Fixed-length format
    if len(s) == 30:
        date_s = s[:10]
        payload = s[10:25]
        event_type = s[25]
        device_id = s[26:30]

        emp_code = payload[:9]
        hhmmss = payload[9:]
        if event_type not in ("I", "O"):
            return None
        if (not emp_code.isdigit()) or (not hhmmss.isdigit()):
            return None
        if (not device_id.isdigit()):
            return None

        try:
            y, m, d = [int(x) for x in date_s.split("-")]
            hh = int(hhmmss[0:2]); mm = int(hhmmss[2:4]); ss = int(hhmmss[4:6])
            dt = datetime(y, m, d, hh, mm, ss)
        except Exception:
            return None

        return {
            "emp_code": emp_code,
            "event_dt": dt,
            "event_type": event_type,
            "device_id": device_id,
            "raw": s,
        }

    # (B) CSV-like fallbacks
    delim = None
    for d in [',', ';', '\t']:
        if d in s:
            delim = d
            break
    if not delim:
        return None

    parts = [p.strip().strip('"').strip("'") for p in s.split(delim)]
    parts = [p for p in parts if p != ""]
    if len(parts) < 3:
        return None

    # Detect event type (supports multiple exports)
    event_type = None
    type_map = {
        'I': 'I', 'IN': 'I', 'A': 'I', 'CHECKIN': 'I', 'CHECK-IN': 'I', '0': 'I',
        'O': 'O', 'OUT': 'O', 'B': 'O', 'CHECKOUT': 'O', 'CHECK-OUT': 'O', '1': 'O',
    }
    for p in parts:
        key = (p or '').strip().upper()
        if key in type_map:
            event_type = type_map[key]
            break
    if event_type is None:
        return None

    device_id = "0000"
    for p in reversed(parts):
        if p.isdigit() and 1 <= len(p) <= 6:
            device_id = p.zfill(4)[-4:]
            break

    emp_code = None
    # Prefer first column as employee/identity number when it is numeric.
    if parts and (parts[0] or '').isdigit():
        emp_code = (parts[0] or '').strip()
    else:
        for p in parts:
            if p.isdigit() and len(p) >= 4:
                # avoid capturing YYYYMMDD
                if len(p) == 8 and p.startswith(('19', '20')):
                    continue
                emp_code = p
                break
    if emp_code is None:
        return None

    dt = None

    # 1) Combined datetime token (e.g. "2/15/2026 7:17 AM")
    dt_formats = [
        "%m/%d/%Y %I:%M %p",
        "%m/%d/%Y %I:%M:%S %p",
        "%d/%m/%Y %I:%M %p",
        "%d/%m/%Y %I:%M:%S %p",
        "%m/%d/%Y %H:%M",
        "%m/%d/%Y %H:%M:%S",
        "%Y-%m-%d %H:%M",
        "%Y-%m-%d %H:%M:%S",
        "%Y/%m/%d %H:%M",
        "%Y/%m/%d %H:%M:%S",
    ]
    for p in parts:
        pp = (p or '').strip()
        if (('/' in pp or '-' in pp) and (':' in pp)):
            for fmt in dt_formats:
                try:
                    dt = datetime.strptime(pp, fmt)
                    break
                except Exception:
                    dt = None
            if dt is not None:
                break

    # 2) date + time columns (two tokens)
    if dt is None:
        for i in range(len(parts) - 1):
            d0 = (parts[i] or '').strip()
            t0 = (parts[i + 1] or '').strip()
            date_s = None
            if len(d0) == 10 and d0[4] in ('-', '/') and d0[7] in ('-', '/'):
                date_s = d0.replace('/', '-')
            elif len(d0) == 8 and d0.isdigit() and d0.startswith(('19', '20')):
                date_s = f"{d0[0:4]}-{d0[4:6]}-{d0[6:8]}"
            if not date_s:
                continue

            # AM/PM time like "7:17 AM"
            t_upper = t0.upper()
            if (('AM' in t_upper) or ('PM' in t_upper)) and (':' in t0):
                tt = ' '.join(t0.split())
                for tfmt in ("%I:%M %p", "%I:%M:%S %p"):
                    try:
                        t_parsed = datetime.strptime(tt, tfmt).time()
                        dt = datetime.fromisoformat(date_s + "T" + t_parsed.strftime("%H:%M:%S"))
                        break
                    except Exception:
                        dt = None
                if dt is not None:
                    break

            time_s = None
            if len(t0) == 8 and t0[2] == ':' and t0[5] == ':':
                time_s = t0
            elif len(t0) == 5 and t0[2] == ':':
                time_s = t0 + ":00"
            elif len(t0) == 6 and t0.isdigit():
                time_s = f"{t0[0:2]}:{t0[2:4]}:{t0[4:6]}"
            if not time_s:
                continue

            try:
                dt = datetime.fromisoformat(f"{date_s}T{time_s}")
                break
            except Exception:
                dt = None

    # 3) ISO token with 'T'
    if dt is None:
        for p in parts:
            pp = (p or '').strip()
            if 'T' in pp and len(pp) >= 16:
                try:
                    dt = datetime.fromisoformat(pp.replace('Z', ''))
                    break
                except Exception:
                    dt = None

    if dt is None:
        return None

    return {
        "emp_code": emp_code,
        "event_dt": dt,
        "event_type": event_type,
        "device_id": device_id,
        "raw": s,
    }


@portal_bp.route("/hr/attendance/import", methods=["GET", "POST"])
@login_required
@_perm(HR_ATT_CREATE)
def hr_attendance_import():
    if request.method == "POST":
        f = request.files.get("file")
        if not f or not getattr(f, "filename", ""):
            flash("اختر ملف ساعة الدوام.", "danger")
            return redirect(request.url)

        filename = Path(f.filename).name

        batch = AttendanceImportBatch(
            filename=filename,
            imported_by_id=current_user.id,
            imported_at=datetime.utcnow(),
            total_lines=0,
            inserted=0,
            skipped=0,
        )
        db.session.add(batch)
        db.session.flush()

        content = f.read()
        try:
            text = content.decode("utf-8", errors="ignore")
        except Exception:
            text = content.decode(errors="ignore")

        lines = [ln for ln in text.splitlines() if (ln or "").strip()]
        batch.total_lines = len(lines)

        match_by = _timeclock_get_match_by()
        code_to_user = _timeclock_build_code_to_user(match_by)

        seen = set()
        errors = []

        for ln in lines:
            parsed = _parse_timeclock_line(ln)
            if not parsed:
                batch.skipped += 1
                errors.append(f"Bad line: {ln!r}")
                continue

            emp_code = parsed["emp_code"]
            user_id = code_to_user.get(emp_code)
            if not user_id and (emp_code or '').isdigit():
                user_id = code_to_user.get((emp_code.lstrip('0') or '0'))
            if not user_id:
                batch.skipped += 1
                errors.append(f"Unknown emp_code={emp_code} line={ln!r}")
                continue

            key = (user_id, parsed["event_dt"], parsed["event_type"], parsed["device_id"])
            if key in seen:
                batch.skipped += 1
                continue
            seen.add(key)

            ev = AttendanceEvent(
                batch_id=batch.id,
                user_id=user_id,
                event_dt=parsed["event_dt"],
                event_type=parsed["event_type"],
                device_id=parsed["device_id"],
                raw_line=parsed["raw"],
                created_at=datetime.utcnow(),
            )
            db.session.add(ev)
            batch.inserted += 1

        # store first N errors
        if errors:
            batch.errors = "\n".join(errors[:200])

        db.session.commit()
        flash(f"تم استيراد الملف: {batch.inserted} سجل، {batch.skipped} تم تجاهله.", "success")
        return redirect(url_for("portal.hr_attendance_batches"))

    return render_template("portal/hr/attendance_import.html")


@portal_bp.route("/hr/attendance/batches")
@login_required
@_perm(HR_ATT_READ)
def hr_attendance_batches():
    q = (request.args.get("q") or "").strip()
    date_from = (request.args.get("date_from") or "").strip()
    date_to = (request.args.get("date_to") or "").strip()

    qry = AttendanceImportBatch.query
    if q:
        qry = apply_search_all_columns(qry, AttendanceImportBatch, q, exclude_columns={"errors"})
    if date_from:
        try:
            dt_from = datetime.fromisoformat(date_from + "T00:00:00")
            qry = qry.filter(AttendanceImportBatch.imported_at >= dt_from)
        except Exception:
            pass
    if date_to:
        try:
            dt_to = datetime.fromisoformat(date_to + "T23:59:59")
            qry = qry.filter(AttendanceImportBatch.imported_at <= dt_to)
        except Exception:
            pass

    batches = qry.order_by(AttendanceImportBatch.imported_at.desc()).limit(300).all()
    return render_template("portal/hr/attendance_batches.html", batches=batches, q=q, date_from=date_from, date_to=date_to)


@portal_bp.route("/hr/attendance/events")
@login_required
@_perm(HR_ATT_READ)
def hr_attendance_events():
    """Admin/events view for raw timeclock events."""
    q = (request.args.get("q") or "").strip()
    date_from = (request.args.get("date_from") or "").strip()
    date_to = (request.args.get("date_to") or "").strip()
    event_type = (request.args.get("event_type") or "").strip().upper()  # I/O
    # UI-required filters (may be used later in more advanced logic)
    work_location_lookup_id = (request.args.get("work_location_lookup_id") or "").strip()
    closing = (request.args.get("closing") or "").strip().upper()  # OPEN/CLOSED (placeholder)
    device_id = (request.args.get("device_id") or "").strip()
    batch_id = (request.args.get("batch_id") or "").strip()
    user_id = (request.args.get("user_id") or "").strip()

    qry = AttendanceEvent.query

    if q:
        qry = apply_search_all_columns(qry, AttendanceEvent, q)
    if date_from:
        try:
            dt_from = datetime.fromisoformat(date_from + "T00:00:00")
            qry = qry.filter(AttendanceEvent.event_dt >= dt_from)
        except Exception:
            pass
    if date_to:
        try:
            dt_to = datetime.fromisoformat(date_to + "T23:59:59")
            qry = qry.filter(AttendanceEvent.event_dt <= dt_to)
        except Exception:
            pass
    if event_type in ("I", "O"):
        qry = qry.filter(AttendanceEvent.event_type == event_type)
    if device_id:
        qry = qry.filter(AttendanceEvent.device_id == device_id)
    if batch_id.isdigit():
        qry = qry.filter(AttendanceEvent.batch_id == int(batch_id))
    if user_id.isdigit():
        qry = qry.filter(AttendanceEvent.user_id == int(user_id))

    # Filter by employee work location (from EmployeeFile)
    if work_location_lookup_id.isdigit():
        try:
            loc_int = int(work_location_lookup_id)
            qry = (
                qry.join(User, User.id == AttendanceEvent.user_id)
                   .join(EmployeeFile, EmployeeFile.user_id == User.id)
                   .filter(EmployeeFile.work_location_lookup_id == loc_int)
            )
        except Exception:
            pass

    events = qry.order_by(AttendanceEvent.event_dt.desc()).limit(500).all()

    # Users map for dropdown + table display (prevents UndefinedError in templates)
    try:
        # Prefer all users for filter dropdown (simple + predictable)
        ulist = User.query.order_by(User.name.asc().nullslast(), User.email.asc()).limit(2000).all()
        users = {u.id: u for u in ulist}
    except Exception:
        users = {}

    # Lookups for dropdowns
    try:
        locs = _hr_lookup_items_for_category('WORK_LOCATION')
    except Exception:
        locs = []

    # Excel export (same filters)
    export = (request.args.get("export") or "").strip().lower()
    if export in ("1", "true", "excel", "xlsx"):
        try:
            from utils.excel import make_xlsx_bytes

            headers = [
                "ID",
                "User ID",
                "Email",
                "Event DT",
                "Type",
                "Device",
                "Batch",
                "Raw",
            ]
            rows = []
            for e in events:
                u = users.get(e.user_id)
                rows.append([
                    e.id,
                    e.user_id,
                    (u.email if u else ""),
                    str(e.event_dt),
                    e.event_type,
                    e.device_id,
                    e.batch_id,
                    e.raw_line,
                ])
            xbytes = make_xlsx_bytes("attendance_events", headers, rows)
            return send_file(
                BytesIO(xbytes),
                as_attachment=True,
                download_name="attendance_events.xlsx",
                mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        except Exception:
            flash("تعذر تصدير Excel حالياً.", "danger")

    return render_template(
        "portal/hr/attendance_events.html",
        events=events,
        users=users,
        locs=locs,
        q=q,
        date_from=date_from,
        date_to=date_to,
        event_type=event_type,
        work_location_lookup_id=work_location_lookup_id,
        closing=closing,
        device_id=device_id,
        batch_id=batch_id,
        user_id=user_id,
    )


# -------------------------
# HR: Org Structure (Managers + Teams)
# -------------------------


def _unit_label(unit_type: str, unit_obj) -> str:
    if not unit_obj:
        return ""
    # Organization/Directorate/Department/Section/Team all have label property or name_ar
    return getattr(unit_obj, "label", None) or getattr(unit_obj, "name_ar", None) or getattr(unit_obj, "name", None) or str(unit_obj)


def _get_mgr(unit_type: str, unit_id: int) -> OrgUnitManager:
    row = OrgUnitManager.query.filter_by(unit_type=unit_type, unit_id=unit_id).first()
    if not row:
        row = OrgUnitManager(unit_type=unit_type, unit_id=unit_id, updated_at=datetime.utcnow(), updated_by_id=current_user.id)
        db.session.add(row)
        db.session.flush()
    return row


def _all_users_for_select():
    return User.query.order_by(User.name.asc().nullslast(), User.email.asc()).all()


def _unit_obj(unit_type: str, unit_id: int):
    unit_type = (unit_type or "").strip().upper()
    if unit_type == "ORGANIZATION":
        return Organization.query.get(unit_id)
    if unit_type == "DIRECTORATE":
        return Directorate.query.get(unit_id)
    if unit_type == "UNIT":
        return Unit.query.get(unit_id)
    if unit_type == "DEPARTMENT":
        return Department.query.get(unit_id)
    if unit_type == "SECTION":
        return Section.query.get(unit_id)
    if unit_type == "TEAM":
        return Team.query.get(unit_id)
    return None


def _unit_parent(unit_type: str, unit_obj):
    """Return (parent_type, parent_id, parent_obj) or (None,None,None)."""
    unit_type = (unit_type or "").strip().upper()
    if not unit_obj:
        return None, None, None

    if unit_type == "TEAM":
        sec = getattr(unit_obj, "section", None)
        return "SECTION", getattr(unit_obj, "section_id", None), sec

    if unit_type == "SECTION":
        # Section can belong to Department OR Unit OR Directorate (exactly one)
        if getattr(unit_obj, "department_id", None):
            dep = getattr(unit_obj, "department", None)
            return "DEPARTMENT", getattr(unit_obj, "department_id", None), dep
        if getattr(unit_obj, "unit_id", None):
            u = getattr(unit_obj, "unit", None)
            return "UNIT", getattr(unit_obj, "unit_id", None), u
        if getattr(unit_obj, "directorate_id", None):
            d = getattr(unit_obj, "directorate", None)
            return "DIRECTORATE", getattr(unit_obj, "directorate_id", None), d
        return None, None, None

    if unit_type == "DEPARTMENT":
        # Department can belong to Unit OR Directorate (exactly one)
        if getattr(unit_obj, "unit_id", None):
            u = getattr(unit_obj, "unit", None)
            return "UNIT", getattr(unit_obj, "unit_id", None), u
        d = getattr(unit_obj, "directorate", None)
        return "DIRECTORATE", getattr(unit_obj, "directorate_id", None), d

    if unit_type == "UNIT":
        d = getattr(unit_obj, "directorate", None)
        return "DIRECTORATE", getattr(unit_obj, "directorate_id", None), d

    if unit_type == "DIRECTORATE":
        org = getattr(unit_obj, "organization", None)
        return "ORGANIZATION", getattr(unit_obj, "organization_id", None), org

    return None, None, None



def _build_unit_chain(unit_type: str, unit_id: int):
    """Return bottom-up chain: [(type,id,obj), ...] up to ORGANIZATION."""
    chain = []
    ut = (unit_type or "").strip().upper()
    try:
        uid = int(unit_id)
    except Exception:
        return chain

    obj = _unit_obj(ut, uid)
    if not obj:
        return chain

    while obj and ut:
        chain.append((ut, uid, obj))
        p_type, p_id, p_obj = _unit_parent(ut, obj)
        if not p_type or not p_id:
            break
        ut = p_type
        uid = int(p_id)
        obj = p_obj or _unit_obj(ut, uid)
    return chain


def _safe_query_org_assignments():
    """Best-effort read for OrgUnitAssignment (creates missing tables if DB is older)."""
    try:
        return OrgUnitAssignment.query.all()
    except OperationalError as e:
        try:
            db.session.rollback()
        except Exception:
            pass
        msg = str(e).lower()
        if "no such table" in msg and "org_unit_assignment" in msg:
            try:
                db.create_all()
                return OrgUnitAssignment.query.all()
            except Exception:
                try:
                    db.session.rollback()
                except Exception:
                    pass
        return []


def _primary_assignment_for_user(user_id: int, rows: list[OrgUnitAssignment]):
    user_rows = [r for r in rows if r.user_id == user_id]
    if not user_rows:
        return None
    prim = [r for r in user_rows if r.is_primary]
    return prim[0] if prim else user_rows[0]


def _manager_name(mgr_map: dict, unit_type: str, unit_id: int):
    row = mgr_map.get((unit_type, unit_id))
    if row and row.manager_user:
        return row.manager_user.full_name
    return "—"


def _build_employee_row(user: User, assignment: OrgUnitAssignment | None, mgr_map: dict) -> dict:
    """Build a row for the 'manager chain' table."""
    out = {
        "user": user,
        "assigned_unit": "—",
        "team_mgr": "—",
        "section_mgr": "—",
        "department_mgr": "—",
        "unit_mgr": "—",
        "directorate_mgr": "—",
        "organization_mgr": "—",
        "chain_text": "—",
    }
    if not assignment:
        return out

    chain = _build_unit_chain(assignment.unit_type, assignment.unit_id)
    ids = {t: i for (t, i, _o) in chain}

    # Friendly assigned unit label
    try:
        uobj = _unit_obj(assignment.unit_type, assignment.unit_id)
        out["assigned_unit"] = f"{assignment.unit_type}: {_unit_label(assignment.unit_type, uobj)}"
    except Exception:
        pass

    if "TEAM" in ids:
        out["team_mgr"] = _manager_name(mgr_map, "TEAM", ids["TEAM"])
    if "SECTION" in ids:
        out["section_mgr"] = _manager_name(mgr_map, "SECTION", ids["SECTION"])
    if "DEPARTMENT" in ids:
        out["department_mgr"] = _manager_name(mgr_map, "DEPARTMENT", ids["DEPARTMENT"])
    if "UNIT" in ids:
        out["unit_mgr"] = _manager_name(mgr_map, "UNIT", ids["UNIT"])
    if "DIRECTORATE" in ids:
        out["directorate_mgr"] = _manager_name(mgr_map, "DIRECTORATE", ids["DIRECTORATE"])
    if "ORGANIZATION" in ids:
        out["organization_mgr"] = _manager_name(mgr_map, "ORGANIZATION", ids["ORGANIZATION"])

    parts = []
    for key, label in (
        ("team_mgr", "مدير الفريق"),
        ("section_mgr", "مدير القسم"),
        ("department_mgr", "مدير الدائرة"),
        ("unit_mgr", "مدير الوحدة"),
        ("directorate_mgr", "مدير الإدارة"),
        ("organization_mgr", "مدير المؤسسة"),
    ):
        val = out.get(key)
        if val and val != "—":
            parts.append(f"{label}: {val}")
    if parts:
        out["chain_text"] = " → ".join(parts)
    return out



def _mgr_pair(mgr_map: dict, unit_type: str, unit_id: int) -> tuple[str, str]:
    row = mgr_map.get((unit_type, unit_id))
    mgr = row.manager_user.full_name if row and getattr(row, 'manager_user', None) else "—"
    dep = row.deputy_user.full_name if row and getattr(row, 'deputy_user', None) else "—"
    return mgr, dep


def _build_org_tree(
    orgs,
    directorates,
    units,
    departments,
    sections,
    teams,
    mgr_map: dict,
    org_assignments: list,
    include_people: bool = False,
):
    """Build a nested tree structure for rendering the org chart in HTML.

    Hierarchy:
      Organization → Directorates + Units
      Directorate → Departments (direct) + Sections (direct)
      Unit → Departments + Sections (direct)
      Department → Sections
      Section → Teams
    """

    def sk(obj):
        # Sort by code then name
        code = (getattr(obj, 'code', None) or '')
        name = (getattr(obj, 'name_ar', None) or getattr(obj, 'name_en', None) or '')
        return (code, name)

    dirs_by_org = {}
    for d in directorates or []:
        dirs_by_org.setdefault(d.organization_id, []).append(d)
    for k in dirs_by_org:
        dirs_by_org[k].sort(key=sk)

    units_by_org = {}
    for u in units or []:
        units_by_org.setdefault(getattr(u, 'organization_id', None), []).append(u)
    for k in units_by_org:
        units_by_org[k].sort(key=sk)

    depts_by_dir = {}
    depts_by_unit = {}
    for d in departments or []:
        if getattr(d, 'unit_id', None):
            depts_by_unit.setdefault(d.unit_id, []).append(d)
        elif getattr(d, 'directorate_id', None):
            depts_by_dir.setdefault(d.directorate_id, []).append(d)
    for k in depts_by_dir:
        depts_by_dir[k].sort(key=sk)
    for k in depts_by_unit:
        depts_by_unit[k].sort(key=sk)

    secs_by_dept = {}
    secs_by_dir = {}
    secs_by_unit = {}
    for s in sections or []:
        if getattr(s, 'department_id', None):
            secs_by_dept.setdefault(s.department_id, []).append(s)
        elif getattr(s, 'unit_id', None):
            secs_by_unit.setdefault(s.unit_id, []).append(s)
        elif getattr(s, 'directorate_id', None):
            secs_by_dir.setdefault(s.directorate_id, []).append(s)
    for k in secs_by_dept:
        secs_by_dept[k].sort(key=sk)
    for k in secs_by_dir:
        secs_by_dir[k].sort(key=sk)
    for k in secs_by_unit:
        secs_by_unit[k].sort(key=sk)

    teams_by_sec = {}
    for t in teams or []:
        teams_by_sec.setdefault(t.section_id, []).append(t)
    for k in teams_by_sec:
        teams_by_sec[k].sort(key=sk)

    members_by_unit = {}
    if include_people:
        for a in org_assignments or []:
            key = (a.unit_type, a.unit_id)
            members_by_unit.setdefault(key, []).append(a)
        for key, lst in members_by_unit.items():
            lst.sort(key=lambda a: (
                not bool(getattr(a, 'is_primary', False)),
                ((a.user.name or '').lower() if getattr(a, 'user', None) else ''),
                ((a.user.email or '').lower() if getattr(a, 'user', None) else ''),
            ))

    def build_node(unit_type: str, obj, children: list):
        mgr, dep = _mgr_pair(mgr_map, unit_type, obj.id)
        node = {
            'type': unit_type,
            'id': obj.id,
            'name_ar': getattr(obj, 'name_ar', str(obj.id)),
            'name_en': getattr(obj, 'name_en', None),
            'code': getattr(obj, 'code', None),
            'manager': mgr,
            'deputy': dep,
            'members': [],
            'children': children or [],
        }
        if include_people:
            for a in members_by_unit.get((unit_type, obj.id), []):
                u = getattr(a, 'user', None)
                if not u:
                    continue
                node['members'].append({
                    'name': u.full_name,
                    'title': (getattr(a, 'title', None) or '').strip() or None,
                    'is_primary': bool(getattr(a, 'is_primary', False)),
                })
        return node

    tree = []
    for o in orgs or []:
        org_children = []

        # Directorates under organization
        for dr in dirs_by_org.get(o.id, []):
            dir_kids = []

            # Departments directly under directorate
            for dp in depts_by_dir.get(dr.id, []):
                sec_children = []
                for sc in secs_by_dept.get(dp.id, []):
                    team_children = [build_node('TEAM', tm, []) for tm in teams_by_sec.get(sc.id, [])]
                    sec_children.append(build_node('SECTION', sc, team_children))
                dir_kids.append(build_node('DEPARTMENT', dp, sec_children))

            # Sections directly under directorate
            for sc in secs_by_dir.get(dr.id, []):
                team_children = [build_node('TEAM', tm, []) for tm in teams_by_sec.get(sc.id, [])]
                dir_kids.append(build_node('SECTION', sc, team_children))

            org_children.append(build_node('DIRECTORATE', dr, dir_kids))

        # Units under organization (parallel to Directorates)
        for un in units_by_org.get(o.id, []):
            unit_kids = []

            # Departments under unit
            for dp in depts_by_unit.get(un.id, []):
                sec_children = []
                for sc in secs_by_dept.get(dp.id, []):
                    team_children = [build_node('TEAM', tm, []) for tm in teams_by_sec.get(sc.id, [])]
                    sec_children.append(build_node('SECTION', sc, team_children))
                unit_kids.append(build_node('DEPARTMENT', dp, sec_children))

            # Sections directly under unit
            for sc in secs_by_unit.get(un.id, []):
                team_children = [build_node('TEAM', tm, []) for tm in teams_by_sec.get(sc.id, [])]
                unit_kids.append(build_node('SECTION', sc, team_children))

            org_children.append(build_node('UNIT', un, unit_kids))

        tree.append(build_node('ORGANIZATION', o, org_children))

    return tree

@portal_bp.route("/hr/org-structure")
@login_required
def hr_org_structure():
    # SUPER_ADMIN must always access this page; otherwise require explicit HR org read/manage.
    try:
        role_raw = (getattr(current_user, 'role', '') or '').strip().upper().replace('-', '_').replace(' ', '_')
    except Exception:
        role_raw = ''
    is_super = bool(role_raw.startswith('SUPER'))
    if not is_super:
        try:
            if not (current_user.has_perm(HR_ORG_READ) or current_user.has_perm(HR_MASTERDATA_MANAGE)):
                abort(403)
        except Exception:
            abort(403)
    orgs = Organization.query.order_by(Organization.name_ar.asc()).all()
    directorates = Directorate.query.order_by(Directorate.name_ar.asc()).all()
    units = Unit.query.order_by(Unit.name_ar.asc()).all()
    departments = Department.query.order_by(Department.name_ar.asc()).all()
    sections = Section.query.order_by(Section.name_ar.asc()).all()

    # Divisions (الشُعب) are part of the org-structure UI and used by the template.
    # (Bugfix) ensure variable is defined to avoid NameError.
    divisions = []
    try:
        divisions = Division.query.order_by(Division.code.asc().nullslast(), Division.name_ar.asc()).all()
    except Exception:
        divisions = []

    teams = []
    try:
        teams = Team.query.order_by(Team.section_id.asc(), Team.name_ar.asc()).all()
    except OperationalError as e:
        try:
            db.session.rollback()
        except Exception:
            pass
        msg = str(e).lower()
        if 'no such table' in msg and 'team' in msg:
            try:
                db.create_all()
                teams = Team.query.order_by(Team.section_id.asc(), Team.name_ar.asc()).all()
                flash('تم تحديث بنية قاعدة البيانات تلقائياً (إنشاء جداول HR). حدّث الصفحة إذا لزم.', 'info')
            except Exception:
                try:
                    db.session.rollback()
                except Exception:
                    pass
                flash('قاعدة البيانات الحالية قديمة ولا تحتوي على جداول HR (Teams). شغّل: python init_db.py أو احذف instance/workflow.db ثم أعد التشغيل.', 'warning')
        else:
            raise

    mgr_rows = OrgUnitManager.query.all()
    mgr_map = {(m.unit_type, m.unit_id): m for m in mgr_rows}

    try:
        can_manage_org = bool(current_user.has_perm(HR_ORG_MANAGE) or current_user.has_perm(HR_MASTERDATA_MANAGE))
    except Exception:
        can_manage_org = False

    users = _all_users_for_select() if can_manage_org else []

    org_assignments = _safe_query_org_assignments()
    employees = User.query.order_by(User.name.asc().nullslast(), User.email.asc()).all()
    employee_rows = []
    for u in employees:
        a = _primary_assignment_for_user(u.id, org_assignments)
        employee_rows.append(_build_employee_row(u, a, mgr_map))

    include_people = (request.args.get('include_people') or '').strip().lower() in ('1', 'true', 'yes')
    org_tree = _build_org_tree(orgs, directorates, units, departments, sections, teams, mgr_map, org_assignments, include_people=include_people)

    return render_template(
        'portal/hr/org_structure.html',
        orgs=orgs,
        directorates=directorates,
        units=units,
        departments=departments,
        sections=sections,
        divisions=divisions,
        teams=teams,
        assignments=mgr_map,
        users=users,
        employee_rows=employee_rows,
        org_tree=org_tree,
        include_people=include_people,
        can_manage_org=can_manage_org,
        legacy_locked=_legacy_org_locked(),
    )



# =========================
# Dynamic OrgNode assignments (assign employees to unified org structure)
# =========================

@portal_bp.route("/hr/org-nodes/assignments", methods=["GET"])
@login_required
def hr_org_node_assignments():
    # SUPER_ADMIN always allowed; otherwise require HR org manage (or HR masterdata manage).
    try:
        role_raw = (getattr(current_user, 'role', '') or '').strip().upper().replace('-', '_').replace(' ', '_')
    except Exception:
        role_raw = ''
    is_super = bool(role_raw.startswith('SUPER'))
    if not is_super:
        try:
            if not (current_user.has_perm(HR_ORG_MANAGE) or current_user.has_perm(HR_MASTERDATA_MANAGE)):
                abort(403)
        except Exception:
            abort(403)

    q = (request.args.get("q") or "").strip()
    users = []
    if q:
        users = (
            User.query
            .filter(or_(User.name.ilike(f"%{q}%"), User.email.ilike(f"%{q}%")))
            .order_by(User.id.desc())
            .limit(25)
            .all()
        )

    node_map = {}
    try:
        ids = [int(x.org_node_id) for x in users if getattr(x, "org_node_id", None)]
        if ids:
            for n in OrgNode.query.filter(OrgNode.id.in_(ids)).all():
                node_map[int(n.id)] = n
    except Exception:
        node_map = {}
    
    return render_template(
        "portal/hr/org_node_assignments.html",
        q=q,
        users=users,
        node_map=node_map,
    )


@portal_bp.route("/hr/org-nodes/assignments/<int:user_id>", methods=["GET", "POST"])
@login_required
def hr_org_node_assign_user(user_id: int):
    # SUPER_ADMIN always allowed; otherwise require HR org manage (or HR masterdata manage).
    try:
        role_raw = (getattr(current_user, 'role', '') or '').strip().upper().replace('-', '_').replace(' ', '_')
    except Exception:
        role_raw = ''
    is_super = bool(role_raw.startswith('SUPER'))
    if not is_super:
        try:
            if not (current_user.has_perm(HR_ORG_MANAGE) or current_user.has_perm(HR_MASTERDATA_MANAGE)):
                abort(403)
        except Exception:
            abort(403)

    u = User.query.get_or_404(user_id)

    if request.method == "POST":
        raw_node_id = (request.form.get("node_id") or "").strip()
        node_id = int(raw_node_id) if raw_node_id.isdigit() else None

        title = (request.form.get("title") or "").strip() or None
        is_primary = bool(request.form.get("is_primary"))

        if not node_id:
            flash("اختر عنصر الهيكلية أولاً.", "warning")
            return redirect(url_for("portal.hr_org_node_assign_user", user_id=user_id))

        node = OrgNode.query.get(int(node_id))
        if not node or not getattr(node, "is_active", True):
            flash("العنصر غير موجود أو غير نشط.", "danger")
            return redirect(url_for("portal.hr_org_node_assign_user", user_id=user_id))

        # Upsert assignment
        a = OrgNodeAssignment.query.filter_by(user_id=u.id, node_id=int(node_id)).first()
        if not a:
            a = OrgNodeAssignment(user_id=u.id, node_id=int(node_id), created_by_id=getattr(current_user, "id", None))
            db.session.add(a)
        a.title = title
        a.is_primary = bool(is_primary)

        try:
            db.session.flush()  # ensure a.id for updates below
        except Exception:
            pass

        if is_primary:
            # unset other primary assignments
            try:
                OrgNodeAssignment.query.filter(
                    OrgNodeAssignment.user_id == u.id,
                    OrgNodeAssignment.id != a.id
                ).update({OrgNodeAssignment.is_primary: False})
            except Exception:
                pass

            # store direct primary pointer on user for faster routing
            try:
                u.org_node_id = int(node_id)
            except Exception:
                pass

        try:
            db.session.commit()
        except Exception as e:
            db.session.rollback()
            flash("تعذر حفظ التعيين. تحقق من البيانات.", "danger")
            return redirect(url_for("portal.hr_org_node_assign_user", user_id=user_id))

        flash("تم حفظ تعيين الموظف على الهيكلية الموحدة.", "success")
        return redirect(url_for("portal.hr_org_node_assign_user", user_id=user_id))

    assignments = (
        OrgNodeAssignment.query
        .filter_by(user_id=u.id)
        .order_by(OrgNodeAssignment.is_primary.desc(), OrgNodeAssignment.id.desc())
        .all()
    )

    tree = build_org_node_picker_tree(mode="all")

    return render_template(
        "portal/hr/org_node_assign_user.html",
        emp=u,
        assignments=assignments,
        tree=tree,
        selected_id=getattr(u, "org_node_id", None),
    )


@portal_bp.route("/hr/org-nodes/assignments/<int:user_id>/delete/<int:assign_id>", methods=["POST"])
@login_required
def hr_org_node_assign_delete(user_id: int, assign_id: int):
    # SUPER_ADMIN always allowed; otherwise require HR org manage (or HR masterdata manage).
    try:
        role_raw = (getattr(current_user, 'role', '') or '').strip().upper().replace('-', '_').replace(' ', '_')
    except Exception:
        role_raw = ''
    is_super = bool(role_raw.startswith('SUPER'))
    if not is_super:
        try:
            if not (current_user.has_perm(HR_ORG_MANAGE) or current_user.has_perm(HR_MASTERDATA_MANAGE)):
                abort(403)
        except Exception:
            abort(403)

    u = User.query.get_or_404(user_id)
    a = OrgNodeAssignment.query.get_or_404(assign_id)
    if int(a.user_id) != int(u.id):
        abort(404)

    was_primary = bool(a.is_primary)
    node_id = int(a.node_id)

    try:
        db.session.delete(a)
        db.session.flush()

        # If removed primary, pick another one as primary (if exists)
        if was_primary and getattr(u, "org_node_id", None) == node_id:
            other = (
                OrgNodeAssignment.query
                .filter(OrgNodeAssignment.user_id == u.id)
                .order_by(OrgNodeAssignment.is_primary.desc(), OrgNodeAssignment.id.desc())
                .first()
            )
            if other:
                other.is_primary = True
                u.org_node_id = int(other.node_id)
            else:
                u.org_node_id = None

        db.session.commit()
        flash("تم حذف التعيين.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر حذف التعيين.", "danger")

    return redirect(url_for("portal.hr_org_node_assign_user", user_id=user_id))


@portal_bp.route("/hr/org-nodes/assignments/<int:user_id>/primary/<int:assign_id>", methods=["POST"])
@login_required
def hr_org_node_assign_make_primary(user_id: int, assign_id: int):
    # SUPER_ADMIN always allowed; otherwise require HR org manage (or HR masterdata manage).
    try:
        role_raw = (getattr(current_user, 'role', '') or '').strip().upper().replace('-', '_').replace(' ', '_')
    except Exception:
        role_raw = ''
    is_super = bool(role_raw.startswith('SUPER'))
    if not is_super:
        try:
            if not (current_user.has_perm(HR_ORG_MANAGE) or current_user.has_perm(HR_MASTERDATA_MANAGE)):
                abort(403)
        except Exception:
            abort(403)

    u = User.query.get_or_404(user_id)
    a = OrgNodeAssignment.query.get_or_404(assign_id)
    if int(a.user_id) != int(u.id):
        abort(404)

    try:
        # unset others
        OrgNodeAssignment.query.filter(
            OrgNodeAssignment.user_id == u.id,
            OrgNodeAssignment.id != a.id
        ).update({OrgNodeAssignment.is_primary: False})

        a.is_primary = True
        u.org_node_id = int(a.node_id)
        db.session.commit()
        flash("تم تعيينه كأساسي.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر تعيينه كأساسي.", "danger")

    return redirect(url_for("portal.hr_org_node_assign_user", user_id=user_id))



@portal_bp.route("/hr/org-structure/team/new", methods=["POST"])
@login_required
@_perm_any(HR_ORG_MANAGE, HR_MASTERDATA_MANAGE)
def hr_team_new():
    if _legacy_org_locked():
        flash('الهيكلية الثابتة مقفلة (قراءة فقط).', 'warning')
        return redirect(url_for('portal.hr_org_structure'))

    section_id = (request.form.get("section_id") or "").strip()
    name_ar = (request.form.get("name_ar") or "").strip()
    name_en = (request.form.get("name_en") or "").strip() or None
    code = (request.form.get("code") or "").strip() or None

    if not section_id.isdigit() or not name_ar:
        flash("البيانات غير مكتملة.", "danger")
        return redirect(url_for("portal.hr_org_structure"))

    sec = Section.query.get(int(section_id))
    if not sec:
        flash("القسم غير موجود.", "danger")
        return redirect(url_for("portal.hr_org_structure"))

    t = Team(section_id=sec.id, name_ar=name_ar, name_en=name_en, code=code, is_active=True, created_at=datetime.utcnow())
    db.session.add(t)
    _portal_audit("HR_TEAM_CREATE", f"إنشاء فريق: {name_ar}", target_type="TEAM", target_id=0)

    try:
        db.session.commit()
        flash("تم إنشاء الفريق.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر إنشاء الفريق.", "danger")

    return redirect(url_for("portal.hr_org_structure"))


@portal_bp.route("/hr/org-structure/team/<int:team_id>/toggle", methods=["POST"])
@login_required
@_perm_any(HR_ORG_MANAGE, HR_MASTERDATA_MANAGE)
def hr_team_toggle(team_id: int):
    if _legacy_org_locked():
        flash('الهيكلية الثابتة مقفلة (قراءة فقط).', 'warning')
        return redirect(url_for('portal.hr_org_structure'))

    t = Team.query.get_or_404(team_id)
    t.is_active = not bool(t.is_active)
    _portal_audit("HR_TEAM_TOGGLE", f"تفعيل/تعطيل فريق: {t.name_ar}", target_type="TEAM", target_id=team_id)

    try:
        db.session.commit()
        flash("تم تحديث حالة الفريق.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر تحديث حالة الفريق.", "danger")

    return redirect(url_for("portal.hr_org_structure"))


@portal_bp.route("/hr/org-structure/team/save", methods=["POST"])
@login_required
@_perm_any(HR_ORG_MANAGE, HR_MASTERDATA_MANAGE)
def hr_team_save():
    if _legacy_org_locked():
        flash('الهيكلية الثابتة مقفلة (قراءة فقط).', 'warning')
        return redirect(url_for('portal.hr_org_structure'))

    """Create or update a Team (from the modal)."""
    team_id = (request.form.get("team_id") or "").strip()
    section_id = (request.form.get("section_id") or "").strip()
    name_ar = (request.form.get("name_ar") or "").strip()
    name_en = (request.form.get("name_en") or "").strip() or None
    code = (request.form.get("code") or "").strip() or None

    if not section_id.isdigit() or not name_ar:
        flash("البيانات غير مكتملة.", "danger")
        return redirect(url_for("portal.hr_org_structure"))

    sec = Section.query.get(int(section_id))
    if not sec:
        flash("القسم غير موجود.", "danger")
        return redirect(url_for("portal.hr_org_structure"))

    is_new = True
    t = None
    if team_id.isdigit():
        t = Team.query.get(int(team_id))
        if t:
            is_new = False

    if not t:
        t = Team(section_id=sec.id, created_at=datetime.utcnow(), is_active=True)
        db.session.add(t)

    t.section_id = sec.id
    t.name_ar = name_ar
    t.name_en = name_en
    t.code = code

    if is_new:
        _portal_audit("HR_TEAM_CREATE", f"إنشاء فريق: {name_ar}", target_type="TEAM", target_id=0)
    else:
        _portal_audit("HR_TEAM_UPDATE", f"تعديل فريق: {name_ar}", target_type="TEAM", target_id=t.id)

    try:
        db.session.commit()
        flash("تم حفظ الفريق.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر حفظ الفريق.", "danger")

    return redirect(url_for("portal.hr_org_structure"))


@portal_bp.route("/hr/org-structure/team/<int:team_id>/delete", methods=["POST"])
@login_required
@_perm_any(HR_ORG_MANAGE, HR_MASTERDATA_MANAGE)
def hr_team_delete(team_id: int):
    if _legacy_org_locked():
        flash('الهيكلية الثابتة مقفلة (قراءة فقط).', 'warning')
        return redirect(url_for('portal.hr_org_structure'))

    t = Team.query.get_or_404(team_id)
    name = t.name_ar
    _portal_audit("HR_TEAM_DELETE", f"حذف فريق: {name}", target_type="TEAM", target_id=team_id)

    try:
        db.session.delete(t)
        db.session.commit()
        flash("تم حذف الفريق.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر حذف الفريق.", "danger")

    return redirect(url_for("portal.hr_org_structure"))


@portal_bp.route("/hr/org-structure/manager", methods=["POST"])
@login_required
@_perm_any(HR_ORG_MANAGE, HR_MASTERDATA_MANAGE)
def hr_org_set_manager():
    if _legacy_org_locked():
        flash('الهيكلية الثابتة مقفلة (قراءة فقط).', 'warning')
        return redirect(url_for('portal.hr_org_structure'))

    unit_type = (request.form.get("unit_type") or "").strip().upper()
    unit_id = (request.form.get("unit_id") or "").strip()
    manager_user_id = (request.form.get("manager_user_id") or "").strip()
    deputy_user_id = (request.form.get("deputy_user_id") or "").strip()

    if unit_type not in ("ORGANIZATION", "DIRECTORATE", "UNIT", "DEPARTMENT", "SECTION", "TEAM"):
        flash("نوع الوحدة غير صحيح.", "danger")
        return redirect(url_for("portal.hr_org_structure"))
    if not unit_id.isdigit():
        flash("معرف الوحدة غير صحيح.", "danger")
        return redirect(url_for("portal.hr_org_structure"))

    uid = int(unit_id)
    row = _get_mgr(unit_type, uid)

    row.manager_user_id = int(manager_user_id) if manager_user_id.isdigit() else None
    row.deputy_user_id = int(deputy_user_id) if deputy_user_id.isdigit() else None
    row.updated_at = datetime.utcnow()
    row.updated_by_id = current_user.id

    _portal_audit("HR_ORG_MANAGER_SET", f"تحديد مدير/بديل للوحدة {unit_type}#{uid}", target_type="ORG_UNIT", target_id=uid)

    try:
        db.session.commit()
        flash("تم تحديث المدير المباشر/البديل.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر حفظ التعديل.", "danger")

    return redirect(url_for("portal.hr_org_structure"))


@portal_bp.route("/hr/org-assignments")
@login_required
def hr_org_assignments():

    if _legacy_org_locked():
        flash('الهيكلية الثابتة مقفلة. استخدم تعيين الموظفين على الهيكلية الموحدة (Dynamic).', 'warning')
        return redirect(url_for('portal.hr_org_node_assignments'))

    # SUPER_ADMIN must always access this page; otherwise require explicit HR org manage.
    try:
        role_raw = (getattr(current_user, 'role', '') or '').strip().upper().replace('-', '_').replace(' ', '_')
    except Exception:
        role_raw = ''
    is_super = bool(role_raw.startswith('SUPER'))
    if not is_super:
        try:
            if not (current_user.has_perm(HR_ORG_MANAGE) or current_user.has_perm(HR_MASTERDATA_MANAGE)):
                abort(403)
        except Exception:
            abort(403)
    # Assign users to org units for Portal HR (membership).
    # This is independent from Workflow org membership.
    orgs = Organization.query.order_by(Organization.name_ar.asc()).all()
    directorates = Directorate.query.order_by(Directorate.name_ar.asc()).all()
    units = Unit.query.order_by(Unit.name_ar.asc()).all()
    departments = Department.query.order_by(Department.name_ar.asc()).all()
    sections = Section.query.order_by(Section.name_ar.asc()).all()
    divisions = Division.query.order_by(Division.name_ar.asc()).all()

    try:
        teams = Team.query.order_by(Team.section_id.asc(), Team.name_ar.asc()).all()
    except Exception:
        teams = []

    assignments = _safe_query_org_assignments()
    users = _all_users_for_select()

    rows = []
    for a in assignments:
        u = a.user
        uobj = _unit_obj(a.unit_type, a.unit_id)
        rows.append({
            "id": a.id,
            "user": u,
            "unit_type": a.unit_type,
            "unit_id": a.unit_id,
            "unit_label": _unit_label(a.unit_type, uobj),
            "is_primary": bool(a.is_primary),
            "title": (a.title or "").strip(),
            "created_at": a.created_at,
        })

    rows.sort(key=lambda r: ((r["user"].name or "").lower(), (r["user"].email or "").lower(), r["unit_type"], r["unit_label"]))

    return render_template(
        "portal/hr/org_assignments.html",
        orgs=orgs,
        directorates=directorates,
        units=units,
        departments=departments,
        sections=sections,
        divisions=divisions,
        teams=teams,
        users=users,
        rows=rows,
    )


@portal_bp.route("/hr/org-assignments/save", methods=["POST"])
@login_required
@_perm_any(HR_ORG_MANAGE, HR_MASTERDATA_MANAGE)
def hr_org_assignments_save():
    if _legacy_org_locked():
        flash('الهيكلية الثابتة مقفلة (قراءة فقط).', 'warning')
        return redirect(url_for('portal.hr_org_node_assignments'))

    user_id = (request.form.get("user_id") or "").strip()
    unit_type = (request.form.get("unit_type") or "").strip().upper()
    unit_id = (request.form.get("unit_id") or "").strip()
    title = (request.form.get("title") or "").strip() or None
    is_primary = bool(request.form.get("is_primary"))

    if not user_id.isdigit() or not unit_id.isdigit():
        flash("البيانات غير مكتملة.", "danger")
        return redirect(url_for("portal.hr_org_assignments"))

    uid = int(user_id)
    u = User.query.get(uid)
    if not u:
        flash("المستخدم غير موجود.", "danger")
        return redirect(url_for("portal.hr_org_assignments"))

    if unit_type not in ("ORGANIZATION", "DIRECTORATE", "UNIT", "DEPARTMENT", "SECTION", "TEAM"):
        flash("نوع الوحدة غير صحيح.", "danger")
        return redirect(url_for("portal.hr_org_assignments"))

    unit_pk = int(unit_id)
    if not _unit_obj(unit_type, unit_pk):
        flash("الوحدة المحددة غير موجودة.", "danger")
        return redirect(url_for("portal.hr_org_assignments"))

    # Upsert by (user_id, unit_type, unit_id)
    row = OrgUnitAssignment.query.filter_by(user_id=uid, unit_type=unit_type, unit_id=unit_pk).first()
    if not row:
        row = OrgUnitAssignment(user_id=uid, unit_type=unit_type, unit_id=unit_pk, created_at=datetime.utcnow(), created_by_id=current_user.id)
        db.session.add(row)

    row.title = title
    row.is_primary = bool(is_primary)

    # Ensure single primary per user
    if row.is_primary:
        (OrgUnitAssignment.query
         .filter(OrgUnitAssignment.user_id == uid)
         .filter(OrgUnitAssignment.id != row.id)
         .update({OrgUnitAssignment.is_primary: False}))

    _portal_audit("HR_ORG_ASSIGN", f"تعيين {u.full_name} على {unit_type}#{unit_pk}", target_type="ORG_ASSIGN", target_id=0)

    try:
        db.session.commit()
        flash("تم حفظ التعيين.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر حفظ التعيين (قد يكون مكررًا).", "danger")

    return redirect(url_for("portal.hr_org_assignments"))


@portal_bp.route("/hr/org-assignments/<int:assign_id>/delete", methods=["POST"])
@login_required
@_perm_any(HR_ORG_MANAGE, HR_MASTERDATA_MANAGE)
def hr_org_assignments_delete(assign_id: int):
    if _legacy_org_locked():
        flash('الهيكلية الثابتة مقفلة (قراءة فقط).', 'warning')
        return redirect(url_for('portal.hr_org_node_assignments'))

    row = OrgUnitAssignment.query.get_or_404(assign_id)
    name = row.user.full_name if row.user else str(row.user_id)
    _portal_audit("HR_ORG_ASSIGN_DELETE", f"حذف تعيين: {name} من {row.unit_type}#{row.unit_id}", target_type="ORG_ASSIGN", target_id=assign_id)
    try:
        db.session.delete(row)
        db.session.commit()
        flash("تم حذف التعيين.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر حذف التعيين.", "danger")
    return redirect(url_for("portal.hr_org_assignments"))


@portal_bp.route("/hr/org-structure/export/<string:fmt>")
@login_required
@_perm_any(HR_ORG_READ, HR_MASTERDATA_MANAGE)
def hr_org_structure_export(fmt: str):
    """Export org structure diagram/data.

    Supported formats:
      - mermaid
      - csv (nodes)
      - json (nodes + managers + people assignments)
    """
    fmt = (fmt or "").strip().lower()

    orgs = Organization.query.order_by(Organization.id.asc()).all()
    dirs = Directorate.query.order_by(Directorate.id.asc()).all()
    units = Unit.query.order_by(Unit.id.asc()).all()
    depts = Department.query.order_by(Department.id.asc()).all()
    secs = Section.query.order_by(Section.id.asc()).all()
    try:
        teams = Team.query.order_by(Team.id.asc()).all()
    except Exception:
        teams = []

    mgr_rows = OrgUnitManager.query.all()
    mgr_map = {(m.unit_type, m.unit_id): m for m in mgr_rows}
    org_assignments = _safe_query_org_assignments()

    def mgr_of(t, i):
        m = mgr_map.get((t, i))
        return {
            "manager_user_id": getattr(m, "manager_user_id", None),
            "manager_name": (m.manager_user.full_name if m and m.manager_user else None),
            "deputy_user_id": getattr(m, "deputy_user_id", None),
            "deputy_name": (m.deputy_user.full_name if m and m.deputy_user else None),
        }

    def dept_parent(d):
        if getattr(d, "unit_id", None):
            return ("UNIT", d.unit_id)
        return ("DIRECTORATE", d.directorate_id)

    def sec_parent(s):
        if getattr(s, "department_id", None):
            return ("DEPARTMENT", s.department_id)
        if getattr(s, "unit_id", None):
            return ("UNIT", s.unit_id)
        return ("DIRECTORATE", s.directorate_id)

    if fmt == "csv":
        output = io.StringIO()
        w = csv.writer(output)
        w.writerow(["type", "id", "parent_type", "parent_id", "name_ar", "code", "manager", "deputy"])

        for o in orgs:
            m = mgr_of("ORGANIZATION", o.id)
            w.writerow(["ORGANIZATION", o.id, "", "", o.name_ar, o.code or "", m.get("manager_name") or "", m.get("deputy_name") or ""])
        for d in dirs:
            m = mgr_of("DIRECTORATE", d.id)
            w.writerow(["DIRECTORATE", d.id, "ORGANIZATION", d.organization_id, d.name_ar, d.code or "", m.get("manager_name") or "", m.get("deputy_name") or ""])
        for u in units:
            m = mgr_of("UNIT", u.id)
            w.writerow(["UNIT", u.id, "ORGANIZATION", u.organization_id, u.name_ar, u.code or "", m.get("manager_name") or "", m.get("deputy_name") or ""])
        for d in depts:
            ptype, pid = dept_parent(d)
            m = mgr_of("DEPARTMENT", d.id)
            w.writerow(["DEPARTMENT", d.id, ptype, pid or "", d.name_ar, d.code or "", m.get("manager_name") or "", m.get("deputy_name") or ""])
        for s in secs:
            ptype, pid = sec_parent(s)
            m = mgr_of("SECTION", s.id)
            w.writerow(["SECTION", s.id, ptype, pid or "", s.name_ar, s.code or "", m.get("manager_name") or "", m.get("deputy_name") or ""])
        for t in teams:
            m = mgr_of("TEAM", t.id)
            w.writerow(["TEAM", t.id, "SECTION", t.section_id, t.name_ar, t.code or "", m.get("manager_name") or "", m.get("deputy_name") or ""])

        bio = BytesIO(output.getvalue().encode("utf-8-sig"))
        bio.seek(0)
        return send_file(bio, as_attachment=True, download_name="org_structure_nodes.csv", mimetype="text/csv")

    if fmt == "json":
        nodes = []
        for o in orgs:
            nodes.append({
                "type": "ORGANIZATION", "id": o.id, "parent": None,
                "name_ar": o.name_ar, "name_en": o.name_en, "code": o.code,
                "managers": mgr_of("ORGANIZATION", o.id),
            })
        for d in dirs:
            nodes.append({
                "type": "DIRECTORATE", "id": d.id, "parent": {"type": "ORGANIZATION", "id": d.organization_id},
                "name_ar": d.name_ar, "name_en": d.name_en, "code": d.code,
                "managers": mgr_of("DIRECTORATE", d.id),
            })
        for u in units:
            nodes.append({
                "type": "UNIT", "id": u.id, "parent": {"type": "ORGANIZATION", "id": u.organization_id},
                "name_ar": u.name_ar, "name_en": u.name_en, "code": u.code,
                "managers": mgr_of("UNIT", u.id),
            })
        for d in depts:
            ptype, pid = dept_parent(d)
            nodes.append({
                "type": "DEPARTMENT", "id": d.id, "parent": {"type": ptype, "id": pid},
                "name_ar": d.name_ar, "name_en": d.name_en, "code": d.code,
                "managers": mgr_of("DEPARTMENT", d.id),
            })
        for s in secs:
            ptype, pid = sec_parent(s)
            nodes.append({
                "type": "SECTION", "id": s.id, "parent": {"type": ptype, "id": pid},
                "name_ar": s.name_ar, "name_en": s.name_en, "code": s.code,
                "managers": mgr_of("SECTION", s.id),
            })
        for t in teams:
            nodes.append({
                "type": "TEAM", "id": t.id, "parent": {"type": "SECTION", "id": t.section_id},
                "name_ar": t.name_ar, "name_en": t.name_en, "code": t.code,
                "managers": mgr_of("TEAM", t.id),
            })

        people = []
        users = User.query.order_by(User.id.asc()).all()
        for u in users:
            a = _primary_assignment_for_user(u.id, org_assignments)
            if not a:
                continue
            people.append({
                "user_id": u.id,
                "name": u.full_name,
                "email": u.email,
                "assignment": {"unit_type": a.unit_type, "unit_id": a.unit_id, "title": a.title, "is_primary": bool(a.is_primary)},
                "manager_chain": _build_employee_row(u, a, mgr_map),
            })

        payload = {"nodes": nodes, "people": people}
        import json
        bio = BytesIO(json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8"))
        bio.seek(0)
        return send_file(bio, as_attachment=True, download_name="org_structure.json", mimetype="application/json")

    if fmt == "mermaid":
        include_people = (request.args.get("include_people") or "").strip().lower() in ("1", "true", "yes")

        def nid(prefix: str, i: int):
            return f"{prefix}_{i}"

        def esc(s: str | None) -> str:
            txt = (s or "")
            txt = txt.replace("\\", "\\\\")
            txt = txt.replace('"', '\\"')
            return txt

        lines = ["flowchart TD"]

        for o in orgs:
            lines.append(f'  {nid("ORG", o.id)}["{esc(o.name_ar)}"]')
        for d in dirs:
            lines.append(f'  {nid("DIR", d.id)}["{esc(d.name_ar)}"]')
            lines.append(f'  {nid("ORG", d.organization_id)} --> {nid("DIR", d.id)}')
        for u in units:
            lines.append(f'  {nid("UNI", u.id)}["{esc(u.name_ar)}"]')
            if getattr(u, "organization_id", None):
                lines.append(f'  {nid("ORG", u.organization_id)} --> {nid("UNI", u.id)}')
        for d in depts:
            lines.append(f'  {nid("DEP", d.id)}["{esc(d.name_ar)}"]')
            ptype, pid = dept_parent(d)
            if ptype == "UNIT":
                lines.append(f'  {nid("UNI", pid)} --> {nid("DEP", d.id)}')
            else:
                lines.append(f'  {nid("DIR", pid)} --> {nid("DEP", d.id)}')
        for s in secs:
            lines.append(f'  {nid("SEC", s.id)}["{esc(s.name_ar)}"]')
            ptype, pid = sec_parent(s)
            if ptype == "DEPARTMENT":
                lines.append(f'  {nid("DEP", pid)} --> {nid("SEC", s.id)}')
            elif ptype == "UNIT":
                lines.append(f'  {nid("UNI", pid)} --> {nid("SEC", s.id)}')
            else:
                lines.append(f'  {nid("DIR", pid)} --> {nid("SEC", s.id)}')
        for t in teams:
            lines.append(f'  {nid("TEA", t.id)}["{esc(t.name_ar)}"]')
            lines.append(f'  {nid("SEC", t.section_id)} --> {nid("TEA", t.id)}')

        def person_id(user_id: int):
            return f"USR_{user_id}"

        for (ut, uid), m in mgr_map.items():
            if not m or not m.manager_user_id:
                continue
            pu = m.manager_user
            if not pu:
                continue
            pid = person_id(pu.id)
            lines.append(f'  {pid}(["👤 {esc(pu.full_name)}"])')
            if ut == "ORGANIZATION":
                lines.append(f"  {nid('ORG', uid)} -. مدير .-> {pid}")
            elif ut == "DIRECTORATE":
                lines.append(f"  {nid('DIR', uid)} -. مدير .-> {pid}")
            elif ut == "UNIT":
                lines.append(f"  {nid('UNI', uid)} -. مدير .-> {pid}")
            elif ut == "DEPARTMENT":
                lines.append(f"  {nid('DEP', uid)} -. مدير .-> {pid}")
            elif ut == "SECTION":
                lines.append(f"  {nid('SEC', uid)} -. مدير .-> {pid}")
            elif ut == "TEAM":
                lines.append(f"  {nid('TEA', uid)} -. مدير .-> {pid}")

        if include_people:
            users = User.query.order_by(User.id.asc()).all()
            for u in users:
                a = _primary_assignment_for_user(u.id, org_assignments)
                if not a:
                    continue
                pid = person_id(u.id)
                lines.append(f'  {pid}(["👥 {esc(u.full_name)}"])')
                if a.unit_type == "ORGANIZATION":
                    lines.append(f"  {nid('ORG', a.unit_id)} -->|موظف| {pid}")
                elif a.unit_type == "DIRECTORATE":
                    lines.append(f"  {nid('DIR', a.unit_id)} -->|موظف| {pid}")
                elif a.unit_type == "UNIT":
                    lines.append(f"  {nid('UNI', a.unit_id)} -->|موظف| {pid}")
                elif a.unit_type == "DEPARTMENT":
                    lines.append(f"  {nid('DEP', a.unit_id)} -->|موظف| {pid}")
                elif a.unit_type == "SECTION":
                    lines.append(f"  {nid('SEC', a.unit_id)} -->|موظف| {pid}")
                elif a.unit_type == "TEAM":
                    lines.append(f"  {nid('TEA', a.unit_id)} -->|موظف| {pid}")

        txt = "\n".join(lines) + "\n"
        bio = BytesIO(txt.encode("utf-8"))
        bio.seek(0)
        return send_file(bio, as_attachment=True, download_name="org_structure.mmd", mimetype="text/plain")

    flash("صيغة تصدير غير مدعومة.", "warning")
    return redirect(url_for("portal.hr_org_structure"))


# -------------------------
# HR: Masterdata (Schedules / Leave Types / Permission Types)
# -------------------------


def _legacy_org_locked() -> bool:
    return str(_setting_get('ORG_LEGACY_LOCKED', '0') or '0').strip() == '1'


def _setting_get(key: str, default: str | None = None) -> str | None:
    row = SystemSetting.query.filter_by(key=key).first()
    return (row.value if row else default)


def _setting_get_int(key: str, default: int = 0) -> int:
    """Read an integer SystemSetting safely.

    Returns `default` when the setting is missing, empty, or not a valid integer.
    """
    v = _setting_get(key, None)
    if v is None:
        return default
    try:
        v = str(v).strip()
        if v == "":
            return default
        return int(float(v))
    except Exception:
        return default


def _setting_set(key: str, value: str | None):
    row = SystemSetting.query.filter_by(key=key).first()
    if not row:
        row = SystemSetting(key=key, value=(value or ""))
        db.session.add(row)
    else:
        row.value = (value or "")


def _parse_hhmm(s: str | None):
    if not s:
        return None
    s = s.strip()
    if len(s) != 5 or ":" not in s:
        return None
    try:
        hh, mm = s.split(":")
        return int(hh) * 60 + int(mm)
    except Exception:
        return None


# -------------------------
# HR: Work policies helpers
# -------------------------


_WEEKDAYS_AR = [
    (6, "الأحد"),
    (0, "الإثنين"),
    (1, "الثلاثاء"),
    (2, "الأربعاء"),
    (3, "الخميس"),
    (4, "الجمعة"),
    (5, "السبت"),
]


def _days_mask_from_list(values) -> int:
    mask = 0
    for v in (values or []):
        try:
            i = int(str(v).strip())
        except Exception:
            continue
        if 0 <= i <= 6:
            mask |= (1 << i)
    return mask


def _days_mask_to_names(mask: int | None) -> list[str]:
    if mask is None:
        return []
    names = []
    for i, name in _WEEKDAYS_AR:
        if mask & (1 << i):
            names.append(name)
    return names


def _work_policy_days_label(p: WorkPolicy | None) -> str:
    if not p:
        return "—"
    dp = (p.days_policy or "FIXED").upper()
    if dp == "HYBRID_WEEKLY_QUOTA":
        off = p.hybrid_office_days or 0
        rem = p.hybrid_remote_days or 0
        return f"مرن أسبوعيًا ({off} مكتب + {rem} عن بعد)"
    # FIXED
    names = _days_mask_to_names(p.fixed_days_mask)
    if not names:
        return "ثابت"
    return "ثابت (" + "، ".join(names) + ")"


def _work_policy_place_label(p: WorkPolicy | None) -> str:
    if not p:
        return "—"
    lp = (p.location_policy or "ONSITE").upper()
    if lp == "REMOTE":
        return "Remote فقط"
    if lp == "HYBRID":
        return "Hybrid"
    return "On-site فقط"


def _ensure_work_policy_tables():
    """Create missing *new* tables (for users on existing SQLite DBs without migrations)."""
    try:
        # touch
        WorkPolicy.query.limit(1).all()
        WorkAssignment.query.limit(1).all()
    except OperationalError:
        try:
            db.session.rollback()
        except Exception:
            pass
        try:
            db.create_all()
        except Exception:
            pass


def _weekday_of(day_str: str) -> int:
    # day_str: YYYY-MM-DD
    try:
        d = date.fromisoformat(day_str)
        return d.weekday()
    except Exception:
        return 0


def _portal_department_id_for_user(user_id: int) -> int | None:
    """Resolve the employee's *Portal* department (for schedule assignments by department).

    Uses the user's primary OrgUnitAssignment (TEAM/SECTION/DEPARTMENT) and walks up to DEPARTMENT.
    """
    try:
        rows = (
            OrgUnitAssignment.query
            .filter_by(user_id=user_id)
            .order_by(OrgUnitAssignment.is_primary.desc(), OrgUnitAssignment.id.desc())
            .all()
        )
    except OperationalError:
        return None
    except Exception:
        return None

    if not rows:
        return None
    a = next((r for r in rows if r.is_primary), rows[0])
    if not a or not a.unit_type or not a.unit_id:
        return None

    ut = (a.unit_type or "").upper()
    try:
        if ut == "DEPARTMENT":
            return int(a.unit_id)
        if ut == "SECTION":
            sec = Section.query.get(int(a.unit_id))
            return int(sec.department_id) if sec and sec.department_id else None
        if ut == "TEAM":
            tm = Team.query.get(int(a.unit_id))
            if not tm or not tm.section_id:
                return None
            sec = Section.query.get(int(tm.section_id))
            return int(sec.department_id) if sec and sec.department_id else None
    except Exception:
        return None
    return None


def _effective_schedule_for_user(user_id: int, day_str: str) -> WorkSchedule | None:
    # 0) New: work assignments (user/role/department) with date ranges
    try:
        _ensure_work_policy_tables()
        u = User.query.get(user_id)
        role = (getattr(u, "role", None) or "").strip() or None
        dept_id = _portal_department_id_for_user(user_id)

        conds = [and_(WorkAssignment.target_type == "USER", WorkAssignment.target_user_id == user_id)]
        if role:
            conds.append(and_(WorkAssignment.target_type == "ROLE", WorkAssignment.target_role == role))
        if dept_id:
            conds.append(and_(WorkAssignment.target_type == "DEPARTMENT", WorkAssignment.target_department_id == dept_id))

        candidates = (
            WorkAssignment.query
            .filter(WorkAssignment.is_active == True)
            .filter(or_(*conds))
            .order_by(WorkAssignment.start_date.desc().nullslast(), WorkAssignment.id.desc())
            .all()
        )

        def _prio(tt: str) -> int:
            tt = (tt or "").upper()
            if tt == "USER":
                return 3
            if tt == "ROLE":
                return 2
            if tt == "DEPARTMENT":
                return 1
            return 0

        best = None
        best_key = None
        for a in candidates:
            if a.start_date and a.start_date > day_str:
                continue
            if a.end_date and a.end_date < day_str:
                continue
            key = (_prio(a.target_type), a.start_date or "", a.id)
            if best_key is None or key > best_key:
                best_key = key
                best = a

        if best:
            return WorkSchedule.query.get(best.schedule_id)
    except Exception:
        # keep old logic if anything fails
        pass

    # 1) legacy assignment range (per-user only)
    ass = (
        EmployeeScheduleAssignment.query
        .filter_by(user_id=user_id, is_active=True)
        .order_by(EmployeeScheduleAssignment.start_date.desc().nullslast())
        .all()
    )
    for a in ass:
        if a.start_date and a.start_date > day_str:
            continue
        if a.end_date and a.end_date < day_str:
            continue
        return WorkSchedule.query.get(a.schedule_id)

    # 2) default schedule id
    default_id = _setting_get("HR_DEFAULT_SCHEDULE_ID")
    if default_id and str(default_id).isdigit():
        return WorkSchedule.query.get(int(default_id))
    return None


@portal_bp.route("/hr/masterdata")
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_masterdata_index():
    schedules = WorkSchedule.query.order_by(WorkSchedule.id.desc()).all()
    perm_types = HRPermissionType.query.order_by(HRPermissionType.code.asc()).all()
    leave_types = HRLeaveType.query.order_by(HRLeaveType.code.asc()).all()
    default_schedule_id = _setting_get("HR_DEFAULT_SCHEDULE_ID")
    return render_template(
        "portal/hr/masterdata_index.html",
        schedules=schedules,
        perm_types=perm_types,
        leave_types=leave_types,
        default_schedule_id=default_schedule_id,
    )


# -------------------------
# HR: Employee File Lookups (Masterdata)
# -------------------------

@portal_bp.route("/hr/masterdata/employee-lookups")
@login_required
@_perm_any(HR_MASTERDATA_MANAGE, HR_REPORTS_VIEW)
def hr_employee_lookups_index():
    # counts
    counts = {}
    try:
        rows = (
            db.session.query(HRLookupItem.category, func.count(HRLookupItem.id))
            .group_by(HRLookupItem.category)
            .all()
        )
        counts = {c: int(n) for c, n in rows}
    except Exception:
        counts = {}
    categories = [(k, EMP_LOOKUP_LABEL.get(k, k), counts.get(k, 0)) for k, _ in EMP_LOOKUP_CATEGORIES]
    return render_template("portal/hr/employee_lookups_index.html", categories=categories, counts=counts)


def _lookup_category_or_404(category: str) -> str:
    cat = (category or '').upper().strip()
    keys = {k for k, _ in EMP_LOOKUP_CATEGORIES}
    if cat not in keys:
        abort(404)
    return cat


@portal_bp.route("/hr/masterdata/employee-lookups/<string:category>", methods=["GET", "POST"])
@login_required
@_perm_any(HR_MASTERDATA_MANAGE, HR_REPORTS_VIEW)
def hr_employee_lookups_category(category: str):
    cat = _lookup_category_or_404(category)

    if request.method == 'POST':
        # Allow view via HR_REPORTS_VIEW, but restrict mutations to HR_MASTERDATA_MANAGE
        if not current_user.has_perm(HR_MASTERDATA_MANAGE):
            abort(403)
        code = (request.form.get('code') or '').strip()
        name_ar = (request.form.get('name_ar') or '').strip()
        name_en = (request.form.get('name_en') or '').strip() or None
        sort_order = (request.form.get('sort_order') or '').strip()
        is_active = bool(request.form.get('is_active'))

        if not code or not name_ar:
            flash('الرمز والاسم العربي مطلوبان.', 'danger')
            return redirect(request.url)

        try:
            so = int(sort_order) if sort_order else 0
        except Exception:
            so = 0

        item = HRLookupItem.query.filter_by(category=cat, code=code).first()
        if item:
            item.name_ar = name_ar
            item.name_en = name_en
            item.sort_order = so
            item.is_active = is_active
            item.updated_at = datetime.utcnow()
        else:
            item = HRLookupItem(
                category=cat,
                code=code,
                name_ar=name_ar,
                name_en=name_en,
                sort_order=so,
                is_active=is_active,
                created_at=datetime.utcnow(),
                updated_at=datetime.utcnow(),
            )
            db.session.add(item)

        _portal_audit(
            action='HR_EMP_LOOKUP_UPSERT',
            note=f"{cat}: {code}",
            target_type='HR_LOOKUP_ITEM',
            target_id=0,
        )
        try:
            db.session.commit()
            flash('تم الحفظ.', 'success')
        except Exception:
            db.session.rollback()
            flash('تعذر الحفظ.', 'danger')
        return redirect(request.url)

    items = HRLookupItem.query.filter_by(category=cat).order_by(
        HRLookupItem.sort_order.asc(),
        HRLookupItem.name_ar.asc(),
        HRLookupItem.id.asc(),
    ).all()
    return render_template(
        "portal/hr/employee_lookups_category.html",
        # preferred names
        category_key=cat,
        category_name=EMP_LOOKUP_LABEL.get(cat, cat),
        # backward-compatible names (older templates)
        category=cat,
        label=EMP_LOOKUP_LABEL.get(cat, cat),
        items=items,
    )


@portal_bp.route("/hr/masterdata/employee-lookups/<string:category>/<int:item_id>/edit", methods=["GET", "POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_employee_lookup_edit(category: str, item_id: int):
    cat = _lookup_category_or_404(category)
    item = HRLookupItem.query.filter_by(id=item_id, category=cat).first_or_404()

    if request.method == 'POST':
        code = (request.form.get('code') or '').strip()
        name_ar = (request.form.get('name_ar') or '').strip()
        name_en = (request.form.get('name_en') or '').strip() or None
        sort_order = (request.form.get('sort_order') or '').strip()
        is_active = bool(request.form.get('is_active'))

        if not code or not name_ar:
            flash('الرمز والاسم العربي مطلوبان.', 'danger')
            return redirect(request.url)

        try:
            so = int(sort_order) if sort_order else 0
        except Exception:
            so = 0

        # prevent duplicate code in same category
        dup = HRLookupItem.query.filter(HRLookupItem.category == cat, HRLookupItem.code == code, HRLookupItem.id != item.id).first()
        if dup:
            flash('هذا الرمز مستخدم مسبقًا ضمن نفس الفئة.', 'danger')
            return redirect(request.url)

        item.code = code
        item.name_ar = name_ar
        item.name_en = name_en
        item.sort_order = so
        item.is_active = is_active
        item.updated_at = datetime.utcnow()

        _portal_audit('HR_EMP_LOOKUP_UPDATE', f"{cat}: {item.code}", target_type='HR_LOOKUP_ITEM', target_id=item.id)
        try:
            db.session.commit()
            flash('تم التحديث.', 'success')
            return redirect(url_for('portal.hr_employee_lookups_category', category=cat))
        except Exception:
            db.session.rollback()
            flash('تعذر الحفظ.', 'danger')



    return render_template(
        "portal/hr/employee_lookups_item_edit.html",
        # preferred names
        category_key=cat,
        category_name=EMP_LOOKUP_LABEL.get(cat, cat),
        # backward-compatible names (older templates)
        category=cat,
        label=EMP_LOOKUP_LABEL.get(cat, cat),
        item=item,
    )


@portal_bp.route("/hr/masterdata/employee-lookups/<string:category>/<int:item_id>/delete", methods=["POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_employee_lookup_delete(category: str, item_id: int):
    cat = _lookup_category_or_404(category)
    item = HRLookupItem.query.filter_by(id=item_id, category=cat).first_or_404()

    # Prefer soft delete to avoid FK issues
    item.is_active = False
    item.updated_at = datetime.utcnow()

    _portal_audit('HR_EMP_LOOKUP_DISABLE', f"{cat}: {item.code}", target_type='HR_LOOKUP_ITEM', target_id=item.id)
    try:
        db.session.commit()
        flash('تم التعطيل.', 'success')
    except Exception:
        db.session.rollback()
        flash('تعذر العملية.', 'danger')

    return redirect(url_for('portal.hr_employee_lookups_category', category=cat))


@portal_bp.route("/hr/masterdata/employee-lookups/<string:category>/export")
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_employee_lookups_export(category: str):
    cat = _lookup_category_or_404(category)
    items = HRLookupItem.query.filter_by(category=cat).order_by(HRLookupItem.sort_order.asc(), HRLookupItem.name_ar.asc(), HRLookupItem.id.asc()).all()

    headers = ["code", "name_ar", "name_en", "sort_order", "is_active"]
    rows = [[i.code, i.name_ar, i.name_en or "", i.sort_order or 0, 1 if i.is_active else 0] for i in items]

    from utils.excel import make_xlsx_bytes
    xbytes = make_xlsx_bytes(cat, headers, rows)
    fn = f"employee_lookup_{cat}.xlsx"
    return send_file(BytesIO(xbytes), as_attachment=True, download_name=fn, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


@portal_bp.route("/hr/masterdata/employee-lookups/<string:category>/template")
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_employee_lookups_template(category: str):
    cat = _lookup_category_or_404(category)
    headers = ["code", "name_ar", "name_en", "sort_order", "is_active"]
    from utils.excel import make_xlsx_bytes
    xbytes = make_xlsx_bytes(cat, headers, [])
    fn = f"employee_lookup_{cat}_template.xlsx"
    return send_file(BytesIO(xbytes), as_attachment=True, download_name=fn, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


@portal_bp.route("/hr/masterdata/employee-lookups/<string:category>/import", methods=["POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_employee_lookups_import(category: str):
    cat = _lookup_category_or_404(category)

    f = request.files.get('file')
    if not f or not getattr(f, 'filename', ''):
        flash('اختر ملف Excel.', 'danger')
        return redirect(url_for('portal.hr_employee_lookups_category', category=cat))

    mode = (request.form.get('mode') or 'upsert').strip().lower()  # upsert / replace / safe_replace

    from utils.importer import read_excel_rows

    try:
        rows = read_excel_rows(f, required_headers={'code', 'name_ar'})
    except Exception:
        flash('تعذر قراءة ملف Excel. تأكد من الأعمدة: code, name_ar ...', 'danger')
        return redirect(url_for('portal.hr_employee_lookups_category', category=cat))

    # Safe replace: deactivate all existing, then upsert from file.
    if mode in {'replace', 'safe_replace'}:
        try:
            HRLookupItem.query.filter_by(category=cat).update({HRLookupItem.is_active: False, HRLookupItem.updated_at: datetime.utcnow()})
            db.session.flush()
        except Exception:
            db.session.rollback()

    upserted = 0
    for r in rows:
        code = (r.get('code') or '').strip()
        name_ar = (r.get('name_ar') or '').strip()
        if not code or not name_ar:
            continue
        name_en = (r.get('name_en') or '').strip() or None
        try:
            sort_order = int((r.get('sort_order') or 0) or 0)
        except Exception:
            sort_order = 0
        is_active = str(r.get('is_active') or '').strip()
        if is_active in {'0', 'false', 'False', 'no', 'No'}:
            active = False
        elif is_active in {'1', 'true', 'True', 'yes', 'Yes'}:
            active = True
        else:
            active = True

        item = HRLookupItem.query.filter_by(category=cat, code=code).first()
        if item:
            item.name_ar = name_ar
            item.name_en = name_en
            item.sort_order = sort_order
            item.is_active = active
            item.updated_at = datetime.utcnow()
        else:
            item = HRLookupItem(
                category=cat,
                code=code,
                name_ar=name_ar,
                name_en=name_en,
                sort_order=sort_order,
                is_active=active,
                created_at=datetime.utcnow(),
                updated_at=datetime.utcnow(),
            )
            db.session.add(item)
        upserted += 1

    _portal_audit('HR_EMP_LOOKUP_IMPORT', f"{cat}: mode={mode} upserted={upserted}", target_type='HR_LOOKUP_ITEM', target_id=0)

    try:
        db.session.commit()
        flash(f"تم الاستيراد بنجاح. عدد الصفوف: {upserted}", 'success')
    except Exception:
        db.session.rollback()
        flash('تعذر الاستيراد.', 'danger')

    return redirect(url_for('portal.hr_employee_lookups_category', category=cat))


@portal_bp.route("/hr/masterdata/default-schedule", methods=["POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_masterdata_set_default_schedule():
    sid = (request.form.get("schedule_id") or "").strip()
    if sid and not sid.isdigit():
        flash("قيمة غير صحيحة.", "danger")
        return redirect(url_for("portal.hr_masterdata_index"))
    _setting_set("HR_DEFAULT_SCHEDULE_ID", sid or "")
    _portal_audit("HR_DEFAULT_SCHEDULE_SET", f"Default schedule={sid}", target_type="SETTING", target_id=0)
    db.session.commit()
    flash("تم حفظ الجدول الافتراضي.", "success")
    return redirect(url_for("portal.hr_masterdata_index"))


@portal_bp.route("/hr/masterdata/schedule/new", methods=["POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_schedule_new():
    name = (request.form.get("name") or "").strip()
    kind = (request.form.get("kind") or "FIXED").strip().upper() or "FIXED"
    start_time = (request.form.get("start_time") or "").strip() or None
    end_time = (request.form.get("end_time") or "").strip() or None

    required_minutes = (request.form.get("required_minutes") or "").strip()
    break_minutes = (request.form.get("break_minutes") or "0").strip()
    grace_minutes = (request.form.get("grace_minutes") or "0").strip()
    overtime_thr = (request.form.get("overtime_threshold_minutes") or "").strip()

    if not name:
        flash("اسم الجدول مطلوب.", "danger")
        return redirect(url_for("portal.hr_masterdata_index"))

    if start_time and _parse_hhmm(start_time) is None:
        flash("بداية الدوام غير صحيحة (HH:MM).", "danger")
        return redirect(url_for("portal.hr_masterdata_index"))
    if end_time and _parse_hhmm(end_time) is None:
        flash("نهاية الدوام غير صحيحة (HH:MM).", "danger")
        return redirect(url_for("portal.hr_masterdata_index"))

    def _i(v, default=None):
        v = (v or "").strip()
        if not v:
            return default
        try:
            return int(v)
        except Exception:
            return default

    ws = WorkSchedule(
        name=name,
        kind=kind,
        start_time=start_time,
        end_time=end_time,
        required_minutes=_i(required_minutes, None),
        break_minutes=_i(break_minutes, 0) or 0,
        grace_minutes=_i(grace_minutes, 0) or 0,
        overtime_threshold_minutes=_i(overtime_thr, None),
        is_active=True,
        created_at=datetime.utcnow(),
        created_by_id=current_user.id,
    )
    db.session.add(ws)
    _portal_audit("HR_SCHEDULE_CREATE", f"إنشاء جدول: {name}", target_type="WORK_SCHEDULE", target_id=0)
    db.session.commit()
    flash("تم إنشاء الجدول.", "success")
    return redirect(url_for("portal.hr_masterdata_index"))


@portal_bp.route("/hr/masterdata/schedule/<int:schedule_id>/toggle", methods=["POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_schedule_toggle(schedule_id: int):
    ws = WorkSchedule.query.get_or_404(schedule_id)
    ws.is_active = not bool(ws.is_active)
    _portal_audit("HR_SCHEDULE_TOGGLE", f"تفعيل/تعطيل جدول: {ws.name}", target_type="WORK_SCHEDULE", target_id=schedule_id)
    db.session.commit()
    flash("تم تحديث حالة الجدول.", "success")
    return redirect(url_for("portal.hr_masterdata_index"))


@portal_bp.route("/hr/masterdata/schedule/<int:schedule_id>/edit", methods=["GET", "POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_schedule_edit(schedule_id: int):
    ws = WorkSchedule.query.get_or_404(schedule_id)
    if request.method == "POST":
        name = (request.form.get("name") or "").strip()
        kind = (request.form.get("kind") or ws.kind or "FIXED").strip().upper()
        start_time = (request.form.get("start_time") or "").strip() or None
        end_time = (request.form.get("end_time") or "").strip() or None
        required_minutes = (request.form.get("required_minutes") or "").strip()
        break_minutes = (request.form.get("break_minutes") or "0").strip()
        grace_minutes = (request.form.get("grace_minutes") or "0").strip()
        overtime_thr = (request.form.get("overtime_threshold_minutes") or "").strip()
        is_active = (request.form.get("is_active") or "") == "1"

        if not name:
            flash("اسم الجدول مطلوب.", "danger")
            return redirect(request.url)

        if start_time and _parse_hhmm(start_time) is None:
            flash("بداية الدوام غير صحيحة (HH:MM).", "danger")
            return redirect(request.url)
        if end_time and _parse_hhmm(end_time) is None:
            flash("نهاية الدوام غير صحيحة (HH:MM).", "danger")
            return redirect(request.url)

        def _i(v, default=None):
            v = (v or "").strip()
            if not v:
                return default
            try:
                return int(v)
            except Exception:
                return default

        ws.name = name
        ws.kind = kind
        ws.start_time = start_time
        ws.end_time = end_time
        ws.required_minutes = _i(required_minutes, None)
        ws.break_minutes = _i(break_minutes, 0) or 0
        ws.grace_minutes = _i(grace_minutes, 0) or 0
        ws.overtime_threshold_minutes = _i(overtime_thr, None)
        ws.is_active = is_active

        _portal_audit("HR_SCHEDULE_UPDATE", f"تعديل جدول: {ws.name}", target_type="WORK_SCHEDULE", target_id=ws.id)
        db.session.commit()
        flash("تم حفظ التعديلات.", "success")
        return redirect(url_for("portal.hr_masterdata_index"))

    return render_template("portal/hr/schedule_edit.html", row=ws)


@portal_bp.route("/hr/masterdata/schedule/<int:schedule_id>/delete", methods=["POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_schedule_delete(schedule_id: int):
    ws = WorkSchedule.query.get_or_404(schedule_id)
    try:
        used = False
        try:
            used = used or (EmployeeScheduleAssignment.query.filter_by(schedule_id=ws.id).first() is not None)
        except Exception:
            pass
        try:
            used = used or (WorkAssignment.query.filter_by(schedule_id=ws.id).first() is not None)
        except Exception:
            pass
        try:
            used = used or (WorkScheduleDay.query.filter_by(schedule_id=ws.id).first() is not None)
        except Exception:
            pass

        if used:
            ws.is_active = False
            # if it's default, clear default
            try:
                default_id = _setting_get("HR_DEFAULT_SCHEDULE_ID")
                if default_id and str(default_id).isdigit() and int(default_id) == ws.id:
                    _setting_set("HR_DEFAULT_SCHEDULE_ID", "")
            except Exception:
                pass
            _portal_audit("HR_SCHEDULE_DISABLE", f"تعطيل جدول مستخدم: {ws.name}", target_type="WORK_SCHEDULE", target_id=ws.id)
            db.session.commit()
            flash("لا يمكن حذف الجدول لأنه مستخدم (تعيينات/أيام/افتراضي). تم تعطيله بدلاً من ذلك.", "warning")
            return redirect(url_for("portal.hr_masterdata_index"))

        # delete days first
        WorkScheduleDay.query.filter_by(schedule_id=ws.id).delete(synchronize_session=False)
        _portal_audit("HR_SCHEDULE_DELETE", f"حذف جدول: {ws.name}", target_type="WORK_SCHEDULE", target_id=ws.id)
        db.session.delete(ws)
        db.session.commit()
        flash("تم حذف الجدول.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"تعذر حذف الجدول: {e}", "danger")
    return redirect(url_for("portal.hr_masterdata_index"))


@portal_bp.route("/hr/masterdata/schedule/<int:schedule_id>/day", methods=["POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_schedule_set_day(schedule_id: int):
    ws = WorkSchedule.query.get_or_404(schedule_id)
    weekday = (request.form.get("weekday") or "").strip()
    start_time = (request.form.get("start_time") or "").strip() or None
    end_time = (request.form.get("end_time") or "").strip() or None
    break_minutes = (request.form.get("break_minutes") or "0").strip()
    grace_minutes = (request.form.get("grace_minutes") or "0").strip()

    if not weekday.isdigit() or int(weekday) not in range(0, 7):
        flash("اليوم غير صحيح.", "danger")
        return redirect(url_for("portal.hr_masterdata_index"))

    if start_time and _parse_hhmm(start_time) is None:
        flash("بداية الدوام غير صحيحة (HH:MM).", "danger")
        return redirect(url_for("portal.hr_masterdata_index"))
    if end_time and _parse_hhmm(end_time) is None:
        flash("نهاية الدوام غير صحيحة (HH:MM).", "danger")
        return redirect(url_for("portal.hr_masterdata_index"))

    def _i(v, default=0):
        try:
            return int((v or "").strip() or default)
        except Exception:
            return default

    wd = WorkScheduleDay.query.filter_by(schedule_id=ws.id, weekday=int(weekday)).first()
    if not wd:
        wd = WorkScheduleDay(schedule_id=ws.id, weekday=int(weekday))
        db.session.add(wd)

    wd.start_time = start_time
    wd.end_time = end_time
    wd.break_minutes = _i(break_minutes, 0) or 0
    wd.grace_minutes = _i(grace_minutes, 0) or 0

    _portal_audit("HR_SCHEDULE_DAY_SET", f"تعديل يوم في جدول: {ws.name} weekday={weekday}", target_type="WORK_SCHEDULE", target_id=ws.id)
    db.session.commit()
    flash("تم حفظ إعدادات اليوم.", "success")
    return redirect(url_for("portal.hr_masterdata_index"))


@portal_bp.route("/hr/masterdata/permission-type/new", methods=["POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_permission_type_new():
    code = (request.form.get("code") or "").strip().upper()
    name_ar = (request.form.get("name_ar") or "").strip()
    name_en = (request.form.get("name_en") or "").strip() or None
    requires_approval = (request.form.get("requires_approval") or "") == "1"
    max_hours = (request.form.get("max_hours") or "").strip()
    counts_as_work = (request.form.get("counts_as_work") or "") == "1"

    if not code or not name_ar:
        flash("الكود والاسم عربي مطلوبان.", "danger")
        return redirect(url_for("portal.hr_masterdata_index"))

    mh = int(max_hours) if max_hours.isdigit() else None

    row = HRPermissionType.query.filter_by(code=code).first()
    if row:
        flash("الكود موجود مسبقًا.", "danger")
        return redirect(url_for("portal.hr_masterdata_index"))

    row = HRPermissionType(
        code=code,
        name_ar=name_ar,
        name_en=name_en,
        requires_approval=requires_approval,
        max_hours=mh,
        counts_as_work=counts_as_work,
        is_active=True,
        created_at=datetime.utcnow(),
        created_by_id=current_user.id,
    )
    db.session.add(row)
    _portal_audit("HR_PERMISSION_TYPE_CREATE", f"إنشاء نوع مغادرة: {code}", target_type="HR_PERMISSION_TYPE", target_id=0)
    db.session.commit()
    flash("تم إضافة نوع المغادرة.", "success")
    return redirect(url_for("portal.hr_masterdata_index"))


@portal_bp.route("/hr/masterdata/permission-type/<int:pt_id>/toggle", methods=["POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_permission_type_toggle(pt_id: int):
    row = HRPermissionType.query.get_or_404(pt_id)
    row.is_active = not bool(row.is_active)
    _portal_audit("HR_PERMISSION_TYPE_TOGGLE", f"تفعيل/تعطيل نوع مغادرة: {row.code}", target_type="HR_PERMISSION_TYPE", target_id=pt_id)
    db.session.commit()
    flash("تم تحديث حالة النوع.", "success")
    return redirect(url_for("portal.hr_masterdata_index"))



@portal_bp.route("/hr/masterdata/permission-type/<int:pt_id>/edit", methods=["GET", "POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_permission_type_edit(pt_id: int):
    row = HRPermissionType.query.get_or_404(pt_id)
    if request.method == "POST":
        code = (request.form.get("code") or "").strip().upper()
        name_ar = (request.form.get("name_ar") or "").strip()
        name_en = (request.form.get("name_en") or "").strip() or None
        requires_approval = (request.form.get("requires_approval") or "") == "1"
        max_hours = (request.form.get("max_hours") or "").strip()
        counts_as_work = (request.form.get("counts_as_work") or "") == "1"
        is_active = (request.form.get("is_active") or "") == "1"

        if not code or not name_ar:
            flash("الكود والاسم عربي مطلوبان.", "danger")
            return redirect(request.url)

        other = HRPermissionType.query.filter(HRPermissionType.code == code, HRPermissionType.id != row.id).first()
        if other:
            flash("الكود موجود مسبقًا.", "danger")
            return redirect(request.url)

        mh = int(max_hours) if max_hours.isdigit() else None

        row.code = code
        row.name_ar = name_ar
        row.name_en = name_en
        row.requires_approval = requires_approval
        row.max_hours = mh
        row.counts_as_work = counts_as_work
        row.is_active = is_active

        _portal_audit("HR_PERMISSION_TYPE_UPDATE", f"تعديل نوع مغادرة: {row.code}", target_type="HR_PERMISSION_TYPE", target_id=pt_id)
        db.session.commit()
        flash("تم حفظ التعديلات.", "success")
        return redirect(url_for("portal.hr_masterdata_index"))

    return render_template("portal/hr/permission_type_edit.html", row=row)


@portal_bp.route("/hr/masterdata/permission-type/<int:pt_id>/delete", methods=["POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_permission_type_delete(pt_id: int):
    row = HRPermissionType.query.get_or_404(pt_id)
    try:
        used = HRPermissionRequest.query.filter_by(permission_type_id=row.id).first() is not None
        if used:
            row.is_active = False
            _portal_audit("HR_PERMISSION_TYPE_DISABLE", f"تعطيل نوع مغادرة مستخدم: {row.code}", target_type="HR_PERMISSION_TYPE", target_id=pt_id)
            db.session.commit()
            flash("لا يمكن حذف النوع لأنه مستخدم سابقاً، تم تعطيله بدلاً من ذلك.", "warning")
        else:
            _portal_audit("HR_PERMISSION_TYPE_DELETE", f"حذف نوع مغادرة: {row.code}", target_type="HR_PERMISSION_TYPE", target_id=pt_id)
            db.session.delete(row)
            db.session.commit()
            flash("تم حذف النوع.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"تعذر حذف النوع: {e}", "danger")
    return redirect(url_for("portal.hr_masterdata_index"))



@portal_bp.route("/hr/masterdata/leave-type/new", methods=["POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_leave_type_new():
    code = (request.form.get("code") or "").strip().upper()
    name_ar = (request.form.get("name_ar") or "").strip()
    name_en = (request.form.get("name_en") or "").strip() or None
    requires_approval = (request.form.get("requires_approval") or "") == "1"
    max_days = (request.form.get("max_days") or "").strip()
    default_balance = (request.form.get("default_balance_days") or "").strip()
    exc_max_days = (request.form.get("exception_max_days") or "").strip()
    exc_requires_hr = (request.form.get("exception_requires_hr") or "1") == "1"
    exc_requires_note = (request.form.get("exception_requires_note") or "0") == "1"
    is_external = (request.form.get("is_external") or "") == "1"
    requires_documents = (request.form.get("requires_documents") or "") == "1"
    documents_hint = (request.form.get("documents_hint") or "").strip() or None

    if not code or not name_ar:
        flash("الكود والاسم عربي مطلوبان.", "danger")
        return redirect(url_for("portal.hr_masterdata_index"))

    md = int(max_days) if max_days.isdigit() else None
    dbd = int(default_balance) if default_balance.isdigit() else None
    exc_md = int(exc_max_days) if exc_max_days.isdigit() else None
    if md and exc_md and exc_md <= md:
        exc_md = None

    row = HRLeaveType.query.filter_by(code=code).first()
    if row:
        flash("الكود موجود مسبقًا.", "danger")
        return redirect(url_for("portal.hr_masterdata_index"))

    row = HRLeaveType(
        code=code,
        name_ar=name_ar,
        name_en=name_en,
        requires_approval=requires_approval,
        max_days=md,
        default_balance_days=dbd,
        exception_max_days=exc_md,
        exception_requires_hr=exc_requires_hr,
        exception_requires_note=exc_requires_note,
        requires_documents=requires_documents,
        documents_hint=documents_hint,
        is_external=is_external,
        is_active=True,
        created_at=datetime.utcnow(),
        created_by_id=current_user.id,
    )
    db.session.add(row)
    _portal_audit("HR_LEAVE_TYPE_CREATE", f"إنشاء نوع إجازة: {code}", target_type="HR_LEAVE_TYPE", target_id=0)
    db.session.commit()
    flash("تم إضافة نوع الإجازة.", "success")
    return redirect(url_for("portal.hr_masterdata_index"))


@portal_bp.route("/hr/masterdata/leave-type/<int:lt_id>/toggle", methods=["POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_leave_type_toggle(lt_id: int):
    row = HRLeaveType.query.get_or_404(lt_id)
    row.is_active = not bool(row.is_active)
    _portal_audit("HR_LEAVE_TYPE_TOGGLE", f"تفعيل/تعطيل نوع إجازة: {row.code}", target_type="HR_LEAVE_TYPE", target_id=lt_id)
    db.session.commit()
    flash("تم تحديث حالة النوع.", "success")
    return redirect(url_for("portal.hr_masterdata_index"))


@portal_bp.route("/hr/masterdata/leave-type/<int:lt_id>/edit", methods=["GET", "POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_leave_type_edit(lt_id: int):
    row = HRLeaveType.query.get_or_404(lt_id)
    if request.method == "POST":
        code = (request.form.get("code") or "").strip().upper()
        name_ar = (request.form.get("name_ar") or "").strip()
        name_en = (request.form.get("name_en") or "").strip() or None
        requires_approval = (request.form.get("requires_approval") or "") == "1"
        max_days = (request.form.get("max_days") or "").strip()
        default_balance = (request.form.get("default_balance_days") or "").strip()
        exc_max_days = (request.form.get("exception_max_days") or "").strip()
        exc_requires_hr = (request.form.get("exception_requires_hr") or "1") == "1"
        exc_requires_note = (request.form.get("exception_requires_note") or "0") == "1"
        is_external = (request.form.get("is_external") or "") == "1"
        requires_documents = (request.form.get("requires_documents") or "") == "1"
        documents_hint = (request.form.get("documents_hint") or "").strip() or None
        is_active = (request.form.get("is_active") or "") == "1"

        if not code or not name_ar:
            flash("الكود والاسم عربي مطلوبان.", "danger")
            return redirect(request.url)

        other = HRLeaveType.query.filter(HRLeaveType.code == code, HRLeaveType.id != row.id).first()
        if other:
            flash("الكود موجود مسبقًا.", "danger")
            return redirect(request.url)

        md = int(max_days) if max_days.isdigit() else None
        dbd = int(default_balance) if default_balance.isdigit() else None
        exc_md = int(exc_max_days) if exc_max_days.isdigit() else None
        if md and exc_md and exc_md <= md:
            exc_md = None

        row.code = code
        row.name_ar = name_ar
        row.name_en = name_en
        row.requires_approval = requires_approval
        row.max_days = md
        row.default_balance_days = dbd
        row.exception_max_days = exc_md
        row.exception_requires_hr = exc_requires_hr
        row.exception_requires_note = exc_requires_note
        row.requires_documents = requires_documents
        row.documents_hint = documents_hint
        row.is_external = is_external
        row.is_active = is_active

        _portal_audit("HR_LEAVE_TYPE_UPDATE", f"تعديل نوع إجازة: {row.code}", target_type="HR_LEAVE_TYPE", target_id=lt_id)
        db.session.commit()
        flash("تم حفظ التعديلات.", "success")
        return redirect(url_for("portal.hr_masterdata_index"))

    entitlements = HRLeaveGradeEntitlement.query.filter_by(leave_type_id=row.id).order_by(HRLeaveGradeEntitlement.grade.asc()).all()
    return render_template("portal/hr/leave_type_edit.html", row=row, entitlements=entitlements)



@portal_bp.route("/hr/masterdata/leave-type/<int:lt_id>/entitlements", methods=["POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_leave_type_entitlement_upsert(lt_id: int):
    """Upsert per-grade entitlement days for a leave type."""
    row = HRLeaveType.query.get_or_404(lt_id)
    grade = (request.form.get("grade") or "").strip().upper()
    allowed_days_s = (request.form.get("allowed_days") or "").strip()

    if not grade:
        flash("أدخل الدرجة.", "danger")
        return redirect(url_for("portal.hr_leave_type_edit", lt_id=lt_id) + "#entitlements")

    try:
        allowed_days = int(allowed_days_s)
        if allowed_days < 0:
            allowed_days = 0
    except Exception:
        flash("أدخل عدد أيام صحيح.", "danger")
        return redirect(url_for("portal.hr_leave_type_edit", lt_id=lt_id) + "#entitlements")

    ge = HRLeaveGradeEntitlement.query.filter_by(leave_type_id=row.id, grade=grade).first()
    if ge is None:
        ge = HRLeaveGradeEntitlement(
            leave_type_id=row.id,
            grade=grade,
            allowed_days=allowed_days,
            created_at=datetime.utcnow(),
            created_by_id=current_user.id,
        )
        db.session.add(ge)
        _portal_audit("HR_LEAVE_GRADE_ENTITLEMENT_CREATE", f"إضافة استحقاق درجة {grade} لنوع {row.code}: {allowed_days}", target_type="HR_LEAVE_TYPE", target_id=lt_id)
    else:
        ge.allowed_days = allowed_days
        _portal_audit("HR_LEAVE_GRADE_ENTITLEMENT_UPDATE", f"تحديث استحقاق درجة {grade} لنوع {row.code}: {allowed_days}", target_type="HR_LEAVE_TYPE", target_id=lt_id)

    db.session.commit()
    flash("تم حفظ استحقاق الدرجة.", "success")
    return redirect(url_for("portal.hr_leave_type_edit", lt_id=lt_id) + "#entitlements")


@portal_bp.route("/hr/masterdata/leave-type/<int:lt_id>/entitlements/<int:ent_id>/delete", methods=["POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_leave_type_entitlement_delete(lt_id: int, ent_id: int):
    row = HRLeaveType.query.get_or_404(lt_id)
    ent = HRLeaveGradeEntitlement.query.get_or_404(ent_id)
    if ent.leave_type_id != row.id:
        abort(404)
    db.session.delete(ent)
    _portal_audit("HR_LEAVE_GRADE_ENTITLEMENT_DELETE", f"حذف استحقاق درجة {ent.grade} لنوع {row.code}", target_type="HR_LEAVE_TYPE", target_id=lt_id)
    db.session.commit()
    flash("تم حذف استحقاق الدرجة.", "success")
    return redirect(url_for("portal.hr_leave_type_edit", lt_id=lt_id) + "#entitlements")


@portal_bp.route("/hr/masterdata/leave-type/<int:lt_id>/delete", methods=["POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_leave_type_delete(lt_id: int):
    row = HRLeaveType.query.get_or_404(lt_id)
    try:
        used = HRLeaveRequest.query.filter_by(leave_type_id=row.id).first() is not None
        if used:
            row.is_active = False
            _portal_audit("HR_LEAVE_TYPE_DISABLE", f"تعطيل نوع إجازة مستخدم: {row.code}", target_type="HR_LEAVE_TYPE", target_id=lt_id)
            db.session.commit()
            flash("لا يمكن حذف النوع لأنه مستخدم سابقاً، تم تعطيله بدلاً من ذلك.", "warning")
        else:
            _portal_audit("HR_LEAVE_TYPE_DELETE", f"حذف نوع إجازة: {row.code}", target_type="HR_LEAVE_TYPE", target_id=lt_id)
            db.session.delete(row)
            db.session.commit()
            flash("تم حذف النوع.", "success")
    except Exception as e:
        db.session.rollback()
        flash(f"تعذر حذف النوع: {e}", "danger")
    return redirect(url_for("portal.hr_masterdata_index"))




# -------------------------


# -------------------------
# HR Leave Balances & Smart Alerts
# -------------------------

def _year_from_datestr(d: str | None) -> int | None:
    try:
        if not d:
            return None
        return int(str(d)[:4])
    except Exception:
        return None


def _leave_entitlement_days(user_id: int, lt: HRLeaveType, year: int) -> int:
    """Return entitlement days (annual allowed days) for a user/type/year.

    Priority:
      1) Explicit HRLeaveBalance row (per user/year)
      2) Per-grade entitlement (HRLeaveGradeEntitlement) using EmployeeFile.grade
      3) Leave type default_balance_days
      4) 0
    """
    # 1) explicit per-user/year balance
    try:
        row = HRLeaveBalance.query.filter_by(user_id=user_id, leave_type_id=lt.id, year=year).first()
        if row and row.total_days is not None:
            return int(row.total_days)
    except Exception:
        pass

    # 2) per-grade entitlement
    try:
        emp = EmployeeFile.query.filter_by(user_id=user_id).first()
        grade = (getattr(emp, 'grade', None) or '').strip().upper()
        if grade:
            ge = HRLeaveGradeEntitlement.query.filter_by(leave_type_id=lt.id, grade=grade).first()
            if ge and ge.allowed_days is not None:
                return int(ge.allowed_days)
    except Exception:
        pass

    # 3) leave type default
    try:
        if lt.default_balance_days is not None:
            return int(lt.default_balance_days)
    except Exception:
        pass

    return 0


def _end_of_month_str(year: int, month: int) -> str:
    """Return last day of month as YYYY-MM-DD."""
    try:
        if month == 12:
            last = date(year + 1, 1, 1) - timedelta(days=1)
        else:
            last = date(year, month + 1, 1) - timedelta(days=1)
        return last.strftime("%Y-%m-%d")
    except Exception:
        return f"{year:04d}-{month:02d}-28"


def _workday_hours() -> int:
    return _setting_get_int("HR_WORKDAY_HOURS", 8)


def _default_allowed_permission_hours() -> int:
    return _setting_get_int("HR_ALLOWED_PERMISSION_HOURS_MONTH", 6)


def _get_monthly_allowed_hours(user_id: int, year: int, month: int) -> int:
    try:
        row = HRMonthlyPermissionAllowance.query.filter_by(user_id=user_id, year=year, month=month).first()
        if row and row.allowed_hours is not None:
            return int(row.allowed_hours)
    except Exception:
        pass
    return int(_default_allowed_permission_hours())


def _permission_hours_in_range(user_id: int, start_day: str, end_day: str) -> int:
    """Sum approved permission hours (departures) in range [start_day, end_day].

    We exclude permission types that are marked counts_as_work.
    """
    if not start_day or not end_day or end_day < start_day:
        return 0
    try:
        q = (HRPermissionRequest.query
             .join(HRPermissionType, HRPermissionRequest.permission_type_id == HRPermissionType.id)
             .filter(HRPermissionRequest.user_id == user_id)
             .filter(HRPermissionRequest.status == 'APPROVED')
             .filter(HRPermissionType.counts_as_work == False)  # noqa: E712
             .filter(HRPermissionRequest.day >= start_day)
             .filter(HRPermissionRequest.day <= end_day))
        rows = q.all()
        total = 0
        for r in rows:
            try:
                total += int(r.hours or 0)
            except Exception:
                continue
        return int(total)
    except Exception:
        return 0


def _permission_excess_hours_year(user_id: int, year: int, as_of_str: str) -> int:
    """Compute total excess permission hours in a year up to as_of_str."""
    try:
        # If year is in the future, no excess.
        if as_of_str[:4] < f"{year:04d}":
            return 0
    except Exception:
        pass

    total_excess = 0
    for m in range(1, 13):
        m_start = f"{year:04d}-{m:02d}-01"
        m_end = _end_of_month_str(year, m)
        # Bound by as_of_str (day-by-day policy)
        end = min(m_end, as_of_str)
        if end < m_start:
            continue
        used = _permission_hours_in_range(user_id, m_start, end)
        allowed = _get_monthly_allowed_hours(user_id, year, m)
        total_excess += max(0, int(used) - int(allowed))
    return int(total_excess)


def _permission_excess_leave_type_id() -> int | None:
    """Leave type that will absorb excess permission hours (converted to leave-days)."""
    # explicit setting first
    try:
        v = (_setting_get('HR_EXCESS_PERM_DEDUCT_LEAVE_TYPE_ID') or '').strip()
        if v.isdigit():
            lt = HRLeaveType.query.get(int(v))
            if lt:
                return int(lt.id)
    except Exception:
        pass

    # fallback by code
    try:
        for code in ('PERSONAL', 'ANNUAL', 'ANNUAL_LEAVE'):
            lt = HRLeaveType.query.filter(func.upper(HRLeaveType.code) == code).first()
            if lt:
                return int(lt.id)
    except Exception:
        pass

    try:
        lt = HRLeaveType.query.filter_by(is_active=True).order_by(HRLeaveType.id.asc()).first()
        return int(lt.id) if lt else None
    except Exception:
        return None


def _leave_used_days_as_of(user_id: int, leave_type_id: int, year: int, as_of: date) -> float:
    """Compute used leave days for a given year up to a specific date (day-by-day)."""
    try:
        # Bound as_of to the requested year
        if as_of.year < year:
            return 0.0
        if as_of.year > year:
            as_of = date(year, 12, 31)

        y_start = date(year, 1, 1)
        y_end = date(year, 12, 31)
        as_of_str = as_of.strftime("%Y-%m-%d")

        total = 0.0
        q = (HRLeaveRequest.query
             .filter(HRLeaveRequest.user_id == user_id)
             .filter(HRLeaveRequest.leave_type_id == leave_type_id)
             .filter(HRLeaveRequest.status.in_(["APPROVED", "CANCELLED"]))
             .order_by(HRLeaveRequest.id.asc()))

        for r in q.all():
            # Only count CANCELLED requests if they were cancelled after approval
            if r.status == "CANCELLED":
                if (r.cancelled_from_status or "").upper() != "APPROVED":
                    continue

            start = _parse_yyyy_mm_dd(r.start_date)
            end = _parse_yyyy_mm_dd(r.end_date)
            if not start or not end:
                continue

            # Day-by-day: only count days up to as_of (and stop at cancellation effective date)
            effective_end = min(end, as_of)
            if r.status == "CANCELLED" and r.cancel_effective_date:
                ce = _parse_yyyy_mm_dd(r.cancel_effective_date)
                if ce:
                    effective_end = min(effective_end, ce)

            # Intersect with requested year
            s = max(start, y_start)
            e = min(effective_end, y_end)
            if e < s:
                continue

            days = (e - s).days + 1
            total += float(days)

        # Add permission-excess deduction (treated as fractional leave) on one selected leave type
        try:
            excess_lt_id = _permission_excess_leave_type_id()
            if excess_lt_id and int(excess_lt_id) == int(leave_type_id):
                workday_hours = _workday_hours()
                excess_hours = _permission_excess_hours_year(user_id, year, as_of_str)
                total += float(excess_hours) / float(workday_hours)
        except Exception:
            pass

        return float(total)
    except Exception:
        return 0.0


def _leave_used_days(user_id: int, leave_type_id: int, year: int) -> float:
    """Used leave days (day-by-day) based on approvals; future days are not deducted."""
    try:
        today = date.today()
        as_of = today
        if year < today.year:
            as_of = date(year, 12, 31)
        elif year > today.year:
            as_of = date(year, 1, 1) - timedelta(days=1)
        return _leave_used_days_as_of(user_id, leave_type_id, year, as_of)
    except Exception:
        return 0.0



def _hr_recipients_user_ids() -> list[int]:
    """Users considered HR recipients for automated HR alerts.

    We include:
      - Users with explicit permissions HR_REQUESTS_VIEW_ALL or HR_REQUESTS_APPROVE
      - All SUPER* roles
    """
    ids: set[int] = set()
    try:
        keys = {HR_REQUESTS_VIEW_ALL, HR_REQUESTS_APPROVE}
        rows = (UserPermission.query
                .filter(UserPermission.key.in_(list(keys)))
                .filter(UserPermission.is_allowed == True)  # noqa: E712
                .all())
        ids.update([r.user_id for r in rows if r.user_id])
    except Exception:
        pass

    try:
        supers = User.query.filter(func.upper(User.role).like('SUPER%')).all()
        ids.update([u.id for u in supers if u and u.id])
    except Exception:
        pass

    return sorted(ids)


def _portal_notify(user_ids: list[int], message: str, ntype: str = 'HR_ALERT'):
    """Create portal notifications (source='portal') for multiple users."""
    if not user_ids:
        return
    now = datetime.utcnow()
    for uid in sorted(set([int(x) for x in user_ids if x])):
        try:
            db.session.add(Notification(
                user_id=uid,
                type=ntype,
                message=message,
                created_at=now,
                is_read=False,
                source='portal',
            ))
        except Exception:
            continue


def _check_pending_leave_requests(send_notifications: bool = True) -> dict:
    """Find leave requests pending longer than configured threshold.

    If send_notifications=True, sends notifications to HR + request approver, but at most once every 24h per request.
    """
    days_thr = _setting_get_int('HR_ALERT_PENDING_DAYS', 2)
    now = datetime.utcnow()
    cutoff = now - timedelta(days=max(1, days_thr))

    pending = []
    try:
        q = (HRLeaveRequest.query
             .filter(HRLeaveRequest.status == 'SUBMITTED')
             .filter(HRLeaveRequest.submitted_at.isnot(None))
             .filter(HRLeaveRequest.submitted_at <= cutoff)
             .order_by(HRLeaveRequest.submitted_at.asc()))
        pending = q.all()
    except Exception:
        pending = []

    if not send_notifications or not pending:
        return {'threshold_days': days_thr, 'pending': pending, 'notified': 0}

    hr_ids = _hr_recipients_user_ids()
    notified = 0

    for r in pending:
        try:
            # Skip if reminded in last 24h
            if r.reminder_sent_at:
                try:
                    if (now - r.reminder_sent_at) < timedelta(hours=24):
                        continue
                except Exception:
                    pass

            # Build recipients
            recips = set(hr_ids)
            if r.approver_user_id:
                recips.add(int(r.approver_user_id))

            emp_name = None
            try:
                emp_name = (r.user.full_name or r.user.name or r.user.email)
            except Exception:
                emp_name = None

            submitted = None
            try:
                submitted = r.submitted_at.strftime('%Y-%m-%d') if r.submitted_at else None
            except Exception:
                pass

            msg = f"تنبيه: طلب إجازة رقم #{r.id} للموظف {emp_name or ''} لم يتم اتخاذ إجراء عليه منذ أكثر من {days_thr} يوم (تقديم: {submitted or '-'})".strip()
            _portal_notify(list(recips), msg, ntype='HR_ALERT')

            # Update reminder metadata
            r.reminder_sent_at = now
            r.reminder_count = int(r.reminder_count or 0) + 1
            notified += 1
        except Exception:
            continue

    try:
        db.session.commit()
    except Exception:
        db.session.rollback()

    return {'threshold_days': days_thr, 'pending': pending, 'notified': notified}


@portal_bp.route('/hr/leaves/balances', methods=['GET', 'POST'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_leave_balances():
    """Manage/display leave balances (entitlements vs used/remaining)."""
    # View allowed for reports viewers, but edits require manage (or view-all).
    can_manage = False
    try:
        can_manage = current_user.has_perm(HR_MASTERDATA_MANAGE) or current_user.has_perm(HR_REQUESTS_VIEW_ALL)
    except Exception:
        can_manage = False
    # Filters
    year_raw = (request.values.get('year') or '').strip()
    year = None
    try:
        year = int(year_raw) if year_raw else datetime.utcnow().year
    except Exception:
        year = datetime.utcnow().year

    user_id_raw = (request.values.get('user_id') or '').strip()
    selected_user = None
    try:
        if user_id_raw and user_id_raw.isdigit():
            selected_user = User.query.get(int(user_id_raw))
    except Exception:
        selected_user = None

    # NOTE: User.full_name is a Python @property (not a SQL column). Also, SQLite doesn't support NULLS LAST.
    # Use a safe SQL ordering for all DBs.
    from sqlalchemy import func
    users = User.query.order_by(func.coalesce(User.name, User.email).asc(), User.email.asc()).all()
    leave_types = HRLeaveType.query.filter(HRLeaveType.is_active == True).order_by(HRLeaveType.code.asc()).all()  # noqa: E712

    # POST: update entitlements for selected user
    if request.method == 'POST':
        if not can_manage:
            abort(403)
        if not selected_user:
            flash('اختر موظفاً أولاً.', 'danger')
            return redirect(url_for('portal.hr_leave_balances', year=year))

        updated = 0
        for lt in leave_types:
            key = f'total_{lt.id}'
            raw = (request.form.get(key) or '').strip()
            if raw == '':
                continue
            try:
                total = int(raw)
            except Exception:
                continue
            row = HRLeaveBalance.query.filter_by(user_id=selected_user.id, leave_type_id=lt.id, year=year).first()
            if not row:
                row = HRLeaveBalance(user_id=selected_user.id, leave_type_id=lt.id, year=year, total_days=total, created_at=datetime.utcnow())
                db.session.add(row)
            else:
                row.total_days = total
            updated += 1

        if updated:
            _portal_audit('HR_LEAVE_BALANCE_UPDATE', f'Update leave balances for user_id={selected_user.id} year={year}', target_type='USER', target_id=selected_user.id)
        db.session.commit()
        flash('تم حفظ أرصدة الإجازات.', 'success')
        return redirect(url_for('portal.hr_leave_balances', user_id=selected_user.id, year=year))

    # Build rows for view
    rows = []
    if selected_user:
        for lt in leave_types:
            total = _leave_entitlement_days(selected_user.id, lt, year)
            used = _leave_used_days(selected_user.id, lt.id, year)
            rem = total - used
            rows.append({'lt': lt, 'total': total, 'used': used, 'remaining': rem})

    return render_template(
        'portal/hr/leave_balances.html',
        users=users,
        selected_user=selected_user,
        year=year,
        leave_types=leave_types,
        rows=rows,
        can_manage=can_manage,
    )


@portal_bp.route('/hr/leaves/monthly', methods=['GET', 'POST'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_monthly_leave_report():
    # NOTE (TEMP): Open this report for testing without permission gates.
    # POST actions remain restricted via can_manage checks below.
    """Monthly report: permissions allowance vs used + leave consumption + balances by leave type.

    This report reflects the 'day-by-day' deduction policy: future days are not deducted until they occur.
    """
    can_manage = False
    try:
        can_manage = current_user.has_perm(HR_MASTERDATA_MANAGE) or current_user.has_perm(HR_REQUESTS_VIEW_ALL)
    except Exception:
        can_manage = False
    today = date.today()

    def _to_int(v, default=None):
        try:
            s = (v or '').strip()
            if s == '':
                return default
            return int(s)
        except Exception:
            return default

    year = _to_int(request.values.get('year'), today.year)
    month = _to_int(request.values.get('month'), today.month)
    month = 1 if not month or month < 1 else (12 if month > 12 else month)

    user_id = _to_int(request.values.get('user_id'), None)
    selected_user = User.query.get(user_id) if user_id else None

    from sqlalchemy import func
    users = User.query.order_by(func.coalesce(User.name, User.email).asc(), User.email.asc()).all()
    leave_types = HRLeaveType.query.filter(HRLeaveType.is_active == True).order_by(HRLeaveType.code.asc()).all()  # noqa: E712

    # Resolve as_of for the selected month
    month_start = date(year, month, 1)
    month_end_str = _end_of_month_str(year, month)
    month_end = _parse_yyyy_mm_dd(month_end_str) or month_start

    if (year < today.year) or (year == today.year and month < today.month):
        as_of = month_end
    elif year == today.year and month == today.month:
        as_of = today
    else:
        as_of = month_start - timedelta(days=1)

    # Settings
    default_allowed_hours = _default_allowed_permission_hours()
    workday_hours = _workday_hours()

    # POST actions: update allowance or global settings
    if request.method == 'POST':
        if not can_manage:
            abort(403)
        action = (request.form.get('action') or '').strip()

        if action == 'save_allowance':
            if not selected_user:
                flash('اختر موظفاً أولاً.', 'warning')
                return redirect(url_for('portal.hr_monthly_leave_report', year=year, month=month))
            allowed = _to_int(request.form.get('allowed_hours'), default_allowed_hours)
            if allowed is None or allowed < 0:
                allowed = 0
            row = HRMonthlyPermissionAllowance.query.filter_by(user_id=selected_user.id, year=year, month=month).first()
            if not row:
                row = HRMonthlyPermissionAllowance(user_id=selected_user.id, year=year, month=month, allowed_hours=allowed)
                db.session.add(row)
            else:
                row.allowed_hours = allowed
            db.session.commit()
            flash('تم حفظ ساعات المغادرة المسموحة للشهر.', 'success')
            return redirect(url_for('portal.hr_monthly_leave_report', user_id=selected_user.id, year=year, month=month))

        if action == 'save_defaults':
            # Save global default allowed hours and excess-deduct leave type
            new_default = _to_int(request.form.get('default_allowed_hours'), default_allowed_hours)
            if new_default is None or new_default < 0:
                new_default = 0
            _setting_set('HR_ALLOWED_PERMISSION_HOURS_MONTH', str(new_default))

            lt_id = _to_int(request.form.get('excess_leave_type_id'), None)
            if lt_id:
                _setting_set('HR_EXCESS_PERM_DEDUCT_LEAVE_TYPE_ID', str(lt_id))
            db.session.commit()
            flash('تم حفظ الإعدادات الافتراضية.', 'success')
            return redirect(url_for('portal.hr_monthly_leave_report', user_id=(selected_user.id if selected_user else None), year=year, month=month))

    # Calculations
    start_str = month_start.strftime('%Y-%m-%d')
    end_str = as_of.strftime('%Y-%m-%d')

    report = None
    balances = []
    leave_days_in_month = []

    if selected_user:
        allowed_hours = _get_monthly_allowed_hours(selected_user.id, year, month)
        used_perm_hours = _permission_hours_in_range(selected_user.id, start_str, end_str) if end_str >= start_str else 0
        excess_hours = max(0, int(used_perm_hours) - int(allowed_hours))
        excess_days = float(excess_hours) / float(workday_hours) if workday_hours else 0.0

        excess_lt_id = _permission_excess_leave_type_id()
        excess_lt = HRLeaveType.query.get(excess_lt_id) if excess_lt_id else None

        # Leave days inside the selected month (day-by-day; bounded by as_of)
        for lt in leave_types:
            days = 0.0
            q = (HRLeaveRequest.query
                 .filter(HRLeaveRequest.user_id == selected_user.id)
                 .filter(HRLeaveRequest.leave_type_id == lt.id)
                 .filter(HRLeaveRequest.status.in_(['APPROVED', 'CANCELLED']))
                 .order_by(HRLeaveRequest.id.asc()))

            for r in q.all():
                if r.status == 'CANCELLED' and (r.cancelled_from_status or '').upper() != 'APPROVED':
                    continue

                s = _parse_yyyy_mm_dd(r.start_date)
                e = _parse_yyyy_mm_dd(r.end_date)
                if not s or not e:
                    continue

                effective_end = min(e, as_of)
                if r.status == 'CANCELLED' and r.cancel_effective_date:
                    ce = _parse_yyyy_mm_dd(r.cancel_effective_date)
                    if ce:
                        effective_end = min(effective_end, ce)

                os = max(s, month_start)
                oe = min(effective_end, month_end)
                if oe < os:
                    continue

                days += float((oe - os).days + 1)

            leave_days_in_month.append({'leave_type': lt, 'days': days})

        # Year-to-date balances as of end_str (as_of)
        for lt in leave_types:
            total = float(_leave_entitlement_days(selected_user.id, lt, year))
            used = float(_leave_used_days_as_of(selected_user.id, lt.id, year, as_of))
            rem = total - used
            balances.append({'leave_type': lt, 'total': total, 'used': used, 'remaining': rem})

        report = {
            'allowed_hours': int(allowed_hours),
            'used_permission_hours': int(used_perm_hours),
            'excess_hours': int(excess_hours),
            'excess_days': float(excess_days),
            'excess_deduct_leave_type': excess_lt,
            'as_of': as_of,
            'start_str': start_str,
            'end_str': end_str,
        }

    return render_template(
        'portal/hr/monthly_leave_report.html',
        users=users,
        selected_user=selected_user,
        year=year,
        month=month,
        leave_types=leave_types,
        report=report,
        balances=balances,
        leave_days_in_month=leave_days_in_month,
        default_allowed_hours=default_allowed_hours,
        excess_leave_type_id=_permission_excess_leave_type_id(),
        workday_hours=workday_hours,
        can_manage=can_manage,
    )




@portal_bp.route('/hr/monthly-leave-report/allowance', methods=['POST'])
@login_required
@_perm_any(HR_MASTERDATA_MANAGE, HR_REQUESTS_VIEW_ALL)
def hr_monthly_leave_report_set_allowance():
    """Upsert allowed permission hours for a specific employee/month."""
    try:
        user_id = int(request.form.get('user_id') or 0)
        year = int(request.form.get('year') or date.today().year)
        month = int(request.form.get('month') or date.today().month)
        allowed_hours = int(request.form.get('allowed_hours') or 0)
    except Exception:
        flash('بيانات غير صحيحة.', 'danger')
        return redirect(url_for('portal.hr_monthly_leave_report'))

    if user_id <= 0 or month < 1 or month > 12:
        flash('بيانات غير صحيحة.', 'danger')
        return redirect(url_for('portal.hr_monthly_leave_report'))

    row = HRMonthlyPermissionAllowance.query.filter_by(user_id=user_id, year=year, month=month).first()
    if not row:
        row = HRMonthlyPermissionAllowance(user_id=user_id, year=year, month=month, allowed_hours=allowed_hours,
                                           created_at=datetime.utcnow(), created_by_id=current_user.id)
        db.session.add(row)
        _portal_audit('HR_MONTHLY_PERMISSION_ALLOWANCE_CREATE', f'تحديد سماح مغادرات {allowed_hours} ساعة للموظف #{user_id} لشهر {year}-{month:02d}',
                      target_type='USER', target_id=user_id)
    else:
        row.allowed_hours = allowed_hours
        row.updated_at = datetime.utcnow() if hasattr(row, 'updated_at') else None
        _portal_audit('HR_MONTHLY_PERMISSION_ALLOWANCE_UPDATE', f'تحديث سماح مغادرات {allowed_hours} ساعة للموظف #{user_id} لشهر {year}-{month:02d}',
                      target_type='USER', target_id=user_id)

    db.session.commit()
    flash('تم حفظ السماح الشهري.', 'success')
    return redirect(url_for('portal.hr_monthly_leave_report', user_id=user_id, year=year, month=month))


@portal_bp.route('/hr/monthly-leave-report/settings', methods=['POST'])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def hr_monthly_leave_report_settings_update():
    """Update global settings for the monthly report (defaults)."""
    try:
        default_allowed_hours = int(request.form.get('default_allowed_hours') or 0)
        workday_hours = int(request.form.get('workday_hours') or 8)
        excess_lt_id = int(request.form.get('excess_leave_type_id') or 0)
        user_id = int(request.form.get('user_id') or 0)
        year = int(request.form.get('year') or date.today().year)
        month = int(request.form.get('month') or date.today().month)
    except Exception:
        flash('بيانات غير صحيحة.', 'danger')
        return redirect(url_for('portal.hr_monthly_leave_report'))

    if default_allowed_hours < 0:
        default_allowed_hours = 0
    if workday_hours <= 0:
        workday_hours = 8

    _setting_set('HR_ALLOWED_PERMISSION_HOURS_MONTH', str(default_allowed_hours))
    _setting_set('HR_WORKDAY_HOURS', str(workday_hours))
    _setting_set('HR_EXCESS_PERM_DEDUCT_LEAVE_TYPE_ID', str(excess_lt_id) if excess_lt_id else '')
    _portal_audit('HR_MONTHLY_REPORT_SETTINGS', f'تحديث إعدادات تقرير الإجازات الشهري: allowed_hours={default_allowed_hours}, workday_hours={workday_hours}, excess_leave_type_id={excess_lt_id or "-"}',
                  target_type='SYSTEM', target_id=None)
    db.session.commit()
    flash('تم حفظ الإعدادات.', 'success')
    return redirect(url_for('portal.hr_monthly_leave_report', user_id=user_id or None, year=year, month=month))

@portal_bp.route('/hr/alerts', methods=['GET'])
@login_required
@_perm_any(HR_ATTENDANCE_READ, HR_REQUESTS_VIEW_ALL)
def hr_alerts():
    '''Show smart alerts (lateness, low leave balance, pending approvals).'''
    # Run pending check (and notify) to satisfy requirement "after 2 days".
    pending_info = _check_pending_leave_requests(send_notifications=True)

    # Thresholds
    late_thr = _setting_get_int('HR_ALERT_LATE_MINUTES_MONTH', 120)
    leave_thr = _setting_get_int('HR_ALERT_LEAVE_REMAIN_DAYS', 2)

    # Optional filters (UI scope filtering)
    def _to_int(v):
        try:
            s = (v or '').strip()
            return int(s) if s.isdigit() else None
        except Exception:
            return None

    directorate_id = _to_int(request.args.get('directorate_id') or request.args.get('directorate'))
    unit_id = _to_int(request.args.get('unit_id') or request.args.get('unit'))
    department_id = _to_int(request.args.get('department_id') or request.args.get('department'))
    section_id = _to_int(request.args.get('section_id') or request.args.get('section'))
    division_id = _to_int(request.args.get('division_id') or request.args.get('division'))

    any_filter = any([directorate_id, unit_id, department_id, section_id, division_id])

    # Org lists for UI
    directorates = Directorate.query.filter(Directorate.is_active == True).order_by(Directorate.name_ar.asc()).all()  # noqa: E712
    units = Unit.query.filter(Unit.is_active == True).order_by(Unit.name_ar.asc()).all()  # noqa: E712
    departments = Department.query.filter(Department.is_active == True).order_by(Department.name_ar.asc()).all()  # noqa: E712
    sections = Section.query.filter(Section.is_active == True).order_by(Section.name_ar.asc()).all()  # noqa: E712
    divisions = Division.query.filter(Division.is_active == True).order_by(Division.name_ar.asc()).all()  # noqa: E712

    dept_map = {d.id: d for d in departments}

    def _effective_directorate_id(u):
        # user.directorate_id is explicit; otherwise derived from department
        try:
            if getattr(u, 'directorate_id', None):
                return int(u.directorate_id)
        except Exception:
            pass
        try:
            if getattr(u, 'department_id', None):
                dept = dept_map.get(int(u.department_id))
                if dept and dept.directorate_id:
                    return int(dept.directorate_id)
        except Exception:
            pass
        return None

    def _effective_unit_id(u):
        # unit derived from department
        try:
            if getattr(u, 'department_id', None):
                dept = dept_map.get(int(u.department_id))
                if dept and dept.unit_id:
                    return int(dept.unit_id)
        except Exception:
            pass
        return None

    def _matches_user(u):
        if not u:
            return False

        if directorate_id and _effective_directorate_id(u) != directorate_id:
            return False
        if unit_id and _effective_unit_id(u) != unit_id:
            return False

        if department_id:
            try:
                if int(getattr(u, 'department_id', 0) or 0) != department_id:
                    return False
            except Exception:
                return False

        # section/division are optional fields (may not exist in older DBs)
        if section_id:
            try:
                if int(getattr(u, 'section_id', 0) or 0) != section_id:
                    return False
            except Exception:
                return False

        if division_id:
            try:
                if int(getattr(u, 'division_id', 0) or 0) != division_id:
                    return False
            except Exception:
                return False

        return True

    # Lateness current month
    now = datetime.utcnow()
    month_prefix = now.strftime('%Y-%m')
    late_rows = []
    try:
        q = (db.session.query(
                AttendanceDailySummary.user_id,
                func.coalesce(func.sum(AttendanceDailySummary.late_minutes), 0).label('late_min')
            )
            .filter(AttendanceDailySummary.day.like(month_prefix + '%'))
            .group_by(AttendanceDailySummary.user_id))

        for uid, late_min in q.all():
            if int(late_min or 0) >= late_thr:
                u = User.query.get(uid)
                if (not any_filter) or _matches_user(u):
                    late_rows.append({'user': u, 'late_minutes': int(late_min or 0)})

        late_rows.sort(key=lambda x: x['late_minutes'], reverse=True)
    except Exception:
        late_rows = []

    # Low leave balance (current year)
    year = now.year
    low_leave = []
    try:
        active_types = HRLeaveType.query.filter(HRLeaveType.is_active == True).order_by(HRLeaveType.code.asc()).all()  # noqa: E712

        # Scope users first when filtering is enabled
        users_q = User.query.order_by(User.id.asc()).all()
        users = [u for u in users_q if _matches_user(u)] if any_filter else users_q

        for u in users:
            for lt in active_types:
                total = _leave_entitlement_days(u.id, lt, year)
                if total <= 0:
                    continue
                used = _leave_used_days(u.id, lt.id, year)
                rem = total - used
                if rem <= leave_thr:
                    low_leave.append({'user': u, 'leave_type': lt, 'total': total, 'used': used, 'remaining': rem})
    except Exception:
        low_leave = []

    # Filter pending list as well (UI scope only)
    try:
        if any_filter and pending_info and isinstance(pending_info, dict):
            pend = pending_info.get('pending') or []
            filtered = []
            for r in pend:
                u = getattr(r, 'user', None)
                if not u:
                    try:
                        u = User.query.get(getattr(r, 'user_id', None))
                    except Exception:
                        u = None
                if _matches_user(u):
                    filtered.append(r)
            pending_info = dict(pending_info)
            pending_info['pending'] = filtered
    except Exception:
        pass

    return render_template(
        'portal/hr/alerts.html',
        pending_info=pending_info,
        late_threshold=late_thr,
        leave_threshold=leave_thr,
        late_rows=late_rows,
        low_leave=low_leave,
        year=year,
        # filter lists + selections
        directorates=directorates,
        units=units,
        departments=departments,
        sections=sections,
        divisions=divisions,
        selected_directorate_id=directorate_id,
        selected_unit_id=unit_id,
        selected_department_id=department_id,
        selected_section_id=section_id,
        selected_division_id=division_id,
    )


@portal_bp.route('/hr/leaves/report', methods=['GET'])

@login_required
@_perm(HR_REPORTS_VIEW)
def hr_leaves_report():
    """Leave requests report (with pending check trigger)."""
    pending_info = _check_pending_leave_requests(send_notifications=True)

    # Filters
    from_day = (request.args.get('from') or '').strip()
    to_day = (request.args.get('to') or '').strip()

    q = HRLeaveRequest.query
    if from_day:
        q = q.filter(HRLeaveRequest.start_date >= from_day)
    if to_day:
        q = q.filter(HRLeaveRequest.start_date <= to_day)

    q = q.order_by(HRLeaveRequest.id.desc())
    rows = q.limit(500).all()

    # Optional export
    if (request.args.get('export') or '').lower() == 'csv':
        out = io.StringIO()
        w = csv.writer(out)
        w.writerow(['id', 'employee', 'leave_type', 'start_date', 'end_date', 'days', 'status', 'submitted_at', 'approver'])
        for r in rows:
            emp = ''
            try:
                emp = r.user.full_name or r.user.name or r.user.email
            except Exception:
                emp = ''
            lt = ''
            try:
                lt = r.leave_type.code
            except Exception:
                lt = ''
            approver = ''
            try:
                a = User.query.get(r.approver_user_id) if r.approver_user_id else None
                approver = (a.full_name or a.name or a.email) if a else ''
            except Exception:
                approver = ''
            w.writerow([r.id, emp, lt, r.start_date, r.end_date, r.days, r.status, r.submitted_at, approver])
        data = out.getvalue().encode('utf-8-sig')
        return send_file(BytesIO(data), mimetype='text/csv', as_attachment=True, download_name='hr_leaves_report.csv')

    return render_template('portal/hr/leaves_report.html', rows=rows, pending_info=pending_info, from_day=from_day, to_day=to_day)

# Portal Admin: Integrations (Timeclock file on server)
# -------------------------


def _timeclock_resolve_source_file(source_path: str) -> str | None:
    """Resolve the configured timeclock source.

    - If a file path is provided -> return it (if exists)
    - If a directory is provided -> pick the latest daily file, preferring names like YYYYMMDD.CSV
    """
    src = (source_path or '').strip()
    if not src:
        return None

    p = Path(src)
    if p.exists() and p.is_file():
        return str(p)

    if p.exists() and p.is_dir():
        # Prefer date-stamped files: 20260215.CSV
        pat = re.compile(r"^(\d{8})\.(csv)$", re.IGNORECASE)
        dated = []
        other = []
        try:
            for child in p.iterdir():
                if not child.is_file():
                    continue
                m = pat.match(child.name)
                if m:
                    dated.append(child)
                elif child.suffix.lower() == '.csv':
                    other.append(child)
        except Exception:
            return None

        if dated:
            # max by filename (YYYYMMDD)
            dated.sort(key=lambda x: x.name)
            return str(dated[-1])

        if other:
            other.sort(key=lambda x: x.stat().st_mtime)
            return str(other[-1])

        return None

    # Not found (could be a UNC path or a file that does not exist yet)
    # Try treating it as a directory anyway (common when admin pastes UNC without pre-check)
    try:
        if src.endswith('\\') or src.endswith('/'):
            dp = Path(src)
            if dp.exists() and dp.is_dir():
                return _timeclock_resolve_source_file(str(dp))
    except Exception:
        pass

    return None


def _timeclock_get_match_by() -> str:
    v = (_setting_get('TIMECLK_MATCH_BY') or '').strip().upper()
    # Allowed values:
    #  - TIMECLK_CODE: match EmployeeFile.timeclock_code
    #  - EMPLOYEE_NO: match EmployeeFile.employee_no
    #  - NATIONAL_ID: match EmployeeFile.national_id
    #  - AUTO: try all (employee_no -> national_id -> timeclock_code)
    return v if v in {'TIMECLK_CODE', 'EMPLOYEE_NO', 'NATIONAL_ID', 'AUTO'} else 'TIMECLK_CODE'


def _timeclock_build_code_to_user(match_by: str) -> dict:
    match_by = (match_by or '').strip().upper()

    def _add_aliases(m: dict):
        # add normalized (no-leading-zeros) alias when safe
        for k, uid in list(m.items()):
            if isinstance(k, str) and k.isdigit():
                alt = k.lstrip('0') or '0'
                if alt not in m:
                    m[alt] = uid

    if match_by == 'EMPLOYEE_NO':
        m = {
            (p.employee_no or '').strip(): p.user_id
            for p in EmployeeFile.query.filter(EmployeeFile.employee_no.isnot(None)).all()
            if (p.employee_no or '').strip()
        }
        _add_aliases(m)
        return m

    if match_by == 'NATIONAL_ID':
        m = {
            (p.national_id or '').strip(): p.user_id
            for p in EmployeeFile.query.filter(EmployeeFile.national_id.isnot(None)).all()
            if (p.national_id or '').strip()
        }
        _add_aliases(m)
        return m

    if match_by == 'AUTO':
        # Priority: employee_no -> national_id -> timeclock_code
        m: dict = {}
        for p in EmployeeFile.query.filter(
            (EmployeeFile.employee_no.isnot(None)) | (EmployeeFile.national_id.isnot(None)) | (EmployeeFile.timeclock_code.isnot(None))
        ).all():
            eno = (p.employee_no or '').strip()
            nid = (p.national_id or '').strip()
            tcc = (p.timeclock_code or '').strip()
            if eno:
                m.setdefault(eno, p.user_id)
            if nid:
                m.setdefault(nid, p.user_id)
            if tcc:
                m.setdefault(tcc, p.user_id)
        _add_aliases(m)
        return m

    # default: TIMECLK_CODE
    m = {
        (p.timeclock_code or '').strip(): p.user_id
        for p in EmployeeFile.query.filter(EmployeeFile.timeclock_code.isnot(None)).all()
        if (p.timeclock_code or '').strip()
    }
    _add_aliases(m)
    return m


def _timeclock_read_incremental(file_path: str, last_size: int | None, append_only: bool):
    p = Path(file_path)
    if not p.exists() or not p.is_file():
        raise FileNotFoundError(str(p))

    size = p.stat().st_size
    if append_only and last_size is not None and last_size >= 0 and size >= last_size:
        # read appended bytes
        with p.open('rb') as f:
            f.seek(last_size)
            data = f.read()
        return data, size

    # fallback full
    with p.open('rb') as f:
        data = f.read()
    return data, size


def _timeclock_sync(file_path: str, imported_by_id: int, append_only: bool = True):
    # returns (inserted, skipped, errors_count)
    resolved = _timeclock_resolve_source_file(file_path)
    if not resolved:
        raise FileNotFoundError((file_path or '').strip() or 'TIMECLK_SOURCE_FILE')

    last_file = (_setting_get('TIMECLK_LAST_FILE') or '').strip()
    last_size = _setting_get('TIMECLK_LAST_SIZE')
    if last_file and (last_file != resolved):
        last_size_i = None
    else:
        last_size_i = int(last_size) if (last_size and str(last_size).isdigit()) else None

    raw_bytes, new_size = _timeclock_read_incremental(resolved, last_size_i, append_only)

    try:
        text = raw_bytes.decode('utf-8', errors='ignore')
    except Exception:
        text = raw_bytes.decode(errors='ignore')

    lines = [ln for ln in text.splitlines() if (ln or '').strip()]
    if not lines:
        _setting_set('TIMECLK_LAST_SIZE', str(new_size))
        _setting_set('TIMECLK_LAST_SYNC_AT', datetime.utcnow().isoformat(timespec='seconds'))
        db.session.commit()
        return 0, 0, 0

    batch = AttendanceImportBatch(
        filename=f"AUTO:{Path(resolved).name}",
        imported_by_id=imported_by_id,
        imported_at=datetime.utcnow(),
        total_lines=len(lines),
        inserted=0,
        skipped=0,
    )
    db.session.add(batch)
    db.session.flush()

    match_by = _timeclock_get_match_by()
    code_to_user = _timeclock_build_code_to_user(match_by)

    errors = []
    for ln in lines:
        parsed = _parse_timeclock_line(ln)
        if not parsed:
            batch.skipped += 1
            errors.append(f"Bad line: {ln!r}")
            continue
        user_id = code_to_user.get(parsed['emp_code'])
        if not user_id and (parsed.get('emp_code') or '').isdigit():
            user_id = code_to_user.get((parsed['emp_code'].lstrip('0') or '0'))
        if not user_id:
            batch.skipped += 1
            errors.append(f"Unknown emp_code={parsed['emp_code']} line={ln!r}")
            continue

        # DB-level unique constraint protects duplicates
        ev = AttendanceEvent(
            batch_id=batch.id,
            user_id=user_id,
            event_dt=parsed['event_dt'],
            event_type=parsed['event_type'],
            device_id=parsed['device_id'],
            raw_line=parsed['raw'],
            created_at=datetime.utcnow(),
        )
        db.session.add(ev)
        try:
            db.session.flush()
            batch.inserted += 1
        except Exception:
            db.session.rollback()
            # recreate batch? keep simple: mark as skipped duplicates
            # NOTE: rollback clears batch changes, so we need a safer approach
            # To keep it robust, we avoid flush in loop and rely on commit error handling.
            raise

    if errors:
        batch.errors = "\n".join(errors[:200])

    _setting_set('TIMECLK_LAST_FILE', str(resolved))
    _setting_set('TIMECLK_LAST_SIZE', str(new_size))
    _setting_set('TIMECLK_LAST_SYNC_AT', datetime.utcnow().isoformat(timespec='seconds'))

    db.session.commit()
    return batch.inserted, batch.skipped, len(errors)


@portal_bp.route('/admin/integrations', methods=['GET', 'POST'])
@login_required
@_perm(PORTAL_INTEGRATIONS_MANAGE)
def portal_admin_integrations():
    if request.method == 'POST':
        file_path = (request.form.get('timeclock_file_path') or '').strip()
        append_only = (request.form.get('append_only') or '1') == '1'
        auto_enabled = (request.form.get('auto_enabled') or '0') == '1'
        auto_interval = (request.form.get('auto_interval') or '').strip()
        imported_by_id = (request.form.get('imported_by_id') or '').strip()
        match_by = (request.form.get('match_by') or '').strip().upper()

        # sanitize values
        try:
            interval_sec = int(auto_interval) if auto_interval else 60
        except Exception:
            interval_sec = 60
        interval_sec = max(10, min(interval_sec, 3600))

        imp_id_int = None
        try:
            imp_id_int = int(imported_by_id) if imported_by_id else None
        except Exception:
            imp_id_int = None

        _setting_set('TIMECLK_SOURCE_FILE', file_path)
        _setting_set('TIMECLK_APPEND_ONLY', '1' if append_only else '0')
        _setting_set('TIMECLK_AUTO_SYNC_ENABLED', '1' if auto_enabled else '0')
        _setting_set('TIMECLK_AUTO_SYNC_INTERVAL', str(interval_sec))
        if imp_id_int is not None:
            _setting_set('TIMECLK_IMPORTED_BY_USER_ID', str(imp_id_int))

        if match_by in {'TIMECLK_CODE', 'EMPLOYEE_NO', 'NATIONAL_ID', 'AUTO'}:
            _setting_set('TIMECLK_MATCH_BY', match_by)

        _portal_audit(
            'PORTAL_INTEGRATION_SAVE',
            f"TIMECLK_SOURCE_FILE={file_path} AUTO={1 if auto_enabled else 0} INTERVAL={interval_sec}",
            target_type='SETTING',
            target_id=0,
        )
        db.session.commit()
        flash('تم حفظ إعدادات التكامل.', 'success')
        return redirect(url_for('portal.portal_admin_integrations'))

    file_path = _setting_get('TIMECLK_SOURCE_FILE') or ''
    append_only = (_setting_get('TIMECLK_APPEND_ONLY') or '1') == '1'
    last_size = _setting_get('TIMECLK_LAST_SIZE') or ''
    last_sync = _setting_get('TIMECLK_LAST_SYNC_AT') or ''

    auto_enabled = (_setting_get('TIMECLK_AUTO_SYNC_ENABLED') or '0') == '1'
    auto_interval = _setting_get('TIMECLK_AUTO_SYNC_INTERVAL') or '60'
    imported_by_id = _setting_get('TIMECLK_IMPORTED_BY_USER_ID') or ''
    match_by = _timeclock_get_match_by()
    last_file = _setting_get('TIMECLK_LAST_FILE') or ''

    return render_template('portal/admin/integrations.html',
                           timeclock_file_path=file_path,
                           append_only=append_only,
                           last_size=last_size,
                           last_sync=last_sync,
                           auto_enabled=auto_enabled,
                           auto_interval=auto_interval,
                           imported_by_id=imported_by_id,
                           match_by=match_by,
                           last_file=last_file)


@portal_bp.route('/hr/attendance/sync-now', methods=['POST'])
@login_required
@_perm(HR_ATT_CREATE)
def hr_attendance_sync_now():
    # Only if portal integrations configured
    if not current_user.has_perm(PORTAL_INTEGRATIONS_MANAGE) and not current_user.has_role('ADMIN'):
        abort(403)

    file_path = (_setting_get('TIMECLK_SOURCE_FILE') or '').strip()
    if not file_path:
        flash('يرجى تحديد مسار ملف ساعة الدوام من صفحة التكاملات.', 'danger')
        return redirect(url_for('portal.portal_admin_integrations'))

    # If a folder is configured, we will pick the latest daily file automatically.
    resolved = _timeclock_resolve_source_file(file_path)
    if not resolved:
        flash('تعذر العثور على ملف ساعة الدوام داخل المسار المحدد. تأكد من مشاركة المجلد وصلاحيات القراءة.', 'danger')
        return redirect(url_for('portal.portal_admin_integrations'))

    append_only = (_setting_get('TIMECLK_APPEND_ONLY') or '1') == '1'

    try:
        # safer insertion: bulk add without flush per-row
        inserted, skipped, errs = _timeclock_sync_simple(file_path, current_user.id, append_only)
        flash(f'تمت المزامنة: {inserted} سجل، {skipped} تم تجاهله.', 'success')
    except FileNotFoundError:
        flash('الملف غير موجود على المسار المحدد.', 'danger')
    except Exception as e:
        flash('تعذر تنفيذ المزامنة.', 'danger')

    return redirect(url_for('portal.hr_attendance_batches'))


def _timeclock_sync_simple(file_path: str, imported_by_id: int, append_only: bool = True):
    # Same as _timeclock_sync but avoids flush/rollback inside loop.
    # Support directory input (daily files like YYYYMMDD.CSV)
    resolved = _timeclock_resolve_source_file(file_path)
    if not resolved:
        raise FileNotFoundError((file_path or '').strip() or 'TIMECLK_SOURCE_FILE')

    last_file = (_setting_get('TIMECLK_LAST_FILE') or '').strip()
    last_size = _setting_get('TIMECLK_LAST_SIZE')

    # If the device rotates files daily, reset incremental pointer when file changes
    if last_file and (last_file != resolved):
        last_size_i = None
    else:
        last_size_i = int(last_size) if (last_size and str(last_size).isdigit()) else None

    raw_bytes, new_size = _timeclock_read_incremental(resolved, last_size_i, append_only)

    try:
        text = raw_bytes.decode('utf-8', errors='ignore')
    except Exception:
        text = raw_bytes.decode(errors='ignore')

    lines = [ln for ln in text.splitlines() if (ln or '').strip()]

    batch = AttendanceImportBatch(
        filename=f"AUTO:{Path(resolved).name}",
        imported_by_id=imported_by_id,
        imported_at=datetime.utcnow(),
        total_lines=len(lines),
        inserted=0,
        skipped=0,
    )
    db.session.add(batch)
    db.session.flush()

    match_by = _timeclock_get_match_by()
    code_to_user = _timeclock_build_code_to_user(match_by)

    seen = set()
    errors = []

    for ln in lines:
        parsed = _parse_timeclock_line(ln)
        if not parsed:
            batch.skipped += 1
            errors.append(f"Bad line: {ln!r}")
            continue

        user_id = code_to_user.get(parsed['emp_code'])
        if not user_id and (parsed.get('emp_code') or '').isdigit():
            user_id = code_to_user.get((parsed['emp_code'].lstrip('0') or '0'))
        if not user_id:
            batch.skipped += 1
            errors.append(f"Unknown emp_code={parsed['emp_code']} line={ln!r}")
            continue

        key = (user_id, parsed['event_dt'], parsed['event_type'], parsed['device_id'])
        if key in seen:
            batch.skipped += 1
            continue
        seen.add(key)

        # avoid duplicates via query check (cheaper than rollback)
        exists_ev = AttendanceEvent.query.filter_by(
            user_id=user_id,
            event_dt=parsed['event_dt'],
            event_type=parsed['event_type'],
            device_id=parsed['device_id'],
        ).first()
        if exists_ev:
            batch.skipped += 1
            continue

        db.session.add(AttendanceEvent(
            batch_id=batch.id,
            user_id=user_id,
            event_dt=parsed['event_dt'],
            event_type=parsed['event_type'],
            device_id=parsed['device_id'],
            raw_line=parsed['raw'],
            created_at=datetime.utcnow(),
        ))
        batch.inserted += 1

    if errors:
        batch.errors = "\n".join(errors[:200])

    _setting_set('TIMECLK_LAST_FILE', str(resolved))
    _setting_set('TIMECLK_LAST_SIZE', str(new_size))
    _setting_set('TIMECLK_LAST_SYNC_AT', datetime.utcnow().isoformat(timespec='seconds'))

    _portal_audit(
        'TIMECLK_SYNC',
        f"TIMECLK sync inserted={batch.inserted} skipped={batch.skipped} errors={len(errors)}",
        target_type='ATTENDANCE_IMPORT',
        target_id=batch.id,
        user_id=imported_by_id,
    )

    db.session.commit()
    return batch.inserted, batch.skipped, len(errors)


# -------------------------
# HR: Daily Summary (Late/Early/Overtime)
# -------------------------


def _summary_compute_one(user_id: int, day_str: str):
    # Collect day events
    dt_from = datetime.fromisoformat(day_str + 'T00:00:00')
    dt_to = datetime.fromisoformat(day_str + 'T23:59:59')

    evs = (
        AttendanceEvent.query
        .filter(AttendanceEvent.user_id == user_id, AttendanceEvent.event_dt >= dt_from, AttendanceEvent.event_dt <= dt_to)
        .order_by(AttendanceEvent.event_dt.asc())
        .all()
    )

    ins = [e.event_dt for e in evs if e.event_type == 'I']
    outs = [e.event_dt for e in evs if e.event_type == 'O']

    first_in = ins[0] if ins else None
    last_out = outs[-1] if outs else None

    schedule = _effective_schedule_for_user(user_id, day_str)
    schedule_id = schedule.id if schedule else None

    break_minutes = int(getattr(schedule, 'break_minutes', 0) or 0) if schedule else 0
    grace_minutes = int(getattr(schedule, 'grace_minutes', 0) or 0) if schedule else 0

    status = 'OK'
    if not first_in or not last_out:
        status = 'INCOMPLETE'

    work_minutes = 0
    late_minutes = 0
    early_leave_minutes = 0
    overtime_minutes = 0

    if first_in and last_out and last_out >= first_in:
        work_minutes = int((last_out - first_in).total_seconds() // 60)
        if break_minutes:
            work_minutes = max(0, work_minutes - break_minutes)

    # Compute late/early/overtime based on schedule times
    if schedule and schedule.kind in ('FIXED', 'RAMADAN', 'SHIFT'):
        weekday = _weekday_of(day_str)
        st = schedule.start_time
        en = schedule.end_time
        if schedule.kind == 'SHIFT':
            day_cfg = WorkScheduleDay.query.filter_by(schedule_id=schedule.id, weekday=weekday).first()
            if day_cfg:
                st = day_cfg.start_time or st
                en = day_cfg.end_time or en
                break_minutes = int(day_cfg.break_minutes or break_minutes)
                grace_minutes = int(day_cfg.grace_minutes or grace_minutes)

        st_min = _parse_hhmm(st)
        en_min = _parse_hhmm(en)

        if first_in and st_min is not None:
            actual_in = first_in.hour * 60 + first_in.minute
            late_minutes = max(0, actual_in - st_min - grace_minutes)

        if last_out and en_min is not None:
            actual_out = last_out.hour * 60 + last_out.minute
            early_leave_minutes = max(0, en_min - actual_out - grace_minutes)

            thr = schedule.overtime_threshold_minutes
            thr = int(thr) if (thr is not None) else 0
            overtime_minutes = max(0, actual_out - en_min - thr)

    if schedule and schedule.kind in ('FLEX', 'REMOTE'):
        # Late/Early undefined; overtime is minutes above required_minutes (if set)
        req = schedule.required_minutes or 0
        if req and work_minutes > req:
            overtime_minutes = work_minutes - req

    return dict(
        user_id=user_id,
        day=day_str,
        schedule_id=schedule_id,
        first_in=first_in,
        last_out=last_out,
        work_minutes=work_minutes,
        break_minutes=break_minutes,
        late_minutes=late_minutes,
        early_leave_minutes=early_leave_minutes,
        overtime_minutes=overtime_minutes,
        status=status,
    )


def _upsert_summary(row: dict):
    existing = AttendanceDailySummary.query.filter_by(user_id=row['user_id'], day=row['day']).first()
    if not existing:
        existing = AttendanceDailySummary(user_id=row['user_id'], day=row['day'])
        db.session.add(existing)

    existing.schedule_id = row.get('schedule_id')
    existing.first_in = row.get('first_in')
    existing.last_out = row.get('last_out')
    existing.work_minutes = row.get('work_minutes', 0) or 0
    existing.break_minutes = row.get('break_minutes', 0) or 0
    existing.late_minutes = row.get('late_minutes', 0) or 0
    existing.early_leave_minutes = row.get('early_leave_minutes', 0) or 0
    existing.overtime_minutes = row.get('overtime_minutes', 0) or 0
    existing.status = row.get('status') or 'OK'
    existing.computed_at = datetime.utcnow()


@portal_bp.route('/hr/attendance/daily')
@login_required
@_perm(HR_ATT_READ)
def hr_attendance_daily():
    day_from = (request.args.get('day_from') or '').strip()
    day_to = (request.args.get('day_to') or '').strip()
    user_id = (request.args.get('user_id') or '').strip()

    qry = AttendanceDailySummary.query

    if day_from:
        qry = qry.filter(AttendanceDailySummary.day >= day_from)
    if day_to:
        qry = qry.filter(AttendanceDailySummary.day <= day_to)
    if user_id.isdigit():
        qry = qry.filter(AttendanceDailySummary.user_id == int(user_id))

    rows = qry.order_by(AttendanceDailySummary.day.desc()).limit(500).all()
    users = User.query.order_by(User.name.asc().nullslast(), User.email.asc()).all()

    return render_template('portal/hr/attendance_daily.html', rows=rows, users=users,
                           day_from=day_from, day_to=day_to, user_id=user_id)


@portal_bp.route('/hr/attendance/daily/recompute', methods=['POST'])
@login_required
@_perm(HR_ATT_CREATE)
def hr_attendance_daily_recompute():
    day_from = (request.form.get('day_from') or '').strip()
    day_to = (request.form.get('day_to') or '').strip()
    user_id = (request.form.get('user_id') or '').strip()

    if not day_from or not day_to:
        flash('حدد تاريخ من/إلى.', 'danger')
        return redirect(url_for('portal.hr_attendance_daily'))

    user_ids = []
    if user_id.isdigit():
        user_ids = [int(user_id)]
    else:
        # only users mapped to timeclock (have timeclock_code)
        user_ids = [p.user_id for p in EmployeeFile.query.filter(EmployeeFile.timeclock_code.isnot(None)).all()]

    days = []
    try:
        d1 = date.fromisoformat(day_from)
        d2 = date.fromisoformat(day_to)
    except Exception:
        flash('صيغة التاريخ غير صحيحة.', 'danger')
        return redirect(url_for('portal.hr_attendance_daily'))

    if d2 < d1:
        d1, d2 = d2, d1

    cur = d1
    while cur <= d2:
        days.append(cur.isoformat())
        cur = cur.fromordinal(cur.toordinal() + 1)

    count = 0
    for uid in user_ids:
        for d in days:
            row = _summary_compute_one(uid, d)
            _upsert_summary(row)
            count += 1

    _portal_audit('HR_ATT_DAILY_RECOMPUTE', f'recompute rows={count} range={day_from}..{day_to}', target_type='ATT_DAILY', target_id=0)

    db.session.commit()
    flash(f'تمت إعادة الحساب ({count}).', 'success')
    return redirect(url_for('portal.hr_attendance_daily', day_from=day_from, day_to=day_to, user_id=user_id))


@portal_bp.route('/hr/attendance/daily/export.xlsx')
@login_required
@_perm(HR_ATT_EXPORT)
def hr_attendance_daily_export_xlsx():
    day_from = (request.args.get('day_from') or '').strip()
    day_to = (request.args.get('day_to') or '').strip()
    user_id = (request.args.get('user_id') or '').strip()

    qry = AttendanceDailySummary.query
    if day_from:
        qry = qry.filter(AttendanceDailySummary.day >= day_from)
    if day_to:
        qry = qry.filter(AttendanceDailySummary.day <= day_to)
    if user_id.isdigit():
        qry = qry.filter(AttendanceDailySummary.user_id == int(user_id))

    rows = qry.order_by(AttendanceDailySummary.day.desc()).limit(5000).all()

    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = 'Daily Attendance'

    headers = ['اليوم', 'الموظف', 'البريد', 'جدول', 'أول دخول', 'آخر خروج', 'ساعات عمل', 'استراحة (د)', 'تأخير (د)', 'خروج مبكر (د)', 'إضافي (د)', 'الحالة']
    ws.append(headers)

    for r in rows:
        name = getattr(r.user, 'name', '') or ''
        email = getattr(r.user, 'email', '') or ''
        sched = getattr(r.schedule, 'name', '') if r.schedule else ''
        fi = r.first_in.isoformat(sep=' ', timespec='minutes') if r.first_in else ''
        lo = r.last_out.isoformat(sep=' ', timespec='minutes') if r.last_out else ''
        hours = round((r.work_minutes or 0) / 60.0, 2)
        ws.append([r.day, name, email, sched, fi, lo, hours, r.break_minutes or 0, r.late_minutes or 0, r.early_leave_minutes or 0, r.overtime_minutes or 0, r.status])

    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)

    return send_file(bio, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name='attendance_daily.xlsx')



# -------------------------
# Correspondence: Inbound / Outbound (independent register)
# -------------------------
@portal_bp.route("/corr")
@login_required
@_perm(CORR_READ)
def corr_index():
    return render_template("portal/corr/index.html")


def _corr_filters_inbound():

    q = (request.args.get("q") or "").strip()
    date_from = (request.args.get("date_from") or "").strip()
    date_to = (request.args.get("date_to") or "").strip()
    categories = _get_multi_arg("category", upper=True)  # list of codes
    senders = _get_multi_arg("sender")  # exact sender names (from lookups)
    sender_like = (request.args.get("sender_like") or "").strip()
    has_attach = (request.args.get("has_attach") or "").strip()  # 1/0

    qry = InboundMail.query

    if q:
        qry = apply_search_all_columns(qry, InboundMail, q)

    # Sender: multi exact OR free-text contains
    if senders:
        qry = qry.filter(InboundMail.sender.in_(senders))
    elif sender_like:
        qry = qry.filter(InboundMail.sender.ilike(f"%{sender_like}%"))

    if categories:
        qry = qry.filter(InboundMail.category.in_(categories))

    if date_from:
        qry = qry.filter(InboundMail.received_date >= date_from)
    if date_to:
        qry = qry.filter(InboundMail.received_date <= date_to)

    if has_attach == "1":
        qry = qry.filter(exists().where(CorrAttachment.inbound_id == InboundMail.id))
    if has_attach == "0":
        qry = qry.filter(~exists().where(CorrAttachment.inbound_id == InboundMail.id))

    filters = dict(
        q=q,
        date_from=date_from,
        date_to=date_to,
        categories=categories,
        senders=senders,
        sender_like=sender_like,
        has_attach=has_attach,
    )

    return qry, filters


def _corr_filters_outbound():

    q = (request.args.get("q") or "").strip()
    date_from = (request.args.get("date_from") or "").strip()
    date_to = (request.args.get("date_to") or "").strip()
    categories = _get_multi_arg("category", upper=True)  # list of codes
    recipients = _get_multi_arg("recipient")  # exact recipient names (from lookups)
    recipient_like = (request.args.get("recipient_like") or "").strip()
    has_attach = (request.args.get("has_attach") or "").strip()  # 1/0

    qry = OutboundMail.query

    if q:
        qry = apply_search_all_columns(qry, OutboundMail, q)

    if recipients:
        qry = qry.filter(OutboundMail.recipient.in_(recipients))
    elif recipient_like:
        qry = qry.filter(OutboundMail.recipient.ilike(f"%{recipient_like}%"))

    if categories:
        qry = qry.filter(OutboundMail.category.in_(categories))

    if date_from:
        qry = qry.filter(OutboundMail.sent_date >= date_from)
    if date_to:
        qry = qry.filter(OutboundMail.sent_date <= date_to)

    if has_attach == "1":
        qry = qry.filter(exists().where(CorrAttachment.outbound_id == OutboundMail.id))
    if has_attach == "0":
        qry = qry.filter(~exists().where(CorrAttachment.outbound_id == OutboundMail.id))

    filters = dict(
        q=q,
        date_from=date_from,
        date_to=date_to,
        categories=categories,
        recipients=recipients,
        recipient_like=recipient_like,
        has_attach=has_attach,
    )

    return qry, filters


@portal_bp.route("/corr/inbound")
@login_required
@_perm(CORR_READ)
def inbound_list():
    from models import CorrCategory, CorrParty, InboundMail

    qry, filters = _corr_filters_inbound()

    # Lookups
    cat_rows = CorrCategory.query.filter_by(is_active=True).order_by(CorrCategory.code.asc()).all()
    sender_rows = (
        CorrParty.query
        .filter(CorrParty.is_active == True)  # noqa: E712
        .filter(CorrParty.kind.in_(["SENDER", "BOTH"]))
        .order_by(CorrParty.name_ar.asc())
        .all()
    )

    # Sorting & pagination
    page = request.args.get("page", type=int, default=1)
    qry = qry.order_by(InboundMail.received_date.desc(), InboundMail.id.desc())

    # Excel export (respects current filters)
    if (request.args.get("export") or "").strip() in {"1", "excel"}:
        from utils.excel import make_xlsx_bytes
        items = qry.limit(5000).all()
        headers = [
            "#", "رقم", "تاريخ الاستلام", "التصنيف", "المرسل", "الموضوع", "عدد المرفقات"
        ]
        rows = []
        for idx, it in enumerate(items, start=1):
            rows.append([
                idx,
                it.ref_no or "",
                it.received_date or "",
                it.category or "",
                it.sender or "",
                it.subject or "",
                int(it.attachments.count()) if hasattr(it, "attachments") else "",
            ])
        xlsx_bytes = make_xlsx_bytes("Inbound", headers, rows)
        return send_file(
            io.BytesIO(xlsx_bytes),
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name="corr_inbound.xlsx",
        )

    pagination = qry.paginate(page=page, per_page=20, error_out=False)
    items = pagination.items

    return render_template(
        "portal/corr/inbound_list.html",
        items=items,
        pagination=pagination,
        cat_rows=cat_rows,
        sender_rows=sender_rows,
        **filters
    )




@portal_bp.route("/corr/inbound/new", methods=["GET", "POST"])
@login_required
@_perm(CORR_CREATE)
def inbound_new():
    cat_rows = CorrCategory.query.filter_by(is_active=True).order_by(CorrCategory.code.asc()).all()
    sender_rows = (
        CorrParty.query
        .filter(CorrParty.is_active == True)  # noqa: E712
        .filter(CorrParty.kind.in_(["SENDER", "BOTH"]))
        .order_by(CorrParty.name_ar.asc())
        .all()
    )

    if request.method == "POST":
        ref_no = (request.form.get("ref_no") or "").strip()

        category_code = (request.form.get("category_code") or "GENERAL").strip().upper()
        category_other = (request.form.get("category_other") or "").strip().upper()
        category = category_other if category_code == "__OTHER__" and category_other else category_code

        sender_val = (request.form.get("sender_val") or "").strip()
        sender_other = (request.form.get("sender_other") or "").strip()
        sender = sender_other if sender_val == "__OTHER__" else sender_val

        subject = (request.form.get("subject") or "").strip()
        body = (request.form.get("body") or "").strip()
        received_date = (request.form.get("received_date") or "").strip()

        if not received_date or not subject:
            flash("التاريخ والموضوع مطلوبان.", "danger")
            return redirect(request.url)

        files = request.files.getlist("files") or []
        if not any(getattr(f, "filename", "") for f in files):
            flash("رفع ملف/ملفات مطلوب لتسجيل الوارد.", "danger")
            return redirect(request.url)

        # Auto reference number (optional)
        if not ref_no:
            try:
                ref_no = _corr_next_ref("IN", received_date, category)
            except Exception:
                ref_no = None

        item = InboundMail(
            ref_no=ref_no or None,
            category=category or "GENERAL",
            sender=sender or None,
            subject=subject,
            body=body or None,
            received_date=received_date,
            created_by_id=current_user.id,
            created_at=datetime.utcnow(),
        )
        db.session.add(item)
        db.session.flush()  # assign item.id without committing

        saved = _save_corr_files(files, inbound_id=item.id, outbound_id=None)
        if not saved:
            db.session.rollback()
            flash("لم يتم رفع أي ملف (تحقق من نوع الملفات).", "danger")
            return redirect(request.url)

        db.session.commit()

        # Audit
        try:
            db.session.add(AuditLog(
                user_id=current_user.id,
                action="CORR_IN_CREATE",
                note=f"category={item.category} sender={item.sender or ''}",
                target_type="CORR_INBOUND",
                target_id=item.id,
                created_at=datetime.utcnow(),
            ))
            if saved:
                db.session.add(AuditLog(
                    user_id=current_user.id,
                    action="CORR_ATTACH_UPLOAD",
                    note=f"inbound_id={item.id} files={saved}",
                    target_type="CORR_INBOUND",
                    target_id=item.id,
                    created_at=datetime.utcnow(),
                ))
            db.session.commit()
        except Exception:
            db.session.rollback()

        flash("تم تسجيل الوارد.", "success")
        return redirect(url_for("portal.inbound_view", inbound_id=item.id))

    return render_template("portal/corr/inbound_new.html", cat_rows=cat_rows, sender_rows=sender_rows)


@portal_bp.route("/corr/outbound")
@login_required
@_perm(CORR_READ)
def outbound_list():
    from models import CorrCategory, CorrParty, OutboundMail

    qry, filters = _corr_filters_outbound()

    # Lookups
    cat_rows = CorrCategory.query.filter_by(is_active=True).order_by(CorrCategory.code.asc()).all()
    recipient_rows = (
        CorrParty.query
        .filter(CorrParty.is_active == True)  # noqa: E712
        .filter(CorrParty.kind.in_(["RECIPIENT", "BOTH"]))
        .order_by(CorrParty.name_ar.asc())
        .all()
    )

    # Sorting & pagination
    page = request.args.get("page", type=int, default=1)
    qry = qry.order_by(OutboundMail.sent_date.desc(), OutboundMail.id.desc())

    # Excel export (respects current filters)
    if (request.args.get("export") or "").strip() in {"1", "excel"}:
        from utils.excel import make_xlsx_bytes
        items = qry.limit(5000).all()
        headers = [
            "#", "رقم", "تاريخ الإرسال", "التصنيف", "الجهة", "الموضوع", "عدد المرفقات"
        ]
        rows = []
        for idx, it in enumerate(items, start=1):
            rows.append([
                idx,
                it.ref_no or "",
                it.sent_date or "",
                it.category or "",
                it.recipient or "",
                it.subject or "",
                int(it.attachments.count()) if hasattr(it, "attachments") else "",
            ])
        xlsx_bytes = make_xlsx_bytes("Outbound", headers, rows)
        return send_file(
            io.BytesIO(xlsx_bytes),
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name="corr_outbound.xlsx",
        )

    pagination = qry.paginate(page=page, per_page=20, error_out=False)
    items = pagination.items

    return render_template(
        "portal/corr/outbound_list.html",
        can_update=bool(current_user.has_perm(CORR_UPDATE) or _can_manage_corr()),
        can_delete=bool(current_user.has_perm(CORR_DELETE) or _can_manage_corr()),
        items=items,
        pagination=pagination,
        cat_rows=cat_rows,
        recipient_rows=recipient_rows,
        **filters
    )


@portal_bp.route("/corr/outbound/new", methods=["GET", "POST"])
@login_required
@_perm(CORR_CREATE)
def outbound_new():
    cat_rows = CorrCategory.query.filter_by(is_active=True).order_by(CorrCategory.code.asc()).all()
    recipient_rows = (
        CorrParty.query
        .filter(CorrParty.is_active == True)  # noqa: E712
        .filter(CorrParty.kind.in_(["RECIPIENT", "BOTH"]))
        .order_by(CorrParty.name_ar.asc())
        .all()
    )

    if request.method == "POST":
        ref_no = (request.form.get("ref_no") or "").strip()

        category_code = (request.form.get("category_code") or "GENERAL").strip().upper()
        category_other = (request.form.get("category_other") or "").strip().upper()
        category = category_other if category_code == "__OTHER__" and category_other else category_code

        recipient_val = (request.form.get("recipient_val") or "").strip()
        recipient_other = (request.form.get("recipient_other") or "").strip()
        recipient = recipient_other if recipient_val == "__OTHER__" else recipient_val

        subject = (request.form.get("subject") or "").strip()
        body = (request.form.get("body") or "").strip()
        sent_date = (request.form.get("sent_date") or "").strip()

        if not sent_date or not subject:
            flash("التاريخ والموضوع مطلوبان.", "danger")
            return redirect(request.url)


        files = request.files.getlist("files") or []
        if not any(getattr(f, "filename", "") for f in files):
            flash("رفع ملف/ملفات مطلوب لتسجيل الصادر.", "danger")
            return redirect(request.url)

        # Auto reference number (optional)
        if not ref_no:
            try:
                ref_no = _corr_next_ref("OUT", sent_date, category)
            except Exception:
                ref_no = None

        item = OutboundMail(
            ref_no=ref_no or None,
            category=category or "GENERAL",
            recipient=recipient or None,
            subject=subject,
            body=body or None,
            sent_date=sent_date,
            created_by_id=current_user.id,
            created_at=datetime.utcnow(),
        )
        db.session.add(item)
        db.session.flush()  # assign item.id without committing

        saved = _save_corr_files(files, inbound_id=None, outbound_id=item.id)
        if not saved:
            db.session.rollback()
            flash("لم يتم رفع أي ملف (تحقق من نوع الملفات).", "danger")
            return redirect(request.url)

        db.session.commit()

        # Audit
        try:
            db.session.add(AuditLog(
                user_id=current_user.id,
                action="CORR_OUT_CREATE",
                note=f"category={item.category} recipient={item.recipient or ''}",
                target_type="CORR_OUTBOUND",
                target_id=item.id,
                created_at=datetime.utcnow(),
            ))
            if saved:
                db.session.add(AuditLog(
                    user_id=current_user.id,
                    action="CORR_ATTACH_UPLOAD",
                    note=f"outbound_id={item.id} files={saved}",
                    target_type="CORR_OUTBOUND",
                    target_id=item.id,
                    created_at=datetime.utcnow(),
                ))
            db.session.commit()
        except Exception:
            db.session.rollback()

        flash("تم تسجيل الصادر.", "success")
        return redirect(url_for("portal.outbound_view", outbound_id=item.id))

    return render_template("portal/corr/outbound_new.html", cat_rows=cat_rows, recipient_rows=recipient_rows)




@portal_bp.route("/corr/outbound/<int:outbound_id>")
@login_required
@_perm(CORR_READ)
def outbound_view(outbound_id: int):
    item = OutboundMail.query.get_or_404(outbound_id)
    attachments = CorrAttachment.query.filter_by(outbound_id=item.id).order_by(CorrAttachment.id.desc()).all()

    logs = (
        AuditLog.query
        .filter(AuditLog.target_type == "CORR_OUTBOUND", AuditLog.target_id == item.id)
        .order_by(AuditLog.created_at.desc())
        .limit(80)
        .all()
    )

    can_update = bool(current_user.has_perm(CORR_UPDATE) or _can_manage_corr())
    can_delete = bool(current_user.has_perm(CORR_DELETE) or _can_manage_corr())

    can_upload = bool(current_user.has_perm(CORR_CREATE) or can_update)

    return render_template(
        "portal/corr/outbound_view.html",
        item=item,
        attachments=attachments,
        logs=logs,
        can_update=can_update,
        can_delete=can_delete,
        can_upload=can_upload,
    )


@portal_bp.route("/corr/inbound/<int:inbound_id>")
@login_required
@_perm(CORR_READ)
def inbound_view(inbound_id: int):
    item = InboundMail.query.get_or_404(inbound_id)
    attachments = CorrAttachment.query.filter_by(inbound_id=item.id).order_by(CorrAttachment.id.desc()).all()

    logs = (
        AuditLog.query
        .filter(AuditLog.target_type == "CORR_INBOUND", AuditLog.target_id == item.id)
        .order_by(AuditLog.created_at.desc())
        .limit(80)
        .all()
    )

    can_update = bool(current_user.has_perm(CORR_UPDATE) or _can_manage_corr())
    can_delete = bool(current_user.has_perm(CORR_DELETE) or _can_manage_corr())

    can_upload = bool(current_user.has_perm(CORR_CREATE) or can_update)

    return render_template(
        "portal/corr/inbound_view.html",
        item=item,
        attachments=attachments,
        logs=logs,
        can_update=can_update,
        can_delete=can_delete,
        can_upload=can_upload,
    )



@portal_bp.route("/corr/inbound/<int:inbound_id>/upload", methods=["POST"])
@login_required
@_perm(CORR_CREATE)
def inbound_upload(inbound_id: int):
    item = InboundMail.query.get_or_404(inbound_id)

    files = request.files.getlist("files") or []
    if not files:
        f1 = request.files.get("file")
        files = [f1] if f1 else []

    if not files or not any(getattr(f, "filename", "") for f in files):
        flash("اختر ملفاً.", "danger")
        return redirect(url_for("portal.inbound_view", inbound_id=item.id))

    saved = _save_corr_files(files, inbound_id=item.id, outbound_id=None)
    if not saved:
        flash("لم يتم رفع أي ملف (تحقق من نوع الملفات).", "danger")
        return redirect(url_for("portal.inbound_view", inbound_id=item.id))

    db.session.commit()

    try:
        db.session.add(AuditLog(
            user_id=current_user.id,
            action="CORR_ATTACH_UPLOAD",
            note=f"inbound_id={item.id} files={saved}",
            target_type="CORR_INBOUND",
            target_id=item.id,
            created_at=datetime.utcnow(),
        ))
        db.session.commit()
    except Exception:
        db.session.rollback()

    flash(f"تم رفع {saved} ملف/ملفات.", "success")
    return redirect(url_for("portal.inbound_view", inbound_id=item.id))



@portal_bp.route("/corr/outbound/<int:outbound_id>/upload", methods=["POST"])
@login_required
@_perm(CORR_CREATE)
def outbound_upload(outbound_id: int):
    item = OutboundMail.query.get_or_404(outbound_id)

    files = request.files.getlist("files") or []
    if not files:
        f1 = request.files.get("file")
        files = [f1] if f1 else []

    if not files or not any(getattr(f, "filename", "") for f in files):
        flash("اختر ملفاً.", "danger")
        return redirect(url_for("portal.outbound_view", outbound_id=item.id))

    saved = _save_corr_files(files, inbound_id=None, outbound_id=item.id)
    if not saved:
        flash("لم يتم رفع أي ملف (تحقق من نوع الملفات).", "danger")
        return redirect(url_for("portal.outbound_view", outbound_id=item.id))

    db.session.commit()

    try:
        db.session.add(AuditLog(
            user_id=current_user.id,
            action="CORR_ATTACH_UPLOAD",
            note=f"outbound_id={item.id} files={saved}",
            target_type="CORR_OUTBOUND",
            target_id=item.id,
            created_at=datetime.utcnow(),
        ))
        db.session.commit()
    except Exception:
        db.session.rollback()

    flash(f"تم رفع {saved} ملف/ملفات.", "success")
    return redirect(url_for("portal.outbound_view", outbound_id=item.id))






@portal_bp.route("/corr/inbound/<int:inbound_id>/edit", methods=["GET", "POST"])
@login_required
def inbound_edit(inbound_id: int):
    item = InboundMail.query.get_or_404(inbound_id)
    if not (current_user.has_perm(CORR_UPDATE) or _can_manage_corr()):
        abort(403)


    cat_rows = CorrCategory.query.filter_by(is_active=True).order_by(CorrCategory.code.asc()).all()
    sender_rows = (
        CorrParty.query
        .filter(CorrParty.is_active == True)  # noqa: E712
        .filter(CorrParty.kind.in_(["SENDER", "BOTH"]))
        .order_by(CorrParty.name_ar.asc())
        .all()
    )

    if request.method == "POST":
        ref_no = (request.form.get("ref_no") or "").strip()

        category_code = (request.form.get("category_code") or (item.category or "GENERAL")).strip().upper()
        category_other = (request.form.get("category_other") or "").strip().upper()
        category = category_other if category_code == "__OTHER__" and category_other else category_code

        sender_val = (request.form.get("sender_val") or "").strip()
        sender_other = (request.form.get("sender_other") or "").strip()
        sender = sender_other if sender_val == "__OTHER__" else sender_val

        subject = (request.form.get("subject") or "").strip()
        body = (request.form.get("body") or "").strip()
        received_date = (request.form.get("received_date") or "").strip()

        if not received_date or not subject:
            flash("التاريخ والموضوع مطلوبان.", "danger")
            return redirect(request.url)

        item.ref_no = ref_no or item.ref_no
        item.category = category or "GENERAL"
        item.sender = sender or None
        item.subject = subject
        item.body = body or None
        item.received_date = received_date

        saved = 0
        files = request.files.getlist("files") or []
        if any(getattr(f, "filename", "") for f in files):
            saved = _save_corr_files(files, inbound_id=item.id, outbound_id=None)

        try:
            db.session.add(AuditLog(
                user_id=current_user.id,
                action="CORR_IN_UPDATE",
                note=f"category={item.category} sender={item.sender or ''}",
                target_type="CORR_INBOUND",
                target_id=item.id,
                created_at=datetime.utcnow(),
            ))
            if saved:
                db.session.add(AuditLog(
                    user_id=current_user.id,
                    action="CORR_ATTACH_UPLOAD",
                    note=f"inbound_id={item.id} files={saved}",
                    target_type="CORR_INBOUND",
                    target_id=item.id,
                    created_at=datetime.utcnow(),
                ))
            db.session.commit()
        except Exception:
            db.session.rollback()
            flash("تعذر حفظ التعديلات.", "danger")
            return redirect(request.url)

        flash("تم تحديث الوارد.", "success")
        return redirect(url_for("portal.inbound_view", inbound_id=item.id))

    attachments = CorrAttachment.query.filter_by(inbound_id=item.id).order_by(CorrAttachment.id.desc()).all()
    logs = (
        AuditLog.query
        .filter(AuditLog.target_type == "CORR_INBOUND", AuditLog.target_id == item.id)
        .order_by(AuditLog.created_at.desc())
        .limit(80)
        .all()
    )
    can_delete = bool(current_user.has_perm(CORR_DELETE) or _can_manage_corr())

    cat_codes = [c.code for c in cat_rows]
    sender_labels = [p.label for p in sender_rows]

    selected_category_code = item.category if (item.category in cat_codes) else "__OTHER__"
    category_other_value = "" if selected_category_code != "__OTHER__" else (item.category or "")

    if item.sender and (item.sender in sender_labels):
        selected_sender_val = item.sender
        sender_other_value = ""
    elif item.sender:
        selected_sender_val = "__OTHER__"
        sender_other_value = item.sender
    else:
        selected_sender_val = ""
        sender_other_value = ""

    return render_template(
        "portal/corr/inbound_edit.html",
        item=item,
        cat_rows=cat_rows,
        sender_rows=sender_rows,
        selected_category_code=selected_category_code,
        category_other_value=category_other_value,
        selected_sender_val=selected_sender_val,
        sender_other_value=sender_other_value,
        attachments=attachments,
        logs=logs,
        can_delete=can_delete,
    )


@portal_bp.route("/corr/outbound/<int:outbound_id>/edit", methods=["GET", "POST"])
@login_required
def outbound_edit(outbound_id: int):
    item = OutboundMail.query.get_or_404(outbound_id)
    if not (current_user.has_perm(CORR_UPDATE) or _can_manage_corr()):
        abort(403)


    cat_rows = CorrCategory.query.filter_by(is_active=True).order_by(CorrCategory.code.asc()).all()
    recipient_rows = (
        CorrParty.query
        .filter(CorrParty.is_active == True)  # noqa: E712
        .filter(CorrParty.kind.in_(["RECIPIENT", "BOTH"]))
        .order_by(CorrParty.name_ar.asc())
        .all()
    )

    if request.method == "POST":
        ref_no = (request.form.get("ref_no") or "").strip()

        category_code = (request.form.get("category_code") or (item.category or "GENERAL")).strip().upper()
        category_other = (request.form.get("category_other") or "").strip().upper()
        category = category_other if category_code == "__OTHER__" and category_other else category_code

        recipient_val = (request.form.get("recipient_val") or "").strip()
        recipient_other = (request.form.get("recipient_other") or "").strip()
        recipient = recipient_other if recipient_val == "__OTHER__" else recipient_val

        subject = (request.form.get("subject") or "").strip()
        body = (request.form.get("body") or "").strip()
        sent_date = (request.form.get("sent_date") or "").strip()

        if not sent_date or not subject:
            flash("التاريخ والموضوع مطلوبان.", "danger")
            return redirect(request.url)

        item.ref_no = ref_no or item.ref_no
        item.category = category or "GENERAL"
        item.recipient = recipient or None
        item.subject = subject
        item.body = body or None
        item.sent_date = sent_date

        saved = 0
        files = request.files.getlist("files") or []
        if any(getattr(f, "filename", "") for f in files):
            saved = _save_corr_files(files, inbound_id=None, outbound_id=item.id)

        try:
            db.session.add(AuditLog(
                user_id=current_user.id,
                action="CORR_OUT_UPDATE",
                note=f"category={item.category} recipient={item.recipient or ''}",
                target_type="CORR_OUTBOUND",
                target_id=item.id,
                created_at=datetime.utcnow(),
            ))
            if saved:
                db.session.add(AuditLog(
                    user_id=current_user.id,
                    action="CORR_ATTACH_UPLOAD",
                    note=f"outbound_id={item.id} files={saved}",
                    target_type="CORR_OUTBOUND",
                    target_id=item.id,
                    created_at=datetime.utcnow(),
                ))
            db.session.commit()
        except Exception:
            db.session.rollback()
            flash("تعذر حفظ التعديلات.", "danger")
            return redirect(request.url)

        flash("تم تحديث الصادر.", "success")
        return redirect(url_for("portal.outbound_view", outbound_id=item.id))

    attachments = CorrAttachment.query.filter_by(outbound_id=item.id).order_by(CorrAttachment.id.desc()).all()
    logs = (
        AuditLog.query
        .filter(AuditLog.target_type == "CORR_OUTBOUND", AuditLog.target_id == item.id)
        .order_by(AuditLog.created_at.desc())
        .limit(80)
        .all()
    )
    can_delete = bool(current_user.has_perm(CORR_DELETE) or _can_manage_corr())

    cat_codes = [c.code for c in cat_rows]
    recipient_labels = [p.label for p in recipient_rows]

    selected_category_code = item.category if (item.category in cat_codes) else "__OTHER__"
    category_other_value = "" if selected_category_code != "__OTHER__" else (item.category or "")

    if item.recipient and (item.recipient in recipient_labels):
        selected_recipient_val = item.recipient
        recipient_other_value = ""
    elif item.recipient:
        selected_recipient_val = "__OTHER__"
        recipient_other_value = item.recipient
    else:
        selected_recipient_val = ""
        recipient_other_value = ""

    return render_template(
        "portal/corr/outbound_edit.html",
        item=item,
        cat_rows=cat_rows,
        recipient_rows=recipient_rows,
        selected_category_code=selected_category_code,
        category_other_value=category_other_value,
        selected_recipient_val=selected_recipient_val,
        recipient_other_value=recipient_other_value,
        attachments=attachments,
        logs=logs,
        can_delete=can_delete,
    )


@portal_bp.route("/corr/inbound/<int:inbound_id>/delete", methods=["POST"])
@login_required
def inbound_delete(inbound_id: int):
    item = InboundMail.query.get_or_404(inbound_id)
    if not (current_user.has_perm(CORR_DELETE) or _can_manage_corr()):
        abort(403)

    storage = _corr_storage_dir()

    atts = CorrAttachment.query.filter_by(inbound_id=item.id).all()
    for a in atts:
        try:
            fp = os.path.join(storage, a.stored_name)
            if os.path.exists(fp):
                os.remove(fp)
        except Exception:
            pass
        try:
            db.session.delete(a)
        except Exception:
            pass

    try:
        db.session.add(AuditLog(
            user_id=current_user.id,
            action="CORR_IN_DELETE",
            note=f"ref={item.ref_no or item.id}",
            target_type="CORR_INBOUND",
            target_id=item.id,
            created_at=datetime.utcnow(),
        ))
        db.session.delete(item)
        db.session.commit()
    except Exception:
        db.session.rollback()
        flash("تعذر حذف الوارد.", "danger")
        return redirect(url_for("portal.inbound_view", inbound_id=item.id))

    flash("تم حذف الوارد.", "success")
    return redirect(url_for("portal.inbound_list"))


@portal_bp.route("/corr/outbound/<int:outbound_id>/delete", methods=["POST"])
@login_required
def outbound_delete(outbound_id: int):
    item = OutboundMail.query.get_or_404(outbound_id)
    if not (current_user.has_perm(CORR_DELETE) or _can_manage_corr()):
        abort(403)

    storage = _corr_storage_dir()

    atts = CorrAttachment.query.filter_by(outbound_id=item.id).all()
    for a in atts:
        try:
            fp = os.path.join(storage, a.stored_name)
            if os.path.exists(fp):
                os.remove(fp)
        except Exception:
            pass
        try:
            db.session.delete(a)
        except Exception:
            pass

    try:
        db.session.add(AuditLog(
            user_id=current_user.id,
            action="CORR_OUT_DELETE",
            note=f"ref={item.ref_no or item.id}",
            target_type="CORR_OUTBOUND",
            target_id=item.id,
            created_at=datetime.utcnow(),
        ))
        db.session.delete(item)
        db.session.commit()
    except Exception:
        db.session.rollback()
        flash("تعذر حذف الصادر.", "danger")
        return redirect(url_for("portal.outbound_view", outbound_id=item.id))

    flash("تم حذف الصادر.", "success")
    return redirect(url_for("portal.outbound_list"))

@portal_bp.route("/corr/attachment/<int:att_id>/download")
@login_required
@_perm(CORR_READ)
def corr_attachment_download(att_id: int):
    att = CorrAttachment.query.get_or_404(att_id)
    storage = _corr_storage_dir()
    file_path = os.path.join(storage, att.stored_name)
    if not os.path.exists(file_path):
        flash("الملف غير موجود.", "danger")
        if att.inbound_id:
            return redirect(url_for("portal.inbound_view", inbound_id=att.inbound_id))
        if att.outbound_id:
            return redirect(url_for("portal.outbound_view", outbound_id=att.outbound_id))
        return redirect(url_for("portal.corr_index"))

    return send_from_directory(storage, att.stored_name, as_attachment=True, download_name=att.original_name)


@portal_bp.route("/corr/attachment/<int:att_id>/view")
@login_required
@_perm(CORR_READ)
def corr_attachment_view(att_id: int):
    att = CorrAttachment.query.get_or_404(att_id)
    storage = _corr_storage_dir()
    file_path = os.path.join(storage, att.stored_name)
    if not os.path.exists(file_path):
        flash("الملف غير موجود.", "danger")
        if att.inbound_id:
            return redirect(url_for("portal.inbound_view", inbound_id=att.inbound_id))
        if att.outbound_id:
            return redirect(url_for("portal.outbound_view", outbound_id=att.outbound_id))
        return redirect(url_for("portal.corr_index"))

    mime, _ = mimetypes.guess_type(file_path)
    return send_from_directory(storage, att.stored_name, as_attachment=False, mimetype=mime or "application/octet-stream")


@portal_bp.route("/corr/attachment/<int:att_id>/delete", methods=["POST"])
@login_required
@_perm(CORR_READ)
def corr_attachment_delete(att_id: int):
    att = CorrAttachment.query.get_or_404(att_id)
    storage = _corr_storage_dir()

    # allow delete by uploader or CORR_MANAGE/admin
    allowed = False
    if getattr(att, "uploaded_by_id", None) and int(att.uploaded_by_id) == int(current_user.id):
        allowed = True
    if _can_manage_corr():
        allowed = True

    if not allowed:
        abort(403)

    file_path = os.path.join(storage, att.stored_name)
    inbound_id = att.inbound_id
    outbound_id = att.outbound_id

    try:
        db.session.delete(att)
        db.session.commit()
    except Exception:
        db.session.rollback()
        flash("تعذر حذف المرفق.", "danger")
        if inbound_id:
            return redirect(url_for("portal.inbound_view", inbound_id=inbound_id))
        if outbound_id:
            return redirect(url_for("portal.outbound_view", outbound_id=outbound_id))
        return redirect(url_for("portal.corr_index"))

    # remove file from disk
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
    except Exception:
        pass

    flash("تم حذف المرفق.", "success")
    if inbound_id:
        return redirect(url_for("portal.inbound_view", inbound_id=inbound_id))
    if outbound_id:
        return redirect(url_for("portal.outbound_view", outbound_id=outbound_id))
    return redirect(url_for("portal.corr_index"))


def _csv_response(filename: str, rows: list[list[str]]):
    import csv, io
    output = io.StringIO()
    writer = csv.writer(output)
    for r in rows:
        writer.writerow(r)
    data = output.getvalue().encode("utf-8-sig")
    from flask import Response
    resp = Response(data, mimetype="text/csv; charset=utf-8")
    resp.headers["Content-Disposition"] = f"attachment; filename={filename}"
    return resp


@portal_bp.route("/corr/inbound/export.csv")
@login_required
@_perm(CORR_READ)
def inbound_export_csv():
    qry, filters = _corr_filters_inbound()
    items = qry.order_by(InboundMail.received_date.desc(), InboundMail.id.desc()).limit(1000).all()

    rows = [["ID", "Received Date", "Ref No", "Sender", "Category", "Subject"]]
    for x in items:
        rows.append([str(x.id), str(x.received_date), str(x.ref_no or ""), str(x.sender or ""), str(x.category or ""), str(x.subject or "")])

    return _csv_response("inbound.csv", rows)


@portal_bp.route("/corr/outbound/export.csv")
@login_required
@_perm(CORR_READ)
def outbound_export_csv():
    qry, filters = _corr_filters_outbound()
    items = qry.order_by(OutboundMail.sent_date.desc(), OutboundMail.id.desc()).limit(1000).all()

    rows = [["ID", "Sent Date", "Ref No", "Recipient", "Category", "Subject"]]
    for x in items:
        rows.append([str(x.id), str(x.sent_date), str(x.ref_no or ""), str(x.recipient or ""), str(x.category or ""), str(x.subject or "")])

    return _csv_response("outbound.csv", rows)




def _ensure_pdf_font():
    """Register DejaVu font for Arabic-friendly PDFs when available."""
    try:
        base_dir = os.path.join(os.getcwd(), "assets", "fonts")
        reg_path = os.path.join(base_dir, "DejaVuSans.ttf")
        bold_path = os.path.join(base_dir, "DejaVuSans-Bold.ttf")
        if os.path.exists(reg_path) and "DejaVuSans" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("DejaVuSans", reg_path))
        if os.path.exists(bold_path) and "DejaVuSans-Bold" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", bold_path))
        return True
    except Exception:
        return False

def _build_corr_list_pdf(title: str, rows: list[list[str]], filters: dict):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=24, leftMargin=24, topMargin=24, bottomMargin=24)
    _ensure_pdf_font()
    styles = getSampleStyleSheet()

    # prefer DejaVu font if registered
    try:
        styles["Title"].fontName = "DejaVuSans-Bold"
        styles["Normal"].fontName = "DejaVuSans"
    except Exception:
        pass

    story = []
    story.append(Paragraph(title, styles["Title"]))
    # filters summary
    flines = []
    for k, v in filters.items():
        if v:
            if isinstance(v, list):
                v = ", ".join([str(x) for x in v])
            flines.append(f"{k}: {v}")
    if flines:
        story.append(Paragraph(" | ".join(flines), styles["Normal"]))
    story.append(Spacer(1, 12))

    tbl = Table(rows, repeatRows=1, hAlign="RIGHT")
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
        ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
        ("FONTNAME", (0,0), (-1,-1), "DejaVuSans"),
        ("FONTSIZE", (0,0), (-1,-1), 9),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
    ]))
    story.append(tbl)
    doc.build(story)
    buffer.seek(0)
    return buffer


@portal_bp.route("/corr/inbound/export.pdf")
@login_required
@_perm(CORR_READ)
def inbound_export_pdf():
    qry, filters = _corr_filters_inbound()
    items = qry.order_by(InboundMail.received_date.desc(), InboundMail.id.desc()).limit(1000).all()

    rows = [["#", "التاريخ", "الرقم", "التصنيف", "الجهة", "الموضوع"]]
    for x in items:
        rows.append([str(x.id), str(x.received_date), str(x.ref_no or ""), str(x.category or ""), str(x.sender or ""), str(x.subject or "")])

    buf = _build_corr_list_pdf("تقرير الوارد - البوابة الإدارية", rows, filters)
    from flask import send_file
    return send_file(buf, mimetype="application/pdf", as_attachment=True, download_name="inbound.pdf")


@portal_bp.route("/corr/outbound/export.pdf")
@login_required
@_perm(CORR_READ)
def outbound_export_pdf():
    qry, filters = _corr_filters_outbound()
    items = qry.order_by(OutboundMail.sent_date.desc(), OutboundMail.id.desc()).limit(1000).all()

    rows = [["#", "التاريخ", "الرقم", "التصنيف", "الجهة", "الموضوع"]]
    for x in items:
        rows.append([str(x.id), str(x.sent_date), str(x.ref_no or ""), str(x.category or ""), str(x.recipient or ""), str(x.subject or "")])

    buf = _build_corr_list_pdf("تقرير الصادر - البوابة الإدارية", rows, filters)
    from flask import send_file
    return send_file(buf, mimetype="application/pdf", as_attachment=True, download_name="outbound.pdf")


def _build_corr_card_pdf(kind: str, ref_no: str, date_s: str, party: str, category: str, subject: str, notes: str | None, url: str):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    _ensure_pdf_font()
    styles = getSampleStyleSheet()
    try:
        styles["Title"].fontName = "DejaVuSans-Bold"
        styles["Normal"].fontName = "DejaVuSans"
        styles["Heading3"].fontName = "DejaVuSans-Bold"
    except Exception:
        pass

    title = "بطاقة وارد" if kind == "IN" else "بطاقة صادر"
    story = []
    story.append(Paragraph(title, styles["Title"]))
    story.append(Spacer(1, 12))

    data = [
        ["الرقم", ref_no],
        ["التاريخ", date_s],
        ["التصنيف", category or ""],
        ["الجهة", party or ""],
        ["الموضوع", subject or ""],
    ]
    t = Table(data, colWidths=[120, 360], hAlign="RIGHT")
    t.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
        ("BACKGROUND", (0,0), (0,-1), colors.whitesmoke),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
        ("FONTNAME", (0,0), (-1,-1), "DejaVuSans"),
        ("FONTSIZE", (0,0), (-1,-1), 11),
    ]))
    story.append(t)

    if notes:
        story.append(Spacer(1, 12))
        story.append(Paragraph("ملاحظات:", styles["Heading3"]))
        story.append(Paragraph(notes.replace("\n", "<br/>"), styles["Normal"]))

    story.append(Spacer(1, 18))
    story.append(Paragraph("QR للوصول للصفحة:", styles["Heading3"]))
    # QR
    qrw = qr.QrCodeWidget(url)
    from reportlab.graphics.shapes import Drawing
    bounds = qrw.getBounds()
    w = bounds[2] - bounds[0]
    h = bounds[3] - bounds[1]
    size = 120
    d = Drawing(size, size, transform=[size / w, 0, 0, size / h, 0, 0])
    d.add(qrw)
    story.append(d)

    doc.build(story)
    buffer.seek(0)
    return buffer


@portal_bp.route("/corr/inbound/<int:inbound_id>/print.pdf")
@login_required
@_perm(CORR_READ)
def inbound_print_pdf(inbound_id: int):
    item = InboundMail.query.get_or_404(inbound_id)
    url = request.host_url.rstrip("/") + url_for("portal.inbound_view", inbound_id=item.id)
    buf = _build_corr_card_pdf(
        "IN",
        item.ref_no or f"#{item.id}",
        item.received_date,
        item.sender or "",
        item.category or "",
        item.subject or "",
        item.body,
        url
    )
    from flask import send_file
    return send_file(buf, mimetype="application/pdf", as_attachment=True, download_name="inbound_card.pdf")


@portal_bp.route("/corr/outbound/<int:outbound_id>/print.pdf")
@login_required
@_perm(CORR_READ)
def outbound_print_pdf(outbound_id: int):
    item = OutboundMail.query.get_or_404(outbound_id)
    url = request.host_url.rstrip("/") + url_for("portal.outbound_view", outbound_id=item.id)
    buf = _build_corr_card_pdf(
        "OUT",
        item.ref_no or f"#{item.id}",
        item.sent_date,
        item.recipient or "",
        item.category or "",
        item.subject or "",
        item.body,
        url
    )
    from flask import send_file
    return send_file(buf, mimetype="application/pdf", as_attachment=True, download_name="outbound_card.pdf")

# -------------------------
# Portal Admin
# -------------------------
from portal.perm_defs import PERMS as PORTAL_PERMS, ALL_KEYS as PORTAL_ALL_KEYS


@portal_bp.route("/admin")
@login_required
@_perm(PORTAL_ADMIN_READ)
def portal_admin_dashboard():
    # UX: show only cards user can access
    cards = []

    def add_card(required_perm, title, desc, icon, endpoint):
        try:
            if current_user.has_perm(required_perm):
                cards.append({
                    "title": title,
                    "desc": desc,
                    "icon": icon,
                    "url": url_for(endpoint),
                })
        except Exception:
            pass

    # Alerts / counts
    pending_access = 0
    try:
        if current_user.has_perm(PORTAL_ADMIN_PERMISSIONS_MANAGE):
            pending_access = _safe_count(PortalAccessRequest.query.filter(PortalAccessRequest.status == 'PENDING'))
    except Exception:
        pending_access = 0

    # Core admin cards
    add_card(PORTAL_ADMIN_PERMISSIONS_MANAGE, "طلبات الصلاحيات", "مراجعة طلبات الموظفين لتفعيل خدمات البوابة.", "bi-inboxes", "portal.admin_access_requests")
    add_card(PORTAL_ADMIN_PERMISSIONS_MANAGE, "صلاحيات البوابة", "تعديل الصلاحيات للأدوار أو للمستخدمين.", "bi-person-gear", "portal.portal_admin_permissions")

    # HR admin
    add_card(HR_MASTERDATA_MANAGE, "الهيكل التنظيمي", "إدارة المنظمات/الإدارات/الدوائر/الأقسام/الفرق.", "bi-diagram-3", "portal.portal_admin_hr_org_structure")
    add_card(HR_MASTERDATA_MANAGE, "إعدادات الدوام", "تعريف أنواع الإجازات/المغادرات والجداول.", "bi-people", "portal.hr_masterdata_index")
    add_card(HR_MASTERDATA_MANAGE, "قوالب الدوام", "تعريف قوالب زمنية (08:00-15:00...) مع السماح.", "bi-clock", "portal.portal_admin_hr_schedule_templates")
    add_card(HR_MASTERDATA_MANAGE, "سياسات الدوام", "تعريف سياسات الأيام/المكان (On-site/Hybrid/Remote).", "bi-shield-check", "portal.portal_admin_hr_work_policies")
    add_card(HR_MASTERDATA_MANAGE, "تعيين الدوام", "تعيين قالب + سياسة لمستخدم/دور/قسم مع تاريخ من/إلى.", "bi-person-check", "portal.portal_admin_hr_work_assignments")
    add_card(HR_PERF_MANAGE, "الأداء والتقييم", "نماذج ودورات وتكليفات 360.", "bi-clipboard2-check", "portal.portal_admin_hr_perf_dashboard")
    add_card(PORTAL_REPORTS_READ, "التقارير والمؤشرات", "لوحة مؤشرات وتقارير HR (تصدير).", "bi-bar-chart-line", "portal.portal_admin_hr_reports_dashboard")
    add_card(PORTAL_AUDIT_READ, "التدقيق والامتثال", "سجلات تدقيق البوابة + تقارير امتثال.", "bi-shield-lock", "portal.portal_admin_compliance")
    add_card(PORTAL_AUDIT_READ, "السجل الزمني", "Timeline موحّد لحركات البوابة.", "bi-clock-history", "portal.portal_admin_timeline")
    add_card(HR_MASTERDATA_MANAGE, "لوحة HR", "لوحة موارد بشرية للإدارة (موظفين/دوام/طلبات/إعدادات).", "bi-clipboard2-data", "portal.portal_admin_hr_dashboard")

    # Delegation
    add_card(PORTAL_ADMIN_PERMISSIONS_MANAGE, "التفويض", "تفويض صلاحيات/مهام لموظف بديل لفترة محددة.", "bi-arrow-left-right", "portal.portal_admin_delegations")

    # Corr admin
    add_card(CORR_LOOKUPS_MANAGE, "إعدادات المراسلات", "إدارة التصنيفات والجهات المرسلة/المستلمة.", "bi-tags", "portal.portal_admin_corr_index")

    # Integrations
    add_card(PORTAL_INTEGRATIONS_MANAGE, "التكاملات", "إعداد مزامنة ملف ساعة الدوام من السيرفر.", "bi-plug", "portal.hr_attendance_import")

    return render_template("portal/admin/index.html", cards=cards, pending_access=pending_access)





@portal_bp.route("/admin/hr", methods=["GET"])
@login_required
@_perm(PORTAL_READ)
def portal_admin_hr_dashboard():
    """HR Admin dashboard (simple, clear cards)."""
    # Gate: require at least one HR admin-like permission
    allowed = False
    try:
        allowed = any([
            current_user.has_perm(HR_MASTERDATA_MANAGE),
            current_user.has_perm(HR_EMP_READ),
            current_user.has_perm(HR_ATT_READ),
            current_user.has_perm(HR_ORG_READ),
            current_user.has_perm(HR_REQUESTS_APPROVE),
            current_user.has_perm(HR_REQUESTS_VIEW_ALL),
        ])
    except Exception:
        allowed = False
    if not allowed:
        abort(403)

    cards = []

    def add_card(required_perm, title, desc, icon, endpoint):
        try:
            if current_user.has_perm(required_perm):
                cards.append({
                    "title": title,
                    "desc": desc,
                    "icon": icon,
                    "url": url_for(endpoint),
                })
        except Exception:
            pass

    # Primary HR admin tools
    add_card(HR_EMP_READ, "ملفات الموظفين", "عرض ملفات الموظفين ورفع المرفقات.", "bi-person-lines-fill", "portal.hr_employees")
    add_card(HR_ATT_READ, "سجل الدوام", "عرض أحداث ساعة الدوام والدفعات.", "bi-clock-history", "portal.hr_attendance_events")
    add_card(HR_REQUESTS_APPROVE, "اعتمادات HR", "مراجعة طلبات الإجازات والمغادرات.", "bi-check2-square", "portal.hr_approvals")
    add_card(HR_MASTERDATA_MANAGE, "إعدادات الدوام", "تعريف أنواع الإجازات/المغادرات والجداول.", "bi-gear", "portal.hr_masterdata_index")
    add_card(HR_MASTERDATA_MANAGE, "الهيكل التنظيمي", "إدارة المنظمات/الإدارات/الدوائر/الأقسام/الفرق.", "bi-diagram-3", "portal.portal_admin_hr_org_structure")
    add_card(HR_MASTERDATA_MANAGE, "قوالب الدوام", "تعريف قوالب زمنية (Schedule Templates).", "bi-clock", "portal.portal_admin_hr_schedule_templates")
    add_card(HR_MASTERDATA_MANAGE, "سياسات الدوام", "سياسات ثابت/Hybrid quota + سياسة مكان.", "bi-shield-check", "portal.portal_admin_hr_work_policies")
    add_card(HR_MASTERDATA_MANAGE, "تعيين الدوام", "ربط (قالب + سياسة) بمستخدم/دور/قسم لفترة محددة.", "bi-person-check", "portal.portal_admin_hr_work_assignments")
    add_card(HR_PERF_MANAGE, "الأداء والتقييم", "نماذج ودورات وتكليفات 360.", "bi-clipboard2-check", "portal.portal_admin_hr_perf_dashboard")
    add_card(PORTAL_REPORTS_READ, "التقارير والمؤشرات", "لوحة مؤشرات وتقارير HR (تصدير).", "bi-bar-chart-line", "portal.portal_admin_hr_reports_dashboard")
    
    add_card(PORTAL_AUDIT_READ, "السجل الزمني", "Timeline موحّد لحركات البوابة.", "bi-clock-history", "portal.portal_admin_timeline")
    add_card(HR_ATT_READ, "استيراد ساعة الدوام", "رفع/مزامنة ملف ساعة الدوام.", "bi-upload", "portal.hr_attendance_import")

    # Counters (best-effort)
    pending = {"leaves": 0, "permissions": 0}
    try:
        if current_user.has_perm(HR_REQUESTS_VIEW_ALL) or current_user.has_perm(HR_REQUESTS_APPROVE):
            pending["leaves"] = _safe_count(LeaveRequest.query.filter(LeaveRequest.status == "PENDING"))
            pending["permissions"] = _safe_count(PermissionRequest.query.filter(PermissionRequest.status == "PENDING"))
    except Exception:
        pass

    return render_template("portal/admin/hr_dashboard.html", cards=cards, pending=pending)


@portal_bp.route("/admin/hr/schedule-templates", methods=["GET", "POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def portal_admin_hr_schedule_templates():
    """Schedule Templates (time templates)."""
    # Ensure new tables won't break older DBs
    _ensure_work_policy_tables()

    def _to_int(v, default=None):
        try:
            return int(v)
        except Exception:
            return default

    if request.method == "POST":
        op = (request.form.get("op") or "").strip().lower()

        if op in ("create", "save"):
            sid = _to_int(request.form.get("id"))
            obj = WorkSchedule.query.get(sid) if (op == "save" and sid) else WorkSchedule()
            obj.name = (request.form.get("name") or "").strip() or "بدون اسم"
            obj.kind = (request.form.get("kind") or "FIXED").strip().upper() or "FIXED"
            obj.start_time = (request.form.get("start_time") or "08:00").strip()
            obj.end_time = (request.form.get("end_time") or "15:00").strip()

            obj.grace_minutes = _to_int(request.form.get("grace_minutes"), 0)
            obj.break_minutes = _to_int(request.form.get("break_minutes"), 0)
            obj.required_minutes = _to_int(request.form.get("required_minutes"))
            obj.overtime_threshold_minutes = _to_int(request.form.get("overtime_threshold_minutes"))
            obj.is_active = bool(request.form.get("is_active"))
            obj.created_by_id = getattr(current_user, "id", None)

            try:
                db.session.add(obj)
                db.session.commit()
                flash("تم الحفظ.", "success")
            except Exception:
                db.session.rollback()
                flash("تعذر الحفظ (تحقق من القيم).", "danger")
            return redirect(url_for("portal.portal_admin_hr_schedule_templates"))

        if op == "toggle":
            sid = _to_int(request.form.get("id"))
            obj = WorkSchedule.query.get_or_404(sid)
            try:
                obj.is_active = not bool(obj.is_active)
                db.session.commit()
                flash("تم التحديث.", "success")
            except Exception:
                db.session.rollback()
                flash("تعذر التحديث.", "danger")
            return redirect(url_for("portal.portal_admin_hr_schedule_templates"))

        if op == "delete":
            sid = _to_int(request.form.get("id"))
            obj = WorkSchedule.query.get_or_404(sid)
            # Safety: do not hard-delete if referenced; fall back to deactivation.
            try:
                refs = 0
                try:
                    refs += WorkScheduleDay.query.filter_by(schedule_id=obj.id).count()
                except Exception:
                    pass
                try:
                    refs += WorkAssignment.query.filter_by(schedule_id=obj.id).count()
                except Exception:
                    pass

                if refs > 0:
                    obj.is_active = False
                    # If it was default, clear default
                    try:
                        if str(_setting_get("HR_DEFAULT_SCHEDULE_ID") or "").strip() == str(obj.id):
                            _setting_set("HR_DEFAULT_SCHEDULE_ID", "")
                    except Exception:
                        pass
                    db.session.commit()
                    flash("لا يمكن حذف قالب مرتبط بتعيينات/أيام. تم إيقافه بدلاً من ذلك.", "warning")
                    return redirect(url_for("portal.portal_admin_hr_schedule_templates"))

                # Not referenced: safe delete
                try:
                    if str(_setting_get("HR_DEFAULT_SCHEDULE_ID") or "").strip() == str(obj.id):
                        _setting_set("HR_DEFAULT_SCHEDULE_ID", "")
                except Exception:
                    pass
                db.session.delete(obj)
                db.session.commit()
                flash("تم الحذف.", "success")
            except Exception:
                db.session.rollback()
                flash("تعذر الحذف (قد يكون مرتبطاً ببيانات أخرى).", "danger")
            return redirect(url_for("portal.portal_admin_hr_schedule_templates"))

        if op == "set_default":
            sid = (request.form.get("id") or "").strip()
            if sid and not sid.isdigit():
                flash("قيمة غير صحيحة.", "danger")
                return redirect(url_for("portal.portal_admin_hr_schedule_templates"))
            _setting_set("HR_DEFAULT_SCHEDULE_ID", sid or "")
            try:
                db.session.commit()
            except Exception:
                db.session.rollback()
            flash("تم حفظ الجدول الافتراضي.", "success")
            return redirect(url_for("portal.portal_admin_hr_schedule_templates"))

        flash("عملية غير معروفة.", "warning")
        return redirect(url_for("portal.portal_admin_hr_schedule_templates"))

    schedules = WorkSchedule.query.order_by(WorkSchedule.id.desc()).all()
    default_schedule_id = _setting_get("HR_DEFAULT_SCHEDULE_ID")
    return render_template(
        "portal/admin/hr_schedule_templates.html",
        schedules=schedules,
        default_schedule_id=default_schedule_id,
    )


@portal_bp.route("/admin/hr/work-policies", methods=["GET", "POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def portal_admin_hr_work_policies():
    """Work Policies (days policy + location policy)."""
    _ensure_work_policy_tables()

    def _to_int(v, default=None):
        try:
            return int(v)
        except Exception:
            return default

    if request.method == "POST":
        op = (request.form.get("op") or "").strip().lower()

        if op in ("create", "save"):
            pid = _to_int(request.form.get("id"))
            obj = WorkPolicy.query.get(pid) if (op == "save" and pid) else WorkPolicy()
            obj.name = (request.form.get("name") or "").strip() or "بدون اسم"
            obj.days_policy = (request.form.get("days_policy") or "FIXED").strip().upper() or "FIXED"
            obj.location_policy = (request.form.get("location_policy") or "ONSITE").strip().upper() or "ONSITE"
            obj.is_active = bool(request.form.get("is_active"))
            obj.created_by_id = getattr(current_user, "id", None)

            # Days settings
            if obj.days_policy == "HYBRID_WEEKLY_QUOTA":
                obj.fixed_days_mask = None
                obj.hybrid_office_days = _to_int(request.form.get("hybrid_office_days"), 0)
                obj.hybrid_remote_days = _to_int(request.form.get("hybrid_remote_days"), 0)
            else:
                obj.hybrid_office_days = None
                obj.hybrid_remote_days = None
                mask = _days_mask_from_list(request.form.getlist("fixed_days"))
                obj.fixed_days_mask = mask or None

            try:
                db.session.add(obj)
                db.session.commit()
                flash("تم الحفظ.", "success")
            except Exception:
                db.session.rollback()
                flash("تعذر الحفظ (تحقق من القيم).", "danger")
            return redirect(url_for("portal.portal_admin_hr_work_policies"))

        if op == "toggle":
            pid = _to_int(request.form.get("id"))
            obj = WorkPolicy.query.get_or_404(pid)
            try:
                obj.is_active = not bool(obj.is_active)
                db.session.commit()
                flash("تم التحديث.", "success")
            except Exception:
                db.session.rollback()
                flash("تعذر التحديث.", "danger")
            return redirect(url_for("portal.portal_admin_hr_work_policies"))

        if op == "delete":
            pid = _to_int(request.form.get("id"))
            obj = WorkPolicy.query.get_or_404(pid)
            try:
                db.session.delete(obj)
                db.session.commit()
                flash("تم الحذف.", "success")
            except Exception:
                db.session.rollback()
                flash("تعذر الحذف (قد يكون مرتبطاً بتعيينات).", "danger")
            return redirect(url_for("portal.portal_admin_hr_work_policies"))

        flash("عملية غير معروفة.", "warning")
        return redirect(url_for("portal.portal_admin_hr_work_policies"))

    policies = WorkPolicy.query.order_by(WorkPolicy.id.desc()).all()
    return render_template(
        "portal/admin/hr_work_policies.html",
        policies=policies,
        weekdays=_WEEKDAYS_AR,
        days_label=_work_policy_days_label,
        place_label=_work_policy_place_label,
    )


@portal_bp.route("/admin/hr/work-assignments", methods=["GET", "POST"])
@login_required
@_perm(HR_MASTERDATA_MANAGE)
def portal_admin_hr_work_assignments():
    """Work Assignments (template + policy) to user/role/department with date range."""
    _ensure_work_policy_tables()

    def _to_int(v, default=None):
        try:
            return int(v)
        except Exception:
            return default

    if request.method == "POST":
        op = (request.form.get("op") or "").strip().lower()

        if op == "create":
            name = (request.form.get("name") or "").strip() or None
            schedule_id = _to_int(request.form.get("schedule_id"))
            policy_id = _to_int(request.form.get("policy_id"))
            target_type = (request.form.get("target_type") or "USER").strip().upper() or "USER"
            start_date = (request.form.get("start_date") or "").strip() or None
            end_date = (request.form.get("end_date") or "").strip() or None

            obj = WorkAssignment()
            obj.name = name
            obj.schedule_id = schedule_id
            obj.policy_id = policy_id
            obj.target_type = target_type
            obj.start_date = start_date
            obj.end_date = end_date
            obj.is_active = True
            obj.created_by_id = getattr(current_user, "id", None)

            # target
            obj.target_user_id = None
            obj.target_role = None
            obj.target_department_id = None
            if target_type == "ROLE":
                obj.target_role = (request.form.get("target_role") or "").strip() or None
            elif target_type == "DEPARTMENT":
                obj.target_department_id = _to_int(request.form.get("target_department_id"))
            else:
                obj.target_user_id = _to_int(request.form.get("target_user_id"))
                obj.target_type = "USER"

            # basic validation
            if not obj.schedule_id:
                flash("اختر قالب الدوام.", "warning")
                return redirect(url_for("portal.portal_admin_hr_work_assignments"))
            if obj.target_type == "ROLE" and not obj.target_role:
                flash("اختر الدور.", "warning")
                return redirect(url_for("portal.portal_admin_hr_work_assignments"))
            if obj.target_type == "DEPARTMENT" and not obj.target_department_id:
                flash("اختر الدائرة.", "warning")
                return redirect(url_for("portal.portal_admin_hr_work_assignments"))
            if obj.target_type == "USER" and not obj.target_user_id:
                flash("اختر المستخدم.", "warning")
                return redirect(url_for("portal.portal_admin_hr_work_assignments"))

            try:
                db.session.add(obj)
                db.session.commit()
                flash("تم إنشاء التعيين.", "success")
            except Exception:
                db.session.rollback()
                flash("تعذر إنشاء التعيين.", "danger")
            return redirect(url_for("portal.portal_admin_hr_work_assignments"))

        if op == "save":
            aid = _to_int(request.form.get("id"))
            obj = WorkAssignment.query.get_or_404(aid)

            obj.name = (request.form.get("name") or "").strip() or None
            obj.schedule_id = _to_int(request.form.get("schedule_id"))
            obj.policy_id = _to_int(request.form.get("policy_id"))
            obj.target_type = (request.form.get("target_type") or "USER").strip().upper() or "USER"
            obj.start_date = (request.form.get("start_date") or "").strip() or None
            obj.end_date = (request.form.get("end_date") or "").strip() or None
            obj.is_active = True if (request.form.get("is_active") == "1" or request.form.get("is_active") == "on") else False

            # Reset targets
            obj.target_user_id = None
            obj.target_role = None
            obj.target_department_id = None
            if obj.target_type == "ROLE":
                obj.target_role = (request.form.get("target_role") or "").strip() or None
            elif obj.target_type == "DEPARTMENT":
                obj.target_department_id = _to_int(request.form.get("target_department_id"))
            else:
                obj.target_user_id = _to_int(request.form.get("target_user_id"))
                obj.target_type = "USER"

            # basic validation
            if not obj.schedule_id:
                flash("اختر قالب الدوام.", "warning")
                return redirect(url_for("portal.portal_admin_hr_work_assignments"))
            if obj.target_type == "ROLE" and not obj.target_role:
                flash("اختر الدور.", "warning")
                return redirect(url_for("portal.portal_admin_hr_work_assignments"))
            if obj.target_type == "DEPARTMENT" and not obj.target_department_id:
                flash("اختر الدائرة.", "warning")
                return redirect(url_for("portal.portal_admin_hr_work_assignments"))
            if obj.target_type == "USER" and not obj.target_user_id:
                flash("اختر المستخدم.", "warning")
                return redirect(url_for("portal.portal_admin_hr_work_assignments"))

            try:
                db.session.commit()
                flash("تم الحفظ.", "success")
            except Exception:
                db.session.rollback()
                flash("تعذر الحفظ.", "danger")
            return redirect(url_for("portal.portal_admin_hr_work_assignments"))

        if op == "toggle":
            aid = _to_int(request.form.get("id"))
            obj = WorkAssignment.query.get_or_404(aid)
            try:
                obj.is_active = not bool(obj.is_active)
                db.session.commit()
                flash("تم التحديث.", "success")
            except Exception:
                db.session.rollback()
                flash("تعذر التحديث.", "danger")
            return redirect(url_for("portal.portal_admin_hr_work_assignments"))

        if op == "delete":
            aid = _to_int(request.form.get("id"))
            obj = WorkAssignment.query.get_or_404(aid)
            try:
                db.session.delete(obj)
                db.session.commit()
                flash("تم الحذف.", "success")
            except Exception:
                db.session.rollback()
                flash("تعذر الحذف.", "danger")
            return redirect(url_for("portal.portal_admin_hr_work_assignments"))

        flash("عملية غير معروفة.", "warning")
        return redirect(url_for("portal.portal_admin_hr_work_assignments"))

    schedules = WorkSchedule.query.order_by(WorkSchedule.name.asc().nullslast(), WorkSchedule.id.desc()).all()
    policies = WorkPolicy.query.order_by(WorkPolicy.id.desc()).all()
    assignments = WorkAssignment.query.order_by(WorkAssignment.id.desc()).all()
    # NOTE: User.full_name is a @property (Python-side), not a SQLAlchemy column.
    # لذلك لا يمكن استخدام .asc() عليه داخل ORDER BY.
    # نرتّب المستخدمين حسب (الاسم إن وُجد وإلا البريد)، مع تجاهل حالة الأحرف.
    name_expr = func.nullif(func.trim(User.name), "")
    users = (
        User.query
        .order_by(
            func.lower(func.coalesce(name_expr, User.email)).asc(),
            User.id.asc(),
        )
        .all()
    )
    departments = Department.query.order_by(Department.name_ar.asc().nullslast(), Department.code.asc().nullslast(), Department.id.asc()).all()
    # Roles (best effort)
    roles = []
    try:
        roles = sorted({(u.role or "").strip() for u in users if (u.role or "").strip()})
    except Exception:
        roles = []

    return render_template(
        "portal/admin/hr_work_assignments.html",
        schedules=schedules,
        policies=policies,
        assignments=assignments,
        users=users,
        roles=roles,
        departments=departments,
        days_label=_work_policy_days_label,
        place_label=_work_policy_place_label,
    )


@portal_bp.route("/admin/hr/org-structure", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_ADMIN_PERMISSIONS_MANAGE)
def portal_admin_hr_org_structure():
    """Full org-structure CRUD + Excel import/export (Portal Admin).

    Tabs:
      - orgs: Organizations
      - dirs: Directorates (belongs to Organization)
      - units: Units (belongs to Organization)
      - depts: Departments/Circles (belongs to either Directorate OR Unit - exactly one)
      - secs: Sections (belongs to either Department OR Directorate OR Unit - exactly one)
      - divs: Divisions (belongs to Section OR Department)
      - teams: Teams (belongs to Section)
    """

    tab = (request.args.get("tab") or "orgs").strip().lower()
    kind = (request.form.get("kind") or tab).strip().lower()
    op = (request.form.get("op") or "").strip().lower() if request.method == "POST" else ""

    def to_code(v):
        if v is None:
            return None
        s = str(v).strip()
        if not s:
            return None
        # normalize floats like 1.0
        try:
            if s.endswith('.0') and s.replace('.', '', 1).isdigit():
                s = str(int(float(s)))
        except Exception:
            pass
        return s

    def to_int(v):
        try:
            if v is None or str(v).strip() == "":
                return None
            return int(str(v).strip())
        except Exception:
            return None

    # -------------------- Export (multi-sheet) --------------------
    if request.method == "GET" and request.args.get("export"):
        from openpyxl import Workbook

        wb = Workbook()
        wb.remove(wb.active)

        orgs = Organization.query.order_by(Organization.code.asc().nullslast(), Organization.name_ar.asc()).all()
        directorates = Directorate.query.order_by(Directorate.code.asc().nullslast(), Directorate.name_ar.asc()).all()
        units = Unit.query.order_by(Unit.code.asc().nullslast(), Unit.name_ar.asc()).all()
        departments = Department.query.order_by(Department.code.asc().nullslast(), Department.name_ar.asc()).all()
        sections = Section.query.order_by(Section.code.asc().nullslast(), Section.name_ar.asc()).all()
        divisions = Division.query.order_by(Division.code.asc().nullslast(), Division.name_ar.asc()).all()
        teams = Team.query.order_by(Team.code.asc().nullslast(), Team.name_ar.asc()).all()

        def add_sheet(title, headers, rows):
            ws = wb.create_sheet(title=title)
            ws.append(headers)
            for r in rows:
                ws.append(r)

        add_sheet(
            "organizations",
            ["code", "name_ar", "name_en", "is_active"],
            [[o.code or "", o.name_ar or "", o.name_en or "", 1 if o.is_active else 0] for o in orgs],
        )

        add_sheet(
            "directorates",
            ["code", "organization_code", "name_ar", "name_en", "is_active"],
            [[d.code or "", (d.organization.code if d.organization else ""), d.name_ar or "", d.name_en or "", 1 if d.is_active else 0] for d in directorates],
        )

        add_sheet(
            "units",
            ["code", "organization_code", "name_ar", "name_en", "is_active"],
            [[u.code or "", (u.organization.code if getattr(u, "organization", None) else ""), u.name_ar or "", u.name_en or "", 1 if u.is_active else 0] for u in units],
        )

        add_sheet(
            "departments",
            ["code", "directorate_code", "unit_code", "name_ar", "name_en", "is_active"],
            [[
                d.code or "",
                (d.directorate.code if d.directorate else ""),
                (d.unit.code if getattr(d, 'unit', None) else ""),
                d.name_ar or "", d.name_en or "", 1 if d.is_active else 0,
            ] for d in departments],
        )

        add_sheet(
            "sections",
            ["code", "department_code", "directorate_code", "unit_code", "name_ar", "name_en", "is_active"],
            [[
                s.code or "",
                (s.department.code if s.department else ""),
                (s.directorate.code if s.directorate else ""),
                (s.unit.code if getattr(s, 'unit', None) else ""),
                s.name_ar or "", s.name_en or "", 1 if s.is_active else 0,
            ] for s in sections],
        )

        add_sheet(
            "divisions",
            ["code", "section_code", "department_code", "name_ar", "name_en", "is_active"],
            [[
                v.code or "",
                (v.section.code if getattr(v, "section", None) else ""),
                (v.department.code if getattr(v, "department", None) else ""),
                v.name_ar or "", v.name_en or "", 1 if v.is_active else 0,
            ] for v in divisions],
        )

        add_sheet(
            "teams",
            ["code", "section_code", "name_ar", "name_en", "is_active"],
            [[t.code or "", (t.section.code if t.section else ""), t.name_ar or "", t.name_en or "", 1 if t.is_active else 0] for t in teams],
        )

        from io import BytesIO
        bio = BytesIO()
        wb.save(bio)
        bio.seek(0)
        return send_file(bio, as_attachment=True, download_name="org_structure.xlsx", mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    # -------------------- POST actions --------------------
    if request.method == "POST":
        if kind not in ("orgs", "dirs", "units", "depts", "secs", "divs", "teams"):
            flash("تبويب غير معروف.", "warning")
            return redirect(url_for("portal.portal_admin_hr_org_structure", tab=tab))

        if op == "import":
            file = request.files.get("file")
            if not file:
                flash("اختر ملف Excel للاستيراد.", "warning")
                return redirect(url_for("portal.portal_admin_hr_org_structure", tab=kind))
            ok, msg = _import_org_structure_excel(kind, file)
            flash(msg, "success" if ok else "danger")
            return redirect(url_for("portal.portal_admin_hr_org_structure", tab=kind))

        # Resolve model
        Model = {"orgs": Organization, "dirs": Directorate, "units": Unit, "depts": Department, "secs": Section, "divs": Division, "teams": Team}.get(kind)

        if op == "delete":
            rid = to_int(request.form.get("id"))
            if not rid:
                flash("معرف غير صالح.", "warning")
                return redirect(url_for("portal.portal_admin_hr_org_structure", tab=kind))
            row = Model.query.get(rid)
            if not row:
                flash("غير موجود.", "warning")
                return redirect(url_for("portal.portal_admin_hr_org_structure", tab=kind))

            # guards
            if kind == "orgs" and (Directorate.query.filter_by(organization_id=row.id).first() or Unit.query.filter_by(organization_id=row.id).first()):
                flash("لا يمكن حذف منظمة مرتبطة بإدارات/وحدات.", "warning")
                return redirect(url_for("portal.portal_admin_hr_org_structure", tab=kind))
            if kind == "dirs" and (Department.query.filter_by(directorate_id=row.id).first() or Section.query.filter_by(directorate_id=row.id).first()):
                flash("لا يمكن حذف إدارة مرتبطة بدوائر/أقسام.", "warning")
                return redirect(url_for("portal.portal_admin_hr_org_structure", tab=kind))
            if kind == "units" and (Department.query.filter_by(unit_id=row.id).first() or Section.query.filter_by(unit_id=row.id).first()):
                flash("لا يمكن حذف وحدة مرتبطة بدوائر/أقسام.", "warning")
                return redirect(url_for("portal.portal_admin_hr_org_structure", tab=kind))
            if kind == "depts" and (Section.query.filter_by(department_id=row.id).first() or Division.query.filter_by(department_id=row.id).first()):
                flash("لا يمكن حذف دائرة مرتبطة بأقسام.", "warning")
                return redirect(url_for("portal.portal_admin_hr_org_structure", tab=kind))
            if kind == "secs" and (Team.query.filter_by(section_id=row.id).first() or Division.query.filter_by(section_id=row.id).first()):
                flash("لا يمكن حذف قسم مرتبط بفرق.", "warning")
                return redirect(url_for("portal.portal_admin_hr_org_structure", tab=kind))
            if kind == "divs" and Team.query.filter_by(division_id=row.id).first():
                flash("لا يمكن حذف شعبة مرتبطة بفرق.", "warning")
                return redirect(url_for("portal.portal_admin_hr_org_structure", tab=kind))

            try:
                db.session.delete(row)
                db.session.commit()
                flash("تم الحذف.", "success")
            except Exception:
                db.session.rollback()
                flash("تعذر الحذف (قد يكون مرتبطاً ببيانات أخرى).", "danger")
            return redirect(url_for("portal.portal_admin_hr_org_structure", tab=kind))

        if op == "save":
            rid = to_int(request.form.get("id"))
            row = Model.query.get(rid) if rid else Model()

            # Parent handling
            if kind == "dirs":
                pid = to_int(request.form.get("parent_id"))
                row.organization_id = pid
            elif kind == "units":
                pid = to_int(request.form.get("parent_id"))
                row.organization_id = pid
            elif kind == "teams":
                pid = to_int(request.form.get("parent_id"))
                row.section_id = pid
            elif kind == "divs":
                ptype = (request.form.get("parent_type") or "section").strip().lower()
                pid_sec = to_int(request.form.get("parent_id_sec"))
                pid_dept = to_int(request.form.get("parent_id_dept"))
                row.section_id = None
                row.department_id = None
                if ptype == "department":
                    row.department_id = pid_dept
                else:
                    row.section_id = pid_sec
            elif kind == "depts":
                ptype = (request.form.get("parent_type") or "directorate").strip().lower()
                pid_dir = to_int(request.form.get("parent_id_dir"))
                pid_unit = to_int(request.form.get("parent_id_unit"))
                if ptype == "unit":
                    row.unit_id = pid_unit
                    row.directorate_id = None
                else:
                    row.directorate_id = pid_dir
                    row.unit_id = None
            elif kind == "secs":
                ptype = (request.form.get("parent_type") or "department").strip().lower()
                pid_dept = to_int(request.form.get("parent_id_dept"))
                pid_dir = to_int(request.form.get("parent_id_dir"))
                pid_unit = to_int(request.form.get("parent_id_unit"))
                row.department_id = None
                row.directorate_id = None
                row.unit_id = None
                if ptype == "directorate":
                    row.directorate_id = pid_dir
                elif ptype == "unit":
                    row.unit_id = pid_unit
                else:
                    row.department_id = pid_dept

            # Common fields
            row.code = to_code(request.form.get("code"))
            row.name_ar = (request.form.get("name_ar") or "").strip() or None
            row.name_en = (request.form.get("name_en") or "").strip() or None
            row.is_active = bool(request.form.get("is_active"))

            # validations (best-effort)
            if kind in ("dirs", "units", "teams"):
                if not getattr(row, {"dirs": "organization_id", "units": "organization_id", "teams": "section_id"}[kind]):
                    flash("اختر التبعية (Parent).", "warning")
                    return redirect(url_for("portal.portal_admin_hr_org_structure", tab=kind))
            if kind == "depts":
                if not (row.directorate_id or row.unit_id) or (row.directorate_id and row.unit_id):
                    flash("يجب اختيار تبعية واحدة فقط: إدارة أو وحدة.", "warning")
                    return redirect(url_for("portal.portal_admin_hr_org_structure", tab=kind))
            if kind == "secs":
                parents = [row.department_id, row.directorate_id, row.unit_id]
                if sum(1 for p in parents if p) != 1:
                    flash("يجب اختيار تبعية واحدة فقط: دائرة أو إدارة أو وحدة.", "warning")
                    return redirect(url_for("portal.portal_admin_hr_org_structure", tab=kind))

            if kind == "divs":
                parents = [row.section_id, row.department_id]
                if sum(1 for p in parents if p) != 1:
                    flash("يجب اختيار تبعية واحدة فقط: قسم أو دائرة.", "warning")
                    return redirect(url_for("portal.portal_admin_hr_org_structure", tab=kind))

            try:
                db.session.add(row)
                db.session.commit()
                flash("تم الحفظ.", "success")
            except Exception:
                db.session.rollback()
                flash("تعذر الحفظ (تحقق من القيم/التكرار).", "danger")

            return redirect(url_for("portal.portal_admin_hr_org_structure", tab=kind))

        flash("عملية غير معروفة.", "warning")
        return redirect(url_for("portal.portal_admin_hr_org_structure", tab=kind))

    # -------------------- GET --------------------
    orgs = Organization.query.order_by(Organization.code.asc().nullslast(), Organization.name_ar.asc()).all()
    directorates = Directorate.query.order_by(Directorate.code.asc().nullslast(), Directorate.name_ar.asc()).all()
    units = Unit.query.order_by(Unit.code.asc().nullslast(), Unit.name_ar.asc()).all()
    departments = Department.query.order_by(Department.code.asc().nullslast(), Department.name_ar.asc()).all()
    sections = Section.query.order_by(Section.code.asc().nullslast(), Section.name_ar.asc()).all()
    divisions = Division.query.order_by(Division.code.asc().nullslast(), Division.name_ar.asc()).all()
    teams = Team.query.order_by(Team.code.asc().nullslast(), Team.name_ar.asc()).all()

    return render_template(
        "portal/admin/hr_org_structure.html",
        tab=tab,
        orgs=orgs,
        directorates=directorates,
        units=units,
        departments=departments,
        sections=sections,
        divisions=divisions,
        teams=teams,
    )

def _import_org_structure_excel(kind: str, file_storage):
    """Import org structure from an Excel sheet.

    Notes:
    - Accepts both single-sheet and the exported multi-sheet workbook.
    - Tries to locate the proper sheet by name based on `kind`.
    - Detects header row automatically (in case the first row is a title).
    - Upserts by code (with a fuzzy lookup for numeric codes with leading zeros).
    """
    kind = (kind or "").strip().lower()
    if kind not in ("orgs", "dirs", "units", "depts", "secs", "divs", "teams"):
        return False, "حدد نوع البيانات للاستيراد."

    try:
        from openpyxl import load_workbook
        from sqlalchemy import func

        # FileStorage -> stream
        stream = getattr(file_storage, "stream", None) or file_storage
        try:
            stream.seek(0)
        except Exception:
            pass

        wb = load_workbook(stream, data_only=True)

        # Prefer the matching sheet if workbook has multiple sheets (export creates multi-sheets)
        sheet_candidates = {
            "orgs": ["organizations", "orgs", "organizations_sheet", "منظمات", "المنظمات"],
            "dirs": ["directorates", "dirs", "الإدارات", "الادارات"],
            "units": ["units", "unit", "الوحدات", "وحدات"],
            "depts": ["departments", "depts", "الدوائر", "دوائر", "الأقسام", "اقسام"],
            "secs": ["sections", "secs", "الأقسام", "الاقسام"],
            "divs": ["divisions", "divs", "الشعب", "شعب"],
            "teams": ["teams", "الفرق", "فرق"],
        }
        ws = wb.active
        if len(wb.sheetnames) > 1:
            sn_lower = {s.lower(): s for s in wb.sheetnames}
            for cand in sheet_candidates.get(kind, []):
                real = sn_lower.get(str(cand).strip().lower())
                if real:
                    ws = wb[real]
                    break

        # -------------------------
        # Header detection + aliases
        # -------------------------
        def _norm_hdr(v):
            if v is None:
                return ""
            s = str(v).strip().lower()
            # remove bidi/formatting marks that can break matching
            for ch in ("‏", "‎", "‪", "‫", "‬", "﻿"):
                s = s.replace(ch, "")
            s = re.sub(r"\s+", "_", s)
            return s

        aliases = {
            "code": {"code", "id", "unit_code", "dept_code", "section_code", "division_code", "كود", "الكود", "رمز", "الرمز", "رقم"},
            "organization_code": {"organization_code", "org_code", "organization", "org", "كود_المنظمة", "رمز_المنظمة", "المنظمة"},
            "directorate_code": {"directorate_code", "dir_code", "directorate", "كود_الادارة", "كود_الإدارة", "الادارة", "الإدارة"},
            "unit_code": {"unit_code", "unit", "الوحدة", "كود_الوحدة"},
            "department_code": {"department_code", "department", "المديرية", "الدائرة", "كود_الدائرة", "كود_القسم"},
            "section_code": {"section_code", "section", "القسم", "كود_القسم"},
            "name_ar": {"name_ar", "arabic_name", "name_arabic", "اسم", "الاسم", "الاسم_ع", "name"},
            "name_en": {"name_en", "english_name", "name_english", "name_eng", "الاسم_بالانجليزية", "الاسم_بالإنجليزية"},
            "is_active": {"is_active", "active", "enabled", "فعال", "نشط"},
        }

        def _canon(h: str) -> str:
            h = _norm_hdr(h)
            for k, opts in aliases.items():
                if h in opts:
                    return k
            return h

        header_row = None
        scan_max = min(10, ws.max_row or 1)
        for r in range(1, scan_max + 1):
            vals = [_canon(c.value) for c in ws[r]]
            has_code = any(v == "code" for v in vals)
            has_name = any(v in ("name_ar", "name_en", "name") for v in vals)
            if has_code and has_name:
                header_row = r
                break
        if header_row is None:
            header_row = 1

        headers = [_canon(c.value) for c in ws[header_row]]
        hidx = {h: i for i, h in enumerate(headers) if h}

        if "code" not in hidx:
            return False, "ملف Excel غير مطابق: لم يتم العثور على عمود (code). حمّل ملف التصدير واستخدم نفس الأعمدة."

        def cell(row, key):
            i = hidx.get(key)
            if i is None:
                return None
            try:
                return row[i].value
            except Exception:
                return None

        def to_code(v):
            if v is None:
                return None
            if isinstance(v, (int,)):
                return str(v).strip()
            s = str(v).strip()
            if s == "":
                return None
            # normalize numeric-like values (e.g. 1.0)
            try:
                if isinstance(v, float) and v.is_integer():
                    return str(int(v))
                if s.endswith(".0") and s.replace(".", "", 1).isdigit():
                    return str(int(float(s)))
            except Exception:
                pass
            return s

        def to_bool(v):
            if v is None:
                return True
            s = str(v).strip().lower()
            if s in ("0", "false", "no", "n", "غير", "لا"):
                return False
            return True

        def find_by_code(Model, code_val):
            """Exact match first, then numeric match ignoring leading zeros."""
            c = to_code(code_val)
            if not c:
                return None
            obj = Model.query.filter_by(code=c).first()
            if obj:
                return obj
            if c.isdigit():
                stripped = c.lstrip("0") or "0"
                try:
                    obj = Model.query.filter(func.ltrim(Model.code, "0") == stripped).first()
                    return obj
                except Exception:
                    return None
            return None

        inserted = 0
        updated = 0
        skipped = 0
        errors = []
        seen_codes = set()

        for ridx, row in enumerate(ws.iter_rows(min_row=header_row + 1), start=header_row + 1):
            code = to_code(cell(row, "code"))
            if not code:
                skipped += 1
                continue

            if code in seen_codes:
                skipped += 1
                errors.append(f"row {ridx}: duplicate code={code}")
                continue
            seen_codes.add(code)

            name_ar = (cell(row, "name_ar") or "")
            name_en = (cell(row, "name_en") or "")
            is_active = to_bool(cell(row, "is_active"))

            # parent codes (optional by kind)
            org_code = to_code(cell(row, "organization_code"))
            dir_code = to_code(cell(row, "directorate_code"))
            unit_code = to_code(cell(row, "unit_code"))
            dept_code = to_code(cell(row, "department_code"))
            sec_code = to_code(cell(row, "section_code"))

            try:
                if kind == "orgs":
                    obj = find_by_code(Organization, code) or Organization(code=code)
                    is_new = obj.id is None
                    obj.name_ar = str(name_ar).strip() or None
                    obj.name_en = str(name_en).strip() or None
                    obj.is_active = bool(is_active)
                    db.session.add(obj)
                    inserted += int(is_new)
                    updated += int(not is_new)

                elif kind == "dirs":
                    org = find_by_code(Organization, org_code) if org_code else None
                    if not org:
                        skipped += 1
                        errors.append(f"row {ridx}: directorate code={code}: organization_code not found ({org_code})")
                        continue
                    obj = find_by_code(Directorate, code) or Directorate(code=code)
                    is_new = obj.id is None
                    obj.organization_id = org.id
                    obj.name_ar = str(name_ar).strip() or None
                    obj.name_en = str(name_en).strip() or None
                    obj.is_active = bool(is_active)
                    db.session.add(obj)
                    inserted += int(is_new)
                    updated += int(not is_new)

                elif kind == "units":
                    org = find_by_code(Organization, org_code) if org_code else None
                    # Backward compatibility: some sheets used directorate_code for units
                    if not org and dir_code:
                        di = find_by_code(Directorate, dir_code)
                        org = di.organization if di else None
                    if not org:
                        skipped += 1
                        errors.append(f"row {ridx}: unit code={code}: organization_code not found ({org_code})")
                        continue

                    obj = find_by_code(Unit, code) or Unit(code=code)
                    is_new = obj.id is None
                    obj.organization_id = org.id
                    obj.name_ar = str(name_ar).strip() or None
                    obj.name_en = str(name_en).strip() or None
                    obj.is_active = bool(is_active)
                    db.session.add(obj)
                    inserted += int(is_new)
                    updated += int(not is_new)

                elif kind == "depts":
                    parent_dir = find_by_code(Directorate, dir_code) if dir_code else None
                    parent_unit = find_by_code(Unit, unit_code) if unit_code else None
                    if bool(parent_dir) == bool(parent_unit):
                        skipped += 1
                        errors.append(f"row {ridx}: department code={code}: provide directorate_code OR unit_code")
                        continue
                    obj = find_by_code(Department, code) or Department(code=code)
                    is_new = obj.id is None
                    if parent_unit:
                        obj.unit_id = parent_unit.id
                        obj.directorate_id = None
                    else:
                        obj.directorate_id = parent_dir.id
                        obj.unit_id = None
                    obj.name_ar = str(name_ar).strip() or None
                    obj.name_en = str(name_en).strip() or None
                    obj.is_active = bool(is_active)
                    db.session.add(obj)
                    inserted += int(is_new)
                    updated += int(not is_new)

                elif kind == "secs":
                    parent_dept = find_by_code(Department, dept_code) if dept_code else None
                    parent_dir = find_by_code(Directorate, dir_code) if dir_code else None
                    parent_unit = find_by_code(Unit, unit_code) if unit_code else None
                    parents = [p for p in (parent_dept, parent_dir, parent_unit) if p]
                    if len(parents) != 1:
                        skipped += 1
                        errors.append(f"row {ridx}: section code={code}: provide exactly one of department_code/directorate_code/unit_code")
                        continue
                    obj = find_by_code(Section, code) or Section(code=code)
                    is_new = obj.id is None
                    if parent_dept:
                        obj.department_id = parent_dept.id
                        obj.directorate_id = None
                        obj.unit_id = None
                    elif parent_unit:
                        obj.unit_id = parent_unit.id
                        obj.department_id = None
                        obj.directorate_id = None
                    else:
                        obj.directorate_id = parent_dir.id
                        obj.department_id = None
                        obj.unit_id = None
                    obj.name_ar = str(name_ar).strip() or None
                    obj.name_en = str(name_en).strip() or None
                    obj.is_active = bool(is_active)
                    db.session.add(obj)
                    inserted += int(is_new)
                    updated += int(not is_new)

                elif kind == "divs":
                    sec_code_given = bool(sec_code)
                    dept_code_given = bool(dept_code)
                    parent_sec = find_by_code(Section, sec_code) if sec_code_given else None
                    parent_dept = find_by_code(Department, dept_code) if dept_code_given else None

                    if sec_code_given and dept_code_given:
                        skipped += 1
                        errors.append(f"row {ridx}: division code={code}: provide section_code OR department_code (not both)")
                        continue
                    if not sec_code_given and not dept_code_given:
                        skipped += 1
                        errors.append(f"row {ridx}: division code={code}: missing section_code/department_code")
                        continue
                    if sec_code_given and not parent_sec:
                        skipped += 1
                        errors.append(f"row {ridx}: division code={code}: section_code not found ({sec_code})")
                        continue
                    if dept_code_given and not parent_dept:
                        skipped += 1
                        errors.append(f"row {ridx}: division code={code}: department_code not found ({dept_code})")
                        continue

                    obj = find_by_code(Division, code) or Division(code=code)
                    is_new = obj.id is None
                    if parent_sec:
                        obj.section_id = parent_sec.id
                        obj.department_id = None
                    else:
                        obj.department_id = parent_dept.id
                        obj.section_id = None
                    obj.name_ar = str(name_ar).strip() or None
                    obj.name_en = str(name_en).strip() or None
                    obj.is_active = bool(is_active)
                    db.session.add(obj)
                    inserted += int(is_new)
                    updated += int(not is_new)

                elif kind == "teams":
                    sec = find_by_code(Section, sec_code) if sec_code else None
                    if not sec:
                        skipped += 1
                        errors.append(f"row {ridx}: team code={code}: section_code not found ({sec_code})")
                        continue
                    obj = find_by_code(Team, code) or Team(code=code)
                    is_new = obj.id is None
                    obj.section_id = sec.id
                    obj.name_ar = str(name_ar).strip() or None
                    obj.name_en = str(name_en).strip() or None
                    obj.is_active = bool(is_active)
                    db.session.add(obj)
                    inserted += int(is_new)
                    updated += int(not is_new)

            except Exception as e:
                skipped += 1
                errors.append(f"row {ridx} code={code}: {e}")

        db.session.commit()

        msg = f"تم الاستيراد: {inserted} إضافة، {updated} تحديث، {skipped} تجاوز."
        if errors:
            msg += " (ملاحظات: " + "; ".join(errors[:6]) + (")" if len(errors) > 6 else ")")
        # If nothing imported/updated and many were skipped, treat as failure to hint wrong sheet/headers
        if inserted == 0 and updated == 0 and skipped > 0:
            return False, msg
        return True, msg

    except Exception:
        db.session.rollback()
        return False, "تعذر قراءة ملف Excel. تأكد أنه .xlsx وبالأعمدة المطلوبة."



@portal_bp.route("/admin/delegations", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_ADMIN_PERMISSIONS_MANAGE)
def portal_admin_delegations():
    # Create / deactivate delegations
    op = (request.form.get("op") or "").strip().lower() if request.method == "POST" else ""

    # Always load users list for dropdown
    users = User.query.order_by(User.name.asc().nullslast(), User.email.asc()).all()

    if request.method == "POST" and op == "create":
        # Accept both old and new form field names (backward compatible)
        from_user_id = (request.form.get("from_user_id") or request.form.get("delegator_user_id") or "").strip()
        to_user_id = (request.form.get("to_user_id") or request.form.get("deputy_user_id") or "").strip()

        starts_at_s = (request.form.get("starts_at") or "").strip()
        expires_at_s = (request.form.get("expires_at") or request.form.get("ends_at") or "").strip()
        note = (request.form.get("note") or "").strip() or None

        if not (from_user_id.isdigit() and to_user_id.isdigit()):
            flash("اختر المفوِّض والمفوَّض إليه.", "danger")
            return redirect(url_for("portal.portal_admin_delegations"))

        if int(from_user_id) == int(to_user_id):
            flash("لا يمكن التفويض لنفس الموظف.", "danger")
            return redirect(url_for("portal.portal_admin_delegations"))

        def _parse_dt(s: str):
            if not s:
                return None
            s = (s or "").strip().replace(" ", "T")
            try:
                return datetime.fromisoformat(s)
            except Exception:
                return None

        starts_at = _parse_dt(starts_at_s) or datetime.utcnow()
        expires_at = _parse_dt(expires_at_s)
        if not expires_at:
            # Default: 30 days
            expires_at = starts_at + timedelta(days=30)

        if expires_at <= starts_at:
            flash("تاريخ النهاية يجب أن يكون بعد البداية.", "danger")
            return redirect(url_for("portal.portal_admin_delegations"))

        existing = (
            Delegation.query
            .filter(
                Delegation.from_user_id == int(from_user_id),
                Delegation.to_user_id == int(to_user_id),
                Delegation.is_active == True,
                Delegation.expires_at >= starts_at,
            )
            .first()
        )
        if existing:
            flash("هناك تفويض نشط مشابه بالفعل.", "warning")
            return redirect(url_for("portal.portal_admin_delegations"))

        row = Delegation(
            from_user_id=int(from_user_id),
            to_user_id=int(to_user_id),
            starts_at=starts_at,
            expires_at=expires_at,
            is_active=True,
            created_by_id=current_user.id,
            created_at=datetime.utcnow(),
            note=note,
        )
        db.session.add(row)
        db.session.commit()
        flash("تم إنشاء التفويض.", "success")
        return redirect(url_for("portal.portal_admin_delegations"))

    if request.method == "POST" and op == "deactivate":
        did = (request.form.get("id") or "").strip()
        if did.isdigit():
            row = Delegation.query.get(int(did))
            if row:
                row.is_active = False
                db.session.commit()
                flash("تم إلغاء التفويض.", "success")
        return redirect(url_for("portal.portal_admin_delegations"))

    delegations = (
        Delegation.query
        .order_by(Delegation.is_active.desc(), Delegation.expires_at.desc(), Delegation.starts_at.desc())
        .all()
    )

    return render_template("portal/admin/delegations.html", delegations=delegations, users=users)


# -------------------------
# Portal Permissions Presets (اختصارات الصلاحيات)
# -------------------------

# Stored in DB table: portal_permission_preset (Text field, no 255-char limit)
PORTAL_PERM_PRESETS_KEY = "PORTAL_PERMISSION_PRESETS"


def _portal_perm_presets_defaults():
    """Default shortcuts shown on the portal permissions page."""
    base_employee_keys = [
        PORTAL_READ,
        HR_READ,
        HR_ATT_READ,
        HR_REQUESTS_READ,
        HR_REQUESTS_CREATE,
        HR_SS_READ,
        HR_SS_CREATE,
        HR_DOCS_READ,
    ]

    def _with_base(keys: list[str]) -> list[str]:
        return _normalize_keys((base_employee_keys or []) + (keys or []))

    presets_main = {
        "EMPLOYEE": {
            "label": "موظف",
            "keys": _with_base([]),
        },
        "MANAGER": {
            "label": "مدير",
            "keys": _with_base([HR_REQUESTS_APPROVE, HR_SS_APPROVE]),
        },
        "HR_ADMIN": {
            "label": "HR / Admin",
            "keys": _with_base([
                HR_ATT_CREATE, HR_ATT_EXPORT,
                HR_EMP_READ, HR_EMP_MANAGE, HR_EMP_ATTACH,
                HR_ORG_READ, HR_ORG_MANAGE,
                HR_MASTERDATA_MANAGE,
                HR_REQUESTS_VIEW_ALL, HR_REQUESTS_APPROVE,
                HR_SS_WORKFLOWS_MANAGE,
                HR_DOCS_MANAGE,
                PORTAL_ADMIN_READ, PORTAL_ADMIN_PERMISSIONS_MANAGE,
            ]),
        },
    }

    presets_extra = {
        "CLEAR": {
            "label": "مسح كل الصلاحيات",
            "keys": [],
        },
        "SYS_EVAL_VIEW": {
            "label": "التقييم النظامي (عرض)",
            "keys": _normalize_keys([PORTAL_READ, HR_SYSTEM_EVALUATION_VIEW]),
        },
        "CORR": {
            "label": "مراسلات",
            "keys": _with_base([CORR_READ, CORR_CREATE, CORR_UPDATE, CORR_EXPORT]),
        },
        "CORR_VIEW": {
            "label": "مراسلات (عرض فقط)",
            "keys": _with_base([CORR_READ]),
        },
        "ATTENDANCE_VIEW": {
            "label": "الدوام (عرض فقط)",
            "keys": _with_base([]),
        },
        "PAYSLIP_VIEW": {
            "label": "قسيمة الراتب (عرض فقط)",
            "keys": _with_base([HR_PAYSLIP_VIEW]),
        },
        "HR_SELF_SERVICE": {
            "label": "خدمات الموظف (طلبات/مستندات/قسيمة)",
            "keys": _with_base([HR_DOCS_VIEW, HR_PAYSLIP_VIEW, HR_SYSTEM_EVALUATION_VIEW]),
        },
        "STORE": {
            "label": "مستودع",
            "keys": _with_base([STORE_READ, STORE_MANAGE]),
        },
        "STORE_VIEW": {
            "label": "مستودع (عرض فقط)",
            "keys": _with_base([STORE_READ]),
        },
        "PORTAL_ADMIN": {
            "label": "مدير البوابة",
            "keys": _with_base([PORTAL_ADMIN_READ, PORTAL_ADMIN_PERMISSIONS_MANAGE]),
        },
        "ALL": {
            "label": "كل الصلاحيات (للاختبار)",
            "keys": list(PORTAL_ALL_KEYS),
        },
    }

    return presets_main, presets_extra


def _portal_perm_presets_normalize(presets_map: dict) -> dict:
    """Normalize/validate a presets dict from user input."""
    out = {}
    for code, obj in (presets_map or {}).items():
        c = (code or "").strip().upper()
        if not c or not isinstance(obj, dict):
            continue
        label = (obj.get("label") or c)
        try:
            label = str(label).strip() or c
        except Exception:
            label = c

        keys = obj.get("keys")
        if not isinstance(keys, list):
            keys = []
        cleaned = []
        for k in keys:
            kk = (str(k) if k is not None else "").strip().upper()
            if not kk:
                continue
            if kk not in PORTAL_ALL_KEYS:
                continue
            cleaned.append(kk)

        out[c] = {
            "label": label,
            "keys": _normalize_keys(cleaned),
        }
    return out


def _portal_perm_presets_load() -> tuple[dict, dict]:
    """Load presets from SystemSetting (if any), otherwise return defaults.

    Structure stored in SystemSetting.value (JSON):
    {
      "main": {"EMPLOYEE": {"label": "...", "keys": [...]}, ...},
      "extra": {...}
    }
    """
    defaults_main, defaults_extra = _portal_perm_presets_defaults()
    presets_main = dict(defaults_main)
    presets_extra = dict(defaults_extra)

    try:
        from models import PortalPermissionPreset

        rows = (
            PortalPermissionPreset.query
            .filter(PortalPermissionPreset.is_active.is_(True))
            .order_by(PortalPermissionPreset.category.asc(), PortalPermissionPreset.sort_order.asc(), PortalPermissionPreset.code.asc())
            .all()
        )
        if not rows:
            return presets_main, presets_extra

        main_override = {}
        extra_override = {}
        for r in rows:
            code = (getattr(r, "code", "") or "").strip().upper()
            if not code:
                continue
            label = (getattr(r, "label", None) or code)
            try:
                label = str(label).strip() or code
            except Exception:
                label = code

            keys = []
            try:
                keys = json.loads(getattr(r, "keys_json", "[]") or "[]")
                if not isinstance(keys, list):
                    keys = []
            except Exception:
                keys = []

            obj = _portal_perm_presets_normalize({code: {"label": label, "keys": keys}}).get(code)
            if not obj:
                continue

            cat = (getattr(r, "category", "") or "").strip().lower()
            if cat == "main":
                main_override[code] = obj
            else:
                extra_override[code] = obj

        if main_override:
            presets_main.update(main_override)
        if extra_override:
            presets_extra.update(extra_override)
        return presets_main, presets_extra
    except Exception:
        return presets_main, presets_extra


@portal_bp.route("/admin/permissions/presets", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_ADMIN_PERMISSIONS_MANAGE)
def portal_admin_permission_presets():
    """Manage the shortcuts (presets) shown on the portal permissions page."""
    if request.method == "POST":
        action = (request.form.get("action") or "save").strip().lower()

        from models import PortalPermissionPreset

        if action == "reset":
            try:
                PortalPermissionPreset.query.delete()
                db.session.commit()
            except Exception:
                db.session.rollback()
            flash("تمت إعادة ضبط اختصارات الصلاحيات للوضع الافتراضي.", "success")
            return redirect(url_for("portal.portal_admin_permission_presets"))

        raw = (request.form.get("presets_json") or "").strip()
        if not raw:
            flash("يرجى لصق JSON صحيح.", "warning")
            return redirect(url_for("portal.portal_admin_permission_presets"))

        try:
            cfg = json.loads(raw)
            if not isinstance(cfg, dict):
                raise ValueError("invalid")

            pm = cfg.get("main") or cfg.get("presets_main")
            pe = cfg.get("extra") or cfg.get("presets_extra")
            if isinstance(pm, dict) or isinstance(pe, dict):
                main = _portal_perm_presets_normalize(pm or {})
                extra = _portal_perm_presets_normalize(pe or {})
            else:
                # flat map
                flat = _portal_perm_presets_normalize(cfg)
                main_codes = {"EMPLOYEE", "MANAGER", "HR_ADMIN"}
                main = {k: v for k, v in flat.items() if k in main_codes}
                extra = {k: v for k, v in flat.items() if k not in main_codes}

            # Replace all presets in the table (simple + deterministic)
            try:
                PortalPermissionPreset.query.delete()
                db.session.flush()
            except Exception:
                db.session.rollback()
                PortalPermissionPreset.query.delete()

            order = 0
            for code, obj in (main or {}).items():
                order += 1
                db.session.add(PortalPermissionPreset(
                    code=code,
                    label=(obj.get("label") or code),
                    category="main",
                    keys_json=json.dumps(obj.get("keys") or [], ensure_ascii=False),
                    sort_order=order,
                    is_active=True,
                    updated_at=datetime.utcnow(),
                ))

            for code, obj in (extra or {}).items():
                order += 1
                db.session.add(PortalPermissionPreset(
                    code=code,
                    label=(obj.get("label") or code),
                    category="extra",
                    keys_json=json.dumps(obj.get("keys") or [], ensure_ascii=False),
                    sort_order=order,
                    is_active=True,
                    updated_at=datetime.utcnow(),
                ))

            db.session.commit()
            flash("تم حفظ اختصارات الصلاحيات.", "success")
            return redirect(url_for("portal.portal_admin_permission_presets"))
        except Exception:
            db.session.rollback()
            flash("JSON غير صالح أو يحتوي مفاتيح غير معروفة.", "danger")
            return redirect(url_for("portal.portal_admin_permission_presets"))

    # GET (show effective config: overrides + defaults)
    p_main, p_extra = _portal_perm_presets_load()
    raw = json.dumps({"main": p_main, "extra": p_extra}, ensure_ascii=False, indent=2)

    # Human-friendly labels for preview
    try:
        from portal.perm_defs import PERMS as PORTAL_PERMS
        perm_labels = {p.key: p.label for group in PORTAL_PERMS.values() for p in group}
        perm_descs = {p.key: p.desc for group in PORTAL_PERMS.values() for p in group}
    except Exception:
        perm_labels = {}
        perm_descs = {}

    return render_template(
        "portal/admin/permission_presets.html",
        presets_json=raw,
        presets_main=p_main,
        presets_extra=p_extra,
        perm_labels=perm_labels,
        perm_descs=perm_descs,
        setting_key=PORTAL_PERM_PRESETS_KEY,
    )


# -------------------------
# Permissions catalog (reference)
# -------------------------
_PORTAL_PERM_CATALOG_CACHE = {"ts": 0.0, "data": None}

def _portal_discover_system_permission_refs() -> dict:
    """Discover non-Portal permission keys used across the codebase.

    Returns: {PERM_KEY: [source1, source2, ...]}
    Cached for 10 minutes to keep the page fast.
    """
    import os
    import re
    import time

    global _PORTAL_PERM_CATALOG_CACHE
    now = time.time()
    try:
        if _PORTAL_PERM_CATALOG_CACHE.get("data") and (now - float(_PORTAL_PERM_CATALOG_CACHE.get("ts") or 0) < 600):
            return _PORTAL_PERM_CATALOG_CACHE["data"] or {}
    except Exception:
        pass

    keys_to_sources: dict[str, set[str]] = {}

    def add_key(k: str, src: str):
        kk = (k or "").strip().upper()
        if not kk:
            return
        keys_to_sources.setdefault(kk, set()).add(src)

    # 1) Legacy matrix (Workflow/roles)
    try:
        from permissions.matrix import PERMISSION_MATRIX
        for _, perms in (PERMISSION_MATRIX or {}).items():
            for p in (perms or []):
                add_key(str(p), "permissions/matrix.py")
    except Exception:
        pass

    # 2) Scan python sources for explicit checks
    try:
        root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
        exclude_dirs = {".git", "__pycache__", "venv", ".venv", "env", "node_modules", "static", "migrations", "uploads", "instance"}
        patterns = [
            (re.compile(r"has_role_perm\(\s*['\"]([A-Za-z0-9_]+)['\"]\s*\)"), "has_role_perm"),
            (re.compile(r"role_perm_required\(\s*['\"]([A-Za-z0-9_]+)['\"]\s*\)"), "role_perm_required"),
        ]

        for dirpath, dirnames, filenames in os.walk(root):
            dirnames[:] = [d for d in dirnames if (d not in exclude_dirs and not d.startswith("."))]
            for fn in filenames:
                if not fn.endswith(".py"):
                    continue
                fpath = os.path.join(dirpath, fn)
                try:
                    if os.path.getsize(fpath) > 2_000_000:  # avoid huge files
                        continue
                except Exception:
                    pass
                try:
                    with open(fpath, "r", encoding="utf-8", errors="ignore") as f:
                        txt = f.read()
                    rel = os.path.relpath(fpath, root).replace("\\", "/")
                    for rex, tag in patterns:
                        for m in rex.finditer(txt):
                            add_key(m.group(1), f"{rel}:{tag}")
                except Exception:
                    continue
    except Exception:
        pass

    out = {k: sorted(list(v))[:6] for k, v in keys_to_sources.items()}
    _PORTAL_PERM_CATALOG_CACHE = {"ts": now, "data": out}
    return out


@portal_bp.route("/admin/permissions/catalog", methods=["GET"])
@login_required
@_perm(PORTAL_ADMIN_PERMISSIONS_MANAGE)
def portal_admin_permissions_catalog():
    """Reference page: list all known permission keys (Portal + System)."""
    # Portal definitions (canonical)
    try:
        from portal.perm_defs import PERMS as PORTAL_PERMS
    except Exception:
        PORTAL_PERMS = {}

    perm_labels = {p.key: p.label for group in PORTAL_PERMS.values() for p in group}
    perm_descs = {p.key: p.desc for group in PORTAL_PERMS.values() for p in group}
    perm_modules = {p.key: p.module for group in PORTAL_PERMS.values() for p in group}

    portal_keys = set(perm_labels.keys())

    # System / legacy
    system_refs = _portal_discover_system_permission_refs()
    system_keys: dict[str, set[str]] = {k: set(v) for k, v in (system_refs or {}).items() if k and k not in portal_keys}

    # Also include any keys already saved in DB (RolePermission/UserPermission), even if not found in scan
    try:
        from models import RolePermission, UserPermission
        from extensions import db as _db
        for (p,) in (_db.session.query(RolePermission.permission).distinct().all() or []):
            kk = (str(p) if p is not None else "").strip().upper()
            if kk and kk not in portal_keys:
                system_keys.setdefault(kk, set()).add("database:RolePermission")
        for (p,) in (_db.session.query(UserPermission.key).distinct().all() or []):
            kk = (str(p) if p is not None else "").strip().upper()
            if kk and kk not in portal_keys:
                system_keys.setdefault(kk, set()).add("database:UserPermission")
    except Exception:
        pass

    system_items = [
        {"key": k, "sources": sorted(list(srcs))[:6]}
        for k, srcs in sorted(system_keys.items(), key=lambda kv: kv[0])
    ]

    return render_template(
        "portal/admin/permissions_catalog.html",
        portal_groups=PORTAL_PERMS,
        system_items=system_items,
        perm_labels=perm_labels,
        perm_descs=perm_descs,
        perm_modules=perm_modules,
    )


@portal_bp.route("/admin/permissions", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_ADMIN_PERMISSIONS_MANAGE)
def portal_admin_permissions():
    """Manage Portal permissions by ROLE or by USER.

    Important:
    - This page modifies ONLY Portal-related permissions (keys in PORTAL_ALL_KEYS).
    - ROLE scope writes to RolePermission.
    - USER scope writes to UserPermission (additive to role permissions).
    """
    from models import Role, RolePermission, User, UserPermission
    from sqlalchemy import func
    from extensions import db

    scope = (request.args.get("scope") or request.form.get("scope") or "role").strip().lower()
    if scope not in ("role", "user"):
        scope = "role"

    # -------------------------
    # Roles list (masterdata roles or fallback)
    # -------------------------
    roles = Role.query.filter_by(is_active=True).order_by(Role.code.asc()).all()
    role_codes = [r.code for r in roles]
    if not role_codes:
        role_codes = [
            r for (r,) in db.session.query(User.role).distinct().order_by(User.role.asc()).all()
            if (r or "").strip()
        ]


    # Meta for role dropdown (names + user counts) to reduce confusion when applying permissions.
    role_items = []
    try:
        _counts_rows = (
            db.session.query(func.lower(User.role), func.count(User.id))
            .group_by(func.lower(User.role))
            .all()
        )
        role_user_counts = {
            ((r or "").strip().lower()): int(c or 0)
            for (r, c) in (_counts_rows or [])
            if (r or "").strip()
        }
    except Exception:
        role_user_counts = {}

    if roles:
        for _r in roles:
            _code = (_r.code or "").strip()
            if not _code:
                continue
            role_items.append({
                "code": _code,
                "name_ar": (getattr(_r, "name_ar", "") or "").strip(),
                "name_en": (getattr(_r, "name_en", "") or "").strip(),
                "count": role_user_counts.get(_code.lower(), 0),
            })
    else:
        for _code in role_codes:
            _c = (_code or "").strip()
            if not _c:
                continue
            role_items.append({
                "code": _c,
                "name_ar": "",
                "name_en": "",
                "count": role_user_counts.get(_c.lower(), 0),
            })

    # -------------------------
    # Users list (only when scope=user)
    # -------------------------
    q = (request.args.get("q") or "").strip()
    users = []
    selected_user = None
    selected_user_id = (request.args.get("user_id") or "").strip()

    if scope == "user":
        uqry = User.query
        if q:
            # Global-ish search for users (excluding sensitive hashes).
            uqry = apply_search_all_columns(uqry, User, q, exclude_columns={"password_hash"})
        uqry = uqry.order_by(User.name.asc(), User.email.asc())
        users = uqry.limit(60).all()

        if selected_user_id:
            try:
                selected_user = User.query.get(int(selected_user_id))
                if selected_user and (selected_user not in users):
                    users = [selected_user] + users
            except Exception:
                selected_user = None

    # -------------------------
    # Excel export (all Portal permissions in one file)
    # -------------------------
    export = (request.args.get("export") or "").strip().lower()
    if request.method == "GET" and export in ("1", "true", "yes", "excel", "xlsx"):
        from utils.excel import make_xlsx_bytes_multi

        # Role permissions (Portal only)
        rperms = (RolePermission.query
                 .filter(RolePermission.permission.in_(PORTAL_ALL_KEYS))
                 .order_by(RolePermission.role.asc(), RolePermission.permission.asc())
                 .all())
        role_rows = [[rp.role, rp.permission] for rp in rperms]

        # User permissions (Portal only)
        q2 = (db.session.query(UserPermission, User.email)
              .join(User, User.id == UserPermission.user_id)
              .filter(UserPermission.key.in_(PORTAL_ALL_KEYS))
              .order_by(func.lower(User.email).asc(), UserPermission.key.asc()))
        user_rows = []
        for up, email in q2.all():
            user_rows.append([
                email or "",
                up.key,
                "Yes" if bool(up.is_allowed) else "No",
            ])

        xbytes = make_xlsx_bytes_multi([
            ("RolePermissions", ["role", "permission"], role_rows),
            ("UserPermissions", ["user_email", "permission", "is_allowed"], user_rows),
        ])
        return send_file(
            io.BytesIO(xbytes),
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name="portal_permissions.xlsx",
        )

    # -------------------------
    # UX presets (server-side)
    # -------------------------
    # NOTE: Built BEFORE POST so preset buttons can submit without relying on JS.
    # Admins can override these via: /portal/admin/permissions/presets
    presets_main, presets_extra = _portal_perm_presets_load()
    presets = {**(presets_main or {}), **(presets_extra or {})}

    # -------------------------
    # Apply changes (POST)
    # -------------------------
    if request.method == "POST":
        scope = (request.form.get("scope") or "role").strip().lower()
        preset_code = (request.form.get("preset_code") or "").strip()
        if preset_code and preset_code in presets:
            perms = [(p or "").strip().upper() for p in (presets[preset_code].get("keys") or []) if (p or "").strip()]
        else:
            perms = [(p or "").strip().upper() for p in request.form.getlist("permissions") if (p or "").strip()]

        if scope == "role":
            selected_role = (request.form.get("role") or "").strip()
            if not selected_role:
                flash("اختر الدور.", "danger")
                return redirect(url_for("portal.portal_admin_permissions", scope="role"))

            # Delete only portal permissions (leave other permissions intact)
            RolePermission.query \
                .filter(func.lower(RolePermission.role) == selected_role.lower()) \
                .filter(RolePermission.permission.in_(PORTAL_ALL_KEYS)) \
                .delete(synchronize_session=False)

            for p in perms:
                if p in PORTAL_ALL_KEYS:
                    db.session.add(RolePermission(role=selected_role, permission=p))

            db.session.commit()

            # Verify saved count (helps diagnose "not saving" reports)
            try:
                saved_count = (
                    RolePermission.query
                    .filter(func.lower(RolePermission.role) == selected_role.lower())
                    .filter(RolePermission.permission.in_(PORTAL_ALL_KEYS))
                    .count()
                )
            except Exception:
                saved_count = None

            try:
                role_uc = role_user_counts.get(selected_role.lower(), None)
            except Exception:
                role_uc = None

            msg = f"تم تحديث صلاحيات البوابة للدور ({selected_role})"
            if saved_count is not None:
                msg += f" — محفوظ: {saved_count}"
            if role_uc is not None:
                msg += f" — عدد المستخدمين على الدور: {role_uc}"
            flash(msg, "success")
            if role_uc == 0:
                flash("تنبيه: هذا الدور لا يوجد عليه مستخدمون، لذلك لن تلاحظ تغييرًا عند تجربة أي حساب.", "warning")

            try:
                role_obj = Role.query.filter(func.lower(Role.code) == selected_role.lower()).first()
                db.session.add(AuditLog(user_id=current_user.id, action='PORTAL_PERMISSIONS_ROLE_UPDATE', note=f'Update portal perms for role {selected_role} ({len(perms)})', target_type='ROLE', target_id=(role_obj.id if role_obj else None), created_at=datetime.utcnow()))
                db.session.commit()
            except Exception:
                db.session.rollback()
            return redirect(url_for("portal.portal_admin_permissions", scope="role", role=selected_role))

        # scope == user
        selected_user_id = (request.form.get("user_id") or "").strip()
        if not selected_user_id:
            flash("اختر المستخدم.", "danger")
            return redirect(url_for("portal.portal_admin_permissions", scope="user", q=q))

        try:
            uid = int(selected_user_id)
        except Exception:
            flash("معرّف المستخدم غير صالح.", "danger")
            return redirect(url_for("portal.portal_admin_permissions", scope="user", q=q))

        user = User.query.get(uid)
        if not user:
            flash("المستخدم غير موجود.", "danger")
            return redirect(url_for("portal.portal_admin_permissions", scope="user", q=q))

        # Delete only portal permissions (leave other user permissions intact)
        UserPermission.query \
            .filter(UserPermission.user_id == uid) \
            .filter(UserPermission.key.in_(PORTAL_ALL_KEYS)) \
            .delete(synchronize_session=False)

        for p in perms:
            if p in PORTAL_ALL_KEYS:
                db.session.add(UserPermission(user_id=uid, key=p, is_allowed=True))

        db.session.commit()

        # Verify saved count
        try:
            saved_count = (
                UserPermission.query
                .filter(UserPermission.user_id == uid)
                .filter(UserPermission.key.in_(PORTAL_ALL_KEYS))
                .filter(UserPermission.is_allowed == True)  # noqa: E712
                .count()
            )
        except Exception:
            saved_count = None

        msg = "تم تحديث صلاحيات البوابة للمستخدم."
        if saved_count is not None:
            msg += f" — محفوظ: {saved_count}"
        flash(msg, "success")

        try:
            db.session.add(AuditLog(user_id=current_user.id, action='PORTAL_PERMISSIONS_USER_UPDATE', note=f'Update portal perms for user {user.email} ({len(perms)})', target_type='USER', target_id=uid, created_at=datetime.utcnow()))
            db.session.commit()
        except Exception:
            db.session.rollback()
        return redirect(url_for("portal.portal_admin_permissions", scope="user", user_id=uid))

    # -------------------------
    # Load current state (GET)
    # -------------------------
    selected_role = (request.args.get("role") or "").strip()
    checked = set()  # checked permissions for the current scope
    role_checked_for_user = set()  # for user scope: role-derived permissions (for badges)

    if scope == "role" and selected_role:
        rows = (
            RolePermission.query
            .filter(func.lower(RolePermission.role) == selected_role.lower())
            .filter(RolePermission.permission.in_(PORTAL_ALL_KEYS))
            .all()
        )
        checked = { (r.permission or "").strip().upper() for r in rows if r.permission }

    if scope == "user" and selected_user:
        # user explicit allows
        urows = (
            UserPermission.query
            .filter(UserPermission.user_id == selected_user.id)
            .filter(UserPermission.key.in_(PORTAL_ALL_KEYS))
            .filter(UserPermission.is_allowed == True)  # noqa: E712
            .all()
        )
        checked = { (r.key or "").strip().upper() for r in urows if r.key }

        # role permissions (badges only)
        role_norm = (selected_user.role or "").strip().lower()
        if role_norm:
            rrows = (
                RolePermission.query
                .filter(func.lower(RolePermission.role) == role_norm)
                .filter(RolePermission.permission.in_(PORTAL_ALL_KEYS))
                .all()
            )
            role_checked_for_user = { (r.permission or "").strip().upper() for r in rrows if r.permission }

    # (presets_main / presets_extra / presets) were built above (server-side),
    # so the quick shortcut buttons can submit without relying on JS.

    return render_template(
        "portal/admin/permissions.html",
        scope=scope,
        role_codes=role_codes,
        role_items=role_items,
        role_user_counts=role_user_counts,
        users=users,
        q=q,
        selected_role=selected_role,
        selected_user=selected_user,
        perm_groups=PORTAL_PERMS,
        checked=checked,
        role_checked_for_user=role_checked_for_user,
        presets=presets,
        presets_main=presets_main,
        presets_extra=presets_extra,
    )


# -------------------------
# Portal Admin: Correspondence lookups (Categories / Parties)
# -------------------------
@portal_bp.route("/admin/corr")
@login_required
@_perm(CORR_LOOKUPS_MANAGE)
def portal_admin_corr_index():
    # Seed basic defaults once (safe if already created)
    try:
        if CorrCategory.query.count() == 0:
            db.session.add(CorrCategory(code="GENERAL", name_ar="عام", name_en="General", is_active=True))
            db.session.commit()
    except Exception:
        db.session.rollback()
    return render_template("portal/admin/corr/index.html")


@portal_bp.route("/admin/corr/categories", methods=["GET", "POST"])
@login_required
@_perm(CORR_LOOKUPS_MANAGE)
def portal_admin_corr_categories():
    q = (request.args.get("q") or "").strip()
    qry = CorrCategory.query
    if q:
        qry = apply_search_all_columns(qry, CorrCategory, q)

    if request.method == "POST":
        code = (request.form.get("code") or "").strip().upper()
        name_ar = (request.form.get("name_ar") or "").strip()
        name_en = (request.form.get("name_en") or "").strip() or None
        is_active = (request.form.get("is_active") or "1") == "1"

        if not code or not name_ar:
            flash("الرمز والاسم العربي مطلوبان.", "danger")
            return redirect(url_for("portal.portal_admin_corr_categories", q=q))

        if CorrCategory.query.filter_by(code=code).first():
            flash("هذا الرمز موجود مسبقًا.", "danger")
            return redirect(url_for("portal.portal_admin_corr_categories", q=q))

        db.session.add(CorrCategory(
            code=code,
            name_ar=name_ar,
            name_en=name_en,
            is_active=is_active,
            created_by_id=current_user.id,
            created_at=datetime.utcnow(),
        ))
        db.session.commit()

        try:
            db.session.add(AuditLog(
                user_id=current_user.id,
                action="CORR_LOOKUP_ADD",
                note=f"category {code}",
                target_type="CORR_CATEGORY",
                target_id=None,
                created_at=datetime.utcnow(),
            ))
            db.session.commit()
        except Exception:
            db.session.rollback()

        flash("تم إضافة التصنيف.", "success")
        return redirect(url_for("portal.portal_admin_corr_categories"))

    rows = qry.order_by(CorrCategory.is_active.desc(), CorrCategory.code.asc()).all()

    # Excel export
    if (request.args.get("export") or "").strip() in ("1", "true", "yes", "excel"):
        data: list[dict] = []
        for r in rows:
            data.append({
                "id": r.id,
                "code": r.code,
                "name_ar": r.name_ar,
                "name_en": r.name_en or "",
                "is_active": "نعم" if r.is_active else "لا",
            })

        xbytes = _xlsx_from_dicts(
            data,
            columns=[
                ("id", "ID"),
                ("code", "الرمز"),
                ("name_ar", "الاسم (AR)"),
                ("name_en", "الاسم (EN)"),
                ("is_active", "نشط"),
            ],
            sheet_name="CorrCategories",
        )
        return send_file(
            io.BytesIO(xbytes),
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name="corr_categories.xlsx",
        )

    return render_template("portal/admin/corr/categories.html", rows=rows, q=q)


@portal_bp.route("/admin/corr/categories/<int:cat_id>/edit", methods=["POST"])
@login_required
@_perm(CORR_LOOKUPS_MANAGE)
def portal_admin_corr_category_edit(cat_id: int):
    row = CorrCategory.query.get_or_404(cat_id)
    row.code = (request.form.get("code") or row.code).strip().upper()
    row.name_ar = (request.form.get("name_ar") or row.name_ar).strip()
    row.name_en = (request.form.get("name_en") or "").strip() or None
    row.is_active = (request.form.get("is_active") or "1") == "1"
    db.session.commit()
    flash("تم تحديث التصنيف.", "success")
    return redirect(url_for("portal.portal_admin_corr_categories"))


@portal_bp.route("/admin/corr/categories/<int:cat_id>/delete", methods=["POST"])
@login_required
@_perm(CORR_LOOKUPS_MANAGE)
def portal_admin_corr_category_delete(cat_id: int):
    row = CorrCategory.query.get_or_404(cat_id)
    try:
        db.session.delete(row)
        db.session.commit()
        flash("تم حذف التصنيف.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر حذف التصنيف.", "danger")
    return redirect(url_for("portal.portal_admin_corr_categories"))


@portal_bp.route("/admin/corr/parties", methods=["GET", "POST"])
@login_required
@_perm(CORR_LOOKUPS_MANAGE)
def portal_admin_corr_parties():
    q = (request.args.get("q") or "").strip()
    kind = (request.args.get("kind") or "").strip().upper()
    qry = CorrParty.query
    if kind in ("SENDER", "RECIPIENT", "BOTH"):
        qry = qry.filter(CorrParty.kind == kind)
    if q:
        qry = apply_search_all_columns(qry, CorrParty, q)

    if request.method == "POST":
        kind_p = (request.form.get("kind") or "").strip().upper()
        name_ar = (request.form.get("name_ar") or "").strip()
        name_en = (request.form.get("name_en") or "").strip() or None
        is_active = (request.form.get("is_active") or "1") == "1"

        if kind_p not in ("SENDER", "RECIPIENT", "BOTH") or not name_ar:
            flash("النوع والاسم العربي مطلوبان.", "danger")
            return redirect(url_for("portal.portal_admin_corr_parties", q=q, kind=kind))

        db.session.add(CorrParty(
            kind=kind_p,
            name_ar=name_ar,
            name_en=name_en,
            is_active=is_active,
            created_by_id=current_user.id,
            created_at=datetime.utcnow(),
        ))
        db.session.commit()

        try:
            db.session.add(AuditLog(
                user_id=current_user.id,
                action="CORR_LOOKUP_ADD",
                note=f"party {kind_p} {name_ar}",
                target_type="CORR_PARTY",
                target_id=None,
                created_at=datetime.utcnow(),
            ))
            db.session.commit()
        except Exception:
            db.session.rollback()

        flash("تم إضافة الجهة.", "success")
        return redirect(url_for("portal.portal_admin_corr_parties"))

    rows = qry.order_by(CorrParty.is_active.desc(), CorrParty.kind.asc(), CorrParty.name_ar.asc()).all()

    # Excel export
    if (request.args.get("export") or "").strip() in ("1", "true", "yes", "excel"):
        data = []
        for r in rows:
            data.append({
                "id": r.id,
                "kind": r.kind,
                "name_ar": r.name_ar,
                "name_en": r.name_en or "",
                "is_active": "نعم" if r.is_active else "لا",
            })
        xbytes = _xlsx_from_dicts(
            data,
            columns=[
                ("id", "ID"),
                ("kind", "النوع"),
                ("name_ar", "الاسم (AR)"),
                ("name_en", "الاسم (EN)"),
                ("is_active", "نشط"),
            ],
            sheet_name="CorrParties",
        )
        return send_file(
            BytesIO(xbytes),
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name="corr_parties.xlsx",
        )

    return render_template("portal/admin/corr/parties.html", rows=rows, q=q, kind=kind)


@portal_bp.route("/admin/corr/parties/<int:party_id>/edit", methods=["POST"])
@login_required
@_perm(CORR_LOOKUPS_MANAGE)
def portal_admin_corr_party_edit(party_id: int):
    row = CorrParty.query.get_or_404(party_id)
    row.kind = (request.form.get("kind") or row.kind).strip().upper()
    row.name_ar = (request.form.get("name_ar") or row.name_ar).strip()
    row.name_en = (request.form.get("name_en") or "").strip() or None
    row.is_active = (request.form.get("is_active") or "1") == "1"
    db.session.commit()
    flash("تم تحديث الجهة.", "success")
    return redirect(url_for("portal.portal_admin_corr_parties"))


@portal_bp.route("/admin/corr/parties/<int:party_id>/delete", methods=["POST"])
@login_required
@_perm(CORR_LOOKUPS_MANAGE)
def portal_admin_corr_party_delete(party_id: int):
    row = CorrParty.query.get_or_404(party_id)
    try:
        db.session.delete(row)
        db.session.commit()
        flash("تم حذف الجهة.", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر حذف الجهة.", "danger")
    return redirect(url_for("portal.portal_admin_corr_parties"))



# -------------------------
# HR Self-Service Requests (Light Workflow)
# -------------------------

def _hr_ss_upload_dir(req_id: int) -> Path:
    base = Path(current_app.instance_path) / "uploads" / "hr_ss" / str(req_id)
    base.mkdir(parents=True, exist_ok=True)
    return base


def _hr_docs_upload_dir(doc_id: int) -> Path:
    base = Path(current_app.instance_path) / "uploads" / "hr_docs" / str(doc_id)
    base.mkdir(parents=True, exist_ok=True)
    return base


def _hr_discipline_upload_dir(case_id: int) -> Path:
    base = Path(current_app.instance_path) / "uploads" / "hr_discipline" / str(case_id)
    base.mkdir(parents=True, exist_ok=True)
    return base


def _ss_safe_ext(filename: str) -> str:
    return _clean_ext(filename)


def _ss_allowed_file(filename: str) -> bool:
    ext = _clean_suffix(filename)
    return ext in (ALLOWED_CORR_EXTS or set())


def _ss_payload_from_form(type_code: str, form) -> dict:
    payload = {"type_code": type_code}
    if type_code == "CERTIFICATE":
        payload["certificate_kind"] = (form.get("certificate_kind") or "").strip()
        payload["language"] = (form.get("language") or "AR").strip()
        payload["to_whom"] = (form.get("to_whom") or "").strip()
        payload["notes"] = (form.get("notes") or "").strip()
    elif type_code == "UPDATE_PROFILE":
        payload["new_phone"] = (form.get("new_phone") or "").strip()
        payload["new_address"] = (form.get("new_address") or "").strip()
        payload["new_bank_name"] = (form.get("new_bank_name") or "").strip()
        payload["new_bank_account"] = (form.get("new_bank_account") or "").strip()
        payload["notes"] = (form.get("notes") or "").strip()
    elif type_code == "UPLOAD_DOCUMENTS":
        payload["doc_kind"] = (form.get("doc_kind") or "").strip()
        payload["expiry_date"] = (form.get("expiry_date") or "").strip()
        payload["notes"] = (form.get("notes") or "").strip()
    else:
        payload["notes"] = (form.get("notes") or "").strip()
    return payload


def _ensure_hr_ss_defaults() -> None:
    """Create minimal default definitions/steps if none exist."""
    if HRSSWorkflowDefinition.query.count() > 0:
        return

    defs = [
        ("CERTIFICATE", "طلب شهادة", "Certificate Request"),
        ("UPDATE_PROFILE", "تحديث بيانات الموظف", "Update Employee Profile"),
        ("UPLOAD_DOCUMENTS", "رفع مستندات HR", "Upload HR Documents"),
    ]
    created = []
    for code, ar, en in defs:
        d = HRSSWorkflowDefinition(code=code, name_ar=ar, name_en=en, is_active=True, created_by_id=getattr(current_user, "id", None))
        db.session.add(d)
        created.append(d)
    db.session.flush()

    # Default: single step approval by role "HR"
    for d in created:
        s1 = HRSSWorkflowStepDefinition(definition_id=d.id, step_no=1, approver_role="HR", approver_user_id=None, sla_hours=None, is_active=True)
        db.session.add(s1)

    db.session.commit()


def _can_view_ss_request(r: HRSSRequest) -> bool:
    if r.requester_id == current_user.id:
        return True
    if current_user.has_perm(HR_SS_APPROVE) or current_user.has_perm(HR_SS_WORKFLOWS_MANAGE) or current_user.has_perm(HR_REQUESTS_VIEW_ALL):
        return True
    return False


def _notify_users(user_ids: list[int], message: str, level: str = "INFO") -> None:
    if not user_ids:
        return
    for uid in user_ids:
        db.session.add(Notification(user_id=uid, message=message, type=level, source='portal'))


def _notify_role(role: str, message: str, level: str = "INFO") -> None:
    role = (role or "").strip()
    if not role:
        return
    users = User.query.filter(func.upper(User.role) == role.upper()).all()
    _notify_users([u.id for u in users], message, level=level)


def _notify_for_step(approval: HRSSRequestApproval, r: HRSSRequest) -> None:
    msg = f"طلب HR Self-Service #{r.id} ({r.type_code}) بانتظار اعتمادك (Step {approval.step_no})."
    if approval.approver_user_id:
        _notify_users([approval.approver_user_id], msg)
    elif approval.approver_role:
        _notify_role(approval.approver_role, msg)


@portal_bp.route("/hr/self-service")
@login_required
@_perm(HR_SS_READ)
def hr_ss_home():
    _ensure_hr_ss_defaults()
    defs = HRSSWorkflowDefinition.query.filter_by(is_active=True).order_by(HRSSWorkflowDefinition.id.asc()).all()
    recent = (
        HRSSRequest.query.filter_by(requester_id=current_user.id)
        .order_by(HRSSRequest.created_at.desc())
        .limit(10)
        .all()
    )
    return render_template("portal/hr/ss_home.html", defs=defs, recent=recent)


@portal_bp.route("/hr/self-service/requests")
@login_required
@_perm(HR_SS_READ)
def hr_ss_my_requests():
    status = (request.args.get("status") or "").strip()
    type_code = (request.args.get("type") or "").strip()

    q = HRSSRequest.query.filter_by(requester_id=current_user.id)
    if status:
        q = q.filter(HRSSRequest.status == status)
    if type_code:
        q = q.filter(HRSSRequest.type_code == type_code)

    rows = q.order_by(HRSSRequest.updated_at.desc()).all()
    types = [r[0] for r in db.session.query(HRSSRequest.type_code).filter_by(requester_id=current_user.id).distinct().all()]
    return render_template("portal/hr/ss_my_requests.html", rows=rows, status=status, type_code=type_code, types=types)


@portal_bp.route("/hr/self-service/requests/new/<string:type_code>", methods=["GET", "POST"])
@login_required
@_perm(HR_SS_CREATE)
def hr_ss_new_request(type_code: str):
    _ensure_hr_ss_defaults()
    defn = HRSSWorkflowDefinition.query.filter_by(code=type_code).first_or_404()

    if request.method == "POST":
        payload = _ss_payload_from_form(type_code, request.form)
        r = HRSSRequest(
            requester_id=current_user.id,
            type_code=type_code,
            status="DRAFT",
            current_step_no=None,
            payload_json=json.dumps(payload, ensure_ascii=False),
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow(),
        )
        db.session.add(r)
        db.session.commit()

        # optional files
        files = request.files.getlist("files") if request.files else []
        if files:
            up_dir = _hr_ss_upload_dir(r.id)
            for f in files:
                if not f or not getattr(f, "filename", ""):
                    continue
                if not _ss_allowed_file(f.filename):
                    flash(f"امتداد غير مسموح: {f.filename}", "warning")
                    continue
                ext = _ss_safe_ext(f.filename)
                stored = f"{uuid.uuid4().hex}{('.' + ext) if ext else ''}"
                safe_orig = secure_filename(f.filename) or f.filename
                f.save(str(up_dir / stored))
                db.session.add(
                    HRSSRequestAttachment(
                        request_id=r.id,
                        doc_type=payload.get("doc_kind") or payload.get("certificate_kind") or "OTHER",
                        original_name=safe_orig,
                        stored_name=stored,
                        note=None,
                        uploaded_by_id=current_user.id,
                    )
                )
            db.session.commit()

        _portal_audit("HR_SS_REQUEST_CREATE", target_type="hr_ss_request", target_id=r.id, note=f"type={type_code}")
        flash("تم إنشاء الطلب كمسودة.", "success")
        return redirect(url_for("portal.hr_ss_view_request", req_id=r.id))

    return render_template("portal/hr/ss_new_request.html", defn=defn)


@portal_bp.route("/hr/self-service/requests/<int:req_id>")
@login_required
@_perm(HR_SS_READ)
def hr_ss_view_request(req_id: int):
    r = HRSSRequest.query.get_or_404(req_id)
    if not _can_view_ss_request(r):
        abort(403)

    payload = r.payload()
    approvals = HRSSRequestApproval.query.filter_by(request_id=r.id).order_by(HRSSRequestApproval.step_no.asc()).all()
    attachments = HRSSRequestAttachment.query.filter_by(request_id=r.id).order_by(HRSSRequestAttachment.uploaded_at.desc()).all()

    can_edit = (r.requester_id == current_user.id) and (r.status in ["DRAFT", "RETURNED"])
    can_submit = can_edit and current_user.has_perm(HR_SS_CREATE)
    can_cancel = (r.requester_id == current_user.id) and (r.status in ["DRAFT", "RETURNED"])

    return render_template(
        "portal/hr/ss_request_view.html",
        r=r,
        payload=payload,
        approvals=approvals,
        attachments=attachments,
        can_edit=can_edit,
        can_submit=can_submit,
        can_cancel=can_cancel,
    )


@portal_bp.route("/hr/self-service/requests/<int:req_id>/update", methods=["POST"])
@login_required
@_perm(HR_SS_CREATE)
def hr_ss_update_request(req_id: int):
    r = HRSSRequest.query.get_or_404(req_id)
    if r.requester_id != current_user.id:
        abort(403)
    if r.status not in ["DRAFT", "RETURNED"]:
        flash("لا يمكن تعديل هذا الطلب في هذه المرحلة.", "warning")
        return redirect(url_for("portal.hr_ss_view_request", req_id=r.id))

    payload = r.payload()
    payload["notes"] = (request.form.get("notes") or "").strip()
    r.payload_json = json.dumps(payload, ensure_ascii=False)
    r.updated_at = datetime.utcnow()
    db.session.commit()

    files = request.files.getlist("files") if request.files else []
    if files:
        up_dir = _hr_ss_upload_dir(r.id)
        for f in files:
            if not f or not getattr(f, "filename", ""):
                continue
            if not _ss_allowed_file(f.filename):
                flash(f"امتداد غير مسموح: {f.filename}", "warning")
                continue
            ext = _ss_safe_ext(f.filename)
            stored = f"{uuid.uuid4().hex}{('.' + ext) if ext else ''}"
            safe_orig = secure_filename(f.filename) or f.filename
            f.save(str(up_dir / stored))
            db.session.add(
                HRSSRequestAttachment(
                    request_id=r.id,
                    doc_type=payload.get("doc_kind") or payload.get("certificate_kind") or "OTHER",
                    original_name=safe_orig,
                    stored_name=stored,
                    note=None,
                    uploaded_by_id=current_user.id,
                )
            )
        db.session.commit()

    _portal_audit("HR_SS_REQUEST_UPDATE", target_type="hr_ss_request", target_id=r.id, note=None)
    flash("تم حفظ التعديل.", "success")
    return redirect(url_for("portal.hr_ss_view_request", req_id=r.id))


@portal_bp.route("/hr/self-service/requests/<int:req_id>/submit", methods=["POST"])
@login_required
@_perm(HR_SS_CREATE)
def hr_ss_submit_request(req_id: int):
    r = HRSSRequest.query.get_or_404(req_id)
    if r.requester_id != current_user.id:
        abort(403)
    if r.status not in ["DRAFT", "RETURNED"]:
        flash("لا يمكن تقديم هذا الطلب.", "warning")
        return redirect(url_for("portal.hr_ss_view_request", req_id=r.id))

    defn = HRSSWorkflowDefinition.query.filter_by(code=r.type_code, is_active=True).first()
    if not defn:
        flash("لا يوجد تعريف سير عمل لهذا النوع.", "danger")
        return redirect(url_for("portal.hr_ss_view_request", req_id=r.id))

    steps = (
        HRSSWorkflowStepDefinition.query.filter_by(definition_id=defn.id, is_active=True)
        .order_by(HRSSWorkflowStepDefinition.step_no.asc())
        .all()
    )
    # reset approvals
    HRSSRequestApproval.query.filter_by(request_id=r.id).delete()
    db.session.flush()

    if not steps:
        r.status = "APPROVED"
        r.closed_at = datetime.utcnow()
        r.submitted_at = datetime.utcnow()
        r.updated_at = datetime.utcnow()
        db.session.commit()
        flash("تمت الموافقة تلقائياً (لا يوجد خطوات).", "success")
        return redirect(url_for("portal.hr_ss_view_request", req_id=r.id))

    for idx, s in enumerate(steps):
        st = "PENDING" if idx == 0 else "QUEUED"
        db.session.add(
            HRSSRequestApproval(
                request_id=r.id,
                step_no=s.step_no,
                approver_role=s.approver_role,
                approver_user_id=s.approver_user_id,
                status=st,
                note=None,
            )
        )

    r.status = "IN_REVIEW"
    r.current_step_no = steps[0].step_no
    r.submitted_at = datetime.utcnow()
    r.updated_at = datetime.utcnow()
    db.session.commit()

    # notify first step
    first = HRSSRequestApproval.query.filter_by(request_id=r.id, step_no=r.current_step_no).first()
    if first:
        _notify_for_step(first, r)
        db.session.commit()

    _portal_audit("HR_SS_REQUEST_SUBMIT", target_type="hr_ss_request", target_id=r.id, note=None)
    flash("تم تقديم الطلب.", "success")
    return redirect(url_for("portal.hr_ss_view_request", req_id=r.id))


@portal_bp.route("/hr/self-service/requests/<int:req_id>/cancel", methods=["POST"])
@login_required
@_perm(HR_SS_READ)
def hr_ss_cancel_request(req_id: int):
    r = HRSSRequest.query.get_or_404(req_id)
    if r.requester_id != current_user.id:
        abort(403)
    if r.status not in ["DRAFT", "RETURNED"]:
        flash("لا يمكن إلغاء هذا الطلب.", "warning")
        return redirect(url_for("portal.hr_ss_view_request", req_id=r.id))

    r.status = "CANCELLED"
    r.updated_at = datetime.utcnow()
    r.closed_at = datetime.utcnow()
    db.session.commit()

    _portal_audit("HR_SS_REQUEST_CANCEL", target_type="hr_ss_request", target_id=r.id, note=None)
    flash("تم إلغاء الطلب.", "success")
    return redirect(url_for("portal.hr_ss_my_requests"))


@portal_bp.route("/hr/self-service/requests/<int:req_id>/attachments/<int:att_id>/download")
@login_required
@_perm(HR_SS_READ)
def hr_ss_download_attachment(req_id: int, att_id: int):
    r = HRSSRequest.query.get_or_404(req_id)
    if not _can_view_ss_request(r):
        abort(403)
    a = HRSSRequestAttachment.query.filter_by(id=att_id, request_id=req_id).first_or_404()
    up_dir = _hr_ss_upload_dir(req_id)
    return send_from_directory(str(up_dir), a.stored_name, as_attachment=True, download_name=a.original_name)


@portal_bp.route("/hr/self-service/approvals")
@login_required
def hr_ss_approvals():
    # Allow access if user can approve, or can manage/view all, or is SUPER.
    try:
        role_raw = (getattr(current_user, 'role', '') or '').strip().upper().replace('-', '_').replace(' ', '_')
    except Exception:
        role_raw = ''
    is_super = bool(role_raw.startswith('SUPER'))

    role = (current_user.role or "").strip()
    try:
        can_view_all = bool(current_user.has_perm(HR_SS_WORKFLOWS_MANAGE) or current_user.has_perm(HR_REQUESTS_VIEW_ALL))
    except Exception:
        can_view_all = False
    try:
        can_approve = bool(current_user.has_perm(HR_SS_APPROVE))
    except Exception:
        can_approve = False

    # If not explicitly allowed, still allow if there are pending approvals assigned to the user/role
    if not (is_super or can_view_all or can_approve):
        role_norm = (role or '').strip().upper()
        role_cmp = func.replace(func.replace(func.replace(func.upper(HRSSRequestApproval.approver_role), '_', ''), '-', ''), ' ', '')
        my_cmp = (role_norm.replace('_', '').replace('-', '').replace(' ', ''))
        exists_q = (
            HRSSRequestApproval.query.join(HRSSRequest)
            .filter(HRSSRequest.status == 'IN_REVIEW', HRSSRequestApproval.status == 'PENDING')
            .filter((HRSSRequestApproval.approver_user_id == current_user.id) | (role_cmp == my_cmp))
        )
        if exists_q.first() is None:
            abort(403)

    q = HRSSRequestApproval.query.join(HRSSRequest).filter(HRSSRequest.status == "IN_REVIEW", HRSSRequestApproval.status == "PENDING")
    if not (is_super or can_view_all):
        # Normal approver: only items assigned to user/role
        role_norm = (role or '').strip().upper()
        role_cmp = func.replace(func.replace(func.replace(func.upper(HRSSRequestApproval.approver_role), '_', ''), '-', ''), ' ', '')
        my_cmp = (role_norm.replace('_', '').replace('-', '').replace(' ', ''))
        q = q.filter((HRSSRequestApproval.approver_user_id == current_user.id) | (role_cmp == my_cmp))
    approvals = q.order_by(HRSSRequest.submitted_at.asc().nullslast(), HRSSRequest.id.asc()).all()

    rows = []
    for a in approvals:
        req = HRSSRequest.query.get(a.request_id)
        if req:
            rows.append(type("Row", (), {"appr": a, "req": req}))

    return render_template("portal/hr/ss_approvals.html", rows=rows)


@portal_bp.route("/hr/self-service/approvals/<int:req_id>", methods=["GET", "POST"])
@login_required
def hr_ss_approval_view(req_id: int):
    r = HRSSRequest.query.get_or_404(req_id)
    if r.status != "IN_REVIEW" or not r.current_step_no:
        flash("هذا الطلب ليس قيد الاعتماد.", "warning")
        return redirect(url_for("portal.hr_ss_approvals"))

    role = (current_user.role or "").strip()
    appr = HRSSRequestApproval.query.filter_by(request_id=r.id, step_no=r.current_step_no, status="PENDING").first()
    if not appr:
        flash("لا يوجد خطوة حالية بانتظار الاعتماد.", "warning")
        return redirect(url_for("portal.hr_ss_approvals"))

    role_norm = (role or '').strip().upper()
    appr_role_norm = (appr.approver_role or '').strip().upper()
    role_cmp = role_norm.replace('_', '').replace('-', '').replace(' ', '')
    appr_cmp = appr_role_norm.replace('_', '').replace('-', '').replace(' ', '')

    allowed = (appr.approver_user_id == current_user.id) or (appr_cmp and appr_cmp == role_cmp)

    # SUPER / managers can access any pending approvals
    try:
        role_raw = (getattr(current_user, 'role', '') or '').strip().upper().replace('-', '_').replace(' ', '_')
    except Exception:
        role_raw = ''
    is_super = bool(role_raw.startswith('SUPER'))
    try:
        can_view_all = bool(current_user.has_perm(HR_SS_WORKFLOWS_MANAGE) or current_user.has_perm(HR_REQUESTS_VIEW_ALL))
    except Exception:
        can_view_all = False

    if not (allowed or is_super or can_view_all):
        abort(403)

    if request.method == "POST":
        action = (request.form.get("action") or "").strip()
        note = (request.form.get("note") or "").strip() or None

        appr.note = note
        appr.acted_by_id = current_user.id
        appr.acted_at = datetime.utcnow()

        if action == "approve":
            appr.status = "APPROVED"
            # move to next
            next_appr = HRSSRequestApproval.query.filter_by(request_id=r.id, step_no=r.current_step_no + 1).first()
            if next_appr:
                next_appr.status = "PENDING"
                r.current_step_no = next_appr.step_no
                r.updated_at = datetime.utcnow()
                _notify_for_step(next_appr, r)
                flash("تمت الموافقة والانتقال للخطوة التالية.", "success")
            else:
                r.status = "APPROVED"
                r.closed_at = datetime.utcnow()
                r.current_step_no = None
                r.updated_at = datetime.utcnow()
                _notify_users([r.requester_id], f"تم اعتماد طلبك HR Self-Service #{r.id}.", level="SUCCESS")
                flash("تمت الموافقة النهائية.", "success")

        elif action == "reject":
            appr.status = "REJECTED"
            r.status = "REJECTED"
            r.closed_at = datetime.utcnow()
            r.current_step_no = None
            r.updated_at = datetime.utcnow()
            _notify_users([r.requester_id], f"تم رفض طلبك HR Self-Service #{r.id}.", level="DANGER")
            flash("تم رفض الطلب.", "danger")

        elif action == "return":
            appr.status = "RETURNED"
            r.status = "RETURNED"
            r.current_step_no = None
            r.updated_at = datetime.utcnow()
            _notify_users([r.requester_id], f"تم إرجاع طلبك HR Self-Service #{r.id} للتعديل.", level="WARNING")
            flash("تم إرجاع الطلب للموظف.", "warning")

        else:
            flash("إجراء غير معروف.", "warning")
            return redirect(url_for("portal.hr_ss_approval_view", req_id=r.id))

        db.session.commit()
        _portal_audit("HR_SS_APPROVAL_ACTION", target_type="hr_ss_request", target_id=r.id, note=f"action={action}")
        return redirect(url_for("portal.hr_ss_approvals"))

    payload = r.payload()
    attachments = HRSSRequestAttachment.query.filter_by(request_id=r.id).order_by(HRSSRequestAttachment.uploaded_at.desc()).all()
    return render_template("portal/hr/ss_approval_view.html", r=r, appr=appr, payload=payload, attachments=attachments)


# -------------------------
# HR Docs
# -------------------------

@portal_bp.route("/hr/docs")
@login_required
@_perm(HR_DOCS_READ)
def hr_docs_home():
    docs = HRDoc.query.filter_by(is_published=True).order_by(HRDoc.category.asc(), HRDoc.title_ar.asc()).all()
    # attach current version
    rows = []
    for d in docs:
        cur = HRDocVersion.query.get(d.current_version_id) if d.current_version_id else None
        rows.append(type("Row", (), {"id": d.id, "title_ar": d.title_ar, "title_en": d.title_en, "category": d.category, "current_ver": cur, "created_at": d.created_at, "created_by": d.created_by}))
    return render_template("portal/hr/docs_index.html", docs=rows)


@portal_bp.route("/hr/docs/<int:doc_id>")
@login_required
@_perm(HR_DOCS_READ)
def hr_docs_view(doc_id: int):
    doc = HRDoc.query.get_or_404(doc_id)
    versions = HRDocVersion.query.filter_by(doc_id=doc.id).order_by(HRDocVersion.version_no.desc()).all()
    return render_template("portal/hr/docs_view.html", doc=doc, versions=versions)


@portal_bp.route("/hr/docs/<int:doc_id>/download/<int:ver_id>")
@login_required
@_perm(HR_DOCS_READ)
def hr_docs_download(doc_id: int, ver_id: int):
    v = HRDocVersion.query.filter_by(id=ver_id, doc_id=doc_id).first_or_404()
    up_dir = _hr_docs_upload_dir(doc_id)
    return send_from_directory(str(up_dir), v.stored_name, as_attachment=True, download_name=v.original_name)


@portal_bp.route("/hr/docs/admin")
@login_required
@_perm(HR_DOCS_MANAGE)
def hr_docs_admin():
    focus = request.args.get("focus")
    docs = HRDoc.query.order_by(HRDoc.created_at.desc()).all()
    rows = []
    for d in docs:
        cur = HRDocVersion.query.get(d.current_version_id) if d.current_version_id else None
        rows.append(type("Row", (), {"id": d.id, "title_ar": d.title_ar, "category": d.category, "current_ver": cur}))
    return render_template("portal/hr/docs_admin.html", docs=rows, focus=focus)


@portal_bp.route("/hr/docs/admin/create", methods=["POST"])
@login_required
@_perm(HR_DOCS_MANAGE)
def hr_docs_create():
    title_ar = (request.form.get("title_ar") or "").strip()
    category = (request.form.get("category") or "POLICY").strip()
    if not title_ar:
        flash("العنوان مطلوب.", "danger")
        return redirect(url_for("portal.hr_docs_admin"))

    d = HRDoc(title_ar=title_ar, title_en="", category=category, is_published=True, created_by_id=current_user.id, created_at=datetime.utcnow())
    db.session.add(d)
    db.session.commit()
    _portal_audit("HR_DOC_CREATE", target_type="hr_doc", target_id=d.id, note=category)
    flash("تم إنشاء الوثيقة.", "success")
    return redirect(url_for("portal.hr_docs_admin", focus=d.id))


@portal_bp.route("/hr/docs/admin/upload", methods=["POST"])
@login_required
@_perm(HR_DOCS_MANAGE)
def hr_docs_upload_version():
    doc_id = int(request.form.get("doc_id") or 0)
    doc = HRDoc.query.get_or_404(doc_id)

    f = request.files.get("file")
    if not f or not getattr(f, "filename", ""):
        flash("الملف مطلوب.", "danger")
        return redirect(url_for("portal.hr_docs_admin", focus=doc.id))

    if not _ss_allowed_file(f.filename):
        flash("امتداد غير مسموح.", "danger")
        return redirect(url_for("portal.hr_docs_admin", focus=doc.id))

    change_log = (request.form.get("change_log") or "").strip() or None
    effective_date = (request.form.get("effective_date") or "").strip() or None

    # new version_no
    max_no = db.session.query(func.max(HRDocVersion.version_no)).filter_by(doc_id=doc.id).scalar() or 0
    new_no = int(max_no) + 1

    ext = _ss_safe_ext(f.filename)
    stored = f"{uuid.uuid4().hex}{('.' + ext) if ext else ''}"
    safe_orig = secure_filename(f.filename) or f.filename
    up_dir = _hr_docs_upload_dir(doc.id)
    f.save(str(up_dir / stored))

    v = HRDocVersion(
        doc_id=doc.id,
        version_no=new_no,
        original_name=safe_orig,
        stored_name=stored,
        change_log=change_log,
        approved_by_id=current_user.id,
        approved_at=datetime.utcnow(),
        effective_date=effective_date,
        created_by_id=current_user.id,
        created_at=datetime.utcnow(),
    )
    db.session.add(v)
    db.session.flush()
    doc.current_version_id = v.id
    db.session.commit()

    _portal_audit("HR_DOC_UPLOAD_VERSION", target_type="hr_doc", target_id=doc.id, note=f"v{new_no}")
    flash("تم رفع الإصدار واعتماده.", "success")
    return redirect(url_for("portal.hr_docs_view", doc_id=doc.id))


# -------------------------
# HR Discipline
# -------------------------

@portal_bp.route("/hr/discipline")
@login_required
@_perm(HR_DISCIPLINE_READ)
def hr_discipline_home():
    rows = HRDisciplinaryCase.query.order_by(HRDisciplinaryCase.created_at.desc()).all()
    return render_template("portal/hr/discipline_index.html", rows=rows)


@portal_bp.route("/hr/discipline/new", methods=["GET", "POST"])
@login_required
@_perm(HR_DISCIPLINE_MANAGE)
def hr_discipline_new():
    if request.method == "POST":
        employee_id = int(request.form.get("employee_id") or 0)
        title = (request.form.get("title") or "").strip()
        category = (request.form.get("category") or "VIOLATION").strip()
        severity = (request.form.get("severity") or "LOW").strip()
        description = (request.form.get("description") or "").strip() or None

        if not employee_id or not title:
            flash("الموظف والعنوان مطلوبان.", "danger")
            return redirect(url_for("portal.hr_discipline_new"))

        c = HRDisciplinaryCase(
            employee_id=employee_id,
            title=title,
            category=category,
            severity=severity,
            description=description,
            created_by_id=current_user.id,
            status="OPEN",
            created_at=datetime.utcnow(),
        )
        db.session.add(c)
        db.session.commit()
        _portal_audit("HR_DISCIPLINE_CREATE", target_type="hr_disciplinary_case", target_id=c.id, note=category)
        flash("تم إنشاء القضية.", "success")
        return redirect(url_for("portal.hr_discipline_view", case_id=c.id))

    name_expr = func.nullif(func.trim(User.name), "")
    users = (
        User.query
        .order_by(func.lower(func.coalesce(name_expr, User.email)).asc(), User.id.asc())
        .all()
    )
    return render_template("portal/hr/discipline_new.html", users=users)


@portal_bp.route("/hr/discipline/<int:case_id>")
@login_required
@_perm(HR_DISCIPLINE_READ)
def hr_discipline_view(case_id: int):
    c = HRDisciplinaryCase.query.get_or_404(case_id)
    can_manage = current_user.has_perm(HR_DISCIPLINE_MANAGE)
    actions = HRDisciplinaryAction.query.filter_by(case_id=c.id).order_by(HRDisciplinaryAction.action_date.desc()).all()
    atts = HRDisciplinaryAttachment.query.filter_by(case_id=c.id).order_by(HRDisciplinaryAttachment.uploaded_at.desc()).all()
    return render_template("portal/hr/discipline_case.html", c=c, can_manage=can_manage, actions=actions, atts=atts)


@portal_bp.route("/hr/discipline/<int:case_id>/action", methods=["POST"])
@login_required
@_perm(HR_DISCIPLINE_MANAGE)
def hr_discipline_add_action(case_id: int):
    c = HRDisciplinaryCase.query.get_or_404(case_id)
    action_type = (request.form.get("action_type") or "NOTE").strip()
    note = (request.form.get("note") or "").strip() or None
    a = HRDisciplinaryAction(case_id=c.id, action_type=action_type, note=note, created_by_id=current_user.id, action_date=datetime.utcnow())
    db.session.add(a)
    db.session.commit()
    _portal_audit("HR_DISCIPLINE_ADD_ACTION", target_type="hr_disciplinary_case", target_id=c.id, note=action_type)
    flash("تمت إضافة الإجراء.", "success")
    return redirect(url_for("portal.hr_discipline_view", case_id=c.id))


@portal_bp.route("/hr/discipline/<int:case_id>/attachments/upload", methods=["POST"])
@login_required
@_perm(HR_DISCIPLINE_MANAGE)
def hr_discipline_upload_attachment(case_id: int):
    c = HRDisciplinaryCase.query.get_or_404(case_id)
    f = request.files.get("file")
    if not f or not getattr(f, "filename", ""):
        flash("الملف مطلوب.", "danger")
        return redirect(url_for("portal.hr_discipline_view", case_id=c.id))
    if not _ss_allowed_file(f.filename):
        flash("امتداد غير مسموح.", "danger")
        return redirect(url_for("portal.hr_discipline_view", case_id=c.id))

    ext = _ss_safe_ext(f.filename)
    stored = f"{uuid.uuid4().hex}{('.' + ext) if ext else ''}"
    safe_orig = secure_filename(f.filename) or f.filename
    up_dir = _hr_discipline_upload_dir(c.id)
    f.save(str(up_dir / stored))

    note = (request.form.get("note") or "").strip() or None
    att = HRDisciplinaryAttachment(case_id=c.id, original_name=safe_orig, stored_name=stored, note=note, uploaded_by_id=current_user.id)
    db.session.add(att)
    db.session.commit()

    _portal_audit("HR_DISCIPLINE_UPLOAD_ATTACHMENT", target_type="hr_disciplinary_case", target_id=c.id, note=None)
    flash("تم رفع المرفق.", "success")
    return redirect(url_for("portal.hr_discipline_view", case_id=c.id))


@portal_bp.route("/hr/discipline/<int:case_id>/attachments/<int:att_id>/download")
@login_required
@_perm(HR_DISCIPLINE_READ)
def hr_discipline_download_attachment(case_id: int, att_id: int):
    c = HRDisciplinaryCase.query.get_or_404(case_id)
    a = HRDisciplinaryAttachment.query.filter_by(id=att_id, case_id=c.id).first_or_404()
    up_dir = _hr_discipline_upload_dir(c.id)
    return send_from_directory(str(up_dir), a.stored_name, as_attachment=True, download_name=a.original_name)


# -------------------------
# Admin: HR SS Workflows
# -------------------------

@portal_bp.route("/admin/hr/self-service-workflows", methods=["GET", "POST"])
@login_required
@_perm(HR_SS_WORKFLOWS_MANAGE)
def portal_admin_hr_ss_workflows():
    _ensure_hr_ss_defaults()

    if request.method == "POST":
        action = (request.form.get("action") or "").strip()

        if action == "add_def":
            code = (request.form.get("code") or "").strip().upper()
            name_ar = (request.form.get("name_ar") or "").strip()
            name_en = (request.form.get("name_en") or "").strip()
            if not code or not name_ar:
                flash("الكود والاسم (AR) مطلوبان.", "danger")
                return redirect(url_for("portal.portal_admin_hr_ss_workflows"))
            if HRSSWorkflowDefinition.query.filter_by(code=code).first():
                flash("الكود موجود مسبقاً.", "warning")
                return redirect(url_for("portal.portal_admin_hr_ss_workflows"))

            d = HRSSWorkflowDefinition(code=code, name_ar=name_ar, name_en=name_en, is_active=True, created_by_id=current_user.id)
            db.session.add(d)
            db.session.commit()
            _portal_audit("HR_SS_WORKFLOW_DEF_CREATE", target_type="hr_ss_workflow_definition", target_id=d.id, note=code)
            flash("تمت إضافة نوع الطلب.", "success")
            return redirect(url_for("portal.portal_admin_hr_ss_workflows"))

        if action == "add_step":
            definition_id = int(request.form.get("definition_id") or 0)
            step_no = int(request.form.get("step_no") or 0)
            sla_hours = request.form.get("sla_hours")
            sla_hours = int(sla_hours) if sla_hours and sla_hours.isdigit() else None
            approver_role = (request.form.get("approver_role") or "").strip() or None
            approver_user_id = request.form.get("approver_user_id")
            approver_user_id = int(approver_user_id) if approver_user_id and approver_user_id.isdigit() else None

            if not definition_id or not step_no:
                flash("نوع الطلب ورقم الخطوة مطلوبان.", "danger")
                return redirect(url_for("portal.portal_admin_hr_ss_workflows"))

            # If user chosen, ignore role
            if approver_user_id:
                approver_role = None

            s = HRSSWorkflowStepDefinition(
                definition_id=definition_id,
                step_no=step_no,
                approver_role=approver_role,
                approver_user_id=approver_user_id,
                sla_hours=sla_hours,
                is_active=True,
            )
            db.session.add(s)
            try:
                db.session.commit()
            except Exception:
                db.session.rollback()
                flash("تعذر إضافة الخطوة (قد يكون رقم الخطوة مكرر).", "danger")
                return redirect(url_for("portal.portal_admin_hr_ss_workflows"))

            _portal_audit("HR_SS_WORKFLOW_STEP_CREATE", target_type="hr_ss_workflow_step_definition", target_id=s.id, note=f"def={definition_id},step={step_no}")
            flash("تمت إضافة الخطوة.", "success")
            return redirect(url_for("portal.portal_admin_hr_ss_workflows"))

        if action == "toggle_step":
            step_id = int(request.form.get("step_id") or 0)
            s = HRSSWorkflowStepDefinition.query.get_or_404(step_id)
            s.is_active = not s.is_active
            db.session.commit()
            flash("تم تحديث حالة الخطوة.", "success")
            return redirect(url_for("portal.portal_admin_hr_ss_workflows"))

        if action == "delete_step":
            step_id = int(request.form.get("step_id") or 0)
            s = HRSSWorkflowStepDefinition.query.get_or_404(step_id)
            db.session.delete(s)
            db.session.commit()
            flash("تم حذف الخطوة.", "success")
            return redirect(url_for("portal.portal_admin_hr_ss_workflows"))

        flash("إجراء غير معروف.", "warning")
        return redirect(url_for("portal.portal_admin_hr_ss_workflows"))

    defs = HRSSWorkflowDefinition.query.order_by(HRSSWorkflowDefinition.code.asc()).all()
    name_expr = func.nullif(func.trim(User.name), "")
    users = (
        User.query
        .order_by(func.lower(func.coalesce(name_expr, User.email)).asc(), User.id.asc())
        .all()
    )
    return render_template("portal/admin/hr_ss_workflows.html", defs=defs, users=users)


# =========================================================
# HR: Performance & Evaluation (360)
# =========================================================


def _primary_unit_for_user(user_id: int):
    """Return (unit_type, unit_id) for user's primary org assignment (best-effort)."""
    try:
        row = (OrgUnitAssignment.query.filter_by(user_id=user_id, is_primary=True).first()
               or OrgUnitAssignment.query.filter_by(user_id=user_id).order_by(OrgUnitAssignment.id.desc()).first())
        if row and row.unit_type and row.unit_id:
            return (row.unit_type.upper(), int(row.unit_id))
    except Exception:
        pass
    return (None, None)


def _unit_parents(unit_type: str, unit_id: int):
    """Yield parent chain (unit_type, unit_id)."""
    try:
        ut = (unit_type or '').upper()
        if ut == 'TEAM':
            t = Team.query.get(unit_id)
            if t and t.section_id:
                yield ('SECTION', int(t.section_id))
                ut, unit_id = 'SECTION', int(t.section_id)
            else:
                return
        if ut == 'SECTION':
            s = Section.query.get(unit_id)
            if s and s.department_id:
                yield ('DEPARTMENT', int(s.department_id))
                ut, unit_id = 'DEPARTMENT', int(s.department_id)
            else:
                return
        if ut == 'DEPARTMENT':
            d = Department.query.get(unit_id)
            if d and d.directorate_id:
                yield ('DIRECTORATE', int(d.directorate_id))
                ut, unit_id = 'DIRECTORATE', int(d.directorate_id)
            else:
                return
        if ut == 'DIRECTORATE':
            dr = Directorate.query.get(unit_id)
            if dr and dr.organization_id:
                yield ('ORGANIZATION', int(dr.organization_id))
    except Exception:
        return


def _find_direct_manager_for_user(user_id: int):
    """Resolve direct manager (or deputy) based on OrgUnitManager for the user's org unit chain."""
    try:
        u = User.query.get(user_id)
        if not u:
            return None
        unit_type, unit_id = _primary_unit_for_user(user_id)
        candidates = []
        if unit_type and unit_id:
            candidates.append((unit_type, unit_id))
            candidates.extend(list(_unit_parents(unit_type, unit_id)))
        # fallback: department/directorate from user
        if getattr(u, 'department_id', None):
            candidates.append(('DEPARTMENT', int(u.department_id)))
        if getattr(u, 'directorate_id', None):
            candidates.append(('DIRECTORATE', int(u.directorate_id)))

        seen = set()
        for ut, uid in candidates:
            if not ut or not uid:
                continue
            key = (ut, uid)
            if key in seen:
                continue
            seen.add(key)
            m = OrgUnitManager.query.filter_by(unit_type=ut, unit_id=uid).first()
            if not m:
                continue
            if m.manager_user_id and int(m.manager_user_id) != int(user_id):
                return m.manager_user
            if m.deputy_user_id and int(m.deputy_user_id) != int(user_id):
                return m.deputy_user
    except Exception:
        pass
    return None


def _pick_peer_candidates(user_id: int, peer_count: int, exclude_ids: set[int] | None = None):
    """Pick peers from the same primary unit (best-effort)."""
    exclude_ids = set(exclude_ids or set())
    exclude_ids.add(int(user_id))

    unit_type, unit_id = _primary_unit_for_user(user_id)
    if not unit_type or not unit_id:
        return []

    try:
        rows = (
            OrgUnitAssignment.query
            .filter_by(unit_type=unit_type, unit_id=unit_id)
            .order_by(OrgUnitAssignment.is_primary.desc(), OrgUnitAssignment.id.desc())
            .all()
        )
        peer_user_ids = []
        for r in rows:
            if not r.user_id:
                continue
            uid = int(r.user_id)
            if uid in exclude_ids:
                continue
            if uid not in peer_user_ids:
                peer_user_ids.append(uid)
        if not peer_user_ids:
            return []

        # deterministic ordering
        name_expr = func.nullif(func.trim(User.name), "")
        peers = (
            User.query
            .filter(User.id.in_(peer_user_ids))
            .order_by(func.lower(func.coalesce(name_expr, User.email)).asc(), User.id.asc())
            .all()
        )
        return peers[: max(0, int(peer_count or 0))]
    except Exception:
        return []


def _perf_questions_for_form(form_id: int):
    sections = HRPerformanceSection.query.filter_by(form_id=form_id).order_by(HRPerformanceSection.order_no.asc(), HRPerformanceSection.id.asc()).all()
    questions_by_section = {}
    for s in sections:
        qs = HRPerformanceQuestion.query.filter_by(section_id=s.id).order_by(HRPerformanceQuestion.order_no.asc(), HRPerformanceQuestion.id.asc()).all()
        questions_by_section[s.id] = qs
    return sections, questions_by_section


def _perf_compute_score(form_id: int, answers: dict):
    """Compute a weighted score from answers (rating/number/yesno)."""
    total_w = 0.0
    total = 0.0

    # Flatten questions
    qrows = (
        HRPerformanceQuestion.query
        .join(HRPerformanceSection, HRPerformanceQuestion.section_id == HRPerformanceSection.id)
        .filter(HRPerformanceSection.form_id == int(form_id))
        .all()
    )

    for q in qrows:
        qid = str(q.id)
        if qid not in answers:
            continue
        val = (answers.get(qid) or {}).get('value')
        if val in (None, ''):
            continue

        try:
            w = float(q.weight or 1.0)
        except Exception:
            w = 1.0

        vnum = None
        try:
            if q.q_type == 'RATING_1_5':
                vnum = float(int(val))
            elif q.q_type == 'NUMBER':
                vnum = float(val)
            elif q.q_type == 'YESNO':
                vv = str(val).strip().lower()
                vnum = 1.0 if vv in ('1','yes','y','true','نعم') else 0.0
        except Exception:
            vnum = None

        if vnum is None:
            continue

        total_w += w
        total += (w * vnum)

    if total_w <= 0:
        return None
    return round(total / total_w, 3)


@portal_bp.route("/hr/performance")
@login_required
@_perm(HR_PERF_READ)
def hr_perf_home():
    """Employee: show my assigned evaluations and my summary."""
    cycles = HRPerformanceCycle.query.order_by(HRPerformanceCycle.id.desc()).limit(30).all()

    assignments = (
        HRPerformanceAssignment.query
        .join(HRPerformanceCycle, HRPerformanceAssignment.cycle_id == HRPerformanceCycle.id)
        .filter(HRPerformanceAssignment.evaluator_user_id == current_user.id)
        .order_by(HRPerformanceAssignment.id.desc())
        .limit(200)
        .all()
    )

    # My summary (average score per cycle where I'm the evaluatee)
    my_summary = []
    try:
        rows = (
            db.session.query(
                HRPerformanceCycle.id,
                HRPerformanceCycle.name,
                func.avg(HRPerformanceAssignment.score_total)
            )
            .join(HRPerformanceAssignment, HRPerformanceAssignment.cycle_id == HRPerformanceCycle.id)
            .filter(HRPerformanceAssignment.evaluatee_user_id == current_user.id)
            .filter(HRPerformanceAssignment.status == 'SUBMITTED')
            .group_by(HRPerformanceCycle.id, HRPerformanceCycle.name)
            .order_by(HRPerformanceCycle.id.desc())
            .all()
        )
        for cid, cname, avg_score in rows:
            my_summary.append({"cycle_id": cid, "cycle_name": cname, "avg_score": float(avg_score) if avg_score is not None else None})
    except Exception:
        my_summary = []

    return render_template(
        "portal/hr/performance_home.html",
        cycles=cycles,
        assignments=assignments,
        my_summary=my_summary,
    )



@portal_bp.route("/hr/performance/assignment/<int:assignment_id>", methods=["GET", "POST"])
@login_required
@_perm(HR_PERF_READ)
def hr_perf_assignment(assignment_id: int):
    a = HRPerformanceAssignment.query.get_or_404(assignment_id)

    # Access control
    is_admin = False
    try:
        is_admin = (current_user.has_role("SUPERADMIN") or current_user.has_role("SUPER_ADMIN") or current_user.has_perm(HR_PERF_MANAGE))
    except Exception:
        is_admin = False

    if not is_admin and int(a.evaluator_user_id) != int(current_user.id):
        abort(403)

    try:
        can_submit = current_user.has_perm(HR_PERF_SUBMIT) or is_admin
    except Exception:
        can_submit = is_admin

    cycle = a.cycle
    form = cycle.form if cycle else None
    if not form:
        flash("نموذج التقييم غير موجود.", "danger")
        return redirect(url_for("portal.hr_perf_home"))

    sections, questions_by_section = _perf_questions_for_form(form.id)

    saved = {}
    try:
        saved = json.loads(a.answers_json) if a.answers_json else {}
        if not isinstance(saved, dict):
            saved = {}
    except Exception:
        saved = {}

    if request.method == "POST":
        if not can_submit:
            abort(403)
        if a.status == 'SUBMITTED':
            flash("تم إرسال هذا التقييم مسبقاً.", "info")
            return redirect(url_for("portal.hr_perf_assignment", assignment_id=a.id))

        answers = {}
        errors = []

        # Collect answers
        for s in sections:
            for q in questions_by_section.get(s.id, []):
                field = f"q_{q.id}"
                cfield = f"c_{q.id}"
                val = (request.form.get(field) or '').strip()
                comment = (request.form.get(cfield) or '').strip()

                if q.is_required and not val:
                    errors.append(f"السؤال مطلوب: {q.prompt}")

                if val or comment:
                    answers[str(q.id)] = {"value": val, "comment": comment}

        overall_comment = (request.form.get('overall_comment') or '').strip()
        if overall_comment:
            answers['overall_comment'] = overall_comment

        if errors:
            for e in errors[:6]:
                flash(e, "danger")
            return render_template(
                "portal/hr/performance_assignment.html",
                a=a,
                cycle=cycle,
                form=form,
                sections=sections,
                questions_by_section=questions_by_section,
                saved=answers,
                can_edit=can_submit,
            )

        a.answers_json = json.dumps(answers, ensure_ascii=False)
        a.score_total = _perf_compute_score(form.id, answers)
        a.status = 'SUBMITTED'
        a.submitted_at = datetime.utcnow()
        try:
            a.created_by_id = a.created_by_id or current_user.id
        except Exception:
            pass

        _portal_audit(
            "HR_PERF_SUBMIT",
            note=f"cycle={a.cycle_id},ee={a.evaluatee_user_id},type={a.evaluator_type}",
            target_type="hr_perf_assignment",
            target_id=a.id,
        )

        try:
            db.session.commit()
            flash("تم إرسال التقييم.", "success")
        except Exception:
            db.session.rollback()
            flash("تعذر حفظ التقييم.", "danger")

        return redirect(url_for("portal.hr_perf_home"))

    return render_template(
        "portal/hr/performance_assignment.html",
        a=a,
        cycle=cycle,
        form=form,
        sections=sections,
        questions_by_section=questions_by_section,
        saved=saved,
        can_edit=can_submit,
    )


@portal_bp.route("/admin/hr/performance")
@login_required
@_perm(HR_PERF_MANAGE)
def portal_admin_hr_perf_dashboard():
    cards = [
        {
            "title": "نماذج التقييم",
            "desc": "إنشاء نموذج تقييم ديناميكي (أقسام + أسئلة).",
            "icon": "bi-ui-checks",
            "url": url_for("portal.portal_admin_hr_perf_forms"),
        },
        {
            "title": "دورات الأداء",
            "desc": "تعريف دورة (فترة) وربطها بنموذج وتوليد تكليفات 360.",
            "icon": "bi-calendar-event",
            "url": url_for("portal.portal_admin_hr_perf_cycles"),
        },
    ]
    return render_template("portal/admin/hr_perf_dashboard.html", cards=cards)


@portal_bp.route("/admin/hr/performance/forms", methods=["GET", "POST"])
@login_required
@_perm(HR_PERF_MANAGE)
def portal_admin_hr_perf_forms():
    if request.method == "POST":
        action = (request.form.get('action') or 'create').strip().lower()

        if action == 'delete':
            fid = int(request.form.get('form_id') or 0)
            f = HRPerformanceForm.query.get_or_404(fid)
            # Safety: if linked to cycles, don't hard-delete; just deactivate.
            try:
                linked = HRPerformanceCycle.query.filter_by(form_id=f.id).count()
            except Exception:
                linked = 0

            try:
                if linked and linked > 0:
                    f.is_active = False
                    _portal_audit("HR_PERF_FORM_DEACTIVATE", note=f.name, target_type='hr_perf_form', target_id=f.id)
                    db.session.commit()
                    flash("النموذج مرتبط بدورات تقييم. تم إيقافه بدلاً من الحذف.", "warning")
                    return redirect(url_for("portal.portal_admin_hr_perf_forms"))

                # delete sections/questions first
                sec_ids = [s.id for s in HRPerformanceSection.query.filter_by(form_id=f.id).all()]
                if sec_ids:
                    HRPerformanceQuestion.query.filter(HRPerformanceQuestion.section_id.in_(sec_ids)).delete(synchronize_session=False)
                    HRPerformanceSection.query.filter(HRPerformanceSection.id.in_(sec_ids)).delete(synchronize_session=False)
                db.session.delete(f)
                _portal_audit("HR_PERF_FORM_DELETE", note=f.name, target_type='hr_perf_form', target_id=f.id)
                db.session.commit()
                flash("تم حذف النموذج.", "success")
            except Exception:
                db.session.rollback()
                flash("تعذر الحذف (قد يكون مرتبطاً ببيانات أخرى).", "danger")
            return redirect(url_for("portal.portal_admin_hr_perf_forms"))

        # create
        name = (request.form.get('name') or '').strip()
        desc = (request.form.get('description') or '').strip()
        if not name:
            flash("اسم النموذج مطلوب.", "danger")
            return redirect(url_for("portal.portal_admin_hr_perf_forms"))
        f = HRPerformanceForm(name=name, description=desc, created_by_id=current_user.id)
        db.session.add(f)
        _portal_audit("HR_PERF_FORM_CREATE", note=name, target_type='hr_perf_form', target_id=None)
        try:
            db.session.commit()
            flash("تم إنشاء النموذج.", "success")
        except Exception:
            db.session.rollback()
            flash("تعذر إنشاء النموذج.", "danger")
        return redirect(url_for("portal.portal_admin_hr_perf_forms"))

    forms = HRPerformanceForm.query.order_by(HRPerformanceForm.id.desc()).all()
    return render_template("portal/admin/hr_perf_forms.html", forms=forms)


@portal_bp.route("/admin/hr/performance/forms/<int:form_id>", methods=["GET", "POST"])
@login_required
@_perm(HR_PERF_MANAGE)
def portal_admin_hr_perf_form_edit(form_id: int):
    f = HRPerformanceForm.query.get_or_404(form_id)

    if request.method == "POST":
        action = (request.form.get('action') or '').strip()

        if action == 'update_form':
            f.name = (request.form.get('name') or '').strip() or f.name
            f.description = (request.form.get('description') or '').strip()
            f.is_active = True if (request.form.get('is_active') == '1') else False
            _portal_audit("HR_PERF_FORM_UPDATE", note=f.name, target_type='hr_perf_form', target_id=f.id)
            db.session.commit()
            flash("تم حفظ النموذج.", "success")
            return redirect(url_for('portal.portal_admin_hr_perf_form_edit', form_id=f.id))

        if action == 'add_section':
            title = (request.form.get('section_title') or '').strip()
            if not title:
                flash("عنوان القسم مطلوب.", "danger")
                return redirect(url_for('portal.portal_admin_hr_perf_form_edit', form_id=f.id))
            order_no = int(request.form.get('section_order') or 0)
            s = HRPerformanceSection(form_id=f.id, title=title, order_no=order_no)
            db.session.add(s)
            _portal_audit("HR_PERF_SECTION_CREATE", note=title, target_type='hr_perf_section', target_id=None)
            db.session.commit()
            flash("تمت إضافة قسم.", "success")
            return redirect(url_for('portal.portal_admin_hr_perf_form_edit', form_id=f.id))

        if action == 'delete_section':
            sid = int(request.form.get('section_id') or 0)
            s = HRPerformanceSection.query.get_or_404(sid)
            # cascade by query
            HRPerformanceQuestion.query.filter_by(section_id=s.id).delete()
            db.session.delete(s)
            _portal_audit("HR_PERF_SECTION_DELETE", note=str(sid), target_type='hr_perf_section', target_id=sid)
            db.session.commit()
            flash("تم حذف القسم.", "success")
            return redirect(url_for('portal.portal_admin_hr_perf_form_edit', form_id=f.id))

        if action == 'add_question':
            sid = int(request.form.get('q_section_id') or 0)
            prompt = (request.form.get('prompt') or '').strip()
            q_type = (request.form.get('q_type') or 'RATING_1_5').strip().upper()
            is_required = True if (request.form.get('is_required') == '1') else False
            weight = float(request.form.get('weight') or 1.0)
            order_no = int(request.form.get('order_no') or 0)
            if not sid or not prompt:
                flash("القسم ونص السؤال مطلوبان.", "danger")
                return redirect(url_for('portal.portal_admin_hr_perf_form_edit', form_id=f.id))
            q = HRPerformanceQuestion(section_id=sid, prompt=prompt, q_type=q_type, is_required=is_required, weight=weight, order_no=order_no)
            db.session.add(q)
            _portal_audit("HR_PERF_QUESTION_CREATE", note=prompt, target_type='hr_perf_question', target_id=None)
            db.session.commit()
            flash("تمت إضافة سؤال.", "success")
            return redirect(url_for('portal.portal_admin_hr_perf_form_edit', form_id=f.id))

        if action == 'delete_question':
            qid = int(request.form.get('question_id') or 0)
            q = HRPerformanceQuestion.query.get_or_404(qid)
            db.session.delete(q)
            _portal_audit("HR_PERF_QUESTION_DELETE", note=str(qid), target_type='hr_perf_question', target_id=qid)
            db.session.commit()
            flash("تم حذف السؤال.", "success")
            return redirect(url_for('portal.portal_admin_hr_perf_form_edit', form_id=f.id))

        flash("إجراء غير معروف.", "warning")
        return redirect(url_for('portal.portal_admin_hr_perf_form_edit', form_id=f.id))

    sections, questions_by_section = _perf_questions_for_form(f.id)
    return render_template(
        "portal/admin/hr_perf_form_edit.html",
        form=f,
        sections=sections,
        questions_by_section=questions_by_section,
    )


@portal_bp.route("/admin/hr/performance/cycles", methods=["GET", "POST"])
@login_required
@_perm(HR_PERF_MANAGE)
def portal_admin_hr_perf_cycles():
    if request.method == "POST":
        action = (request.form.get('action') or 'create').strip().lower()

        if action == 'toggle':
            cid = int(request.form.get('id') or 0)
            c = HRPerformanceCycle.query.get_or_404(cid)
            try:
                if c.status == 'ACTIVE':
                    c.status = 'CLOSED'
                else:
                    c.status = 'ACTIVE'
                _portal_audit("HR_PERF_CYCLE_TOGGLE", note=f"{c.name}:{c.status}", target_type='hr_perf_cycle', target_id=c.id)
                db.session.commit()
                flash("تم تحديث حالة الدورة.", "success")
            except Exception:
                db.session.rollback()
                flash("تعذر تحديث الحالة.", "danger")
            return redirect(url_for('portal.portal_admin_hr_perf_cycles'))

        if action == 'delete':
            cid = int(request.form.get('id') or 0)
            c = HRPerformanceCycle.query.get_or_404(cid)
            try:
                used = HRPerformanceAssignment.query.filter_by(cycle_id=c.id).count()
            except Exception:
                used = 0

            try:
                if used and used > 0:
                    c.status = 'CLOSED'
                    _portal_audit("HR_PERF_CYCLE_CLOSE_ON_DELETE", note=c.name, target_type='hr_perf_cycle', target_id=c.id)
                    db.session.commit()
                    flash("لا يمكن حذف دورة لديها تكليفات. تم إغلاقها بدلاً من الحذف.", "warning")
                    return redirect(url_for('portal.portal_admin_hr_perf_cycles'))

                db.session.delete(c)
                _portal_audit("HR_PERF_CYCLE_DELETE", note=c.name, target_type='hr_perf_cycle', target_id=c.id)
                db.session.commit()
                flash("تم حذف الدورة.", "success")
            except Exception:
                db.session.rollback()
                flash("تعذر الحذف.", "danger")
            return redirect(url_for('portal.portal_admin_hr_perf_cycles'))

        # create
        name = (request.form.get('name') or '').strip()
        form_id = int(request.form.get('form_id') or 0)
        start_date = (request.form.get('start_date') or '').strip()
        end_date = (request.form.get('end_date') or '').strip()
        peer_count = int(request.form.get('peer_count') or 2)
        if not name or not form_id:
            flash("الاسم والنموذج مطلوبان.", "danger")
            return redirect(url_for('portal.portal_admin_hr_perf_cycles'))

        c = HRPerformanceCycle(
            name=name,
            form_id=form_id,
            start_date=start_date or None,
            end_date=end_date or None,
            status='DRAFT',
            peer_count=max(0, peer_count),
            created_by_id=current_user.id,
        )
        db.session.add(c)
        _portal_audit("HR_PERF_CYCLE_CREATE", note=name, target_type='hr_perf_cycle', target_id=None)
        try:
            db.session.commit()
            flash("تم إنشاء الدورة.", "success")
        except Exception:
            db.session.rollback()
            flash("تعذر إنشاء الدورة.", "danger")
        return redirect(url_for('portal.portal_admin_hr_perf_cycles'))

    forms = HRPerformanceForm.query.filter(HRPerformanceForm.is_active.is_(True)).order_by(HRPerformanceForm.id.desc()).all()
    cycles = HRPerformanceCycle.query.order_by(HRPerformanceCycle.id.desc()).all()
    return render_template("portal/admin/hr_perf_cycles.html", forms=forms, cycles=cycles)


@portal_bp.route("/admin/hr/performance/cycles/<int:cycle_id>")
@login_required
@_perm(HR_PERF_MANAGE)
def portal_admin_hr_perf_cycle_detail(cycle_id: int):
    c = HRPerformanceCycle.query.get_or_404(cycle_id)
    assignments = HRPerformanceAssignment.query.filter_by(cycle_id=c.id).order_by(HRPerformanceAssignment.id.desc()).all()

    stats = {"total": 0, "submitted": 0, "pending": 0}
    try:
        stats["total"] = len(assignments)
        stats["submitted"] = len([a for a in assignments if a.status == 'SUBMITTED'])
        stats["pending"] = stats["total"] - stats["submitted"]
    except Exception:
        pass

    return render_template(
        "portal/admin/hr_perf_cycle_detail.html",
        cycle=c,
        assignments=assignments,
        stats=stats,
    )


@portal_bp.route("/admin/hr/performance/cycles/<int:cycle_id>/set-status", methods=["POST"])
@login_required
@_perm(HR_PERF_MANAGE)
def portal_admin_hr_perf_cycle_set_status(cycle_id: int):
    c = HRPerformanceCycle.query.get_or_404(cycle_id)
    status = (request.form.get('status') or '').strip().upper()
    if status not in ('DRAFT','ACTIVE','CLOSED'):
        flash("حالة غير صالحة.", "danger")
        return redirect(url_for('portal.portal_admin_hr_perf_cycle_detail', cycle_id=c.id))
    c.status = status
    _portal_audit("HR_PERF_CYCLE_STATUS", note=f"{c.name}:{status}", target_type='hr_perf_cycle', target_id=c.id)
    db.session.commit()
    flash("تم تحديث الحالة.", "success")
    return redirect(url_for('portal.portal_admin_hr_perf_cycle_detail', cycle_id=c.id))


@portal_bp.route("/admin/hr/performance/cycles/<int:cycle_id>/generate", methods=["POST"])
@login_required
@_perm(HR_PERF_MANAGE)
def portal_admin_hr_perf_cycle_generate(cycle_id: int):
    """Generate 360 assignments for all users (best-effort).

    Creates:
      - SELF per user
      - MANAGER per user (if resolvable)
      - PEER per user (up to peer_count)
    """
    c = HRPerformanceCycle.query.get_or_404(cycle_id)

    if c.status == 'CLOSED':
        flash("الدورة مغلقة.", "warning")
        return redirect(url_for('portal.portal_admin_hr_perf_cycle_detail', cycle_id=c.id))

    peer_count = int(c.peer_count or 0)
    created = 0

    users = User.query.order_by(User.id.asc()).all()
    existing = set()
    for a in HRPerformanceAssignment.query.filter_by(cycle_id=c.id).all():
        existing.add((a.evaluatee_user_id, a.evaluator_user_id, a.evaluator_type))

    for u in users:
        if not u or not getattr(u, 'id', None):
            continue
        uid = int(u.id)

        # SELF
        key = (uid, uid, 'SELF')
        if key not in existing:
            db.session.add(HRPerformanceAssignment(
                cycle_id=c.id,
                evaluatee_user_id=uid,
                evaluator_user_id=uid,
                evaluator_type='SELF',
                status='PENDING',
                created_by_id=current_user.id,
            ))
            existing.add(key)
            created += 1

        # MANAGER
        mgr = _find_direct_manager_for_user(uid)
        if mgr and getattr(mgr, 'id', None):
            mid = int(mgr.id)
            key = (uid, mid, 'MANAGER')
            if key not in existing:
                db.session.add(HRPerformanceAssignment(
                    cycle_id=c.id,
                    evaluatee_user_id=uid,
                    evaluator_user_id=mid,
                    evaluator_type='MANAGER',
                    status='PENDING',
                    created_by_id=current_user.id,
                ))
                existing.add(key)
                created += 1

        # PEERS
        exclude = {uid}
        if mgr and getattr(mgr, 'id', None):
            exclude.add(int(mgr.id))
        peers = _pick_peer_candidates(uid, peer_count, exclude_ids=exclude)
        for p in peers:
            pid = int(p.id)
            key = (uid, pid, 'PEER')
            if key in existing:
                continue
            db.session.add(HRPerformanceAssignment(
                cycle_id=c.id,
                evaluatee_user_id=uid,
                evaluator_user_id=pid,
                evaluator_type='PEER',
                status='PENDING',
                created_by_id=current_user.id,
            ))
            existing.add(key)
            created += 1

    _portal_audit("HR_PERF_ASSIGNMENTS_GENERATE", note=f"cycle={c.id},created={created}", target_type='hr_perf_cycle', target_id=c.id)

    try:
        db.session.commit()
        flash(f"تم توليد التكليفات. (الجديد: {created})", "success")
    except Exception:
        db.session.rollback()
        flash("تعذر توليد التكليفات.", "danger")

    return redirect(url_for('portal.portal_admin_hr_perf_cycle_detail', cycle_id=c.id))


@portal_bp.route("/admin/hr/performance/cycles/<int:cycle_id>/export.csv")
@login_required
@_perm_any(HR_PERF_MANAGE, HR_PERF_EXPORT)
def portal_admin_hr_perf_cycle_export_csv(cycle_id: int):
    c = HRPerformanceCycle.query.get_or_404(cycle_id)
    assigns = HRPerformanceAssignment.query.filter_by(cycle_id=c.id).order_by(HRPerformanceAssignment.id.asc()).all()

    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["cycle_id","cycle_name","evaluatee_id","evaluatee_name","evaluator_id","evaluator_name","type","status","score","submitted_at"])
    for a in assigns:
        w.writerow([
            c.id,
            c.name,
            a.evaluatee_user_id,
            (a.evaluatee.full_name if a.evaluatee else ''),
            a.evaluator_user_id,
            (a.evaluator.full_name if a.evaluator else ''),
            a.evaluator_type,
            a.status,
            a.score_total if a.score_total is not None else '',
            a.submitted_at.isoformat() if a.submitted_at else '',
        ])

    data = buf.getvalue().encode('utf-8-sig')
    return send_file(
        BytesIO(data),
        mimetype='text/csv',
        as_attachment=True,
        download_name=f"hr_perf_cycle_{c.id}.csv",
    )


# =========================================================
# Portal Admin: Reports & Compliance (light)
# =========================================================


@portal_bp.route("/admin/hr/reports")
@login_required
@_perm_any(PORTAL_REPORTS_READ, PORTAL_REPORTS_EXPORT)
def portal_admin_hr_reports_dashboard():
    """A lightweight KPI dashboard for HR (best-effort with available data)."""

    # Date scope: today
    today = date.today()
    day_str = today.strftime('%Y-%m-%d')

    employees_count = _safe_count(User.query)

    # Pending HR requests
    pending_leaves = _safe_count(HRLeaveRequest.query.filter(HRLeaveRequest.status == 'PENDING'))
    pending_perms = _safe_count(HRPermissionRequest.query.filter(HRPermissionRequest.status == 'PENDING'))

    # Attendance: absent/late today (best-effort)
    absent_today = 0
    late_today = 0

    try:
        users = User.query.order_by(User.id.asc()).all()
        for u in users:
            if not u or not getattr(u, 'id', None):
                continue
            sched = _effective_schedule_for_user(int(u.id), day_str)
            if not sched:
                continue

            weekday = today.weekday()  # Mon=0..Sun=6

            # Default workdays: Sun-Thu
            def_is_workday = weekday in (6, 0, 1, 2, 3)

            # If we can find policy from assignment, use it
            pol = None
            try:
                wa = (WorkAssignment.query.filter_by(user_id=u.id)
                      .filter(or_(WorkAssignment.start_date == None, WorkAssignment.start_date <= day_str))
                      .filter(or_(WorkAssignment.end_date == None, WorkAssignment.end_date >= day_str))
                      .order_by(WorkAssignment.id.desc()).first())
                pol = wa.policy if wa else None
            except Exception:
                pol = None

            is_workday = def_is_workday
            try:
                if pol and (pol.days_policy == 'FIXED') and (pol.fixed_days_mask is not None):
                    is_workday = bool(int(pol.fixed_days_mask) & (1 << weekday))
            except Exception:
                pass

            if not is_workday:
                continue

            # Check first check-in
            evs = (
                AttendanceEvent.query
                .filter(AttendanceEvent.user_id == u.id)
                .filter(AttendanceEvent.event_dt >= f"{day_str} 00:00:00")
                .filter(AttendanceEvent.event_dt <= f"{day_str} 23:59:59")
                .filter(AttendanceEvent.event_type.in_(['I','IN','CHECKIN']))
                .order_by(AttendanceEvent.event_dt.asc())
                .all()
            )

            if not evs:
                absent_today += 1
                continue

            first_dt = evs[0].event_dt
            # Compute schedule start and grace
            start_t = None
            grace = 0
            try:
                grace = int((pol.start_grace_min if pol else 0) or 0)
            except Exception:
                grace = 0

            try:
                if sched.schedule_type == 'FIXED':
                    start_t = sched.start_time
                elif sched.schedule_type == 'SHIFT':
                    drow = WorkScheduleDay.query.filter_by(schedule_id=sched.id, weekday=weekday).first()
                    start_t = drow.start_time if drow else None
            except Exception:
                start_t = None

            if start_t and first_dt:
                try:
                    # first_dt might be datetime or string
                    if isinstance(first_dt, str):
                        fdt = datetime.fromisoformat(first_dt)
                    else:
                        fdt = first_dt
                    sched_start = datetime.combine(today, start_t)
                    if fdt > (sched_start + timedelta(minutes=grace)):
                        late_today += 1
                except Exception:
                    pass
    except Exception:
        pass

    kpis = {
        'employees': employees_count,
        'pending_leaves': pending_leaves,
        'pending_permissions': pending_perms,
        'absent_today': absent_today,
        'late_today': late_today,
        'day_str': day_str,
    }

    return render_template('portal/admin/hr_reports_dashboard.html', kpis=kpis)




# -------------------------
# Audit helpers (Portal Timeline)
# -------------------------

def _audit_guess_module(action: str | None, target_type: str | None) -> str:
    a = (action or "").upper()
    t = (target_type or "").upper()

    if a.startswith("HR_") or t.startswith("HR_") or t in ("EMPLOYEE_FILE", "EMPLOYEE_ATTACHMENT", "LEAVE_REQUEST", "MISSION_REQUEST", "PERMISSION_REQUEST"):
        return "HR"
    if a.startswith("INV_") or a.startswith("STORE_") or t.startswith("INV_") or t.startswith("STORE_"):
        return "STORE"
    if a.startswith("TRANSPORT_") or t.startswith("TRANSPORT_"):
        return "TRANSPORT"
    if a.startswith("CORR_") or t.startswith("CORR_") or t in ("INBOUND_MAIL", "OUTBOUND_MAIL"):
        return "CORR"
    if a.startswith("PORTAL_") or a.startswith("ACCESS_") or t in ("PORTAL_ACCESS_REQUEST", "ROLE", "USER"):
        return "PERMS"
    return "OTHER"


def _audit_target_url_label(r):
    """Best-effort: returns (url,label) for an AuditLog row."""
    t = (getattr(r, 'target_type', None) or '').upper()
    tid = getattr(r, 'target_id', None)
    if not tid:
        return (None, None)

    try:
        # HR: employee file
        if t in ("EMPLOYEE_FILE",):
            ef = EmployeeFile.query.get(int(tid))
            label = f"ملف موظف: {(getattr(ef, 'employee_no', None) or '')} {(getattr(ef, 'full_name_quad', None) or '')}".strip() if ef else f"ملف موظف #{tid}"
            return (url_for('portal.hr_employee_file', user_id=int(tid)), label)

        if t in ("EMPLOYEE_ATTACHMENT",):
            att = EmployeeAttachment.query.get(int(tid))
            if att:
                label = f"مرفق موظف: {att.original_name}"
                return (url_for('portal.hr_employee_attachment_edit', user_id=int(att.user_id), att_id=int(att.id)), label)
            return (None, f"مرفق موظف #{tid}")

        # Store: files
        if t in ("STORE_FILE",):
            sf = StoreFile.query.get(int(tid))
            label = f"ملف مستودع: {sf.original_name}" if sf else f"ملف مستودع #{tid}"
            return (url_for('portal.store_file_view', file_id=int(tid)), label)

        # Inventory vouchers
        if t in ("INV_ISSUE_VOUCHER",):
            v = InvIssueVoucher.query.get(int(tid))
            label = f"سند صرف: {v.voucher_no}" if v else f"سند صرف #{tid}"
            return (url_for('portal.inventory_issue_voucher_view', v_id=int(tid)), label)
        if t in ("INV_INBOUND_VOUCHER",):
            v = InvInboundVoucher.query.get(int(tid))
            label = f"سند إدخال: {v.voucher_no}" if v else f"سند إدخال #{tid}"
            return (url_for('portal.inventory_inbound_voucher_view', v_id=int(tid)), label)
        if t in ("INV_SCRAP_VOUCHER",):
            v = InvScrapVoucher.query.get(int(tid))
            label = f"سند إتلاف: {v.voucher_no}" if v else f"سند إتلاف #{tid}"
            return (url_for('portal.inventory_scrap_voucher_view', v_id=int(tid)), label)
        if t in ("INV_RETURN_VOUCHER",):
            v = InvReturnVoucher.query.get(int(tid))
            label = f"سند إرجاع: {v.voucher_no}" if v else f"سند إرجاع #{tid}"
            return (url_for('portal.inventory_return_voucher_view', v_id=int(tid)), label)
        if t in ("INV_STOCKTAKE_VOUCHER", "INV_STOCKTAKE"):
            v = InvStocktakeVoucher.query.get(int(tid))
            label = f"سند جرد: {v.voucher_no}" if v else f"سند جرد #{tid}"
            return (url_for('portal.inventory_stocktake_voucher_view', v_id=int(tid)), label)
        if t in ("INV_CUSTODY_VOUCHER", "INV_CUSTODY"):
            v = InvCustodyVoucher.query.get(int(tid))
            label = f"سند عهدة: {v.voucher_no}" if v else f"سند عهدة #{tid}"
            return (url_for('portal.inventory_custody_voucher_view', v_id=int(tid)), label)

        # Correspondence
        if t in ("INBOUND_MAIL", "CORR_INBOUND"):
            m = InboundMail.query.get(int(tid))
            ref = (m.ref_no or f"#{m.id}") if m else f"#{tid}"
            label = f"وارد {ref}" if m else f"وارد {ref}"
            return (url_for('portal.inbound_view', inbound_id=int(tid)), label)
        if t in ("OUTBOUND_MAIL", "CORR_OUTBOUND"):
            m = OutboundMail.query.get(int(tid))
            ref = (m.ref_no or f"#{m.id}") if m else f"#{tid}"
            label = f"صادر {ref}" if m else f"صادر {ref}"
            return (url_for('portal.outbound_view', outbound_id=int(tid)), label)

        # Transport
        if t == "TRANSPORT_PERMIT":
            return (url_for('portal.transport_permit_view', permit_id=int(tid)), f"إذن حركة #{tid}")
        if t == "TRANSPORT_TRIP":
            return (url_for('portal.transport_trips', ) + f"?highlight={tid}", f"رحلة #{tid}")
        if t == "TRANSPORT_TASK":
            return (url_for('portal.transport_tasks', ) + f"?highlight={tid}", f"مهمة سائق #{tid}")

        # Permissions
        if t == "PORTAL_ACCESS_REQUEST":
            return (url_for('portal.admin_access_request_view', req_id=int(tid)), f"طلب صلاحية #{tid}")
        if t == "USER":
            return (url_for('portal.portal_admin_permissions', scope='user', user_id=int(tid)), f"مستخدم #{tid}")
        if t == "ROLE":
            from models import Role
            role = Role.query.get(int(tid))
            code = (role.code if role else str(tid))
            label = (role.label if role else f"Role #{tid}")
            return (url_for('portal.portal_admin_permissions', scope='role', role=code), label)

    except Exception:
        return (None, None)

    return (None, None)

@portal_bp.route("/admin/compliance", methods=["GET"])
@login_required
@_perm_any(PORTAL_AUDIT_READ, PORTAL_ADMIN_PERMISSIONS_MANAGE)
def portal_admin_compliance():
    """Portal compliance view (audit logs + quick filters)."""
    q_action = (request.args.get('action') or '').strip().upper()
    q_user = (request.args.get('user') or '').strip()
    q_from = (request.args.get('from') or '').strip()  # YYYY-MM-DD
    q_to = (request.args.get('to') or '').strip()

    qry = AuditLog.query

    if q_action:
        qry = qry.filter(AuditLog.action.ilike(f"%{q_action}%"))

    if q_user:
        like = f"%{q_user}%"
        qry = qry.join(User, AuditLog.user_id == User.id).filter(or_(User.email.ilike(like), User.name.ilike(like)))

    if q_from:
        qry = qry.filter(AuditLog.created_at >= f"{q_from} 00:00:00")
    if q_to:
        qry = qry.filter(AuditLog.created_at <= f"{q_to} 23:59:59")

    logs = qry.order_by(AuditLog.id.desc()).limit(500).all()

    # Users for filter

    # enrich rows (target links/labels)
    for r in logs:
        try:
            r._module = _audit_guess_module(getattr(r, 'action', None), getattr(r, 'target_type', None))
            u, lbl = _audit_target_url_label(r)
            r._target_url = u
            r._target_label = lbl
        except Exception:
            pass

    # Users for filter
    name_expr = func.nullif(func.trim(User.name), "")
    users = (
        User.query
        .order_by(func.lower(func.coalesce(name_expr, User.email)).asc(), User.id.asc())
        .limit(200)
        .all()
    )

    return render_template('portal/admin/compliance.html', logs=logs, users=users, q_action=q_action, q_user=q_user, q_from=q_from, q_to=q_to)

@portal_bp.route("/admin/timeline", methods=["GET"])
@login_required
@_perm_any(PORTAL_AUDIT_READ, PORTAL_ADMIN_PERMISSIONS_MANAGE)
def portal_admin_timeline():
    selected_module = (request.args.get('module') or 'ALL').strip().upper()
    q = (request.args.get('q') or '').strip()
    q_action = (request.args.get('action') or '').strip().upper()
    q_user = (request.args.get('user') or '').strip()
    q_from = (request.args.get('from') or '').strip()
    q_to = (request.args.get('to') or '').strip()
    try:
        page = int((request.args.get('page') or '1').strip())
    except Exception:
        page = 1
    if page < 1:
        page = 1

    per_page = 40

    qry = AuditLog.query

    # module filter
    if selected_module and selected_module != 'ALL':
        m = selected_module
        if m == 'HR':
            qry = qry.filter(or_(AuditLog.action.ilike('HR_%'), AuditLog.target_type.ilike('HR_%'), AuditLog.target_type.in_(['EMPLOYEE_FILE','EMPLOYEE_ATTACHMENT','LEAVE_REQUEST','MISSION_REQUEST','PERMISSION_REQUEST'])))
        elif m == 'STORE':
            qry = qry.filter(or_(AuditLog.action.ilike('INV_%'), AuditLog.action.ilike('STORE_%'), AuditLog.target_type.ilike('INV_%'), AuditLog.target_type.ilike('STORE_%')))
        elif m == 'TRANSPORT':
            qry = qry.filter(or_(AuditLog.action.ilike('TRANSPORT_%'), AuditLog.target_type.ilike('TRANSPORT_%')))
        elif m == 'CORR':
            qry = qry.filter(or_(AuditLog.action.ilike('CORR_%'), AuditLog.target_type.ilike('CORR_%'), AuditLog.target_type.in_(['INBOUND_MAIL','OUTBOUND_MAIL'])))
        elif m == 'PERMS':
            qry = qry.filter(or_(AuditLog.action.ilike('PORTAL_%'), AuditLog.action.ilike('ACCESS_%'), AuditLog.target_type.in_(['PORTAL_ACCESS_REQUEST','ROLE','USER'])))
        else:
            pass

    if q_action:
        qry = qry.filter(AuditLog.action.ilike(f"%{q_action}%"))

    if q_user:
        like = f"%{q_user}%"
        qry = qry.join(User, AuditLog.user_id == User.id).filter(or_(User.email.ilike(like), User.name.ilike(like)))

    if q:
        like = f"%{q}%"
        qry = qry.filter(or_(AuditLog.action.ilike(like), AuditLog.note.ilike(like), AuditLog.target_type.ilike(like)))

    if q_from:
        qry = qry.filter(AuditLog.created_at >= f"{q_from} 00:00:00")
    if q_to:
        qry = qry.filter(AuditLog.created_at <= f"{q_to} 23:59:59")

    rows = qry.order_by(AuditLog.id.desc()).offset((page-1)*per_page).limit(per_page + 1).all()
    has_next = len(rows) > per_page
    logs = rows[:per_page]
    has_prev = page > 1

    for r in logs:
        try:
            r._module = _audit_guess_module(getattr(r, 'action', None), getattr(r, 'target_type', None))
            u, lbl = _audit_target_url_label(r)
            r._target_url = u
            r._target_label = lbl
        except Exception:
            pass

    from types import SimpleNamespace
    modules = [
        SimpleNamespace(key='ALL', label='الكل', icon='bi-grid'),
        SimpleNamespace(key='HR', label='الموارد البشرية', icon='bi-people'),
        SimpleNamespace(key='STORE', label='المستودع', icon='bi-box-seam'),
        SimpleNamespace(key='TRANSPORT', label='الحركة', icon='bi-truck'),
        SimpleNamespace(key='CORR', label='الصادر والوارد', icon='bi-envelope-paper'),
        SimpleNamespace(key='PERMS', label='الصلاحيات', icon='bi-shield-lock'),
    ]

    return render_template(
        'portal/admin/timeline.html',
        modules=modules,
        selected_module=selected_module,
        logs=logs,
        q=q,
        q_action=q_action,
        q_user=q_user,
        q_from=q_from,
        q_to=q_to,
        page=page,
        has_prev=has_prev,
        has_next=has_next,
    )



@portal_bp.route("/admin/compliance/export.csv")
@login_required
@_perm_any(PORTAL_AUDIT_READ, PORTAL_REPORTS_EXPORT, PORTAL_ADMIN_PERMISSIONS_MANAGE)
def portal_admin_compliance_export_csv():
    q_action = (request.args.get('action') or '').strip().upper()
    q_user = (request.args.get('user') or '').strip()
    q_from = (request.args.get('from') or '').strip()  # YYYY-MM-DD
    q_to = (request.args.get('to') or '').strip()

    qry = AuditLog.query
    if q_action:
        qry = qry.filter(AuditLog.action.ilike(f"%{q_action}%"))
    if q_user:
        like = f"%{q_user}%"
        qry = qry.join(User, AuditLog.user_id == User.id).filter(or_(User.email.ilike(like), User.name.ilike(like)))
    if q_from:
        qry = qry.filter(AuditLog.created_at >= f"{q_from} 00:00:00")
    if q_to:
        qry = qry.filter(AuditLog.created_at <= f"{q_to} 23:59:59")

    rows = qry.order_by(AuditLog.id.desc()).limit(5000).all()

    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["id","created_at","user_id","user","action","note","target_type","target_id","on_behalf_of","delegation_id"])
    for r in rows:
        w.writerow([
            r.id,
            r.created_at.isoformat() if r.created_at else '',
            r.user_id,
            (r.user.full_name if getattr(r, 'user', None) else ''),
            r.action,
            r.note or '',
            r.target_type or '',
            r.target_id or '',
            r.on_behalf_of_id or '',
            r.delegation_id or '',
        ])

    data = buf.getvalue().encode('utf-8-sig')
    return send_file(BytesIO(data), mimetype='text/csv', as_attachment=True, download_name='portal_compliance.csv')


# HR_LEAVES_MISSIONS_EVENTS_PATCH_V1

# ===== Helpers =====
def _weekly_mask() -> int:
    key = 'HR_WEEKLY_HOLIDAYS_MASK'
    row = SystemSetting.query.filter_by(key=key).first()
    try:
        return int((row.value or '0').strip()) if row else 0
    except Exception:
        return 0

def _is_weekly_off(d: date, mask: int) -> bool:
    return bool(mask & (1 << d.weekday()))

def _is_official_day_off(day_s: str) -> bool:
    from models import HROfficialOccasionRange
    r1 = HROfficialOccasion.query.filter_by(day=day_s).first()
    if r1 and bool(r1.is_day_off):
        return True
    ranges = HROfficialOccasionRange.query.filter(
        HROfficialOccasionRange.start_day <= day_s,
        HROfficialOccasionRange.end_day >= day_s
    ).all()
    return any(bool(x.is_day_off) for x in ranges)

def _calc_leave_days_excluding_off(start_s: str, end_s: str) -> int:
    d0 = _parse_yyyy_mm_dd(start_s)
    d1 = _parse_yyyy_mm_dd(end_s)
    if not d0 or not d1:
        return 0
    if d1 < d0:
        d0, d1 = d1, d0
    mask = _weekly_mask()
    n = 0
    cur = d0
    while cur <= d1:
        ds = _as_yyyy_mm_dd(cur)
        if (not _is_weekly_off(cur, mask)) and (not _is_official_day_off(ds)):
            n += 1
        cur += timedelta(days=1)
    return n

def _mission_upload_dir(mission_id: int) -> Path:
    base = Path(current_app.instance_path) / "uploads" / "missions" / str(mission_id)
    base.mkdir(parents=True, exist_ok=True)
    return base

def _leaves_upload_dir(req_id: int) -> Path:
    base = Path(current_app.instance_path) / "uploads" / "leaves" / str(req_id)
    base.mkdir(parents=True, exist_ok=True)
    return base

def _ensure_status_defs(entity: str):
    from models import HRStatusDef
    entity = (entity or '').strip().upper()
    if HRStatusDef.query.filter_by(entity=entity).count() == 0:
        defaults = []
        if entity == "LEAVE":
            defaults = [
                ("NEW","جديدة","New",10),
                ("APPROVED_BY_MANAGER","معتمدة من المسؤول","Approved by manager",20),
                ("REJECTED_BY_MANAGER","مرفوضة من المسؤول","Rejected by manager",30),
                ("CONFIRMED","مثبتة","Confirmed",40),
                ("CANCELLED","ملغاة","Cancelled",90),
            ]
        elif entity == "MISSION":
            defaults = [
                ("NEW","جديدة","New",10),
                ("CONFIRMED","مثبتة","Confirmed",40),
                ("CANCELLED","ملغاة","Cancelled",90),
            ]
        for code, ar, en, so in defaults:
            db.session.add(HRStatusDef(entity=entity, code=code, name_ar=ar, name_en=en, sort_order=so, is_active=True))
        db.session.commit()
    return HRStatusDef.query.filter_by(entity=entity, is_active=True).order_by(HRStatusDef.sort_order.asc()).all()

def _ensure_occasion_types():
    from models import HROfficialOccasionType
    if HROfficialOccasionType.query.count() == 0:
        defaults = [
            ("عطلة رسمية","Official Holiday",True,10),
            ("مناسبة وطنية","National Event",False,20),
            ("مناسبة دينية","Religious Event",True,30),
        ]
        for ar,en,isoff,so in defaults:
            db.session.add(HROfficialOccasionType(name_ar=ar, name_en=en, is_day_off_default=isoff, sort_order=so, is_active=True))
        db.session.commit()
    return HROfficialOccasionType.query.filter_by(is_active=True).order_by(HROfficialOccasionType.sort_order.asc()).all()

def _load_work_location_lookups():
    govs = _hr_lookup_items_for_category('WORK_GOVERNORATE')
    locs = _hr_lookup_items_for_category('WORK_LOCATION')
    return {"govs": govs, "locs": locs}

def _save_mission_attachment(mission_id: int, f):
    from models import HROfficialMissionAttachment
    folder = _mission_upload_dir(mission_id)
    orig = (f.filename or '').strip()
    stored = f"{uuid.uuid4().hex}_{orig}"
    full = folder / stored
    f.save(str(full))
    att = HROfficialMissionAttachment(
        mission_id=mission_id,
        original_name=orig,
        stored_name=stored,
        content_type=getattr(f, "mimetype", None),
        size_bytes=full.stat().st_size if full.exists() else None,
        uploaded_by_id=getattr(current_user, "id", None),
    )
    db.session.add(att)
    db.session.commit()

# ===== Leaves Admin =====
@portal_bp.route('/hr/leaves/admin', methods=['GET'])
@login_required
@_perm_any(HR_READ, HR_REQUESTS_VIEW_ALL, HR_MASTERDATA_MANAGE)
def hr_leaves_admin_log():
    from models import EmployeeFile
    q = HRLeaveRequest.query
    user_id = (request.args.get('user_id') or '').strip()
    leave_type_id = (request.args.get('leave_type_id') or '').strip()
    admin_status_id = (request.args.get('admin_status_id') or '').strip()
    leave_place = (request.args.get('leave_place') or '').strip()

    if user_id.isdigit():
        q = q.filter(HRLeaveRequest.user_id == int(user_id))
    if leave_type_id.isdigit():
        q = q.filter(HRLeaveRequest.leave_type_id == int(leave_type_id))
    if admin_status_id.isdigit():
        q = q.filter(HRLeaveRequest.admin_status_id == int(admin_status_id))
    if leave_place in ("INTERNAL","EXTERNAL"):
        q = q.filter(HRLeaveRequest.leave_place == leave_place)

    rows = q.order_by(HRLeaveRequest.start_date.desc()).limit(500).all()
    uids = [r.user_id for r in rows]
    files = EmployeeFile.query.filter(EmployeeFile.user_id.in_(uids)).all() if uids else []
    file_map = {f.user_id: f for f in files}

    return render_template(
        'portal/hr/leaves_admin_log.html',
        rows=rows,
        file_map=file_map,
        users=_list_hr_users(),
        leave_types=HRLeaveType.query.filter_by(is_active=True).order_by(HRLeaveType.id.asc()).all(),
        status_defs=_ensure_status_defs("LEAVE"),
        can_manage=_hr_can_manage(),
        filters=dict(user_id=user_id, leave_type_id=leave_type_id, admin_status_id=admin_status_id, leave_place=leave_place),
    )

@portal_bp.route('/hr/leaves/admin/new', methods=['GET','POST'])
@login_required
@_perm_any(HR_READ, HR_REQUESTS_VIEW_ALL, HR_MASTERDATA_MANAGE)
def hr_leaves_admin_new():
    if request.method == 'POST':
        if not _hr_can_manage():
            abort(403)

        user_id = (request.form.get('user_id') or '').strip()
        leave_type_id = (request.form.get('leave_type_id') or '').strip()
        start_date = (request.form.get('start_date') or '').strip()
        end_date = (request.form.get('end_date') or '').strip()
        days_s = (request.form.get('days') or '').strip()
        note = (request.form.get('note') or '').strip()
        leave_place = (request.form.get('leave_place') or '').strip()
        admin_status_id = (request.form.get('admin_status_id') or '').strip()

        if not user_id.isdigit() or not leave_type_id.isdigit():
            flash('اختر موظفاً ونوع الإجازة.', 'danger')
            return redirect(url_for('portal.hr_leaves_admin_new'))
        if not _parse_yyyy_mm_dd(start_date) or not _parse_yyyy_mm_dd(end_date):
            flash('التاريخ غير صحيح.', 'danger')
            return redirect(url_for('portal.hr_leaves_admin_new'))

        auto_days = _calc_leave_days_excluding_off(start_date, end_date)
        try:
            days_val = int(days_s) if days_s else auto_days
        except Exception:
            days_val = auto_days


        lt = None
        try:
            lt = HRLeaveType.query.get(int(leave_type_id))
        except Exception:
            lt = None
        is_external = bool(getattr(lt, 'is_external', False)) if lt else False

        # External leave fields (optional)
        travel_country = (request.form.get('travel_country') or '').strip() or None
        travel_city = (request.form.get('travel_city') or '').strip() or None
        travel_address = (request.form.get('travel_address') or '').strip() or None
        travel_contact_phone = (request.form.get('travel_contact_phone') or '').strip() or None
        travel_purpose = (request.form.get('travel_purpose') or '').strip() or None
        border_crossing = (request.form.get('border_crossing') or '').strip() or None

        lp = (leave_place or '').strip().upper()
        leave_place_val = lp if lp in ('INTERNAL','EXTERNAL') else None
        if is_external and not leave_place_val:
            leave_place_val = 'EXTERNAL'

        try:
            sdef_id = int(admin_status_id) if admin_status_id.isdigit() else None
        except Exception:
            sdef_id = None

        row = HRLeaveRequest(
            user_id=int(user_id),
            leave_type_id=int(leave_type_id),
            start_date=start_date,
            end_date=end_date,
            days=days_val,
            note=note or None,
            leave_place=leave_place_val,
            travel_country=travel_country if is_external else None,
            travel_city=travel_city if is_external else None,
            travel_address=travel_address if is_external else None,
            travel_contact_phone=travel_contact_phone if is_external else None,
            travel_purpose=travel_purpose if is_external else None,
            border_crossing=border_crossing if is_external else None,
            entered_by="ADMIN",
            created_by_id=getattr(current_user,'id',None),
            admin_status_id=sdef_id,
            status="APPROVED",
            decided_at=datetime.utcnow(),
            decided_by_id=getattr(current_user,'id',None),
        )
        db.session.add(row)
        db.session.commit()
        files = []
        try:
            files = request.files.getlist('attachments') or []
        except Exception:
            files = []
        valid = [f for f in files if f and (getattr(f,'filename','') or '').strip()]
        if not valid:
            f1 = request.files.get('attachment')
            if f1 and (getattr(f1,'filename','') or '').strip():
                valid = [f1]

        if valid:
            folder = _leaves_upload_dir(row.id)
            for f in valid:
                orig = (f.filename or '').strip()
                stored = f"{uuid.uuid4().hex}_{orig}"
                f.save(str(folder / stored))
                att = HRLeaveAttachment(
                    request_id=row.id,
                    doc_type="ADMIN_DOC",
                    original_name=orig,
                    stored_name=stored,
                    uploaded_by_id=getattr(current_user,'id',None),
                )
                db.session.add(att)
            db.session.commit()

        flash('تم حفظ الإجازة.', 'success')
        return redirect(url_for('portal.hr_leaves_admin_log'))

    leave_types = HRLeaveType.query.filter_by(is_active=True).order_by(HRLeaveType.id.asc()).all()
    types_meta = {str(t.id): {"is_external": bool(getattr(t, "is_external", False)),
                              "requires_documents": bool(getattr(t, "requires_documents", False)),
                              "documents_hint": (getattr(t, "documents_hint", None) or "")} for t in leave_types}

    return render_template(
        'portal/hr/leaves_admin_form.html',
        users=_list_hr_users(),
        leave_types=leave_types,
        types_meta=types_meta,
        status_defs=_ensure_status_defs("LEAVE"),
        today=_as_yyyy_mm_dd(date.today()),
        edit_mode=False,
    )

@portal_bp.route('/hr/leaves/admin/<int:row_id>/edit', methods=['GET','POST'])
@login_required
@_perm_any(HR_READ, HR_REQUESTS_VIEW_ALL, HR_MASTERDATA_MANAGE)
def hr_leaves_admin_edit(row_id: int):
    if not _hr_can_manage():
        abort(403)
    row = HRLeaveRequest.query.get_or_404(row_id)

    if request.method == 'POST':
        row.leave_type_id = int(request.form.get('leave_type_id') or row.leave_type_id)
        row.start_date = (request.form.get('start_date') or row.start_date).strip()
        row.end_date = (request.form.get('end_date') or row.end_date).strip()
        row.note = (request.form.get('note') or '').strip() or None
        lp = (request.form.get('leave_place') or '').strip().upper()
        row.leave_place = lp if lp in ("INTERNAL","EXTERNAL") else None

        # External leave fields (optional)
        lt = None
        try:
            lt = HRLeaveType.query.get(int(row.leave_type_id))
        except Exception:
            lt = None
        is_external = bool(getattr(lt, 'is_external', False)) if lt else False

        row.travel_country = ((request.form.get('travel_country') or '').strip() or None) if is_external else None
        row.travel_city = ((request.form.get('travel_city') or '').strip() or None) if is_external else None
        row.travel_address = ((request.form.get('travel_address') or '').strip() or None) if is_external else None
        row.travel_contact_phone = ((request.form.get('travel_contact_phone') or '').strip() or None) if is_external else None
        row.travel_purpose = ((request.form.get('travel_purpose') or '').strip() or None) if is_external else None
        row.border_crossing = ((request.form.get('border_crossing') or '').strip() or None) if is_external else None

        # If the type is external and place not set, default it
        if is_external and not row.leave_place:
            row.leave_place = 'EXTERNAL'

        days_s = (request.form.get('days') or '').strip()
        if days_s:
            try:
                row.days = int(days_s)
            except Exception:
                pass
        else:
            row.days = _calc_leave_days_excluding_off(row.start_date, row.end_date)

        admin_status_id = (request.form.get('admin_status_id') or '').strip()
        row.admin_status_id = int(admin_status_id) if admin_status_id.isdigit() else None

        db.session.commit()

        files = []
        try:
            files = request.files.getlist('attachments') or []
        except Exception:
            files = []
        valid = [f for f in files if f and (getattr(f,'filename','') or '').strip()]
        if not valid:
            f1 = request.files.get('attachment')
            if f1 and (getattr(f1,'filename','') or '').strip():
                valid = [f1]

        if valid:
            folder = _leaves_upload_dir(row.id)
            for f in valid:
                orig = (f.filename or '').strip()
                stored = f"{uuid.uuid4().hex}_{orig}"
                f.save(str(folder / stored))
                att = HRLeaveAttachment(
                    request_id=row.id,
                    doc_type="ADMIN_DOC",
                    original_name=orig,
                    stored_name=stored,
                    uploaded_by_id=getattr(current_user,'id',None),
                )
                db.session.add(att)
            db.session.commit()

        flash('تم تحديث الإجازة.', 'success')
        return redirect(url_for('portal.hr_leaves_admin_log'))

    leave_types = HRLeaveType.query.filter_by(is_active=True).order_by(HRLeaveType.id.asc()).all()
    types_meta = {str(t.id): {"is_external": bool(getattr(t, "is_external", False)),
                              "requires_documents": bool(getattr(t, "requires_documents", False)),
                              "documents_hint": (getattr(t, "documents_hint", None) or "")} for t in leave_types}

    return render_template(
        'portal/hr/leaves_admin_form.html',
        users=_list_hr_users(),
        leave_types=leave_types,
        types_meta=types_meta,
        status_defs=_ensure_status_defs("LEAVE"),
        today=_as_yyyy_mm_dd(date.today()),
        edit_mode=True,
        row=row,
    )

@portal_bp.route('/hr/leaves/admin/attachments/<int:att_id>/download')
@login_required
@_perm_any(HR_READ, HR_REQUESTS_VIEW_ALL, HR_MASTERDATA_MANAGE)
def hr_leave_attachment_download_admin(att_id: int):
    att = HRLeaveAttachment.query.get_or_404(att_id)
    folder = _leaves_upload_dir(att.request_id)
    return send_from_directory(str(folder), att.stored_name, as_attachment=True, download_name=att.original_name or att.stored_name)

# ===== Missions =====
@portal_bp.route('/hr/missions/new', methods=['GET','POST'])
@login_required
@_perm_any(HR_READ, HR_REQUESTS_VIEW_ALL, HR_MASTERDATA_MANAGE)
def hr_official_mission_new():
    if request.method == 'POST':
        if not _hr_can_manage():
            abort(403)

        user_id = request.form.get('user_id')
        title = (request.form.get('title') or '').strip()
        start_day = (request.form.get('start_day') or '').strip()
        end_day = (request.form.get('end_day') or '').strip()
        days_s = (request.form.get('days') or '').strip()
        destination = (request.form.get('destination') or '').strip()
        note = (request.form.get('note') or '').strip()
        status_def_id = (request.form.get('status_def_id') or '').strip()

        if not title:
            flash('العنوان مطلوب.', 'danger')
            return redirect(url_for('portal.hr_official_mission_new'))
        if not _parse_yyyy_mm_dd(start_day) or not _parse_yyyy_mm_dd(end_day):
            flash('التاريخ غير صحيح.', 'danger')
            return redirect(url_for('portal.hr_official_mission_new'))

        try:
            uid = int(user_id)
        except Exception:
            flash('اختر موظفاً.', 'danger')
            return redirect(url_for('portal.hr_official_mission_new'))

        d0 = _parse_yyyy_mm_dd(start_day)
        d1 = _parse_yyyy_mm_dd(end_day)
        auto_days = (d1 - d0).days + 1 if d0 and d1 else None
        try:
            days_val = int(days_s) if days_s else auto_days
        except Exception:
            days_val = auto_days

        try:
            sdef_id = int(status_def_id) if status_def_id else None
        except Exception:
            sdef_id = None

        row = HROfficialMission(
            user_id=uid,
            title=title,
            start_day=start_day,
            end_day=end_day,
            days=days_val,
            destination=destination or None,
            note=note or None,
            entered_by="ADMIN",
            status_def_id=sdef_id,
            created_by_id=getattr(current_user, 'id', None),
        )
        db.session.add(row)
        db.session.commit()

        f = request.files.get('attachment')
        if f and (f.filename or '').strip():
            _save_mission_attachment(row.id, f)

        flash('تم حفظ المهمة الرسمية.', 'success')
        return redirect(url_for('portal.hr_official_missions'))

    return render_template(
        'portal/hr/missions_new.html',
        users=_list_hr_users(),
        today=_as_yyyy_mm_dd(date.today()),
        status_defs=_ensure_status_defs("MISSION"),
    )

@portal_bp.route('/hr/missions/<int:mission_id>/edit', methods=['GET','POST'])
@login_required
@_perm_any(HR_READ, HR_REQUESTS_VIEW_ALL, HR_MASTERDATA_MANAGE)
def hr_official_mission_edit(mission_id: int):
    if not _hr_can_manage():
        abort(403)
    row = HROfficialMission.query.get_or_404(mission_id)

    if request.method == 'POST':
        row.title = (request.form.get('title') or '').strip()
        row.start_day = (request.form.get('start_day') or '').strip()
        row.end_day = (request.form.get('end_day') or '').strip()
        row.destination = (request.form.get('destination') or '').strip() or None
        row.note = (request.form.get('note') or '').strip() or None

        days_s = (request.form.get('days') or '').strip()
        try:
            row.days = int(days_s) if days_s else row.days
        except Exception:
            pass

        status_def_id = (request.form.get('status_def_id') or '').strip()
        try:
            row.status_def_id = int(status_def_id) if status_def_id else None
        except Exception:
            row.status_def_id = None

        db.session.commit()

        f = request.files.get('attachment')
        if f and (f.filename or '').strip():
            _save_mission_attachment(row.id, f)

        flash('تم تحديث المهمة.', 'success')
        return redirect(url_for('portal.hr_official_missions'))

    return render_template(
        'portal/hr/missions_new.html',
        edit_mode=True,
        row=row,
        users=_list_hr_users(),
        today=_as_yyyy_mm_dd(date.today()),
        status_defs=_ensure_status_defs("MISSION"),
    )

@portal_bp.route('/hr/missions/attachments/<int:att_id>/download')
@login_required
@_perm_any(HR_READ, HR_REQUESTS_VIEW_ALL, HR_MASTERDATA_MANAGE)
def hr_mission_attachment_download(att_id: int):
    from models import HROfficialMissionAttachment
    att = HROfficialMissionAttachment.query.get_or_404(att_id)
    folder = _mission_upload_dir(att.mission_id)
    return send_from_directory(str(folder), att.stored_name, as_attachment=True, download_name=att.original_name or att.stored_name)

@portal_bp.route('/hr/missions/<int:mission_id>/delete', methods=['POST'])
@login_required
@_perm_any(HR_READ, HR_REQUESTS_VIEW_ALL, HR_MASTERDATA_MANAGE)
def hr_official_mission_delete(mission_id: int):
    if not _hr_can_manage():
        abort(403)
    row = HROfficialMission.query.get_or_404(mission_id)
    db.session.delete(row)
    db.session.commit()
    flash('تم حذف المهمة.', 'success')
    return redirect(url_for('portal.hr_official_missions'))

# ===== Occasions =====
@portal_bp.route('/hr/occasions')
@login_required
@_perm_any(HR_READ, HR_REQUESTS_VIEW_ALL, HR_MASTERDATA_MANAGE)
def hr_official_occasions():
    from models import HROfficialOccasionRange
    rows = HROfficialOccasionRange.query.order_by(HROfficialOccasionRange.start_day.desc()).limit(300).all()
    return render_template('portal/hr/occasions_log.html', rows=rows, can_manage=_hr_can_manage())

@portal_bp.route('/hr/occasions/new', methods=['GET','POST'])
@login_required
@_perm_any(HR_READ, HR_REQUESTS_VIEW_ALL, HR_MASTERDATA_MANAGE)
def hr_official_occasion_new():
    from models import HROfficialOccasionRange
    if request.method == 'POST':
        if not _hr_can_manage():
            abort(403)

        gov_id = (request.form.get('work_governorate_lookup_id') or '').strip()
        loc_id = (request.form.get('work_location_lookup_id') or '').strip()
        type_id = (request.form.get('type_id') or '').strip()

        title = (request.form.get('title') or '').strip()
        note = (request.form.get('note') or '').strip()
        start_day = (request.form.get('start_day') or '').strip()
        end_day = (request.form.get('end_day') or '').strip()
        is_day_off = (request.form.get('is_day_off') or '1').strip() == '1'

        if not title or not _parse_yyyy_mm_dd(start_day) or not _parse_yyyy_mm_dd(end_day):
            flash('العنوان والتواريخ مطلوبة.', 'danger')
            return redirect(url_for('portal.hr_official_occasion_new'))

        try:
            type_id_i = int(type_id)
        except Exception:
            flash('اختر نوع المناسبة.', 'danger')
            return redirect(url_for('portal.hr_official_occasion_new'))

        def _int_or_none(x):
            try:
                return int(x) if x else None
            except Exception:
                return None

        row = HROfficialOccasionRange(
            work_governorate_lookup_id=_int_or_none(gov_id),
            work_location_lookup_id=_int_or_none(loc_id),
            type_id=type_id_i,
            title=title,
            note=note or None,
            start_day=start_day,
            end_day=end_day,
            is_day_off=bool(is_day_off),
            created_by_id=getattr(current_user, 'id', None),
        )
        db.session.add(row)
        db.session.commit()
        flash('تم حفظ المناسبة الرسمية.', 'success')
        return redirect(url_for('portal.hr_official_occasions'))

    return render_template(
        'portal/hr/occasions_new.html',
        today=_as_yyyy_mm_dd(date.today()),
        types=_ensure_occasion_types(),
        lookups=_load_work_location_lookups(),
    )

@portal_bp.route('/hr/occasions/<int:row_id>/delete', methods=['POST'])
@login_required
@_perm_any(HR_READ, HR_REQUESTS_VIEW_ALL, HR_MASTERDATA_MANAGE)
def hr_official_occasion_delete(row_id: int):
    from models import HROfficialOccasionRange
    if not _hr_can_manage():
        abort(403)
    row = HROfficialOccasionRange.query.get_or_404(row_id)
    db.session.delete(row)
    db.session.commit()
    flash('تم حذف المناسبة.', 'success')
    return redirect(url_for('portal.hr_official_occasions'))

# ===== Masterdata: Status Defs + Occasion Types =====
@portal_bp.route('/hr/masterdata/status-defs', methods=['GET','POST'])
@login_required
@_perm_any(HR_MASTERDATA_MANAGE, HR_REQUESTS_VIEW_ALL)
def hr_status_defs():
    from models import HRStatusDef
    if request.method == 'POST':
        if not _hr_can_manage():
            abort(403)
        entity = (request.form.get('entity') or '').strip().upper()
        code = (request.form.get('code') or '').strip().upper()
        name_ar = (request.form.get('name_ar') or '').strip()
        name_en = (request.form.get('name_en') or '').strip()
        sort_order = (request.form.get('sort_order') or '100').strip()

        if entity not in ('LEAVE','MISSION') or not code or not name_ar:
            flash('البيانات غير مكتملة.', 'danger')
            return redirect(url_for('portal.hr_status_defs'))

        try:
            so = int(sort_order)
        except Exception:
            so = 100

        row = HRStatusDef.query.filter_by(entity=entity, code=code).first()
        if row:
            row.name_ar = name_ar
            row.name_en = name_en or None
            row.sort_order = so
            row.is_active = True
        else:
            row = HRStatusDef(entity=entity, code=code, name_ar=name_ar, name_en=name_en or None,
                              sort_order=so, is_active=True, created_by_id=getattr(current_user,'id',None))
            db.session.add(row)
        db.session.commit()
        flash('تم حفظ الحالة.', 'success')
        return redirect(url_for('portal.hr_status_defs'))

    rows = HRStatusDef.query.order_by(HRStatusDef.entity.asc(), HRStatusDef.sort_order.asc()).all()
    return render_template('portal/hr/status_defs.html', rows=rows, can_manage=_hr_can_manage())

@portal_bp.route('/hr/masterdata/status-defs/<int:row_id>/toggle', methods=['POST'])
@login_required
@_perm_any(HR_MASTERDATA_MANAGE, HR_REQUESTS_VIEW_ALL)
def hr_status_def_toggle(row_id: int):
    from models import HRStatusDef
    if not _hr_can_manage():
        abort(403)
    row = HRStatusDef.query.get_or_404(row_id)
    row.is_active = not bool(row.is_active)
    db.session.commit()
    flash('تم تحديث الحالة.', 'success')
    return redirect(url_for('portal.hr_status_defs'))

@portal_bp.route('/hr/masterdata/occasion-types', methods=['GET','POST'])
@login_required
@_perm_any(HR_MASTERDATA_MANAGE, HR_REQUESTS_VIEW_ALL)
def hr_occasion_types():
    from models import HROfficialOccasionType
    if request.method == 'POST':
        if not _hr_can_manage():
            abort(403)
        name_ar = (request.form.get('name_ar') or '').strip()
        name_en = (request.form.get('name_en') or '').strip()
        is_day_off_default = (request.form.get('is_day_off_default') or '1').strip() == '1'
        sort_order = (request.form.get('sort_order') or '100').strip()
        if not name_ar:
            flash('الاسم بالعربية مطلوب.', 'danger')
            return redirect(url_for('portal.hr_occasion_types'))
        try:
            so = int(sort_order)
        except Exception:
            so = 100
        row = HROfficialOccasionType(
            name_ar=name_ar, name_en=name_en or None,
            is_day_off_default=bool(is_day_off_default),
            sort_order=so, is_active=True,
            created_by_id=getattr(current_user,'id',None),
        )
        db.session.add(row)
        db.session.commit()
        flash('تم حفظ النوع.', 'success')
        return redirect(url_for('portal.hr_occasion_types'))

    rows = HROfficialOccasionType.query.order_by(HROfficialOccasionType.sort_order.asc()).all()
    return render_template('portal/hr/occasion_types.html', rows=rows, can_manage=_hr_can_manage())

@portal_bp.route('/hr/masterdata/occasion-types/<int:row_id>/toggle', methods=['POST'])
@login_required
@_perm_any(HR_MASTERDATA_MANAGE, HR_REQUESTS_VIEW_ALL)
def hr_occasion_type_toggle(row_id: int):
    from models import HROfficialOccasionType
    if not _hr_can_manage():
        abort(403)
    row = HROfficialOccasionType.query.get_or_404(row_id)
    row.is_active = not bool(row.is_active)
    db.session.commit()
    flash('تم تحديث النوع.', 'success')
    return redirect(url_for('portal.hr_occasion_types'))


# -------------------------
# HR Reports - Leaves (Additional)
# -------------------------


@portal_bp.route('/hr/reports/leaves/employee-balances', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_leave_employee_balances():
    """تقرير أرصدة الموظفين: فلترة متقدمة (الموظف/الموقع/نوع التعيين/السنة/الرصيد)."""
    users = _list_hr_users()
    work_locations = _hr_lookup_options('WORK_LOCATION')
    appointment_types = _hr_lookup_options('APPOINTMENT_TYPE')
    loc_map = {x.id: x.name for x in (work_locations or [])}
    app_map = {x.id: x.name for x in (appointment_types or [])}

    user_id_raw = (request.args.get('user_id') or '').strip()
    work_location_id_raw = (request.args.get('work_location_id') or '').strip()
    appointment_type_id_raw = (request.args.get('appointment_type_id') or '').strip()
    year_raw = (request.args.get('year') or '').strip()
    within_opening = (request.args.get('within_opening') or '').strip()  # ''/yes/no
    bal_op = (request.args.get('bal_op') or '').strip()  # ''/eq/gt/gte/lt/lte
    days_raw = (request.args.get('days') or '').strip()

    user_id = int(user_id_raw) if user_id_raw.isdigit() else None
    work_location_id = int(work_location_id_raw) if work_location_id_raw.isdigit() else None
    appointment_type_id = int(appointment_type_id_raw) if appointment_type_id_raw.isdigit() else None

    try:
        year = int(year_raw) if year_raw else date.today().year
    except Exception:
        year = date.today().year

    try:
        days_thr = float(days_raw) if days_raw else None
    except Exception:
        days_thr = None

    # Candidate users
    user_ids = _filtered_user_ids(employee_id=user_id, work_location_id=work_location_id, appointment_type_id=appointment_type_id)
    if not user_ids and any([user_id, work_location_id, appointment_type_id]):
        user_ids = []
    elif not user_ids:
        # No filters provided: avoid very heavy report
        try:
            user_ids = [u.id for u in users[:300]]
        except Exception:
            user_ids = []

    # Leave types
    leave_types = HRLeaveType.query.filter_by(is_active=True).order_by(HRLeaveType.name_ar.asc(), HRLeaveType.id.asc()).all()

    rows = []
    today = date.today()
    if user_ids:
        # Prefetch user objects (employee_file is selectin)
        users_map = {u.id: u for u in User.query.filter(User.id.in_(user_ids)).all()}

        for uid in user_ids:
            u = users_map.get(uid)
            if not u:
                continue
            ef = getattr(u, 'employee_file', None)
            wl_name = loc_map.get(getattr(ef, 'work_location_lookup_id', None), '') if ef else ''
            ap_name = app_map.get(getattr(ef, 'appointment_type_lookup_id', None), '') if ef else ''
            for lt in leave_types:
                total = int(_leave_entitlement_days(uid, lt, year) or 0)
                used = float(_leave_used_days_as_of(uid, lt.id, year, today) or 0.0)
                if total == 0 and used == 0:
                    continue
                remaining = float(total) - float(used)

                # Within opening balance filter
                if within_opening == 'yes' and remaining < 0:
                    continue
                if within_opening == 'no' and remaining >= 0:
                    continue

                # Balance comparator filter
                if bal_op and days_thr is not None:
                    if not _cmp_ok(float(remaining), bal_op, float(days_thr)):
                        continue

                rows.append({
                    'user': u,
                    'leave_type': lt,
                    'year': year,
                    'total': total,
                    'used': round(float(used), 2),
                    'remaining': round(float(remaining), 2),
                    'work_location_name': wl_name,
                    'appointment_type_name': ap_name,
                })

    # Hard cap to keep UI safe
    rows = rows[:5000]

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)
        loc_map = {x.id: x.name for x in (work_locations or [])}
        app_map = {x.id: x.name for x in (appointment_types or [])}
        headers = ['السنة', 'الموظف', 'الرقم الوظيفي', 'موقع العمل', 'نوع التعيين', 'نوع الإجازة', 'الرصيد الافتتاحي/الاستحقاق', 'المستخدم حتى اليوم', 'الرصيد الحالي']
        xrows = []
        for r in rows:
            u = r.get('user')
            ef = getattr(u, 'employee_file', None) if u else None
            name = (getattr(ef, 'full_name_quad', None) or getattr(u, 'full_name', None) or getattr(u, 'name', None) or getattr(u, 'email', None) or '') if u else ''
            emp_no = getattr(ef, 'employee_no', '') if ef else ''
            loc = loc_map.get(getattr(ef, 'work_location_lookup_id', None), '') if ef else ''
            app = app_map.get(getattr(ef, 'appointment_type_lookup_id', None), '') if ef else ''
            lt = r.get('leave_type')
            lt_name = (getattr(lt, 'name_ar', None) or getattr(lt, 'name_en', None) or getattr(lt, 'code', None) or '') if lt else ''
            xrows.append([
                r.get('year') or '',
                name,
                emp_no,
                loc,
                app,
                lt_name,
                r.get('total') or 0,
                r.get('used') or 0,
                r.get('remaining') or 0,
            ])
        return _export_xlsx(f'leave_employee_balances_{year}.xlsx', headers, xrows)

    return render_template(
        'portal/hr/reports_leave_employee_balances.html',
        users=users,
        work_locations=work_locations,
        appointment_types=appointment_types,
        rows=rows,
        selected_user_id=user_id,
        selected_work_location_id=work_location_id,
        selected_appointment_type_id=appointment_type_id,
        year=year,
        within_opening=within_opening,
        bal_op=bal_op,
        days=days_raw,
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
    )


@portal_bp.route('/hr/reports/leaves/salary-deductions', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_salary_deductions():
    """تقرير الخصم من الراتب: يعتمد على نتائج تنفيذ الخصم (hr_att_deduction_run/items)."""
    users = _list_hr_users()
    work_locations = _hr_lookup_options('WORK_LOCATION')
    loc_map = {x.id: x.name for x in (work_locations or [])}

    user_id_raw = (request.args.get('user_id') or '').strip()
    work_location_id_raw = (request.args.get('work_location_id') or '').strip()
    from_date = _parse_yyyy_mm_dd(request.args.get('from') or request.args.get('from_date'))
    to_date = _parse_yyyy_mm_dd(request.args.get('to') or request.args.get('to_date'))

    user_id = int(user_id_raw) if user_id_raw.isdigit() else None
    work_location_id = int(work_location_id_raw) if work_location_id_raw.isdigit() else None

    user_ids = _filtered_user_ids(employee_id=user_id, work_location_id=work_location_id)
    if not user_ids and any([user_id, work_location_id]):
        user_ids = []

    # Month key boundaries (YYYYMM)
    from_key = (from_date.year * 100 + from_date.month) if from_date else None
    to_key = (to_date.year * 100 + to_date.month) if to_date else None

    rows = []
    if user_ids:
        q = (HRAttendanceDeductionItem.query
             .join(HRAttendanceDeductionRun, HRAttendanceDeductionRun.id == HRAttendanceDeductionItem.run_id)
             .filter(HRAttendanceDeductionItem.user_id.in_(user_ids)))

        if from_key is not None:
            q = q.filter((HRAttendanceDeductionRun.year * 100 + HRAttendanceDeductionRun.month) >= int(from_key))
        if to_key is not None:
            q = q.filter((HRAttendanceDeductionRun.year * 100 + HRAttendanceDeductionRun.month) <= int(to_key))

        q = q.order_by(HRAttendanceDeductionRun.year.desc(), HRAttendanceDeductionRun.month.desc(), HRAttendanceDeductionItem.amount.desc())
        rows = q.limit(5000).all()

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)
        headers = ['الشهر', 'الموظف', 'الرقم الوظيفي', 'موقع العمل', 'دقائق التأخير', 'دقائق الخروج المبكر', 'أيام الغياب', 'أيام الخصم', 'الحالة', 'ملاحظات']
        xrows = []
        for it in rows:
            run = getattr(it, 'run', None)
            u = getattr(it, 'user', None)
            ef = getattr(u, 'employee_file', None) if u else None
            name = (getattr(ef, 'full_name_quad', None) or getattr(u, 'full_name', None) or getattr(u, 'name', None) or getattr(u, 'email', None) or '') if u else ''
            emp_no = getattr(ef, 'employee_no', '') if ef else ''
            loc = loc_map.get(getattr(ef, 'work_location_lookup_id', None), '') if ef else ''
            month_lbl = ''
            if run:
                month_lbl = f"{int(getattr(run,'year',0) or 0):04d}-{int(getattr(run,'month',0) or 0):02d}"
            xrows.append([
                month_lbl,
                name,
                emp_no,
                loc,
                int(getattr(it, 'late_minutes', 0) or 0),
                int(getattr(it, 'early_leave_minutes', 0) or 0),
                int(getattr(it, 'absent_days', 0) or 0),
                float(getattr(it, 'amount', 0) or 0),
                getattr(run, 'status', '') if run else '',
                getattr(it, 'note', '') or '',
            ])
        return _export_xlsx('salary_deductions.xlsx', headers, xrows)

    return render_template(
        'portal/hr/reports_salary_deductions.html',
        users=users,
        work_locations=work_locations,
        rows=rows,
        selected_user_id=user_id,
        selected_work_location_id=work_location_id,
        from_date=from_date.strftime('%Y-%m-%d') if from_date else '',
        to_date=to_date.strftime('%Y-%m-%d') if to_date else '',
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
    )


@portal_bp.route('/hr/reports/leaves/overtime', methods=['GET'])
@login_required
@_perm(HR_REPORTS_VIEW)
def hr_report_overtime_leaves():
    """تقرير إجازات الإضافي: الموظف/محافظة العمل/موقع العمل + فلتر (فترات عمل مدخلة)."""
    users = _list_hr_users()
    work_govs = _hr_lookup_options('WORK_GOVERNORATE')
    work_locations = _hr_lookup_options('WORK_LOCATION')

    user_id_raw = (request.args.get('user_id') or '').strip()
    gov_id_raw = (request.args.get('work_governorate_id') or '').strip()
    loc_id_raw = (request.args.get('work_location_id') or '').strip()
    has_periods = (request.args.get('has_periods') or '').strip()  # ''/yes/no
    p_from = _parse_yyyy_mm_dd(request.args.get('period_from') or request.args.get('from_date'))
    p_to = _parse_yyyy_mm_dd(request.args.get('period_to') or request.args.get('to_date'))

    user_id = int(user_id_raw) if user_id_raw.isdigit() else None
    gov_id = int(gov_id_raw) if gov_id_raw.isdigit() else None
    loc_id = int(loc_id_raw) if loc_id_raw.isdigit() else None

    # Candidate users (EmployeeFile-based)
    qef = EmployeeFile.query
    if user_id:
        qef = qef.filter(EmployeeFile.user_id == int(user_id))
    if gov_id:
        qef = qef.filter(EmployeeFile.work_governorate_lookup_id == int(gov_id))
    if loc_id:
        qef = qef.filter(EmployeeFile.work_location_lookup_id == int(loc_id))
    user_ids = [r.user_id for r in qef.with_entities(EmployeeFile.user_id).all()]

    # Optional: filter by having overtime "work periods" (AttendanceDailySummary.overtime_minutes)
    overtime_sum = {}
    if user_ids and has_periods in ('yes', 'no'):
        aq = AttendanceDailySummary.query.filter(AttendanceDailySummary.user_id.in_(user_ids))
        if p_from:
            aq = aq.filter(AttendanceDailySummary.day >= p_from.strftime('%Y-%m-%d'))
        if p_to:
            aq = aq.filter(AttendanceDailySummary.day <= p_to.strftime('%Y-%m-%d'))
        aq = aq.filter(AttendanceDailySummary.overtime_minutes > 0)
        ids_with = set([r[0] for r in aq.with_entities(AttendanceDailySummary.user_id).distinct().all()])

        if has_periods == 'yes':
            user_ids = [uid for uid in user_ids if uid in ids_with]
        else:
            user_ids = [uid for uid in user_ids if uid not in ids_with]

        # Sum overtime minutes for display when "yes"
        if has_periods == 'yes' and user_ids:
            sums = (AttendanceDailySummary.query
                    .with_entities(AttendanceDailySummary.user_id, func.coalesce(func.sum(AttendanceDailySummary.overtime_minutes), 0))
                    .filter(AttendanceDailySummary.user_id.in_(user_ids))
                    .filter(AttendanceDailySummary.overtime_minutes > 0))
            if p_from:
                sums = sums.filter(AttendanceDailySummary.day >= p_from.strftime('%Y-%m-%d'))
            if p_to:
                sums = sums.filter(AttendanceDailySummary.day <= p_to.strftime('%Y-%m-%d'))
            sums = sums.group_by(AttendanceDailySummary.user_id).all()
            overtime_sum = {int(uid): int(m or 0) for uid, m in (sums or [])}

    # Identify overtime/extra leave types
    overtime_type_ids = []
    try:
        qlt = HRLeaveType.query.filter_by(is_active=True)
        qlt = qlt.filter(or_(
            func.upper(HRLeaveType.code).in_(['OVERTIME', 'OT', 'EXTRA', 'ADDITIONAL', 'OVT']),
            HRLeaveType.name_ar.ilike('%اضافي%'),
            HRLeaveType.name_ar.ilike('%إضافي%'),
            HRLeaveType.name_ar.ilike('%الإضافي%'),
        ))
        overtime_type_ids = [int(x.id) for x in qlt.all() if x and x.id]
    except Exception:
        overtime_type_ids = []

    rows = []
    overtime_types_found = bool(overtime_type_ids)
    if user_ids and overtime_type_ids:
        q = (HRLeaveRequest.query
             .filter(HRLeaveRequest.user_id.in_(user_ids))
             .filter(HRLeaveRequest.leave_type_id.in_(overtime_type_ids))
             .order_by(HRLeaveRequest.start_date.desc(), HRLeaveRequest.id.desc()))
        rows = q.limit(5000).all()

    if (request.args.get('export') or '').lower() == 'xlsx':
        if not current_user.has_perm(HR_REPORTS_EXPORT):
            abort(403)
        gov_map = {x.id: x.name for x in (work_govs or [])}
        loc_map = {x.id: x.name for x in (work_locations or [])}
        headers = ['الموظف', 'الرقم الوظيفي', 'محافظة العمل', 'موقع العمل', 'نوع الإجازة', 'من تاريخ', 'إلى تاريخ', 'عدد الأيام', 'الحالة', 'دقائق عمل إضافي (ضمن الفترة)']
        xrows = []
        for r in rows:
            u = getattr(r, 'user', None)
            ef = getattr(u, 'employee_file', None) if u else None
            name = (getattr(ef, 'full_name_quad', None) or getattr(u, 'full_name', None) or getattr(u, 'name', None) or getattr(u, 'email', None) or '') if u else ''
            emp_no = getattr(ef, 'employee_no', '') if ef else ''
            gov = gov_map.get(getattr(ef, 'work_governorate_lookup_id', None), '') if ef else ''
            loc = loc_map.get(getattr(ef, 'work_location_lookup_id', None), '') if ef else ''
            lt = getattr(r, 'leave_type', None)
            lt_name = (getattr(lt, 'name_ar', None) or getattr(lt, 'name_en', None) or getattr(lt, 'code', None) or '') if lt else ''
            xrows.append([
                name,
                emp_no,
                gov,
                loc,
                lt_name,
                r.start_date or '',
                r.end_date or '',
                int(r.days or 0) if (r.days is not None) else '',
                r.status or '',
                overtime_sum.get(getattr(u, 'id', None), 0) if u else 0,
            ])
        return _export_xlsx('overtime_leaves.xlsx', headers, xrows)

    return render_template(
        'portal/hr/reports_overtime_leaves.html',
        users=users,
        work_govs=work_govs,
        work_locations=work_locations,
        rows=rows,
        selected_user_id=user_id,
        selected_work_governorate_id=gov_id,
        selected_work_location_id=loc_id,
        has_periods=has_periods,
        period_from=p_from.strftime('%Y-%m-%d') if p_from else '',
        period_to=p_to.strftime('%Y-%m-%d') if p_to else '',
        overtime_sum=overtime_sum,
        overtime_types_found=overtime_types_found,
        can_export=current_user.has_perm(HR_REPORTS_EXPORT),
    )

# ==========================================================
# Inventory Store (Warehouse Module)
# ==========================================================

@portal_bp.route("/inventory")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_home():
    return render_template("portal/inventory/home.html")


@portal_bp.route("/inventory/_p/<key>")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_placeholder(key: str):
    titles = {
        "inbound_voucher": "سندات الإدخال - إدخال سند",
        "stocktake_new": "سندات الجرد - إدخال سند",
        "stocktake_log": "سندات الجرد - سجل سندات الجرد",
        "scrap": "سندات الإتلاف",
        "returns": "سندات الإرجاع",
        "report_warehouses": "تقرير المخازن ومحتوياتها",
        "report_items_all": "تقرير الأصناف والكميات في كافة المخازن",
        "report_item_card": "تقرير بطاقة الصنف",
        "report_yearly_vouchers": "الجدول السنوي للسندات",
        "report_yearly_qty": "الجدول السنوي للكميات",
        "report_issue_log": "تقرير سجلات الصرف",
        "custody_stocktake": "العهدة - سند جرد عهدة",
        "custody_vouchers": "العهدة - سجل سندات العهدة",
        "custody_view": "العهدة - عرض العهدة",
        "items_custody": "العهدة - الأصناف والعهدة",
        "room_requesters": "لوحة التحكم - المستخدمون القادرون على طلب مواد للغرف",
        "warehouse_perms": "لوحة التحكم - صلاحيات الموظفين على المخازن",
        "inventory_settings": "لوحة التحكم - الإعدادات العامة",
    }
    title = titles.get(key, "المستودع")
    return render_template("portal/inventory/placeholder.html", title=title)


@portal_bp.route("/inventory/requests")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_requests_log():
    departments = Department.query.order_by(Department.name_ar.asc()).all()
    work_locations = HRLookupItem.query.filter(HRLookupItem.category == "WORK_LOCATION").order_by(HRLookupItem.sort_order.asc(), HRLookupItem.name_ar.asc()).all()

    department_id = request.args.get("department_id") or ""
    work_location_id = request.args.get("work_location_id") or ""
    entered_by = (request.args.get("entered_by") or "").strip()
    contains = (request.args.get("contains") or "").strip()
    from_date = request.args.get("from_date") or ""
    to_date = request.args.get("to_date") or ""
    status = request.args.get("status") or ""
    manager_approval = request.args.get("manager_approval") or ""

    qry = InvRequest.query

    if department_id:
        try:
            qry = qry.filter(InvRequest.department_id == int(department_id))
        except Exception:
            pass
    if work_location_id:
        try:
            qry = qry.filter(InvRequest.work_location_lookup_id == int(work_location_id))
        except Exception:
            pass

    if entered_by:
        like = f"%{entered_by}%"
        qry = qry.join(User, InvRequest.entered_by_id == User.id, isouter=True).filter(
            or_(
                User.full_name.ilike(like),
                User.email.ilike(like),
            )
        )

    if contains:
        qry = qry.filter(InvRequest.items_text.ilike(f"%{contains}%"))

    if from_date:
        qry = qry.filter(InvRequest.request_date >= from_date)
    if to_date:
        qry = qry.filter(InvRequest.request_date <= to_date)

    if status:
        qry = qry.filter(InvRequest.status == status)
    if manager_approval:
        qry = qry.filter(InvRequest.manager_approval == manager_approval)

    rows = qry.order_by(InvRequest.request_date.desc(), InvRequest.id.desc()).limit(500).all()

    selected = {
        "department_id": int(department_id) if department_id.isdigit() else None,
        "work_location_id": int(work_location_id) if work_location_id.isdigit() else None,
        "entered_by": entered_by,
        "contains": contains,
        "from_date": from_date,
        "to_date": to_date,
        "status": status,
        "manager_approval": manager_approval,
    }

    return render_template(
        "portal/inventory/requests_log.html",
        departments=departments,
        work_locations=work_locations,
        rows=rows,
        selected=selected,
    )


@portal_bp.route("/inventory/vouchers/issue/new", methods=["GET", "POST"])
@login_required
@_perm(STORE_MANAGE)
def inventory_issue_voucher_new():
    warehouses = InvWarehouse.query.filter(InvWarehouse.is_active == True).order_by(InvWarehouse.name.asc()).all()  # noqa: E712
    categories = InvItemCategory.query.filter(InvItemCategory.is_active == True).order_by(InvItemCategory.name.asc()).all()  # noqa: E712
    items = InvItem.query.filter(InvItem.is_active == True).order_by(InvItem.name.asc()).all()  # noqa: E712
    rooms = InvRoom.query.filter(InvRoom.is_active == True).order_by(InvRoom.name.asc()).all()  # noqa: E712

    if request.method == "POST":
        issue_kind = (request.form.get("issue_kind") or "ROOM").strip().upper()
        voucher_no = (request.form.get("voucher_no") or "").strip()
        voucher_date = (request.form.get("voucher_date") or "").strip()
        from_warehouse_id = request.form.get("from_warehouse_id")
        to_warehouse_id = request.form.get("to_warehouse_id")
        to_room_name = (request.form.get("to_room_name") or "").strip()
        note = (request.form.get("note") or "").strip()

        if not voucher_no or not voucher_date or not from_warehouse_id:
            flash("يرجى تعبئة رقم السند + تاريخ السند + من مستودع.", "warning")
            return redirect(url_for("portal.inventory_issue_voucher_new"))

        try:
            from_warehouse_id = int(from_warehouse_id)
        except Exception:
            flash("من مستودع غير صالح.", "warning")
            return redirect(url_for("portal.inventory_issue_voucher_new"))

        if issue_kind == "WAREHOUSE":
            try:
                to_warehouse_id_val = int(to_warehouse_id) if to_warehouse_id else None
            except Exception:
                to_warehouse_id_val = None
            if not to_warehouse_id_val:
                flash("يرجى اختيار إلى مستودع.", "warning")
                return redirect(url_for("portal.inventory_issue_voucher_new"))
            if to_warehouse_id_val == from_warehouse_id:
                flash("لا يمكن أن يكون (إلى مستودع) نفس (من مستودع).", "warning")
                return redirect(url_for("portal.inventory_issue_voucher_new"))
        else:
            to_warehouse_id_val = None
            if not to_room_name:
                flash("يرجى تعبئة اسم الغرفة.", "warning")
                return redirect(url_for("portal.inventory_issue_voucher_new"))

        # Collect line arrays (support both item_id[] and item_id)
        item_ids = request.form.getlist("item_id[]") or request.form.getlist("item_id")
        qtys = request.form.getlist("qty[]") or request.form.getlist("qty")
        serials = request.form.getlist("serial[]") or request.form.getlist("serial")
        w_starts = request.form.getlist("w_start[]") or request.form.getlist("w_start")
        w_ends = request.form.getlist("w_end[]") or request.form.getlist("w_end")
        details_list = request.form.getlist("details[]") or request.form.getlist("details")

        if not item_ids:
            flash("يرجى إضافة صنف واحد على الأقل.", "warning")
            return redirect(url_for("portal.inventory_issue_voucher_new"))

        # Create voucher
        v = InvIssueVoucher(
            issue_kind=issue_kind,
            voucher_no=voucher_no,
            voucher_date=voucher_date,
            from_warehouse_id=from_warehouse_id,
            to_warehouse_id=to_warehouse_id_val,
            to_room_name=to_room_name if issue_kind != "WAREHOUSE" else None,
            note=note or None,
            created_by_id=current_user.id,
            created_at=datetime.utcnow(),
        )
        db.session.add(v)
        db.session.flush()

        any_line = False
        for i, raw_item_id in enumerate(item_ids):
            if not raw_item_id:
                continue
            try:
                item_id = int(raw_item_id)
            except Exception:
                continue

            raw_qty = (qtys[i] if i < len(qtys) else "1")
            try:
                qty = float(raw_qty) if raw_qty not in (None, "") else 1.0
            except Exception:
                qty = 1.0

            ln = InvIssueVoucherLine(
                voucher_id=v.id,
                item_id=item_id,
                qty=qty,
                serial=(serials[i] if i < len(serials) else None) or None,
                warranty_start=(w_starts[i] if i < len(w_starts) else None) or None,
                warranty_end=(w_ends[i] if i < len(w_ends) else None) or None,
                details=(details_list[i] if i < len(details_list) else None) or None,
            )
            db.session.add(ln)
            any_line = True

        if not any_line:
            db.session.rollback()
            flash("يرجى إضافة صنف واحد على الأقل.", "warning")
            return redirect(url_for("portal.inventory_issue_voucher_new"))

        # Attachments
        # Audit + commit voucher/lines first (protect from attachments failure)
        try:
            db.session.add(AuditLog(
                user_id=current_user.id,
                action="INV_ISSUE_CREATE",
                note=f"إنشاء سند صرف ({voucher_no})",
                target_type="INV_ISSUE_VOUCHER",
                target_id=v.id,
                created_at=datetime.utcnow(),
            ))
        except Exception:
            pass

        db.session.commit()

        # Attachments (best effort)
        files = request.files.getlist("attachments")
        if files:
            try:
                from uuid import uuid4

                base = os.path.join(current_app.instance_path, "uploads", "inventory", "issue", str(v.id))
                os.makedirs(base, exist_ok=True)

                for f in files:
                    if not f or not getattr(f, "filename", ""):
                        continue
                    original = f.filename
                    stored = f"{uuid4().hex}_{secure_filename(original) or 'file'}"
                    full_path = os.path.join(base, stored)
                    f.save(full_path)

                    rel = os.path.relpath(full_path, current_app.instance_path)
                    db.session.add(
                        InvIssueVoucherAttachment(
                            voucher_id=v.id,
                            original_name=original,
                            stored_name=stored,
                            file_path=rel,
                            uploaded_by_id=current_user.id,
                            uploaded_at=datetime.utcnow(),
                        )
                    )
                db.session.commit()
            except Exception:
                db.session.rollback()
                flash("تم حفظ السند، لكن تعذّر حفظ المرفقات.", "warning")

        flash("تم حفظ سند الصرف.", "success")
        return redirect(url_for("portal.inventory_issue_voucher_view", v_id=v.id))

    # Default form values for initial GET (template expects `form.*`)
    form = {
        "issue_kind": "ROOM",
        "voucher_no": "",
        "voucher_date": "",
        "from_warehouse_id": None,
        "to_warehouse_id": None,
        "to_room_name": "",
        "note": "",
        "lines": [],
    }

    return render_template(
        "portal/inventory/issue_voucher_new.html",
        warehouses=warehouses,
        categories=categories,
        items=items,
        rooms=rooms,
        form=form,
    )


@portal_bp.route("/inventory/vouchers/issue/list")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_issue_voucher_list():
    warehouses = InvWarehouse.query.filter(InvWarehouse.is_active == True).order_by(InvWarehouse.name.asc()).all()  # noqa: E712

    voucher_no = (request.args.get("voucher_no") or "").strip()
    from_date = request.args.get("from_date") or ""
    to_date = request.args.get("to_date") or ""
    from_warehouse_id = request.args.get("from_warehouse_id") or ""
    to_warehouse_id = request.args.get("to_warehouse_id") or ""
    contains = (request.args.get("contains") or "").strip()
    issue_kind = (request.args.get("issue_kind") or "").strip().upper()

    qry = InvIssueVoucher.query

    if voucher_no:
        qry = qry.filter(InvIssueVoucher.voucher_no.ilike(f"%{voucher_no}%"))

    if from_date:
        qry = qry.filter(InvIssueVoucher.voucher_date >= from_date)
    if to_date:
        qry = qry.filter(InvIssueVoucher.voucher_date <= to_date)

    if from_warehouse_id and from_warehouse_id.isdigit():
        qry = qry.filter(InvIssueVoucher.from_warehouse_id == int(from_warehouse_id))
    if to_warehouse_id and to_warehouse_id.isdigit():
        qry = qry.filter(InvIssueVoucher.to_warehouse_id == int(to_warehouse_id))

    if issue_kind:
        qry = qry.filter(InvIssueVoucher.issue_kind == issue_kind)

    if contains:
        like = f"%{contains}%"
        sub = (
            db.session.query(InvIssueVoucherLine.id)
            .join(InvItem, InvIssueVoucherLine.item_id == InvItem.id)
            .filter(InvIssueVoucherLine.voucher_id == InvIssueVoucher.id)
            .filter(or_(InvItem.name.ilike(like), InvItem.code.ilike(like)))
            .exists()
        )
        qry = qry.filter(sub)

    rows = qry.order_by(InvIssueVoucher.voucher_date.desc(), InvIssueVoucher.id.desc()).limit(500).all()

    selected = {
        "voucher_no": voucher_no,
        "from_date": from_date,
        "to_date": to_date,
        "from_warehouse_id": int(from_warehouse_id) if from_warehouse_id.isdigit() else None,
        "to_warehouse_id": int(to_warehouse_id) if to_warehouse_id.isdigit() else None,
        "contains": contains,
        "issue_kind": issue_kind,
    }

    return render_template(
        "portal/inventory/issue_voucher_list.html",
        warehouses=warehouses,
        rows=rows,
        selected=selected,
    )


@portal_bp.route("/inventory/vouchers/issue/<int:v_id>")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_issue_voucher_view(v_id: int):
    v = InvIssueVoucher.query.get_or_404(v_id)

    # Permission scope: STORE_READ allows view-all; STORE_MANAGE allows view-all.
    # For now, we also allow creator to view.
    if not (current_user.has_perm(STORE_READ) or current_user.has_perm(STORE_MANAGE) or (v.created_by_id == current_user.id)):
        abort(403)

    lines = InvIssueVoucherLine.query.filter(InvIssueVoucherLine.voucher_id == v.id).order_by(InvIssueVoucherLine.id.asc()).all()
    attachments = InvIssueVoucherAttachment.query.filter(InvIssueVoucherAttachment.voucher_id == v.id).order_by(InvIssueVoucherAttachment.id.asc()).all()

    return render_template(
        "portal/inventory/issue_voucher_view.html",
        v=v,
        lines=lines,
        attachments=attachments,
    )


@portal_bp.route("/inventory/vouchers/issue/attachment/<int:att_id>/download")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_issue_attachment_download(att_id: int):
    att = InvIssueVoucherAttachment.query.get_or_404(att_id)
    v = InvIssueVoucher.query.get(att.voucher_id)

    if not v:
        abort(404)

    if not (current_user.has_perm(STORE_READ) or current_user.has_perm(STORE_MANAGE) or (v.created_by_id == current_user.id)):
        abort(403)

    abs_path = os.path.join(current_app.instance_path, att.file_path)
    if not os.path.exists(abs_path):
        abort(404)

    directory = os.path.dirname(abs_path)
    filename = os.path.basename(abs_path)
    return send_from_directory(directory, filename, as_attachment=True, download_name=att.original_name)


# -------------------------


# -------------------------
# Inventory: Inbound (Input) Vouchers
# -------------------------


@portal_bp.route("/inventory/vouchers/inbound/new", methods=["GET", "POST"])
@login_required
@_perm(STORE_MANAGE)
def inventory_inbound_voucher_new():
    warehouses = InvWarehouse.query.filter(InvWarehouse.is_active == True).order_by(InvWarehouse.name.asc()).all()  # noqa: E712
    items = InvItem.query.filter(InvItem.is_active == True).order_by(InvItem.name.asc()).all()  # noqa: E712

    if request.method == "POST":
        voucher_no = (request.form.get("voucher_no") or "").strip()
        voucher_date = (request.form.get("voucher_date") or "").strip()
        to_warehouse_id = request.form.get("to_warehouse_id")
        note = (request.form.get("note") or "").strip()

        if not voucher_no or not voucher_date or not to_warehouse_id:
            flash("يرجى تعبئة رقم السند + تاريخ السند + إلى مستودع.", "warning")
            return redirect(url_for("portal.inventory_inbound_voucher_new"))

        try:
            to_warehouse_id = int(to_warehouse_id)
        except Exception:
            flash("إلى مستودع غير صالح.", "warning")
            return redirect(url_for("portal.inventory_inbound_voucher_new"))

        item_ids = request.form.getlist("item_id[]") or request.form.getlist("item_id")
        qtys = request.form.getlist("qty[]") or request.form.getlist("qty")
        serials = request.form.getlist("serial[]") or request.form.getlist("serial")
        w_starts = request.form.getlist("w_start[]") or request.form.getlist("w_start")
        w_ends = request.form.getlist("w_end[]") or request.form.getlist("w_end")
        details_list = request.form.getlist("details[]") or request.form.getlist("details")

        if not item_ids:
            flash("يرجى إضافة صنف واحد على الأقل.", "warning")
            return redirect(url_for("portal.inventory_inbound_voucher_new"))

        v = InvInboundVoucher(
            voucher_no=voucher_no,
            voucher_date=voucher_date,
            to_warehouse_id=to_warehouse_id,
            note=note or None,
            created_by_id=current_user.id,
            created_at=datetime.utcnow(),
        )
        db.session.add(v)
        db.session.flush()

        any_line = False
        for i, raw_item_id in enumerate(item_ids):
            if not raw_item_id:
                continue
            try:
                item_id = int(raw_item_id)
            except Exception:
                continue

            raw_qty = (qtys[i] if i < len(qtys) else "1")
            try:
                qty = float(raw_qty) if raw_qty not in (None, "") else 1.0
            except Exception:
                qty = 1.0

            ln = InvInboundVoucherLine(
                voucher_id=v.id,
                item_id=item_id,
                qty=qty,
                serial=(serials[i] if i < len(serials) else None) or None,
                warranty_start=(w_starts[i] if i < len(w_starts) else None) or None,
                warranty_end=(w_ends[i] if i < len(w_ends) else None) or None,
                details=(details_list[i] if i < len(details_list) else None) or None,
            )
            db.session.add(ln)
            any_line = True

        if not any_line:
            db.session.rollback()
            flash("يرجى إضافة صنف واحد على الأقل.", "warning")
            return redirect(url_for("portal.inventory_inbound_voucher_new"))

        # Audit + commit voucher/lines first (protect from attachments failure)
        try:
            db.session.add(AuditLog(
                user_id=current_user.id,
                action="INV_INBOUND_CREATE",
                note=f"إنشاء سند إدخال ({voucher_no})",
                target_type="INV_INBOUND_VOUCHER",
                target_id=v.id,
                created_at=datetime.utcnow(),
            ))
        except Exception:
            pass

        db.session.commit()

        # Attachments (best effort)
        files = request.files.getlist("attachments")
        if files:
            try:
                from uuid import uuid4

                base = os.path.join(current_app.instance_path, "uploads", "inventory", "inbound", str(v.id))
                os.makedirs(base, exist_ok=True)

                for f in files:
                    if not f or not getattr(f, "filename", ""):
                        continue
                    original = f.filename
                    stored = f"{uuid4().hex}_{secure_filename(original) or 'file'}"
                    full_path = os.path.join(base, stored)
                    f.save(full_path)

                    rel = os.path.relpath(full_path, current_app.instance_path)
                    db.session.add(
                        InvInboundVoucherAttachment(
                            voucher_id=v.id,
                            original_name=original,
                            stored_name=stored,
                            file_path=rel,
                            uploaded_by_id=current_user.id,
                            uploaded_at=datetime.utcnow(),
                        )
                    )
                db.session.commit()
            except Exception:
                db.session.rollback()
                flash("تم حفظ السند، لكن تعذّر حفظ المرفقات.", "warning")

        flash("تم حفظ سند الإدخال.", "success")
        return redirect(url_for("portal.inventory_inbound_voucher_view", v_id=v.id))

    form = {
        "voucher_no": "",
        "voucher_date": "",
        "to_warehouse_id": None,
        "note": "",
        "lines": [],
    }

    return render_template(
        "portal/inventory/inbound_voucher_new.html",
        warehouses=warehouses,
        items=items,
        form=form,
    )


@portal_bp.route("/inventory/vouchers/inbound/list")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_inbound_voucher_list():
    warehouses = InvWarehouse.query.filter(InvWarehouse.is_active == True).order_by(InvWarehouse.name.asc()).all()  # noqa: E712

    voucher_no = (request.args.get("voucher_no") or "").strip()
    from_date = request.args.get("from_date") or ""
    to_date = request.args.get("to_date") or ""
    to_warehouse_id = request.args.get("to_warehouse_id") or ""
    contains = (request.args.get("contains") or "").strip()

    qry = InvInboundVoucher.query

    if voucher_no:
        qry = qry.filter(InvInboundVoucher.voucher_no.ilike(f"%{voucher_no}%"))

    if from_date:
        qry = qry.filter(InvInboundVoucher.voucher_date >= from_date)
    if to_date:
        qry = qry.filter(InvInboundVoucher.voucher_date <= to_date)

    if to_warehouse_id and to_warehouse_id.isdigit():
        qry = qry.filter(InvInboundVoucher.to_warehouse_id == int(to_warehouse_id))

    if contains:
        like = f"%{contains}%"
        sub = (
            db.session.query(InvInboundVoucherLine.id)
            .join(InvItem, InvItem.id == InvInboundVoucherLine.item_id)
            .filter(InvInboundVoucherLine.voucher_id == InvInboundVoucher.id)
            .filter(or_(InvItem.name.ilike(like), InvItem.code.ilike(like)))
            .exists()
        )
        qry = qry.filter(sub)

    rows = qry.order_by(InvInboundVoucher.voucher_date.desc(), InvInboundVoucher.id.desc()).limit(500).all()

    selected = {
        "voucher_no": voucher_no,
        "from_date": from_date,
        "to_date": to_date,
        "to_warehouse_id": int(to_warehouse_id) if to_warehouse_id.isdigit() else None,
        "contains": contains,
    }

    return render_template(
        "portal/inventory/inbound_voucher_list.html",
        warehouses=warehouses,
        rows=rows,
        selected=selected,
    )


@portal_bp.route("/inventory/vouchers/inbound/<int:v_id>")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_inbound_voucher_view(v_id: int):
    v = InvInboundVoucher.query.get_or_404(v_id)
    if not (current_user.has_perm(STORE_READ) or current_user.has_perm(STORE_MANAGE) or (v.created_by_id == current_user.id)):
        abort(403)

    lines = InvInboundVoucherLine.query.filter(InvInboundVoucherLine.voucher_id == v.id).order_by(InvInboundVoucherLine.id.asc()).all()
    attachments = InvInboundVoucherAttachment.query.filter(InvInboundVoucherAttachment.voucher_id == v.id).order_by(InvInboundVoucherAttachment.id.asc()).all()

    return render_template(
        "portal/inventory/inbound_voucher_view.html",
        v=v,
        lines=lines,
        attachments=attachments,
    )


@portal_bp.route("/inventory/vouchers/inbound/attachment/<int:att_id>/download")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_inbound_attachment_download(att_id: int):
    att = InvInboundVoucherAttachment.query.get_or_404(att_id)
    v = InvInboundVoucher.query.get(att.voucher_id)
    if not v:
        abort(404)

    if not (current_user.has_perm(STORE_READ) or current_user.has_perm(STORE_MANAGE) or (v.created_by_id == current_user.id)):
        abort(403)

    abs_path = os.path.join(current_app.instance_path, att.file_path)
    if not os.path.exists(abs_path):
        abort(404)

    directory = os.path.dirname(abs_path)
    filename = os.path.basename(abs_path)
    return send_from_directory(directory, filename, as_attachment=True, download_name=att.original_name)

# Inventory: Scrap (Destruction) Vouchers
# -------------------------


@portal_bp.route("/inventory/vouchers/scrap/new", methods=["GET", "POST"])
@login_required
@_perm(STORE_MANAGE)
def inventory_scrap_voucher_new():
    warehouses = InvWarehouse.query.filter(InvWarehouse.is_active == True).order_by(InvWarehouse.name.asc()).all()  # noqa: E712
    items = InvItem.query.filter(InvItem.is_active == True).order_by(InvItem.name.asc()).all()  # noqa: E712

    if request.method == "POST":
        voucher_no = (request.form.get("voucher_no") or "").strip()
        voucher_date = (request.form.get("voucher_date") or "").strip()
        from_warehouse_id = request.form.get("from_warehouse_id")
        note = (request.form.get("note") or "").strip()

        if not voucher_no or not voucher_date or not from_warehouse_id:
            flash("يرجى تعبئة رقم السند + تاريخ السند + من مستودع.", "warning")
            return redirect(url_for("portal.inventory_scrap_voucher_new"))

        try:
            from_warehouse_id = int(from_warehouse_id)
        except Exception:
            flash("من مستودع غير صالح.", "warning")
            return redirect(url_for("portal.inventory_scrap_voucher_new"))

        item_ids = request.form.getlist("item_id[]") or request.form.getlist("item_id")
        qtys = request.form.getlist("qty[]") or request.form.getlist("qty")
        serials = request.form.getlist("serial[]") or request.form.getlist("serial")
        w_starts = request.form.getlist("w_start[]") or request.form.getlist("w_start")
        w_ends = request.form.getlist("w_end[]") or request.form.getlist("w_end")
        details_list = request.form.getlist("details[]") or request.form.getlist("details")

        if not item_ids:
            flash("يرجى إضافة صنف واحد على الأقل.", "warning")
            return redirect(url_for("portal.inventory_scrap_voucher_new"))

        v = InvScrapVoucher(
            voucher_no=voucher_no,
            voucher_date=voucher_date,
            from_warehouse_id=from_warehouse_id,
            note=note or None,
            created_by_id=current_user.id,
            created_at=datetime.utcnow(),
        )
        db.session.add(v)
        db.session.flush()

        any_line = False
        for i, raw_item_id in enumerate(item_ids):
            if not raw_item_id:
                continue
            try:
                item_id = int(raw_item_id)
            except Exception:
                continue

            raw_qty = (qtys[i] if i < len(qtys) else "1")
            try:
                qty = float(raw_qty) if raw_qty not in (None, "") else 1.0
            except Exception:
                qty = 1.0

            ln = InvScrapVoucherLine(
                voucher_id=v.id,
                item_id=item_id,
                qty=qty,
                serial=(serials[i] if i < len(serials) else None) or None,
                warranty_start=(w_starts[i] if i < len(w_starts) else None) or None,
                warranty_end=(w_ends[i] if i < len(w_ends) else None) or None,
                details=(details_list[i] if i < len(details_list) else None) or None,
            )
            db.session.add(ln)
            any_line = True

        if not any_line:
            db.session.rollback()
            flash("يرجى إضافة صنف واحد على الأقل.", "warning")
            return redirect(url_for("portal.inventory_scrap_voucher_new"))

        # Audit + commit voucher/lines first (protect from attachments failure)
        try:
            db.session.add(AuditLog(
                user_id=current_user.id,
                action="INV_SCRAP_CREATE",
                note=f"إنشاء سند إتلاف ({voucher_no})",
                target_type="INV_SCRAP_VOUCHER",
                target_id=v.id,
                created_at=datetime.utcnow(),
            ))
        except Exception:
            pass

        db.session.commit()

        # Attachments (best effort)
        files = request.files.getlist("attachments")
        if files:
            try:
                from uuid import uuid4

                base = os.path.join(current_app.instance_path, "uploads", "inventory", "scrap", str(v.id))
                os.makedirs(base, exist_ok=True)

                for f in files:
                    if not f or not getattr(f, "filename", ""):
                        continue
                    original = f.filename
                    stored = f"{uuid4().hex}_{secure_filename(original) or 'file'}"
                    full_path = os.path.join(base, stored)
                    f.save(full_path)

                    rel = os.path.relpath(full_path, current_app.instance_path)
                    db.session.add(
                        InvScrapVoucherAttachment(
                            voucher_id=v.id,
                            original_name=original,
                            stored_name=stored,
                            file_path=rel,
                            uploaded_by_id=current_user.id,
                            uploaded_at=datetime.utcnow(),
                        )
                    )
                db.session.commit()
            except Exception:
                db.session.rollback()
                flash("تم حفظ السند، لكن تعذّر حفظ المرفقات.", "warning")

        flash("تم حفظ سند الإتلاف.", "success")
        return redirect(url_for("portal.inventory_scrap_voucher_view", v_id=v.id))

    form = {
        "voucher_no": "",
        "voucher_date": "",
        "from_warehouse_id": None,
        "note": "",
        "lines": [],
    }

    return render_template(
        "portal/inventory/scrap_voucher_new.html",
        warehouses=warehouses,
        items=items,
        form=form,
    )


@portal_bp.route("/inventory/vouchers/scrap/list")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_scrap_voucher_list():
    warehouses = InvWarehouse.query.filter(InvWarehouse.is_active == True).order_by(InvWarehouse.name.asc()).all()  # noqa: E712

    voucher_no = (request.args.get("voucher_no") or "").strip()
    from_date = request.args.get("from_date") or ""
    to_date = request.args.get("to_date") or ""
    from_warehouse_id = request.args.get("from_warehouse_id") or ""
    contains = (request.args.get("contains") or "").strip()

    qry = InvScrapVoucher.query

    if voucher_no:
        qry = qry.filter(InvScrapVoucher.voucher_no.ilike(f"%{voucher_no}%"))

    if from_date:
        qry = qry.filter(InvScrapVoucher.voucher_date >= from_date)
    if to_date:
        qry = qry.filter(InvScrapVoucher.voucher_date <= to_date)

    if from_warehouse_id and from_warehouse_id.isdigit():
        qry = qry.filter(InvScrapVoucher.from_warehouse_id == int(from_warehouse_id))

    if contains:
        like = f"%{contains}%"
        sub = (
            db.session.query(InvScrapVoucherLine.id)
            .join(InvItem, InvItem.id == InvScrapVoucherLine.item_id)
            .filter(InvScrapVoucherLine.voucher_id == InvScrapVoucher.id)
            .filter(or_(InvItem.name.ilike(like), InvItem.code.ilike(like)))
            .exists()
        )
        qry = qry.filter(sub)

    rows = qry.order_by(InvScrapVoucher.voucher_date.desc(), InvScrapVoucher.id.desc()).limit(500).all()

    selected = {
        "voucher_no": voucher_no,
        "from_date": from_date,
        "to_date": to_date,
        "from_warehouse_id": int(from_warehouse_id) if from_warehouse_id.isdigit() else None,
        "contains": contains,
    }

    return render_template(
        "portal/inventory/scrap_voucher_list.html",
        warehouses=warehouses,
        rows=rows,
        selected=selected,
    )


@portal_bp.route("/inventory/vouchers/scrap/<int:v_id>")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_scrap_voucher_view(v_id: int):
    v = InvScrapVoucher.query.get_or_404(v_id)
    if not (current_user.has_perm(STORE_READ) or current_user.has_perm(STORE_MANAGE) or (v.created_by_id == current_user.id)):
        abort(403)

    lines = InvScrapVoucherLine.query.filter(InvScrapVoucherLine.voucher_id == v.id).order_by(InvScrapVoucherLine.id.asc()).all()
    attachments = InvScrapVoucherAttachment.query.filter(InvScrapVoucherAttachment.voucher_id == v.id).order_by(InvScrapVoucherAttachment.id.asc()).all()

    return render_template(
        "portal/inventory/scrap_voucher_view.html",
        v=v,
        lines=lines,
        attachments=attachments,
    )


@portal_bp.route("/inventory/vouchers/scrap/attachment/<int:att_id>/download")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_scrap_attachment_download(att_id: int):
    att = InvScrapVoucherAttachment.query.get_or_404(att_id)
    v = InvScrapVoucher.query.get(att.voucher_id)
    if not v:
        abort(404)

    if not (current_user.has_perm(STORE_READ) or current_user.has_perm(STORE_MANAGE) or (v.created_by_id == current_user.id)):
        abort(403)

    abs_path = os.path.join(current_app.instance_path, att.file_path)
    if not os.path.exists(abs_path):
        abort(404)

    directory = os.path.dirname(abs_path)
    filename = os.path.basename(abs_path)
    return send_from_directory(directory, filename, as_attachment=True, download_name=att.original_name)


# -------------------------
# Inventory: Return Vouchers
# -------------------------


@portal_bp.route("/inventory/vouchers/return/new", methods=["GET", "POST"])
@login_required
@_perm(STORE_MANAGE)
def inventory_return_voucher_new():
    warehouses = InvWarehouse.query.filter(InvWarehouse.is_active == True).order_by(InvWarehouse.name.asc()).all()  # noqa: E712
    items = InvItem.query.filter(InvItem.is_active == True).order_by(InvItem.name.asc()).all()  # noqa: E712
    rooms = InvRoom.query.filter(InvRoom.is_active == True).order_by(InvRoom.name.asc()).all()  # noqa: E712

    if request.method == "POST":
        voucher_no = (request.form.get("voucher_no") or "").strip()
        voucher_date = (request.form.get("voucher_date") or "").strip()
        to_warehouse_id = request.form.get("to_warehouse_id")
        from_room_name = (request.form.get("from_room_name") or "").strip()
        note = (request.form.get("note") or "").strip()

        if not voucher_no or not voucher_date or not to_warehouse_id:
            flash("يرجى تعبئة رقم السند + تاريخ السند + إلى مستودع.", "warning")
            return redirect(url_for("portal.inventory_return_voucher_new"))

        try:
            to_warehouse_id = int(to_warehouse_id)
        except Exception:
            flash("إلى مستودع غير صالح.", "warning")
            return redirect(url_for("portal.inventory_return_voucher_new"))

        if not from_room_name:
            flash("يرجى تعبئة (من جهة/غرفة).", "warning")
            return redirect(url_for("portal.inventory_return_voucher_new"))

        item_ids = request.form.getlist("item_id[]") or request.form.getlist("item_id")
        qtys = request.form.getlist("qty[]") or request.form.getlist("qty")
        serials = request.form.getlist("serial[]") or request.form.getlist("serial")
        w_starts = request.form.getlist("w_start[]") or request.form.getlist("w_start")
        w_ends = request.form.getlist("w_end[]") or request.form.getlist("w_end")
        details_list = request.form.getlist("details[]") or request.form.getlist("details")

        if not item_ids:
            flash("يرجى إضافة صنف واحد على الأقل.", "warning")
            return redirect(url_for("portal.inventory_return_voucher_new"))

        v = InvReturnVoucher(
            voucher_no=voucher_no,
            voucher_date=voucher_date,
            to_warehouse_id=to_warehouse_id,
            from_room_name=from_room_name,
            note=note or None,
            created_by_id=current_user.id,
            created_at=datetime.utcnow(),
        )
        db.session.add(v)
        db.session.flush()

        any_line = False
        for i, raw_item_id in enumerate(item_ids):
            if not raw_item_id:
                continue
            try:
                item_id = int(raw_item_id)
            except Exception:
                continue

            raw_qty = (qtys[i] if i < len(qtys) else "1")
            try:
                qty = float(raw_qty) if raw_qty not in (None, "") else 1.0
            except Exception:
                qty = 1.0

            ln = InvReturnVoucherLine(
                voucher_id=v.id,
                item_id=item_id,
                qty=qty,
                serial=(serials[i] if i < len(serials) else None) or None,
                warranty_start=(w_starts[i] if i < len(w_starts) else None) or None,
                warranty_end=(w_ends[i] if i < len(w_ends) else None) or None,
                details=(details_list[i] if i < len(details_list) else None) or None,
            )
            db.session.add(ln)
            any_line = True

        if not any_line:
            db.session.rollback()
            flash("يرجى إضافة صنف واحد على الأقل.", "warning")
            return redirect(url_for("portal.inventory_return_voucher_new"))

        # Audit + commit voucher/lines first (protect from attachments failure)
        try:
            db.session.add(AuditLog(
                user_id=current_user.id,
                action="INV_RETURN_CREATE",
                note=f"إنشاء سند إرجاع ({voucher_no})",
                target_type="INV_RETURN_VOUCHER",
                target_id=v.id,
                created_at=datetime.utcnow(),
            ))
        except Exception:
            pass

        db.session.commit()

        # Attachments (best effort)
        files = request.files.getlist("attachments")
        if files:
            try:
                from uuid import uuid4

                base = os.path.join(current_app.instance_path, "uploads", "inventory", "return", str(v.id))
                os.makedirs(base, exist_ok=True)

                for f in files:
                    if not f or not getattr(f, "filename", ""):
                        continue
                    original = f.filename
                    stored = f"{uuid4().hex}_{secure_filename(original) or 'file'}"
                    full_path = os.path.join(base, stored)
                    f.save(full_path)

                    rel = os.path.relpath(full_path, current_app.instance_path)
                    db.session.add(
                        InvReturnVoucherAttachment(
                            voucher_id=v.id,
                            original_name=original,
                            stored_name=stored,
                            file_path=rel,
                            uploaded_by_id=current_user.id,
                            uploaded_at=datetime.utcnow(),
                        )
                    )
                db.session.commit()
            except Exception:
                db.session.rollback()
                flash("تم حفظ السند، لكن تعذّر حفظ المرفقات.", "warning")

        flash("تم حفظ سند الإرجاع.", "success")
        return redirect(url_for("portal.inventory_return_voucher_view", v_id=v.id))

    form = {
        "voucher_no": "",
        "voucher_date": "",
        "to_warehouse_id": None,
        "from_room_name": "",
        "note": "",
        "lines": [],
    }

    return render_template(
        "portal/inventory/return_voucher_new.html",
        warehouses=warehouses,
        items=items,
        rooms=rooms,
        form=form,
    )


@portal_bp.route("/inventory/vouchers/return/list")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_return_voucher_list():
    warehouses = InvWarehouse.query.filter(InvWarehouse.is_active == True).order_by(InvWarehouse.name.asc()).all()  # noqa: E712

    voucher_no = (request.args.get("voucher_no") or "").strip()
    from_date = request.args.get("from_date") or ""
    to_date = request.args.get("to_date") or ""
    to_warehouse_id = request.args.get("to_warehouse_id") or ""
    contains = (request.args.get("contains") or "").strip()

    qry = InvReturnVoucher.query

    if voucher_no:
        qry = qry.filter(InvReturnVoucher.voucher_no.ilike(f"%{voucher_no}%"))

    if from_date:
        qry = qry.filter(InvReturnVoucher.voucher_date >= from_date)
    if to_date:
        qry = qry.filter(InvReturnVoucher.voucher_date <= to_date)

    if to_warehouse_id and to_warehouse_id.isdigit():
        qry = qry.filter(InvReturnVoucher.to_warehouse_id == int(to_warehouse_id))

    if contains:
        like = f"%{contains}%"
        sub = (
            db.session.query(InvReturnVoucherLine.id)
            .join(InvItem, InvItem.id == InvReturnVoucherLine.item_id)
            .filter(InvReturnVoucherLine.voucher_id == InvReturnVoucher.id)
            .filter(or_(InvItem.name.ilike(like), InvItem.code.ilike(like)))
            .exists()
        )
        qry = qry.filter(sub)

    rows = qry.order_by(InvReturnVoucher.voucher_date.desc(), InvReturnVoucher.id.desc()).limit(500).all()

    selected = {
        "voucher_no": voucher_no,
        "from_date": from_date,
        "to_date": to_date,
        "to_warehouse_id": int(to_warehouse_id) if to_warehouse_id.isdigit() else None,
        "contains": contains,
    }

    return render_template(
        "portal/inventory/return_voucher_list.html",
        warehouses=warehouses,
        rows=rows,
        selected=selected,
    )


@portal_bp.route("/inventory/vouchers/return/<int:v_id>")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_return_voucher_view(v_id: int):
    v = InvReturnVoucher.query.get_or_404(v_id)
    if not (current_user.has_perm(STORE_READ) or current_user.has_perm(STORE_MANAGE) or (v.created_by_id == current_user.id)):
        abort(403)

    lines = InvReturnVoucherLine.query.filter(InvReturnVoucherLine.voucher_id == v.id).order_by(InvReturnVoucherLine.id.asc()).all()
    attachments = InvReturnVoucherAttachment.query.filter(InvReturnVoucherAttachment.voucher_id == v.id).order_by(InvReturnVoucherAttachment.id.asc()).all()

    return render_template(
        "portal/inventory/return_voucher_view.html",
        v=v,
        lines=lines,
        attachments=attachments,
    )


@portal_bp.route("/inventory/vouchers/return/attachment/<int:att_id>/download")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_return_attachment_download(att_id: int):
    att = InvReturnVoucherAttachment.query.get_or_404(att_id)
    v = InvReturnVoucher.query.get(att.voucher_id)
    if not v:
        abort(404)

    if not (current_user.has_perm(STORE_READ) or current_user.has_perm(STORE_MANAGE) or (v.created_by_id == current_user.id)):
        abort(403)

    abs_path = os.path.join(current_app.instance_path, att.file_path)
    if not os.path.exists(abs_path):
        abort(404)

    directory = os.path.dirname(abs_path)
    filename = os.path.basename(abs_path)
    return send_from_directory(directory, filename, as_attachment=True, download_name=att.original_name)


@portal_bp.route("/inventory/admin/warehouses", methods=["GET", "POST"])
@login_required
@_perm(STORE_MANAGE)
def inventory_admin_warehouses():
    edit_id = request.args.get("edit_id")
    edit = InvWarehouse.query.get(int(edit_id)) if (edit_id and edit_id.isdigit()) else None

    if request.method == "POST":
        action = request.form.get("action") or "create"
        if action == "delete":
            wid = request.form.get("id")
            w = InvWarehouse.query.get(int(wid)) if (wid and wid.isdigit()) else None
            if w:
                db.session.delete(w)
                db.session.commit()
                flash("تم حذف المخزن.", "success")
            return redirect(url_for("portal.inventory_admin_warehouses"))

        name = (request.form.get("name") or "").strip()
        code = (request.form.get("code") or "").strip() or None
        note = (request.form.get("note") or "").strip() or None
        is_active = bool(request.form.get("is_active"))

        if not name:
            flash("يرجى تعبئة اسم المخزن.", "warning")
            return redirect(url_for("portal.inventory_admin_warehouses", edit_id=edit.id) if edit else url_for("portal.inventory_admin_warehouses"))

        if action == "update":
            wid = request.form.get("id")
            w = InvWarehouse.query.get(int(wid)) if (wid and wid.isdigit()) else None
            if not w:
                flash("المخزن غير موجود.", "warning")
                return redirect(url_for("portal.inventory_admin_warehouses"))
            w.name = name
            w.code = code
            w.note = note
            w.is_active = is_active
            db.session.commit()
            flash("تم تحديث المخزن.", "success")
            return redirect(url_for("portal.inventory_admin_warehouses"))

        w = InvWarehouse(name=name, code=code, note=note, is_active=is_active, created_at=datetime.utcnow())
        db.session.add(w)
        db.session.commit()
        flash("تم إضافة المخزن.", "success")
        return redirect(url_for("portal.inventory_admin_warehouses"))

    rows = InvWarehouse.query.order_by(InvWarehouse.id.desc()).all()
    return render_template("portal/inventory/admin_warehouses.html", rows=rows, edit=edit)


@portal_bp.route("/inventory/admin/categories", methods=["GET", "POST"])
@login_required
@_perm(STORE_MANAGE)
def inventory_admin_categories():
    edit_id = request.args.get("edit_id")
    edit = InvItemCategory.query.get(int(edit_id)) if (edit_id and edit_id.isdigit()) else None

    if request.method == "POST":
        action = request.form.get("action") or "create"
        if action == "delete":
            cid = request.form.get("id")
            c = InvItemCategory.query.get(int(cid)) if (cid and cid.isdigit()) else None
            if c:
                db.session.delete(c)
                db.session.commit()
                flash("تم حذف التصنيف.", "success")
            return redirect(url_for("portal.inventory_admin_categories"))

        name = (request.form.get("name") or "").strip()
        is_active = bool(request.form.get("is_active"))

        if not name:
            flash("يرجى تعبئة الاسم.", "warning")
            return redirect(url_for("portal.inventory_admin_categories", edit_id=edit.id) if edit else url_for("portal.inventory_admin_categories"))

        if action == "update":
            cid = request.form.get("id")
            c = InvItemCategory.query.get(int(cid)) if (cid and cid.isdigit()) else None
            if not c:
                flash("التصنيف غير موجود.", "warning")
                return redirect(url_for("portal.inventory_admin_categories"))
            c.name = name
            c.is_active = is_active
            db.session.commit()
            flash("تم تحديث التصنيف.", "success")
            return redirect(url_for("portal.inventory_admin_categories"))

        c = InvItemCategory(name=name, is_active=is_active, created_at=datetime.utcnow())
        db.session.add(c)
        db.session.commit()
        flash("تم إضافة التصنيف.", "success")
        return redirect(url_for("portal.inventory_admin_categories"))

    rows = InvItemCategory.query.order_by(InvItemCategory.id.desc()).all()
    return render_template("portal/inventory/admin_categories.html", rows=rows, edit=edit)


@portal_bp.route("/inventory/admin/items", methods=["GET", "POST"])
@login_required
@_perm(STORE_MANAGE)
def inventory_admin_items():
    categories = InvItemCategory.query.filter(InvItemCategory.is_active == True).order_by(InvItemCategory.name.asc()).all()  # noqa: E712

    edit_id = request.args.get("edit_id")
    edit = InvItem.query.get(int(edit_id)) if (edit_id and edit_id.isdigit()) else None

    if request.method == "POST":
        action = request.form.get("action") or "create"
        if action == "delete":
            iid = request.form.get("id")
            it = InvItem.query.get(int(iid)) if (iid and iid.isdigit()) else None
            if it:
                db.session.delete(it)
                db.session.commit()
                flash("تم حذف الصنف.", "success")
            return redirect(url_for("portal.inventory_admin_items"))

        name = (request.form.get("name") or "").strip()
        code = (request.form.get("code") or "").strip() or None
        unit = (request.form.get("unit") or "").strip() or None
        note = (request.form.get("note") or "").strip() or None
        cat_id = request.form.get("category_id")
        is_active = bool(request.form.get("is_active"))

        cat_id_val = int(cat_id) if (cat_id and cat_id.isdigit()) else None

        if not name:
            flash("يرجى تعبئة اسم الصنف.", "warning")
            return redirect(url_for("portal.inventory_admin_items", edit_id=edit.id) if edit else url_for("portal.inventory_admin_items"))

        if action == "update":
            iid = request.form.get("id")
            it = InvItem.query.get(int(iid)) if (iid and iid.isdigit()) else None
            if not it:
                flash("الصنف غير موجود.", "warning")
                return redirect(url_for("portal.inventory_admin_items"))
            it.name = name
            it.code = code
            it.unit = unit
            it.note = note
            it.category_id = cat_id_val
            it.is_active = is_active
            db.session.commit()
            flash("تم تحديث الصنف.", "success")
            return redirect(url_for("portal.inventory_admin_items"))

        it = InvItem(name=name, code=code, unit=unit, note=note, category_id=cat_id_val, is_active=is_active, created_at=datetime.utcnow())
        db.session.add(it)
        db.session.commit()
        flash("تم إضافة الصنف.", "success")
        return redirect(url_for("portal.inventory_admin_items"))

    rows = InvItem.query.order_by(InvItem.id.desc()).all()
    return render_template("portal/inventory/admin_items.html", rows=rows, categories=categories, edit=edit)


# ==========================================================
# Inventory: Stocktake (سندات الجرد)
# ==========================================================

@portal_bp.route("/inventory/vouchers/stocktake/new", methods=["GET", "POST"])
@login_required
@_perm(STORE_MANAGE)
def inventory_stocktake_voucher_new():
    warehouses = InvWarehouse.query.filter(InvWarehouse.is_active == True).order_by(InvWarehouse.name.asc()).all()  # noqa: E712
    categories = InvItemCategory.query.filter(InvItemCategory.is_active == True).order_by(InvItemCategory.name.asc()).all()  # noqa: E712
    items = InvItem.query.filter(InvItem.is_active == True).order_by(InvItem.name.asc()).all()  # noqa: E712

    if request.method == "POST":
        voucher_no = (request.form.get("voucher_no") or "").strip()
        voucher_date = (request.form.get("voucher_date") or "").strip()
        warehouse_id = request.form.get("warehouse_id")
        note = (request.form.get("note") or "").strip()

        if not voucher_no or not voucher_date or not warehouse_id:
            flash("يرجى تعبئة رقم السند + تاريخ السند + المخزن.", "warning")
            return redirect(url_for("portal.inventory_stocktake_voucher_new"))

        try:
            warehouse_id = int(warehouse_id)
        except Exception:
            flash("المخزن غير صالح.", "warning")
            return redirect(url_for("portal.inventory_stocktake_voucher_new"))

        item_ids = request.form.getlist("item_id[]") or request.form.getlist("item_id")
        qtys = request.form.getlist("qty[]") or request.form.getlist("qty")
        serials = request.form.getlist("serial[]") or request.form.getlist("serial")
        w_starts = request.form.getlist("w_start[]") or request.form.getlist("w_start")
        w_ends = request.form.getlist("w_end[]") or request.form.getlist("w_end")
        details_list = request.form.getlist("details[]") or request.form.getlist("details")

        if not item_ids:
            flash("يرجى إضافة صنف واحد على الأقل.", "warning")
            return redirect(url_for("portal.inventory_stocktake_voucher_new"))

        v = InvStocktakeVoucher(
            voucher_no=voucher_no,
            voucher_date=voucher_date,
            warehouse_id=warehouse_id,
            note=note or None,
            created_by_id=current_user.id,
            created_at=datetime.utcnow(),
        )
        db.session.add(v)
        db.session.flush()

        any_line = False
        for i, raw_item_id in enumerate(item_ids):
            if not raw_item_id:
                continue
            try:
                item_id = int(raw_item_id)
            except Exception:
                continue

            raw_qty = (qtys[i] if i < len(qtys) else "0")
            try:
                qty = float(raw_qty) if raw_qty not in (None, "") else 0.0
            except Exception:
                qty = 0.0

            ln = InvStocktakeVoucherLine(
                voucher_id=v.id,
                item_id=item_id,
                qty=qty,
                serial=(serials[i] if i < len(serials) else None) or None,
                warranty_start=(w_starts[i] if i < len(w_starts) else None) or None,
                warranty_end=(w_ends[i] if i < len(w_ends) else None) or None,
                details=(details_list[i] if i < len(details_list) else None) or None,
            )
            db.session.add(ln)
            any_line = True

        if not any_line:
            db.session.rollback()
            flash("يرجى إضافة صنف واحد على الأقل.", "warning")
            return redirect(url_for("portal.inventory_stocktake_voucher_new"))

        # Audit + commit voucher/lines first (protect from attachments failure)
        try:
            db.session.add(AuditLog(
                user_id=current_user.id,
                action="INV_STOCKTAKE_CREATE",
                note=f"إنشاء سند جرد ({voucher_no})",
                target_type="INV_STOCKTAKE_VOUCHER",
                target_id=v.id,
                created_at=datetime.utcnow(),
            ))
        except Exception:
            pass

        db.session.commit()

        # Attachments (best effort)
        files = request.files.getlist("attachments")
        if files:
            try:
                from uuid import uuid4

                base = os.path.join(current_app.instance_path, "uploads", "inventory", "stocktake", str(v.id))
                os.makedirs(base, exist_ok=True)

                for f in files:
                    if not f or not getattr(f, "filename", ""):
                        continue
                    original = f.filename
                    stored = f"{uuid4().hex}_{secure_filename(original) or 'file'}"
                    full_path = os.path.join(base, stored)
                    f.save(full_path)

                    rel = os.path.relpath(full_path, current_app.instance_path)
                    db.session.add(
                        InvStocktakeVoucherAttachment(
                            voucher_id=v.id,
                            original_name=original,
                            stored_name=stored,
                            file_path=rel,
                            uploaded_by_id=current_user.id,
                            uploaded_at=datetime.utcnow(),
                        )
                    )
                db.session.commit()
            except Exception:
                db.session.rollback()
                flash("تم حفظ السند، لكن تعذّر حفظ المرفقات.", "warning")

        flash("تم حفظ سند الجرد.", "success")
        return redirect(url_for("portal.inventory_stocktake_voucher_view", v_id=v.id))

    form = {
        "voucher_no": "",
        "voucher_date": "",
        "warehouse_id": None,
        "note": "",
        "lines": [],
    }
    return render_template(
        "portal/inventory/stocktake_voucher_new.html",
        warehouses=warehouses,
        categories=categories,
        items=items,
        form=form,
    )


@portal_bp.route("/inventory/vouchers/stocktake/list")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_stocktake_voucher_list():
    warehouses = InvWarehouse.query.filter(InvWarehouse.is_active == True).order_by(InvWarehouse.name.asc()).all()  # noqa: E712

    voucher_no = (request.args.get("voucher_no") or "").strip()
    from_date = request.args.get("from_date") or ""
    to_date = request.args.get("to_date") or ""
    warehouse_id = request.args.get("warehouse_id") or ""

    qry = InvStocktakeVoucher.query
    if voucher_no:
        qry = qry.filter(InvStocktakeVoucher.voucher_no.ilike(f"%{voucher_no}%"))
    if from_date:
        qry = qry.filter(InvStocktakeVoucher.voucher_date >= from_date)
    if to_date:
        qry = qry.filter(InvStocktakeVoucher.voucher_date <= to_date)
    if warehouse_id and warehouse_id.isdigit():
        qry = qry.filter(InvStocktakeVoucher.warehouse_id == int(warehouse_id))

    rows = qry.order_by(InvStocktakeVoucher.id.desc()).limit(500).all()

    selected = {
        "voucher_no": voucher_no,
        "from_date": from_date,
        "to_date": to_date,
        "warehouse_id": int(warehouse_id) if warehouse_id.isdigit() else None,
    }

    return render_template(
        "portal/inventory/stocktake_voucher_list.html",
        warehouses=warehouses,
        rows=rows,
        selected=selected,
    )


@portal_bp.route("/inventory/vouchers/stocktake/<int:v_id>")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_stocktake_voucher_view(v_id: int):
    v = InvStocktakeVoucher.query.get_or_404(v_id)
    if not (current_user.has_perm(STORE_READ) or current_user.has_perm(STORE_MANAGE) or (v.created_by_id == current_user.id)):
        abort(403)

    lines = InvStocktakeVoucherLine.query.filter(InvStocktakeVoucherLine.voucher_id == v.id).order_by(InvStocktakeVoucherLine.id.asc()).all()
    attachments = InvStocktakeVoucherAttachment.query.filter(InvStocktakeVoucherAttachment.voucher_id == v.id).order_by(InvStocktakeVoucherAttachment.id.asc()).all()

    return render_template(
        "portal/inventory/stocktake_voucher_view.html",
        v=v,
        lines=lines,
        attachments=attachments,
    )


@portal_bp.route("/inventory/vouchers/stocktake/attachment/<int:att_id>/download")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_stocktake_attachment_download(att_id: int):
    att = InvStocktakeVoucherAttachment.query.get_or_404(att_id)
    v = InvStocktakeVoucher.query.get(att.voucher_id)
    if not v:
        abort(404)
    if not (current_user.has_perm(STORE_READ) or current_user.has_perm(STORE_MANAGE) or (v.created_by_id == current_user.id)):
        abort(403)

    abs_path = os.path.join(current_app.instance_path, att.file_path)
    if not os.path.exists(abs_path):
        abort(404)

    directory = os.path.dirname(abs_path)
    filename = os.path.basename(abs_path)
    return send_from_directory(directory, filename, as_attachment=True, download_name=att.original_name)


# ==========================================================
# Inventory: Custody (العهدة)
# ==========================================================

@portal_bp.route("/inventory/custody/vouchers/new", methods=["GET", "POST"])
@login_required
@_perm(STORE_MANAGE)
def inventory_custody_voucher_new():
    categories = InvItemCategory.query.filter(InvItemCategory.is_active == True).order_by(InvItemCategory.name.asc()).all()  # noqa: E712
    items = InvItem.query.filter(InvItem.is_active == True).order_by(InvItem.name.asc()).all()  # noqa: E712
    rooms = InvRoom.query.filter(InvRoom.is_active == True).order_by(InvRoom.name.asc()).all()  # noqa: E712
    users = User.query.order_by(func.coalesce(User.name, User.email).asc(), User.id.asc()).all()

    if request.method == "POST":
        holder_kind = (request.form.get("holder_kind") or "EMPLOYEE").strip().upper()
        voucher_no = (request.form.get("voucher_no") or "").strip()
        voucher_date = (request.form.get("voucher_date") or "").strip()
        holder_user_id = request.form.get("holder_user_id")
        holder_room_id = request.form.get("holder_room_id")
        note = (request.form.get("note") or "").strip()

        if not voucher_no or not voucher_date:
            flash("يرجى تعبئة رقم السند + تاريخ السند.", "warning")
            return redirect(url_for("portal.inventory_custody_voucher_new"))

        holder_user_id_val = None
        holder_room_id_val = None
        if holder_kind == "ROOM":
            if holder_room_id and holder_room_id.isdigit():
                holder_room_id_val = int(holder_room_id)
            if not holder_room_id_val:
                flash("يرجى اختيار الغرفة.", "warning")
                return redirect(url_for("portal.inventory_custody_voucher_new"))
        else:
            if holder_user_id and holder_user_id.isdigit():
                holder_user_id_val = int(holder_user_id)
            if not holder_user_id_val:
                flash("يرجى اختيار الموظف.", "warning")
                return redirect(url_for("portal.inventory_custody_voucher_new"))

        item_ids = request.form.getlist("item_id[]") or request.form.getlist("item_id")
        qtys = request.form.getlist("qty[]") or request.form.getlist("qty")
        serials = request.form.getlist("serial[]") or request.form.getlist("serial")
        w_starts = request.form.getlist("w_start[]") or request.form.getlist("w_start")
        w_ends = request.form.getlist("w_end[]") or request.form.getlist("w_end")
        details_list = request.form.getlist("details[]") or request.form.getlist("details")

        if not item_ids:
            flash("يرجى إضافة صنف واحد على الأقل.", "warning")
            return redirect(url_for("portal.inventory_custody_voucher_new"))

        v = InvCustodyVoucher(
            holder_kind=holder_kind,
            voucher_no=voucher_no,
            voucher_date=voucher_date,
            holder_user_id=holder_user_id_val,
            holder_room_id=holder_room_id_val,
            note=note or None,
            created_by_id=current_user.id,
            created_at=datetime.utcnow(),
        )
        db.session.add(v)
        db.session.flush()

        any_line = False
        for i, raw_item_id in enumerate(item_ids):
            if not raw_item_id:
                continue
            try:
                item_id = int(raw_item_id)
            except Exception:
                continue

            raw_qty = (qtys[i] if i < len(qtys) else "1")
            try:
                qty = float(raw_qty) if raw_qty not in (None, "") else 1.0
            except Exception:
                qty = 1.0

            ln = InvCustodyVoucherLine(
                voucher_id=v.id,
                item_id=item_id,
                qty=qty,
                serial=(serials[i] if i < len(serials) else None) or None,
                warranty_start=(w_starts[i] if i < len(w_starts) else None) or None,
                warranty_end=(w_ends[i] if i < len(w_ends) else None) or None,
                details=(details_list[i] if i < len(details_list) else None) or None,
            )
            db.session.add(ln)
            any_line = True

        if not any_line:
            db.session.rollback()
            flash("يرجى إضافة صنف واحد على الأقل.", "warning")
            return redirect(url_for("portal.inventory_custody_voucher_new"))

        # Audit + commit voucher/lines first (protect from attachments failure)
        try:
            db.session.add(AuditLog(
                user_id=current_user.id,
                action="INV_CUSTODY_CREATE",
                note=f"إنشاء سند عهدة ({voucher_no})",
                target_type="INV_CUSTODY_VOUCHER",
                target_id=v.id,
                created_at=datetime.utcnow(),
            ))
        except Exception:
            pass

        db.session.commit()

        # Attachments (best effort)
        files = request.files.getlist("attachments")
        if files:
            try:
                from uuid import uuid4

                base = os.path.join(current_app.instance_path, "uploads", "inventory", "custody", str(v.id))
                os.makedirs(base, exist_ok=True)

                for f in files:
                    if not f or not getattr(f, "filename", ""):
                        continue
                    original = f.filename
                    stored = f"{uuid4().hex}_{secure_filename(original) or 'file'}"
                    full_path = os.path.join(base, stored)
                    f.save(full_path)

                    rel = os.path.relpath(full_path, current_app.instance_path)
                    db.session.add(
                        InvCustodyVoucherAttachment(
                            voucher_id=v.id,
                            original_name=original,
                            stored_name=stored,
                            file_path=rel,
                            uploaded_by_id=current_user.id,
                            uploaded_at=datetime.utcnow(),
                        )
                    )
                db.session.commit()
            except Exception:
                db.session.rollback()
                flash("تم حفظ السند، لكن تعذّر حفظ المرفقات.", "warning")

        flash("تم حفظ سند العهدة.", "success")
        return redirect(url_for("portal.inventory_custody_voucher_view", v_id=v.id))

    form = {
        "holder_kind": "EMPLOYEE",
        "voucher_no": "",
        "voucher_date": "",
        "holder_user_id": None,
        "holder_room_id": None,
        "note": "",
        "lines": [],
    }
    return render_template(
        "portal/inventory/custody_voucher_new.html",
        categories=categories,
        items=items,
        rooms=rooms,
        users=users,
        form=form,
    )


@portal_bp.route("/inventory/custody/vouchers/list")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_custody_voucher_list():
    rooms = InvRoom.query.filter(InvRoom.is_active == True).order_by(InvRoom.name.asc()).all()  # noqa: E712
    users = User.query.order_by(func.coalesce(User.name, User.email).asc(), User.id.asc()).all()

    voucher_no = (request.args.get("voucher_no") or "").strip()
    from_date = request.args.get("from_date") or ""
    to_date = request.args.get("to_date") or ""
    holder_kind = (request.args.get("holder_kind") or "").strip().upper()
    holder_user_id = request.args.get("holder_user_id") or ""
    holder_room_id = request.args.get("holder_room_id") or ""

    qry = InvCustodyVoucher.query
    if voucher_no:
        qry = qry.filter(InvCustodyVoucher.voucher_no.ilike(f"%{voucher_no}%"))
    if from_date:
        qry = qry.filter(InvCustodyVoucher.voucher_date >= from_date)
    if to_date:
        qry = qry.filter(InvCustodyVoucher.voucher_date <= to_date)
    if holder_kind in ("EMPLOYEE", "ROOM"):
        qry = qry.filter(InvCustodyVoucher.holder_kind == holder_kind)
    if holder_user_id and holder_user_id.isdigit():
        qry = qry.filter(InvCustodyVoucher.holder_user_id == int(holder_user_id))
    if holder_room_id and holder_room_id.isdigit():
        qry = qry.filter(InvCustodyVoucher.holder_room_id == int(holder_room_id))

    rows = qry.order_by(InvCustodyVoucher.id.desc()).limit(500).all()

    selected = {
        "voucher_no": voucher_no,
        "from_date": from_date,
        "to_date": to_date,
        "holder_kind": holder_kind,
        "holder_user_id": int(holder_user_id) if holder_user_id.isdigit() else None,
        "holder_room_id": int(holder_room_id) if holder_room_id.isdigit() else None,
    }

    return render_template(
        "portal/inventory/custody_voucher_list.html",
        rows=rows,
        rooms=rooms,
        users=users,
        selected=selected,
    )


@portal_bp.route("/inventory/custody/vouchers/<int:v_id>")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_custody_voucher_view(v_id: int):
    v = InvCustodyVoucher.query.get_or_404(v_id)
    if not (current_user.has_perm(STORE_READ) or current_user.has_perm(STORE_MANAGE) or (v.created_by_id == current_user.id)):
        abort(403)

    lines = InvCustodyVoucherLine.query.filter(InvCustodyVoucherLine.voucher_id == v.id).order_by(InvCustodyVoucherLine.id.asc()).all()
    attachments = InvCustodyVoucherAttachment.query.filter(InvCustodyVoucherAttachment.voucher_id == v.id).order_by(InvCustodyVoucherAttachment.id.asc()).all()

    return render_template(
        "portal/inventory/custody_voucher_view.html",
        v=v,
        lines=lines,
        attachments=attachments,
    )


@portal_bp.route("/inventory/custody/vouchers/attachment/<int:att_id>/download")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_custody_attachment_download(att_id: int):
    att = InvCustodyVoucherAttachment.query.get_or_404(att_id)
    v = InvCustodyVoucher.query.get(att.voucher_id)
    if not v:
        abort(404)
    if not (current_user.has_perm(STORE_READ) or current_user.has_perm(STORE_MANAGE) or (v.created_by_id == current_user.id)):
        abort(403)

    abs_path = os.path.join(current_app.instance_path, att.file_path)
    if not os.path.exists(abs_path):
        abort(404)

    directory = os.path.dirname(abs_path)
    filename = os.path.basename(abs_path)
    return send_from_directory(directory, filename, as_attachment=True, download_name=att.original_name)


@portal_bp.route("/inventory/custody/overview")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_custody_overview():
    """عرض العهدة: يظهر آخر قيد لكل (حامل، صنف) بحسب التاريخ/المعرف."""

    # optional filters
    holder_kind = (request.args.get("holder_kind") or "").strip().upper()
    holder_user_id = request.args.get("holder_user_id") or ""
    holder_room_id = request.args.get("holder_room_id") or ""

    rooms = InvRoom.query.filter(InvRoom.is_active == True).order_by(InvRoom.name.asc()).all()  # noqa: E712
    users = User.query.order_by(func.coalesce(User.name, User.email).asc(), User.id.asc()).all()

    vouchers_q = InvCustodyVoucher.query
    if holder_kind in ("EMPLOYEE", "ROOM"):
        vouchers_q = vouchers_q.filter(InvCustodyVoucher.holder_kind == holder_kind)
    if holder_user_id and holder_user_id.isdigit():
        vouchers_q = vouchers_q.filter(InvCustodyVoucher.holder_user_id == int(holder_user_id))
    if holder_room_id and holder_room_id.isdigit():
        vouchers_q = vouchers_q.filter(InvCustodyVoucher.holder_room_id == int(holder_room_id))

    vouchers = vouchers_q.order_by(InvCustodyVoucher.voucher_date.desc(), InvCustodyVoucher.id.desc()).limit(1500).all()
    if not vouchers:
        return render_template(
            "portal/inventory/custody_overview.html",
            rows=[],
            rooms=rooms,
            users=users,
            selected={
                "holder_kind": holder_kind,
                "holder_user_id": int(holder_user_id) if holder_user_id.isdigit() else None,
                "holder_room_id": int(holder_room_id) if holder_room_id.isdigit() else None,
            },
        )

    v_ids = [v.id for v in vouchers]
    lines = InvCustodyVoucherLine.query.filter(InvCustodyVoucherLine.voucher_id.in_(v_ids)).all()

    # choose last entry per (holder_kind, holder_id, item_id)
    last = {}
    v_by_id = {v.id: v for v in vouchers}
    for ln in lines:
        v = v_by_id.get(ln.voucher_id)
        if not v:
            continue
        holder_id = v.holder_user_id if v.holder_kind == "EMPLOYEE" else v.holder_room_id
        key = (v.holder_kind, holder_id, ln.item_id)
        prev = last.get(key)
        if prev is None:
            last[key] = (v, ln)
        else:
            pv, _pln = prev
            if (v.voucher_date, v.id) > (pv.voucher_date, pv.id):
                last[key] = (v, ln)

    rows = []
    for (hk, hid, item_id), (v, ln) in last.items():
        rows.append({
            "holder_kind": hk,
            "holder_id": hid,
            "holder_label": (v.holder_user.full_name if hk == "EMPLOYEE" and v.holder_user else (v.holder_room.label if v.holder_room else "-")),
            "item": ln.item,
            "qty": ln.qty,
            "serial": ln.serial,
            "voucher_date": v.voucher_date,
            "voucher_id": v.id,
        })

    rows.sort(key=lambda r: (r["holder_kind"], r["holder_label"], (r["item"].name if r["item"] else "")))

    selected = {
        "holder_kind": holder_kind,
        "holder_user_id": int(holder_user_id) if holder_user_id.isdigit() else None,
        "holder_room_id": int(holder_room_id) if holder_room_id.isdigit() else None,
    }

    return render_template(
        "portal/inventory/custody_overview.html",
        rows=rows,
        rooms=rooms,
        users=users,
        selected=selected,
    )


@portal_bp.route("/inventory/custody/items")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_custody_items():
    """الأصناف والعهدة: تجميع على مستوى الصنف (آخر قيود لكل حامل)."""
    # reuse overview rows then aggregate by item
    vouchers = InvCustodyVoucher.query.order_by(InvCustodyVoucher.voucher_date.desc(), InvCustodyVoucher.id.desc()).limit(2000).all()
    if not vouchers:
        return render_template("portal/inventory/custody_items.html", rows=[])

    v_ids = [v.id for v in vouchers]
    lines = InvCustodyVoucherLine.query.filter(InvCustodyVoucherLine.voucher_id.in_(v_ids)).all()

    last = {}
    v_by_id = {v.id: v for v in vouchers}
    for ln in lines:
        v = v_by_id.get(ln.voucher_id)
        if not v:
            continue
        holder_id = v.holder_user_id if v.holder_kind == "EMPLOYEE" else v.holder_room_id
        key = (v.holder_kind, holder_id, ln.item_id)
        prev = last.get(key)
        if prev is None or (v.voucher_date, v.id) > (prev[0].voucher_date, prev[0].id):
            last[key] = (v, ln)

    agg = {}
    for _k, (_v, ln) in last.items():
        if not ln.item:
            continue
        agg.setdefault(ln.item_id, {"item": ln.item, "qty": 0.0, "holders": 0})
        agg[ln.item_id]["qty"] += float(ln.qty or 0)
        agg[ln.item_id]["holders"] += 1

    rows = list(agg.values())
    rows.sort(key=lambda r: (r["item"].name if r["item"] else ""))
    return render_template("portal/inventory/custody_items.html", rows=rows)


# ==========================================================
# Inventory: Reports (التقارير)
# ==========================================================

def _inv_build_balances():
    """Compute current balances per (warehouse_id, item_id).

    Rules:
    - Inbound/Return: +
    - Issue (ROOM): - from warehouse
    - Issue (WAREHOUSE transfer): - from warehouse AND + to warehouse
    - Scrap: -
    - Stocktake: treated as a baseline per (warehouse,item): take latest stocktake qty, then apply movements after it.
    """
    balances: dict[tuple[int, int], float] = {}

    # 1) Stocktake baseline per (warehouse,item)
    baseline_date: dict[tuple[int, int], str] = {}
    baseline_qty: dict[tuple[int, int], float] = {}

    st_rows = (
        db.session.query(
            InvStocktakeVoucher.warehouse_id,
            InvStocktakeVoucherLine.item_id,
            InvStocktakeVoucher.voucher_date,
            InvStocktakeVoucher.id,
            func.sum(InvStocktakeVoucherLine.qty),
        )
        .join(InvStocktakeVoucherLine, InvStocktakeVoucherLine.voucher_id == InvStocktakeVoucher.id)
        .group_by(
            InvStocktakeVoucher.warehouse_id,
            InvStocktakeVoucherLine.item_id,
            InvStocktakeVoucher.voucher_date,
            InvStocktakeVoucher.id,
        )
        .all()
    )
    # keep (date, id) for baseline selection
    baseline_key: dict[tuple[int, int], tuple[str, int]] = {}
    for wh_id, item_id, v_date, v_id, s in st_rows:
        if not wh_id or not item_id or not v_date:
            continue
        k = (int(wh_id), int(item_id))
        cand = (str(v_date), int(v_id))
        prev = baseline_key.get(k)
        if prev is None or cand > prev:
            baseline_key[k] = cand
            baseline_date[k] = str(v_date)
            baseline_qty[k] = float(s or 0)

    for k, qty in baseline_qty.items():
        balances[k] = float(qty or 0)

    def _after_baseline(wh_id: int, item_id: int, v_date: str) -> bool:
        bd = baseline_date.get((wh_id, item_id))
        return (not bd) or (v_date and v_date > bd)

    def _add(wh_id: int, item_id: int, v_date: str, delta: float):
        if not wh_id or not item_id or not v_date:
            return
        wh_id = int(wh_id)
        item_id = int(item_id)
        if not _after_baseline(wh_id, item_id, str(v_date)):
            return
        balances[(wh_id, item_id)] = balances.get((wh_id, item_id), 0.0) + float(delta or 0)

    # 2) Inbound (+)
    inbound_rows = (
        db.session.query(InvInboundVoucher.to_warehouse_id, InvInboundVoucherLine.item_id, InvInboundVoucher.voucher_date, func.sum(InvInboundVoucherLine.qty))
        .join(InvInboundVoucherLine, InvInboundVoucherLine.voucher_id == InvInboundVoucher.id)
        .group_by(InvInboundVoucher.to_warehouse_id, InvInboundVoucherLine.item_id, InvInboundVoucher.voucher_date)
        .all()
    )
    for wh_id, item_id, v_date, s in inbound_rows:
        _add(wh_id, item_id, v_date, float(s or 0))

    # 3) Return (+)
    return_rows = (
        db.session.query(InvReturnVoucher.to_warehouse_id, InvReturnVoucherLine.item_id, InvReturnVoucher.voucher_date, func.sum(InvReturnVoucherLine.qty))
        .join(InvReturnVoucherLine, InvReturnVoucherLine.voucher_id == InvReturnVoucher.id)
        .group_by(InvReturnVoucher.to_warehouse_id, InvReturnVoucherLine.item_id, InvReturnVoucher.voucher_date)
        .all()
    )
    for wh_id, item_id, v_date, s in return_rows:
        _add(wh_id, item_id, v_date, float(s or 0))

    # 4) Issue (ROOM: -, WAREHOUSE transfer: -/+)
    issue_rows = (
        db.session.query(
            InvIssueVoucher.issue_kind,
            InvIssueVoucher.from_warehouse_id,
            InvIssueVoucher.to_warehouse_id,
            InvIssueVoucher.voucher_date,
            InvIssueVoucherLine.item_id,
            func.sum(InvIssueVoucherLine.qty),
        )
        .join(InvIssueVoucherLine, InvIssueVoucherLine.voucher_id == InvIssueVoucher.id)
        .group_by(
            InvIssueVoucher.issue_kind,
            InvIssueVoucher.from_warehouse_id,
            InvIssueVoucher.to_warehouse_id,
            InvIssueVoucher.voucher_date,
            InvIssueVoucherLine.item_id,
        )
        .all()
    )
    for kind, from_wh, to_wh, v_date, item_id, s in issue_rows:
        qty = float(s or 0)
        _add(from_wh, item_id, v_date, -qty)
        if (kind or "").upper() == "WAREHOUSE" and to_wh:
            _add(to_wh, item_id, v_date, qty)

    # 5) Scrap (-)
    scrap_rows = (
        db.session.query(InvScrapVoucher.from_warehouse_id, InvScrapVoucherLine.item_id, InvScrapVoucher.voucher_date, func.sum(InvScrapVoucherLine.qty))
        .join(InvScrapVoucherLine, InvScrapVoucherLine.voucher_id == InvScrapVoucher.id)
        .group_by(InvScrapVoucher.from_warehouse_id, InvScrapVoucherLine.item_id, InvScrapVoucher.voucher_date)
        .all()
    )
    for wh_id, item_id, v_date, s in scrap_rows:
        _add(wh_id, item_id, v_date, -float(s or 0))

    return balances


@portal_bp.route("/inventory/reports/warehouses")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_report_warehouses():
    warehouses = InvWarehouse.query.filter(InvWarehouse.is_active == True).order_by(InvWarehouse.name.asc()).all()  # noqa: E712
    balances = _inv_build_balances()

    rows = []
    for w in warehouses:
        total_qty = 0.0
        item_count = 0
        for (wh_id, _item_id), qty in balances.items():
            if wh_id != w.id:
                continue
            if abs(qty) < 1e-9:
                continue
            item_count += 1
            total_qty += float(qty or 0)
        rows.append({"warehouse": w, "item_count": item_count, "total_qty": total_qty})

    return render_template("portal/inventory/report_warehouses.html", rows=rows)


@portal_bp.route("/inventory/reports/items")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_report_items_all():
    items = InvItem.query.filter(InvItem.is_active == True).order_by(InvItem.name.asc()).all()  # noqa: E712
    balances = _inv_build_balances()

    agg = {}
    for (wh_id, item_id), qty in balances.items():
        agg[item_id] = agg.get(item_id, 0.0) + float(qty or 0)

    rows = []
    for it in items:
        rows.append({"item": it, "qty": agg.get(it.id, 0.0)})

    return render_template("portal/inventory/report_items_all.html", rows=rows)


@portal_bp.route("/inventory/reports/item-card")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_report_item_card():
    warehouses = InvWarehouse.query.filter(InvWarehouse.is_active == True).order_by(InvWarehouse.name.asc()).all()  # noqa: E712
    items = InvItem.query.filter(InvItem.is_active == True).order_by(InvItem.name.asc()).all()  # noqa: E712

    item_id = request.args.get("item_id") or ""
    warehouse_id = request.args.get("warehouse_id") or ""
    from_date = request.args.get("from_date") or ""
    to_date = request.args.get("to_date") or ""

    movements = []
    sel_item = None

    if item_id and item_id.isdigit():
        iid = int(item_id)
        sel_item = InvItem.query.get(iid)

        # inbound
        q = db.session.query(InvInboundVoucher.voucher_date, InvInboundVoucher.voucher_no, InvInboundVoucher.to_warehouse_id, InvInboundVoucherLine.qty)
        q = q.join(InvInboundVoucherLine, InvInboundVoucherLine.voucher_id == InvInboundVoucher.id)
        q = q.filter(InvInboundVoucherLine.item_id == iid)
        if warehouse_id and warehouse_id.isdigit():
            q = q.filter(InvInboundVoucher.to_warehouse_id == int(warehouse_id))
        if from_date:
            q = q.filter(InvInboundVoucher.voucher_date >= from_date)
        if to_date:
            q = q.filter(InvInboundVoucher.voucher_date <= to_date)
        for d, no, wh, qty in q.all():
            movements.append({"date": d, "type": "إدخال", "voucher_no": no, "warehouse_id": wh, "qty": float(qty or 0)})

        # return
        q = db.session.query(InvReturnVoucher.voucher_date, InvReturnVoucher.voucher_no, InvReturnVoucher.to_warehouse_id, InvReturnVoucherLine.qty)
        q = q.join(InvReturnVoucherLine, InvReturnVoucherLine.voucher_id == InvReturnVoucher.id)
        q = q.filter(InvReturnVoucherLine.item_id == iid)
        if warehouse_id and warehouse_id.isdigit():
            q = q.filter(InvReturnVoucher.to_warehouse_id == int(warehouse_id))
        if from_date:
            q = q.filter(InvReturnVoucher.voucher_date >= from_date)
        if to_date:
            q = q.filter(InvReturnVoucher.voucher_date <= to_date)
        for d, no, wh, qty in q.all():
            movements.append({"date": d, "type": "إرجاع", "voucher_no": no, "warehouse_id": wh, "qty": float(qty or 0)})

        # issue (negative)
        q = db.session.query(InvIssueVoucher.voucher_date, InvIssueVoucher.voucher_no, InvIssueVoucher.from_warehouse_id, InvIssueVoucherLine.qty)
        q = q.join(InvIssueVoucherLine, InvIssueVoucherLine.voucher_id == InvIssueVoucher.id)
        q = q.filter(InvIssueVoucherLine.item_id == iid)
        if warehouse_id and warehouse_id.isdigit():
            q = q.filter(InvIssueVoucher.from_warehouse_id == int(warehouse_id))
        if from_date:
            q = q.filter(InvIssueVoucher.voucher_date >= from_date)
        if to_date:
            q = q.filter(InvIssueVoucher.voucher_date <= to_date)
        for d, no, wh, qty in q.all():
            movements.append({"date": d, "type": "صرف", "voucher_no": no, "warehouse_id": wh, "qty": -float(qty or 0)})

        # scrap (negative)
        q = db.session.query(InvScrapVoucher.voucher_date, InvScrapVoucher.voucher_no, InvScrapVoucher.from_warehouse_id, InvScrapVoucherLine.qty)
        q = q.join(InvScrapVoucherLine, InvScrapVoucherLine.voucher_id == InvScrapVoucher.id)
        q = q.filter(InvScrapVoucherLine.item_id == iid)
        if warehouse_id and warehouse_id.isdigit():
            q = q.filter(InvScrapVoucher.from_warehouse_id == int(warehouse_id))
        if from_date:
            q = q.filter(InvScrapVoucher.voucher_date >= from_date)
        if to_date:
            q = q.filter(InvScrapVoucher.voucher_date <= to_date)
        for d, no, wh, qty in q.all():
            movements.append({"date": d, "type": "إتلاف", "voucher_no": no, "warehouse_id": wh, "qty": -float(qty or 0)})

        movements.sort(key=lambda m: (m["date"], m["voucher_no"]))

    selected = {
        "item_id": int(item_id) if item_id.isdigit() else None,
        "warehouse_id": int(warehouse_id) if warehouse_id.isdigit() else None,
        "from_date": from_date,
        "to_date": to_date,
    }

    wh_map = {w.id: w for w in warehouses}
    for m in movements:
        m["warehouse"] = wh_map.get(m["warehouse_id"])

    return render_template(
        "portal/inventory/report_item_card.html",
        warehouses=warehouses,
        items=items,
        selected=selected,
        item=sel_item,
        movements=movements,
    )


@portal_bp.route("/inventory/reports/yearly-vouchers")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_report_yearly_vouchers():
    year = (request.args.get("year") or "").strip()

    def _year_of(d):
        return (d or "")[:4]

    counts = {}

    for cls, label in [(InvIssueVoucher, "صرف"), (InvInboundVoucher, "إدخال"), (InvReturnVoucher, "إرجاع"), (InvScrapVoucher, "إتلاف"), (InvStocktakeVoucher, "جرد")]:
        rows = db.session.query(cls.voucher_date, func.count(cls.id)).group_by(cls.voucher_date).all()
        for d, c in rows:
            y = _year_of(d)
            if not y:
                continue
            if year and y != year:
                continue
            counts.setdefault(y, {})
            counts[y][label] = counts[y].get(label, 0) + int(c or 0)

    years = sorted(counts.keys())
    rows_out = []
    for y in years:
        rows_out.append({"year": y, "data": counts[y]})

    return render_template("portal/inventory/report_yearly_vouchers.html", rows=rows_out, selected_year=year)


@portal_bp.route("/inventory/reports/yearly-qty")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_report_yearly_qty():
    year = (request.args.get("year") or "").strip()

    def _year_of(d):
        return (d or "")[:4]

    sums = {}

    # inbound
    rows = db.session.query(InvInboundVoucher.voucher_date, func.sum(InvInboundVoucherLine.qty)).join(InvInboundVoucherLine, InvInboundVoucherLine.voucher_id == InvInboundVoucher.id).group_by(InvInboundVoucher.voucher_date).all()
    for d, s in rows:
        y = _year_of(d)
        if not y:
            continue
        if year and y != year:
            continue
        sums.setdefault(y, {})
        sums[y]["إدخال"] = sums[y].get("إدخال", 0.0) + float(s or 0)

    # return
    rows = db.session.query(InvReturnVoucher.voucher_date, func.sum(InvReturnVoucherLine.qty)).join(InvReturnVoucherLine, InvReturnVoucherLine.voucher_id == InvReturnVoucher.id).group_by(InvReturnVoucher.voucher_date).all()
    for d, s in rows:
        y = _year_of(d)
        if not y:
            continue
        if year and y != year:
            continue
        sums.setdefault(y, {})
        sums[y]["إرجاع"] = sums[y].get("إرجاع", 0.0) + float(s or 0)

    # issue
    rows = db.session.query(InvIssueVoucher.voucher_date, func.sum(InvIssueVoucherLine.qty)).join(InvIssueVoucherLine, InvIssueVoucherLine.voucher_id == InvIssueVoucher.id).group_by(InvIssueVoucher.voucher_date).all()
    for d, s in rows:
        y = _year_of(d)
        if not y:
            continue
        if year and y != year:
            continue
        sums.setdefault(y, {})
        sums[y]["صرف"] = sums[y].get("صرف", 0.0) + float(s or 0)

    # scrap
    rows = db.session.query(InvScrapVoucher.voucher_date, func.sum(InvScrapVoucherLine.qty)).join(InvScrapVoucherLine, InvScrapVoucherLine.voucher_id == InvScrapVoucher.id).group_by(InvScrapVoucher.voucher_date).all()
    for d, s in rows:
        y = _year_of(d)
        if not y:
            continue
        if year and y != year:
            continue
        sums.setdefault(y, {})
        sums[y]["إتلاف"] = sums[y].get("إتلاف", 0.0) + float(s or 0)

    # stocktake (qty recorded, informational)
    rows = db.session.query(InvStocktakeVoucher.voucher_date, func.sum(InvStocktakeVoucherLine.qty)).join(InvStocktakeVoucherLine, InvStocktakeVoucherLine.voucher_id == InvStocktakeVoucher.id).group_by(InvStocktakeVoucher.voucher_date).all()
    for d, s in rows:
        y = _year_of(d)
        if not y:
            continue
        if year and y != year:
            continue
        sums.setdefault(y, {})
        sums[y]["جرد"] = sums[y].get("جرد", 0.0) + float(s or 0)

    years = sorted(sums.keys())
    rows_out = []
    for y in years:
        rows_out.append({"year": y, "data": sums[y]})

    return render_template("portal/inventory/report_yearly_qty.html", rows=rows_out, selected_year=year)


@portal_bp.route("/inventory/reports/issue-log")
@login_required
@_perm_any(STORE_READ, STORE_MANAGE)
def inventory_report_issue_log():
    warehouses = InvWarehouse.query.filter(InvWarehouse.is_active == True).order_by(InvWarehouse.name.asc()).all()  # noqa: E712

    from_date = request.args.get("from_date") or ""
    to_date = request.args.get("to_date") or ""
    warehouse_id = request.args.get("warehouse_id") or ""
    contains = (request.args.get("contains") or "").strip()

    qry = db.session.query(InvIssueVoucher, InvIssueVoucherLine).join(InvIssueVoucherLine, InvIssueVoucherLine.voucher_id == InvIssueVoucher.id)

    if from_date:
        qry = qry.filter(InvIssueVoucher.voucher_date >= from_date)
    if to_date:
        qry = qry.filter(InvIssueVoucher.voucher_date <= to_date)
    if warehouse_id and warehouse_id.isdigit():
        qry = qry.filter(InvIssueVoucher.from_warehouse_id == int(warehouse_id))
    if contains:
        # filter by item name/code
        qry = qry.join(InvItem, InvItem.id == InvIssueVoucherLine.item_id).filter(or_(InvItem.name.ilike(f"%{contains}%"), InvItem.code.ilike(f"%{contains}%")))

    rows = qry.order_by(InvIssueVoucher.voucher_date.desc(), InvIssueVoucher.id.desc()).limit(1000).all()

    selected = {
        "from_date": from_date,
        "to_date": to_date,
        "warehouse_id": int(warehouse_id) if warehouse_id.isdigit() else None,
        "contains": contains,
    }

    return render_template(
        "portal/inventory/report_issue_log.html",
        warehouses=warehouses,
        rows=rows,
        selected=selected,
    )


# ==========================================================
# Inventory: Control Panel (لوحة التحكم)
# ==========================================================

@portal_bp.route("/inventory/admin/suppliers", methods=["GET", "POST"])
@login_required
@_perm(STORE_MANAGE)
def inventory_admin_suppliers():
    edit_id = request.args.get("edit_id")
    edit = InvSupplier.query.get(int(edit_id)) if (edit_id and edit_id.isdigit()) else None

    if request.method == "POST":
        action = request.form.get("action") or "create"
        if action == "delete":
            sid = request.form.get("id")
            s = InvSupplier.query.get(int(sid)) if (sid and sid.isdigit()) else None
            if s:
                db.session.delete(s)
                db.session.commit()
                flash("تم حذف المورد.", "success")
            return redirect(url_for("portal.inventory_admin_suppliers"))

        name = (request.form.get("name") or "").strip()
        phone = (request.form.get("phone") or "").strip() or None
        email = (request.form.get("email") or "").strip() or None
        address = (request.form.get("address") or "").strip() or None
        note = (request.form.get("note") or "").strip() or None
        is_active = bool(request.form.get("is_active"))

        if not name:
            flash("يرجى تعبئة اسم المورد.", "warning")
            return redirect(url_for("portal.inventory_admin_suppliers", edit_id=edit.id) if edit else url_for("portal.inventory_admin_suppliers"))

        if action == "update":
            sid = request.form.get("id")
            s = InvSupplier.query.get(int(sid)) if (sid and sid.isdigit()) else None
            if not s:
                flash("المورد غير موجود.", "warning")
                return redirect(url_for("portal.inventory_admin_suppliers"))
            s.name = name
            s.phone = phone
            s.email = email
            s.address = address
            s.note = note
            s.is_active = is_active
            db.session.commit()
            flash("تم تحديث المورد.", "success")
            return redirect(url_for("portal.inventory_admin_suppliers"))

        s = InvSupplier(name=name, phone=phone, email=email, address=address, note=note, is_active=is_active, created_at=datetime.utcnow())
        db.session.add(s)
        db.session.commit()
        flash("تم إضافة المورد.", "success")
        return redirect(url_for("portal.inventory_admin_suppliers"))

    rows = InvSupplier.query.order_by(InvSupplier.id.desc()).all()
    return render_template("portal/inventory/admin_suppliers.html", rows=rows, edit=edit)


@portal_bp.route("/inventory/admin/units", methods=["GET", "POST"])
@login_required
@_perm(STORE_MANAGE)
def inventory_admin_units():
    edit_id = request.args.get("edit_id")
    edit = InvUnit.query.get(int(edit_id)) if (edit_id and edit_id.isdigit()) else None

    if request.method == "POST":
        action = request.form.get("action") or "create"
        if action == "delete":
            uid = request.form.get("id")
            u = InvUnit.query.get(int(uid)) if (uid and uid.isdigit()) else None
            if u:
                db.session.delete(u)
                db.session.commit()
                flash("تم حذف الوحدة.", "success")
            return redirect(url_for("portal.inventory_admin_units"))

        name = (request.form.get("name") or "").strip()
        note = (request.form.get("note") or "").strip() or None
        is_active = bool(request.form.get("is_active"))

        if not name:
            flash("يرجى تعبئة اسم الوحدة.", "warning")
            return redirect(url_for("portal.inventory_admin_units", edit_id=edit.id) if edit else url_for("portal.inventory_admin_units"))

        if action == "update":
            uid = request.form.get("id")
            u = InvUnit.query.get(int(uid)) if (uid and uid.isdigit()) else None
            if not u:
                flash("الوحدة غير موجودة.", "warning")
                return redirect(url_for("portal.inventory_admin_units"))
            u.name = name
            u.note = note
            u.is_active = is_active
            db.session.commit()
            flash("تم تحديث الوحدة.", "success")
            return redirect(url_for("portal.inventory_admin_units"))

        u = InvUnit(name=name, note=note, is_active=is_active, created_at=datetime.utcnow())
        db.session.add(u)
        db.session.commit()
        flash("تم إضافة الوحدة.", "success")
        return redirect(url_for("portal.inventory_admin_units"))

    rows = InvUnit.query.order_by(InvUnit.id.desc()).all()
    return render_template("portal/inventory/admin_units.html", rows=rows, edit=edit)


@portal_bp.route("/inventory/admin/rooms", methods=["GET", "POST"])
@login_required
@_perm(STORE_MANAGE)
def inventory_admin_rooms():
    units = InvUnit.query.filter(InvUnit.is_active == True).order_by(InvUnit.name.asc()).all()  # noqa: E712

    edit_id = request.args.get("edit_id")
    edit = InvRoom.query.get(int(edit_id)) if (edit_id and edit_id.isdigit()) else None

    if request.method == "POST":
        action = request.form.get("action") or "create"
        if action == "delete":
            rid = request.form.get("id")
            r = InvRoom.query.get(int(rid)) if (rid and rid.isdigit()) else None
            if r:
                db.session.delete(r)
                db.session.commit()
                flash("تم حذف الغرفة.", "success")
            return redirect(url_for("portal.inventory_admin_rooms"))

        name = (request.form.get("name") or "").strip()
        code = (request.form.get("code") or "").strip() or None
        unit_id = request.form.get("unit_id")
        note = (request.form.get("note") or "").strip() or None
        is_active = bool(request.form.get("is_active"))

        unit_id_val = int(unit_id) if (unit_id and unit_id.isdigit()) else None

        if not name:
            flash("يرجى تعبئة اسم الغرفة.", "warning")
            return redirect(url_for("portal.inventory_admin_rooms", edit_id=edit.id) if edit else url_for("portal.inventory_admin_rooms"))

        if action == "update":
            rid = request.form.get("id")
            r = InvRoom.query.get(int(rid)) if (rid and rid.isdigit()) else None
            if not r:
                flash("الغرفة غير موجودة.", "warning")
                return redirect(url_for("portal.inventory_admin_rooms"))
            r.name = name
            r.code = code
            r.unit_id = unit_id_val
            r.note = note
            r.is_active = is_active
            db.session.commit()
            flash("تم تحديث الغرفة.", "success")
            return redirect(url_for("portal.inventory_admin_rooms"))

        r = InvRoom(name=name, code=code, unit_id=unit_id_val, note=note, is_active=is_active, created_at=datetime.utcnow())
        db.session.add(r)
        db.session.commit()
        flash("تم إضافة الغرفة.", "success")
        return redirect(url_for("portal.inventory_admin_rooms"))

    rows = InvRoom.query.order_by(InvRoom.id.desc()).all()
    return render_template("portal/inventory/admin_rooms.html", rows=rows, edit=edit, units=units)


@portal_bp.route("/inventory/admin/room-requesters", methods=["GET", "POST"])
@login_required
@_perm(STORE_MANAGE)
def inventory_admin_room_requesters():
    users = User.query.order_by(func.coalesce(User.name, User.email).asc(), User.id.asc()).all()
    rooms = InvRoom.query.filter(InvRoom.is_active == True).order_by(InvRoom.name.asc()).all()  # noqa: E712

    if request.method == "POST":
        action = request.form.get("action") or "create"
        if action == "delete":
            rid = request.form.get("id")
            rr = InvRoomRequester.query.get(int(rid)) if (rid and rid.isdigit()) else None
            if rr:
                db.session.delete(rr)
                db.session.commit()
                flash("تم حذف السجل.", "success")
            return redirect(url_for("portal.inventory_admin_room_requesters"))

        user_id = request.form.get("user_id")
        room_id = request.form.get("room_id")
        is_active = bool(request.form.get("is_active"))

        try:
            user_id_val = int(user_id) if user_id else None
        except Exception:
            user_id_val = None
        room_id_val = int(room_id) if (room_id and room_id.isdigit()) else None

        if not user_id_val:
            flash("يرجى اختيار المستخدم.", "warning")
            return redirect(url_for("portal.inventory_admin_room_requesters"))

        rr = InvRoomRequester(user_id=user_id_val, room_id=room_id_val, is_active=is_active, created_at=datetime.utcnow())
        db.session.add(rr)
        db.session.commit()
        flash("تم حفظ السجل.", "success")
        return redirect(url_for("portal.inventory_admin_room_requesters"))

    rows = InvRoomRequester.query.order_by(InvRoomRequester.id.desc()).all()
    return render_template("portal/inventory/admin_room_requesters.html", rows=rows, users=users, rooms=rooms)


@portal_bp.route("/inventory/admin/warehouse-permissions", methods=["GET", "POST"])
@login_required
@_perm(STORE_MANAGE)
def inventory_admin_warehouse_perms():
    users = User.query.order_by(func.coalesce(User.name, User.email).asc(), User.id.asc()).all()
    warehouses = InvWarehouse.query.filter(InvWarehouse.is_active == True).order_by(InvWarehouse.name.asc()).all()  # noqa: E712

    if request.method == "POST":
        action = request.form.get("action") or "create"
        if action == "delete":
            pid = request.form.get("id")
            p = InvWarehousePermission.query.get(int(pid)) if (pid and pid.isdigit()) else None
            if p:
                db.session.delete(p)
                db.session.commit()
                flash("تم حذف الصلاحية.", "success")
            return redirect(url_for("portal.inventory_admin_warehouse_perms"))

        user_id = request.form.get("user_id")
        warehouse_id = request.form.get("warehouse_id")
        if not (user_id and user_id.isdigit() and warehouse_id and warehouse_id.isdigit()):
            flash("يرجى اختيار المستخدم والمخزن.", "warning")
            return redirect(url_for("portal.inventory_admin_warehouse_perms"))

        can_view = bool(request.form.get("can_view"))
        can_issue = bool(request.form.get("can_issue"))
        can_inbound = bool(request.form.get("can_inbound"))
        can_stocktake = bool(request.form.get("can_stocktake"))
        can_manage = bool(request.form.get("can_manage"))

        # upsert
        p = InvWarehousePermission.query.filter_by(user_id=int(user_id), warehouse_id=int(warehouse_id)).first()
        if not p:
            p = InvWarehousePermission(user_id=int(user_id), warehouse_id=int(warehouse_id), created_at=datetime.utcnow())
            db.session.add(p)

        p.can_view = can_view
        p.can_issue = can_issue
        p.can_inbound = can_inbound
        p.can_stocktake = can_stocktake
        p.can_manage = can_manage

        db.session.commit()
        flash("تم حفظ الصلاحية.", "success")
        return redirect(url_for("portal.inventory_admin_warehouse_perms"))

    rows = InvWarehousePermission.query.order_by(InvWarehousePermission.id.desc()).all()
    return render_template("portal/inventory/admin_warehouse_perms.html", rows=rows, users=users, warehouses=warehouses)


@portal_bp.route("/inventory/admin/settings", methods=["GET", "POST"])
@login_required
@_perm(STORE_MANAGE)
def inventory_admin_settings():
    prefix = "INV_"

    if request.method == "POST":
        key = (request.form.get("key") or "").strip().upper()
        value = (request.form.get("value") or "").strip()
        description = (request.form.get("description") or "").strip() or None

        if not key:
            flash("يرجى تعبئة المفتاح.", "warning")
            return redirect(url_for("portal.inventory_admin_settings"))

        if not key.startswith(prefix):
            key = prefix + key

        row = SystemSetting.query.filter(SystemSetting.key == key).first()
        if not row:
            row = SystemSetting(key=key, value=value)
            # NOTE: some deployments do not have `description` column on SystemSetting
            if hasattr(SystemSetting, "description"):
                try:
                    row.description = description
                except Exception:
                    pass
            db.session.add(row)
        else:
            row.value = value
            if hasattr(SystemSetting, "description"):
                try:
                    row.description = description
                except Exception:
                    pass

        db.session.commit()
        flash("تم حفظ الإعداد.", "success")
        return redirect(url_for("portal.inventory_admin_settings"))

    rows = SystemSetting.query.filter(SystemSetting.key.ilike(f"{prefix}%")).order_by(SystemSetting.key.asc()).all()
    return render_template("portal/inventory/admin_settings.html", rows=rows)



# =========================
# HR Training (التدريب)
# =========================

TRAINING_CATEGORY = "TRAINING_CATEGORY"          # تصنيفات الدورات
TRAINING_SPONSOR = "TRAINING_SPONSOR"            # الجهة الممولة/الداعية
TRAINING_CONDITION_FIELD = "TRAINING_CONDITION_FIELD"  # حقول الشروط
COUNTRY_CATEGORY = "COUNTRY"                     # الدول (من ملف لاحقًا)

# System settings keys (defaults apply when per-training override is None)
SS_TRAINING_REQUIRE_MANAGER_APPROVAL = "TRAINING_REQUIRE_MANAGER_APPROVAL"
SS_TRAINING_NEEDS_WINDOW = "TRAINING_NEEDS_WINDOW_OPEN"  # نافذة إدخال الاحتياجات التدريبية
SS_TRAINING_EMPLOYEE_NOTIFICATIONS = "TRAINING_EMPLOYEE_NOTIFICATIONS_ENABLED"
SS_TRAINING_APPLY_COND_PORTAL = "TRAINING_APPLY_CONDITIONS_ON_PORTAL"
SS_TRAINING_APPLY_COND_ADMIN = "TRAINING_APPLY_CONDITIONS_IN_PROGRAM"


def _training_can_manage() -> bool:
    try:
        return bool(
            current_user.has_perm(HR_MASTERDATA_MANAGE)
            or current_user.has_perm(HR_EMP_MANAGE)
            or current_user.has_perm(HR_REQUESTS_VIEW_ALL)
            or current_user.has_perm(HR_REPORTS_VIEW)
        )
    except Exception:
        return False


def _ss_get_bool(key: str, default: bool = False) -> bool:
    try:
        row = SystemSetting.query.filter_by(key=key).first()
        if not row or row.value is None:
            return default
        v = str(row.value).strip().lower()
        if v in ("1", "true", "yes", "on", "نعم", "فعال", "enable", "enabled"):
            return True
        if v in ("0", "false", "no", "off", "لا", "غير فعال", "disable", "disabled"):
            return False
        return default
    except Exception:
        return default


def _ss_set_bool(key: str, val: bool):
    v = "1" if val else "0"
    row = SystemSetting.query.filter_by(key=key).first()
    if not row:
        row = SystemSetting(key=key, value=v)
        db.session.add(row)
    else:
        row.value = v


def _training_effective_bool(program: HRTrainingProgram, field_name: str, sys_key: str, default: bool = False) -> bool:
    try:
        pv = getattr(program, field_name, None)
        if pv is None:
            return _ss_get_bool(sys_key, default)
        return bool(pv)
    except Exception:
        return _ss_get_bool(sys_key, default)


def _training_storage_dir() -> str:
    base = os.path.join(current_app.instance_path, "uploads", "training")
    os.makedirs(base, exist_ok=True)
    return base


def _training_upload_dir(program_id: int) -> str:
    base = os.path.join(_training_storage_dir(), str(program_id))
    os.makedirs(base, exist_ok=True)
    return base


def _save_training_files(files, program_id: int, title: str | None = None) -> int:
    saved = 0
    if not files:
        return 0
    folder = _training_upload_dir(program_id)
    for f in files:
        if not f or not getattr(f, "filename", ""):
            continue
        if not _allowed_file(f.filename):
            continue
        original_name = f.filename
        ext = _clean_suffix(original_name)
        stored_name = f"TRN_{program_id}_{uuid.uuid4().hex}{ext}"
        file_path = os.path.join(folder, stored_name)
        f.save(file_path)

        att = HRTrainingAttachment(
            program_id=program_id,
            title=(title or None),
            original_name=original_name,
            stored_name=stored_name,
            uploaded_by_id=current_user.id,
            uploaded_at=datetime.utcnow(),
        )
        db.session.add(att)
        saved += 1
    return saved


def _training_place_label(country_lookup: HRLookupItem | None) -> str:
    try:
        name_ar = (getattr(country_lookup, "name_ar", "") or "").strip()
        code = (getattr(country_lookup, "code", "") or "").strip().upper()
        if code in ("PS", "PSE", "PALESTINE") or ("فلسطين" in name_ar):
            return "داخل الوطن"
        if not (name_ar or code):
            return "-"
        return "خارج الوطن"
    except Exception:
        return "-"


def _training_notify_published(p: HRTrainingProgram, actor_id: int | None = None):
    """Notify employees (Portal notifications) when a training program is published.

    - If the program is set to "publish by conditions only", we notify only eligible employees.
    """
    try:
        # Respect per-program override or system setting
        if not _training_effective_bool(p, "employee_notifications_enabled", SS_TRAINING_EMPLOYEE_NOTIFICATIONS, True):
            return

        try:
            title = (p.course.name_ar if p.course else "") or (p.program_no or f"#{p.id}")
        except Exception:
            title = p.program_no or f"#{p.id}"

        period = ""
        try:
            if p.nomination_start or p.nomination_end:
                period = f" (الترشيح: {p.nomination_start or '-'} إلى {p.nomination_end or '-'})"
        except Exception:
            period = ""

        scope = ""
        try:
            if getattr(p, "publish_conditions_only", False):
                scope = " (حسب الشروط فقط)"
        except Exception:
            scope = ""

        msg = f"تم نشر تدريب جديد{scope}: {title}{period}. راجع سجل التدريبات للتفاصيل والانتساب."

        # Recipients: all non-admin users (best-effort)
        role_upper = func.upper(func.coalesce(User.role, ""))
        rows = (
            User.query
            .filter(role_upper != "ADMIN")
            .filter(~role_upper.like("SUPER%"))
            .with_entities(User.id)
            .all()
        )
        user_ids = [int(uid) for (uid,) in rows if uid]
        if not user_ids:
            return

        # If published by conditions only -> filter recipients
        try:
            if getattr(p, "publish_conditions_only", False):
                conds = HRTrainingCondition.query.filter_by(program_id=p.id).all()
                if conds:
                    # preload employee files
                    efs = EmployeeFile.query.filter(EmployeeFile.user_id.in_(user_ids)).all()
                    ef_map = {int(ef.user_id): ef for ef in efs if getattr(ef, 'user_id', None) is not None}

                    qual_cache = {}
                    eligible = []
                    for uid in user_ids:
                        ef = ef_map.get(uid)
                        ok = True
                        for cond in conds:
                            if not _training_eval_condition_strict(ef, cond, qual_cache=qual_cache):
                                ok = False
                                break
                        if ok:
                            eligible.append(uid)
                    user_ids = eligible

                if not user_ids:
                    return
        except Exception:
            # If anything goes wrong, fallback to notifying all employees
            pass

        now = datetime.utcnow()
        event_key = uuid.uuid4().hex
        notifs = [
            Notification(
                user_id=uid,
                message=msg,
                type="PORTAL",
                source="portal",
                is_read=False,
                created_at=now,
                actor_id=(int(actor_id) if actor_id else None),
                event_key=event_key,
                is_mirror=False,
            )
            for uid in user_ids
        ]
        db.session.add_all(notifs)
        db.session.commit()
    except Exception:
        try:
            db.session.rollback()
        except Exception:
            pass


def _ensure_training_condition_fields_seeded():
    """Seed default condition fields if none exist."""
    try:
        exists_any = HRLookupItem.query.filter_by(category=TRAINING_CONDITION_FIELD).first()
        if exists_any:
            return
        seed = [
            ("AGE", "العمر", "Age"),
            ("JOB_TITLE", "المسمى الوظيفي", "Job title"),
            ("APPOINTMENT_TYPE", "نوع التعيين", "Appointment type"),
            ("JOB_CATEGORY", "الفئة", "Job category"),
            ("JOB_GRADE", "الدرجة", "Grade"),
            ("DIRECTORATE", "الإدارة العامة", "Directorate"),
            ("DEPARTMENT", "الدائرة/القسم", "Department"),
            ("WORK_GOV", "محافظة العمل", "Work governorate"),
            ("WORK_LOC", "موقع العمل", "Work location"),
            ("ADMIN_TITLE", "المسمى الإداري", "Administrative title"),
            ("SERVICE_YEARS", "سنوات الخدمة", "Years of service"),
            ("EDU_DEGREE", "الدرجة العلمية", "Degree"),
            ("EDU_SPEC", "التخصص", "Specialization"),
        ]
        for i, (code, ar, en) in enumerate(seed):
            db.session.add(HRLookupItem(category=TRAINING_CONDITION_FIELD, code=code, name_ar=ar, name_en=en, sort_order=i*10, is_active=True))
        db.session.commit()
    except Exception:
        try:
            db.session.rollback()
        except Exception:
            pass


def _training_eval_condition(ef: EmployeeFile | None, cond: HRTrainingCondition) -> bool:
    """Best-effort condition evaluation.

    Notes:
    - We support a subset of fields reliably.
    - If a field is unknown or data is missing, we treat it as PASS to avoid blocking enrollment unexpectedly.
    """
    try:
        op = (cond.operator or "EQ").strip().upper()
        v1 = (cond.value1 or "").strip()
        v2 = (cond.value2 or "").strip()

        code = ""
        try:
            code = ((cond.field_lookup.code if cond.field_lookup else "") or "").strip().upper()
        except Exception:
            code = ""

        if not ef or not code:
            return True

        # map user value
        user_val = None
        if code == "AGE":
            bd = _parse_yyyy_mm_dd(getattr(ef, "birth_date", None))
            if not bd:
                return True
            today = date.today()
            user_val = today.year - bd.year - ((today.month, today.day) < (bd.month, bd.day))
        elif code == "SERVICE_YEARS":
            hd = _parse_yyyy_mm_dd(getattr(ef, "hire_date", None))
            if not hd:
                return True
            today = date.today()
            user_val = today.year - hd.year - ((today.month, today.day) < (hd.month, hd.day))
        elif code == "JOB_TITLE":
            user_val = getattr(ef, "job_title_lookup_id", None)
        elif code == "APPOINTMENT_TYPE":
            user_val = getattr(ef, "appointment_type_lookup_id", None)
        elif code == "JOB_CATEGORY":
            user_val = getattr(ef, "job_category_lookup_id", None)
        elif code == "JOB_GRADE":
            user_val = getattr(ef, "job_grade_lookup_id", None)
        elif code == "DIRECTORATE":
            user_val = getattr(ef, "directorate_id", None)
        elif code == "DEPARTMENT":
            user_val = getattr(ef, "department_id", None)
        elif code == "WORK_GOV":
            user_val = getattr(ef, "work_governorate_lookup_id", None)
        elif code == "WORK_LOC":
            user_val = getattr(ef, "work_location_lookup_id", None)
        elif code == "ADMIN_TITLE":
            user_val = getattr(ef, "admin_title_lookup_id", None)
        elif code == "EDU_DEGREE" or code == "EDU_SPEC":
            # Most recent qualification (best-effort)
            try:
                q = (
                    EmployeeQualification.query
                    .filter_by(user_id=getattr(ef, "user_id", None))
                    .order_by(
                        EmployeeQualification.qualification_date.desc().nullslast(),
                        EmployeeQualification.created_at.desc(),
                    )
                    .first()
                )
            except Exception:
                q = None

            if not q:
                return True

            if code == "EDU_DEGREE":
                user_val = getattr(q, "degree_lookup_id", None)
            else:
                user_val = getattr(q, "specialization_lookup_id", None)
        else:
            return True

        # try parse v1/v2 numeric/int
        def _to_num(x: str):
            try:
                if x is None:
                    return None
                s = str(x).strip()
                if not s:
                    return None
                if s.isdigit() or (s.startswith('-') and s[1:].isdigit()):
                    return int(s)
                return float(s)
            except Exception:
                return None

        # comparisons
        if op == "BETWEEN":
            a = _to_num(v1)
            b = _to_num(v2)
            if a is None or b is None or user_val is None:
                # fallback string compare
                return True
            try:
                uv = float(user_val)
            except Exception:
                return True
            lo, hi = (a, b) if a <= b else (b, a)
            return lo <= uv <= hi

        if user_val is None:
            return True

        # numeric
        uv_num = None
        if isinstance(user_val, (int, float)):
            uv_num = float(user_val)

        if op in ("GT", "LT", "GTE", "LTE"):
            a = _to_num(v1)
            if a is None or uv_num is None:
                return True
            if op == "GT":
                return uv_num > float(a)
            if op == "LT":
                return uv_num < float(a)
            if op == "GTE":
                return uv_num >= float(a)
            if op == "LTE":
                return uv_num <= float(a)

        # equality / not equality
        # Allow v1 to be an ID for lookup fields.
        if op == "NEQ":
            try:
                if str(user_val) == v1:
                    return False
                if isinstance(user_val, (int,)) and v1.isdigit() and int(v1) == int(user_val):
                    return False
            except Exception:
                pass
            return True

        # default EQ
        try:
            if str(user_val) == v1:
                return True
            if isinstance(user_val, (int,)) and v1.isdigit() and int(v1) == int(user_val):
                return True
        except Exception:
            pass
        return False

    except Exception:
        return True


def _training_eval_condition_strict(ef: EmployeeFile | None, cond: HRTrainingCondition, qual_cache: dict | None = None) -> bool:
    """Strict evaluation for condition-based publishing/notification.

    Unlike _training_eval_condition(), missing data is treated as FAIL.
    """
    try:
        op = (cond.operator or "EQ").strip().upper()
        v1 = (cond.value1 or "").strip()
        v2 = (cond.value2 or "").strip()

        code = ""
        try:
            code = ((cond.field_lookup.code if cond.field_lookup else "") or "").strip().upper()
        except Exception:
            code = ""

        if not ef or not code:
            return False

        # map user value
        user_val = None
        if code == "AGE":
            bd = _parse_yyyy_mm_dd(getattr(ef, "birth_date", None))
            if not bd:
                return False
            today = date.today()
            user_val = today.year - bd.year - ((today.month, today.day) < (bd.month, bd.day))
        elif code == "SERVICE_YEARS":
            hd = _parse_yyyy_mm_dd(getattr(ef, "hire_date", None))
            if not hd:
                return False
            today = date.today()
            user_val = today.year - hd.year - ((today.month, today.day) < (hd.month, hd.day))
        elif code == "JOB_TITLE":
            user_val = getattr(ef, "job_title_lookup_id", None)
        elif code == "APPOINTMENT_TYPE":
            user_val = getattr(ef, "appointment_type_lookup_id", None)
        elif code == "JOB_CATEGORY":
            user_val = getattr(ef, "job_category_lookup_id", None)
        elif code == "JOB_GRADE":
            user_val = getattr(ef, "job_grade_lookup_id", None)
        elif code == "DIRECTORATE":
            user_val = getattr(ef, "directorate_id", None)
        elif code == "DEPARTMENT":
            user_val = getattr(ef, "department_id", None)
        elif code == "WORK_GOV":
            user_val = getattr(ef, "work_governorate_lookup_id", None)
        elif code == "WORK_LOC":
            user_val = getattr(ef, "work_location_lookup_id", None)
        elif code == "ADMIN_TITLE":
            user_val = getattr(ef, "admin_title_lookup_id", None)
        elif code in ("EDU_DEGREE", "EDU_SPEC"):
            uid = getattr(ef, "user_id", None)
            if not uid:
                return False

            if qual_cache is None:
                qual_cache = {}

            q = qual_cache.get(int(uid))
            if q is None and int(uid) not in qual_cache:
                try:
                    q = (
                        EmployeeQualification.query
                        .filter_by(user_id=int(uid))
                        .order_by(
                            EmployeeQualification.qualification_date.desc().nullslast(),
                            EmployeeQualification.created_at.desc(),
                        )
                        .first()
                    )
                except Exception:
                    q = None
                qual_cache[int(uid)] = q

            if not q:
                return False

            if code == "EDU_DEGREE":
                user_val = getattr(q, "degree_lookup_id", None)
            else:
                user_val = getattr(q, "specialization_lookup_id", None)
        else:
            # Unknown field: fail in strict mode
            return False

        if user_val is None:
            return False

        def _to_num(x: str):
            try:
                if x is None:
                    return None
                s = str(x).strip()
                if not s:
                    return None
                if s.isdigit() or (s.startswith('-') and s[1:].isdigit()):
                    return int(s)
                return float(s)
            except Exception:
                return None

        # BETWEEN
        if op == "BETWEEN":
            a = _to_num(v1)
            b = _to_num(v2)
            if a is None or b is None:
                return False
            try:
                uv = float(user_val)
            except Exception:
                return False
            lo, hi = (a, b) if a <= b else (b, a)
            return lo <= uv <= hi

        # numeric comparisons
        if op in ("GT", "LT", "GTE", "LTE"):
            a = _to_num(v1)
            if a is None:
                return False
            try:
                uv = float(user_val)
            except Exception:
                return False
            if op == "GT":
                return uv > float(a)
            if op == "LT":
                return uv < float(a)
            if op == "GTE":
                return uv >= float(a)
            if op == "LTE":
                return uv <= float(a)

        # equality / not equality
        if op == "NEQ":
            try:
                if isinstance(user_val, int) and v1.isdigit():
                    return int(user_val) != int(v1)
            except Exception:
                pass
            try:
                return str(user_val) != v1
            except Exception:
                return False

        # default EQ
        try:
            if str(user_val) == v1:
                return True
        except Exception:
            pass
        try:
            if isinstance(user_val, int) and v1.isdigit() and int(v1) == int(user_val):
                return True
        except Exception:
            pass
        return False

    except Exception:
        return False


@portal_bp.route("/hr/training/log")
@login_required
@_perm(PORTAL_READ)
def hr_training_log():
    """Training log.

    - HR admins see all programs with filters.
    - Employees see their enrollments + available published programs.
    """
    try:
        if not (current_user.has_perm(HR_READ) or current_user.has_perm(HR_SS_READ) or current_user.has_perm(HR_SS_CREATE)):
            abort(403)
    except Exception:
        abort(403)

    can_manage = _training_can_manage()

    # Filters
    course_id = request.args.get("course_id")
    published = request.args.get("published")
    place = request.args.get("place")
    year = request.args.get("year")
    employee_id = request.args.get("employee_id")

    q = HRTrainingProgram.query

    if course_id and str(course_id).isdigit():
        q = q.filter(HRTrainingProgram.course_id == int(course_id))

    if published in ("1", "0"):
        q = q.filter(HRTrainingProgram.is_published == (published == "1"))

    if year and str(year).isdigit():
        y = int(year)
        q = q.filter(or_(HRTrainingProgram.start_date.like(f"{y}-%"), HRTrainingProgram.end_date.like(f"{y}-%")))

    # employee filter (admin only)
    if can_manage and employee_id and str(employee_id).isdigit():
        eid = int(employee_id)
        q = q.join(HRTrainingEnrollment, HRTrainingEnrollment.program_id == HRTrainingProgram.id).filter(HRTrainingEnrollment.user_id == eid)

    q = q.order_by(func.coalesce(HRTrainingProgram.start_date, "").desc(), HRTrainingProgram.id.desc())

    programs = q.all()

    # Derived maps (labels/availability)
    today_str = date.today().strftime('%Y-%m-%d')
    place_map = {}
    open_map = {}
    for _p in programs:
        try:
            place_map[_p.id] = _training_place_label(getattr(_p, 'country_lookup', None))
        except Exception:
            place_map[_p.id] = '-'
        try:
            s = (_p.nomination_start or '0000-00-00')
            e = (_p.nomination_end or '9999-99-99')
            open_map[_p.id] = bool(_p.is_published and s <= today_str <= e)
        except Exception:
            open_map[_p.id] = False

    # Place filter (IN/OUT)
    if place in ('IN', 'OUT'):
        want = 'داخل الوطن' if place == 'IN' else 'خارج الوطن'
        programs = [pp for pp in programs if place_map.get(pp.id) == want]

    # Employee mode: keep only my enrollments + published
    my_map = {}
    if not can_manage:
        try:
            rows = HRTrainingEnrollment.query.filter_by(user_id=current_user.id).all()
            my_map = {r.program_id: r for r in rows}
        except Exception:
            my_map = {}

        # If program is published by conditions only, hide it from employees who do not match.
        eligible_open = set()
        try:
            ef_me = EmployeeFile.query.filter_by(user_id=current_user.id).first()
        except Exception:
            ef_me = None

        cond_map = {}
        try:
            cond_prog_ids = [pp.id for pp in programs if open_map.get(pp.id) and getattr(pp, 'publish_conditions_only', False)]
            if cond_prog_ids:
                conds = HRTrainingCondition.query.filter(HRTrainingCondition.program_id.in_(cond_prog_ids)).all()
                for c in conds:
                    cond_map.setdefault(int(c.program_id), []).append(c)

            if cond_map:
                qual_cache = {}
                for pid, conds in cond_map.items():
                    ok = True
                    for cond in conds:
                        if not _training_eval_condition_strict(ef_me, cond, qual_cache=qual_cache):
                            ok = False
                            break
                    if ok:
                        eligible_open.add(int(pid))
        except Exception:
            eligible_open = set()

        def _is_visible_open(pp):
            if not open_map.get(pp.id):
                return False
            try:
                if getattr(pp, 'publish_conditions_only', False):
                    # If no conditions defined -> visible; else visible only when eligible
                    if int(pp.id) not in cond_map:
                        return True
                    return int(pp.id) in eligible_open
            except Exception:
                return open_map.get(pp.id)
            return True

        programs = [p for p in programs if (p.id in my_map) or _is_visible_open(p)]

    # lists
    courses = HRTrainingCourse.query.order_by(HRTrainingCourse.sort_order.asc(), HRTrainingCourse.id.asc()).all()

    employees = []
    if can_manage:
        try:
            employees = User.query.order_by(User.full_name.asc()).limit(250).all()
        except Exception:
            employees = []

    return render_template(
        "portal/hr/training/log.html",
        can_manage=can_manage,
        programs=programs,
        courses=courses,
        employees=employees,
        my_map=my_map,
        place_map=place_map,
        open_map=open_map,
        today_str=today_str,
        filters={
            "course_id": course_id or "",
            "published": published or "",
            "place": place or "",
            "year": year or "",
            "employee_id": employee_id or "",
        },
    )


@portal_bp.route("/hr/training/programs/new")
@login_required
@_perm(PORTAL_READ)
def hr_training_program_new():
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    if not _training_can_manage():
        abort(403)

    p = HRTrainingProgram(program_no="")
    p.created_by_id = current_user.id
    p.updated_by_id = current_user.id
    db.session.add(p)
    db.session.commit()
    flash("تم إنشاء تدريب جديد. أكمل البيانات.", "success")
    return redirect(url_for("portal.hr_training_program_info", program_id=p.id))


def _training_wizard_ctx(program: HRTrainingProgram, step: str):
    steps = [
        ("info", "معلومات التدريب", "portal.hr_training_program_info"),
        ("conditions", "شروط الانتساب", "portal.hr_training_program_conditions"),
        ("attachments", "المرفقات", "portal.hr_training_program_attachments"),
        ("settings", "إعدادات التعميم", "portal.hr_training_program_settings"),
        ("enrollments", "المنتسبين", "portal.hr_training_program_enrollments"),
    ]
    return {"program": program, "step": step, "steps": steps}


@portal_bp.route("/hr/training/programs/<int:program_id>/info", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_READ)
def hr_training_program_info(program_id: int):
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    p = HRTrainingProgram.query.get_or_404(program_id)

    if not _training_can_manage():
        abort(403)

    # Lists
    categories = HRLookupItem.query.filter_by(category=TRAINING_CATEGORY).order_by(HRLookupItem.sort_order.asc(), HRLookupItem.id.asc()).all()
    sponsors = HRLookupItem.query.filter_by(category=TRAINING_SPONSOR).order_by(HRLookupItem.sort_order.asc(), HRLookupItem.id.asc()).all()
    countries = HRLookupItem.query.filter_by(category=COUNTRY_CATEGORY).order_by(HRLookupItem.sort_order.asc(), HRLookupItem.id.asc()).all()
    courses = HRTrainingCourse.query.order_by(HRTrainingCourse.sort_order.asc(), HRTrainingCourse.id.asc()).all()

    if request.method == "POST":
        p.program_no = (request.form.get("program_no") or "").strip()
        course_id = request.form.get("course_id")
        p.course_id = int(course_id) if (course_id and course_id.isdigit()) else None

        p.start_date = (request.form.get("start_date") or "").strip() or None
        p.end_date = (request.form.get("end_date") or "").strip() or None

        country_id = request.form.get("country_lookup_id")
        p.country_lookup_id = int(country_id) if (country_id and country_id.isdigit()) else None

        sponsor_id = request.form.get("sponsor_lookup_id")
        p.sponsor_lookup_id = int(sponsor_id) if (sponsor_id and sponsor_id.isdigit()) else None

        p.venue = (request.form.get("venue") or "").strip() or None

        hours = (request.form.get("hours") or "").strip()
        try:
            p.hours = int(hours) if hours else None
        except Exception:
            p.hours = None

        p.is_hosted = True if (request.form.get("is_hosted") == "1") else False
        p.trainer_name = (request.form.get("trainer_name") or "").strip() or None
        p.notes = (request.form.get("notes") or "").strip() or None

        p.updated_by_id = current_user.id

        db.session.commit()
        flash("تم حفظ معلومات التدريب.", "success")
        return redirect(url_for("portal.hr_training_program_conditions", program_id=p.id))

    ctx = _training_wizard_ctx(p, "info")
    return render_template(
        "portal/hr/training/program_info.html",
        **ctx,
        categories=categories,
        sponsors=sponsors,
        countries=countries,
        courses=courses,
    )


@portal_bp.route("/hr/training/programs/<int:program_id>/conditions", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_READ)
def hr_training_program_conditions(program_id: int):
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    p = HRTrainingProgram.query.get_or_404(program_id)
    if not _training_can_manage():
        abort(403)

    _ensure_training_condition_fields_seeded()

    fields = HRLookupItem.query.filter_by(category=TRAINING_CONDITION_FIELD).order_by(HRLookupItem.sort_order.asc(), HRLookupItem.id.asc()).all()

    # Build lookup options for condition values (best-effort)
    field_code_map = {int(f.id): ((f.code or "").strip().upper()) for f in (fields or []) if getattr(f, "id", None)}

    lookup_code_to_category = {
        "JOB_TITLE": "JOB_TITLE",
        "APPOINTMENT_TYPE": "APPOINTMENT_TYPE",
        "JOB_CATEGORY": "JOB_CATEGORY",
        "JOB_GRADE": "JOB_GRADE",
        "WORK_GOV": "WORK_GOV",
        "WORK_LOC": "WORK_LOCATION",
        "ADMIN_TITLE": "ADMIN_TITLE",
        "EDU_DEGREE": "QUAL_DEGREE",
        "EDU_SPEC": "QUAL_SPECIALIZATION",
    }

    value_options = {}
    value_label_map = {}
    try:
        for code, cat in lookup_code_to_category.items():
            items = _lookup_items(cat)
            value_options[code] = [
                {"id": int(it.id), "label": (it.name_ar or it.name_en or str(it.id))}
                for it in (items or [])
                if getattr(it, "id", None)
            ]
            value_label_map[code] = {
                int(it.id): (it.name_ar or it.name_en or str(it.id))
                for it in (items or [])
                if getattr(it, "id", None)
            }
    except Exception:
        pass

    # Directorates / Departments
    try:
        dirs = Directorate.query.order_by(Directorate.name_ar.asc(), Directorate.id.asc()).all()
        value_options["DIRECTORATE"] = [{"id": int(d.id), "label": (d.name_ar or d.name_en or str(d.id))} for d in (dirs or []) if getattr(d, "id", None)]
        value_label_map["DIRECTORATE"] = {int(d.id): (d.name_ar or d.name_en or str(d.id)) for d in (dirs or []) if getattr(d, "id", None)}
    except Exception:
        pass

    try:
        deps = Department.query.order_by(Department.name_ar.asc(), Department.id.asc()).all()
        value_options["DEPARTMENT"] = [{"id": int(d.id), "label": (d.name_ar or d.name_en or str(d.id))} for d in (deps or []) if getattr(d, "id", None)]
        value_label_map["DEPARTMENT"] = {int(d.id): (d.name_ar or d.name_en or str(d.id)) for d in (deps or []) if getattr(d, "id", None)}
    except Exception:
        pass

    numeric_codes = ["AGE", "SERVICE_YEARS"]

    if request.method == "POST":
        action = (request.form.get("action") or "").strip().lower()
        if action == "add":
            field_id = request.form.get("field_lookup_id")
            op = (request.form.get("operator") or "EQ").strip().upper()
            v1 = (request.form.get("value1") or "").strip() or None
            v2 = (request.form.get("value2") or "").strip() or None

            c = HRTrainingCondition(
                program_id=p.id,
                field_lookup_id=int(field_id) if (field_id and field_id.isdigit()) else None,
                operator=op,
                value1=v1,
                value2=v2,
            )
            db.session.add(c)
            db.session.commit()
            flash("تمت إضافة الشرط.", "success")
            return redirect(url_for("portal.hr_training_program_conditions", program_id=p.id))

        if action == "delete":
            cid = request.form.get("cond_id")
            if cid and cid.isdigit():
                row = HRTrainingCondition.query.filter_by(id=int(cid), program_id=p.id).first()
                if row:
                    db.session.delete(row)
                    db.session.commit()
                    flash("تم حذف الشرط.", "success")
            return redirect(url_for("portal.hr_training_program_conditions", program_id=p.id))

    ctx = _training_wizard_ctx(p, "conditions")
    return render_template(
        "portal/hr/training/program_conditions.html",
        **ctx,
        fields=fields,
        field_code_map=field_code_map,
        value_options=value_options,
        value_label_map=value_label_map,
        numeric_codes=numeric_codes,
    )


@portal_bp.route("/hr/training/programs/<int:program_id>/attachments", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_READ)
def hr_training_program_attachments(program_id: int):
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    p = HRTrainingProgram.query.get_or_404(program_id)
    if not _training_can_manage():
        abort(403)

    if request.method == "POST":
        action = (request.form.get("action") or "").strip().lower()

        if action == "upload":
            title = (request.form.get("title") or "").strip() or None
            files = request.files.getlist("files")
            try:
                saved = _save_training_files(files, p.id, title=title)
                db.session.commit()
                flash(f"تم رفع {saved} مرفق(ات).", "success")
            except Exception:
                db.session.rollback()
                flash("تعذر رفع المرفقات.", "danger")
            return redirect(url_for("portal.hr_training_program_attachments", program_id=p.id))

        if action == "delete":
            att_id = request.form.get("att_id")
            if att_id and att_id.isdigit():
                att = HRTrainingAttachment.query.filter_by(id=int(att_id), program_id=p.id).first()
                if att:
                    # remove file best-effort
                    try:
                        folder = _training_upload_dir(p.id)
                        fp = os.path.join(folder, att.stored_name or "")
                        if os.path.exists(fp):
                            os.remove(fp)
                    except Exception:
                        pass
                    db.session.delete(att)
                    db.session.commit()
                    flash("تم حذف المرفق.", "success")
            return redirect(url_for("portal.hr_training_program_attachments", program_id=p.id))

    ctx = _training_wizard_ctx(p, "attachments")
    atts = HRTrainingAttachment.query.filter_by(program_id=p.id).order_by(HRTrainingAttachment.uploaded_at.desc()).all()
    return render_template(
        "portal/hr/training/program_attachments.html",
        **ctx,
        attachments=atts,
    )


@portal_bp.route("/hr/training/programs/<int:program_id>/attachments/<int:att_id>/download")
@login_required
@_perm(PORTAL_READ)
def hr_training_attachment_download(program_id: int, att_id: int):
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    p = HRTrainingProgram.query.get_or_404(program_id)
    att = HRTrainingAttachment.query.filter_by(id=att_id, program_id=program_id).first_or_404()

    allowed = False
    if _training_can_manage():
        allowed = True
    else:
        # employee allowed if enrolled OR program published
        try:
            if p.is_published:
                allowed = True
        except Exception:
            pass
        try:
            r = HRTrainingEnrollment.query.filter_by(program_id=p.id, user_id=current_user.id).first()
            if r:
                allowed = True
        except Exception:
            pass

    if not allowed:
        abort(403)

    folder = _training_upload_dir(program_id)
    return send_from_directory(folder, att.stored_name, as_attachment=True, download_name=(att.original_name or att.stored_name))


@portal_bp.route("/hr/training/programs/<int:program_id>/settings", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_READ)
def hr_training_program_settings(program_id: int):
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    p = HRTrainingProgram.query.get_or_404(program_id)
    if not _training_can_manage():
        abort(403)

    if request.method == "POST":
        was_published = bool(p.is_published)
        was_cond = bool(getattr(p, 'publish_conditions_only', False))

        # publish_mode: 0=لا, 1=نعم (للجميع), 2=حسب الشروط فقط
        pmode = (request.form.get('publish_mode') or request.form.get('is_published') or '').strip()
        if pmode in ('2', 'COND', 'CONDITIONS'):
            new_published, new_cond = True, True
        elif pmode in ('1', 'true', 'on'):
            new_published, new_cond = True, False
        else:
            new_published, new_cond = False, False

        p.is_published = new_published
        try:
            p.publish_conditions_only = bool(new_cond)
        except Exception:
            pass
        p.nomination_start = (request.form.get("nomination_start") or "").strip() or None
        p.nomination_end = (request.form.get("nomination_end") or "").strip() or None

        def _tri(name: str):
            v = (request.form.get(name) or "").strip()
            if v == "":
                return None
            return True if v == "1" else False

        p.require_manager_approval = _tri("require_manager_approval")
        p.needs_training_needs_window = _tri("needs_training_needs_window")
        p.employee_notifications_enabled = _tri("employee_notifications_enabled")
        p.apply_conditions_on_portal = _tri("apply_conditions_on_portal")
        p.apply_conditions_in_program = _tri("apply_conditions_in_program")

        p.updated_by_id = current_user.id
        db.session.commit()

        # If published now, optionally notify employees via Portal notifications
        try:
            if (not was_published) and new_published:
                _training_notify_published(p, actor_id=current_user.id)
        except Exception:
            pass

        flash("تم حفظ إعدادات التعميم.", "success")
        return redirect(url_for("portal.hr_training_program_enrollments", program_id=p.id))

    ctx = _training_wizard_ctx(p, "settings")
    sys_defaults = {
        "require_manager_approval": _ss_get_bool(SS_TRAINING_REQUIRE_MANAGER_APPROVAL, False),
        "needs_training_needs_window": _ss_get_bool(SS_TRAINING_NEEDS_WINDOW, False),
        "employee_notifications_enabled": _ss_get_bool(SS_TRAINING_EMPLOYEE_NOTIFICATIONS, True),
        "apply_conditions_on_portal": _ss_get_bool(SS_TRAINING_APPLY_COND_PORTAL, False),
        "apply_conditions_in_program": _ss_get_bool(SS_TRAINING_APPLY_COND_ADMIN, False),
    }
    return render_template(
        "portal/hr/training/program_settings.html",
        **ctx,
        sys_defaults=sys_defaults,
    )


@portal_bp.route("/hr/training/programs/<int:program_id>/enrollments", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_READ)
def hr_training_program_enrollments(program_id: int):
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    p = HRTrainingProgram.query.get_or_404(program_id)
    if not _training_can_manage():
        abort(403)

    if request.method == "POST":
        action = (request.form.get("action") or "").strip().lower()

        if action == "add":
            email = (request.form.get("employee_email") or "").strip().lower()
            status = (request.form.get("status") or "CANDIDATE").strip().upper()
            notes = (request.form.get("notes") or "").strip() or None

            u = None
            if email:
                u = User.query.filter(func.lower(User.email) == email).first()

            if not u:
                flash("لم يتم العثور على الموظف (تحقق من الإيميل).", "warning")
                return redirect(url_for("portal.hr_training_program_enrollments", program_id=p.id))

            # apply conditions (admin nomination) if enabled
            if _training_effective_bool(p, "apply_conditions_in_program", SS_TRAINING_APPLY_COND_ADMIN, False):
                ef = EmployeeFile.query.filter_by(user_id=u.id).first()
                ok = True
                for cond in HRTrainingCondition.query.filter_by(program_id=p.id).all():
                    if not _training_eval_condition(ef, cond):
                        ok = False
                        break
                if not ok:
                    flash("لا تنطبق شروط الانتساب على هذا الموظف.", "danger")
                    return redirect(url_for("portal.hr_training_program_enrollments", program_id=p.id))

            row = HRTrainingEnrollment.query.filter_by(program_id=p.id, user_id=u.id).first()
            if row:
                flash("الموظف موجود مسبقًا ضمن المنتسبين.", "info")
                return redirect(url_for("portal.hr_training_program_enrollments", program_id=p.id))

            row = HRTrainingEnrollment(program_id=p.id, user_id=u.id, status=status, notes=notes)
            db.session.add(row)
            db.session.commit()
            flash("تمت إضافة المنتسب.", "success")
            return redirect(url_for("portal.hr_training_program_enrollments", program_id=p.id))

        if action == "update":
            rid = request.form.get("enroll_id")
            status = (request.form.get("status") or "").strip().upper()
            notes = (request.form.get("notes") or "").strip() or None
            if rid and rid.isdigit():
                row = HRTrainingEnrollment.query.filter_by(id=int(rid), program_id=p.id).first()
                if row:
                    if status:
                        row.status = status
                    row.notes = notes
                    db.session.commit()
                    flash("تم تحديث حالة المنتسب.", "success")
            return redirect(url_for("portal.hr_training_program_enrollments", program_id=p.id))

        if action == "delete":
            rid = request.form.get("enroll_id")
            if rid and rid.isdigit():
                row = HRTrainingEnrollment.query.filter_by(id=int(rid), program_id=p.id).first()
                if row:
                    db.session.delete(row)
                    db.session.commit()
                    flash("تم حذف المنتسب.", "success")
            return redirect(url_for("portal.hr_training_program_enrollments", program_id=p.id))

    rows = HRTrainingEnrollment.query.filter_by(program_id=p.id).order_by(HRTrainingEnrollment.created_at.desc()).all()

    # Employees lookup (for faster selection in the form)
    try:
        employees = (
            User.query
            .order_by(func.coalesce(User.name, User.email).asc(), User.id.asc())
            .limit(500)
            .all()
        )
    except Exception:
        employees = []

    ctx = _training_wizard_ctx(p, "enrollments")
    return render_template(
        "portal/hr/training/program_enrollments.html",
        **ctx,
        enrollments=rows,
        employees=employees,
    )


@portal_bp.route("/hr/training/programs/<int:program_id>/apply", methods=["POST"])
@login_required
@_perm(PORTAL_READ)
def hr_training_program_apply(program_id: int):
    """Employee apply (self nomination)."""
    try:
        if not (current_user.has_perm(HR_READ) or current_user.has_perm(HR_SS_READ) or current_user.has_perm(HR_SS_CREATE)):
            abort(403)
    except Exception:
        abort(403)

    p = HRTrainingProgram.query.get_or_404(program_id)

    # published + within window
    today = date.today()
    s = _parse_yyyy_mm_dd(p.nomination_start) or date.min
    e = _parse_yyyy_mm_dd(p.nomination_end) or date.max
    if not p.is_published or not (s <= today <= e):
        flash("هذا التدريب غير متاح للانتساب الآن.", "warning")
        return redirect(url_for("portal.hr_training_log"))

    # apply conditions on portal if enabled OR if training is published by conditions only
    cond_required = False
    try:
        cond_required = bool(getattr(p, "publish_conditions_only", False)) or _training_effective_bool(p, "apply_conditions_on_portal", SS_TRAINING_APPLY_COND_PORTAL, False)
    except Exception:
        cond_required = _training_effective_bool(p, "apply_conditions_on_portal", SS_TRAINING_APPLY_COND_PORTAL, False)

    if cond_required:
        ef = EmployeeFile.query.filter_by(user_id=current_user.id).first()
        ok = True
        qual_cache = {}
        for cond in HRTrainingCondition.query.filter_by(program_id=p.id).all():
            try:
                if getattr(p, "publish_conditions_only", False):
                    if not _training_eval_condition_strict(ef, cond, qual_cache=qual_cache):
                        ok = False
                        break
                else:
                    if not _training_eval_condition(ef, cond):
                        ok = False
                        break
            except Exception:
                # fallback to existing behavior
                if not _training_eval_condition(ef, cond):
                    ok = False
                    break
        if not ok:
            flash("لا تنطبق شروط الانتساب عليك.", "danger")
            return redirect(url_for("portal.hr_training_log"))



    row = HRTrainingEnrollment.query.filter_by(program_id=p.id, user_id=current_user.id).first()
    if row:
        flash("أنت منتسب/مرشح مسبقًا لهذا التدريب.", "info")
        return redirect(url_for("portal.hr_training_log"))

    row = HRTrainingEnrollment(program_id=p.id, user_id=current_user.id, status="CANDIDATE")
    db.session.add(row)
    db.session.commit()
    flash("تم إرسال طلب الانتساب.", "success")
    return redirect(url_for("portal.hr_training_log"))


@portal_bp.route("/hr/training/programs/<int:program_id>/withdraw", methods=["POST"])
@login_required
@_perm(PORTAL_READ)
def hr_training_program_withdraw(program_id: int):
    try:
        if not (current_user.has_perm(HR_READ) or current_user.has_perm(HR_SS_READ) or current_user.has_perm(HR_SS_CREATE)):
            abort(403)
    except Exception:
        abort(403)

    row = HRTrainingEnrollment.query.filter_by(program_id=program_id, user_id=current_user.id).first()
    if not row:
        flash("لا يوجد انتساب لهذا التدريب.", "info")
        return redirect(url_for("portal.hr_training_log"))

    row.status = "WITHDRAWN"
    db.session.commit()
    flash("تم الانسحاب من التدريب.", "success")
    return redirect(url_for("portal.hr_training_log"))


@portal_bp.route("/hr/training/admin", methods=["GET", "POST"])
@login_required
@_perm(PORTAL_READ)
def hr_training_admin_dashboard():
    """Training admin dashboard (lookups, courses, settings)."""
    try:
        if not current_user.has_perm(HR_READ):
            abort(403)
    except Exception:
        abort(403)

    if not _training_can_manage():
        abort(403)

    tab = (request.args.get("tab") or "settings").strip().lower()

    if request.method == "POST":
        action = (request.form.get("action") or "").strip().lower()

        # Settings
        if action == "save_settings":
            _ss_set_bool(SS_TRAINING_REQUIRE_MANAGER_APPROVAL, request.form.get("require_manager_approval") == "1")
            _ss_set_bool(SS_TRAINING_NEEDS_WINDOW, request.form.get("needs_window") == "1")
            _ss_set_bool(SS_TRAINING_EMPLOYEE_NOTIFICATIONS, request.form.get("employee_notifications") == "1")
            _ss_set_bool(SS_TRAINING_APPLY_COND_PORTAL, request.form.get("apply_cond_portal") == "1")
            _ss_set_bool(SS_TRAINING_APPLY_COND_ADMIN, request.form.get("apply_cond_admin") == "1")
            db.session.commit()
            flash("تم حفظ الإعدادات العامة.", "success")
            return redirect(url_for("portal.hr_training_admin_dashboard", tab="settings"))

        # Lookups: categories/sponsors/countries
        if action in ("add_lookup", "edit_lookup", "delete_lookup"):
            cat = (request.form.get("category") or "").strip().upper()
            if cat not in (TRAINING_CATEGORY, TRAINING_SPONSOR, COUNTRY_CATEGORY):
                flash("تصنيف غير صحيح.", "warning")
                return redirect(url_for("portal.hr_training_admin_dashboard", tab="lookups"))

            if action == "add_lookup":
                code = (request.form.get("code") or "").strip() or None
                name_ar = (request.form.get("name_ar") or "").strip()
                name_en = (request.form.get("name_en") or "").strip() or None
                if not name_ar:
                    flash("يرجى تعبئة النص.", "warning")
                    return redirect(url_for("portal.hr_training_admin_dashboard", tab="lookups"))
                it = HRLookupItem(category=cat, code=(code or name_ar[:30].upper()), name_ar=name_ar, name_en=name_en, is_active=True)
                db.session.add(it)
                db.session.commit()
                flash("تمت الإضافة.", "success")
                return redirect(url_for("portal.hr_training_admin_dashboard", tab="lookups"))

            if action == "edit_lookup":
                lid = request.form.get("id")
                if lid and lid.isdigit():
                    it = HRLookupItem.query.filter_by(id=int(lid), category=cat).first()
                    if it:
                        it.code = (request.form.get("code") or it.code or "").strip() or it.code
                        it.name_ar = (request.form.get("name_ar") or it.name_ar or "").strip() or it.name_ar
                        it.name_en = (request.form.get("name_en") or it.name_en or "").strip() or None
                        db.session.commit()
                        flash("تم التعديل.", "success")
                return redirect(url_for("portal.hr_training_admin_dashboard", tab="lookups"))

            if action == "delete_lookup":
                lid = request.form.get("id")
                if lid and lid.isdigit():
                    it = HRLookupItem.query.filter_by(id=int(lid), category=cat).first()
                    if it:
                        db.session.delete(it)
                        db.session.commit()
                        flash("تم الحذف.", "success")
                return redirect(url_for("portal.hr_training_admin_dashboard", tab="lookups"))

        # Courses
        if action in ("add_course", "edit_course", "delete_course"):
            if action == "add_course":
                name_ar = (request.form.get("name_ar") or "").strip()
                name_en = (request.form.get("name_en") or "").strip() or None
                category_id = request.form.get("category_lookup_id")
                if not name_ar:
                    flash("يرجى تعبئة اسم التدريب.", "warning")
                    return redirect(url_for("portal.hr_training_admin_dashboard", tab="courses"))
                row = HRTrainingCourse(
                    name_ar=name_ar,
                    name_en=name_en,
                    category_lookup_id=int(category_id) if (category_id and category_id.isdigit()) else None,
                    created_by_id=current_user.id,
                )
                db.session.add(row)
                db.session.commit()
                flash("تمت إضافة التدريب.", "success")
                return redirect(url_for("portal.hr_training_admin_dashboard", tab="courses"))

            if action == "edit_course":
                cid = request.form.get("id")
                if cid and cid.isdigit():
                    row = HRTrainingCourse.query.get(int(cid))
                    if row:
                        row.name_ar = (request.form.get("name_ar") or row.name_ar or "").strip() or row.name_ar
                        row.name_en = (request.form.get("name_en") or "").strip() or None
                        category_id = request.form.get("category_lookup_id")
                        row.category_lookup_id = int(category_id) if (category_id and category_id.isdigit()) else None
                        db.session.commit()
                        flash("تم تعديل التدريب.", "success")
                return redirect(url_for("portal.hr_training_admin_dashboard", tab="courses"))

            if action == "delete_course":
                cid = request.form.get("id")
                if cid and cid.isdigit():
                    row = HRTrainingCourse.query.get(int(cid))
                    if row:
                        db.session.delete(row)
                        db.session.commit()
                        flash("تم حذف التدريب.", "success")
                return redirect(url_for("portal.hr_training_admin_dashboard", tab="courses"))

    # Data for UI
    categories = HRLookupItem.query.filter_by(category=TRAINING_CATEGORY).order_by(HRLookupItem.sort_order.asc(), HRLookupItem.id.asc()).all()
    sponsors = HRLookupItem.query.filter_by(category=TRAINING_SPONSOR).order_by(HRLookupItem.sort_order.asc(), HRLookupItem.id.asc()).all()
    countries = HRLookupItem.query.filter_by(category=COUNTRY_CATEGORY).order_by(HRLookupItem.sort_order.asc(), HRLookupItem.id.asc()).all()

    courses = HRTrainingCourse.query.order_by(HRTrainingCourse.sort_order.asc(), HRTrainingCourse.id.asc()).all()

    settings = {
        "require_manager_approval": _ss_get_bool(SS_TRAINING_REQUIRE_MANAGER_APPROVAL, False),
        "needs_window": _ss_get_bool(SS_TRAINING_NEEDS_WINDOW, False),
        "employee_notifications": _ss_get_bool(SS_TRAINING_EMPLOYEE_NOTIFICATIONS, True),
        "apply_cond_portal": _ss_get_bool(SS_TRAINING_APPLY_COND_PORTAL, False),
        "apply_cond_admin": _ss_get_bool(SS_TRAINING_APPLY_COND_ADMIN, False),
    }

    return render_template(
        "portal/hr/training/admin_dashboard.html",
        tab=tab,
        settings=settings,
        categories=categories,
        sponsors=sponsors,
        countries=countries,
        courses=courses,
    )
